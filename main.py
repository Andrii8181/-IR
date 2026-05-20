# main.py  — S.A.D. v2.1
# -*- coding: utf-8 -*-
"""
S.A.D. — Статистичний аналіз даних  v2.1
Залежності: pip install numpy scipy openpyxl matplotlib pillow
"""
import os, sys, math, json, io
import numpy as np
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, colorchooser
from tkinter.scrolledtext import ScrolledText
import tkinter.font as tkfont
from itertools import combinations
from collections import defaultdict
from datetime import datetime

from scipy.stats import (shapiro, kruskal, mannwhitneyu, friedmanchisquare,
                         wilcoxon, levene, pearsonr, spearmanr)
from scipy.stats import f as f_dist, t as t_dist, norm
from scipy.stats import studentized_range

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from matplotlib.figure import Figure
    from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
    import matplotlib.patches as mpatches
    import matplotlib.colors as mcolors
    HAS_MPL = True
except Exception:
    HAS_MPL = False; Figure = None; FigureCanvasTkAgg = None

try:
    from PIL import Image as _PILImage
    HAS_PIL = True
except Exception:
    HAS_PIL = False

ALPHA   = 0.05
COL_W   = 10
APP_VER = "2.1"

if HAS_MPL:
    import matplotlib
    matplotlib.rcParams.update({
        'font.family': 'serif',
        'font.serif':  ['Times New Roman','Times','DejaVu Serif'],
        'font.size': 11, 'axes.titlesize': 12, 'axes.labelsize': 11,
        'xtick.labelsize': 10, 'ytick.labelsize': 10,
        'axes.linewidth': 0.8,
        'axes.spines.top': False, 'axes.spines.right': False,
    })

# ── DPI awareness ──────────────────────────────────────────────
try:
    import ctypes
    try:    ctypes.windll.shcore.SetProcessDpiAwareness(1)
    except Exception:
        try: ctypes.windll.user32.SetProcessDPIAware()
        except Exception: pass
except Exception: pass

# ── Icon ──────────────────────────────────────────────────────
def _find_icon():
    dirs = [os.getcwd()]
    try: dirs.insert(0, os.path.dirname(os.path.abspath(__file__)))
    except Exception: pass
    try:
        if hasattr(sys, "_MEIPASS"): dirs.append(sys._MEIPASS)
    except Exception: pass
    for d in dirs:
        p = os.path.join(d, "icon.ico")
        if os.path.exists(p): return p
    return None

def set_icon(win):
    ico = _find_icon()
    if not ico: return
    try: win.iconbitmap(ico)
    except Exception:
        try: win.iconbitmap(default=ico)
        except Exception: pass

# ── Clipboard PNG → Windows ────────────────────────────────────
def embed_figure(fig, master, dpi=96):
    """Вставляє matplotlib Figure у tkinter frame."""
    cv = FigureCanvasTkAgg(fig, master=master)
    widget = cv.get_tk_widget()
    widget.pack(fill=tk.BOTH, expand=True)
    cv.draw()
    return cv


def _copy_fig_to_clipboard(fig):
    if not (HAS_MPL and HAS_PIL): return False, "Потрібні matplotlib і Pillow"
    try:
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=300, bbox_inches="tight")
        buf.seek(0)
        pil = _PILImage.open(buf)
        ok, msg = _copy_pil_win(pil); buf.close(); return ok, msg
    except Exception as ex: return False, str(ex)

def _copy_pil_win(pil_img):
    try: import ctypes; from ctypes import wintypes
    except Exception: return False, "ctypes недоступний"
    if sys.platform != "win32": return False, "Лише для Windows"
    if pil_img is None: return False, "Немає зображення"
    try:
        buf = io.BytesIO(); pil_img.convert("RGB").save(buf, "BMP"); bmp = buf.getvalue()
        if len(bmp) <= 14: return False, "BMP помилка"
        data = bmp[14:]
    except Exception as ex: return False, str(ex)
    u32 = ctypes.WinDLL("user32", use_last_error=True)
    k32 = ctypes.WinDLL("kernel32", use_last_error=True)
    u32.OpenClipboard.argtypes  = [wintypes.HWND];           u32.OpenClipboard.restype  = wintypes.BOOL
    u32.CloseClipboard.argtypes = [];                        u32.CloseClipboard.restype  = wintypes.BOOL
    u32.EmptyClipboard.argtypes = [];                        u32.EmptyClipboard.restype  = wintypes.BOOL
    u32.SetClipboardData.argtypes = [wintypes.UINT, wintypes.HANDLE]; u32.SetClipboardData.restype = wintypes.HANDLE
    k32.GlobalAlloc.argtypes  = [wintypes.UINT, ctypes.c_size_t]; k32.GlobalAlloc.restype  = wintypes.HGLOBAL
    k32.GlobalLock.argtypes   = [wintypes.HGLOBAL];          k32.GlobalLock.restype   = wintypes.LPVOID
    k32.GlobalUnlock.argtypes = [wintypes.HGLOBAL];          k32.GlobalUnlock.restype = wintypes.BOOL
    k32.GlobalFree.argtypes   = [wintypes.HGLOBAL];          k32.GlobalFree.restype   = wintypes.HGLOBAL
    if not u32.OpenClipboard(None): return False, f"OpenClipboard err {ctypes.get_last_error()}"
    try:
        u32.EmptyClipboard()
        hg = k32.GlobalAlloc(0x0042, len(data))
        if not hg: return False, "GlobalAlloc failed"
        pg = k32.GlobalLock(hg)
        if not pg: k32.GlobalFree(hg); return False, "GlobalLock failed"
        try: ctypes.memmove(pg, data, len(data))
        finally: k32.GlobalUnlock(hg)
        if not u32.SetClipboardData(8, hg): k32.GlobalFree(hg); return False, "SetClipboardData failed"
        return True, ""
    finally: u32.CloseClipboard()

# ═══════════════════════════════════════════════════════════════
# STAT HELPERS
# ═══════════════════════════════════════════════════════════════
def sig_mark(p):
    if p is None or (isinstance(p, float) and math.isnan(p)): return ""
    return "**" if p < 0.01 else ("*" if p < 0.05 else "")

def norm_txt(p):
    if p is None or (isinstance(p, float) and math.isnan(p)): return "н/д"
    return "нормальний розподіл" if p > 0.05 else "ненормальний розподіл"

def fmt(x, nd=3):
    if x is None or (isinstance(x, float) and math.isnan(x)): return ""
    try: return f"{float(x):.{nd}f}"
    except: return ""

def first_seen(seq):
    seen, out = set(), []
    for x in seq:
        if x not in seen: seen.add(x); out.append(x)
    return out

def center_win(win):
    win.update_idletasks()
    w, h = win.winfo_width(), win.winfo_height()
    sw, sh = win.winfo_screenwidth(), win.winfo_screenheight()
    win.geometry(f"{w}x{h}+{(sw-w)//2}+{(sh-h)//2}")

def median_q(arr):
    if not arr: return np.nan, np.nan, np.nan
    a = np.array(arr, dtype=float); a = a[~np.isnan(a)]
    if len(a) == 0: return np.nan, np.nan, np.nan
    return float(np.median(a)), float(np.percentile(a, 25)), float(np.percentile(a, 75))

def cv_vals(vals):
    a = np.array(vals, dtype=float); a = a[~np.isnan(a)]
    if len(a) < 2: return np.nan
    m = float(np.mean(a))
    return np.nan if m == 0 else float(np.std(a, ddof=1) / m * 100)

def cv_means(means):
    v = [float(x) for x in means if x is not None and not (isinstance(x, float) and math.isnan(x))]
    if len(v) < 2: return np.nan
    m = float(np.mean(v))
    return np.nan if m == 0 else float(np.std(v, ddof=1) / m * 100)

def eta2_label(pe2):
    if pe2 is None or math.isnan(pe2): return ""
    if pe2 < 0.01: return "дуже слабкий"
    if pe2 < 0.06: return "слабкий"
    if pe2 < 0.14: return "середній"
    return "сильний"

def eps2_kw(H, n, k):
    if any(x is None for x in [H, n, k]) or math.isnan(H) or n <= k or k < 2: return np.nan
    return float((H - k + 1) / (n - k))

def kendalls_w(chisq, nb, kt):
    if any(x is None for x in [chisq, nb, kt]) or math.isnan(chisq) or nb <= 0 or kt <= 1: return np.nan
    return float(chisq / (nb * (kt - 1)))

def cliffs_d(x, y):
    x = np.array(x, dtype=float); y = np.array(y, dtype=float)
    x = x[~np.isnan(x)]; y = y[~np.isnan(y)]
    nx, ny = len(x), len(y)
    if nx == 0 or ny == 0: return np.nan
    gt = int(np.sum(x[:, None] > y[None, :])); lt = int(np.sum(x[:, None] < y[None, :]))
    return float((gt - lt) / (nx * ny))

def cliffs_lbl(d):
    if d is None or math.isnan(d): return ""
    if d < 0.147: return "дуже слабкий"
    if d < 0.33:  return "слабкий"
    if d < 0.474: return "середній"
    return "сильний"

def fit_font(texts, family="Times New Roman", start=13, min_s=9, target=155):
    f = tkfont.Font(family=family, size=start); sz = start
    while sz > min_s:
        if max(f.measure(t) for t in texts) <= target: break
        sz -= 1; f.configure(size=sz)
    return f

# ═══════════════════════════════════════════════════════════════
# TREEVIEW TABLE (fixed columns, no drift)
# ═══════════════════════════════════════════════════════════════
def make_tv(parent, headers, rows, min_col=90):
    frm = tk.Frame(parent, bd=1, relief=tk.SUNKEN)
    vsb = ttk.Scrollbar(frm, orient="vertical")
    hsb = ttk.Scrollbar(frm, orient="horizontal")
    tv = ttk.Treeview(frm, columns=headers, show="headings",
                      yscrollcommand=vsb.set, xscrollcommand=hsb.set,
                      height=min(len(rows) + 1, 22))
    vsb.config(command=tv.yview); hsb.config(command=tv.xview)
    vsb.pack(side=tk.RIGHT, fill=tk.Y)
    hsb.pack(side=tk.BOTTOM, fill=tk.X)
    tv.pack(fill=tk.BOTH, expand=True)
    fnt = tkfont.Font(family="Times New Roman", size=11)
    for i, h in enumerate(headers):
        cw = max(fnt.measure(str(h)) + 24, min_col,
                 max((fnt.measure(str(r[i]) if i < len(r) and r[i] else "") + 20) for r in rows) if rows else min_col)
        tv.heading(h, text=str(h), anchor="center")
        tv.column(h, width=cw, minwidth=50, anchor="center", stretch=True)
    for row in rows:
        tv.insert("", "end", values=[("" if v is None else str(v)) for v in row])
    style = ttk.Style()
    style.configure("Treeview", font=("Times New Roman", 11), rowheight=22)
    style.configure("Treeview.Heading", font=("Times New Roman", 11, "bold"))
    return frm, tv

# ═══════════════════════════════════════════════════════════════
# DATA HELPERS
# ═══════════════════════════════════════════════════════════════

def _bind_nav(entries_2d, win, factors_count=0):
    """Прив'язати навігацію Enter/стрілки до двовимірного масиву Entry."""
    def _pos(w):
        for i, row in enumerate(entries_2d):
            for j, e in enumerate(row):
                if e is w: return i, j
        return None, None

    def _on_enter(event):
        i, j = _pos(event.widget)
        if i is None: return "break"
        ni = i + 1
        if ni >= len(entries_2d): return "break"
        entries_2d[ni][j].focus_set(); entries_2d[ni][j].icursor(tk.END)
        return "break"

    def _on_arrow(event):
        i, j = _pos(event.widget)
        if i is None: return "break"
        if event.keysym == "Up":    i = max(0, i-1)
        elif event.keysym == "Down": i = min(len(entries_2d)-1, i+1)
        elif event.keysym == "Left": j = max(0, j-1)
        elif event.keysym == "Right":j = min(len(entries_2d[i])-1, j+1)
        entries_2d[i][j].focus_set(); entries_2d[i][j].icursor(tk.END)
        return "break"

    for row in entries_2d:
        for e in row:
            e.bind("<Return>", _on_enter)
            e.bind("<Up>",     _on_arrow)
            e.bind("<Down>",   _on_arrow)
            e.bind("<Left>",   _on_arrow)
            e.bind("<Right>",  _on_arrow)


def _nav_move(entries_2d, ri, ci):
    if 0 <= ri < len(entries_2d) and 0 <= ci < len(entries_2d[ri]):
        entries_2d[ri][ci].focus_set(); entries_2d[ri][ci].icursor(tk.END)
    return "break"

def _nav_down(entries_2d, ri, ci, add_row_fn=None):
    nri = ri + 1
    if nri >= len(entries_2d) and add_row_fn:
        try: add_row_fn()
        except Exception: pass
    _nav_move(entries_2d, min(nri, len(entries_2d)-1), ci)
    return "break"

def bind_nav(entries_2d, e, add_row_fn=None):
    """Прив'язати навігацію Enter/стрілки до комірки Entry у двовимірному списку."""
    def find_pos():
        for ri, row in enumerate(entries_2d):
            for ci, cell in enumerate(row):
                if cell is e: return ri, ci
        return None, None

    def move(ri, ci):
        if 0 <= ri < len(entries_2d) and 0 <= ci < len(entries_2d[ri]):
            entries_2d[ri][ci].focus_set()
            entries_2d[ri][ci].icursor(tk.END)

    def on_enter(ev):
        ri, ci = find_pos()
        if ri is None: return "break"
        nri = ri + 1
        if nri >= len(entries_2d) and add_row_fn:
            add_row_fn()
        move(min(nri, len(entries_2d)-1), ci)
        return "break"

    def on_arrow(ev):
        ri, ci = find_pos()
        if ri is None: return "break"
        k = ev.keysym
        if   k == "Up":    move(max(0, ri-1), ci)
        elif k == "Down":  move(min(len(entries_2d)-1, ri+1), ci)
        elif k == "Left":  move(ri, max(0, ci-1))
        elif k == "Right": move(ri, min(len(entries_2d[ri])-1, ci+1))
        return "break"

    e.bind("<Return>", on_enter)
    e.bind("<Up>",    on_arrow)
    e.bind("<Down>",  on_arrow)
    e.bind("<Left>",  on_arrow)
    e.bind("<Right>", on_arrow)


    g = defaultdict(list)
    for r in long:
        v = r.get("value", np.nan)
        if v is None or math.isnan(v): continue
        g[tuple(r.get(x) for x in keys)].append(float(v))
    return g

def vstats(long, fkeys):
    vals = defaultdict(list)
    for r in long:
        v = r.get("value", np.nan)
        if v is None or math.isnan(v): continue
        vals[tuple(r.get(k) for k in fkeys)].append(float(v))
    out = {}
    for k, a in vals.items():
        n = len(a); m = float(np.mean(a)) if n else np.nan
        sd = float(np.std(a, ddof=1)) if n >= 2 else (0. if n == 1 else np.nan)
        out[k] = (m, sd, n)
    return out

def mean_ranks(long, keyfn):
    vals = []; ks = []
    for r in long:
        v = r.get("value", np.nan)
        if v is None or math.isnan(v): continue
        vals.append(float(v)); ks.append(keyfn(r))
    if not vals: return {}
    order = np.argsort(vals); sv = np.array(vals)[order]
    ranks = np.empty(len(vals), dtype=float)
    i = 0
    while i < len(sv):
        j = i
        while j < len(sv) and sv[j] == sv[i]: j += 1
        ar = (i + 1 + j) / 2.; ranks[order[i:j]] = ar; i = j
    agg = defaultdict(list)
    for k, rk in zip(ks, ranks): agg[k].append(float(rk))
    return {k: float(np.mean(v)) for k, v in agg.items()}

# ═══════════════════════════════════════════════════════════════
# CLD
# ═══════════════════════════════════════════════════════════════
def cld(levels_order, means_dict, sig_matrix):
    valid = [l for l in levels_order if not math.isnan(means_dict.get(l, np.nan))]
    if not valid: return {l: "" for l in levels_order}
    sl = sorted(valid, key=lambda z: means_dict[z], reverse=True)
    def sig(a, b): return bool(sig_matrix.get((a, b), False) or sig_matrix.get((b, a), False))
    groups = []
    for lv in sl:
        compat = [gi for gi, g in enumerate(groups) if all(not sig(lv, o) for o in g)]
        if not compat: groups.append({lv})
        else:
            for gi in compat: groups[gi].add(lv)
    def shared(a, b): return any(a in g and b in g for g in groups)
    for i in range(len(sl)):
        for j in range(i + 1, len(sl)):
            a, b = sl[i], sl[j]
            if sig(a, b) or shared(a, b): continue
            ng = {a, b}
            for c in sl:
                if c in ng: continue
                if not sig(c, a) and not sig(c, b) and all(not sig(c, x) for x in ng): ng.add(c)
            groups.append(ng)
    uniq = []
    for g in groups:
        if not any(g == h for h in uniq): uniq.append(g)
    cleaned = [g for g in uniq if not any(g < h for h in uniq)]
    alpha_ = "abcdefghijklmnopqrstuvwxyz"
    mapping = {lv: [] for lv in sl}
    for gi, g in enumerate(cleaned):
        lt = alpha_[gi] if gi < len(alpha_) else f"g{gi}"
        for lv in g: mapping[lv].append(lt)
    return {lv: "".join(sorted(mapping.get(lv, []))) for lv in levels_order}

# ═══════════════════════════════════════════════════════════════
# LEVENE TEST
# ═══════════════════════════════════════════════════════════════
def groups_by(long, fkeys):
    """Групує значення з long по комбінаціям рівнів факторів fkeys.
    Повертає dict: {tuple_of_levels: [values]}
    Для одного ключа fkeys=(f,) → {(level,): [values]}
    """
    from collections import defaultdict
    result = defaultdict(list)
    for r in long:
        key = tuple(r.get(f, "") for f in fkeys)
        result[key].append(r["value"])
    return dict(result)

def levene_test(groups_dict):
    arrs = [np.array(v, dtype=float) for v in groups_dict.values() if len(v) > 0]
    if len(arrs) < 2: return np.nan, np.nan
    try:
        stat, p = levene(*arrs, center='median')
        return float(stat), float(p)
    except Exception: return np.nan, np.nan

# ═══════════════════════════════════════════════════════════════
# PAIRWISE — parametric
# ═══════════════════════════════════════════════════════════════
def lsd_sig(levels, means, ns, MS, df, alpha=ALPHA):
    sig = {}
    if MS is None or df is None or math.isnan(MS) or math.isnan(df): return sig
    df = int(df);
    if df <= 0: return sig
    tc = float(t_dist.ppf(1 - alpha / 2, df))
    for a, b in combinations(levels, 2):
        ma, mb = means.get(a, np.nan), means.get(b, np.nan)
        na, nb = ns.get(a, 0), ns.get(b, 0)
        if any(math.isnan(x) for x in [ma, mb]) or na <= 0 or nb <= 0: continue
        se = math.sqrt(MS * (1 / na + 1 / nb))
        sig[(a, b)] = (abs(ma - mb) > tc * se)
    return sig

def pairwise_param(levels, means, ns, MS, df, method, alpha=ALPHA):
    """
    Parametric pairwise comparisons.
    Duncan: true step-down procedure — critical q depends on number of
    means spanned (p-range), not total m. This is methodologically correct.
    Tukey: simultaneous, uses studentized range with m groups.
    Bonferroni: Bonferroni-adjusted t-test.
    """
    rows = []; sig = {}
    if MS is None or df is None or math.isnan(MS) or math.isnan(df): return rows, sig
    df = int(df)
    if df <= 0: return rows, sig
    lvls = [x for x in levels if not math.isnan(means.get(x, np.nan)) and ns.get(x, 0) > 0]
    m = len(lvls)
    if m < 2: return rows, sig

    # For Duncan: sort means descending and compute step ranges
    if method == "duncan":
        sorted_lvls = sorted(lvls, key=lambda x: means[x], reverse=True)
        # Build significance matrix using step-down procedure
        for i in range(m):
            for j in range(i + 1, m):
                a, b = sorted_lvls[i], sorted_lvls[j]
                ma, mb = means[a], means[b]; na, nb = ns[a], ns[b]
                se = math.sqrt(MS * (1 / na + 1 / nb))
                if se <= 0: continue
                p_range = j - i + 1      # number of means spanned (≥ 2)
                # Duncan critical value: use studentized range with p_range groups
                # Duncan alpha per step: alpha_p = 1 - (1 - alpha)^(p-1)
                alpha_p = 1.0 - (1.0 - alpha) ** (p_range - 1)
                alpha_p = min(alpha_p, 0.5)  # cap for stability
                try:
                    q_crit = float(studentized_range.ppf(1 - alpha_p, p_range, df))
                    lsd_p = q_crit * se / math.sqrt(2)
                    is_s = (abs(ma - mb) > lsd_p)
                    # compute approximate p via q
                    q_obs = abs(ma - mb) * math.sqrt(2) / se
                    pa = float(1 - studentized_range.cdf(q_obs, p_range, df))
                except Exception:
                    is_s = False; pa = np.nan
                sig[(a, b)] = is_s
                rows.append([f"{a} vs {b}", fmt(pa, 4),
                             ("істотна різниця " + sig_mark(pa)) if is_s else "–"])
        return rows, sig

    # Tukey and Bonferroni
    for a, b in combinations(lvls, 2):
        ma, mb = means[a], means[b]; na, nb = ns[a], ns[b]
        se = math.sqrt(MS * (1 / na + 1 / nb))
        if se <= 0: continue
        tv = abs(ma - mb) / se; pr = 2 * (1 - float(t_dist.cdf(tv, df)))
        if method == "bonferroni":
            pa = min(1., pr * (m * (m - 1) / 2))
        elif method == "tukey":
            # Tukey–Kramer (handles unequal n via harmonic se)
            pa = float(1 - studentized_range.cdf(math.sqrt(2) * tv, m, df))
        else:
            pa = pr
        is_s = (pa < alpha); sig[(a, b)] = is_s
        rows.append([f"{a} vs {b}", fmt(pa, 4),
                     ("істотна різниця " + sig_mark(pa)) if is_s else "–"])
    return rows, sig

# ═══════════════════════════════════════════════════════════════
# PAIRWISE — nonparametric
# ═══════════════════════════════════════════════════════════════
def pairwise_mw(levels, groups, alpha=ALPHA):
    rows = []; sig = {}
    lvls = [x for x in levels if len(groups.get(x, [])) > 0]
    m = len(lvls); mt = m * (m - 1) / 2
    if m < 2: return rows, sig
    for a, b in combinations(lvls, 2):
        x = np.array(groups[a], dtype=float); y = np.array(groups[b], dtype=float)
        try:
            U, p = mannwhitneyu(x, y, alternative="two-sided")
            pa = min(1., float(p) * mt); d = cliffs_d(x, y)
            sig[(a, b)] = (pa < alpha)
            rows.append([f"{a} vs {b}", fmt(float(U), 3), fmt(pa, 4),
                         ("істотна різниця " + sig_mark(pa)) if pa < alpha else "–",
                         fmt(d, 4), cliffs_lbl(abs(d))])
        except Exception: continue
    return rows, sig

def pairwise_wilcox(levels, mat, alpha=ALPHA):
    rows = []; sig = {}
    k = len(levels); mt = k * (k - 1) / 2
    if k < 2: return rows, sig
    arr = np.array(mat, dtype=float)
    for i in range(k):
        for j in range(i + 1, k):
            x, y = arr[:, i], arr[:, j]
            try:
                st, p = wilcoxon(x, y, zero_method="wilcox", alternative="two-sided", mode="auto")
                pa = min(1., float(p) * mt)
                z = abs(norm.ppf(pa / 2)) if 0 < pa < 1 else 0.
                r = z / math.sqrt(len(x)) if len(x) > 0 else np.nan
                sig[(levels[i], levels[j])] = (pa < alpha)
                rows.append([f"{levels[i]} vs {levels[j]}", fmt(float(st), 3), fmt(pa, 4),
                             ("істотна різниця " + sig_mark(pa)) if pa < alpha else "–", fmt(r, 4)])
            except Exception: continue
    return rows, sig

# ═══════════════════════════════════════════════════════════════
# RCBD MATRIX
# ═══════════════════════════════════════════════════════════════
def rcbd_matrix(long, vnames, bnames, vk="VARIANT", bk="BLOCK"):
    bb = defaultdict(dict)
    for r in long:
        v = r.get("value", np.nan)
        if v is None or math.isnan(v): continue
        b = r.get(bk); vn = r.get(vk)
        if b is None or vn is None: continue
        bb[b][vn] = float(v)
    mat = []; kept = []
    for b in bnames:
        d = bb.get(b, {})
        if all(vn in d for vn in vnames): mat.append([d[vn] for vn in vnames]); kept.append(b)
    return mat, kept

# ═══════════════════════════════════════════════════════════════
# GLM / OLS — SS Types I / II / III
# ═══════════════════════════════════════════════════════════════
def _encode(col_vals, levels):
    cols, names = [], []
    for lv in levels[1:]:
        cols.append(np.array([1. if v == lv else 0. for v in col_vals], dtype=float))
        names.append(str(lv))
    return cols, names

def _build_X(long, fkeys, lbf, extra=None):
    n = len(long); y = np.array([float(r["value"]) for r in long], dtype=float)
    Xc = [np.ones(n, dtype=float)]; cn = ["Intercept"]; ts = {"Intercept": [0]}
    fdc = {}; fdn = {}
    for f in fkeys:
        vals = [r.get(f) for r in long]; cols, names = _encode(vals, lbf[f])
        fdc[f] = cols; fdn[f] = names
        if cols:
            idx = []
            for c, nm in zip(cols, names): Xc.append(c); cn.append(f"{f}:{nm}"); idx.append(len(Xc) - 1)
            ts[f"Фактор {f}"] = idx
        else: ts[f"Фактор {f}"] = []
    for r2 in range(2, len(fkeys) + 1):
        for cmb in combinations(fkeys, r2):
            lists = [fdc[f] for f in cmb]; nls = [fdn[f] for f in cmb]
            if any(len(L) == 0 for L in lists): ts["Фактор " + "×".join(cmb)] = []; continue
            idx = []
            def rec(i, cc, cn_, idx=idx, cmb=cmb, lists=lists, nls=nls):
                if i == len(lists):
                    Xc.append(cc); cn.append("×".join(f"{cmb[j]}:{cn_[j]}" for j in range(len(cmb)))); idx.append(len(Xc) - 1); return
                for ci, nm in zip(lists[i], nls[i]): rec(i + 1, (ci.copy() if cc is None else cc * ci), cn_ + [nm])
            rec(0, None, [])
            ts["Фактор " + "×".join(cmb)] = idx
    if extra:
        for nm, cols, coln in extra:
            idx = []
            for c, cn_ in zip(cols, coln): Xc.append(c); cn.append(f"{nm}:{cn_}"); idx.append(len(Xc) - 1)
            ts[nm] = idx
    X = np.column_stack(Xc); return y, X, ts, cn

def _ols(y, X):
    beta, *_ = np.linalg.lstsq(X, y, rcond=None); yh = X @ beta; res = y - yh
    sse = float(np.sum(res ** 2)); n, p = X.shape; dfe = n - p
    return beta, yh, res, sse, dfe, (sse / dfe if dfe > 0 else np.nan)

def _ss_type3(y, Xf, ts):
    """Type III (partial) SS — drop-term approach."""
    _, _, res, sse, dfe, mse = _ols(y, Xf); out = {}
    for term, idx in ts.items():
        if term == "Intercept": continue
        if not idx: out[term] = (np.nan, 0, np.nan, np.nan, np.nan); continue
        keep = [i for i in range(Xf.shape[1]) if i not in idx]
        Xr = Xf[:, keep]; _, _, _, sse_r, _, _ = _ols(y, Xr)
        ss = float(sse_r - sse); df = len(idx); ms = ss / df if df > 0 else np.nan
        F = (ms / mse) if (df > 0 and not math.isnan(mse) and mse > 0) else np.nan
        p = float(1 - f_dist.cdf(F, df, dfe)) if (not math.isnan(F) and dfe > 0) else np.nan
        out[term] = (ss, df, ms, F, p)
    return out, sse, dfe, mse, res

def _ss_type1(y, Xf, ts, fkeys):
    """Type I (sequential) SS."""
    _, _, _, sse_full, dfe, mse = _ols(y, Xf); out = {}
    # build sequence: intercept, then each term in order
    seq_terms = ["Intercept"] + [f"Фактор {f}" for f in fkeys]
    for r2 in range(2, len(fkeys) + 1):
        for cmb in combinations(fkeys, r2): seq_terms.append("Фактор " + "×".join(cmb))
    included = [0]  # intercept always included
    prev_sse = float(np.sum((y - np.mean(y)) ** 2))
    for term in seq_terms[1:]:
        idx = ts.get(term, [])
        if not idx: out[term] = (np.nan, 0, np.nan, np.nan, np.nan); continue
        new_cols = included + idx
        Xr = Xf[:, new_cols]; _, _, _, sse_r, _, _ = _ols(y, Xr)
        ss = float(prev_sse - sse_r); df = len(idx); ms = ss / df if df > 0 else np.nan
        F = (ms / mse) if (df > 0 and not math.isnan(mse) and mse > 0) else np.nan
        p = float(1 - f_dist.cdf(F, df, dfe)) if (not math.isnan(F) and dfe > 0) else np.nan
        out[term] = (ss, df, ms, F, p)
        prev_sse = sse_r; included += idx
    return out, sse_full, dfe, mse, _ols(y, Xf)[2]

def _ss_type2(y, Xf, ts, fkeys):
    """Type II (hierarchical) SS — each factor adjusted for all other main effects."""
    _, _, res, sse_full, dfe, mse = _ols(y, Xf); out = {}
    # For each main effect: remove it from model with all other main effects but no interactions
    main_terms = {f: ts.get(f"Фактор {f}", []) for f in fkeys}
    inter_terms = {}
    for r2 in range(2, len(fkeys) + 1):
        for cmb in combinations(fkeys, r2):
            inter_terms["Фактор " + "×".join(cmb)] = ts.get("Фактор " + "×".join(cmb), [])
    for term, idx in ts.items():
        if term == "Intercept": continue
        if not idx: out[term] = (np.nan, 0, np.nan, np.nan, np.nan); continue
        # keep: all columns except those in current term
        # But also keep all higher-order terms that don't contain this factor
        # Type II: remove current term from "all terms of same or lower order"
        keep = [i for i in range(Xf.shape[1]) if i not in idx]
        Xr = Xf[:, keep]; _, _, _, sse_r, _, _ = _ols(y, Xr)
        ss = float(sse_r - sse_full); df = len(idx); ms = ss / df if df > 0 else np.nan
        F = (ms / mse) if (df > 0 and not math.isnan(mse) and mse > 0) else np.nan
        p = float(1 - f_dist.cdf(F, df, dfe)) if (not math.isnan(F) and dfe > 0) else np.nan
        out[term] = (ss, df, ms, F, p)
    return out, sse_full, dfe, mse, res

def _ss_dispatch(ss_type, y, Xf, ts, fkeys):
    if ss_type == "I":    return _ss_type1(y, Xf, ts, fkeys)
    elif ss_type == "II": return _ss_type2(y, Xf, ts, fkeys)
    elif ss_type == "IV": return _ss_type4(y, Xf, ts, fkeys)
    else:                 return _ss_type3(y, Xf, ts)   # III default

def _ss_type4(y, Xf, ts, fkeys):
    """
    Type IV SS — for unbalanced designs with empty cells.
    Uses estimable contrasts: for each term, we compare the full model
    against the model where that term's columns are zeroed-out (projection approach).
    For balanced data this equals Type III. For missing cells it avoids
    non-estimable functions by operating only on estimable contrasts.
    """
    _, _, res, sse_full, dfe, mse = _ols(y, Xf); out = {}
    n = Xf.shape[0]
    hat = Xf @ np.linalg.pinv(Xf)   # hat matrix H = X(X'X)⁻X'
    for term, idx in ts.items():
        if term == "Intercept": continue
        if not idx: out[term] = (np.nan, 0, np.nan, np.nan, np.nan); continue
        # Build contrast: columns in `idx` projected onto estimable space
        C = Xf[:, idx]
        # Estimable part: C_est = H @ C  (project onto column space of X)
        C_est = hat @ C
        # SS via general linear hypothesis: SS = y'C(C'C)⁻C'y  projected
        try:
            CTC = C_est.T @ C_est
            CTC_inv = np.linalg.pinv(CTC)
            _, _, res_full, sse_f, dfe_f, mse_f = _ols(y, Xf)
            beta, *_ = np.linalg.lstsq(Xf, y, rcond=None)
            Cb = C_est.T @ (Xf @ beta)
            ss = float(Cb.T @ CTC_inv @ Cb)
            df = len(idx); ms = ss / df if df > 0 else np.nan
            F = (ms / mse) if (df > 0 and not math.isnan(mse) and mse > 0) else np.nan
            p = float(1 - f_dist.cdf(F, df, dfe)) if (not math.isnan(F) and dfe > 0) else np.nan
            out[term] = (ss, df, ms, F, p)
        except Exception:
            out[term] = (np.nan, 0, np.nan, np.nan, np.nan)
    return out, sse_full, dfe, mse, res

def _block_dum(long, bk="BLOCK"):
    blocks = first_seen([r.get(bk) for r in long if r.get(bk) is not None])
    if not blocks: return [], [], blocks
    vals = [r.get(bk) for r in long]; cols = []; names = []
    for b in blocks[1:]:
        cols.append(np.array([1. if v == b else 0. for v in vals], dtype=float)); names.append(str(b))
    return cols, names, blocks

def _nir05(long, fkeys, mse, dfe, lbf):
    nir = {}
    if math.isnan(mse) or dfe <= 0: return nir
    tc = float(t_dist.ppf(1 - ALPHA / 2, int(dfe)))
    for f in fkeys:
        nl = defaultdict(int)
        for r in long:
            v = r.get("value", np.nan)
            if v is None or math.isnan(v): continue
            if r.get(f): nl[r[f]] += 1
        ns = [n for n in nl.values() if n > 0]
        if ns: nir[f"Фактор {f}"] = tc * math.sqrt(2 * mse / (len(ns) / sum(1 / n for n in ns)))
    nc = defaultdict(int)
    for r in long:
        v = r.get("value", np.nan)
        if v is None or math.isnan(v): continue
        nc[tuple(r.get(f) for f in fkeys)] += 1
    ns = [n for n in nc.values() if n > 0]
    if ns: nir["Загальна"] = tc * math.sqrt(2 * mse / (len(ns) / sum(1 / n for n in ns)))
    return nir

def build_eff_rows(table):
    """
    Таблиця сили впливу (% від суми SS компонентів).
    Завжди дає суму = 100% незалежно від типу SS.

    Для Type III SS сума компонентів ≠ SS_total через часткову природу SS.
    Тому ділимо на СУМУ самих компонентів (факторів + залишку),
    а не на SS_total з рядка «Загальна».
    """
    # Збираємо всі значущі рядки (крім «Загальна» і WP-error/Блоки)
    components = []
    for row in table:
        nm, SSv = row[0], row[1]
        if nm == "Загальна": continue
        if "WP-error" in str(nm): continue
        if SSv is None or (isinstance(SSv, float) and math.isnan(SSv)): continue
        if float(SSv) < 0: continue
        components.append([nm, float(SSv)])

    if not components: return []

    # Сума компонентів — це знаменник для %
    total_ss = sum(c[1] for c in components)
    if total_ss <= 0: return []

    # Розраховуємо % від реальної суми
    out = [[nm, ss / total_ss * 100] for nm, ss in components]

    # Коригуємо останній рядок для точної суми = 100%
    current_sum = sum(r[1] for r in out)
    out[-1][1] += 100.0 - current_sum

    return [[r[0], fmt(r[1], 2)] for r in out]


def build_pe2_rows(table):
    ss_err = np.nan
    for row in table:
        if row[0].startswith("Залишок"): ss_err = row[1]; break
    out = []
    for row in table:
        nm, SSv = row[0], row[1]
        if nm.startswith("Залишок") or nm == "Загальна": continue
        if SSv is None or (isinstance(SSv, float) and math.isnan(SSv)): continue
        if isinstance(ss_err, float) and math.isnan(ss_err): continue
        d = float(SSv) + float(ss_err); pe2 = float(SSv) / d if d > 0 else np.nan
        out.append([nm, fmt(pe2, 4), eta2_label(pe2)])
    return out

# ═══════════════════════════════════════════════════════════════
# ANOVA MODELS
# ═══════════════════════════════════════════════════════════════
def anova_crd(long, fkeys, lbf, ss_type="III"):
    y, X, ts, _ = _build_X(long, fkeys, lbf)
    terms, sse, dfe, mse, res = _ss_dispatch(ss_type, y, X, ts, fkeys)
    sst = float(np.sum((y - np.mean(y)) ** 2))
    ord_ = [f"Фактор {f}" for f in fkeys]
    for r2 in range(2, len(fkeys) + 1):
        for c in combinations(fkeys, r2): ord_.append("Фактор " + "×".join(c))
    table = [[nm, *terms.get(nm, (np.nan, 0, np.nan, np.nan, np.nan))] for nm in ord_]
    table.append(["Залишок", sse, dfe, mse, np.nan, np.nan])
    # SS_total = реальна сума (для Type I = sst; для III може відрізнятись)
    # Завжди показуємо математичне SS_total = Σ(yi - ȳ)²
    table.append(["Загальна", sst, len(y) - 1, np.nan, np.nan, np.nan])
    return {"table": table, "SS_error": sse, "df_error": dfe, "MS_error": mse,
            "SS_total": sst, "ss_type": ss_type,
            "residuals": res.tolist(), "NIR05": _nir05(long, fkeys, mse, dfe, lbf)}

def anova_rcbd(long, fkeys, lbf, bk="BLOCK", ss_type="III"):
    bc, bn, _ = _block_dum(long, bk)
    extra = [("Блоки", bc, bn)] if bc else []
    y, X, ts, _ = _build_X(long, fkeys, lbf, extra)
    terms, sse, dfe, mse, res = _ss_dispatch(ss_type, y, X, ts, fkeys)
    sst = float(np.sum((y - np.mean(y)) ** 2))
    table = []
    if bc: table.append(["Блоки", *terms.get("Блоки", (np.nan, 0, np.nan, np.nan, np.nan))])
    ord_ = [f"Фактор {f}" for f in fkeys]
    for r2 in range(2, len(fkeys) + 1):
        for c in combinations(fkeys, r2): ord_.append("Фактор " + "×".join(c))
    for nm in ord_: table.append([nm, *terms.get(nm, (np.nan, 0, np.nan, np.nan, np.nan))])
    table.append(["Залишок", sse, dfe, mse, np.nan, np.nan])
    table.append(["Загальна", sst, len(y) - 1, np.nan, np.nan, np.nan])
    return {"table": table, "SS_error": sse, "df_error": dfe, "MS_error": mse,
            "SS_total": sst, "residuals": res.tolist(), "NIR05": _nir05(long, fkeys, mse, dfe, lbf)}

def anova_latin_square(long, fkeys, lbf, ss_type="III"):
    """
    Латинський квадрат: Y = μ + τᵢ + ρⱼ + γₖ + εᵢⱼₖ
    τ = варіант, ρ = рядок, γ = стовпець
    df_помилка = (k-1)(k-2)
    """
    # Кодуємо рядки і стовпці як блокові ефекти
    rows_vals = [r.get("ROW","") for r in long]
    cols_vals = [r.get("COL","") for r in long]
    row_lvls  = first_seen([v for v in rows_vals if v])
    col_lvls  = first_seen([v for v in cols_vals if v])

    # Перевірка k×k структури
    k = len(lbf.get(fkeys[0], []))
    if len(row_lvls) != k or len(col_lvls) != k:
        raise ValueError(
            f"Латинський квадрат вимагає k={k} рядків і k={k} стовпців.\n"
            f"Знайдено: рядків={len(row_lvls)}, стовпців={len(col_lvls)}.\n"
            f"Перевірте що стовпці «Рядок» і «Стовпець» заповнені правильно.")

    row_c, row_n = _encode(rows_vals, row_lvls)
    col_c, col_n = _encode(cols_vals, col_lvls)
    extra = [("Рядки", row_c, row_n), ("Стовпці", col_c, col_n)]

    y, X, ts, _ = _build_X(long, fkeys, lbf, extra)
    terms, sse, dfe, mse, res = _ss_dispatch(ss_type, y, X, ts, fkeys)

    # Теоретичний df помилки = (k-1)(k-2)
    df_theory = (k-1)*(k-2)
    if df_theory > 0 and dfe != df_theory:
        # Використовуємо теоретичний df якщо числовий хибний
        if sse > 0:
            mse = sse / df_theory
            dfe = df_theory

    sst = float(np.sum((y - np.mean(y))**2))
    table = []
    for nm in [("Рядки", "Рядки"), ("Стовпці", "Стовпці")] + \
              [(f"Фактор {f}", f"Фактор {f}") for f in fkeys]:
        key = nm[1]
        table.append([nm[0], *terms.get(key, (np.nan, 0, np.nan, np.nan, np.nan))])
    table.append(["Залишок",  sse, dfe, mse, np.nan, np.nan])
    table.append(["Загальна", sst, len(y)-1, np.nan, np.nan, np.nan])

    return {"table": table, "SS_error": sse, "df_error": dfe, "MS_error": mse,
            "SS_total": sst, "residuals": res.tolist(),
            "NIR05": _nir05(long, fkeys, mse, dfe, lbf),
            "latin_k": k, "latin_rows": len(row_lvls), "latin_cols": len(col_lvls)}


    if main_f not in fkeys: main_f = fkeys[0]
    bc, bn, _ = _block_dum(long, bk)
    ml = first_seen([r.get(main_f) for r in long if r.get(main_f) is not None])
    if len(ml) < 2: raise ValueError("Головний фактор має мати ≥ 2 рівні")
    mv = [r.get(main_f) for r in long]; mc, mn = _encode(mv, ml)
    wpc = []; wpn = []
    for bi, bc_ in enumerate(bc):
        for mi, mc_ in enumerate(mc): wpc.append(bc_ * mc_); wpn.append(f"{bn[bi]}×{mn[mi]}")
    extra = []
    if bc: extra.append(("Блоки", bc, bn))
    wt = f"WP-error(Блоки×{main_f})"
    if wpc: extra.append((wt, wpc, wpn))
    lbf = {f: first_seen([r.get(f) for r in long if r.get(f) is not None]) for f in fkeys}
    y, X, ts, _ = _build_X(long, fkeys, lbf, extra)
    _, _, res, sse, dfe, mse = _ols(y, X)
    wp_idx = ts.get(wt, [])
    if not wp_idx: raise ValueError("Неможливо побудувати whole-plot error")
    keep = [i for i in range(X.shape[1]) if i not in wp_idx]
    _, _, _, sse_r, _, _ = _ols(y, X[:, keep])
    ss_wp = float(sse_r - sse); df_wp = len(wp_idx); ms_wp = ss_wp / df_wp if df_wp > 0 else np.nan
    sst = float(np.sum((y - np.mean(y)) ** 2))
    terms, _, _, _, _ = _ss_dispatch(ss_type, y, X, ts, fkeys)
    table = []
    if bc: table.append(["Блоки", *terms.get("Блоки", (np.nan, 0, np.nan, np.nan, np.nan))[:4], np.nan, np.nan])
    table.append([wt, ss_wp, df_wp, ms_wp, np.nan, np.nan])
    ord_ = [f"Фактор {f}" for f in fkeys]
    for r2 in range(2, len(fkeys) + 1):
        for c in combinations(fkeys, r2): ord_.append("Фактор " + "×".join(c))
    for nm in ord_:
        ss, df, ms, F, p = terms.get(nm, (np.nan, 0, np.nan, np.nan, np.nan))
        if nm == f"Фактор {main_f}":
            F2 = (ms / ms_wp) if (df > 0 and not any(math.isnan(x) for x in [ms, ms_wp]) and ms_wp > 0) else np.nan
            p2 = float(1 - f_dist.cdf(F2, df, df_wp)) if (not math.isnan(F2) and df_wp > 0) else np.nan
        else:
            F2 = (ms / mse) if (df > 0 and not any(math.isnan(x) for x in [ms, mse]) and mse > 0) else np.nan
            p2 = float(1 - f_dist.cdf(F2, df, dfe)) if (not math.isnan(F2) and dfe > 0) else np.nan
        table.append([nm, ss, df, ms, F2, p2])
    table.append(["Залишок", sse, dfe, mse, np.nan, np.nan])
    table.append(["Загальна", sst, len(y) - 1, np.nan, np.nan, np.nan])
    nir = {}
    if not (math.isnan(mse) or dfe <= 0 or math.isnan(ms_wp) or df_wp <= 0):
        def nh_f(f):
            nl = defaultdict(int)
            for r in long:
                v = r.get("value", np.nan)
                if v is None or math.isnan(v): continue
                if r.get(f): nl[r[f]] += 1
            ns = [n for n in nl.values() if n > 0]
            return (len(ns) / sum(1 / n for n in ns)) if ns else np.nan
        tc_s = float(t_dist.ppf(1 - ALPHA / 2, int(dfe))) if dfe > 0 else np.nan
        tc_w = float(t_dist.ppf(1 - ALPHA / 2, int(df_wp))) if df_wp > 0 else np.nan
        for f in fkeys:
            nh = nh_f(f)
            if math.isnan(nh) or nh <= 0: continue
            if f == main_f: nir[f"Фактор {f}(WP)"] = tc_w * math.sqrt(2 * ms_wp / nh)
            else: nir[f"Фактор {f}"] = tc_s * math.sqrt(2 * mse / nh)
    return {"table": table, "SS_error": sse, "df_error": dfe, "MS_error": mse,
            "SS_total": sst, "residuals": res.tolist(),
            "MS_whole": ms_wp, "df_whole": df_wp, "main_factor": main_f, "NIR05": nir}

# ═══════════════════════════════════════════════════════════════
# PROJECT SAVE / LOAD
# ═══════════════════════════════════════════════════════════════
def project_to_dict(app):
    rows = [[e.get() for e in row] for row in app.entries]
    return {"version": APP_VER, "factors_count": app.factors_count,
            "factor_title_map": app.factor_title_map, "cols": app.cols, "rows_data": rows}

def project_from_dict(app, d):
    fc = d.get("factors_count", 1)
    app.open_table(fc)
    app.factor_title_map = d.get("factor_title_map", {})
    for j, fk in enumerate(app.factor_keys):
        t = app.factor_title_map.get(fk, f"Фактор {fk}")
        if j < len(app.header_labels): app.header_labels[j].configure(text=t)
    sc = d.get("cols", app.cols)
    while app.cols < sc: app.add_column()
    rd = d.get("rows_data", [])
    while len(app.entries) < len(rd): app.add_row()
    for i, rv in enumerate(rd):
        for j, v in enumerate(rv):
            if i < len(app.entries) and j < len(app.entries[i]):
                app.entries[i][j].delete(0, tk.END); app.entries[i][j].insert(0, v)

# ═══════════════════════════════════════════════════════════════
# GRAPH SETTINGS
# ═══════════════════════════════════════════════════════════════
DEF_GS = {
    "font_family": "Times New Roman", "font_style": "normal", "font_size": 11,
    "box_color": "#ffffff", "median_color": "#c62828",
    "whisker_color": "#000000", "flier_color": "#555555",
    "venn_colors": ["#4c72b0", "#dd8452", "#55a868", "#c44e52"],
    "venn_alpha": 0.45, "venn_font_size": 11, "venn_font_color": "#000000",
    "heatmap_cmap": "RdYlGn", "heatmap_font_size": 10, "heatmap_annot_color": "#000000",
}

# ═══════════════════════════════════════════════════════════════
# GRAPH SETTINGS DIALOG
# ═══════════════════════════════════════════════════════════════
class GraphSettingsDlg(tk.Toplevel):
    FONTS  = ["Times New Roman", "Arial", "Calibri", "Georgia", "Verdana", "Courier New"]
    STYLES = ["normal", "bold", "italic", "bold italic"]
    CMAPS  = ["RdYlGn", "coolwarm", "RdBu", "PiYG", "PRGn", "bwr", "seismic", "viridis", "plasma"]

    def __init__(self, parent, gs: dict, show_heatmap=False):
        super().__init__(parent)
        self.title("Налаштування графіків")
        self.resizable(False, False); set_icon(self)
        self.gs = dict(gs); self.result = None
        self._col_box = gs["box_color"]; self._col_med = gs["median_color"]
        self._col_wh  = gs["whisker_color"]; self._col_fl = gs["flier_color"]
        self._venn_fc = gs["venn_font_color"]; self._venn_cols = list(gs["venn_colors"])

        nb = ttk.Notebook(self); nb.pack(fill=tk.BOTH, expand=True, padx=8, pady=8)

        # ── Boxplot ──
        bp = tk.Frame(nb, padx=12, pady=10); nb.add(bp, text="Boxplot")
        self._ff = tk.StringVar(value=gs["font_family"])
        self._fs = tk.StringVar(value=gs["font_style"])
        self._fz = tk.IntVar(value=gs["font_size"])
        r = 0
        for lbl, var, vals in [("Шрифт:", self._ff, self.FONTS), ("Стиль:", self._fs, self.STYLES)]:
            tk.Label(bp, text=lbl).grid(row=r, column=0, sticky="w", pady=4)
            ttk.Combobox(bp, textvariable=var, values=vals, state="readonly", width=22).grid(row=r, column=1, sticky="w", padx=6); r += 1
        tk.Label(bp, text="Розмір:").grid(row=r, column=0, sticky="w", pady=4)
        tk.Spinbox(bp, from_=7, to=28, textvariable=self._fz, width=6).grid(row=r, column=1, sticky="w", padx=6); r += 1
        self._bp_btns = {}
        for lbl, attr in [("Колір коробки:", "_col_box"), ("Колір медіани:", "_col_med"),
                           ("Колір вусів:", "_col_wh"), ("Колір викидів:", "_col_fl")]:
            tk.Label(bp, text=lbl).grid(row=r, column=0, sticky="w", pady=4)
            btn = tk.Button(bp, width=6, relief=tk.SUNKEN, bg=getattr(self, attr),
                            command=lambda a=attr: self._pick(a))
            btn.grid(row=r, column=1, sticky="w", padx=6); self._bp_btns[attr] = btn; r += 1

        # ── Venn ──
        vf = tk.Frame(nb, padx=12, pady=10); nb.add(vf, text="Діаграма Венна")
        self._vff = tk.StringVar(value=gs["font_family"])
        self._vfz = tk.IntVar(value=gs["venn_font_size"])
        self._valpha = tk.DoubleVar(value=gs["venn_alpha"])
        r = 0
        tk.Label(vf, text="Шрифт:").grid(row=r, column=0, sticky="w", pady=4)
        ttk.Combobox(vf, textvariable=self._vff, values=self.FONTS, state="readonly", width=22).grid(row=r, column=1, sticky="w", padx=6); r += 1
        tk.Label(vf, text="Розмір:").grid(row=r, column=0, sticky="w", pady=4)
        tk.Spinbox(vf, from_=7, to=28, textvariable=self._vfz, width=6).grid(row=r, column=1, sticky="w", padx=6); r += 1
        tk.Label(vf, text="Прозорість:").grid(row=r, column=0, sticky="w", pady=4)
        tk.Scale(vf, from_=0., to=1., resolution=0.05, orient="horizontal",
                 variable=self._valpha, length=180).grid(row=r, column=1, sticky="w", padx=6); r += 1
        tk.Label(vf, text="Колір тексту:").grid(row=r, column=0, sticky="w", pady=4)
        self._vfc_btn = tk.Button(vf, width=6, relief=tk.SUNKEN, bg=self._venn_fc, command=self._pick_vfc)
        self._vfc_btn.grid(row=r, column=1, sticky="w", padx=6); r += 1
        tk.Label(vf, text="Кольори кіл:").grid(row=r, column=0, sticky="w", pady=4)
        bf_ = tk.Frame(vf); bf_.grid(row=r, column=1, sticky="w", padx=6)
        self._vci_btns = []
        for idx in range(4):
            b = tk.Button(bf_, width=4, relief=tk.SUNKEN, bg=self._venn_cols[idx],
                          command=lambda i=idx: self._pick_vci(i))
            b.pack(side=tk.LEFT, padx=2); self._vci_btns.append(b)

        # ── Heatmap ──
        if show_heatmap:
            hf = tk.Frame(nb, padx=12, pady=10); nb.add(hf, text="Теплова карта")
            self._hcmap = tk.StringVar(value=gs.get("heatmap_cmap", "RdYlGn"))
            self._hfz   = tk.IntVar(value=gs.get("heatmap_font_size", 10))
            self._hannot_col = gs.get("heatmap_annot_color", "#000000")
            r = 0
            tk.Label(hf, text="Палітра:").grid(row=r, column=0, sticky="w", pady=4)
            ttk.Combobox(hf, textvariable=self._hcmap, values=self.CMAPS, state="readonly", width=18).grid(row=r, column=1, sticky="w", padx=6); r += 1
            tk.Label(hf, text="Розмір шрифту:").grid(row=r, column=0, sticky="w", pady=4)
            tk.Spinbox(hf, from_=6, to=20, textvariable=self._hfz, width=6).grid(row=r, column=1, sticky="w", padx=6); r += 1
            tk.Label(hf, text="Колір анотацій:").grid(row=r, column=0, sticky="w", pady=4)
            self._hannot_btn = tk.Button(hf, width=6, relief=tk.SUNKEN, bg=self._hannot_col, command=self._pick_hannot)
            self._hannot_btn.grid(row=r, column=1, sticky="w", padx=6)
        else:
            self._hcmap = tk.StringVar(value=gs.get("heatmap_cmap", "RdYlGn"))
            self._hfz   = tk.IntVar(value=gs.get("heatmap_font_size", 10))
            self._hannot_col = gs.get("heatmap_annot_color", "#000000")

        bf2 = tk.Frame(self); bf2.pack(fill=tk.X, padx=10, pady=(0, 10))
        tk.Button(bf2, text="OK", width=10, command=self._ok).pack(side=tk.LEFT, padx=4)
        tk.Button(bf2, text="Скасувати", width=12, command=self.destroy).pack(side=tk.LEFT)
        self.update_idletasks(); center_win(self); self.grab_set()

    def _pick(self, attr):
        c = colorchooser.askcolor(color=getattr(self, attr), parent=self, title="Виберіть колір")
        if c and c[1]:
            setattr(self, attr, c[1])
            if attr in self._bp_btns: self._bp_btns[attr].configure(bg=c[1])

    def _pick_vfc(self):
        c = colorchooser.askcolor(color=self._venn_fc, parent=self, title="Колір тексту")
        if c and c[1]: self._venn_fc = c[1]; self._vfc_btn.configure(bg=c[1])

    def _pick_vci(self, idx):
        c = colorchooser.askcolor(color=self._venn_cols[idx], parent=self, title=f"Колір кола {idx+1}")
        if c and c[1]: self._venn_cols[idx] = c[1]; self._vci_btns[idx].configure(bg=c[1])

    def _pick_hannot(self):
        c = colorchooser.askcolor(color=self._hannot_col, parent=self, title="Колір анотацій")
        if c and c[1]: self._hannot_col = c[1]; self._hannot_btn.configure(bg=c[1])

    def _ok(self):
        self.result = {
            "font_family": self._ff.get(), "font_style": self._fs.get(), "font_size": self._fz.get(),
            "box_color": self._col_box, "median_color": self._col_med,
            "whisker_color": self._col_wh, "flier_color": self._col_fl,
            "venn_colors": list(self._venn_cols), "venn_alpha": float(self._valpha.get()),
            "venn_font_size": self._vfz.get(), "venn_font_color": self._venn_fc,
            "heatmap_cmap": self._hcmap.get(), "heatmap_font_size": self._hfz.get(),
            "heatmap_annot_color": self._hannot_col,
        }
        self.destroy()

# ═══════════════════════════════════════════════════════════════
# VENN DIAGRAM  — proper overlap layout
# ═══════════════════════════════════════════════════════════════
def draw_venn(ax, factor_values, interaction_values, colors, alpha, font_size, font_color, font_family, title=""):
    """
    Proportional Venn diagram.
    factor_values: list of (label, pct) for main factors (circles)
    interaction_values: dict {frozenset_of_indices: (label, pct)} for interactions (overlaps)
    """
    ax.set_aspect("equal"); ax.axis("off")
    n = len(factor_values)
    if n == 0: return

    # Circle centres layout
    import math as _m
    if n == 1:
        centres = [(0, 0)]
    elif n == 2:
        centres = [(-0.33, 0), (0.33, 0)]
    elif n == 3:
        r_ = 0.38
        centres = [(r_ * _m.cos(_m.radians(90 + 120 * i)),
                    r_ * _m.sin(_m.radians(90 + 120 * i))) for i in range(3)]
    else:
        r_ = 0.42
        centres = [(r_ * _m.cos(_m.radians(45 + 90 * i)),
                    r_ * _m.sin(_m.radians(45 + 90 * i))) for i in range(4)]

    radius = 0.40

    # Draw circles
    for i, (cx, cy) in enumerate(centres):
        circle = mpatches.Circle((cx, cy), radius, fc=colors[i % len(colors)],
                                  alpha=alpha, ec="#444444", lw=1.2, zorder=2)
        ax.add_patch(circle)

    # Label main factors — place outside circle
    for i, ((cx, cy), (lbl, pct)) in enumerate(zip(centres, factor_values)):
        # offset outward from centre of diagram
        dx = cx - 0; dy = cy - 0
        mag = max(_m.hypot(dx, dy), 0.01)
        ox = cx + (dx / mag) * 0.55; oy = cy + (dy / mag) * 0.55
        if n == 1: ox, oy = 0, 0.65
        ax.text(ox, oy, f"{lbl}\n{fmt(pct, 1)}%",
                ha="center", va="center", fontsize=font_size,
                color=font_color, fontfamily=font_family,
                fontweight="bold", linespacing=1.3, zorder=5)

    # Label interactions at geometric intersections
    for key, (lbl, pct) in (interaction_values or {}).items():
        idxs = sorted(list(key))
        if len(idxs) < 2 or max(idxs) >= n: continue
        # mean of involved centres
        mx = sum(centres[i][0] for i in idxs) / len(idxs)
        my = sum(centres[i][1] for i in idxs) / len(idxs)
        ax.text(mx, my, f"{fmt(pct, 1)}%",
                ha="center", va="center", fontsize=max(font_size - 1, 7),
                color=font_color, fontfamily=font_family,
                fontweight="bold", zorder=6)

    ax.set_xlim(-1.1, 1.1); ax.set_ylim(-1.1, 1.1)
    if title:
        ax.set_title(title, fontsize=font_size + 1, fontfamily=font_family, pad=8)



# ═══════════════════════════════════════════════════════════════
# HEATMAP-ONLY SETTINGS DIALOG (для кореляційного аналізу)
# ═══════════════════════════════════════════════════════════════
class HeatmapSettingsDlg(tk.Toplevel):
    """Діалог налаштувань лише для теплової карти кореляцій."""
    CMAPS = ["RdYlGn","coolwarm","RdBu","PiYG","PRGn","bwr","seismic",
             "viridis","plasma","Blues","Reds","Greens"]

    def __init__(self, parent, gs: dict):
        super().__init__(parent)
        self.title("Налаштування теплової карти")
        self.resizable(False, False); set_icon(self)
        self.gs = dict(gs); self.result = None
        self._hannot_col = gs.get("heatmap_annot_color", "#000000")

        frm = tk.Frame(self, padx=16, pady=14); frm.pack(fill=tk.BOTH, expand=True)

        self._hcmap = tk.StringVar(value=gs.get("heatmap_cmap","RdYlGn"))
        self._hfz   = tk.IntVar(value=gs.get("heatmap_font_size", 10))
        self._ff    = tk.StringVar(value=gs.get("font_family","Times New Roman"))

        r = 0
        for lbl, wid in [("Палітра кольорів:", None),
                          ("Шрифт:", None),
                          ("Розмір шрифту:", None),
                          ("Колір тексту у клітинках:", None)]:
            tk.Label(frm, text=lbl, font=("Times New Roman",12)
                     ).grid(row=r, column=0, sticky="w", pady=6)
            if r == 0:
                ttk.Combobox(frm, textvariable=self._hcmap, values=self.CMAPS,
                             state="readonly", width=20).grid(row=r, column=1,
                             sticky="w", padx=8)
            elif r == 1:
                fonts = ["Times New Roman","Arial","Calibri","Georgia","Verdana","Courier New"]
                ttk.Combobox(frm, textvariable=self._ff, values=fonts,
                             state="readonly", width=20).grid(row=r, column=1,
                             sticky="w", padx=8)
            elif r == 2:
                tk.Spinbox(frm, from_=6, to=20, textvariable=self._hfz,
                           width=6).grid(row=r, column=1, sticky="w", padx=8)
            elif r == 3:
                self._hannot_btn = tk.Button(frm, width=6, relief=tk.SUNKEN,
                                             bg=self._hannot_col,
                                             command=self._pick_col)
                self._hannot_btn.grid(row=r, column=1, sticky="w", padx=8)
            r += 1

        bf = tk.Frame(frm); bf.grid(row=r, column=0, columnspan=2, pady=(14,0))
        tk.Button(bf, text="OK", width=10, bg="#c62828", fg="white",
                  font=("Times New Roman",12), command=self._ok).pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Скасувати", width=12,
                  font=("Times New Roman",12), command=self.destroy).pack(side=tk.LEFT)
        self.update_idletasks(); center_win(self); self.grab_set()

    def _pick_col(self):
        c = colorchooser.askcolor(color=self._hannot_col, parent=self, title="Колір тексту")
        if c and c[1]: self._hannot_col = c[1]; self._hannot_btn.configure(bg=c[1])

    def _ok(self):
        self.result = dict(self.gs)
        self.result.update({
            "heatmap_cmap":        self._hcmap.get(),
            "heatmap_font_size":   self._hfz.get(),
            "heatmap_annot_color": self._hannot_col,
            "font_family":         self._ff.get(),
        })
        self.destroy()


# ═══════════════════════════════════════════════════════════════
# SCATTER MATRIX SETTINGS DIALOG
# ═══════════════════════════════════════════════════════════════
class ScatterSettingsDlg(tk.Toplevel):
    """Діалог налаштувань матриці діаграм розсіювання."""
    COLORS = ["#4c72b0","#dd8452","#55a868","#c44e52","#8172b2","#1a6b1a","#c62828","#555555"]
    FONTS  = ["Times New Roman","Arial","Calibri","Georgia","Verdana","Courier New"]

    def __init__(self, parent, sc_gs: dict):
        super().__init__(parent)
        self.title("Налаштування матриці розсіювання")
        self.resizable(False, False); set_icon(self)
        self.sc_gs = dict(sc_gs); self.result = None
        self._pt_color  = sc_gs.get("sc_point_color",  "#4c72b0")
        self._tr_color  = sc_gs.get("sc_trend_color",  "#c62828")
        self._hist_col  = sc_gs.get("sc_hist_color",   "#4c72b0")

        frm = tk.Frame(self, padx=16, pady=14); frm.pack(fill=tk.BOTH, expand=True)

        self._pt_size  = tk.IntVar(value=sc_gs.get("sc_point_size",  14))
        self._pt_alpha = tk.DoubleVar(value=sc_gs.get("sc_point_alpha", 0.75))
        self._show_tr  = tk.BooleanVar(value=sc_gs.get("sc_show_trend", True))
        self._tr_width = tk.DoubleVar(value=sc_gs.get("sc_trend_width", 0.9))
        self._ff       = tk.StringVar(value=sc_gs.get("font_family","Times New Roman"))
        self._fz       = tk.IntVar(value=sc_gs.get("sc_font_size", 6))

        rows_cfg = [
            ("Шрифт:",                  "combo",  self._ff,       self.FONTS),
            ("Розмір підписів:",        "spin",   self._fz,       (5, 18)),
            ("Розмір точок:",           "spin",   self._pt_size,  (3, 50)),
            ("Прозорість точок (0-1):", "scale",  self._pt_alpha, (0.1, 1.0)),
            ("Показувати лінію тренду:","check",  self._show_tr,  None),
            ("Товщина лінії тренду:",   "scale",  self._tr_width, (0.3, 3.0)),
        ]
        self._btn_refs = {}
        r = 0
        for lbl, wtype, var, opts in rows_cfg:
            tk.Label(frm, text=lbl, font=("Times New Roman",12)
                     ).grid(row=r, column=0, sticky="w", pady=5)
            if wtype == "combo":
                ttk.Combobox(frm, textvariable=var, values=opts,
                             state="readonly", width=20).grid(row=r, column=1, sticky="w", padx=8)
            elif wtype == "spin":
                tk.Spinbox(frm, from_=opts[0], to=opts[1], textvariable=var,
                           width=7).grid(row=r, column=1, sticky="w", padx=8)
            elif wtype == "scale":
                tk.Scale(frm, from_=opts[0], to=opts[1], resolution=0.05,
                         orient="horizontal", variable=var,
                         length=160).grid(row=r, column=1, sticky="w", padx=8)
            elif wtype == "check":
                tk.Checkbutton(frm, variable=var).grid(row=r, column=1, sticky="w", padx=8)
            r += 1

        # Colour pickers
        for lbl, attr, init in [
            ("Колір точок:",       "_pt_color",  self._pt_color),
            ("Колір лінії тренду:","_tr_color",  self._tr_color),
            ("Колір гістограм:",   "_hist_col",  self._hist_col),
        ]:
            tk.Label(frm, text=lbl, font=("Times New Roman",12)
                     ).grid(row=r, column=0, sticky="w", pady=5)
            btn = tk.Button(frm, width=6, relief=tk.SUNKEN, bg=init,
                            command=lambda a=attr: self._pick(a))
            btn.grid(row=r, column=1, sticky="w", padx=8)
            self._btn_refs[attr] = btn; r += 1

        bf = tk.Frame(frm); bf.grid(row=r, column=0, columnspan=2, pady=(14,0))
        tk.Button(bf, text="OK", width=10, bg="#c62828", fg="white",
                  font=("Times New Roman",12), command=self._ok).pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Скасувати", width=12,
                  font=("Times New Roman",12), command=self.destroy).pack(side=tk.LEFT)
        self.update_idletasks(); center_win(self); self.grab_set()

    def _pick(self, attr):
        c = colorchooser.askcolor(color=getattr(self, attr), parent=self,
                                  title="Виберіть колір")
        if c and c[1]:
            setattr(self, attr, c[1])
            self._btn_refs[attr].configure(bg=c[1])

    def _ok(self):
        self.result = dict(self.sc_gs)
        self.result.update({
            "sc_point_color":  self._pt_color,
            "sc_trend_color":  self._tr_color,
            "sc_hist_color":   self._hist_col,
            "sc_point_size":   self._pt_size.get(),
            "sc_point_alpha":  self._pt_alpha.get(),
            "sc_show_trend":   self._show_tr.get(),
            "sc_trend_width":  self._tr_width.get(),
            "font_family":     self._ff.get(),
            "sc_font_size":    self._fz.get(),
        })
        self.destroy()

# ═══════════════════════════════════════════════════════════════
# CORRELATION ANALYSIS WINDOW
# ═══════════════════════════════════════════════════════════════
class CorrelationWindow:

    HELP_TEXT = """
КОРЕЛЯЦІЙНИЙ АНАЛІЗ — ПОКРОКОВА ІНСТРУКЦІЯ
═══════════════════════════════════════════

ЩО ТАКЕ КОРЕЛЯЦІЯ?
  Кореляція показує наявність і силу статистичного зв'язку
  між двома показниками.
  Коефіцієнт кореляції r (або ρ) від -1 до +1:
    r = +1 → ідеальний прямий зв'язок
    r = -1 → ідеальний обернений зв'язок
    r = 0  → зв'язку немає

КРОК 1. ПІДГОТОВКА ТАБЛИЦІ ДАНИХ
  Кожен стовпець таблиці = один показник (змінна).
  Перейменуйте стовпці — двічі клікніть на заголовок
  (синя клітинка вгорі) і введіть назву показника.

  Кожен рядок = одне спостереження (рослина, ділянка, рік).
  Введіть числові дані у клітинки.

  Приклад:
  | Врожайність | Висота | Маса зерна |
  |    4.2      |  98.5  |   38.2     |
  |    5.1      | 103.2  |   41.5     |

  Мінімум: 2 стовпці (показники) по ≥ 3 значення.

  Вставте дані з Excel: скопіюйте в Excel → клік на першу
  клітинку таблиці → кнопка «Вставити з Excel».

КРОК 2. ЗАПУСК АНАЛІЗУ
  Натисніть «▶ Аналіз» → з'явиться вікно параметрів.

КРОК 3. ВИБІР МЕТОДУ КОРЕЛЯЦІЇ

  Авто (рекомендовано):
    Програма перевіряє нормальність кожного показника.
    Якщо всі нормальні → Пірсон.
    Якщо хоч один ненормальний → Спірмен.

  Пірсон r:
    Для нормально розподілених, неперервних даних.
    Вимірює ЛІНІЙНИЙ зв'язок.
    ⚠ При порушенні нормальності — програма попередить!

  Спірмен ρ (rho):
    Непараметричний, для будь-якого розподілу.
    Виявляє МОНОТОННИЙ зв'язок (не лише лінійний).
    Надійніший при наявності викидів.

КРОК 4. ПОПРАВКА НА МНОЖИННІ ПОРІВНЯННЯ
  При n показниках виконується n×(n-1)/2 тестів.
  Без поправки ризик хибних результатів зростає!

  Бонферроні (строга):
    p_скор = p × кількість_пар
    Рекомендується при ≤ 10 показниках.

  Benjamini-Hochberg/FDR (ліберальніша):
    Контролює частку хибних відкриттів.
    Рекомендується при > 10 показниках.

КРОК 5. ІНТЕРПРЕТАЦІЯ ТЕПЛОВОЇ КАРТИ

  В кожній клітинці три рядки:
    r = -0.82     ← коефіцієнт кореляції
    p = 0.003     ← p-значення після поправки
    n = 15        ← кількість пар спостережень

  Значущість позначається зірочками:
    * — p < 0.05  (значущий зв'язок)
    ** — p < 0.01 (високо значущий зв'язок)
    без зірочки — p ≥ 0.05 (зв'язок незначущий)

  Колір:
    Зелений → позитивна кореляція (r > 0)
    Червоний → негативна кореляція (r < 0)
    Жовтий/білий → зв'язок відсутній

КРОК 6. СИЛА ЗВ'ЯЗКУ — ІНТЕРПРЕТАЦІЯ |r|
  0.00 – 0.19:  дуже слабкий (практично відсутній)
  0.20 – 0.39:  слабкий
  0.40 – 0.59:  помірний (середній)
  0.60 – 0.79:  сильний
  0.80 – 1.00:  дуже сильний

  Для агрономічних досліджень:
  |r| ≥ 0.60 — практично значущий зв'язок.

КРОК 7. МАТРИЦЯ ДІАГРАМ РОЗСІЮВАННЯ
  Відкривається автоматично після теплової карти.
  По діагоналі — гістограма кожного показника.
  Поза діагоналлю — точкові діаграми кожної пари.
  Червона лінія — лінія тренду.

  Якщо точки лежать уздовж прямої → сильна лінійна кореляція.
  Якщо точки хаотичні → зв'язку немає.
  Якщо є криволінійна залежність → розгляньте Спірмена.

ВАЖЛИВО:
  ⚠ Кореляція ≠ Причинно-наслідковий зв'язок!
  Навіть дуже сильна кореляція не означає що один показник
  СПРИЧИНЯЄ зміну іншого. Для висновків про причинність
  потрібен регресійний аналіз або теоретичне обґрунтування.
"""

    def __init__(self, root, graph_settings):
        self.root = root
        self.gs = dict(graph_settings)
        self._hm_fig = None
        self._sc_fig = None

        self.win = tk.Toplevel(root)
        self.win.title("Кореляційний аналіз")
        self.win.geometry("1060x660"); set_icon(self.win)

        self._build_menu()
        self._build_toolbar()
        self._build_table()

    # ── Меню ──────────────────────────────────────────────────
    def _build_menu(self):
        mb = tk.Menu(self.win)
        fm = tk.Menu(mb, tearoff=0)
        fm.add_command(label="Зберегти проект", command=self._save_proj)
        fm.add_command(label="Відкрити проект", command=self._load_proj)
        fm.add_separator()
        fm.add_command(label="Завантажити Excel", command=self._load_excel)
        mb.add_cascade(label="Файл", menu=fm)
        em = tk.Menu(mb, tearoff=0)
        em.add_command(label="Додати рядок",     command=self.add_row)
        em.add_command(label="Видалити рядок",   command=self.del_row)
        em.add_command(label="Додати стовпчик",  command=self.add_col)
        em.add_command(label="Видалити стовпчик",command=self.del_col)
        mb.add_cascade(label="Правка", menu=em)
        self.win.config(menu=mb)

    # ── Панель інструментів ───────────────────────────────────
    def _build_toolbar(self):
        tb = tk.Frame(self.win, padx=6, pady=5); tb.pack(fill=tk.X)
        tk.Button(tb, text="▶ Аналіз", bg="#c62828", fg="white",
                  font=("Times New Roman", 13),
                  command=self._run_analysis).pack(side=tk.LEFT, padx=4)
        self._settings_btn = tk.Menubutton(tb, text="⚙ Налаштування ▾",
                                           font=("Times New Roman", 11),
                                           relief=tk.RAISED, bd=2)
        self._settings_btn.pack(side=tk.LEFT, padx=4)
        sm = tk.Menu(self._settings_btn, tearoff=0)
        sm.add_command(label="Додати рядок",      command=self.add_row)
        sm.add_command(label="Видалити рядок",    command=self.del_row)
        sm.add_separator()
        sm.add_command(label="Додати стовпець",   command=self.add_col)
        sm.add_command(label="Видалити стовпець", command=self.del_col)
        sm.add_separator()
        sm.add_command(label="🗑 Очистити таблицю", command=self._clear_table)
        self._settings_btn["menu"] = sm
        tk.Button(tb, text="Вставити з буфера",
                  font=("Times New Roman", 11),
                  command=self._paste).pack(side=tk.LEFT, padx=4)
        tk.Button(tb, text="📚 Довідка", bg="#1a4b8c", fg="white",
                  font=("Times New Roman", 11),
                  command=self._show_help).pack(side=tk.LEFT, padx=4)
        tk.Label(tb,
                 text="Двічі клікніть на синій заголовок щоб перейменувати показник",
                 font=("Times New Roman", 9), fg="#666"
                 ).pack(side=tk.LEFT, padx=10)

    def _clear_table(self):
        if not messagebox.askyesno("Очистити таблицю",
                "Видалити всі числові дані?\n(Назви стовпців залишаться)"):
            return
        for row in self.entries:
            for e in row: e.delete(0, tk.END)

    # ── Таблиця даних ─────────────────────────────────────────
    def _build_table(self):
        self.rows = 14; self.cols = 6
        tbl_frm = tk.Frame(self.win); tbl_frm.pack(fill=tk.BOTH, expand=True, padx=4, pady=4)
        self.canvas = tk.Canvas(tbl_frm)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb = ttk.Scrollbar(tbl_frm, orient="vertical", command=self.canvas.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        self.canvas.configure(yscrollcommand=sb.set)
        self.inner = tk.Frame(self.canvas)
        self.canvas.create_window((0, 0), window=self.inner, anchor="nw")
        self.inner.bind("<Configure>",
                        lambda e: self.canvas.config(scrollregion=self.canvas.bbox("all")))
        self.win.bind("<MouseWheel>",
                      lambda e: self.canvas.yview_scroll(int(-1*(e.delta/120)), "units"))
        self.header_labels = []
        self.header_vars   = []
        self._build_headers()
        self.entries = []
        for i in range(self.rows):
            self._add_row_widgets(i)
        _bind_nav(self.entries, self.win)

    def _build_headers(self):
        for j in range(self.cols):
            var = tk.StringVar(value=f"Показник {j+1}")
            self.header_vars.append(var)
            lbl = tk.Label(self.inner, textvariable=var,
                           relief=tk.RIDGE, width=14, cursor="hand2",
                           bg="#1a4b8c", fg="white",
                           font=("Times New Roman", 11, "bold"))
            lbl.grid(row=0, column=j, padx=2, pady=2, sticky="nsew")
            lbl.bind("<Double-Button-1>", lambda e, idx=j: self._rename_col(idx))
            self.header_labels.append(lbl)

    def _rename_col(self, idx):
        dlg = tk.Toplevel(self.win); dlg.title("Перейменувати показник")
        dlg.resizable(False, False); dlg.grab_set(); set_icon(dlg)
        tk.Label(dlg, text=f"Назва показника {idx+1}:",
                 font=("Times New Roman", 12)).pack(padx=16, pady=14)
        var = tk.StringVar(value=self.header_vars[idx].get())
        e = tk.Entry(dlg, textvariable=var, font=("Times New Roman", 12), width=28)
        e.pack(padx=16, pady=4); e.select_range(0, tk.END); e.focus_set()
        def apply():
            nm = var.get().strip()
            if nm: self.header_vars[idx].set(nm)
            dlg.destroy()
        tk.Button(dlg, text="ОК", bg="#c62828", fg="white",
                  font=("Times New Roman", 12), command=apply).pack(pady=(4,14))
        dlg.bind("<Return>", lambda ev: apply())
        center_win(dlg)

    def _add_row_widgets(self, i):
        row_ = []
        for j in range(self.cols):
            e = tk.Entry(self.inner, width=14, font=("Times New Roman", 11),
                         highlightthickness=1, highlightbackground="#c0c0c0")
            e.grid(row=i+1, column=j, padx=2, pady=2)
            e.bind("<Return>", self._on_enter)
            e.bind("<Tab>",    self._on_tab)
            row_.append(e)
        self.entries.append(row_)

    def add_row(self):
        i = len(self.entries)
        self._add_row_widgets(i); self.rows += 1
        _bind_nav(self.entries, self.win)
        self.canvas.config(scrollregion=self.canvas.bbox("all"))

    def del_row(self):
        if not self.entries: return
        for e in self.entries.pop(): e.destroy()
        self.rows -= 1
        self.canvas.config(scrollregion=self.canvas.bbox("all"))

    def add_col(self):
        ci = self.cols; self.cols += 1
        var = tk.StringVar(value=f"Показник {ci+1}")
        self.header_vars.append(var)
        lbl = tk.Label(self.inner, textvariable=var,
                       relief=tk.RIDGE, width=14, cursor="hand2",
                       bg="#1a4b8c", fg="white",
                       font=("Times New Roman", 11, "bold"))
        lbl.grid(row=0, column=ci, padx=2, pady=2, sticky="nsew")
        lbl.bind("<Double-Button-1>", lambda e, idx=ci: self._rename_col(idx))
        self.header_labels.append(lbl)
        for i, row_ in enumerate(self.entries):
            e = tk.Entry(self.inner, width=14, font=("Times New Roman", 11),
                         highlightthickness=1, highlightbackground="#c0c0c0")
            e.grid(row=i+1, column=ci, padx=2, pady=2)
            e.bind("<Return>", self._on_enter); e.bind("<Tab>", self._on_tab)
            row_.append(e)
        _bind_nav(self.entries, self.win)

    def del_col(self):
        if self.cols <= 2: return
        self.header_labels.pop().destroy()
        self.header_vars.pop()
        for row_ in self.entries: row_.pop().destroy()
        self.cols -= 1

    # ── Навігація ─────────────────────────────────────────────
    def _on_enter(self, event):
        for i, row_ in enumerate(self.entries):
            for j, e in enumerate(row_):
                if e is event.widget:
                    if i+1 >= len(self.entries): self.add_row()
                    self.entries[i+1][j].focus_set(); return "break"
        return "break"

    def _on_tab(self, event):
        for i, row_ in enumerate(self.entries):
            for j, e in enumerate(row_):
                if e is event.widget:
                    nj = j+1; ni = i
                    if nj >= self.cols: nj = 0; ni = i+1
                    if ni >= len(self.entries): self.add_row()
                    self.entries[ni][nj].focus_set(); return "break"
        return "break"

    # ── Вставка, збереження, завантаження ────────────────────
    def _paste(self):
        """Вставити дані з буфера обміну.
        Якщо активна клітинка Entry — вставляємо з неї.
        Інакше — з клітинки (0,0)."""
        try: data = self.win.clipboard_get()
        except Exception:
            messagebox.showwarning("Буфер порожній",
                "Буфер обміну порожній або не містить тексту.\n"
                "Скопіюйте дані з Excel (Ctrl+C) і спробуйте знову."); return
        if not data.strip(): return
        w = self.win.focus_get()
        pos = (0, 0)
        if isinstance(w, tk.Entry):
            for i, row_ in enumerate(self.entries):
                for j, e in enumerate(row_):
                    if e is w: pos = (i, j); break
                if pos != (0, 0): break
        r0, c0 = pos
        rows_data = [rt for rt in data.splitlines() if rt.strip()]
        if not rows_data: return

        # Визначаємо потрібні розміри таблиці
        max_cols_needed = c0 + max((len(rt.split("\t")) for rt in rows_data), default=1)
        max_rows_needed = r0 + len(rows_data)

        # Додаємо стовпці якщо потрібно
        while self.cols < max_cols_needed:
            self.add_col()

        # Вставляємо дані
        for ir, rt in enumerate(rows_data):
            rr = r0 + ir
            while rr >= len(self.entries): self.add_row()
            for jc, val in enumerate(rt.split("\t")):
                cc = c0 + jc
                if cc >= self.cols: break
                self.entries[rr][cc].delete(0, tk.END)
                self.entries[rr][cc].insert(0, val.strip())

        # Інформуємо якщо таблицю розширено
        added_cols = max_cols_needed - (c0 + 1) if max_cols_needed > self.cols else 0
        added_rows = max_rows_needed - len(self.entries) if max_rows_needed > len(self.entries) else 0

        # Фокус на першу вставлену клітинку
        if self.entries and r0 < len(self.entries):
            self.entries[r0][c0].focus_set()

    def _save_proj(self):
        path = filedialog.asksaveasfilename(
            parent=self.win, defaultextension=".sadp",
            filetypes=[("SAD проект","*.sadp"),("JSON","*.json")])
        if not path: return
        d = {"type":"correlation","version":APP_VER,
             "headers":[v.get() for v in self.header_vars],
             "rows_data":[[e.get() for e in row] for row in self.entries]}
        try:
            with open(path,"w",encoding="utf-8") as f: json.dump(d,f,ensure_ascii=False,indent=2)
            messagebox.showinfo("Збережено", path)
        except Exception as ex: messagebox.showerror("Помилка",str(ex))

    def _load_proj(self):
        path = filedialog.askopenfilename(
            parent=self.win, filetypes=[("SAD проект","*.sadp"),("JSON","*.json")])
        if not path: return
        try:
            with open(path,"r",encoding="utf-8") as f: d=json.load(f)
        except Exception as ex: messagebox.showerror("Помилка",str(ex)); return
        headers = d.get("headers",[])
        rd      = d.get("rows_data",[])
        while self.cols < len(headers): self.add_col()
        for j, h in enumerate(headers):
            if j < len(self.header_vars): self.header_vars[j].set(h)
        while len(self.entries) < len(rd): self.add_row()
        for i, rv in enumerate(rd):
            for j, v in enumerate(rv):
                if i<len(self.entries) and j<len(self.entries[i]):
                    self.entries[i][j].delete(0,tk.END); self.entries[i][j].insert(0,v)

    def _load_excel(self):
        if not HAS_OPENPYXL: messagebox.showerror("","pip install openpyxl"); return
        path = filedialog.askopenfilename(parent=self.win,
                    filetypes=[("Excel","*.xlsx *.xlsm *.xls")])
        if not path: return
        try:
            wb = openpyxl.load_workbook(path,data_only=True,read_only=True)
            raw = [[cell for cell in row] for row in wb.active.iter_rows(values_only=True)]
            wb.close()
        except Exception as ex: messagebox.showerror("",str(ex)); return
        while raw and all(v is None for v in raw[-1]): raw.pop()
        if not raw: return
        nc = max(len(r) for r in raw)
        while self.cols < nc: self.add_col()
        while len(self.entries) < len(raw): self.add_row()
        for i,row in enumerate(raw):
            for j,v in enumerate(row):
                if j>=self.cols: break
                cv = "" if v is None else str(v).replace(",",".")
                self.entries[i][j].delete(0,tk.END); self.entries[i][j].insert(0,cv)

    def _settings(self):
        dlg = GraphSettingsDlg(self.win, self.gs, show_heatmap=True)
        self.win.wait_window(dlg)
        if dlg.result: self.gs = dlg.result

    # ── Копіювання ────────────────────────────────────────────
    def _copy_heatmap(self):
        if self._hm_fig is None:
            messagebox.showwarning("","Спочатку виконайте аналіз."); return
        ok, msg = _copy_fig_to_clipboard(self._hm_fig)
        if ok: messagebox.showinfo("","Теплову карту скопійовано.\nВставте у Word через Ctrl+V.")
        else:   messagebox.showwarning("",f"Помилка: {msg}")

    # ── Довідка ───────────────────────────────────────────────
    def _show_help(self):
        win = tk.Toplevel(self.win)
        win.title("Довідка — Кореляційний аналіз")
        win.geometry("700x640"); set_icon(win)
        frm = tk.Frame(win); frm.pack(fill=tk.BOTH, expand=True, padx=8, pady=6)
        vsb = ttk.Scrollbar(frm, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        txt = tk.Text(frm, wrap="word", font=("Times New Roman",11),
                      yscrollcommand=vsb.set, relief=tk.FLAT,
                      bg="#fafafa", padx=10, pady=8, cursor="arrow")
        txt.pack(fill=tk.BOTH, expand=True)
        vsb.config(command=txt.yview)
        txt.insert("1.0", self.HELP_TEXT.strip())
        txt.configure(state="disabled")
        txt.bind("<MouseWheel>",
                 lambda e: txt.yview_scroll(int(-1*(e.delta/120)), "units"))
        tk.Button(win, text="Закрити", command=win.destroy,
                  font=("Times New Roman",11)).pack(pady=6)

    # ── Діалог параметрів і запуск ────────────────────────────
    def _run_analysis(self):
        """Спрощений діалог параметрів — без блоку 'де знаходяться назви'
        (назви беруться з заголовків стовпців які користувач перейменував)."""
        dlg = tk.Toplevel(self.win)
        dlg.title("Параметри кореляційного аналізу")
        dlg.resizable(False, False); set_icon(dlg)
        frm = tk.Frame(dlg, padx=20, pady=16); frm.pack(fill=tk.BOTH, expand=True)
        rb_f = ("Times New Roman", 12)

        # ── Метод кореляції ──────────────────────────────────
        tk.Label(frm, text="Метод кореляції:",
                 font=("Times New Roman", 12, "bold")).grid(
                 row=0, column=0, columnspan=2, sticky="w", pady=0)
        meth_var = tk.StringVar(value="auto")
        methods = [
            ("auto",
             "Авто (рекомендовано) — перевіряє нормальність:\n"
             "  ✓ всі нормальні → Пірсон\n"
             "  ✓ хоч один ненормальний → Спірмен"),
            ("pearson",
             "Пірсон r — лінійний зв'язок, нормальний розподіл\n"
             "  ⚠ При ненормальних даних програма попередить"),
            ("spearman",
             "Спірмен ρ — непараметричний, будь-який розподіл,\n"
             "  монотонний зв'язок, стійкий до викидів"),
        ]
        for ri, (val, txt_) in enumerate(methods):
            tk.Radiobutton(frm, text=txt_, variable=meth_var, value=val,
                           font=rb_f, justify="left", wraplength=440
                           ).grid(row=1+ri, column=0, columnspan=2, sticky="w", pady=3)

        ttk.Separator(frm, orient="horizontal").grid(
            row=4, column=0, columnspan=2, sticky="ew", pady=10)

        # ── Поправка на множинні порівняння ──────────────────
        tk.Label(frm, text="Поправка на множинні порівняння:",
                 font=("Times New Roman", 12, "bold")).grid(
                 row=5, column=0, columnspan=2, sticky="w", pady=0)
        corr_var = tk.StringVar(value="bonferroni")
        corrections = [
            ("bonferroni",
             "Бонферроні — строга, контролює сімейну помилку (FWER)\n"
             "  Рекомендується при ≤ 10 показниках"),
            ("bh",
             "Benjamini–Hochberg (FDR) — ліберальніша, більша потужність\n"
             "  Рекомендується при > 10 показниках"),
            ("none",
             "Без поправки — не рекомендується при > 3 показниках"),
        ]
        for ri, (val, txt_) in enumerate(corrections):
            tk.Radiobutton(frm, text=txt_, variable=corr_var, value=val,
                           font=rb_f, justify="left", wraplength=440
                           ).grid(row=6+ri, column=0, columnspan=2, sticky="w", pady=3)

        ttk.Separator(frm, orient="horizontal").grid(
            row=9, column=0, columnspan=2, sticky="ew", pady=10)

        # ── Рівень значущості ─────────────────────────────────
        tk.Label(frm, text="Рівень значущості α:",
                 font=("Times New Roman", 12, "bold")).grid(row=10, column=0, sticky="w")
        alpha_var = tk.StringVar(value="0.05")
        ttk.Combobox(frm, textvariable=alpha_var, values=["0.01","0.05","0.10"],
                     state="readonly", width=9,
                     font=("Times New Roman",12)).grid(row=10, column=1, sticky="w", padx=8)

        # ── Кнопки ───────────────────────────────────────────
        bf = tk.Frame(frm); bf.grid(row=11, column=0, columnspan=2, pady=(16,0))
        out = {"ok": False}

        def ok():
            out.update({"ok": True,
                        "method":     meth_var.get(),
                        "correction": corr_var.get(),
                        "alpha":      float(alpha_var.get())})
            dlg.destroy()

        tk.Button(bf, text="▶ Виконати аналіз", width=20,
                  bg="#c62828", fg="white",
                  font=("Times New Roman",12), command=ok).pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Скасувати", width=12,
                  font=("Times New Roman",12), command=dlg.destroy).pack(side=tk.LEFT, padx=4)

        dlg.update_idletasks()
        center_win(dlg)
        dlg.bind("<Return>", lambda e: ok())
        dlg.grab_set(); self.win.wait_window(dlg)
        if not out["ok"]: return

        # Назви показників — завжди з заголовків стовпців
        self._compute_and_show(out["method"], out["alpha"], out["correction"])

    # ── Обчислення кореляцій ──────────────────────────────────
    def _compute_and_show(self, method, alpha, correction="bonferroni"):
        """
        Читає дані зі стовпців таблиці.
        Назви показників = заголовки стовпців (header_vars).
        Кожен стовпець = один показник, рядки = спостереження.
        """
        # Збираємо дані по стовпцях
        labels = []; data_cols = []
        for j in range(self.cols):
            col_name = self.header_vars[j].get().strip() or f"Показник {j+1}"
            col_vals = []
            for row in self.entries:
                v = row[j].get().strip() if j < len(row) else ""
                if not v: continue
                try: col_vals.append(float(v.replace(",",".")))
                except Exception: continue
            if len(col_vals) >= 3:
                labels.append(col_name)
                data_cols.append(col_vals)

        if len(data_cols) < 2:
            messagebox.showwarning("Замало даних",
                "Потрібно ≥ 2 стовпці з даними (≥ 3 значення у кожному).\n\n"
                "Переконайтесь що дані введені у числовому форматі "
                "і кожен стовпець містить хоча б 3 числа."); return

        n = len(labels)
        arrays = [np.array(d, dtype=float) for d in data_cols]
        n_mat  = np.zeros((n, n), dtype=int)

        # ── Авто-вибір / перевірка Пірсона ───────────────────
        actual_method = method
        if method in ("auto", "pearson"):
            non_normal = []
            for i, arr in enumerate(arrays):
                if len(arr) < 3: continue
                try:
                    _, p_sw = shapiro(arr)
                    if p_sw <= 0.05: non_normal.append(labels[i])
                except Exception: pass

            if method == "auto":
                if non_normal:
                    messagebox.showinfo("Авто-вибір методу",
                        f"Показники з ненормальним розподілом:\n"
                        f"{', '.join(non_normal[:5])}\n\n"
                        "Автоматично обрано: Спірмен (непараметричний).")
                    actual_method = "spearman"
                else:
                    messagebox.showinfo("Авто-вибір методу",
                        "Всі показники відповідають нормальному розподілу.\n"
                        "Автоматично обрано: Пірсон.")
                    actual_method = "pearson"
            elif method == "pearson" and non_normal:
                ans = messagebox.askyesno(
                    "Увага: нормальність порушена",
                    f"Показники з ненормальним розподілом:\n"
                    f"{', '.join(non_normal[:5])}\n\n"
                    "Кореляція Пірсона передбачає нормальний розподіл.\n"
                    "Рекомендується: Спірмен.\n\n"
                    "Продовжити з Пірсоном попри порушення?")
                if not ans: return

        # ── Попарна кореляційна матриця ───────────────────────
        r_mat = np.full((n, n), np.nan)
        p_mat = np.full((n, n), np.nan)
        np.fill_diagonal(r_mat, 1.0)
        np.fill_diagonal(p_mat, 1.0)
        np.fill_diagonal(n_mat, [len(a) for a in arrays])

        raw_p_pairs = []
        for i in range(n):
            for j in range(i+1, n):
                a = arrays[i]; b = arrays[j]
                min_len = min(len(a), len(b))
                a2 = a[:min_len]; b2 = b[:min_len]
                mask = ~(np.isnan(a2) | np.isnan(b2))
                a2 = a2[mask]; b2 = b2[mask]
                pair_n = len(a2)
                n_mat[i,j] = n_mat[j,i] = pair_n
                if pair_n < 3: continue
                try:
                    if actual_method == "pearson":
                        r_, p_ = pearsonr(a2, b2)
                    else:
                        r_, p_ = spearmanr(a2, b2)
                    r_mat[i,j] = r_mat[j,i] = float(r_)
                    raw_p_pairs.append((i, j, float(p_)))
                except Exception: pass

        if not raw_p_pairs:
            messagebox.showwarning("Замало даних",
                "Жодної пари показників не має ≥ 3 спільних спостережень."); return

        # ── Поправка на множинні порівняння ───────────────────
        m = len(raw_p_pairs)
        if correction == "bonferroni":
            for i, j, p_raw in raw_p_pairs:
                p_mat[i,j] = p_mat[j,i] = min(1.0, p_raw * m)
        elif correction == "bh":
            sp = sorted(raw_p_pairs, key=lambda x: x[2])
            bh = np.array([p for _,_,p in sp])
            for k in range(len(bh)-1,-1,-1):
                bh[k] = min(1.0, bh[k]*len(bh)/(k+1))
                if k < len(bh)-1: bh[k] = min(bh[k], bh[k+1])
            for idx,(i,j,_) in enumerate(sp):
                p_mat[i,j] = p_mat[j,i] = float(bh[idx])
        else:
            for i,j,p_raw in raw_p_pairs:
                p_mat[i,j] = p_mat[j,i] = p_raw

        corr_label = {"bonferroni":"Бонферроні","bh":"BH/FDR","none":"без поправки"
                      }.get(correction, correction)

        # ── Показуємо результати в одному вікні з вкладками ──
        self._show_results(labels, r_mat, p_mat, n_mat, alpha,
                           actual_method, corr_label, arrays)

    # ── Об'єднане вікно результатів ───────────────────────────
    def _show_results(self, labels, r_mat, p_mat, n_mat, alpha,
                      method, corr_label, arrays):
        if not HAS_MPL:
            messagebox.showwarning("","matplotlib недоступний."); return

        win = tk.Toplevel(self.win)
        win.title("Кореляційний аналіз — результати")
        win.geometry("1100x800"); set_icon(win)
        win.resizable(True, True)

        # Зберігаємо дані для перебудови
        self._res_win    = win
        self._res_labels = labels; self._res_r = r_mat
        self._res_p      = p_mat;  self._res_n = n_mat
        self._res_alpha  = alpha;  self._res_method = method
        self._res_corr   = corr_label; self._res_arrays = arrays
        self._hm_title   = getattr(self, '_hm_title', "")
        self._sc_title   = getattr(self, '_sc_title', "")

        meth_lbl = "Пірсон" if method == "pearson" else "Спірмен"

        # ── Перемикач вкладок через кнопки (не ttk.Notebook) ─
        # ttk.Notebook на Windows приховує текст активної вкладки
        switch_f = tk.Frame(win, bg="#1a4b8c"); switch_f.pack(fill=tk.X)

        content_f = tk.Frame(win); content_f.pack(fill=tk.BOTH, expand=True)

        self._hm_frame = None
        self._sc_frame = None
        _hm_outer = tk.Frame(content_f); _hm_outer.pack(fill=tk.BOTH, expand=True)
        _sc_outer = tk.Frame(content_f); _sc_outer.pack(fill=tk.BOTH, expand=True)
        _sc_outer.pack_forget()   # ховаємо другу

        _active_tab = [0]
        _btn_hm = _btn_sc = None

        def _show_hm():
            _sc_outer.pack_forget()
            _hm_outer.pack(fill=tk.BOTH, expand=True)
            _active_tab[0] = 0
            _btn_hm.configure(bg="white", fg="#1a4b8c", relief=tk.FLAT,
                               font=("Times New Roman",11,"bold"))
            _btn_sc.configure(bg="#1a4b8c", fg="#d0d8e8", relief=tk.FLAT,
                               font=("Times New Roman",11))
        def _show_sc():
            _hm_outer.pack_forget()
            _sc_outer.pack(fill=tk.BOTH, expand=True)
            _active_tab[0] = 1
            _btn_sc.configure(bg="white", fg="#1a4b8c", relief=tk.FLAT,
                               font=("Times New Roman",11,"bold"))
            _btn_hm.configure(bg="#1a4b8c", fg="#d0d8e8", relief=tk.FLAT,
                               font=("Times New Roman",11))

        _btn_hm = tk.Button(switch_f, text="  🌡  Теплова карта  ",
                            bg="white", fg="#1a4b8c",
                            font=("Times New Roman",11,"bold"),
                            relief=tk.FLAT, padx=16, pady=8,
                            cursor="hand2", command=_show_hm)
        _btn_hm.pack(side=tk.LEFT)
        _btn_sc = tk.Button(switch_f, text="  ⬡  Матриця розсіювання  ",
                            bg="#1a4b8c", fg="#d0d8e8",
                            font=("Times New Roman",11),
                            relief=tk.FLAT, padx=16, pady=8,
                            cursor="hand2", command=_show_sc)
        _btn_sc.pack(side=tk.LEFT)
        tk.Label(switch_f, bg="#1a4b8c",
                 text=f"  Метод: {meth_lbl}  |  α={alpha}  |  Поправка: {corr_label}",
                 font=("Times New Roman",9), fg="#a0b8cc").pack(side=tk.LEFT, padx=8)

        # ── Теплова карта ──────────────────────────────────────
        tb1 = tk.Frame(_hm_outer, bg="#f0f0f0", padx=6, pady=4); tb1.pack(fill=tk.X)
        for btxt, bcmd, bcol in [
            ("💾 Зберегти PNG", lambda: self._save_fig_png(self._hm_fig,"теплова_карта"), None),
            ("📋 Копіювати",    self._copy_heatmap, None),
            ("⚙ Налаштування", lambda: self._settings_heatmap(), "#1a4b8c"),
        ]:
            kw = {"bg": bcol, "fg": "white"} if bcol else {}
            tk.Button(tb1, text=btxt, font=("Times New Roman",10),
                      relief=tk.FLAT, padx=8, pady=3, cursor="hand2",
                      command=bcmd, **kw).pack(side=tk.RIGHT, padx=3)
        self._hm_frame = tk.Frame(_hm_outer)
        self._hm_frame.pack(fill=tk.BOTH, expand=True)

        # ── Матриця розсіювання ────────────────────────────────
        tb2 = tk.Frame(_sc_outer, bg="#f0f0f0", padx=6, pady=4); tb2.pack(fill=tk.X)
        for btxt, bcmd, bcol in [
            ("💾 Зберегти PNG", lambda: self._save_fig_png(self._sc_fig,"матриця_розсіювання"), None),
            ("📋 Копіювати",    lambda: self._copy_scatter(), None),
            ("⚙ Налаштування", lambda: self._settings_scatter(labels,arrays,method), "#1a4b8c"),
        ]:
            kw = {"bg": bcol, "fg": "white"} if bcol else {}
            tk.Button(tb2, text=btxt, font=("Times New Roman",10),
                      relief=tk.FLAT, padx=8, pady=3, cursor="hand2",
                      command=bcmd, **kw).pack(side=tk.RIGHT, padx=3)
        self._sc_frame = tk.Frame(_sc_outer)
        self._sc_frame.pack(fill=tk.BOTH, expand=True)

        # Будуємо обидва графіки одразу
        self._draw_heatmap(self._hm_frame, labels, r_mat, p_mat, n_mat,
                           alpha, method, corr_label, self.gs)
        self._draw_scatter(self._sc_frame, labels, arrays, method)

    def _save_fig_png(self, fig, name="графік"):
        if fig is None: messagebox.showwarning("","Спочатку виконайте аналіз."); return
        path = filedialog.asksaveasfilename(
            defaultextension=".png",
            filetypes=[("PNG","*.png"),("SVG","*.svg")],
            title=f"Зберегти {name}")
        if not path: return
        try:
            fig.savefig(path, dpi=150, bbox_inches="tight")
            messagebox.showinfo("Збережено", f"Збережено:\n{path}")
        except Exception as ex:
            messagebox.showerror("Помилка", str(ex))

    def _settings_heatmap(self):
        dlg = HeatmapSettingsDlg(self._res_win, self.gs)
        self._res_win.wait_window(dlg)
        if not dlg.result: return
        self.gs = dlg.result
        # Заголовок
        self._ask_title_hm()

    def _ask_title_hm(self):
        dlg = tk.Toplevel(self._res_win); dlg.title("Заголовок теплової карти")
        dlg.resizable(False, False); dlg.grab_set(); set_icon(dlg)
        tk.Label(dlg, text="Заголовок графіка:",
                 font=("Times New Roman",12)).pack(padx=16, pady=(14,4))
        tv = tk.StringVar(value=self._hm_title)
        te = tk.Entry(dlg, textvariable=tv, font=("Times New Roman",12), width=40)
        te.pack(padx=16, pady=4); te.focus_set()
        def _ok():
            self._hm_title = tv.get().strip()
            dlg.destroy()
            self._redraw_heatmap()
        tk.Button(dlg, text="Застосувати", bg="#c62828", fg="white",
                  font=("Times New Roman",12), command=_ok).pack(pady=(4,14))
        dlg.bind("<Return>", lambda e: _ok())
        center_win(dlg)

    def _redraw_heatmap(self):
        for w in self._hm_frame.winfo_children(): w.destroy()
        self._draw_heatmap(self._hm_frame,
                           self._res_labels, self._res_r, self._res_p,
                           self._res_n, self._res_alpha, self._res_method,
                           self._res_corr, self.gs)

    def _settings_scatter(self, labels, arrays, method):
        dlg = tk.Toplevel(self._res_win); dlg.title("Налаштування матриці розсіювання")
        dlg.resizable(False, False); dlg.grab_set(); set_icon(dlg)
        rf = ("Times New Roman",11)
        frm = tk.Frame(dlg, padx=16, pady=12); frm.pack()

        # Заголовок
        tk.Label(frm, text="Заголовок графіка:", font=rf
                 ).grid(row=0, column=0, sticky="w", pady=4)
        tv = tk.StringVar(value=self._sc_title)
        tk.Entry(frm, textvariable=tv, width=34, font=rf
                 ).grid(row=0, column=1, sticky="w", padx=8)

        sc = self._sc_gs
        params = [
            ("Колір точок:",  "sc_point_color", sc.get("sc_point_color","#4c72b0"), "color"),
            ("Колір тренду:", "sc_trend_color", sc.get("sc_trend_color","#c62828"), "color"),
            ("Колір гістограм:", "sc_hist_color", sc.get("sc_hist_color","#4c72b0"), "color"),
            ("Розмір точок:", "sc_point_size",  sc.get("sc_point_size",14), "spin"),
            ("Показати тренд:", "sc_show_trend", sc.get("sc_show_trend",True), "check"),
        ]
        vars_ = {}
        for ri, (lbl, key, default, wtype) in enumerate(params, 1):
            tk.Label(frm, text=lbl, font=rf).grid(row=ri, column=0, sticky="w", pady=3)
            if wtype == "color":
                v = tk.StringVar(value=default)
                vars_[key] = v
                def _pick(v2=v):
                    c = colorchooser.askcolor(color=v2.get(), parent=dlg)
                    if c and c[1]: v2.set(c[1])
                tk.Button(frm, text="Обрати", command=_pick, font=rf
                          ).grid(row=ri, column=1, sticky="w", padx=8)
            elif wtype == "spin":
                v = tk.IntVar(value=int(default))
                vars_[key] = v
                tk.Spinbox(frm, from_=2, to=30, textvariable=v, width=6, font=rf
                           ).grid(row=ri, column=1, sticky="w", padx=8)
            elif wtype == "check":
                v = tk.BooleanVar(value=bool(default))
                vars_[key] = v
                tk.Checkbutton(frm, variable=v
                               ).grid(row=ri, column=1, sticky="w", padx=8)

        def _apply():
            self._sc_title = tv.get().strip()
            for key, v in vars_.items():
                self._sc_gs[key] = v.get()
            dlg.destroy()
            self._redraw_scatter(labels, arrays, method)
        bf = tk.Frame(frm); bf.grid(row=len(params)+2, column=0, columnspan=2, pady=(12,0))
        tk.Button(bf, text="Застосувати", bg="#c62828", fg="white",
                  font=rf, command=_apply).pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Скасувати", font=rf,
                  command=dlg.destroy).pack(side=tk.LEFT)
        center_win(dlg)

    def _redraw_scatter(self, labels, arrays, method):
        for w in self._sc_frame.winfo_children(): w.destroy()
        self._draw_scatter(self._sc_frame, labels, arrays, method)

    def _restyle(self, win, labels, r_mat, p_mat, n_mat, alpha, method, corr_label):
        """Залишаємо для сумісності."""
        self._settings_heatmap()

    def _draw_heatmap(self, frame, labels, r_mat, p_mat, n_mat, alpha, method, corr_label, gs):
        n = len(labels)
        dpi = 96
        # Стартовий розмір — буде масштабований через resize_event
        fig  = Figure(figsize=(8, 7), dpi=dpi)
        ax   = fig.add_subplot(111)

        cmap_name = gs.get("heatmap_cmap","RdYlGn")
        fsize     = gs.get("heatmap_font_size", 9)
        acol      = gs.get("heatmap_annot_color","#000000")
        ff        = gs.get("font_family","Times New Roman")

        try:    cmap = matplotlib.cm.get_cmap(cmap_name)
        except: cmap = matplotlib.cm.get_cmap("RdYlGn")

        masked = np.ma.array(r_mat, mask=np.isnan(r_mat))
        im = ax.imshow(masked, cmap=cmap, vmin=-1, vmax=1, aspect="auto")

        meth_full = "Пірсон" if method=="pearson" else "Спірмен"
        custom_title = getattr(self, '_hm_title', "")
        ax.set_title(
            custom_title if custom_title else
            f"Кореляційна матриця ({meth_full}, {corr_label}, α={alpha})\nКлітинки: r / p-скор / n",
            fontsize=fsize+1, fontfamily=ff)
        ax.set_xticks(range(n))
        ax.set_xticklabels(labels, rotation=45, ha="right", fontsize=fsize, fontfamily=ff)
        ax.set_yticks(range(n))
        ax.set_yticklabels(labels, fontsize=fsize, fontfamily=ff)

        for i in range(n):
            for j in range(n):
                r_ = r_mat[i,j]; p_ = p_mat[i,j]
                if i == j:
                    ax.text(j,i,"—",ha="center",va="center",
                            fontsize=fsize,color=acol,fontfamily=ff)
                    continue
                if math.isnan(r_): continue
                mark   = sig_mark(p_) if not math.isnan(p_) else ""
                p_str  = fmt(p_,3) if not math.isnan(p_) else "н/д"
                n_ij   = int(n_mat[i,j]) if n_mat is not None else 0
                txt_   = f"{r_:.2f}{mark}\np={p_str}\nn={n_ij}"
                ax.text(j,i,txt_,ha="center",va="center",
                        fontsize=max(6,fsize-1),color=acol,fontfamily=ff,linespacing=1.3)

        cbar = fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
        cbar.set_label("r", fontsize=fsize, fontfamily=ff)
        fig.tight_layout()
        self._hm_fig = fig

        embed_figure(fig, frame, dpi=dpi)

    # ── Матриця діаграм розсіювання ───────────────────────────
    def _draw_scatter(self, frame, labels, arrays, method):
        if not HAS_MPL: return
        n = len(labels)
        if n > 8: labels = labels[:6]; arrays = arrays[:6]; n = 6
        if n < 2: return
        self._sc_gs = getattr(self, "_sc_gs", {
            "sc_point_color":  "#4c72b0",
            "sc_trend_color":  "#c62828",
            "sc_hist_color":   "#4c72b0",
            "sc_point_size":   14,
            "sc_point_alpha":  0.75,
            "sc_show_trend":   True,
            "sc_trend_width":  0.9,
            "font_family":     "Times New Roman",
            "sc_font_size":    6,
        })
        self._sc_labels = labels; self._sc_arrays = arrays; self._sc_method = method
        # Малюємо безпосередньо в переданий frame
        for w in frame.winfo_children(): w.destroy()
        sc = self._sc_gs
        pt_col   = sc.get("sc_point_color",  "#4c72b0")
        tr_col   = sc.get("sc_trend_color",  "#c62828")
        hi_col   = sc.get("sc_hist_color",   "#4c72b0")
        pt_size  = sc.get("sc_point_size",   14)
        pt_alpha = sc.get("sc_point_alpha",  0.75)
        show_tr  = sc.get("sc_show_trend",   True)
        tr_width = sc.get("sc_trend_width",  0.9)
        ff       = sc.get("font_family",     "Times New Roman")
        fz       = sc.get("sc_font_size",    6)

        dpi = 96
        fig = Figure(dpi=dpi)

        for i in range(n):
            for j in range(n):
                ax = fig.add_subplot(n, n, i*n+j+1)
                if i == j:
                    a = arrays[i][~np.isnan(arrays[i])]
                    if len(a) > 0:
                        ax.hist(a, bins=max(4, int(np.sqrt(len(a)))),
                                color=hi_col, alpha=0.8, edgecolor="white", linewidth=0.4)
                    ax.set_title(labels[i], fontsize=fz+1, pad=2, fontfamily=ff)
                else:
                    xi = arrays[j]; yi = arrays[i]
                    mn = min(len(xi), len(yi))
                    xi = xi[:mn]; yi = yi[:mn]
                    mask = ~(np.isnan(xi) | np.isnan(yi))
                    xi = xi[mask]; yi = yi[mask]
                    ax.scatter(xi, yi, s=pt_size, alpha=pt_alpha,
                               color=pt_col, edgecolors="none")
                    if show_tr and len(xi) >= 3:
                        try:
                            z  = np.polyfit(xi, yi, 1)
                            xl = np.linspace(xi.min(), xi.max(), 50)
                            ax.plot(xl, np.poly1d(z)(xl),
                                    color=tr_col, lw=tr_width, alpha=0.9)
                        except Exception: pass
                ax.tick_params(labelsize=fz)
                if j == 0: ax.set_ylabel(labels[i], fontsize=fz, fontfamily=ff)
                if i == n-1: ax.set_xlabel(labels[j], fontsize=fz, fontfamily=ff)
                ax.spines["top"].set_visible(False)
                ax.spines["right"].set_visible(False)

        meth_full = "Пірсон" if method == "pearson" else "Спірмен"
        custom_sc = getattr(self, '_sc_title', '')
        fig.suptitle(
            custom_sc if custom_sc else f"Матриця діаграм розсіювання ({meth_full})",
            fontsize=9, fontfamily=ff, y=1.0)
        try: fig.tight_layout()
        except Exception: pass
        self._sc_fig = fig

        embed_figure(fig, frame, dpi=dpi)

    def _copy_scatter(self, win=None):
        if self._sc_fig is None:
            messagebox.showwarning("","Графік ще не побудований."); return
        ok_, msg = _copy_fig_to_clipboard(self._sc_fig)
        if ok_: messagebox.showinfo("","Матрицю розсіювання скопійовано.\nВставте у Word через Ctrl+V.")
        else:   messagebox.showwarning("",f"Помилка: {msg}")

    def _restyle_scatter(self, win, labels, arrays, method):
        dlg = ScatterSettingsDlg(win, self._sc_gs)
        win.wait_window(dlg)
        if dlg.result:
            self._sc_gs = dlg.result
            self._draw_scatter(self._sc_frame, labels, arrays, method)



# ═══════════════════════════════════════════════════════════════
# GUI — SADTk
# ═══════════════════════════════════════════════════════════════
class SADTk:
    SEL_BG = "#cce5ff"; SEL_ANC = "#99ccff"; ACT_BG = "#fff3c4"

    def __init__(self, root):
        self.root = root
        root.title("S.A.D. — Статистичний аналіз даних")
        root.geometry("1000x580"); set_icon(root)
        try: ttk.Style().theme_use("clam")
        except Exception: pass

        mf = tk.Frame(root, bg="white"); mf.pack(expand=True, fill=tk.BOTH)
        tk.Label(mf, text="S.A.D. — Статистичний аналіз даних",
                 font=("Times New Roman", 20, "bold"), fg="#000000", bg="white").pack(pady=16)

        bf = tk.Frame(mf, bg="white"); bf.pack(pady=8)
        for i, (txt, fc) in enumerate([("Однофакторний аналіз", 1), ("Двофакторний аналіз", 2),
                                        ("Трифакторний аналіз", 3), ("Чотирифакторний аналіз", 4)]):
            tk.Button(bf, text=txt, width=22, height=2, font=("Times New Roman", 13),
                      command=lambda f=fc: self.open_table(f)).grid(row=i // 2, column=i % 2, padx=10, pady=6)

        tk.Button(mf, text="Кореляційний аналіз", width=30, height=2, font=("Times New Roman", 13),
                  bg="#1a4b8c", fg="white",
                  command=self.open_correlation).pack(pady=(4, 10))

        tk.Label(mf, text="Виберіть тип аналізу → Введіть дані → Аналіз даних",
                 font=("Times New Roman", 12), fg="#555555", bg="white").pack(pady=4)

        self.table_win = None; self.report_win = None; self.graph_win = None
        self._graph_figs = {}
        self._active_cell = None; self._active_prev = None
        self._sel_anchor = None; self._sel_cells = set(); self._sel_orig = {}
        self._fill_drag = False; self._fill_rows = []; self._fill_cols = []
        self.factor_title_map = {}
        self.graph_settings = dict(DEF_GS)
        self._current_project_path = None
        self._lbf_cache = {}

    # ── open correlation ──────────────────────────────────────
    def open_correlation(self):
        CorrelationWindow(self.root, self.graph_settings)

    def _show_anova_help(self, parent, fc):
        """Довідка специфічна для кожного типу ANOVA."""
        help_texts = {
            1: """
ОДНОФАКТОРНИЙ ДИСПЕРСІЙНИЙ АНАЛІЗ
═══════════════════════════════════

ЩО РОБИТЬ?
  Порівнює СЕРЕДНІ по 3 і більше групах (варіантах).
  Перевіряє: «Чи хоча б одна група відрізняється від інших?»

  Перевага над t-тестом: при кількох порівняннях одночасно
  не накопичується помилка першого роду.

СТРУКТУРА ТАБЛИЦІ:
  Стовпець 1: Фактор A (назви варіантів/обробок)
  Стовпці 2+: Повторності (числові значення)

  Приклад (3 дози добрива, 4 повторності):
  | Доза    | Повт.1 | Повт.2 | Повт.3 | Повт.4 |
  | Контроль|  5.2   |  4.8   |  5.5   |  4.9   |
  | Доза 1  |  6.1   |  6.4   |  5.9   |  6.3   |
  | Доза 2  |  7.2   |  6.8   |  7.5   |  7.1   |

  Перейменуйте «Фактор A» (подвійний клік) на назву вашого фактора.

ВИБІР ДИЗАЙНУ:
  CRD — ділянки однорідні, варіанти розміщені випадково
  RCBD — є блоки (рельєф, родючість) — ефективніший
  Split-plot — лише для 2+ факторів

ІНТЕРПРЕТАЦІЯ:
  F-тест значущий → є різниця між варіантами
  Переходьте до пост-хок для визначення ЯКИХ САМЕ
""",
            2: """
ДВОФАКТОРНИЙ ДИСПЕРСІЙНИЙ АНАЛІЗ
══════════════════════════════════

ЩО РОБИТЬ?
  Оцінює вплив ДВОХ факторів (A і B) та їх взаємодії (A×B).

  Три питання одночасно:
  1. Чи значущий головний ефект фактора A?
  2. Чи значущий головний ефект фактора B?
  3. Чи є взаємодія A×B? (найважливіше!)

СТРУКТУРА ТАБЛИЦІ:
  Стовпець 1: Фактор A (рівні першого фактора)
  Стовпець 2: Фактор B (рівні другого фактора)
  Стовпці 3+: Повторності (числові значення)

  Приклад (сорт × доза добрива):
  | Сорт    | Доза  | Повт.1 | Повт.2 | Повт.3 |
  | Сорт А  | Доза1 |  5.2   |  4.8   |  5.5   |
  | Сорт А  | Доза2 |  6.1   |  6.4   |  5.9   |
  | Сорт Б  | Доза1 |  4.9   |  5.1   |  4.7   |
  | Сорт Б  | Доза2 |  7.2   |  6.8   |  7.5   |

ВЗАЄМОДІЯ A×B:
  Значуща → ефект фактора A залежить від рівня B.
  Тобто: один сорт краще реагує на добрива, інший — ні.
  При значущій взаємодії — аналізуйте прості ефекти!
  Незначуща → ефекти факторів незалежні.
""",
            3: """
ТРИФАКТОРНИЙ ДИСПЕРСІЙНИЙ АНАЛІЗ
══════════════════════════════════

ЩО РОБИТЬ?
  Оцінює вплив трьох факторів (A, B, C) та їх взаємодій:
  A×B, A×C, B×C, і потрійну взаємодію A×B×C.

СТРУКТУРА ТАБЛИЦІ:
  Стовпець 1: Фактор A
  Стовпець 2: Фактор B
  Стовпець 3: Фактор C
  Стовпці 4+: Повторності

  Кожна унікальна комбінація A×B×C = один рядок.
  При 3 рівнях кожного фактора: 3×3×3 = 27 рядків.

ВЗАЄМОДІЯ A×B×C:
  Значуща → ефект пари факторів залежить від третього.
  Наприклад: вплив добрива і сорту різний в різні роки.

ПОРАДА:
  При 3 факторах рекомендується Тип III SS.
  Кількість рядків = k_A × k_B × k_C (де k = кількість рівнів).

═══════════════════════════════════════════════════════════
ОСОБЛИВИЙ ВИПАДОК: ЛАТИНСЬКИЙ КВАДРАТ (у 3-факторному)
═══════════════════════════════════════════════════════════

Якщо у вас Латинський квадрат — використовуйте саме
3-факторний аналіз з таким призначенням факторів:

  Фактор A = Варіант  (назви обробок: N60, N90, N120...)
  Фактор B = Рядок    (номери рядів поля: 1, 2, 3...)
  Фактор C = Стовпець (номери стовпців поля: 1, 2, 3...)

Структура таблиці (k=4 варіанти, 4×4 квадрат):
  | Варіант | Рядок | Стовп | Значення |
  | N60     |   1   |   1   |   18.4   |
  | N90     |   1   |   2   |   21.3   |
  | N120    |   1   |   3   |   22.8   |
  | N0      |   1   |   4   |   15.2   |
  | N0      |   2   |   1   |   14.8   |
  ...

Кожен рядок таблиці = одна ділянка = одне значення.
Стовпець "Значення" вводьте у колонку «Повт.1».

У вікні параметрів оберіть:  Дизайн → Латинський квадрат

Програма автоматично:
  ✓ Виносить SS_рядки і SS_стовпці з помилки
  ✓ Розраховує правильний df = (k-1)(k-2)
  ✓ Перевіряє що k варіантів = k рядків = k стовпців
""",
            4: """
ЧОТИРИФАКТОРНИЙ ДИСПЕРСІЙНИЙ АНАЛІЗ
═════════════════════════════════════

ЩО РОБИТЬ?
  Оцінює вплив чотирьох факторів (A, B, C, D) та всіх
  можливих взаємодій (4 парних + 4 потрійних + 1 четверна).

СТРУКТУРА ТАБЛИЦІ:
  Стовпці 1-4: Фактори A, B, C, D
  Стовпці 5+: Повторності

  При 2 рівнях кожного: 2×2×2×2 = 16 комбінацій.
  При 3 рівнях кожного: 3×3×3×3 = 81 комбінація!

ВАЖЛИВО:
  Чотирифакторний аналіз вимагає дуже великого досліду.
  Мінімум: 2 повторності × 16 комбінацій = 32 спостереження.
  Рекомендується Тип III SS.
  Потрійні і четверна взаємодії рідко бувають значущими
  і важко інтерпретуються.

ПОРАДА:
  При незначущих вищих взаємодіях — спростіть до трифакторного.
"""
        }
        text = help_texts.get(fc, "Довідка недоступна.")

        win = tk.Toplevel(parent)
        win.title(f"Довідка — {['','Одно','Дво','Три','Чотири'][fc]}факторний ANOVA")
        win.geometry("680x540"); set_icon(win)
        frm = tk.Frame(win); frm.pack(fill=tk.BOTH, expand=True, padx=8, pady=6)
        vsb = ttk.Scrollbar(frm, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        txt = tk.Text(frm, wrap="word", font=("Times New Roman",11),
                      yscrollcommand=vsb.set, relief=tk.FLAT,
                      bg="#fafafa", padx=10, pady=8, cursor="arrow")
        txt.pack(fill=tk.BOTH, expand=True); vsb.config(command=txt.yview)
        txt.insert("1.0", text.strip()); txt.configure(state="disabled")
        txt.bind("<MouseWheel>",
                 lambda e: txt.yview_scroll(int(-1*(e.delta/120)),"units"))
        tk.Button(win, text="Закрити", command=win.destroy,
                  font=("Times New Roman",11)).pack(pady=6)

    # ── factor titles ─────────────────────────────────────────
    def ftitle(self, fk): return self.factor_title_map.get(fk, f"Фактор {fk}")
    def _set_ftitle(self, fk, t): self.factor_title_map[fk] = t.strip() or f"Фактор {fk}"

    # ── project save/load ─────────────────────────────────────
    def save_project(self):
        if not hasattr(self, "entries") or not self.entries:
            messagebox.showwarning("", "Відкрийте таблицю."); return
        path = filedialog.asksaveasfilename(
            parent=self.table_win or self.root, title="Зберегти проект",
            defaultextension=".sadp",
            filetypes=[("SAD проект", "*.sadp"), ("JSON", "*.json"), ("Усі", "*.*")])
        if not path: return
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(project_to_dict(self), f, ensure_ascii=False, indent=2)
            self._current_project_path = path
            messagebox.showinfo("Збережено", f"Проект збережено:\n{path}")
        except Exception as ex: messagebox.showerror("Помилка", str(ex))

    def load_project(self):
        path = filedialog.askopenfilename(
            parent=self.table_win or self.root, title="Відкрити проект",
            filetypes=[("SAD проект", "*.sadp"), ("JSON", "*.json"), ("Усі", "*.*")])
        if not path: return
        try:
            with open(path, "r", encoding="utf-8") as f: d = json.load(f)
            project_from_dict(self, d)
            self._current_project_path = path
            messagebox.showinfo("Відкрито", f"Проект відкрито:\n{path}")
        except Exception as ex: messagebox.showerror("Помилка", str(ex))

    def clear_project(self):
        if not messagebox.askyesno("Очистити", "Очистити всі дані таблиці?"): return
        for row in self.entries:
            for e in row: e.delete(0, tk.END)

    # ── selection ─────────────────────────────────────────────
    def _clear_sel(self):
        # No visual highlighting — just clear tracking state
        self._sel_cells.clear(); self._sel_anchor = None; self._sel_orig.clear()

    def _restore_bg(self, r, c):
        pass   # no-op: no coloring was applied

    def _apply_sel(self, cells):
        pass   # no-op: selection tracking only, no visual highlight

    def _sel_range(self, r1, c1, r2, c2):
        new = {(r, c) for r in range(min(r1, r2), max(r1, r2) + 1)
               for c in range(min(c1, c2), max(c1, c2) + 1)
               if r < len(self.entries) and c < len(self.entries[r])}
        self._sel_cells = new

    def _sel_bounds(self):
        if not self._sel_cells: return None
        rs = [r for r, c in self._sel_cells]; cs = [c for r, c in self._sel_cells]
        return min(rs), min(cs), max(rs), max(cs)

    def _near_br(self, w, mg=6):
        try:
            px = w.winfo_pointerx(); py = w.winfo_pointery()
            x0 = w.winfo_rootx(); y0 = w.winfo_rooty()
            return (x0 + w.winfo_width() - mg <= px <= x0 + w.winfo_width()) and \
                   (y0 + w.winfo_height() - mg <= py <= y0 + w.winfo_height())
        except Exception: return False

    def _sel_handle_cell(self):
        b = self._sel_bounds()
        if b is None: return None
        try: return self.entries[b[2]][b[3]]
        except Exception: return None

    # ── active cell ───────────────────────────────────────────
    def _set_active(self, w):
        if self._active_cell is w: return
        if isinstance(self._active_cell, tk.Entry) and self._active_prev:
            try: self._active_cell.configure(**self._active_prev)
            except Exception: pass
        self._active_cell = w
        if isinstance(w, tk.Entry):
            self._active_prev = {"bg": w.cget("bg"), "highlightthickness": int(w.cget("highlightthickness")),
                                  "highlightbackground": w.cget("highlightbackground"),
                                  "highlightcolor": w.cget("highlightcolor"),
                                  "relief": w.cget("relief"),
                                  "bd": int(w.cget("bd")) if str(w.cget("bd")).isdigit() else 1}
            try: w.configure(bg=self.ACT_BG, highlightthickness=3,
                              highlightbackground="#c62828", highlightcolor="#c62828",
                              relief=tk.SOLID, bd=1)
            except Exception: pass

    # ── bind cell ─────────────────────────────────────────────
    def bind_cell(self, e):
        e.bind("<Return>",               self._on_enter)
        e.bind("<Up>",                   self._on_arrow)
        e.bind("<Down>",                 self._on_arrow)
        e.bind("<Left>",                 self._on_arrow)
        e.bind("<Right>",                self._on_arrow)
        e.bind("<Control-c>",            self._on_copy)
        e.bind("<Control-C>",            self._on_copy)
        e.bind("<Control-v>",            self._on_paste)
        e.bind("<Control-V>",            self._on_paste)
        e.bind("<FocusIn>",              lambda ev: self._set_active(ev.widget))
        e.bind("<ButtonPress-1>",        self._on_press)
        e.bind("<B1-Motion>",            self._on_drag)
        e.bind("<ButtonRelease-1>",      self._on_release)
        e.bind("<Motion>",               self._on_motion)
        e.bind("<Leave>",                lambda ev: ev.widget.configure(cursor=""))
        e.bind("<Shift-ButtonPress-1>",  self._on_shift_click)

    # ── mouse events ──────────────────────────────────────────
    def _on_motion(self, event):
        w = event.widget
        if not isinstance(w, tk.Entry): return
        pos = self._pos(w)
        if not pos: return
        r, c = pos
        if c >= self.factors_count: w.configure(cursor=""); return
        if self._sel_cells and self._sel_handle_cell() is w and self._near_br(w):
            w.configure(cursor="crosshair")
        elif not self._sel_cells and self._near_br(w):
            w.configure(cursor="crosshair")
        else:
            w.configure(cursor="")

    def _on_press(self, event):
        w = event.widget
        if not isinstance(w, tk.Entry): return
        pos = self._pos(w)
        if not pos: return
        r, c = pos
        if c < self.factors_count:
            if self._sel_cells and self._sel_handle_cell() is w and self._near_br(w):
                self._start_fill(use_sel=True); return "break"
            if not self._sel_cells and self._near_br(w):
                self._clear_sel()
                self._sel_anchor = (r, c); self._sel_cells = {(r, c)}
                self._sel_orig[(r, c)] = w.cget("bg"); self._apply_sel({(r, c)})
                self._start_fill(use_sel=False); return "break"
        self._fill_drag = False
        self._clear_sel()
        self._sel_anchor = (r, c); self._sel_cells = {(r, c)}
        self._sel_orig[(r, c)] = w.cget("bg"); self._apply_sel({(r, c)})
        w.focus_set()

    def _on_shift_click(self, event):
        w = event.widget; pos = self._pos(w)
        if not pos or self._sel_anchor is None: return
        ar, ac = self._sel_anchor; r, c = pos
        self._sel_range(ar, ac, r, c); return "break"

    def _on_drag(self, event):
        w = event.widget
        if not isinstance(w, tk.Entry): return
        if self._fill_drag: self._do_fill(event); return "break"
        if self._sel_anchor is None: return
        ar, ac = self._sel_anchor
        pos = self._pos(w)
        if pos:
            r, c = pos
        else:
            py = w.winfo_pointery(); px = w.winfo_pointerx(); r, c = ar, ac
            for ri in range(len(self.entries)):
                for ci in range(len(self.entries[ri])):
                    cell = self.entries[ri][ci]
                    if (cell.winfo_rootx() <= px <= cell.winfo_rootx() + cell.winfo_width() and
                            cell.winfo_rooty() <= py <= cell.winfo_rooty() + cell.winfo_height()):
                        r, c = ri, ci; break
        self._sel_range(ar, ac, r, c)

    def _on_release(self, event):
        if self._fill_drag:
            self._fill_drag = False; self._fill_rows = []; self._fill_cols = []; return "break"

    # ── fill drag ─────────────────────────────────────────────
    def _start_fill(self, use_sel):
        self._fill_drag = True
        if use_sel and self._sel_cells:
            b = self._sel_bounds()
            if b is None: self._fill_drag = False; return
            self._fill_rows = list(range(b[0], b[2] + 1))
            self._fill_cols = list(range(b[1], b[3] + 1))
        elif self._sel_anchor:
            self._fill_rows = [self._sel_anchor[0]]
            self._fill_cols = [self._sel_anchor[1]]
        else:
            self._fill_drag = False

    def _do_fill(self, event):
        w = event.widget
        if not isinstance(w, tk.Entry) or not self._fill_rows or not self._fill_cols: return
        last_src = self._fill_rows[-1]
        py = w.winfo_pointery(); target = last_src
        for rr in range(last_src, len(self.entries)):
            cell = self.entries[rr][self._fill_cols[0]]
            y0 = cell.winfo_rooty()
            if y0 <= py <= y0 + cell.winfo_height(): target = rr; break
        else:
            if py > self.entries[-1][self._fill_cols[0]].winfo_rooty(): target = len(self.entries)
        if target <= last_src: return
        n_src = len(self._fill_rows)
        dst = last_src + 1
        while dst <= target:
            while dst >= len(self.entries): self.add_row()
            src_r = self._fill_rows[(dst - last_src - 1) % n_src]
            for c in self._fill_cols:
                if c >= self.factors_count: break
                self.entries[dst][c].delete(0, tk.END)
                self.entries[dst][c].insert(0, self.entries[src_r][c].get())
            dst += 1
        self.inner.update_idletasks()
        self.canvas.config(scrollregion=self.canvas.bbox("all"))

    # ── copy selection Ctrl+C ─────────────────────────────────
    def _on_copy(self, event=None):
        if not self._sel_cells:
            w = event.widget if event else self.table_win.focus_get()
            if isinstance(w, tk.Entry):
                try: sel = w.get("sel.first", "sel.last")
                except Exception: sel = w.get()
                self.table_win.clipboard_clear(); self.table_win.clipboard_append(sel)
            return "break"
        b = self._sel_bounds()
        if b is None: return "break"
        r1, c1, r2, c2 = b
        lines = []
        for r in range(r1, r2 + 1):
            row = []
            for c in range(c1, c2 + 1):
                try: row.append(self.entries[r][c].get())
                except Exception: row.append("")
            lines.append("\t".join(row))
        self.table_win.clipboard_clear(); self.table_win.clipboard_append("\n".join(lines))
        return "break"

    # ── navigation ────────────────────────────────────────────
    def _on_enter(self, event=None):
        pos = self._pos(event.widget)
        if not pos: return "break"
        i, j = pos; ni = i + 1
        if ni >= len(self.entries): self.add_row()
        self.entries[ni][j].focus_set(); self.entries[ni][j].icursor(tk.END); return "break"

    def _on_arrow(self, event=None):
        pos = self._pos(event.widget)
        if not pos: return "break"
        i, j = pos
        if event.keysym == "Up":    i = max(0, i - 1)
        elif event.keysym == "Down":  i = min(len(self.entries) - 1, i + 1)
        elif event.keysym == "Left":  j = max(0, j - 1)
        elif event.keysym == "Right": j = min(len(self.entries[i]) - 1, j + 1)
        self.entries[i][j].focus_set(); self.entries[i][j].icursor(tk.END); return "break"

    def _on_paste(self, event=None):
        widget = event.widget if event else self.table_win.focus_get()
        if not isinstance(widget, tk.Entry): return "break"
        try: data = self.table_win.clipboard_get()
        except Exception: return "break"
        pos = self._pos(widget)
        if not pos: return "break"
        r0, c0 = pos
        for ir, rt in enumerate([r for r in data.splitlines() if r != ""]):
            for jc, val in enumerate(rt.split("\t")):
                rr = r0 + ir; cc = c0 + jc
                while rr >= len(self.entries): self.add_row()
                if cc >= self.cols: continue
                self.entries[rr][cc].delete(0, tk.END); self.entries[rr][cc].insert(0, val)
        return "break"

    def _pos(self, widget):
        for i, row in enumerate(self.entries):
            for j, cell in enumerate(row):
                if cell is widget: return i, j
        return None

    def _mk_entry(self, parent):
        return tk.Entry(parent, width=COL_W, fg="#000000", font=("Times New Roman", 12),
                        highlightthickness=1, highlightbackground="#c0c0c0", highlightcolor="#c0c0c0")

    # ── rename factor ─────────────────────────────────────────
    def rename_factor(self, col):
        if col < 0 or col >= self.factors_count: return
        fk = self.factor_keys[col]; old = self.ftitle(fk)
        dlg = tk.Toplevel(self.table_win or self.root)
        dlg.title("Перейменувати фактор"); dlg.resizable(False, False)
        set_icon(dlg); dlg.grab_set()
        tk.Label(dlg, text=f"Назва фактору {fk}:",
                 font=("Times New Roman",12)).pack(padx=16, pady=14)
        var = tk.StringVar(value=old)
        e = tk.Entry(dlg, textvariable=var, font=("Times New Roman",12), width=28)
        e.pack(padx=16, pady=4); e.select_range(0, tk.END); e.focus_set()
        def ok():
            new = var.get().strip()
            if not new: return
            self._set_ftitle(fk, new)
            if col < len(self.header_labels):
                self.header_labels[col].configure(text=new)
            dlg.destroy()
        tk.Button(dlg, text="OK", bg="#c62828", fg="white",
                  font=("Times New Roman",12), command=ok).pack(pady=(4,14))
        dlg.bind("<Return>", lambda ev: ok())
        center_win(dlg)
        dlg.update_idletasks(); center_win(dlg); dlg.bind("<Return>", lambda ev: ok()); dlg.grab_set()

    def show_design_help(self):
        w = tk.Toplevel(self.root); w.title("Пояснення дизайнів"); w.resizable(False, False); set_icon(w)
        frm = tk.Frame(w, padx=16, pady=14); frm.pack()
        txt = ("CRD — Повна рандомізація\nRCBD — Блочна рандомізація\nSplit-plot — Спліт-плот\n\n"
               "SS Тип I: послідовний (порядок важливий)\n"
               "SS Тип II: ієрархічний (без взаємодій)\n"
               "SS Тип III: частковий (з урахуванням всіх ефектів) — рекомендований")
        t = tk.Text(frm, width=55, height=10, wrap="word"); t.insert("1.0", txt)
        t.configure(state="disabled"); t.pack()
        tk.Button(frm, text="OK", width=10, command=w.destroy).pack(pady=(10, 0))
        w.update_idletasks(); center_win(w); w.grab_set()

    def show_about(self):
        messagebox.showinfo("Розробник",
            f"S.A.D. — Статистичний аналіз даних  v{APP_VER}\n"
            "Розробник: Чаплоуцький Андрій Миколайович\n"
            "Уманський національний університет")

    # ══════════════════════════════════════════════════════════
    # OPEN TABLE  — with menu bar
    # ══════════════════════════════════════════════════════════
    def open_table(self, fc):
        if self.table_win and tk.Toplevel.winfo_exists(self.table_win):
            self.table_win.destroy()
        self.factors_count = fc
        self.factor_keys   = ["A", "B", "C", "D"][:fc]
        for fk in self.factor_keys:
            if fk not in self.factor_title_map: self._set_ftitle(fk, f"Фактор {fk}")

        self.table_win = tw = tk.Toplevel(self.root)
        factor_names = {1:"Однофакторний", 2:"Двофакторний",
                        3:"Трифакторний",  4:"Чотирифакторний"}
        # Для 3-факторного — додаємо опис ЛК у довідку
        tw.title(f"S.A.D. — {factor_names.get(fc,str(fc)+'-факторний')} дисперсійний аналіз")
        tw.geometry("1280x720"); set_icon(tw)

        # ── Toolbar ───────────────────────────────────────────
        ctl = tk.Frame(tw, padx=6, pady=4); ctl.pack(fill=tk.X)

        tk.Button(ctl, text="▶ Аналіз", bg="#c62828", fg="white",
                  font=("Times New Roman", 13),
                  command=self.analyze).pack(side=tk.LEFT, padx=4)

        # Налаштування — спадне меню
        mb2 = tk.Menubutton(ctl, text="⚙ Налаштування ▾",
                            font=("Times New Roman", 11),
                            relief=tk.RAISED, bd=2)
        mb2.pack(side=tk.LEFT, padx=4)
        sm2 = tk.Menu(mb2, tearoff=0)
        sm2.add_command(label="Додати рядок",      command=self.add_row)
        sm2.add_command(label="Видалити рядок",    command=self.delete_row)
        sm2.add_separator()
        sm2.add_command(label="Додати стовпчик",   command=self.add_column)
        sm2.add_command(label="Видалити стовпчик", command=self.delete_column)
        sm2.add_separator()
        sm2.add_command(label="💾 Зберегти проект", command=self.save_project)
        sm2.add_command(label="📂 Відкрити проект", command=self.load_project)
        sm2.add_separator()
        sm2.add_command(label="🗑 Очистити таблицю", command=self.clear_project)
        mb2["menu"] = sm2

        tk.Button(ctl, text="Вставити з буфера",
                  font=("Times New Roman", 11),
                  command=self._paste_from_focus).pack(side=tk.LEFT, padx=4)
        tk.Button(ctl, text="📚 Довідка",
                  bg="#1a4b8c", fg="white",
                  font=("Times New Roman", 11),
                  command=lambda: self._show_anova_help(tw, fc)).pack(side=tk.LEFT, padx=4)

        # Підказка
        tk.Label(ctl,
                 text="Подвійний клік на синьому заголовку фактора → перейменувати",
                 font=("Times New Roman",9), fg="#666").pack(side=tk.LEFT, padx=8)

        # ── Table canvas ──────────────────────────────────────
        self.canvas = tk.Canvas(tw)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb = ttk.Scrollbar(tw, orient="vertical", command=self.canvas.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y); self.canvas.configure(yscrollcommand=sb.set)
        self.inner = tk.Frame(self.canvas)
        self.canvas.create_window((0, 0), window=self.inner, anchor="nw")
        # mousewheel scroll
        def _mw(e): self.canvas.yview_scroll(int(-1 * (e.delta / 120)), "units")
        self.canvas.bind_all("<MouseWheel>", _mw)

        self.rows = 12; self.cols = fc + 6
        self.entries = []; self.header_labels = []
        col_names = [self.ftitle(fk) for fk in self.factor_keys] + [f"Повт.{i+1}" for i in range(6)]
        for j, nm in enumerate(col_names):
            is_factor = j < fc
            lbl = tk.Label(self.inner, text=nm, relief=tk.RIDGE, width=COL_W,
                           bg="#1a4b8c" if is_factor else "#2e6b2e",
                           fg="white",
                           cursor="hand2" if is_factor else "arrow",
                           font=("Times New Roman", 12, "bold"))
            lbl.grid(row=0, column=j, padx=2, pady=2, sticky="nsew")
            self.header_labels.append(lbl)
            if is_factor:
                lbl.bind("<Double-Button-1>", lambda e, c=j: self.rename_factor(c))
        for i in range(self.rows):
            row_e = []
            for j in range(self.cols):
                e = self._mk_entry(self.inner); e.grid(row=i + 1, column=j, padx=2, pady=2)
                self.bind_cell(e); row_e.append(e)
            self.entries.append(row_e)
        self.inner.update_idletasks(); self.canvas.config(scrollregion=self.canvas.bbox("all"))
        self.entries[0][0].focus_set()
        tw.bind("<Control-v>", self._on_paste); tw.bind("<Control-V>", self._on_paste)

    # ── table editing ─────────────────────────────────────────
    def add_row(self):
        i = len(self.entries); row_e = []
        for j in range(self.cols):
            e = self._mk_entry(self.inner); e.grid(row=i + 1, column=j, padx=2, pady=2)
            self.bind_cell(e); row_e.append(e)
        self.entries.append(row_e); self.rows += 1
        self.inner.update_idletasks(); self.canvas.config(scrollregion=self.canvas.bbox("all"))

    def delete_row(self):
        if not self.entries: return
        for e in self.entries.pop(): e.destroy()
        self.rows -= 1; self.inner.update_idletasks()
        self.canvas.config(scrollregion=self.canvas.bbox("all"))

    def add_column(self):
        self.cols += 1; ci = self.cols - 1
        nm = f"Повт.{ci - self.factors_count + 1}"
        lbl = tk.Label(self.inner, text=nm, relief=tk.RIDGE, width=COL_W,
                       bg="#f0f0f0", fg="#000000", font=("Times New Roman", 12, "bold"))
        lbl.grid(row=0, column=ci, padx=2, pady=2, sticky="nsew"); self.header_labels.append(lbl)
        for i, row in enumerate(self.entries):
            e = self._mk_entry(self.inner); e.grid(row=i + 1, column=ci, padx=2, pady=2)
            self.bind_cell(e); row.append(e)
        self.inner.update_idletasks(); self.canvas.config(scrollregion=self.canvas.bbox("all"))

    def delete_column(self):
        if self.cols <= self.factors_count + 1: return
        self.header_labels.pop().destroy()
        for row in self.entries: row.pop().destroy()
        self.cols -= 1; self.inner.update_idletasks()
        self.canvas.config(scrollregion=self.canvas.bbox("all"))

    def _paste_from_focus(self):
        w = self.table_win.focus_get()
        if isinstance(w, tk.Entry):
            class _E: widget = w
            self._on_paste(_E())

    def _used_rep(self):
        rep_cols = []
        for c in range(self.factors_count, self.cols):
            for r in range(len(self.entries)):
                s = self.entries[r][c].get().strip()
                if not s: continue
                try: float(s.replace(",", ".")); rep_cols.append(c); break
                except Exception: continue
        return rep_cols

    def collect_long(self, design):
        long = []; rep = self._used_rep()
        if not rep: return long, rep
        for i, row in enumerate(self.entries):
            lvls = [row[k].get().strip() or f"рядок{i+1}" for k in range(self.factors_count)]
            for ic, c in enumerate(rep):
                s = row[c].get().strip()
                if not s: continue
                try: val = float(s.replace(",", "."))
                except Exception: continue
                rec = {"value": val}
                for ki, fk in enumerate(self.factor_keys): rec[fk] = lvls[ki]
                if design in ("rcbd", "split"):
                    rec["BLOCK"] = f"Блок {ic+1}"
                elif design == "latin":
                    # Для латинського квадрату: перший фактор = варіант,
                    # другий фактор = рядок (ROW), третій = стовпець (COL)
                    # якщо factors_count >= 3 — беремо з таблиці
                    if self.factors_count >= 3:
                        rec["ROW"] = lvls[1] if len(lvls) > 1 else f"Ряд {i+1}"
                        rec["COL"] = lvls[2] if len(lvls) > 2 else f"Стовп {ic+1}"
                    else:
                        # factors_count == 1: автогенерація ROW/COL з номерів
                        rec["ROW"] = f"Ряд {i+1}"
                        rec["COL"] = f"Стовп {ic+1}"
                long.append(rec)
        return long, rep

    # ── dialogs ───────────────────────────────────────────────
    def ask_params(self):
        parent = self.table_win or self.root
        dlg = tk.Toplevel(parent); dlg.title("Параметри аналізу")
        dlg.resizable(False, False); set_icon(dlg)
        frm = tk.Frame(dlg, padx=16, pady=14); frm.pack()
        rf = ("Times New Roman", 12)
        rb_f = ("Times New Roman", 13)

        tk.Label(frm, text="Назва показника (необов'язково):",
                 font=rf).grid(row=0, column=0, sticky="w", pady=4)
        e_ind = tk.Entry(frm, width=32, font=rf)
        e_ind.grid(row=0, column=1, pady=4, padx=6)
        tk.Label(frm, text="Одиниці виміру (необов'язково):",
                 font=rf).grid(row=1, column=0, sticky="w", pady=4)
        e_un = tk.Entry(frm, width=32, font=rf)
        e_un.grid(row=1, column=1, pady=4, padx=6)

        # ── Дизайн ───────────────────────────────────────────
        tk.Label(frm, text="Дизайн досліду:", font=("Times New Roman",12,"bold")
                 ).grid(row=2, column=0, columnspan=2, sticky="w", pady=14)

        design_info = tk.Frame(frm); design_info.grid(row=3, column=0, columnspan=2, sticky="w")
        tk.Label(design_info,
                 text=(
                     "CRD — Повністю рандомізований:\n"
                     "  Варіанти розміщені випадково по всіх ділянках.\n"
                     "  Використовується на однорідному фоні.\n\n"
                     "RCBD — Рандомізовані повні блоки:\n"
                     "  Ділянки розбиті на блоки (повторності) однорідних умов.\n"
                     "  Блок = одна повторність. Всередині блоку — випадкове розміщення.\n"
                     "  Рекомендується при неоднорідності фону.\n\n"
                     "Split-plot — Розщеплені ділянки:\n"
                     "  Головний фактор (whole-plot) = великі ділянки,\n"
                     "  Другорядний фактор (sub-plot) = дрібніші, всередині великих.\n"
                     "  Типово: фактор A = обробка всього поля, B = сорт на підділянці."
                 ),
                 font=("Times New Roman",10), justify="left",
                 bg="#f0f4ff", relief=tk.FLAT, padx=8, pady=6
                 ).pack(fill=tk.X)

        dv = tk.StringVar(value="crd")
        df = tk.Frame(frm); df.grid(row=4, column=0, columnspan=2, sticky="w", pady=8)
        for txt, val in [("CRD", "crd"), ("RCBD", "rcbd"),
                          ("Латинський квадрат", "latin"),
                          ("Split-plot (лише параметричний)", "split")]:
            tk.Radiobutton(df, text=txt, variable=dv, value=val,
                           font=rb_f).pack(side=tk.LEFT, padx=8)

        # Підказка для латинського квадрату
        latin_hint = tk.Label(frm,
            text="ℹ ЛАТИНСЬКИЙ КВАДРАТ — інструкція:\n"
                 "  1. Відкрийте 3-ФАКТОРНИЙ аналіз (кнопка «3 фактори» на головній).\n"
                 "  2. Перейменуйте фактори (подвійний клік):\n"
                 "     Фактор A = Варіант  (назви обробок або сортів)\n"
                 "     Фактор B = Рядок    (1, 2, 3... — ряди поля)\n"
                 "     Фактор C = Стовпець (1, 2, 3... — стовпці поля)\n"
                 "  3. Кожен рядок таблиці = одна ділянка (k²  рядків).\n"
                 "     Значення вводьте у колонку «Повт.1».\n"
                 "  4. Вимога: k варіантів = k рядів = k стовпців (k ≥ 3).\n"
                 "  Модель: Y = μ + Варіант + Рядок + Стовпець + ε\n"
                 "  df(похибка) = (k-1)(k-2)",
            font=("Times New Roman",10), fg="#1a4b8c",
            bg="#eef4ff", relief=tk.FLAT, padx=8, pady=6, justify="left")
        latin_hint.grid(row=5, column=0, columnspan=2, sticky="ew", pady=(0,4))
        latin_hint.grid_remove()

        # Тепер реєструємо trace (latin_hint вже існує)
        def _on_design(*_):
            v = dv.get()
            if v == "latin": latin_hint.grid()
            else:            latin_hint.grid_remove()
            if v == "split": sp_frm.grid()
            else:            sp_frm.grid_remove()
        dv.trace_add("write", _on_design)

        mfv = tk.StringVar(value=self.factor_keys[0] if self.factor_keys else "A")
        sp_frm = tk.Frame(frm); sp_frm.grid(row=5, column=0, columnspan=2, sticky="w", pady=(0,4))
        tk.Label(sp_frm, text="Головний фактор (whole-plot):", font=rf).pack(side=tk.LEFT)
        ttk.Combobox(sp_frm, textvariable=mfv, width=6, state="readonly",
                     values=self.factor_keys).pack(side=tk.LEFT, padx=6)
        sp_frm.grid_remove()


        # ── Тип SS ───────────────────────────────────────────
        ttk.Separator(frm, orient="horizontal").grid(
            row=6, column=0, columnspan=2, sticky="ew", pady=8)
        tk.Label(frm, text="Тип суми квадратів (SS):", font=("Times New Roman",12,"bold")
                 ).grid(row=7, column=0, columnspan=2, sticky="w", pady=0)
        tk.Label(frm,
                 text=(
                     "Тип I — Послідовний: кожен фактор після попередніх.\n"
                     "  Порядок факторів важливий. Тільки для збалансованих дизайнів.\n\n"
                     "Тип II — Ієрархічний: кожен фактор після решти головних ефектів\n"
                     "  (без взаємодій). Для незбалансованих даних без взаємодій.\n\n"
                     "Тип III — Частковий ← РЕКОМЕНДУЄТЬСЯ: кожен ефект при всіх інших.\n"
                     "  Стандарт SPSS/SAS. Не залежить від порядку. Взаємодії враховані.\n\n"
                     "Тип IV — Для сильно незбалансованих дизайнів з порожніми клітинками."
                 ),
                 font=("Times New Roman",10), justify="left",
                 bg="#f0f4ff", relief=tk.FLAT, padx=8, pady=4
                 ).grid(row=8, column=0, columnspan=2, sticky="ew")
        ssv = tk.StringVar(value="III")
        ssf = tk.Frame(frm); ssf.grid(row=9, column=0, columnspan=2, sticky="w", pady=4)
        for ss in ["I","II","III","IV"]:
            tk.Radiobutton(ssf, text=f"Тип {ss}", variable=ssv, value=ss,
                           font=rb_f).pack(side=tk.LEFT, padx=8)

        # ── Кнопки ───────────────────────────────────────────
        out = {"ok": False}
        def ok():
            out.update({"ok": True,
                        "indicator": e_ind.get().strip() or "Показник",
                        "units":     e_un.get().strip()  or "–",
                        "design":    dv.get(),
                        "split_main": mfv.get(),
                        "ss_type":   ssv.get()})
            dlg.destroy()

        bf = tk.Frame(frm); bf.grid(row=10, column=0, columnspan=2, pady=(14,0))
        tk.Button(bf, text="▶ Виконати аналіз", width=18,
                  bg="#c62828", fg="white", font=rf, command=ok).pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Скасувати", width=12,
                  font=rf, command=dlg.destroy).pack(side=tk.LEFT)

        dlg.update_idletasks(); center_win(dlg)
        e_ind.focus_set()
        dlg.bind("<Return>", lambda e: ok())
        dlg.grab_set(); parent.wait_window(dlg)
        return out

    def choose_method(self, p_norm, design, n_var):
        parent = self.table_win or self.root
        dlg = tk.Toplevel(parent); dlg.title("Вибір методу")
        dlg.resizable(False, False); set_icon(dlg)
        frm = tk.Frame(dlg, padx=16, pady=14); frm.pack()
        normal = (p_norm is not None) and (not math.isnan(p_norm)) and (p_norm > 0.05)
        rb_f = ("Times New Roman", 13)
        ordinal = getattr(self, '_ordinal_mode', False)

        if ordinal:
            # ── Режим бальних даних: ЛИШЕ непараметричні ────
            tk.Label(frm,
                     text="⚠ Бальна шкала → ЛИШЕ непараметричні методи",
                     fg="#c62828", font=("Times New Roman",12,"bold")
                     ).pack(anchor="w", pady=0)
            tk.Label(frm,
                     text=(
                         "Параметричні методи (НІР, Тьюкі, Дункан) заблоковані.\n"
                         "Причина: бальні дані є порядковими — середнє і\n"
                         "стандартне відхилення методично некоректні для них.\n"
                         "Оберіть непараметричний критерій:"
                     ),
                     fg="#555", font=("Times New Roman",11), justify="left"
                     ).pack(anchor="w", pady=(0,10))
            if design == "crd":
                options = [
                    ("Краскел-Уолліс + Манн-Уітні (Бонферроні)  ← рекомендовано", "kw"),
                ]
            elif n_var == 2:
                options = [("Вілкоксон (парний)  ← рекомендовано для 2 варіантів", "wilcoxon")]
            else:
                options = [("Фрідман + Вілкоксон (Бонферроні)  ← рекомендовано", "friedman")]

        elif normal:
            tk.Label(frm, text="✓ Дані відповідають нормальному розподілу (Shapiro–Wilk).",
                     justify="left", fg="#1a6b1a",
                     font=("Times New Roman",11)).pack(anchor="w", pady=0)
            options = [("НІР₀₅ (LSD)", "lsd"),
                       ("Тест Тьюкі", "tukey"),
                       ("Тест Дункана", "duncan"),
                       ("Бонферроні", "bonferroni"),
                       ("🔁 arcsin(√p) + параметричний", "arcsin_param")]
        else:
            if design == "split":
                tk.Label(frm,
                         text="Split-plot: лише параметричний.\n"
                              "Залишки не нормальні → аналіз некоректний.\n"
                              "Рекомендація: трансформуйте або оберіть CRD/RCBD.",
                         fg="#c62828", justify="left",
                         font=("Times New Roman",11)).pack(anchor="w")
                options = []
            else:
                tk.Label(frm,
                         text="⚠ Дані НЕ відповідають нормальному розподілу.\n"
                              "Оберіть метод:",
                         fg="#c62828", justify="left",
                         font=("Times New Roman",11)).pack(anchor="w", pady=(0, 8))
                if design == "crd":
                    options = [
                        ("Краскела–Уолліса", "kw"),
                        ("Манна-Уітні", "mw"),
                        ("🔁 arcsin(√p) + параметричний", "arcsin_param"),
                        ("🔁 ln(x) + параметричний",      "log_param"),
                        ("🔁 √x + параметричний",          "sqrt_param"),
                        ("🔁 log₁₀(x) + параметричний",   "log10_param"),
                    ]
                else:
                    if n_var == 2:
                        options = [
                            ("Wilcoxon (парний)", "wilcoxon"),
                            ("🔁 arcsin(√p) + параметричний", "arcsin_param"),
                            ("🔁 ln(x) + параметричний",      "log_param"),
                            ("🔁 √x + параметричний",          "sqrt_param"),
                            ("🔁 log₁₀(x) + параметричний",   "log10_param"),
                        ]
                    else:
                        options = [
                            ("Friedman", "friedman"),
                            ("🔁 arcsin(√p) + параметричний", "arcsin_param"),
                            ("🔁 ln(x) + параметричний",      "log_param"),
                            ("🔁 √x + параметричний",          "sqrt_param"),
                            ("🔁 log₁₀(x) + параметричний",   "log10_param"),
                        ]
        out = {"ok": False, "method": None}
        if not options:
            tk.Button(frm, text="OK", width=10, command=dlg.destroy).pack(pady=(10, 0))
            dlg.update_idletasks(); center_win(dlg); dlg.grab_set(); parent.wait_window(dlg); return out
        var = tk.StringVar(value=options[0][1])
        for txt, val in options:
            tk.Radiobutton(frm, text=txt, variable=var, value=val, font=rb_f).pack(anchor="w", pady=2)
        def ok():
            out.update({"ok": True, "method": var.get()}); dlg.destroy()
        bf = tk.Frame(frm); bf.pack(pady=(12, 0))
        tk.Button(bf, text="▶ Виконати", width=14, bg="#c62828", fg="white",
                  font=("Times New Roman",12), command=ok).pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Скасувати", width=12,
                  font=("Times New Roman",12), command=dlg.destroy).pack(side=tk.LEFT)
        dlg.update_idletasks(); center_win(dlg); dlg.bind("<Return>", lambda e: ok())
        dlg.grab_set(); parent.wait_window(dlg); return out

    # ════════════════════════════════════════════════════════════
    # ANALYZE
    # ════════════════════════════════════════════════════════════
    def analyze(self):
        try:
            self._analyze_inner()
        except Exception as _ae:
            import traceback
            messagebox.showerror("Помилка аналізу",
                f"Виникла помилка при виконанні аналізу:\n\n"
                f"{str(_ae)}\n\n"
                f"Деталі:\n{traceback.format_exc()[-600:]}")

    def _analyze_inner(self):
        created = datetime.now()
        params = self.ask_params()
        if not params["ok"]: return
        indicator = params["indicator"]; units = params["units"]
        design = params["design"]; split_main = params["split_main"]
        ss_type = params.get("ss_type", "III")

        long, used_rep = self.collect_long(design)
        if not long: messagebox.showwarning("Помилка даних", "Немає числових даних."); return
        values = np.array([r["value"] for r in long], dtype=float)

        # ── Автоматичне виявлення відсоткових даних (arcsin) ───
        is_percent_units = any(u in units.lower() for u in
                               ["%", "відсот", "percent", "частк", "зав'яз",
                                "схожість", "схожіст", "ураж", "виживан",
                                "товарн", "вихід"])
        arcsin_applied = False
        arcsin_was_suggested = False

        if is_percent_units:
            vmin = float(np.min(values)); vmax = float(np.max(values))
            # Перевіряємо чи дані у діапазоні 0-100 (відсотки) або 0-1 (частки)
            in_pct_range  = (vmin >= 0 and vmax <= 100 and vmax > 1)
            in_frac_range = (vmin >= 0 and vmax <= 1)
            is_fraction   = in_pct_range or in_frac_range
            needs_arcsin  = is_fraction and (vmin < 20 or vmax > 80)

            if is_fraction and not needs_arcsin:
                # Всі значення між 20% і 80% — підказка але не наполягаємо
                messagebox.showinfo("Інформація про дані",
                    "Одиниці вимірювання: " + units + "\n"
                    "Діапазон значень: " + fmt(vmin,1) + "% – " + fmt(vmax,1) + "%\n\n"
                    "Всі значення у діапазоні 20–80% \u2192 arcsin трансформація\n"
                    "не є критично необхідною, але дозволена.\n\n"
                    "Аналіз буде виконано без трансформації.\n"
                    "Якщо Shapiro–Wilk покаже ненормальність \u2014\n"
                    "оберіть arcsin(\u221ap) у вікні вибору методу.")

            elif is_fraction and needs_arcsin:
                arcsin_was_suggested = True
                ans = messagebox.askyesno("Виявлено відсоткові дані",
                    "Одиниці вимірювання: " + units + "\n"
                    "Діапазон значень: " + fmt(vmin,1) + "% – " + fmt(vmax,1) + "%\n\n"
                    "Значення поза діапазоном 20–80%\n"
                    "виявляють нерівномірність дисперсій.\n\n"
                    "РЕКОМЕНДОВАНА трансформація: arcsin(\u221ap)\n"
                    "Формула: y = arcsin(\u221a(p/100))\n\n"
                    "Що робить:\n"
                    "\u2022 Вирівнює дисперсії між варіантами\n"
                    "\u2022 Наближає розподіл до нормального\n"
                    "\u2022 Дозволяє коректно застосувати ANOVA\n\n"
                    "Середні у звіті будуть у вихідних %.\n\n"
                    "Застосувати arcsin(\u221ap) трансформацію?")

                if ans:
                    # Застосовуємо arcsin трансформацію
                    def _arcsin_transform(v):
                        p = v / 100.0 if v > 1 else v  # конвертуємо у частки
                        p = max(0.0, min(1.0, p))       # обмежуємо [0,1]
                        return math.asin(math.sqrt(p))

                    long = [dict(r, value=_arcsin_transform(r["value"])) for r in long]
                    values = np.array([r["value"] for r in long], dtype=float)
                    arcsin_applied = True

                    messagebox.showinfo("Трансформацію застосовано",
                        "arcsin(\u221ap) трансформацію застосовано.\n"
                        "Аналіз виконується на трансформованих даних.\n"
                        "У звіті середні наведено у вихідних відсотках.\n"
                        "Літери значущості (НІР) — за трансформованими даними.")

        # ── Автоматичне виявлення бальних (порядкових) даних ──
        is_ordinal_units = any(u in units.lower() for u in
                               ["бал", "score", "rank", "rang", "ранг",
                                "очко", "пункт", "клас", "ступін", "ступен"])
        ordinal_detected = False
        ordinal_forced   = False   # чи примусово переведено на непараметричний

        if is_ordinal_units and not arcsin_applied:
            vmin_o = float(np.min(values)); vmax_o = float(np.max(values))
            # Додаткова перевірка: всі значення цілі і діапазон ≤ 20
            all_int = np.all(values == np.floor(values))
            small_range = (vmax_o - vmin_o) <= 20

            if all_int and small_range:
                ordinal_detected = True
                messagebox.showwarning(
                    "⚠ Виявлено бальну (порядкову) шкалу",
                    f"Одиниці вимірювання: «{units}»\n"
                    f"Діапазон значень: {int(vmin_o)} – {int(vmax_o)} балів\n\n"
                    "МЕТОДИЧНА ВИМОГА:\n"
                    "Бальні шкали є ПОРЯДКОВИМИ (ordinal) даними.\n"
                    "Це означає що:\n"
                    "  • Різниця між балами нерівномірна\n"
                    "  • Розподіл зазвичай ненормальний\n"
                    "  • Параметрична ANOVA методично НЕКОРЕКТНА\n\n"
                    "Програма автоматично застосує:\n"
                    "  ✓ Непараметричний аналіз (Краскел-Уолліс / Фрідман)\n"
                    "  ✓ Медіана [Q1; Q3] замість Mean ± SD у звіті\n"
                    "  ✓ Boxplot або Dot plot для візуалізації\n\n"
                    "Параметричні методи заблоковано для цих даних.\n"
                    "Натисніть OK щоб продовжити з правильним методом.")
                ordinal_forced = True

        # Зберігаємо прапор для choose_method і show_report
        self._ordinal_mode  = ordinal_detected and ordinal_forced
        self._ordinal_units = units

        # ── Мінімальна кількість спостережень ──────────────────
        if len(values) < 6:
            messagebox.showwarning("Замало даних",
                f"Для дисперсійного аналізу потрібно щонайменше 6 спостережень.\n"
                f"Наразі: {len(values)}."); return

        lbf = {f: first_seen([r.get(f) for r in long]) for f in self.factor_keys}

        # ── Перевірка мінімальної кількості рівнів кожного фактора ──
        for f in self.factor_keys:
            if len(lbf[f]) < 2:
                messagebox.showwarning("Помилка даних",
                    f"Фактор «{self.ftitle(f)}» має лише 1 рівень.\n"
                    "Для аналізу потрібно щонайменше 2 рівні."); return

        # ── Перевірка мінімальної кількості повторностей для RCBD/Split ──
        if design in ("rcbd", "split") and len(used_rep) < 2:
            messagebox.showwarning("Помилка дизайну",
                f"Для дизайну {design.upper()} потрібно щонайменше 2 повторності (блоки).\n"
                f"Наразі: {len(used_rep)}."); return

        # ── Перевірка Латинського квадрату ─────────────────────
        if design == "latin":
            if self.factors_count < 3:
                messagebox.showerror("Латинський квадрат — помилка структури",
                    "Для Латинського квадрату відкрийте 3-факторний аналіз:\n\n"
                    "  Фактор A = Варіант (назви варіантів)\n"
                    "  Фактор B = Рядок (номери рядів: 1, 2, 3...)\n"
                    "  Фактор C = Стовпець (номери стовпців: 1, 2, 3...)\n\n"
                    "Кожен рядок таблиці = одна ділянка, одне значення у колонці «Повт.1»"); return
            # Перевіряємо k = n_рядків = n_стовпців
            fk = self.factor_keys
            k_var  = len(lbf.get(fk[0], []))
            k_rows = len(lbf.get(fk[1], [])) if len(fk) > 1 else 0
            k_cols = len(lbf.get(fk[2], [])) if len(fk) > 2 else 0
            if not (k_var == k_rows == k_cols):
                messagebox.showerror("Латинський квадрат — помилка k×k",
                    f"Латинський квадрат вимагає k варіантів = k рядків = k стовпців.\n\n"
                    f"Знайдено:\n"
                    f"  Варіантів (Фактор A): {k_var}\n"
                    f"  Рядків (Фактор B):    {k_rows}\n"
                    f"  Стовпців (Фактор C):  {k_cols}\n\n"
                    f"Перевірте що всі три значення однакові."); return
            if k_var < 3:
                messagebox.showwarning("Замало варіантів",
                    f"Латинський квадрат вимагає щонайменше 3 варіанти (k≥3).\n"
                    f"Знайдено: {k_var}. Рекомендується k=4-6."); return
            if k_var > 8:
                ans = messagebox.askyesno("Попередження: великий квадрат",
                    f"k={k_var} варіантів → {k_var}×{k_var} = {k_var**2} ділянок.\n"
                    f"При k>8 Латинський квадрат стає громіздким.\n"
                    f"Рекомендується RCBD для великої кількості варіантів.\n\n"
                    "Продовжити?")
                if not ans: return

        var_order = first_seen([tuple(r.get(f) for f in self.factor_keys) for r in long])
        v_names = [" | ".join(map(str, k)) for k in var_order]
        n_var = len(var_order)

        # ── Перевірка збалансованості при Тип I SS ─────────────
        if ss_type == "I":
            from collections import Counter
            cell_counts = Counter(tuple(r.get(f) for f in self.factor_keys) for r in long)
            counts = list(cell_counts.values())
            if len(set(counts)) > 1:
                ans = messagebox.askyesno("Увага: незбалансовані дані + Тип I SS",
                    "Дані незбалансовані (різна кількість спостережень у клітинках).\n\n"
                    "При Тип I SS (послідовний) результат залежить від ПОРЯДКУ введення факторів.\n"
                    "Зміна порядку дає інші значення SS та p.\n\n"
                    "Рекомендація: використовуйте Тип III для незбалансованих даних.\n\n"
                    "Продовжити з Тип I SS?")
                if not ans: return

        try:
            if design == "crd":
                res = anova_crd(long, self.factor_keys, lbf, ss_type)
            elif design == "rcbd":
                res = anova_rcbd(long, self.factor_keys, lbf, ss_type=ss_type)
            elif design == "latin":
                # Для ЛК перший ключ = варіант, другий = рядок, третій = стовпець
                fk_var = [self.factor_keys[0]]
                res = anova_latin_square(long, fk_var, lbf, ss_type)
            else:
                if split_main not in self.factor_keys: split_main = self.factor_keys[0]
                res = anova_split(long, self.factor_keys, split_main, ss_type=ss_type)
        except Exception as ex:
            import traceback
            messagebox.showerror("Помилка моделі",
                str(ex) + "\n\nДетальніше:\n" + traceback.format_exc()[-500:])
            return

        residuals = np.array(res.get("residuals", []), dtype=float)
        n_res = len(residuals)
        self._last_residuals = residuals.tolist()  # для вкладки залишків
        try: W, p_norm = shapiro(residuals) if n_res >= 3 else (np.nan, np.nan)
        except Exception: W, p_norm = np.nan, np.nan

        # ── Попередження про обмеження Shapiro-Wilk ────────────
        sw_warning = ""
        if n_res < 8:
            sw_warning = f"\n⚠ Увага: n={n_res} — надто мало для надійного тесту нормальності."
        elif n_res > 100:
            sw_warning = (f"\n⚠ Увага: n={n_res} — при великих вибірках Shapiro–Wilk виявляє\n"
                          "  навіть мінімальні відхилення як значущі. Оцінюйте разом з Q-Q графіком.")

        normal = (not math.isnan(p_norm)) and (p_norm > 0.05)
        if design == "split" and not normal:
            messagebox.showwarning("Split-plot: аналіз неможливий",
                f"Залишки моделі не відповідають нормальному розподілу\n"
                f"(Shapiro–Wilk: W={fmt(W,4)}, p={fmt(p_norm,4)}).\n\n"
                "Split-plot реалізований лише для параметричних методів.\n"
                "Рекомендації:\n"
                "• трансформуйте дані (логарифмування) і повторіть;\n"
                "• або оберіть CRD/RCBD для непараметричного аналізу."
                + sw_warning); return

        choice = self.choose_method(p_norm, design, n_var)
        if not choice or not choice.get("ok"): return
        method = choice["method"]

        log_applied = False
        transform_label = ""
        if method == "arcsin_param" and not arcsin_applied:
            # Користувач обрав arcsin вручну у choose_method
            # (не було автоматичного застосування)
            vmin_ = float(np.min(values)); vmax_ = float(np.max(values))
            if vmin_ < 0 or vmax_ > 100:
                messagebox.showwarning("Помилка трансформації",
                    f"arcsin(√p) застосовна лише для даних у діапазоні 0–100% або 0–1.\n"
                    f"Ваші дані: мін={fmt(vmin_,2)}, макс={fmt(vmax_,2)}.\n\n"
                    "Оберіть інший метод трансформації."); return
            def _asin(v):
                p = v/100.0 if v > 1 else v
                return math.asin(math.sqrt(max(0., min(1., p))))
            long = [dict(r, value=_asin(r["value"])) for r in long]
            values = np.array([r["value"] for r in long], dtype=float)
            arcsin_applied = True
            transform_label = "arcsin(√p)"
            log_applied = True
            # Перераховуємо модель на трансформованих даних
            try:
                if design == "crd":      res = anova_crd(long, self.factor_keys, lbf, ss_type)
                elif design == "rcbd":   res = anova_rcbd(long, self.factor_keys, lbf, ss_type=ss_type)
                elif design == "latin":  res = anova_latin_square(long, self.factor_keys, lbf, ss_type)
                else:                    res = anova_split(long, self.factor_keys, split_main, ss_type=ss_type)
            except Exception as ex: messagebox.showerror("Помилка моделі", str(ex)); return
            residuals = np.array(res.get("residuals",[]), dtype=float)
            try: W, p_norm = shapiro(residuals) if len(residuals)>=3 else (np.nan,np.nan)
            except Exception: W, p_norm = np.nan, np.nan
            # Вибираємо пост-хок метод після arcsin
            messagebox.showinfo("arcsin(√p) застосовано",
                f"Трансформацію arcsin(√p) застосовано.\n"
                f"Shapiro–Wilk після трансформації: W={fmt(W,4)}, p={fmt(p_norm,4)}\n"
                f"{'✓ розподіл нормальний' if not math.isnan(p_norm) and p_norm>0.05 else '⚠ розподіл все ще ненормальний — розгляньте непараметричний метод'}\n\n"
                "Оберіть метод пост-хок порівнянь:")
            choice2 = self.choose_method(p_norm, design, n_var)
            if not choice2["ok"]: return
            method = choice2["method"]
            if method == "arcsin_param": method = "lsd"

        elif method in ("log_param", "sqrt_param", "log10_param"):
            # ── Validate data for chosen transformation ──────────
            if method in ("log_param", "log10_param") and np.any(values <= 0):
                messagebox.showwarning("Трансформація неможлива",
                    "Дані містять нулі або від'ємні значення.\n"
                    "Логарифмування неможливе.\n"
                    "Оберіть √x або непараметричний метод."); return
            if method == "sqrt_param" and np.any(values < 0):
                messagebox.showwarning("Трансформація неможлива",
                    "Дані містять від'ємні значення.\n"
                    "Трансформація √x неможлива."); return

            # ── Apply transformation ─────────────────────────────
            if method == "log_param":
                long = [dict(r, value=math.log(r["value"])) for r in long]
                transform_label = "ln(x)"
            elif method == "sqrt_param":
                long = [dict(r, value=math.sqrt(r["value"])) for r in long]
                transform_label = "√x"
            elif method == "log10_param":
                long = [dict(r, value=math.log10(r["value"])) for r in long]
                transform_label = "log₁₀(x)"

            values = np.array([r["value"] for r in long], dtype=float)
            log_applied = True

            # ── Re-run model on transformed data ─────────────────
            try:
                if design == "crd":      res = anova_crd(long, self.factor_keys, lbf, ss_type)
                elif design == "rcbd":   res = anova_rcbd(long, self.factor_keys, lbf, ss_type=ss_type)
                elif design == "latin":  res = anova_latin_square(long, self.factor_keys, lbf, ss_type)
                else:                    res = anova_split(long, self.factor_keys, split_main, ss_type=ss_type)
            except Exception as ex: messagebox.showerror("Помилка моделі", str(ex)); return

            residuals = np.array(res.get("residuals", []), dtype=float)
            try: W, p_norm = shapiro(residuals) if len(residuals) >= 3 else (np.nan, np.nan)
            except Exception: W, p_norm = np.nan, np.nan

            if math.isnan(p_norm) or p_norm <= 0.05:
                messagebox.showwarning("Трансформація не допомогла",
                    f"Застосовано трансформацію {transform_label}, але залишки\n"
                    f"все одно не відповідають нормальному розподілу\n"
                    f"(Shapiro–Wilk: W={fmt(W,4)}, p={fmt(p_norm,4)}).\n\n"
                    "Параметричний аналіз неможливий.\n"
                    "Оберіть непараметричний метод (Kruskal–Wallis, Mann–Whitney тощо)."); return

            messagebox.showinfo("Трансформація успішна",
                f"Застосовано трансформацію {transform_label}.\n"
                f"Shapiro–Wilk після трансформації:\n"
                f"W={fmt(W,4)},  p={fmt(p_norm,4)}  ✓ нормальний розподіл.\n\n"
                "Оберіть метод парних порівнянь:")
            choice2 = self.choose_method(p_norm, design, n_var)
            if not choice2["ok"]: return
            method = choice2["method"]
            if method in ("log_param", "sqrt_param", "log10_param"): method = "lsd"

        MS_err = res.get("MS_error", np.nan); df_err = res.get("df_error", np.nan)
        MS_wp  = res.get("MS_whole", np.nan); df_wp  = res.get("df_whole", np.nan)
        split_mf = res.get("main_factor", split_main) if design == "split" else None

        vs_ = vstats(long, self.factor_keys)
        v_means = {k: vs_[k][0] for k in vs_}; v_sds = {k: vs_[k][1] for k in vs_}; v_ns = {k: vs_[k][2] for k in vs_}
        means1 = {v_names[i]: v_means.get(var_order[i], np.nan) for i in range(n_var)}
        ns1    = {v_names[i]: v_ns.get(var_order[i], 0)         for i in range(n_var)}
        gv = groups_by(long, tuple(self.factor_keys))
        groups1 = {v_names[i]: gv.get(var_order[i], []) for i in range(n_var)}

        fg = {f: {k[0]: v for k, v in groups_by(long, (f,)).items()} for f in self.factor_keys}
        fm = {f: {lv: float(np.mean(arr)) if arr else np.nan for lv, arr in fg[f].items()} for f in self.factor_keys}
        fn = {f: {lv: len(arr) for lv, arr in fg[f].items()} for f in self.factor_keys}
        fsd= {f: {lv: float(np.std(arr, ddof=1)) if len(arr) >= 2 else (0. if len(arr)==1 else np.nan)
                  for lv, arr in fg[f].items()} for f in self.factor_keys}
        fmed = {f: {lv: median_q(arr)[0] for lv, arr in fg[f].items()} for f in self.factor_keys}
        fq   = {f: {lv: median_q(arr)[1:] for lv, arr in fg[f].items()} for f in self.factor_keys}

        vmed = {var_order[i]: median_q(groups1[v_names[i]])[0] for i in range(n_var)}
        vq   = {var_order[i]: median_q(groups1[v_names[i]])[1:] for i in range(n_var)}
        rkv  = mean_ranks(long, lambda r: " | ".join(str(r.get(f)) for f in self.factor_keys))
        rkf  = {f: mean_ranks(long, lambda r, ff=f: r.get(ff)) for f in self.factor_keys}

        lev_F, lev_p = (np.nan, np.nan)
        if method in ("lsd", "tukey", "duncan", "bonferroni"):
            lev_F, lev_p = levene_test(groups1)
            # ── Блокування при неоднорідних дисперсіях ─────────
            if not math.isnan(lev_p) and lev_p < ALPHA:
                ans = messagebox.askyesno(
                    "Неоднорідність дисперсій (тест Левена)",
                    f"Тест Левена виявив неоднорідність дисперсій\n"
                    f"(F={fmt(lev_F,4)}, p={fmt(lev_p,4)}) — умова ANOVA порушена.\n\n"
                    "Параметричний аналіз при неоднорідних дисперсіях дає\n"
                    "недостовірні F-значення та p-значення.\n\n"
                    "Рекомендації:\n"
                    "• застосуйте трансформацію даних (логарифмування);\n"
                    "• або оберіть непараметричний метод (Kruskal–Wallis).\n\n"
                    "Продовжити параметричний аналіз попри порушення умови?")
                if not ans: return

        kw_H = kw_p = kw_df = kw_eps = np.nan; do_ph = True
        fr_chi = fr_p = fr_df = fr_W = np.nan; wil_s = wil_p = np.nan
        rcbd_ph = []; rcbd_sig = {}
        lf = {f: {lv: "" for lv in lbf[f]} for f in self.factor_keys}
        lnamed = {nm: "" for nm in v_names}
        ph_rows = []; fpt = {}

        if method == "lsd":
            # ── Protected LSD (Fisher): перевірити значущість глобального F ──
            # Знаходимо найменше p-значення серед головних ефектів та взаємодій
            global_p_values = []
            for raw_row in res["table"]:
                nm_, _ss, _df, _ms, _F, pv_ = raw_row
                if any(x in str(nm_) for x in ["Залишок", "Загальна", "Блоки", "WP-error"]):
                    continue
                if pv_ is not None and not (isinstance(pv_, float) and math.isnan(pv_)):
                    global_p_values.append(float(pv_))

            global_F_sig = any(p < ALPHA for p in global_p_values) if global_p_values else False

            if not global_F_sig:
                messagebox.showwarning("Protected LSD: пост-хок заблоковано",
                    "Жоден з ефектів у дисперсійному аналізі не є статистично значущим\n"
                    "(p ≥ 0.05 для всіх факторів та їх взаємодій).\n\n"
                    "Відповідно до принципу Protected LSD (Fisher, 1935),\n"
                    "попарні порівняння можна виконувати ЛИШЕ після значущого F-тесту.\n\n"
                    "Висновок: між варіантами немає статистично значущої різниці.\n"
                    "Звіт сформовано з таблицями описової статистики без пост-хок аналізу.")
                # Продовжуємо — формуємо звіт без літер CLD
            else:
                for f in self.factor_keys:
                    MS_ = MS_wp if (design == "split" and f == split_mf) else MS_err
                    df_ = df_wp if (design == "split" and f == split_mf) else df_err
                    lf[f] = cld(lbf[f], fm[f], lsd_sig(lbf[f], fm[f], fn[f], MS_, df_))
                if design != "split":
                    lnamed = cld(v_names, means1, lsd_sig(v_names, means1, ns1, MS_err, df_err))

        elif method in ("tukey", "duncan", "bonferroni"):
            # ── Перевірка значущості глобального F перед пост-хок ──
            global_p_values_2 = []
            for raw_row in res["table"]:
                nm_, _ss, _df, _ms, _F, pv_ = raw_row
                if any(x in str(nm_) for x in ["Залишок", "Загальна", "Блоки", "WP-error"]):
                    continue
                if pv_ is not None and not (isinstance(pv_, float) and math.isnan(pv_)):
                    global_p_values_2.append(float(pv_))
            global_F_sig_2 = any(p < ALPHA for p in global_p_values_2) if global_p_values_2 else False

            if not global_F_sig_2:
                messagebox.showwarning("Пост-хок заблоковано",
                    "Жоден ефект не є статистично значущим (p ≥ 0.05).\n\n"
                    "Виконання пост-хок порівнянь без значущого F-тесту\n"
                    "призводить до надмірного числа хибнопозитивних результатів.\n\n"
                    "Висновок: між варіантами немає статистично значущої різниці.")
            else:
                if design != "split":
                    ph_rows, sig_ = pairwise_param(v_names, means1, ns1, MS_err, df_err, method)
                    lnamed = cld(v_names, means1, sig_)
                    for f in self.factor_keys:
                        r_, s_ = pairwise_param(lbf[f], fm[f], fn[f], MS_err, df_err, method)
                        fpt[f] = r_; lf[f] = cld(lbf[f], fm[f], s_)
                else:
                    for f in self.factor_keys:
                        MS_ = MS_wp if f == split_mf else MS_err
                        df_ = df_wp if f == split_mf else df_err
                        r_, s_ = pairwise_param(lbf[f], fm[f], fn[f], MS_, df_, method)
                        fpt[f] = r_; lf[f] = cld(lbf[f], fm[f], s_)

        elif method == "kw":
            try:
                smp = [groups1[n] for n in v_names if groups1[n]]
                if len(smp) >= 2:
                    kwr = kruskal(*smp); kw_H = float(kwr.statistic); kw_p = float(kwr.pvalue)
                    kw_df = len(smp) - 1; kw_eps = eps2_kw(kw_H, len(long), len(smp))
            except Exception: pass
            if not math.isnan(kw_p) and kw_p >= ALPHA: do_ph = False
            if do_ph:
                ph_rows, sig_ = pairwise_mw(v_names, groups1)
                lnamed = cld(v_names, {n: float(np.median(groups1[n])) if groups1[n] else np.nan for n in v_names}, sig_)

        elif method == "mw":
            ph_rows, sig_ = pairwise_mw(v_names, groups1)
            lnamed = cld(v_names, {n: float(np.median(groups1[n])) if groups1[n] else np.nan for n in v_names}, sig_)

        elif method == "friedman":
            bnames = first_seen([f"Блок {i+1}" for i in range(len(used_rep))])
            long2  = [dict(r, VARIANT=" | ".join(str(r.get(f)) for f in self.factor_keys)) for r in long]
            mat, _ = rcbd_matrix(long2, v_names, bnames)
            if len(mat) < 2: messagebox.showwarning("", "Потрібні ≥ 2 повних блоки."); return
            try:
                fr = friedmanchisquare(*[np.array(c, dtype=float) for c in zip(*mat)])
                fr_chi = float(fr.statistic); fr_p = float(fr.pvalue)
                fr_df = n_var - 1; fr_W = kendalls_w(fr_chi, len(mat), n_var)
            except Exception: pass
            if not math.isnan(fr_p) and fr_p < ALPHA:
                rcbd_ph, rcbd_sig = pairwise_wilcox(v_names, mat)
                lnamed = cld(v_names, {n: float(np.median(groups1[n])) if groups1[n] else np.nan for n in v_names}, rcbd_sig)

        elif method == "wilcoxon":
            if n_var != 2: messagebox.showwarning("", "Wilcoxon лише для 2 варіантів."); return
            bnames = first_seen([f"Блок {i+1}" for i in range(len(used_rep))])
            long2  = [dict(r, VARIANT=" | ".join(str(r.get(f)) for f in self.factor_keys)) for r in long]
            mat, _ = rcbd_matrix(long2, v_names, bnames)
            if len(mat) < 2: messagebox.showwarning("", "Потрібні ≥ 2 блоки."); return
            arr = np.array(mat, dtype=float)
            try:
                st, p = wilcoxon(arr[:, 0], arr[:, 1], zero_method="wilcox", alternative="two-sided", mode="auto")
                wil_s = float(st); wil_p = float(p)
            except Exception: pass
            if not math.isnan(wil_p) and wil_p < ALPHA:
                lnamed = cld(v_names, {n: float(np.median(groups1[n])) if groups1[n] else np.nan for n in v_names},
                             {(v_names[0], v_names[1]): True})

        lv_var = {var_order[i]: lnamed.get(v_names[i], "") for i in range(n_var)}
        SS_tot = res.get("SS_total", np.nan); SS_err = res.get("SS_error", np.nan)
        R2 = (1 - (SS_err / SS_tot)) if not any(math.isnan(x) for x in [SS_tot, SS_err]) and SS_tot > 0 else np.nan

        cv_r = [[self.ftitle(f), fmt(cv_means([fm[f].get(lv, np.nan) for lv in lbf[f]]), 2)]
                for f in self.factor_keys]
        cv_r.append(["Загальний", fmt(cv_vals(values), 2)])

        nonparam = method in ("mw", "kw", "friedman", "wilcoxon")
        # Якщо arcsin застосовано автоматично — відображаємо у звіті
        if arcsin_applied and not transform_label:
            transform_label = "arcsin(√p)"
            log_applied = True

        def _rn(nm):
            if not isinstance(nm, str) or not nm.startswith("Фактор "): return nm
            rest = nm.replace("Фактор ", "")
            parts = rest.split("×")
            return "×".join(self.ftitle(p) if p in self.factor_keys else p for p in parts)

        anova_rows = []
        for raw_row in res["table"]:
            nm, SSv, dfv, MSv, Fv, pv = raw_row
            df_s = str(int(dfv)) if dfv is not None and not (isinstance(dfv, float) and math.isnan(dfv)) else ""
            nm2 = _rn(nm)
            if any(x in nm2 for x in ["Залишок", "WP-error", "Блоки"]) or nm2 == "Загальна":
                anova_rows.append([nm2, fmt(SSv, 3), df_s, fmt(MSv, 3), "", "", ""])
            else:
                mk = sig_mark(pv); concl = f"різниця {mk}" if mk else "–"
                anova_rows.append([nm2, fmt(SSv, 3), df_s, fmt(MSv, 3), fmt(Fv, 3), fmt(pv, 4), concl])

        eff_rows  = [[_rn(r[0]), r[1]] for r in build_eff_rows(res["table"])]
        pe2_rows  = [[_rn(r[0]), r[1], r[2]] for r in build_pe2_rows(res["table"])]

        # cache for graph redraw
        self._lbf_cache = lbf

        self.show_report(
            created=created, indicator=indicator, units=units, design=design,
            arcsin_applied=arcsin_applied,
            ordinal_mode=getattr(self, '_ordinal_mode', False),
            ss_type=ss_type, method=method, log_applied=log_applied,
            transform_label=transform_label,
            n_var=n_var, n_rep=len(used_rep), n_obs=len(long),
            split_mf=split_mf, W=W, p_norm=p_norm,
            lev_F=lev_F, lev_p=lev_p,
            kw_H=kw_H, kw_p=kw_p, kw_df=kw_df, kw_eps=kw_eps, do_ph=do_ph,
            fr_chi=fr_chi, fr_p=fr_p, fr_df=fr_df, fr_W=fr_W,
            wil_s=wil_s, wil_p=wil_p,
            anova_rows=anova_rows, eff_rows=eff_rows, pe2_rows=pe2_rows,
            cv_r=cv_r, R2=R2,
            lf=lf, lv_var=lv_var, lbf=lbf,
            fm=fm, fsd=fsd, fmed=fmed, fq=fq, fn=fn,
            var_order=var_order, v_names=v_names,
            v_means=v_means, v_sds=v_sds, v_ns=v_ns,
            vmed=vmed, vq=vq, rkv=rkv, rkf=rkf,
            groups1=groups1, ph_rows=ph_rows, fpt=fpt,
            rcbd_ph=rcbd_ph, nonparam=nonparam, res=res,
        )
        self.show_graphs(long, lf, indicator, units, eff_rows, pe2_rows,
                         parent_win=self.report_win if (
                             self.report_win and
                             tk.Toplevel.winfo_exists(self.report_win)) else None)

    # ════════════════════════════════════════════════════════════
    # REPORT WINDOW
    # ════════════════════════════════════════════════════════════
    def show_report(self, **kw):
        if self.report_win and tk.Toplevel.winfo_exists(self.report_win):
            self.report_win.destroy()
        self.report_win = rw = tk.Toplevel(self.table_win or self.root)
        rw.title("Звіт ANOVA — " + kw.get("indicator",""))
        try: rw.state("zoomed")
        except Exception: rw.geometry("1400x900")
        set_icon(rw)

        # ── Бокова панель + контент ──────────────────────────
        main_frame = tk.Frame(rw); main_frame.pack(fill=tk.BOTH, expand=True)
        sidebar = tk.Frame(main_frame, width=190, bg="#2c3e50")
        sidebar.pack(side=tk.LEFT, fill=tk.Y); sidebar.pack_propagate(False)
        content = tk.Frame(main_frame); content.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        tk.Label(sidebar, text="ЗВІТ", bg="#2c3e50", fg="#ecf0f1",
                 font=("Times New Roman",12,"bold"), pady=12).pack(fill=tk.X)

        self._active_panel = None
        self._active_rpt_btn = None

        def _show_panel(frame, btn):
            if self._active_panel: self._active_panel.pack_forget()
            if self._active_rpt_btn:
                self._active_rpt_btn.configure(bg="#2c3e50", fg="#bdc3c7")
            frame.pack(fill=tk.BOTH, expand=True)
            self._active_panel = frame
            btn.configure(bg="#c62828", fg="white")
            self._active_rpt_btn = btn

        def _sidebar_btn(text, tooltip):
            fr = tk.Frame(sidebar, bg="#2c3e50"); fr.pack(fill=tk.X)
            b = tk.Button(fr, text=f"  {text}", bg="#2c3e50", fg="#bdc3c7",
                          font=("Times New Roman",11), relief=tk.FLAT,
                          anchor="w", padx=12, pady=6,
                          activebackground="#c62828", activeforeground="white")
            b.pack(fill=tk.X)
            tk.Label(fr, text=f"    {tooltip}", bg="#2c3e50", fg="#7f8c8d",
                     font=("Times New Roman",8), anchor="w").pack(fill=tk.X)
            tk.Frame(sidebar, bg="#3d5166", height=1).pack(fill=tk.X)
            return b

        # ── Панель 1: Текстовий звіт ─────────────────────────
        rpt_frame = tk.Frame(content)

        # Toolbar звіту
        rpt_tb = tk.Frame(rpt_frame, padx=6, pady=4); rpt_tb.pack(fill=tk.X)
        self._report_buf = []
        def copy_all():
            rw.clipboard_clear()
            rw.clipboard_append("\n".join(self._report_buf))
            messagebox.showinfo("","Текстовий звіт скопійовано у буфер.\nВставте у Word через Ctrl+V.")
        tk.Button(rpt_tb, text="📋 Копіювати звіт",
                  font=("Times New Roman",11), command=copy_all).pack(side=tk.LEFT, padx=4)

        # Scrollable body звіту
        outer = tk.Frame(rpt_frame); outer.pack(fill=tk.BOTH, expand=True)
        vsb = ttk.Scrollbar(outer, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        hsb = ttk.Scrollbar(outer, orient="horizontal"); hsb.pack(side=tk.BOTTOM, fill=tk.X)
        cv = tk.Canvas(outer, yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        cv.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.config(command=cv.yview); hsb.config(command=cv.xview)
        body = tk.Frame(cv); cv.create_window((0,0), window=body, anchor="nw")
        def _mw(e): cv.yview_scroll(int(-1*(e.delta/120)),"units")
        body.bind("<Configure>",
                  lambda e: (cv.configure(scrollregion=cv.bbox("all"))))
        rw.bind("<MouseWheel>", _mw); cv.bind("<MouseWheel>", _mw)

        # Кнопки бокової панелі
        b_rpt = _sidebar_btn("📄 Текстовий звіт",  "ANOVA таблиця, НІР, висновки")
        b_rpt.configure(command=lambda: _show_panel(rpt_frame, b_rpt))

        buf = self._report_buf
        def _txt(s):
            tk.Label(body, text=s, font=("Times New Roman", 12), fg="#000000",
                     justify="left", anchor="w", wraplength=1100).pack(fill=tk.X, padx=12, pady=1)
            buf.append(s)
        def _head(s):
            tk.Label(body, text=s, font=("Times New Roman", 13, "bold"), fg="#000000",
                     justify="left", anchor="w").pack(fill=tk.X, padx=12, pady=8)
            buf.append("\n" + s)
        def _sep():
            ttk.Separator(body, orient="horizontal").pack(fill=tk.X, padx=12, pady=4)
            buf.append("-" * 80)
        def _table(headers, rows, min_col=90):
            frm, tv_widget = make_tv(body, headers, rows, min_col)
            frm.pack(fill=tk.X, padx=12, pady=(2, 8))
            # bind mousewheel on treeview and its frame so scroll works
            frm.bind("<MouseWheel>", _mw)
            tv_widget.bind("<MouseWheel>", _mw)
            # plain text for clipboard: TAB-separated (Word auto-formats as table)
            buf.append("\t".join(str(h) for h in headers))
            buf.append("\t".join("-" * max(3, len(str(h))) for h in headers))
            for row in rows:
                buf.append("\t".join("" if v is None else str(v) for v in row))
            buf.append("")

        d = kw
        _txt("З В І Т   С Т А Т И С Т И Ч Н О Г О   А Н А Л І З У")
        _txt(f"Показник: {d['indicator']}   |   Одиниці: {d['units']}")
        _txt(f"Дата: {d['created'].strftime('%d.%m.%Y %H:%M')}")
        _sep()
        design_lbl = {"crd": "CRD (повна рандомізація)", "rcbd": "RCBD (блочна рандомізація)", "split": "Split-plot"}[d['design']]
        _txt(f"Дизайн: {design_lbl}   |   Тип SS: {d['ss_type']}   |   Варіантів: {d['n_var']}   |   Повт.: {d['n_rep']}   |   Спостережень: {d['n_obs']}")
        if d['design'] == "split": _txt(f"Головний фактор (WP): {d['split_mf']}")
        if d.get('arcsin_applied'):
            _txt("⚠ Застосовано трансформацію arcsin(√p) для відсоткових даних. "
                 "Середні у звіті наведено у вихідних відсотках. "
                 "Літери значущості (НІР) визначено за трансформованими даними.")
        elif d.get('ordinal_mode'):
            _txt("ℹ БАЛЬНА ШКАЛА: у звіті наведено медіану [Q1; Q3] замість Mean ± SD. "
                 "Параметричні методи заблоковані. "
                 "Для візуалізації використовуйте Boxplot або Dot plot.")
        elif d['log_applied']:
            tl = d.get('transform_label', 'ln(x)')
            _txt(f"⚠ Застосовано трансформацію {tl}. Середні у звіті — у трансформованій шкалі.")
        method_lbl = {"lsd": "НІР₀₅ (LSD)", "tukey": "Тест Тьюкі", "duncan": "Тест Дункана",
                      "bonferroni": "Бонферроні", "kw": "Kruskal–Wallis", "mw": "Mann–Whitney",
                      "friedman": "Friedman", "wilcoxon": "Wilcoxon"}.get(d['method'], "")
        _txt(f"Метод: {method_lbl}")
        _txt("** — p<0.01; * — p<0.05; різні літери → істотна різниця.")
        _sep()
        _txt(f"Shapiro–Wilk (залишки): {norm_txt(d['p_norm'])}   W={fmt(d['W'], 4)}   p={fmt(d['p_norm'], 4)}")

        if d['method'] == "kw" and not math.isnan(d['kw_p']):
            c_ = ("різниця " + sig_mark(d['kw_p'])) if d['kw_p'] < ALPHA else "–"
            _txt(f"Kruskal–Wallis:  H={fmt(d['kw_H'],4)}  df={d['kw_df']}  p={fmt(d['kw_p'],4)}  {c_}   ε²={fmt(d['kw_eps'],4)}")
        if d['method'] == "friedman" and not math.isnan(d['fr_p']):
            c_ = ("різниця " + sig_mark(d['fr_p'])) if d['fr_p'] < ALPHA else "–"
            _txt(f"Friedman:  χ²={fmt(d['fr_chi'],4)}  df={d['fr_df']}  p={fmt(d['fr_p'],4)}  {c_}   Kendall's W={fmt(d['fr_W'],4)}")
        if d['method'] == "wilcoxon" and not math.isnan(d['wil_p']):
            c_ = ("різниця " + sig_mark(d['wil_p'])) if d['wil_p'] < ALPHA else "–"
            _txt(f"Wilcoxon:  W={fmt(d['wil_s'],4)}  p={fmt(d['wil_p'],4)}  {c_}")

        if not d['nonparam']:
            if not math.isnan(d['lev_p']):
                lc = "умова виконується" if d['lev_p'] >= ALPHA else f"умова порушена {sig_mark(d['lev_p'])}"
                _txt(f"Тест Левена (однорідність дисперсій):  F={fmt(d['lev_F'],4)}  p={fmt(d['lev_p'],4)}  {lc}")
            _sep()
            _head("ТАБЛИЦЯ 1. Дисперсійний аналіз (ANOVA)")
            _table(["Джерело", "SS", "df", "MS", "F", "p", "Висновок"], d['anova_rows'])
            _head("ТАБЛИЦЯ 2. Сила впливу факторів (% від SS)")
            _table(["Джерело", "%"], d['eff_rows'])
            _head("ТАБЛИЦЯ 3. Розмір ефекту (partial η²)")
            _table(["Джерело", "partial η²", "Сила ефекту"], d['pe2_rows'])
            _head("ТАБЛИЦЯ 4. Коефіцієнт варіації (CV, %)")
            _table(["Елемент", "CV, %"], d['cv_r'])
            _txt(f"Коефіцієнт детермінації R² = {fmt(d['R2'], 4)}")
            tno = 5
            if d['method'] == "lsd":
                nir_r = [[k, fmt(v, 4)] for k, v in d['res'].get("NIR05", {}).items()]
                if nir_r:
                    _head(f"ТАБЛИЦЯ {tno}. НІР₀₅"); _table(["Елемент", "НІР₀₅"], nir_r); tno += 1

            for f in self.factor_keys:
                _head(f"ТАБЛИЦЯ {tno}. Середні по фактору: {self.ftitle(f)}")
                rows_f = [[str(lv), fmt(d['fm'][f].get(lv, np.nan), 3),
                           fmt(d['fsd'][f].get(lv, np.nan), 3),
                           d['lf'][f].get(lv, "") or "–"]
                          for lv in d['lbf'][f]]
                _table([self.ftitle(f), "Середнє", "± SD", "Літери CLD"], rows_f); tno += 1

            _head(f"ТАБЛИЦЯ {tno}. Середні по варіантах")
            rows_v = [[nm, fmt(d['v_means'].get(d['var_order'][i], np.nan), 3),
                       fmt(d['v_sds'].get(d['var_order'][i], np.nan), 3),
                       d['lv_var'].get(d['var_order'][i], "") or "–"]
                      for i, nm in enumerate(d['v_names'])]
            _table(["Варіант", "Середнє", "± SD", "Літери CLD"], rows_v); tno += 1

            if d['design'] != "split":
                if d['method'] in ("tukey", "duncan", "bonferroni") and d['ph_rows']:
                    _head(f"ТАБЛИЦЯ {tno}. Парні порівняння варіантів")
                    _table(["Пара", "p", "Висновок"], d['ph_rows']); tno += 1
            else:
                for f in self.factor_keys:
                    rr = d['fpt'].get(f, [])
                    if rr:
                        _head(f"ТАБЛИЦЯ {tno}. Парні порівняння: {self.ftitle(f)}")
                        _table(["Пара", "p", "Висновок"], rr); tno += 1
        else:
            tno = 1
            for f in self.factor_keys:
                _head(f"ТАБЛИЦЯ {tno}. Описова (непараметрична): {self.ftitle(f)}")
                rows = [[str(lv), str(d['fn'][f].get(lv, 0)), fmt(d['fmed'][f].get(lv, np.nan), 3),
                         f"{fmt(d['fq'][f].get(lv,(np.nan,np.nan))[0],3)}–{fmt(d['fq'][f].get(lv,(np.nan,np.nan))[1],3)}",
                         fmt(d['rkf'][f].get(lv, np.nan), 2)]
                        for lv in d['lbf'][f]]
                _table([self.ftitle(f), "n", "Медіана", "Q1–Q3", "Сер. ранг"], rows); tno += 1
            _head(f"ТАБЛИЦЯ {tno}. Описова (непараметрична): варіанти")
            rows = [[d['v_names'][i], str(d['v_ns'].get(d['var_order'][i], 0)),
                     fmt(d['vmed'].get(d['var_order'][i], np.nan), 3),
                     f"{fmt(d['vq'].get(d['var_order'][i],(np.nan,np.nan))[0],3)}–{fmt(d['vq'].get(d['var_order'][i],(np.nan,np.nan))[1],3)}",
                     fmt(d['rkv'].get(d['v_names'][i], np.nan), 2)]
                    for i in range(d['n_var'])]
            _table(["Варіант", "n", "Медіана", "Q1–Q3", "Сер. ранг"], rows); tno += 1
            if d['method'] == "kw":
                if not d['do_ph']: _txt("Kruskal–Wallis p ≥ 0.05 → пост-хок не виконувався.")
                elif d['ph_rows']:
                    _head(f"ТАБЛИЦЯ {tno}. Парні порівняння (MWU + Bonferroni, Cliff's δ)")
                    _table(["Пара", "U", "p(Bonf.)", "Висновок", "δ", "Ефект"], d['ph_rows'])
            if d['method'] == "mw" and d['ph_rows']:
                _head(f"ТАБЛИЦЯ {tno}. Парні порівняння (MWU + Bonferroni, Cliff's δ)")
                _table(["Пара", "U", "p(Bonf.)", "Висновок", "δ", "Ефект"], d['ph_rows'])
            if d['method'] == "friedman":
                if not math.isnan(d['fr_p']) and d['fr_p'] >= ALPHA:
                    _txt("Friedman p ≥ 0.05 → пост-хок не виконувався.")
                elif d['rcbd_ph']:
                    _head(f"ТАБЛИЦЯ {tno}. Парні порівняння (Wilcoxon + Bonferroni)")
                    _table(["Пара", "W", "p(Bonf.)", "Висновок", "r"], d['rcbd_ph'])
        _sep()
        _txt(f"Звіт сформовано: {d['created'].strftime('%d.%m.%Y, %H:%M')}")

        # Показуємо текстовий звіт як першу вкладку
        _show_panel(rpt_frame, b_rpt)

        # Зберігаємо sidebar/content/show_panel для show_graphs
        self._rpt_sidebar = sidebar
        self._rpt_content = content
        self._rpt_show_panel = _show_panel
        self._rpt_sidebar_btn = _sidebar_btn

    # ════════════════════════════════════════════════════════════
    # GRAPHICAL REPORT  — sidebar tabs
    # ════════════════════════════════════════════════════════════
    def show_graphs(self, long, letters_factor, indicator, units,
                    eff_rows, pe2_rows, parent_win=None):
        if not HAS_MPL:
            messagebox.showwarning("", "matplotlib недоступний."); return

        # Якщо є вікно звіту — вбудовуємось у його бокову панель
        if (hasattr(self, '_rpt_sidebar') and
                self.report_win and tk.Toplevel.winfo_exists(self.report_win)):
            gw = self.report_win
            sidebar   = self._rpt_sidebar
            content   = self._rpt_content
            _show_panel = self._rpt_show_panel
            _sidebar_btn = self._rpt_sidebar_btn
            # Роздільник між звітом і графіками
            tk.Frame(sidebar, bg="#1a3a4a", height=2).pack(fill=tk.X)
            tk.Label(sidebar, text="ГРАФІКИ", bg="#2c3e50", fg="#ecf0f1",
                     font=("Times New Roman",11,"bold"),
                     pady=8).pack(fill=tk.X)
        else:
            # Відкриваємо окреме вікно
            if self.graph_win and tk.Toplevel.winfo_exists(self.graph_win):
                self.graph_win.destroy()
            self.graph_win = gw = tk.Toplevel(self.table_win or self.root)
            gw.title(f"Графічний звіт — {indicator}")
            try: gw.state("zoomed")
            except Exception: gw.geometry("1400x900")
            set_icon(gw)
            main = tk.Frame(gw); main.pack(fill=tk.BOTH, expand=True)
            sidebar = tk.Frame(main, width=190, bg="#2c3e50")
            sidebar.pack(side=tk.LEFT, fill=tk.Y); sidebar.pack_propagate(False)
            content = tk.Frame(main); content.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            tk.Label(sidebar, text="ГРАФІКИ", bg="#2c3e50", fg="#ecf0f1",
                     font=("Times New Roman",11,"bold"), pady=12).pack(fill=tk.X)
            self._active_panel = None; self._active_rpt_btn = None
            def _show_panel(frame, btn):
                if self._active_panel: self._active_panel.pack_forget()
                if self._active_rpt_btn:
                    self._active_rpt_btn.configure(bg="#2c3e50", fg="#bdc3c7")
                frame.pack(fill=tk.BOTH, expand=True)
                self._active_panel = frame
                btn.configure(bg="#c62828", fg="white")
                self._active_rpt_btn = btn
            def _sidebar_btn(text, tooltip):
                fr = tk.Frame(sidebar, bg="#2c3e50"); fr.pack(fill=tk.X)
                b = tk.Button(fr, text=f"  {text}", bg="#2c3e50", fg="#bdc3c7",
                              font=("Times New Roman",11), relief=tk.FLAT,
                              anchor="w", padx=12, pady=6,
                              activebackground="#c62828", activeforeground="white")
                b.pack(fill=tk.X)
                tk.Label(fr, text=f"    {tooltip}", bg="#2c3e50", fg="#7f8c8d",
                         font=("Times New Roman",8), anchor="w").pack(fill=tk.X)
                tk.Frame(sidebar, bg="#3d5166", height=1).pack(fill=tk.X)
                return b

        # Зберігаємо дані для перебудови
        self._g_long = long; self._g_lf = letters_factor
        self._g_ind  = indicator; self._g_units = units
        self._g_eff  = eff_rows; self._g_pe2 = pe2_rows
        self._graph_figs = {}
        if not hasattr(self, '_gs_titles'): self._gs_titles = {}
        self._lbf_cache = {f: list(letters_factor.get(f, {}).keys())
                           for f in self.factor_keys}

        ordinal = getattr(self, '_ordinal_mode', False)

        # Якщо вбудовуємось у вікно звіту — використовуємо його sidebar і content
        if (hasattr(self, '_rpt_sidebar') and
                self.report_win and tk.Toplevel.winfo_exists(self.report_win)):
            sidebar      = self._rpt_sidebar
            content      = self._rpt_content
            _show_panel  = self._rpt_show_panel
            gw = self.report_win
        else:
            # Окреме вікно (запасний варіант)
            if self.graph_win and tk.Toplevel.winfo_exists(self.graph_win):
                self.graph_win.destroy()
            self.graph_win = gw = tk.Toplevel(self.table_win or self.root)
            gw.title(f"Графічний звіт — {indicator}")
            try: gw.state("zoomed")
            except Exception: gw.geometry("1400x900")
            set_icon(gw)
            outer_f = tk.Frame(gw); outer_f.pack(fill=tk.BOTH, expand=True)
            sidebar = tk.Frame(outer_f, width=195, bg="#2c3e50")
            sidebar.pack(side=tk.LEFT, fill=tk.Y); sidebar.pack_propagate(False)
            content = tk.Frame(outer_f)
            content.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            tk.Label(sidebar, text="ГРАФІКИ", bg="#2c3e50", fg="#ecf0f1",
                     font=("Times New Roman",11,"bold"), pady=10).pack(fill=tk.X)
            self._active_panel = None; self._active_rpt_btn = None
            def _show_panel(frame, btn):
                if self._active_panel: self._active_panel.pack_forget()
                if self._active_rpt_btn:
                    self._active_rpt_btn.configure(bg="#2c3e50", fg="#bdc3c7")
                frame.pack(fill=tk.BOTH, expand=True)
                self._active_panel = frame
                btn.configure(bg="#c62828", fg="white")
                self._active_rpt_btn = btn

        gs = self.graph_settings

        def _make_btn(lbl, tooltip):
            fr = tk.Frame(sidebar, bg="#2c3e50"); fr.pack(fill=tk.X)
            b = tk.Button(fr, text=f"  {lbl}", bg="#2c3e50", fg="#bdc3c7",
                          font=("Times New Roman",11), relief=tk.FLAT,
                          anchor="w", padx=10, pady=5,
                          activebackground="#c62828", activeforeground="white")
            b.pack(fill=tk.X)
            tk.Label(fr, text=f"    {tooltip}", bg="#2c3e50", fg="#7f8c8d",
                     font=("Times New Roman",8), anchor="w").pack(fill=tk.X)
            tk.Frame(sidebar, bg="#3d5166", height=1).pack(fill=tk.X)
            return b

        if ordinal:
            self._ordinal_graph_var = tk.StringVar(value="Boxplot")
            ord_frame = tk.Frame(content)
            b_box = _make_btn("📦 Boxplot", "Медіана + квартилі")
            b_dot = _make_btn("● Dot plot", "Точки + медіана")
            def _show_ord(val, btn):
                self._ordinal_graph_var.set(val)
                _show_panel(ord_frame, btn)
                self._rebuild_ordinal_graph(ord_frame, long, letters_factor,
                                            indicator, units)
            b_box.configure(command=lambda: _show_ord("Boxplot", b_box))
            b_dot.configure(command=lambda: _show_ord("Dot plot", b_dot))
            self._rebuild_ordinal_graph(ord_frame, long, letters_factor,
                                        indicator, units)
            # НЕ показуємо автоматично — чекаємо кліку користувача
            return

        # ── Звичайні вкладки ────────────────────────────────────
        tab_defs = [
            ("bp",   "📦 Boxplot",           "Розподіл даних"),
            ("bar",  "📊 Середні ± SE",       "Стовпчики з планками"),
            ("int",  "↗ Взаємодія",           "Профіль взаємодії"),
            ("line", "📈 Динаміка рівнів",    "Лінійний по рівнях"),
            ("hist", "〰 Залишки",            "Гістограма + Q-Q"),
            ("vn",   "🎯 Сила впливу",        "% від SS"),
            ("pe",   "💡 Розмір ефекту",      "Partial η²"),
        ]

        frames = {}
        for key, lbl, tooltip in tab_defs:
            f = tk.Frame(content); frames[key] = f
            b = _make_btn(lbl, tooltip)
            b.configure(command=lambda k=key, bt=b: _show_panel(frames[k], bt))

        tk.Label(sidebar, text=f"{indicator}\n{units}",
                 bg="#2c3e50", fg="#95a5a6",
                 font=("Times New Roman",8), justify="center",
                 wraplength=180).pack(side=tk.BOTTOM, pady=6)

        # Будуємо всі графіки (але не показуємо жодного)
        self._build_bp_tab(  frames["bp"],   long, letters_factor,
                             indicator, units, gs, gw)
        self._build_bar_tab( frames["bar"],  long, letters_factor,
                             indicator, units, gs, gw)
        self._build_int_tab( frames["int"],  long, letters_factor,
                             indicator, units, gs, gw)
        self._build_line_tab(frames["line"], long, letters_factor,
                             indicator, units, gs, gw)
        self._build_hist_tab(frames["hist"], long, gs, indicator, units, gw)
        self._build_vn_tab(  frames["vn"],   eff_rows, gs, gw)
        self._build_pe_tab(  frames["pe"],   pe2_rows, gs, gw)
        # НЕ показуємо жодну вкладку автоматично



    # ── Toolbar кожної вкладки з PNG і налаштуваннями ─────────
    def _tab_toolbar(self, frame, fig_key, rebuild_fn=None, settings_fn=None):
        tb = tk.Frame(frame, bg="#f0f0f0", padx=4, pady=4)
        tb.pack(fill=tk.X, side=tk.BOTTOM)
        tk.Button(tb, text="💾 Зберегти PNG",
                  font=("Times New Roman",10),
                  command=lambda: self._save_fig_png(fig_key)
                  ).pack(side=tk.LEFT, padx=4)
        tk.Button(tb, text="📋 Копіювати",
                  font=("Times New Roman",10),
                  command=lambda: self._copy_fig(fig_key)
                  ).pack(side=tk.LEFT, padx=4)
        if settings_fn:
            tk.Button(tb, text="⚙ Налаштування",
                      font=("Times New Roman",10),
                      bg="#1a4b8c", fg="white",
                      command=settings_fn
                      ).pack(side=tk.LEFT, padx=4)
        return tb

    def _save_fig_png(self, key):
        fig = self._graph_figs.get(key)
        if fig is None: messagebox.showwarning("","Графік відсутній."); return
        path = filedialog.asksaveasfilename(
            defaultextension=".png",
            filetypes=[("PNG","*.png"),("SVG","*.svg")],
            title="Зберегти графік")
        if not path: return
        try:
            fig.savefig(path, dpi=150, bbox_inches="tight")
            messagebox.showinfo("Збережено", f"Збережено:\n{path}")
        except Exception as ex:
            messagebox.showerror("Помилка", str(ex))

    def _settings_dialog(self, gw, key, rebuild_fn, extra_params=None):
        """Уніфікований діалог налаштувань графіка з заголовком + специфічні опції."""
        if not hasattr(self, '_gs_titles'): self._gs_titles = {}
        dlg = tk.Toplevel(gw or self.report_win or self.root)
        dlg.title("Налаштування графіка"); dlg.resizable(False, False)
        set_icon(dlg); dlg.grab_set()
        rf = ("Times New Roman",12)
        frm = tk.Frame(dlg, padx=16, pady=12); frm.pack()

        # Заголовок графіка
        tk.Label(frm, text="Заголовок графіка:", font=rf
                 ).grid(row=0, column=0, sticky="w", pady=4)
        cur = self._gs_titles.get(key, "")
        tv = tk.StringVar(value=cur)
        tk.Entry(frm, textvariable=tv, width=36, font=rf
                 ).grid(row=0, column=1, sticky="w", padx=8)

        # Шрифт і розмір
        gs = self.graph_settings
        tk.Label(frm, text="Шрифт:", font=rf
                 ).grid(row=1, column=0, sticky="w", pady=4)
        ff_v = tk.StringVar(value=gs.get("font_family","Times New Roman"))
        ttk.Combobox(frm, textvariable=ff_v,
                     values=["Times New Roman","Arial","Calibri","Georgia"],
                     state="readonly", width=18
                     ).grid(row=1, column=1, sticky="w", padx=8)

        tk.Label(frm, text="Розмір шрифту:", font=rf
                 ).grid(row=2, column=0, sticky="w", pady=4)
        fz_v = tk.IntVar(value=gs.get("font_size",10))
        tk.Spinbox(frm, from_=7, to=18, textvariable=fz_v, width=7
                   ).grid(row=2, column=1, sticky="w", padx=8)

        # Специфічні параметри для цього графіка
        extra_vars = {}
        row_offset = 3
        if extra_params:
            for ri, (lbl, key2, default, wtype, opts) in enumerate(extra_params):
                tk.Label(frm, text=lbl, font=rf
                         ).grid(row=row_offset+ri, column=0, sticky="w", pady=4)
                var = tk.StringVar(value=str(gs.get(key2, default)))
                extra_vars[key2] = var
                if wtype == "combo":
                    ttk.Combobox(frm, textvariable=var, values=opts,
                                 state="readonly", width=16
                                 ).grid(row=row_offset+ri, column=1, sticky="w", padx=8)
                elif wtype == "color":
                    def _pick(v=var):
                        c = colorchooser.askcolor(color=v.get(), parent=dlg)
                        if c and c[1]: v.set(c[1])
                    tk.Button(frm, text="Обрати колір",
                              command=_pick, font=rf
                              ).grid(row=row_offset+ri, column=1, sticky="w", padx=8)
                elif wtype == "check":
                    bv = tk.BooleanVar(value=bool(gs.get(key2, default)))
                    extra_vars[key2] = bv
                    tk.Checkbutton(frm, variable=bv
                                   ).grid(row=row_offset+ri, column=1, sticky="w", padx=8)
                else:
                    tk.Entry(frm, textvariable=var, width=10, font=rf
                             ).grid(row=row_offset+ri, column=1, sticky="w", padx=8)

        def _apply():
            # Зберігаємо заголовок
            self._gs_titles[key] = tv.get().strip()
            # Зберігаємо налаштування шрифту
            self.graph_settings["font_family"] = ff_v.get()
            self.graph_settings["font_size"]   = fz_v.get()
            # Специфічні
            for k2, v2 in extra_vars.items():
                self.graph_settings[k2] = v2.get()
            dlg.destroy()
            if rebuild_fn: rebuild_fn()

        bf = tk.Frame(frm); bf.grid(row=row_offset+len(extra_params or [])+1,
                                    column=0, columnspan=2, pady=(12,0))
        tk.Button(bf, text="✓ Застосувати", bg="#c62828", fg="white",
                  font=rf, command=_apply).pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Скасувати", font=rf,
                  command=dlg.destroy).pack(side=tk.LEFT)
        center_win(dlg)

    # ── TAB 1: Boxplot ─────────────────────────────────────────
    def _build_bp_tab(self, frame, long, lf, indicator, units, gs, gw=None):
        for w in frame.winfo_children(): w.destroy()

        def _rebuild():
            self._build_bp_tab(frame, long, lf, indicator, units,
                               self.graph_settings, gw)
        def _settings():
            self._settings_dialog(gw, "bp", _rebuild, extra_params=[
                ("Колір боксів:", "box_color", "#aed6f1", "color", None),
                ("Колір медіани:", "median_color", "#c62828", "color", None),
                ("Показати сітку:", "show_grid", True, "check", None),
            ])
        self._tab_toolbar(frame, "bp", _rebuild, _settings)
        plot_f = tk.Frame(frame); plot_f.pack(fill=tk.BOTH, expand=True)

        fp = {"fontsize": gs["font_size"], "fontfamily": gs["font_family"]}
        ff = gs["font_family"]; fz = gs["font_size"]
        title = self._gs_titles.get("bp", f"{indicator}, {units}")
        fig = Figure(figsize=(10, 6), dpi=100); ax = fig.add_subplot(111)
        positions=[]; data=[]; xlbls=[]; let_list=[]; fcentres=[]
        x=1.; gap=1.
        for f in self.factor_keys:
            lvls = list(lf.get(f,{}).keys()) or first_seen(
                [r.get(f) for r in long if r.get(f)])
            if not lvls: continue
            sx=x
            for lv in lvls:
                arr=[float(r["value"]) for r in long if r.get(f)==lv
                     and not math.isnan(float(r.get("value",float("nan"))))]
                data.append(arr); positions.append(x)
                xlbls.append(str(lv)); let_list.append((f,lv)); x+=1.
            fcentres.append(((sx+x-1)/2., self.ftitle(f))); x+=gap
        if data:
            bp=ax.boxplot(data, positions=positions, widths=0.6,
                          showfliers=True, patch_artist=True)
            for p in bp["boxes"]:
                p.set(facecolor=gs.get("box_color","#aed6f1"), alpha=0.85)
            for m in bp["medians"]:
                m.set(color=gs.get("median_color","#c62828"), linewidth=2)
            for w in bp["whiskers"]+bp["caps"]:
                w.set(color=gs.get("whisker_color","#555"), linewidth=1.2)
            for fl in bp["fliers"]:
                fl.set(markerfacecolor=gs.get("flier_color","#c62828"),
                       marker="o", markersize=4)
            ax.set_xticks(positions)
            ax.set_xticklabels(xlbls, rotation=30, ha="right",
                               fontfamily=ff, fontsize=max(7,fz-1))
            allv=[v for a in data for v in a]
            if len(allv)>1 and max(allv)>min(allv):
                off=0.04*(max(allv)-min(allv))
                for i,(f_,lv_) in enumerate(let_list):
                    lt=lf.get(f_,{}).get(lv_,"")
                    if lt and data[i]:
                        ax.text(positions[i], max(data[i])+off, lt,
                                ha="center", va="bottom", **fp)
            for cx,fnm in fcentres:
                ax.text(cx,-0.22,fnm,ha="center",va="top",
                        transform=ax.get_xaxis_transform(),**fp)
            fig.subplots_adjust(bottom=0.28,top=0.91,left=0.08,right=0.98)
        ax.set_title(title, **fp); ax.set_ylabel(units, **fp)
        if gs.get("show_grid", True):
            ax.yaxis.grid(True, linestyle="--", alpha=0.35)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        self._graph_figs["bp"]=fig
        embed_figure(fig, plot_f)

    # ── TAB 2: Середні ± SE ────────────────────────────────────
    def _build_bar_tab(self, frame, long, lf, indicator, units, gs, gw=None):
        for w in frame.winfo_children(): w.destroy()

        def _rebuild():
            self._build_bar_tab(frame, long, lf, indicator, units,
                                self.graph_settings, gw)
        def _settings():
            self._settings_dialog(gw, "bar", _rebuild, extra_params=[
                ("Колір стовпців:", "bar_color", "#4c72b0", "color", None),
                ("Показати сітку:", "show_grid", True, "check", None),
            ])
        self._tab_toolbar(frame, "bar", _rebuild, _settings)
        plot_f = tk.Frame(frame); plot_f.pack(fill=tk.BOTH, expand=True)

        fp = {"fontsize": gs["font_size"], "fontfamily": gs["font_family"]}
        ff = gs["font_family"]; fz = gs["font_size"]
        colors_list = ["#4c72b0","#dd8452","#55a868","#c44e52",
                       "#8172b2","#937860","#da8bc3","#8c8c8c"]
        title = self._gs_titles.get("bar", f"{indicator}, {units}")
        fig = Figure(figsize=(10, 6), dpi=100); ax = fig.add_subplot(111)
        positions=[]; means=[]; ses=[]; xlbls=[]; let_list=[]; fcentres=[]
        bar_colors=[]; x=1.; gap=1.; ci=0
        for f in self.factor_keys:
            lvls = list(lf.get(f,{}).keys()) or first_seen(
                [r.get(f) for r in long if r.get(f)])
            if not lvls: continue
            sx=x
            for lv in lvls:
                arr=[float(r["value"]) for r in long if r.get(f)==lv
                     and not math.isnan(float(r.get("value",float("nan"))))]
                n=len(arr); m=float(np.mean(arr)) if arr else 0.
                se=float(np.std(arr,ddof=1)/math.sqrt(n)) if n>1 else 0.
                means.append(m); ses.append(se); positions.append(x)
                xlbls.append(str(lv)); let_list.append((f,lv))
                bar_colors.append(colors_list[ci%len(colors_list)])
                x+=1.; ci+=1
            fcentres.append(((sx+x-1)/2., self.ftitle(f))); x+=gap
        if means:
            ax.bar(positions, means, yerr=ses, capsize=4, width=0.65,
                   color=bar_colors, edgecolor="white", linewidth=0.8,
                   error_kw={"ecolor":"#444","lw":1.2,"capthick":1.2})
            allv=[m+se for m,se in zip(means,ses)]
            if allv and means:
                off=max(0.02*(max(allv)-min(means)) if len(allv)>1 else 0.3, 0.3)
                for i,(f_,lv_) in enumerate(let_list):
                    lt=lf.get(f_,{}).get(lv_,"")
                    if lt:
                        ax.text(positions[i], means[i]+ses[i]+off, lt,
                                ha="center", va="bottom", **fp)
            ax.set_xticks(positions)
            ax.set_xticklabels(xlbls, rotation=30, ha="right",
                               fontfamily=ff, fontsize=max(7,fz-1))
            for cx,fnm in fcentres:
                ax.text(cx,-0.22,fnm,ha="center",va="top",
                        transform=ax.get_xaxis_transform(),**fp)
            fig.subplots_adjust(bottom=0.28,top=0.91,left=0.08,right=0.98)
        ax.set_title(title, **fp); ax.set_ylabel(units, **fp)
        if gs.get("show_grid", True):
            ax.yaxis.grid(True, linestyle="--", alpha=0.35)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        self._graph_figs["bar"]=fig
        embed_figure(fig, plot_f)

    # ── TAB 3: Взаємодія ───────────────────────────────────────
    def _build_int_tab(self, frame, long, lf, indicator, units, gs, gw=None):
        for w in frame.winfo_children(): w.destroy()

        def _rebuild():
            self._build_int_tab(frame, long, lf, indicator, units,
                                self.graph_settings, gw)
        def _settings():
            self._settings_dialog(gw, "int", _rebuild, extra_params=[
                ("Товщина лінії:", "line_width", 1.8, "entry", None),
                ("Маркер:", "marker_style", "o", "combo",
                 ["o","s","^","D","v","*"]),
                ("Показати сітку:", "show_grid", True, "check", None),
            ])
        self._tab_toolbar(frame, "int", _rebuild, _settings)
        plot_f = tk.Frame(frame); plot_f.pack(fill=tk.BOTH, expand=True)

        fp = {"fontsize": gs["font_size"], "fontfamily": gs["font_family"]}
        colors_list = ["#4c72b0","#dd8452","#55a868","#c44e52",
                       "#8172b2","#937860","#da8bc3","#8c8c8c"]
        title = self._gs_titles.get("int",
            "Графік взаємодії факторів (профіль середніх)")
        fkeys = self.factor_keys
        fig = Figure(figsize=(10, 6), dpi=100); ax = fig.add_subplot(111)
        lw = float(gs.get("line_width", 1.8))
        mk = gs.get("marker_style", "o")
        if len(fkeys) >= 2:
            f1, f2 = fkeys[0], fkeys[1]
            lvls1 = list(lf.get(f1,{}).keys()) or first_seen(
                [r.get(f1) for r in long if r.get(f1)])
            lvls2 = list(lf.get(f2,{}).keys()) or first_seen(
                [r.get(f2) for r in long if r.get(f2)])
            for gi, lv2 in enumerate(lvls2):
                means_=[float(np.mean([r["value"] for r in long
                    if r.get(f1)==lv1 and r.get(f2)==lv2
                    and not math.isnan(r.get("value",float("nan")))] or [float("nan")]))
                    for lv1 in lvls1]
                ax.plot(range(len(lvls1)), means_, marker=mk, label=str(lv2),
                        color=colors_list[gi%len(colors_list)],
                        linewidth=lw, markersize=7)
            ax.set_xticks(range(len(lvls1)))
            ax.set_xticklabels([str(l) for l in lvls1],
                               rotation=20, ha="right", **fp)
            ax.set_xlabel(self.ftitle(f1), **fp)
            ax.legend(title=self.ftitle(f2), fontsize=fp["fontsize"]-1,
                      title_fontsize=fp["fontsize"]-1)
        elif len(fkeys) == 1:
            f1 = fkeys[0]
            lvls1 = list(lf.get(f1,{}).keys()) or first_seen(
                [r.get(f1) for r in long if r.get(f1)])
            means_=[float(np.mean([r["value"] for r in long if r.get(f1)==lv
                and not math.isnan(r.get("value",float("nan")))] or [float("nan")]))
                for lv in lvls1]
            ax.plot(range(len(lvls1)), means_, marker=mk,
                    color=colors_list[0], linewidth=lw, markersize=8)
            ax.set_xticks(range(len(lvls1)))
            ax.set_xticklabels([str(l) for l in lvls1], **fp)
        else:
            ax.text(0.5,0.5,"Потрібно ≥ 2 фактори",
                    ha="center",va="center",transform=ax.transAxes,**fp)
            ax.axis("off")
        ax.set_title(title, **fp); ax.set_ylabel(units, **fp)
        if gs.get("show_grid", True):
            ax.yaxis.grid(True, linestyle="--", alpha=0.35)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        fig.tight_layout()
        self._graph_figs["int"]=fig
        embed_figure(fig, plot_f)

    # ── TAB 4: Динаміка по рівнях ──────────────────────────────
    def _build_line_tab(self, frame, long, lf, indicator, units, gs, gw=None):
        for w in frame.winfo_children(): w.destroy()

        def _rebuild():
            self._build_line_tab(frame, long, lf, indicator, units,
                                 self.graph_settings, gw)
        def _settings():
            self._settings_dialog(gw, "line", _rebuild, extra_params=[
                ("Товщина лінії:", "line_width", 1.8, "entry", None),
                ("Маркер:", "marker_style", "o", "combo",
                 ["o","s","^","D","v","*"]),
                ("Показати сітку:", "show_grid", True, "check", None),
            ])
        self._tab_toolbar(frame, "line", _rebuild, _settings)
        plot_f = tk.Frame(frame); plot_f.pack(fill=tk.BOTH, expand=True)

        fp = {"fontsize": gs["font_size"], "fontfamily": gs["font_family"]}
        colors_list = ["#4c72b0","#dd8452","#55a868","#c44e52",
                       "#8172b2","#937860","#da8bc3","#8c8c8c"]
        title = self._gs_titles.get("line",
            "Середні ± SE по рівнях кожного фактора")
        fkeys = self.factor_keys; n = len(fkeys)
        lw = float(gs.get("line_width", 1.8))
        mk = gs.get("marker_style", "o")
        fig = Figure(figsize=(10, 6), dpi=100)
        if n == 0:
            ax = fig.add_subplot(111)
            ax.text(0.5,0.5,"Немає факторів",ha="center",va="center")
            ax.axis("off")
        else:
            for fi, f in enumerate(fkeys):
                ax = fig.add_subplot(1, n, fi+1)
                lvls = list(lf.get(f,{}).keys()) or first_seen(
                    [r.get(f) for r in long if r.get(f)])
                means_=[]; ses_=[]
                for lv in lvls:
                    arr=[r["value"] for r in long if r.get(f)==lv
                         and not math.isnan(r.get("value",float("nan")))]
                    means_.append(float(np.mean(arr)) if arr else float("nan"))
                    ses_.append(float(np.std(arr,ddof=1)/math.sqrt(len(arr)))
                                if len(arr)>1 else 0.)
                ax.errorbar(range(len(lvls)), means_, yerr=ses_,
                           marker=mk, color=colors_list[fi%len(colors_list)],
                           linewidth=lw, markersize=7, capsize=4, ecolor="#555")
                ax.set_xticks(range(len(lvls)))
                ax.set_xticklabels([str(l) for l in lvls],
                                   rotation=20, ha="right", **fp)
                ax.set_title(self.ftitle(f), **fp)
                ax.set_ylabel(units if fi==0 else "", **fp)
                if gs.get("show_grid", True):
                    ax.yaxis.grid(True, linestyle="--", alpha=0.35)
                ax.spines["top"].set_visible(False)
                ax.spines["right"].set_visible(False)
                for i,lv in enumerate(lvls):
                    lt=lf.get(f,{}).get(lv,"")
                    if lt and not math.isnan(means_[i]):
                        ax.text(i, means_[i]+ses_[i]+0.3, lt,
                                ha="center", va="bottom", **fp)
            fig.suptitle(title, **fp)
            fig.tight_layout()
        self._graph_figs["line"]=fig
        embed_figure(fig, plot_f)

    # ── TAB 5: Залишки ─────────────────────────────────────────
    def _build_hist_tab(self, frame, long, gs, indicator, units, gw=None):
        for w in frame.winfo_children(): w.destroy()

        def _rebuild():
            self._build_hist_tab(frame, long, self.graph_settings,
                                 indicator, units, gw)
        def _settings():
            self._settings_dialog(gw, "hist", _rebuild, extra_params=[
                ("Колір гістограми:", "hist_color", "#4c72b0", "color", None),
            ])
        self._tab_toolbar(frame, "hist", _rebuild, _settings)
        plot_f = tk.Frame(frame); plot_f.pack(fill=tk.BOTH, expand=True)

        fp = {"fontsize": gs["font_size"], "fontfamily": gs["font_family"]}
        title = self._gs_titles.get("hist", "Аналіз залишків")
        fig = Figure(figsize=(10, 6), dpi=100)
        residuals = getattr(self, '_last_residuals', None)
        if residuals and len(residuals) > 2:
            res = np.array(residuals)
            ax1 = fig.add_subplot(121)
            ax1.hist(res, bins="auto",
                     color=gs.get("hist_color","#4c72b0"),
                     edgecolor="white", alpha=0.85)
            ax1.set_title("Гістограма залишків", **fp)
            ax1.set_xlabel("Залишок", **fp); ax1.set_ylabel("Частота", **fp)
            ax1.yaxis.grid(True, linestyle="--", alpha=0.3)
            ax1.spines["top"].set_visible(False); ax1.spines["right"].set_visible(False)
            from scipy.stats import probplot
            ax2 = fig.add_subplot(122)
            (osm,osr),(slope,intercept,r)=probplot(res, plot=None)
            ax2.plot(osm, osr, "o",
                     color=gs.get("box_color","#aed6f1"),
                     markersize=5, alpha=0.8)
            ax2.plot([min(osm),max(osm)],
                     [slope*min(osm)+intercept, slope*max(osm)+intercept],
                     color=gs.get("median_color","#c62828"), lw=1.5)
            ax2.set_title(f"Q-Q графік (R²={r**2:.3f})", **fp)
            ax2.set_xlabel("Теоретичні квантилі", **fp)
            ax2.set_ylabel("Вибіркові квантилі", **fp)
            ax2.spines["top"].set_visible(False); ax2.spines["right"].set_visible(False)
        else:
            ax = fig.add_subplot(111)
            ax.text(0.5,0.5,"Залишки недоступні.\nВиконайте аналіз і відкрийте знову.",
                    ha="center",va="center",transform=ax.transAxes,**fp)
            ax.axis("off")
        fig.suptitle(title, **fp); fig.tight_layout()
        self._graph_figs["hist"]=fig
        embed_figure(fig, plot_f)

    # ── TAB 6: Сила впливу ─────────────────────────────────────
    def _build_vn_tab(self, frame, eff_rows, gs, gw=None):
        for w in frame.winfo_children(): w.destroy()

        def _rebuild():
            self._build_vn_tab(frame, self._g_eff, self.graph_settings, gw)
        def _settings():
            self._settings_dialog(gw, "vn", _rebuild, extra_params=[
                ("Колір головних ефектів:", "vn_main_color", "#1a4b8c",
                 "color", None),
                ("Колір взаємодій:", "vn_inter_color", "#c62828",
                 "color", None),
            ])
        self._tab_toolbar(frame, "vn", _rebuild, _settings)
        plot_f = tk.Frame(frame); plot_f.pack(fill=tk.BOTH, expand=True)

        fp = {"fontsize": gs["font_size"], "fontfamily": gs["font_family"]}
        title = self._gs_titles.get("vn", "Сила впливу факторів (% від суми SS)")
        fig = Figure(figsize=(10, 6), dpi=100); ax = fig.add_subplot(111)
        valid = [(str(nm), float(pct)) for nm,pct in eff_rows
                 if pct and not math.isnan(float(pct)) and float(pct)>0]
        if valid:
            valid.sort(key=lambda x: x[1])
            labels_ = [v[0] for v in valid]
            values_ = [v[1] for v in valid]
            main_c  = gs.get("vn_main_color","#1a4b8c")
            inter_c = gs.get("vn_inter_color","#c62828")
            colors_ = [inter_c if "×" in l else main_c for l in labels_]
            bars = ax.barh(range(len(labels_)), values_,
                           color=colors_, edgecolor="white", height=0.6)
            for i,(bar,val) in enumerate(zip(bars,values_)):
                ax.text(val+0.3, i, f"{val:.1f}%", va="center", **fp)
            ax.set_yticks(range(len(labels_)))
            ax.set_yticklabels(labels_, **fp)
            ax.set_xlabel("% від суми SS", **fp)
            ax.xaxis.grid(True, linestyle="--", alpha=0.35)
            from matplotlib.patches import Patch
            ax.legend(handles=[
                Patch(color=main_c,  label="Головний ефект"),
                Patch(color=inter_c, label="Взаємодія"),
            ], fontsize=fp["fontsize"]-1, loc="lower right")
        else:
            ax.text(0.5,0.5,"Немає даних",ha="center",va="center",
                    transform=ax.transAxes,**fp); ax.axis("off")
        ax.set_title(title, **fp)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        fig.tight_layout()
        self._graph_figs["vn"]=fig
        embed_figure(fig, plot_f)

    # ── TAB 7: Розмір ефекту ───────────────────────────────────
    def _build_pe_tab(self, frame, pe2_rows, gs, gw=None):
        for w in frame.winfo_children(): w.destroy()

        def _rebuild():
            self._build_pe_tab(frame, self._g_pe2, self.graph_settings, gw)
        def _settings():
            self._settings_dialog(gw, "pe", _rebuild, extra_params=[
                ("Колір головних ефектів:", "pe_main_color", "#1a6b1a",
                 "color", None),
                ("Колір взаємодій:", "pe_inter_color", "#c62828",
                 "color", None),
            ])
        self._tab_toolbar(frame, "pe", _rebuild, _settings)
        plot_f = tk.Frame(frame); plot_f.pack(fill=tk.BOTH, expand=True)

        fp = {"fontsize": gs["font_size"], "fontfamily": gs["font_family"]}
        title = self._gs_titles.get("pe", "Розмір ефекту (partial η²)")
        fig = Figure(figsize=(10, 6), dpi=100); ax = fig.add_subplot(111)
        valid = [(str(nm), float(pct)) for nm,pct,_ in pe2_rows
                 if pct and not math.isnan(float(pct)) and float(pct)>0]
        if valid:
            valid.sort(key=lambda x: x[1])
            labels_ = [v[0] for v in valid]
            values_ = [v[1] for v in valid]
            main_c  = gs.get("pe_main_color","#1a6b1a")
            inter_c = gs.get("pe_inter_color","#c62828")
            colors_ = [inter_c if "×" in l else main_c for l in labels_]
            ax.barh(range(len(labels_)), values_,
                    color=colors_, edgecolor="white", height=0.6)
            for i,val in enumerate(values_):
                strength = ("дуже слабкий" if val<0.01 else
                            "слабкий" if val<0.06 else
                            "середній" if val<0.14 else "сильний")
                ax.text(val+0.002, i, f"η²={val:.3f} ({strength})",
                        va="center", **fp)
            ax.set_yticks(range(len(labels_)))
            ax.set_yticklabels(labels_, **fp)
            ax.set_xlabel("partial η²", **fp)
            for thresh, lbl, col in [(0.01,"мала","#aaa"),
                                     (0.06,"середня","#888"),
                                     (0.14,"велика","#555")]:
                ax.axvline(thresh, color=col, lw=0.8, linestyle="--")
                ax.text(thresh, len(labels_)-0.5, lbl,
                        color=col, fontsize=max(7,fp["fontsize"]-2),
                        ha="center", va="bottom")
            ax.xaxis.grid(True, linestyle="--", alpha=0.25)
        else:
            ax.text(0.5,0.5,"Немає даних",ha="center",va="center",
                    transform=ax.transAxes,**fp); ax.axis("off")
        ax.set_title(title, **fp)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        fig.tight_layout()
        self._graph_figs["pe"]=fig
        embed_figure(fig, plot_f)

    def _tab_settings(self, *args, **kwargs): pass  # замінено _settings_dialog

    def _open_gs(self, gw, long, lf, ind, units, eff, pe2):
        dlg = GraphSettingsDlg(gw, self.graph_settings)
        gw.wait_window(dlg)
        if dlg.result: self.graph_settings = dlg.result


        plot_f = tk.Frame(frame); plot_f.pack(fill=tk.BOTH, expand=True)
        self._tab_toolbar(frame, "bp",
            on_settings=lambda: self._tab_settings(
                frame, "bp", indicator, units,
                rebuild=lambda: self._build_bp_tab(
                    frame, long, lf, indicator, units,
                    self.graph_settings)))
        fp = {"fontsize": gs["font_size"], "fontfamily": gs["font_family"]}
        ff = gs["font_family"]; fz = gs["font_size"]
        fig = Figure(figsize=(10, 6), dpi=100); ax = fig.add_subplot(111)
        positions=[]; data=[]; xlbls=[]; let_list=[]; fcentres=[]
        x=1.; gap=1.
        for f in self.factor_keys:
            lvls = list(lf.get(f,{}).keys()) or first_seen(
                [r.get(f) for r in long if r.get(f)])
            if not lvls: continue
            sx=x
            for lv in lvls:
                arr=[float(r["value"]) for r in long if r.get(f)==lv
                     and not math.isnan(float(r.get("value",float("nan"))))]
                data.append(arr); positions.append(x)
                xlbls.append(str(lv)); let_list.append((f,lv)); x+=1.
            fcentres.append(((sx+x-1)/2., self.ftitle(f))); x+=gap
        title = getattr(self, '_gs_titles', {}).get("bp", f"{indicator}, {units}")
        if data:
            bp=ax.boxplot(data,positions=positions,widths=0.6,showfliers=True,patch_artist=True)
            for p in bp["boxes"]: p.set(facecolor=gs["box_color"])
            for m in bp["medians"]: m.set(color=gs["median_color"],linewidth=2)
            for w in bp["whiskers"]+bp["caps"]: w.set(color=gs["whisker_color"])
            for fl in bp["fliers"]: fl.set(markerfacecolor=gs["flier_color"],marker="o",markersize=4)
            ax.set_xticks(positions)
            ax.set_xticklabels(xlbls,rotation=90,fontfamily=ff,fontsize=max(7,fz-1))
            allv=[v for a in data for v in a]
            if len(allv)>1:
                off=max(0.01*(max(allv)-min(allv)),0.3)
                for i,(f_,lv_) in enumerate(let_list):
                    lt=(lf.get(f_,{})).get(lv_,"")
                    if lt and data[i]:
                        ax.text(positions[i],max(data[i])+off,lt,
                                ha="center",va="bottom",**fp)
            for cx,fnm in fcentres:
                ax.text(cx,-0.22,fnm,ha="center",va="top",
                        transform=ax.get_xaxis_transform(),**fp)
            fig.subplots_adjust(bottom=0.32,top=0.91,left=0.08,right=0.98)
        ax.set_title(title,**fp); ax.set_ylabel(units,**fp)
        ax.yaxis.grid(True,linestyle="--",alpha=0.35)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        self._graph_figs["bp"]=fig
        embed_figure(fig, plot_f)
        fig.set_size_inches(plot_f.winfo_width()/100 or 10,
                            plot_f.winfo_height()/100 or 5)

    # ── TAB 2: Середні ± SE ────────────────────────────────────
    def _open_gs(self, gw, long, lf, ind, units, eff, pe2):
        dlg = GraphSettingsDlg(gw, self.graph_settings)
        gw.wait_window(dlg)
        if dlg.result:
            self.graph_settings = dlg.result



    def _copy_fig(self, key):
        fig = self._graph_figs.get(key)
        if fig is None: messagebox.showwarning("","Графік відсутній."); return
        ok, msg = _copy_fig_to_clipboard(fig)
        if ok: messagebox.showinfo("","Графік скопійовано (PNG).\nВставте у Word через Ctrl+V.")
        else:  messagebox.showwarning("",f"Помилка: {msg}")



# ═══════════════════════════════════════════════════════════════
# EXPORT  — Word (.docx) and PDF via reportlab
# ═══════════════════════════════════════════════════════════════
def export_report_docx(text_lines, tables, filepath):
    """Export plain-text + table data to .docx using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("Встановіть python-docx:\n  pip install python-docx")
    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Times New Roman'; style.font.size = Pt(12)
    for section in doc.sections:
        section.top_margin = Cm(2); section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2)
    for item in text_lines:
        kind = item.get("kind", "text")
        if kind == "heading":
            p = doc.add_heading(item["text"], level=item.get("level", 2))
            p.runs[0].font.name = 'Times New Roman'
        elif kind == "table":
            headers = item["headers"]; rows = item["rows"]
            tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
            tbl.style = 'Table Grid'
            hrow = tbl.rows[0]
            for j, h in enumerate(headers):
                hrow.cells[j].text = str(h)
                run = hrow.cells[j].paragraphs[0].runs[0]
                run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(11)
            for i, row in enumerate(rows):
                for j, val in enumerate(row):
                    cell = tbl.rows[i+1].cells[j]
                    cell.text = "" if val is None else str(val)
                    cell.paragraphs[0].runs[0].font.name = 'Times New Roman' if cell.paragraphs[0].runs else None
            doc.add_paragraph()
        else:
            p = doc.add_paragraph(item.get("text", ""))
            p.runs[0].font.name = 'Times New Roman' if p.runs else None
    doc.save(filepath)

def export_report_pdf(text_lines, filepath):
    """Export plain text to PDF using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
    except ImportError:
        raise RuntimeError("Встановіть reportlab:\n  pip install reportlab")
    doc = SimpleDocTemplate(filepath, pagesize=A4,
                            rightMargin=2*cm, leftMargin=2.5*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle('h1', parent=styles['Heading2'], fontName='Times-Roman', fontSize=12, spaceAfter=4)
    normal = ParagraphStyle('n', parent=styles['Normal'], fontName='Times-Roman', fontSize=11, spaceAfter=2)
    story = []
    for item in text_lines:
        kind = item.get("kind","text")
        if kind == "heading":
            story.append(Paragraph(item["text"], h1))
        elif kind == "table":
            headers = item["headers"]; rows = item["rows"]
            data = [headers] + [[("" if v is None else str(v)) for v in r] for r in rows]
            t = Table(data, repeatRows=1)
            t.setStyle(TableStyle([
                ('FONTNAME',(0,0),(-1,0),'Times-Roman'), ('FONTSIZE',(0,0),(-1,0),10),
                ('FONTNAME',(0,1),(-1,-1),'Times-Roman'), ('FONTSIZE',(0,1),(-1,-1),9),
                ('GRID',(0,0),(-1,-1),0.5,colors.black),
                ('BACKGROUND',(0,0),(-1,0),colors.HexColor('#e8e8e8')),
            ]))
            story.append(t); story.append(Spacer(1,6))
        else:
            txt = item.get("text","").replace("\n","<br/>")
            if txt.strip(): story.append(Paragraph(txt, normal))
    doc.build(story)


# ═══════════════════════════════════════════════════════════════
# DESCRIPTIVE STATISTICS — standalone module
# ═══════════════════════════════════════════════════════════════
class DescriptiveWindow:
    """Описова статистика — окремий модуль."""

    HELP_TEXT = """
ОПИСОВА СТАТИСТИКА — ПОКРОКОВА ІНСТРУКЦІЯ
══════════════════════════════════════════

ЩО ЦЕ І НАВІЩО?
  Описова статистика — перший крок будь-якого аналізу.
  Вона описує ваші дані числовими характеристиками і
  допомагає виявити аномалії до проведення статистичних тестів.

КРОК 1. ПІДГОТОВКА ДАНИХ
  Кожен стовпець = один показник (змінна).
  Кожен рядок = одне спостереження.
  Двічі клікніть на синій заголовок стовпця щоб задати
  назву показника (наприклад: «Врожайність, т/га»).

  Приклад:
  | Врожайність | Висота | Маса 1000 зерен |
  |    4.2      |  98.5  |      38.2       |
  |    5.1      | 103.2  |      41.5       |

КРОК 2. ЗАПУСК
  Натисніть «▶ Аналіз».
  Мінімум: 2 значення у стовпці для розрахунку SD і SE.

КРОК 3. ІНТЕРПРЕТАЦІЯ ТАБЛИЦІ РЕЗУЛЬТАТІВ

  n — кількість числових значень у стовпці.
    Менше 5 — результати ненадійні.

  Середнє (Mean) — середньоарифметичне.
    Чутливе до викидів: одне екстремальне значення
    може суттєво змінити середнє.

  SD (стандартне відхилення) — середнє відхилення
    від середнього. Велике SD = великий розкид.

  СП (SE, стандартна похибка середнього):
    SE = SD / √n
    Показує точність оцінки середнього.
    Чим більше n, тим менше SE.

  Мін / Макс — найменше і найбільше значення.
    Перевіряйте на помилки введення!

  Медіана — середнє значення впорядкованого ряду.
    Стійка до викидів. Якщо Середнє >> Медіани —
    розподіл правоскошений.

  Q1 / Q3 — 25-й і 75-й перцентилі.
    IQR = Q3 - Q1 (міжквартильний розмах).

  CV% (коефіцієнт варіації):
    CV = SD/Середнє × 100%
    Оцінка точності польового досліду:
    < 10%: відмінна | 10-15%: хороша
    15-20%: задовільна | > 20%: низька

  Асиметрія (Skewness):
    = 0: симетричний розподіл
    > 0: правостороня асиметрія (хвіст праворуч)
    < 0: лівостороня асиметрія

  Ексцес (Kurtosis):
    = 0: нормальний розподіл
    > 0: гостроверхий (більше значень поблизу середнього)
    < 0: пласковерхий

  95% ДІ — довірчий інтервал середнього.
    З ймовірністю 95% «справжнє» середнє знаходиться
    в цьому діапазоні.

  SW p (Shapiro-Wilk):
    Тест нормальності розподілу.
    p > 0.05: розподіл нормальний ✓
    p ≤ 0.05: розподіл ненормальний ⚠
    При n > 50 тест може бути надто чутливим —
    оцінюйте QQ-графік візуально.

КРОК 4. ГРАФІКИ

  Боксплот (коробка з вусами):
    Показує розподіл кожного показника.
    Коробка: Q1 - Q3 | Лінія: медіана
    Вуса: Q1-1.5×IQR до Q3+1.5×IQR
    Точки поза вусами: викиди (outliers)

  QQ-графік (квантиль-квантиль):
    Перевірка нормальності візуально.
    Точки лежать на прямій → нормальний розподіл ✓
    Точки відхиляються від прямої → ненормальний ⚠
"""

    def __init__(self, parent, gs):
        self.win = tk.Toplevel(parent)
        self.win.title("Описова статистика")
        self.win.geometry("1000x640"); set_icon(self.win)
        self.gs = dict(gs)
        self._bp_fig  = None   # боксплот
        self._qq_fig  = None   # QQ-графіки
        self._bp_gs   = {      # налаштування боксплоту
            "font_family": gs.get("font_family","Times New Roman"),
            "font_size": 11,
            "box_color":    gs.get("box_color","#ffffff"),
            "median_color": gs.get("median_color","#c62828"),
            "whisker_color":gs.get("whisker_color","#000000"),
            "flier_color":  gs.get("flier_color","#555555"),
        }
        self._qq_gs   = {      # налаштування QQ
            "font_family": gs.get("font_family","Times New Roman"),
            "font_size": 9,
            "pt_color":    "#4c72b0",
            "line_color":  "#c62828",
        }
        self._build()

    # ── Побудова вікна ───────────────────────────────────────
    def _build(self):
        # ── Меню ──
        mb = tk.Menu(self.win); self.win.config(menu=mb)
        fm = tk.Menu(mb, tearoff=0)
        fm.add_command(label="Завантажити Excel", command=self._load_excel)
        mb.add_cascade(label="Файл", menu=fm)

        # ── Панель інструментів ──
        tb = tk.Frame(self.win, padx=6, pady=5); tb.pack(fill=tk.X)

        tk.Button(tb, text="▶ Аналіз", bg="#c62828", fg="white",
                  font=("Times New Roman", 13),
                  command=self._analyze).pack(side=tk.LEFT, padx=4)

        # Налаштування — спадне меню
        mb2 = tk.Menubutton(tb, text="⚙ Налаштування ▾",
                            font=("Times New Roman", 11),
                            relief=tk.RAISED, bd=2)
        mb2.pack(side=tk.LEFT, padx=4)
        sm = tk.Menu(mb2, tearoff=0)
        sm.add_command(label="Додати рядок",      command=self._add_row)
        sm.add_command(label="Видалити рядок",    command=self._del_row)
        sm.add_separator()
        sm.add_command(label="Додати стовпець",   command=self._add_col)
        sm.add_command(label="Видалити стовпець", command=self._del_col)
        sm.add_separator()
        sm.add_command(label="🗑 Очистити таблицю", command=self._clear_table)
        mb2["menu"] = sm

        tk.Button(tb, text="Вставити з буфера",
                  font=("Times New Roman", 11),
                  command=self._paste).pack(side=tk.LEFT, padx=4)
        tk.Button(tb, text="📚 Довідка", bg="#1a4b8c", fg="white",
                  font=("Times New Roman", 11),
                  command=self._show_help).pack(side=tk.LEFT, padx=4)
        tk.Label(tb,
                 text="Двічі клікніть синій заголовок щоб перейменувати показник",
                 font=("Times New Roman", 9), fg="#666").pack(side=tk.LEFT, padx=10)

        # ── Таблиця ──
        tf = tk.Frame(self.win); tf.pack(fill=tk.BOTH, expand=True, padx=6, pady=4)
        self.rows = 20; self.cols = 8
        canvas = tk.Canvas(tf); canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb = ttk.Scrollbar(tf, orient="vertical", command=canvas.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y); canvas.configure(yscrollcommand=sb.set)
        self.inner = tk.Frame(canvas)
        canvas.create_window((0, 0), window=self.inner, anchor="nw")
        self.inner.bind("<Configure>", lambda e: canvas.config(scrollregion=canvas.bbox("all")))
        self.win.bind("<MouseWheel>", lambda e: canvas.yview_scroll(int(-1*(e.delta/120)),"units"))
        self._canvas = canvas

        # Заголовки (сині, з перейменуванням)
        self.header_vars   = []
        self.header_labels = []
        for j in range(self.cols):
            var = tk.StringVar(value=f"Показник {j+1}")
            self.header_vars.append(var)
            lbl = tk.Label(self.inner, textvariable=var,
                           relief=tk.RIDGE, width=13, cursor="hand2",
                           bg="#1a4b8c", fg="white",
                           font=("Times New Roman", 11, "bold"))
            lbl.grid(row=0, column=j, padx=2, pady=2, sticky="nsew")
            lbl.bind("<Double-Button-1>", lambda e, idx=j: self._rename_col(idx))
            self.header_labels.append(lbl)

        self.entries = []
        for i in range(self.rows):
            row_ = []
            for j in range(self.cols):
                e = tk.Entry(self.inner, width=13, font=("Times New Roman", 11),
                             highlightthickness=1, highlightbackground="#c0c0c0")
                e.grid(row=i+1, column=j, padx=2, pady=2)
                row_.append(e)
            self.entries.append(row_)
        _bind_nav(self.entries, self.win)

    # ── Перейменування стовпця ───────────────────────────────
    def _rename_col(self, idx):
        dlg = tk.Toplevel(self.win); dlg.title("Перейменувати показник")
        dlg.resizable(False, False); dlg.grab_set()
        tk.Label(dlg, text=f"Назва показника {idx+1}:",
                 font=("Times New Roman", 12)).pack(padx=16, pady=14)
        var = tk.StringVar(value=self.header_vars[idx].get())
        e = tk.Entry(dlg, textvariable=var, font=("Times New Roman", 12), width=28)
        e.pack(padx=16, pady=4); e.select_range(0, tk.END); e.focus_set()
        def apply():
            nm = var.get().strip()
            if nm: self.header_vars[idx].set(nm)
            dlg.destroy()
        tk.Button(dlg, text="ОК", bg="#c62828", fg="white",
                  font=("Times New Roman", 12), command=apply).pack(pady=(4, 14))
        dlg.bind("<Return>", lambda ev: apply()); center_win(dlg)

    # ── Управління таблицею ───────────────────────────────────
    def _add_row(self):
        i = len(self.entries); row_ = []
        for j in range(self.cols):
            e = tk.Entry(self.inner, width=13, font=("Times New Roman", 11),
                         highlightthickness=1, highlightbackground="#c0c0c0")
            e.grid(row=i+1, column=j, padx=2, pady=2); row_.append(e)
        self.entries.append(row_); self.rows += 1
        _bind_nav(self.entries, self.win)
        self._canvas.config(scrollregion=self._canvas.bbox("all"))

    def _del_row(self):
        if not self.entries: return
        for e in self.entries.pop(): e.destroy()
        self.rows -= 1
        self._canvas.config(scrollregion=self._canvas.bbox("all"))

    def _add_col(self):
        ci = self.cols; self.cols += 1
        var = tk.StringVar(value=f"Показник {ci+1}")
        self.header_vars.append(var)
        lbl = tk.Label(self.inner, textvariable=var, relief=tk.RIDGE, width=13,
                       cursor="hand2", bg="#1a4b8c", fg="white",
                       font=("Times New Roman", 11, "bold"))
        lbl.grid(row=0, column=ci, padx=2, pady=2, sticky="nsew")
        lbl.bind("<Double-Button-1>", lambda e, idx=ci: self._rename_col(idx))
        self.header_labels.append(lbl)
        for i, row_ in enumerate(self.entries):
            e = tk.Entry(self.inner, width=13, font=("Times New Roman", 11),
                         highlightthickness=1, highlightbackground="#c0c0c0")
            e.grid(row=i+1, column=ci, padx=2, pady=2); row_.append(e)
        _bind_nav(self.entries, self.win)

    def _del_col(self):
        if self.cols <= 1: return
        self.header_labels.pop().destroy(); self.header_vars.pop()
        for row_ in self.entries: row_.pop().destroy()
        self.cols -= 1

    def _clear_table(self):
        if not messagebox.askyesno("Очистити", "Видалити всі числові дані?\n(Назви стовпців залишаться)"): return
        for row in self.entries:
            for e in row: e.delete(0, tk.END)

    # ── Вставка і завантаження ───────────────────────────────
    def _paste(self):
        """Вставити з буфера обміну. Починає з активної клітинки або (0,0)."""
        try: data = self.win.clipboard_get()
        except Exception:
            messagebox.showwarning("Буфер порожній",
                "Буфер обміну порожній або не містить тексту.\n"
                "Скопіюйте дані з Excel (Ctrl+C) і спробуйте знову."); return
        if not data.strip(): return
        w = self.win.focus_get()
        # Знаходимо позицію активної клітинки
        pos = (0, 0)
        if isinstance(w, tk.Entry):
            for i, row_ in enumerate(self.entries):
                for j, e in enumerate(row_):
                    if e is w: pos = (i, j); break
        try: data = self.win.clipboard_get()
        except Exception: return
        r0, c0 = pos
        for ir, line in enumerate(data.splitlines()):
            for jc, val in enumerate(line.split("\t")):
                rr = r0+ir; cc = c0+jc
                while rr >= len(self.entries): self._add_row()
                if cc >= self.cols: continue
                self.entries[rr][cc].delete(0, tk.END)
                self.entries[rr][cc].insert(0, val.strip())

    def _load_excel(self):
        if not HAS_OPENPYXL: messagebox.showerror("","pip install openpyxl"); return
        path = filedialog.askopenfilename(filetypes=[("Excel","*.xlsx *.xlsm"),("All","*.*")])
        if not path: return
        try:
            wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
            raw = list(wb.active.iter_rows(values_only=True)); wb.close()
        except Exception as ex: messagebox.showerror("", str(ex)); return
        if not raw: return
        nc = max(len(r) for r in raw)
        while self.cols < nc: self._add_col()
        while len(self.entries) < len(raw): self._add_row()
        for i, row in enumerate(raw):
            for j, v in enumerate(row):
                if j >= self.cols: break
                self.entries[i][j].delete(0, tk.END)
                self.entries[i][j].insert(0, "" if v is None else str(v).replace(",","."))

    # ── Довідка ──────────────────────────────────────────────
    def _show_help(self):
        win = tk.Toplevel(self.win)
        win.title("Довідка — Описова статистика")
        win.geometry("700x640"); set_icon(win)
        frm = tk.Frame(win); frm.pack(fill=tk.BOTH, expand=True, padx=8, pady=6)
        vsb = ttk.Scrollbar(frm, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        txt = tk.Text(frm, wrap="word", font=("Times New Roman", 11),
                      yscrollcommand=vsb.set, relief=tk.FLAT,
                      bg="#fafafa", padx=10, pady=8, cursor="arrow")
        txt.pack(fill=tk.BOTH, expand=True)
        vsb.config(command=txt.yview)
        txt.insert("1.0", self.HELP_TEXT.strip())
        txt.configure(state="disabled")
        txt.bind("<MouseWheel>", lambda e: txt.yview_scroll(int(-1*(e.delta/120)), "units"))
        tk.Button(win, text="Закрити", command=win.destroy,
                  font=("Times New Roman", 11)).pack(pady=6)

    # ── Аналіз ───────────────────────────────────────────────
    def _analyze(self):
        from scipy.stats import skew, kurtosis
        # Назви беремо з заголовків; дані — з клітинок
        names = []; data_cols = []
        for j in range(self.cols):
            col_name = self.header_vars[j].get().strip() or f"Показник {j+1}"
            col_vals = []
            for row in self.entries:
                v = row[j].get().strip() if j < len(row) else ""
                if not v: continue
                try: col_vals.append(float(v.replace(",",".")))
                except Exception: continue
            if len(col_vals) >= 2:
                names.append(col_name)
                data_cols.append(np.array(col_vals, dtype=float))

        if not data_cols:
            messagebox.showwarning("Замало даних",
                "Не знайдено числових даних.\n"
                "Переконайтесь що введено числа і кожен стовпець містить ≥ 2 значення."); return

        headers = ["Показник","n","Середнє","SD","СП","Мін","Макс","Медіана",
                   "Q1","Q3","CV%","Асиметрія","Ексцес","95% ДІ нижній","95% ДІ верхній","SW p"]
        rows = []
        for nm, arr in zip(names, data_cols):
            a = arr[~np.isnan(arr)]; n = len(a)
            if n < 2: rows.append([nm, n] + ["–"]*14); continue
            m  = float(np.mean(a)); sd = float(np.std(a, ddof=1))
            se = sd / math.sqrt(n)
            ci_lo = m - float(t_dist.ppf(0.975, n-1)) * se
            ci_hi = m + float(t_dist.ppf(0.975, n-1)) * se
            sk = float(skew(a)); ku = float(kurtosis(a))
            q1  = float(np.percentile(a, 25)); q3 = float(np.percentile(a, 75))
            cv  = sd/m*100 if m != 0 else np.nan
            try: _, sw_p = shapiro(a)
            except Exception: sw_p = np.nan
            rows.append([nm, n, fmt(m,3), fmt(sd,3), fmt(se,3),
                         fmt(float(np.min(a)),3), fmt(float(np.max(a)),3),
                         fmt(float(np.median(a)),3), fmt(q1,3), fmt(q3,3),
                         fmt(cv,2), fmt(sk,3), fmt(ku,3),
                         fmt(ci_lo,3), fmt(ci_hi,3), fmt(sw_p,4)])

        self._show_result(headers, rows, data_cols, names)

    # ── Результати ───────────────────────────────────────────
    def _show_result(self, headers, rows, arrays, names):
        win = tk.Toplevel(self.win)
        win.title("Описова статистика — Результати")
        win.geometry("1340x600"); set_icon(win)

        tb = tk.Frame(win, padx=6, pady=5); tb.pack(fill=tk.X)
        tk.Button(tb, text="📋 Копіювати таблицю",
                  font=("Times New Roman", 11),
                  command=lambda: self._copy_table(win, headers, rows)
                  ).pack(side=tk.LEFT, padx=4)
        tk.Button(tb, text="📊 Боксплоти",
                  font=("Times New Roman", 11),
                  command=lambda: self._plot_boxes(arrays, names)
                  ).pack(side=tk.LEFT, padx=4)
        tk.Button(tb, text="📈 QQ-графіки",
                  font=("Times New Roman", 11),
                  command=lambda: self._plot_qq(arrays, names)
                  ).pack(side=tk.LEFT, padx=4)

        frm, _ = make_tv(win, headers, rows, min_col=80)
        frm.pack(fill=tk.BOTH, expand=True, padx=6, pady=6)

    def _copy_table(self, win, headers, rows):
        """Копіює таблицю у буфер обміну у форматі TSV (для вставки у Excel/Word)."""
        lines = ["\t".join(str(h) for h in headers)]
        for row in rows:
            lines.append("\t".join("" if v is None else str(v) for v in row))
        text = "\n".join(lines)
        win.clipboard_clear(); win.clipboard_append(text)
        messagebox.showinfo("Скопійовано",
            "Таблицю скопійовано у буфер обміну.\n"
            "Вставте у Word або Excel через Ctrl+V.")

    # ── Боксплот ─────────────────────────────────────────────
    def _plot_boxes(self, arrays, names):
        if not HAS_MPL: return
        win = tk.Toplevel(self.win)
        win.title("Боксплоти"); win.geometry("920x580"); set_icon(win)

        tb = tk.Frame(win, padx=6, pady=5); tb.pack(fill=tk.X)
        tk.Button(tb, text="📋 Копіювати PNG",
                  font=("Times New Roman", 11),
                  command=lambda: self._copy_fig(self._bp_fig)
                  ).pack(side=tk.LEFT, padx=4)
        tk.Button(tb, text="⚙ Налаштування",
                  font=("Times New Roman", 11),
                  command=lambda: self._restyle_bp(win, arrays, names)
                  ).pack(side=tk.LEFT, padx=4)

        self._bp_frame = tk.Frame(win); self._bp_frame.pack(fill=tk.BOTH, expand=True)
        self._bp_arrays = arrays; self._bp_names = names
        self._draw_boxes(self._bp_frame, arrays, names)

    def _draw_boxes(self, frame, arrays, names):
        for w in frame.winfo_children(): w.destroy()
        gs = self._bp_gs
        ff  = gs.get("font_family", "Times New Roman")
        fz  = gs.get("font_size", 11)
        n   = len(arrays)
        fig = Figure(figsize=(10, 6), dpi=100)
        ax  = fig.add_subplot(111)
        bp  = ax.boxplot([a[~np.isnan(a)] for a in arrays],
                         labels=names, patch_artist=True, widths=0.55)
        for patch in bp["boxes"]:    patch.set(facecolor=gs.get("box_color","#ffffff"))
        for line  in bp["medians"]:  line.set(color=gs.get("median_color","#c62828"), linewidth=2)
        for line  in bp["whiskers"]+bp["caps"]:
            line.set(color=gs.get("whisker_color","#000000"))
        for fl    in bp["fliers"]:
            fl.set(markerfacecolor=gs.get("flier_color","#555555"), marker="o", markersize=4)
        ax.set_ylabel("Значення", fontsize=fz, fontfamily=ff)
        ax.set_title("Boxplot показників", fontsize=fz+1, fontfamily=ff)
        ax.tick_params(axis="x", labelsize=max(8, fz-1))
        ax.yaxis.grid(True, linestyle="--", alpha=0.4)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        fig.tight_layout()
        self._bp_fig = fig
        embed_figure(fig, frame)

    def _restyle_bp(self, win, arrays, names):
        """Dedicated boxplot settings dialog — no KeyError on missing DEF_GS keys."""
        dlg = tk.Toplevel(win); dlg.title("Налаштування боксплоту")
        dlg.resizable(False, False); set_icon(dlg); dlg.grab_set()
        gs = self._bp_gs
        frm = tk.Frame(dlg, padx=16, pady=14); frm.pack()
        rb_f = ("Times New Roman", 12)

        ff_var = tk.StringVar(value=gs.get("font_family","Times New Roman"))
        fz_var = tk.IntVar(value=gs.get("font_size", 11))
        col_box = [gs.get("box_color",    "#ffffff")]
        col_med = [gs.get("median_color", "#c62828")]
        col_wh  = [gs.get("whisker_color","#000000")]
        col_fl  = [gs.get("flier_color",  "#555555")]

        tk.Label(frm, text="Шрифт:", font=rb_f).grid(row=0, column=0, sticky="w", pady=5)
        ttk.Combobox(frm, textvariable=ff_var,
                     values=["Times New Roman","Arial","Calibri","Georgia","Verdana"],
                     state="readonly", width=22).grid(row=0, column=1, sticky="w", padx=8)
        tk.Label(frm, text="Розмір шрифту:", font=rb_f).grid(row=1, column=0, sticky="w", pady=5)
        tk.Spinbox(frm, from_=7, to=24, textvariable=fz_var, width=6).grid(row=1, column=1, sticky="w", padx=8)

        btn_refs = {}
        color_cfg = [
            ("Колір коробки:", col_box, "box"),
            ("Колір медіани:", col_med, "med"),
            ("Колір вусів:",   col_wh,  "wh"),
            ("Колір викидів:", col_fl,  "fl"),
        ]
        for ri, (lbl, col_lst, key) in enumerate(color_cfg):
            tk.Label(frm, text=lbl, font=rb_f).grid(row=2+ri, column=0, sticky="w", pady=5)
            btn = tk.Button(frm, width=6, relief=tk.SUNKEN, bg=col_lst[0])
            btn.grid(row=2+ri, column=1, sticky="w", padx=8)
            btn_refs[key] = (btn, col_lst)
            def _pick(c=col_lst, b=btn):
                ch = colorchooser.askcolor(color=c[0], parent=dlg, title="Виберіть колір")
                if ch and ch[1]: c[0] = ch[1]; b.configure(bg=ch[1])
            btn.configure(command=_pick)

        def apply():
            self._bp_gs.update({
                "font_family":   ff_var.get(),
                "font_size":     fz_var.get(),
                "box_color":     col_box[0],
                "median_color":  col_med[0],
                "whisker_color": col_wh[0],
                "flier_color":   col_fl[0],
            })
            self._draw_boxes(self._bp_frame, arrays, names)
            dlg.destroy()

        bf = tk.Frame(frm); bf.grid(row=6, column=0, columnspan=2, pady=(14,0))
        tk.Button(bf, text="OK", bg="#c62828", fg="white",
                  font=rb_f, command=apply).pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Скасувати", font=rb_f, command=dlg.destroy).pack(side=tk.LEFT)
        center_win(dlg)

    # ── QQ-графіки ───────────────────────────────────────────
    def _plot_qq(self, arrays, names):
        if not HAS_MPL: return
        from scipy.stats import probplot
        n = len(arrays); cols_ = min(n, 4); rows_n = math.ceil(n / cols_)
        win = tk.Toplevel(self.win)
        win.title("QQ-графіки"); win.geometry("1000x640"); set_icon(win)

        tb = tk.Frame(win, padx=6, pady=5); tb.pack(fill=tk.X)
        tk.Button(tb, text="📋 Копіювати PNG",
                  font=("Times New Roman", 11),
                  command=lambda: self._copy_fig(self._qq_fig)
                  ).pack(side=tk.LEFT, padx=4)
        tk.Button(tb, text="⚙ Налаштування",
                  font=("Times New Roman", 11),
                  command=lambda: self._restyle_qq(win, arrays, names)
                  ).pack(side=tk.LEFT, padx=4)
        tk.Label(tb,
                 text="Точки на прямій → нормальний розподіл ✓   |   "
                      "Відхилення від прямої → ненормальний ⚠",
                 font=("Times New Roman", 9), fg="#555").pack(side=tk.LEFT, padx=8)

        self._qq_frame   = tk.Frame(win); self._qq_frame.pack(fill=tk.BOTH, expand=True)
        self._qq_arrays  = arrays; self._qq_names = names
        self._draw_qq(self._qq_frame, arrays, names)

    def _draw_qq(self, frame, arrays, names):
        from scipy.stats import probplot
        for w in frame.winfo_children(): w.destroy()
        gs    = self._qq_gs
        ff    = gs.get("font_family", "Times New Roman")
        fz    = gs.get("font_size", 9)
        pt_c  = gs.get("pt_color",   "#4c72b0")
        ln_c  = gs.get("line_color", "#c62828")
        n     = len(arrays); cols_ = min(n, 4); rows_n = math.ceil(n / cols_)
        fig   = Figure(figsize=(cols_*2.5+0.5, rows_n*2.5+0.5), dpi=100)
        for i, (arr, nm) in enumerate(zip(arrays, names)):
            a = arr[~np.isnan(arr)]
            if len(a) < 3: continue
            ax  = fig.add_subplot(rows_n, cols_, i+1)
            res = probplot(a, dist="norm")
            ax.plot(res[0][0], res[0][1], "o", markersize=4, color=pt_c, alpha=0.8)
            ax.plot(res[0][0], res[1][1] + res[1][0]*res[0][0], "-", color=ln_c, lw=1.5)
            ax.set_title(nm, fontsize=fz+1, fontfamily=ff)
            ax.set_xlabel("Теоретичні квантилі", fontsize=fz, fontfamily=ff)
            ax.set_ylabel("Вибіркові квантилі",  fontsize=fz, fontfamily=ff)
            ax.yaxis.grid(True, linestyle="--", alpha=0.35)
            ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        fig.tight_layout()
        self._qq_fig = fig
        embed_figure(fig, frame)

    def _restyle_qq(self, win, arrays, names):
        """Простий діалог налаштувань QQ-графіків."""
        dlg = tk.Toplevel(win); dlg.title("Налаштування QQ-графіків")
        dlg.resizable(False, False); set_icon(dlg); dlg.grab_set()
        frm = tk.Frame(dlg, padx=16, pady=14); frm.pack()
        rb_f = ("Times New Roman", 12)

        ff_var  = tk.StringVar(value=self._qq_gs.get("font_family","Times New Roman"))
        fz_var  = tk.IntVar(value=self._qq_gs.get("font_size", 9))
        pt_col  = [self._qq_gs.get("pt_color","#4c72b0")]
        ln_col  = [self._qq_gs.get("line_color","#c62828")]

        tk.Label(frm, text="Шрифт:", font=rb_f).grid(row=0, column=0, sticky="w", pady=4)
        ttk.Combobox(frm, textvariable=ff_var,
                     values=["Times New Roman","Arial","Calibri","Georgia"],
                     state="readonly", width=20).grid(row=0, column=1, sticky="w", padx=8)
        tk.Label(frm, text="Розмір шрифту:", font=rb_f).grid(row=1, column=0, sticky="w", pady=4)
        tk.Spinbox(frm, from_=6, to=18, textvariable=fz_var, width=6).grid(row=1, column=1, sticky="w", padx=8)

        pt_btn = tk.Button(frm, width=6, relief=tk.SUNKEN, bg=pt_col[0])
        ln_btn = tk.Button(frm, width=6, relief=tk.SUNKEN, bg=ln_col[0])

        def pick_pt():
            c = colorchooser.askcolor(color=pt_col[0], parent=dlg, title="Колір точок")
            if c and c[1]: pt_col[0]=c[1]; pt_btn.configure(bg=c[1])
        def pick_ln():
            c = colorchooser.askcolor(color=ln_col[0], parent=dlg, title="Колір лінії")
            if c and c[1]: ln_col[0]=c[1]; ln_btn.configure(bg=c[1])

        tk.Label(frm, text="Колір точок:", font=rb_f).grid(row=2, column=0, sticky="w", pady=4)
        pt_btn.configure(command=pick_pt); pt_btn.grid(row=2, column=1, sticky="w", padx=8)
        tk.Label(frm, text="Колір лінії:", font=rb_f).grid(row=3, column=0, sticky="w", pady=4)
        ln_btn.configure(command=pick_ln); ln_btn.grid(row=3, column=1, sticky="w", padx=8)

        def apply():
            self._qq_gs.update({"font_family": ff_var.get(), "font_size": fz_var.get(),
                                 "pt_color": pt_col[0], "line_color": ln_col[0]})
            self._draw_qq(self._qq_frame, arrays, names); dlg.destroy()
        bf = tk.Frame(frm); bf.grid(row=4, column=0, columnspan=2, pady=(14,0))
        tk.Button(bf, text="OK", bg="#c62828", fg="white",
                  font=rb_f, command=apply).pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Скасувати", font=rb_f, command=dlg.destroy).pack(side=tk.LEFT)
        center_win(dlg)

    # ── Копіювання PNG ───────────────────────────────────────
    def _copy_fig(self, fig):
        if fig is None:
            messagebox.showwarning("", "Спочатку побудуйте графік."); return
        ok, msg = _copy_fig_to_clipboard(fig)
        if ok: messagebox.showinfo("", "Графік скопійовано.\nВставте у Word через Ctrl+V.")
        else:   messagebox.showwarning("", f"Помилка: {msg}")




# ═══════════════════════════════════════════════════════════════
# T-TEST MODULE
# ═══════════════════════════════════════════════════════════════
class TTestWindow:
    """t-тест / Критерій Манна-Уітні."""

    HELP_TEXT = """
t-ТЕСТ / МАНН-УІТНІ — ПОКРОКОВА ІНСТРУКЦІЯ
════════════════════════════════════════════

ЩО ЦЕЙ АНАЛІЗ РОБИТЬ?
  Порівнює ДВІ групи і відповідає:
  «Чи є різниця між середніми статистично значущою,
  чи це просто випадкові коливання?»

  Програма автоматично обирає правильний тест:
  ✓ Нормальний розподіл + рівні дисперсії → t-тест Стьюдента
  ✓ Нормальний розподіл + нерівні дисперсії → t-тест Велша
  ✓ Ненормальний розподіл → Манн-Уітні (непараметричний)
  ✓ Парні + нормальний → Парний t-тест
  ✓ Парні + ненормальний → Вілкоксон

РЕЖИМ 1: НЕЗАЛЕЖНІ ВИБІРКИ
  Дві різні групи. Спостереження не пов'язані.
  Приклад: врожайність Сорту А і Сорту Б.
  Введіть значення у поля «Група 1» і «Група 2».

РЕЖИМ 2: ПАРНІ ВИБІРКИ (до/після)
  Ті самі об'єкти вимірюються двічі.
  Приклад: маса рослин до і після обробки.
  ВАЖЛИВО: порядок значень має бути однаковим!
  Перше значення Групи 1 пов'язане з першим Групи 2.
  Кількість значень в обох групах має бути однаковою.

РЕЖИМ 3: ОДНА ВИБІРКА (проти відомого μ)
  Порівняння середнього вибірки з відомим значенням.
  Приклад: чи відрізняється врожайність від нормативу 5 т/га?
  Введіть дані у «Група 1» і вкажіть μ₀.

ЯК ВВОДИТИ ДАНІ:
  Значення через кому, пробіл або кожне з нового рядка.
  Приклад: 4.2, 5.1, 4.8, 5.3, 4.9
  Або:
  4.2
  5.1
  4.8

ІНТЕРПРЕТАЦІЯ РЕЗУЛЬТАТІВ:
  p < α → різниця значуща ✓ (реальна різниця між групами)
  p ≥ α → різниця незначуща ✗ (можлива випадковість)

  РОЗМІР ЕФЕКТУ (Cliff's delta для Манн-Уітні):
  |δ| < 0.15: дуже слабкий
  0.15-0.33:  слабкий
  0.33-0.47:  середній
  > 0.47:     сильний

  Значущий p ≠ велика різниця!
  При великих n навіть мізерна різниця буде значущою.
  Завжди оцінюйте розмір ефекту разом з p.

SHAPIRO-WILK:
  p > 0.05 → нормальний розподіл (параметричний тест)
  p ≤ 0.05 → ненормальний (непараметричний тест)
"""

    def __init__(self, parent):
        self.win = tk.Toplevel(parent)
        self.win.title("t-тест / Критерій Манна-Уітні")
        self.win.geometry("700x660"); set_icon(self.win)
        self._build()

    def _build(self):
        # ── Toolbar ──────────────────────────────────────────
        top = tk.Frame(self.win, padx=8, pady=6); top.pack(fill=tk.X)
        tk.Button(top, text="▶ Виконати", bg="#c62828", fg="white",
                  font=("Times New Roman",13),
                  command=self._run).pack(side=tk.LEFT, padx=4)
        tk.Label(top, text="α:", font=("Times New Roman",12)).pack(side=tk.LEFT, padx=(10,2))
        self.alpha_var = tk.StringVar(value="0.05")
        ttk.Combobox(top, textvariable=self.alpha_var, values=["0.01","0.05","0.10"],
                     state="readonly", width=7).pack(side=tk.LEFT)
        tk.Button(top, text="Вставити з буфера",
                  font=("Times New Roman",11),
                  command=self._paste).pack(side=tk.LEFT, padx=8)
        tk.Button(top, text="📋 Копіювати результат",
                  font=("Times New Roman",11),
                  command=self._copy_result).pack(side=tk.LEFT, padx=4)
        tk.Button(top, text="📚 Довідка", bg="#1a4b8c", fg="white",
                  font=("Times New Roman",11),
                  command=self._show_help).pack(side=tk.LEFT, padx=8)

        # ── Тип тесту ────────────────────────────────────────
        frm = tk.Frame(self.win, padx=12); frm.pack(fill=tk.BOTH, expand=True)
        tk.Label(frm, text="Тип тесту:",
                 font=("Times New Roman",12,"bold")).grid(row=0, column=0, sticky="w", pady=4)
        self.test_var = tk.StringVar(value="ind")
        rf = ("Times New Roman",12)
        tests = [("Незалежні вибірки (2 різні групи)", "ind"),
                 ("Парні вибірки (до/після, однакові об'єкти)", "paired"),
                 ("Одна вибірка (проти відомого μ₀)", "one")]
        for ri, (txt, val) in enumerate(tests):
            tk.Radiobutton(frm, text=txt, variable=self.test_var, value=val,
                           font=rf, command=self._update_ui
                           ).grid(row=ri+1, column=0, columnspan=2, sticky="w")

        # ── Поля введення ─────────────────────────────────────
        tk.Label(frm, text="Група 1 / Вибірка:",
                 font=("Times New Roman",12)).grid(row=4, column=0, sticky="w", pady=10)
        self.e1 = tk.Text(frm, width=55, height=5, font=("Times New Roman",11))
        self.e1.grid(row=5, column=0, columnspan=2, sticky="ew")
        tk.Label(frm, text="Вводьте через кому, пробіл або кожне значення з нового рядка",
                 font=("Times New Roman",9), fg="#666"
                 ).grid(row=6, column=0, columnspan=2, sticky="w")

        self.lbl2 = tk.Label(frm, text="Група 2:", font=("Times New Roman",12))
        self.lbl2.grid(row=7, column=0, sticky="w", pady=8)
        self.e2 = tk.Text(frm, width=55, height=5, font=("Times New Roman",11))
        self.e2.grid(row=8, column=0, columnspan=2, sticky="ew")

        self.lbl_mu = tk.Label(frm, text="Відоме середнє (μ₀):", font=("Times New Roman",12))
        self.e_mu = tk.Entry(frm, width=12, font=("Times New Roman",12))
        self.e_mu.insert(0, "0")

        # ── Результати (scrollable) ───────────────────────────
        res_frm = tk.Frame(frm); res_frm.grid(row=13, column=0, columnspan=2,
                                               sticky="nsew", pady=(8,4))
        frm.rowconfigure(13, weight=1); frm.columnconfigure(0, weight=1)
        vsb = ttk.Scrollbar(res_frm, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        self.res_txt = tk.Text(res_frm, wrap="word",
                               font=("Times New Roman",11),
                               yscrollcommand=vsb.set,
                               relief=tk.FLAT, bg="#f8f8f8",
                               padx=8, pady=6, cursor="arrow",
                               state="disabled", height=8)
        self.res_txt.pack(fill=tk.BOTH, expand=True)
        vsb.config(command=self.res_txt.yview)

        self._update_ui()

    def _update_ui(self):
        t = self.test_var.get()
        if t == "one":
            self.lbl2.grid_remove(); self.e2.grid_remove()
            self.lbl_mu.grid(row=7, column=0, sticky="w", pady=(8,2))
            self.e_mu.grid(row=7, column=1, sticky="w", padx=6)
        else:
            self.lbl_mu.grid_remove(); self.e_mu.grid_remove()
            self.lbl2.grid(row=7, column=0, sticky="w", pady=(8,2))
            self.e2.grid(row=8, column=0, columnspan=2, sticky="ew")
            txt = ("Група 2 (парна — той самий порядок що й Група 1):"
                   if t == "paired" else "Група 2:")
            self.lbl2.configure(text=txt)

    def _parse(self, widget):
        import re
        txt = widget.get("1.0", tk.END).strip().replace(",",".")
        nums = re.findall(r"[-+]?\d*\.?\d+(?:[eE][-+]?\d+)?", txt)
        return np.array([float(x) for x in nums], dtype=float)

    def _set_result(self, text, color="#000000"):
        self.res_txt.configure(state="normal")
        self.res_txt.delete("1.0", tk.END)
        self.res_txt.insert("1.0", text)
        self.res_txt.configure(state="disabled", fg=color)

    def _show_help(self):
        win = tk.Toplevel(self.win)
        win.title("Довідка — t-тест / Манн-Уітні")
        win.geometry("680x640"); set_icon(win)
        frm = tk.Frame(win); frm.pack(fill=tk.BOTH, expand=True, padx=8, pady=6)
        vsb = ttk.Scrollbar(frm, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        txt = tk.Text(frm, wrap="word", font=("Times New Roman",11),
                      yscrollcommand=vsb.set, relief=tk.FLAT,
                      bg="#fafafa", padx=10, pady=8, cursor="arrow")
        txt.pack(fill=tk.BOTH, expand=True); vsb.config(command=txt.yview)
        txt.insert("1.0", self.HELP_TEXT.strip()); txt.configure(state="disabled")
        txt.bind("<MouseWheel>",
                 lambda e: txt.yview_scroll(int(-1*(e.delta/120)),"units"))
        tk.Button(win, text="Закрити", command=win.destroy,
                  font=("Times New Roman",11)).pack(pady=6)

    def _run(self):
        from scipy.stats import ttest_ind, ttest_rel, ttest_1samp
        alpha = float(self.alpha_var.get())
        x1 = self._parse(self.e1)
        t = self.test_var.get()
        p = None  # ініціалізуємо явно
        if len(x1) < 2:
            self._set_result("Група 1 потребує ≥ 2 значень.", "#c62828"); return

        lines = []
        sep = "─" * 46
        lines.append(f"n₁ = {len(x1)}   Середнє₁ = {fmt(np.mean(x1),4)}   SD₁ = {fmt(np.std(x1,ddof=1),4)}")

        try: _, sw1 = shapiro(x1)
        except Exception: sw1 = np.nan
        normal1 = not math.isnan(sw1) and sw1 > 0.05
        lines.append(f"Shapiro–Wilk (Група 1): W = {fmt(sw1,4)}"
                     f"  →  {'✓ нормальний' if normal1 else '⚠ НЕ нормальний'}")

        if t == "one":
            try: mu0 = float(self.e_mu.get())
            except Exception: mu0 = 0.0
            stat, p = ttest_1samp(x1, mu0)
            lines.append(sep)
            lines.append(f"Одновибірковий t-тест (μ₀ = {mu0})")
            lines.append(f"t = {fmt(stat,4)},   df = {len(x1)-1},   p = {fmt(p,4)}")
        else:
            x2 = self._parse(self.e2)
            if len(x2) < 2:
                self._set_result("Група 2 потребує ≥ 2 значень.", "#c62828"); return
            try: _, sw2 = shapiro(x2)
            except Exception: sw2 = np.nan
            normal2 = not math.isnan(sw2) and sw2 > 0.05
            lines.append(f"n₂ = {len(x2)}   Середнє₂ = {fmt(np.mean(x2),4)}   SD₂ = {fmt(np.std(x2,ddof=1),4)}")
            lines.append(f"Shapiro–Wilk (Група 2): W = {fmt(sw2,4)}"
                         f"  →  {'✓ нормальний' if normal2 else '⚠ НЕ нормальний'}")
            lines.append(sep)

            if t == "paired":
                if len(x1) != len(x2):
                    self._set_result(
                        "Парний тест вимагає однакового розміру вибірок.\n"
                        f"Група 1: {len(x1)} значень, Група 2: {len(x2)} значень.",
                        "#c62828"); return
                if normal1 and normal2:
                    stat, p = ttest_rel(x1, x2)
                    lines.append(f"Парний t-тест")
                    lines.append(f"t = {fmt(stat,4)},   df = {len(x1)-1},   p = {fmt(p,4)}")
                else:
                    stat, p = wilcoxon(x1, x2, zero_method="wilcox",
                                       alternative="two-sided", mode="auto")
                    lines.append("Критерій Вілкоксона (знакових рангів)")
                    lines.append(f"W = {fmt(stat,4)},   p = {fmt(p,4)}")
            else:
                try: lev_s, lev_p = levene(x1, x2, center='median')
                except Exception: lev_p = np.nan
                equal_var = not math.isnan(lev_p) and lev_p >= 0.05
                lines.append(f"Тест Левена: p = {fmt(lev_p,4)}"
                             f"  →  {'✓ рівні дисперсії' if equal_var else '⚠ нерівні дисперсії'}")

                if normal1 and normal2:
                    stat, p = ttest_ind(x1, x2, equal_var=equal_var)
                    n1, n2 = len(x1), len(x2)
                    if not equal_var:
                        df_w = ((np.var(x1,ddof=1)/n1 + np.var(x2,ddof=1)/n2)**2 /
                                ((np.var(x1,ddof=1)/n1)**2/(n1-1) +
                                 (np.var(x2,ddof=1)/n2)**2/(n2-1)))
                        test_name = "t-тест Велша (нерівні дисперсії)"
                    else:
                        df_w = n1+n2-2
                        test_name = "t-тест Стьюдента (незалежні)"
                    lines.append(f"{test_name}")
                    lines.append(f"t = {fmt(stat,4)},   df ≈ {fmt(df_w,1)},   p = {fmt(p,4)}")
                else:
                    U, p = mannwhitneyu(x1, x2, alternative="two-sided")
                    d = cliffs_d(x1, x2)
                    lines.append("Критерій Манна-Уітні (непараметричний)")
                    lines.append(f"U = {fmt(U,3)},   p = {fmt(p,4)}")
                    lines.append(f"Cliff's δ = {fmt(d,4)}   ({cliffs_lbl(abs(d))} ефект)")

        lines.append(sep)
        if p is not None and not math.isnan(p):
            sig = p < alpha
            lines.append(
                f"{'✓ Різниця ЗНАЧУЩА' if sig else '✗ Різниця НЕЗНАЧУЩА'}"
                f"   (p = {fmt(p,4)},  α = {alpha})")
            if sig and t != "one":
                x2_arr = self._parse(self.e2)
                diff = float(np.mean(x1) - np.mean(x2_arr)) if len(x2_arr) > 0 else float("nan")
                lines.append(f"Різниця середніх: {fmt(diff,4)}")
            elif sig and t == "one":
                try: mu0v = float(self.e_mu.get())
                except Exception: mu0v = 0.0
                lines.append(f"Різниця від μ₀: {fmt(float(np.mean(x1))-mu0v,4)}")

        self._set_result("\n".join(lines))

    def _paste(self):
        """Вставити дані з буфера у Групу 1 або Групу 2 залежно від фокусу."""
        try: data = self.win.clipboard_get()
        except Exception:
            messagebox.showwarning("Буфер порожній",
                "Скопіюйте дані з Excel і спробуйте знову."); return
        if not data.strip(): return
        # Нормалізуємо — замінюємо табуляції і пробіли на нові рядки
        data = data.replace("\t","\n").replace(",",".")
        w = self.win.focus_get()
        target = self.e2 if w is self.e2 else self.e1
        target.delete("1.0", tk.END)
        target.insert("1.0", data.strip())

    def _copy_result(self):
        text = self.res_txt.get("1.0", tk.END).strip()
        if not text:
            messagebox.showwarning("","Спочатку виконайте аналіз."); return
        self.win.clipboard_clear()
        self.win.clipboard_append(text)
        messagebox.showinfo("Скопійовано","Результат скопійовано у буфер обміну.")






# ═══════════════════════════════════════════════════════════════
# OUTLIER DETECTION
# ═══════════════════════════════════════════════════════════════
def detect_outliers_grubbs(arr, alpha=0.05):
    """Grubbs test for single outlier. Returns (idx_of_outlier or None, G, p)."""
    a = np.array(arr, dtype=float); n = len(a)
    if n < 3: return None, np.nan, np.nan
    m = np.mean(a); s = np.std(a, ddof=1)
    if s == 0: return None, np.nan, np.nan
    G = np.max(np.abs(a - m)) / s
    idx = int(np.argmax(np.abs(a - m)))
    # critical value via t-distribution
    t_crit = float(t_dist.ppf(1 - alpha/(2*n), n-2))
    G_crit = ((n-1)/math.sqrt(n)) * math.sqrt(t_crit**2 / (n-2+t_crit**2))
    p_approx = 2 * n * (1 - float(t_dist.cdf(G * math.sqrt(n) * math.sqrt(n-2) /
               math.sqrt(n-1+G**2*n/(n-1)), n-2))) if n > 2 else np.nan
    return (idx if G > G_crit else None), float(G), float(p_approx)

def detect_outliers_iqr(arr):
    """IQR method. Returns list of indices."""
    a = np.array(arr, dtype=float)
    q1, q3 = np.percentile(a, 25), np.percentile(a, 75)
    iqr = q3 - q1
    lo, hi = q1 - 1.5*iqr, q3 + 1.5*iqr
    return [i for i, v in enumerate(a) if v < lo or v > hi]


# ═══════════════════════════════════════════════════════════════
# REGRESSION MODULE
# ═══════════════════════════════════════════════════════════════
class RegressionWindow:
    MODELS = ["Лінійна:  y = a + bx",
              "Квадратична:  y = a + bx + cx²",
              "Кубічна:  y = a + bx + cx² + dx³",
              "Степенева:  y = a·xᵇ",
              "Експоненційна:  y = a·eᵇˣ",
              "Логарифмічна:  y = a + b·ln(x)",
              "Логістична (4-пар.):  y = d + (a-d)/(1+(x/c)ᵇ)"]

    HELP_TEXT = """
РЕГРЕСІЙНИЙ АНАЛІЗ — ПОКРОКОВА ІНСТРУКЦІЯ
══════════════════════════════════════════

КРОК 1. ВВЕДЕННЯ ДАНИХ
  • Ліве поле: значення незалежної змінної x
    (фактор, що ви змінюєте — доза добрива, час, температура тощо)
  • Праве поле: значення залежної змінної y
    (показник, що вимірюєте — врожайність, висота, маса тощо)
  • Вводьте по одному значенню на рядок або через кому
  • Або натисніть «Вставити дані» — два стовпці з Excel (x | y)
  • Мінімум: 4 пари значень

КРОК 2. ВИБІР МОДЕЛІ
  Лінійна (y = a + bx):
    Коли залежність пряма — з ростом x, y рівномірно росте або спадає.
    Найпоширеніша. Починайте з неї.

  Квадратична (y = a + bx + cx²):
    Коли є оптимум — крива з одним піком або западиною.
    Типово для доз добрив, щільності посіву — є оптимальна доза.

  Кубічна (y = a + bx + cx² + dx³):
    Складніша крива з S-подібним характером.
    Використовуйте якщо квадратична дає погану підгонку.

  Степенева (y = a·xᵇ):
    Вимагає x > 0. Для алометричних залежностей (маса-розмір).

  Експоненційна (y = a·eᵇˣ):
    Для процесів росту (b > 0) або спаду (b < 0).

  Логарифмічна (y = a + b·ln(x)):
    Вимагає x > 0. Коли ефект насичується з ростом x.

  Логістична 4-параметрична:
    S-подібна крива. Для доза-відповідь, росту популяцій.
    Параметри: a = верхня асимптота, d = нижня,
               c = точка перегину, b = крутизна.

КРОК 3. ВИКОНАННЯ АНАЛІЗУ
  Натисніть «▶ Виконати» і переглядайте результати.

КРОК 4. ІНТЕРПРЕТАЦІЯ РЕЗУЛЬТАТІВ

  Рівняння регресії:
    Математична формула залежності y від x.
    Підставте будь-яке значення x щоб отримати прогноз y.

  R² (коефіцієнт детермінації):
    Від 0 до 1. Показує яку частку варіації y пояснює модель.
    R² = 0.85 → модель пояснює 85% мінливості y.
    R² > 0.90 → відмінна підгонка для агрономічних даних.
    R² < 0.50 → модель слабка, шукайте інші фактори.

  R²adj (скоригований R²):
    Враховує кількість параметрів моделі.
    При порівнянні моделей з різною кількістю параметрів
    орієнтуйтесь саме на R²adj, а не на R².
    Якщо R²adj < R² — модель надто складна для ваших даних.

  RMSE (середньоквадратична похибка):
    В одиницях вимірювання y.
    Середнє відхилення прогнозу від факту.
    Менше RMSE → точніший прогноз.
    Наприклад: RMSE = 0.3 т/га означає що модель
    помиляється в середньому на ±0.3 т/га.

  F-тест (значущість моделі):
    p < 0.05 → модель значуща, залежність існує ✓
    p ≥ 0.05 → модель незначуща (можливо замало даних
               або залежності взагалі немає) ✗

  Shapiro–Wilk залишків:
    Перевіряє нормальність відхилень від моделі.
    p > 0.05 → залишки нормальні → модель коректна ✓
    p ≤ 0.05 → залишки ненормальні → перевірте наявність
               викидів або спробуйте іншу модель.

КРОК 5. ОЦІНКА ГРАФІКІВ

  Графік «Точкові дані + Крива регресії»:
    Точки — ваші спостереження.
    Червона лінія — підібрана модель.
    Рожева смуга — 95% довірчий інтервал прогнозу.
    Чим ближче точки до лінії → краща підгонка.

  Графік «Залишки vs Підібрані значення»:
    Залишки = різниця між фактом і прогнозом.
    Ідеальний випадок: точки хаотично розкидані
    навколо нуля без жодного патерну.
    ⚠ Якщо є патерн (дуга, воронка) → модель некоректна!

КРОК 6. ПОРІВНЯННЯ КІЛЬКОХ МОДЕЛЕЙ
  Запустіть аналіз послідовно для різних моделей.
  Оберіть ту де:
  1. R²adj найвищий
  2. RMSE найменший
  3. Залишки нормальні (SW p > 0.05)
  4. Немає патерну на графіку залишків
"""

    def __init__(self, parent, gs):
        self.win = tk.Toplevel(parent)
        self.win.title("Регресійний аналіз")
        self.win.geometry("1280x760"); set_icon(self.win)
        self.win.resizable(True, True)
        self.gs = gs
        self._fig = None
        self._graph_title = ""
        self._build()

    def _build(self):
        rf = ("Times New Roman", 11)

        # ── Toolbar ───────────────────────────────────────────
        top = tk.Frame(self.win, padx=8, pady=5, bg="#f5f5f5")
        top.pack(fill=tk.X)
        tk.Frame(top, bg="#e0e0e0", height=1).pack(fill=tk.X, side=tk.BOTTOM)

        tk.Label(top, text="Модель:", font=rf, bg="#f5f5f5").pack(side=tk.LEFT)
        self.model_var = tk.StringVar(value=self.MODELS[0])
        ttk.Combobox(top, textvariable=self.model_var, values=self.MODELS,
                     state="readonly", width=42, font=rf).pack(side=tk.LEFT, padx=6)
        tk.Label(top, text="α:", font=rf, bg="#f5f5f5").pack(side=tk.LEFT, padx=(8,2))
        self.alpha_var = tk.StringVar(value="0.05")
        ttk.Combobox(top, textvariable=self.alpha_var,
                     values=["0.01","0.05","0.10"],
                     state="readonly", width=7).pack(side=tk.LEFT)
        tk.Button(top, text="▶ Виконати", bg="#c62828", fg="white",
                  font=("Times New Roman",13), relief=tk.FLAT, padx=14, pady=3,
                  cursor="hand2", command=self._run).pack(side=tk.LEFT, padx=(10,4))
        tk.Button(top, text="📋 Вставити",
                  font=rf, relief=tk.FLAT, padx=8, pady=3, cursor="hand2",
                  command=self._paste).pack(side=tk.LEFT, padx=2)
        tk.Button(top, text="📚 Довідка", bg="#1a4b8c", fg="white",
                  font=rf, relief=tk.FLAT, padx=8, pady=3, cursor="hand2",
                  command=self._show_help).pack(side=tk.LEFT, padx=4)

        # ── ОСНОВНА ОБЛАСТЬ ───────────────────────────────────
        main = tk.Frame(self.win); main.pack(fill=tk.BOTH, expand=True, padx=6, pady=4)

        # ── ЛІВО: поля x і y ──────────────────────────────────
        left = tk.Frame(main, width=230); left.pack(side=tk.LEFT, fill=tk.Y)
        left.pack_propagate(False)
        hdr_f = tk.Frame(left, bg="#1a4b8c"); hdr_f.pack(fill=tk.X, pady=(0,4))
        tk.Label(hdr_f, text="  Дані", bg="#1a4b8c", fg="white",
                 font=("Times New Roman",11,"bold"), pady=5).pack(side=tk.LEFT)
        cf = tk.Frame(left); cf.pack(fill=tk.BOTH, expand=True)
        for ci, lbl in enumerate(["x  (незалежна)", "y  (залежна)"]):
            tk.Label(cf, text=lbl, font=("Times New Roman",10,"bold"),
                     fg="#1a4b8c").grid(row=0, column=ci, padx=3, pady=2)
        self.tx = tk.Text(cf, width=10, font=("Times New Roman",11),
                          relief=tk.FLAT, highlightthickness=1,
                          highlightbackground="#c0c0c0", highlightcolor="#1a4b8c")
        self.tx.grid(row=1, column=0, padx=3, pady=2, sticky="nsew")
        self.ty = tk.Text(cf, width=10, font=("Times New Roman",11),
                          relief=tk.FLAT, highlightthickness=1,
                          highlightbackground="#c0c0c0", highlightcolor="#1a4b8c")
        self.ty.grid(row=1, column=1, padx=3, pady=2, sticky="nsew")
        cf.rowconfigure(1, weight=1); cf.columnconfigure(0, weight=1); cf.columnconfigure(1, weight=1)
        tk.Label(left, text="Одне значення на рядок\nабо вставте два стовпці з Excel.",
                 font=("Times New Roman",8), fg="#888", justify="left"
                 ).pack(anchor="w", padx=4, pady=2)

        # Роздільник
        tk.Frame(main, bg="#e0e0e0", width=1).pack(side=tk.LEFT, fill=tk.Y, padx=4)

        # ── ПРАВО: результати (прокручувані) ──────────────────
        right_outer = tk.Frame(main); right_outer.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        _vsb = ttk.Scrollbar(right_outer, orient="vertical")
        _vsb.pack(side=tk.RIGHT, fill=tk.Y)
        self._res_cv = tk.Canvas(right_outer, highlightthickness=0, yscrollcommand=_vsb.set)
        self._res_cv.pack(fill=tk.BOTH, expand=True)
        _vsb.config(command=self._res_cv.yview)
        self.res_frame = tk.Frame(self._res_cv)
        _wid = self._res_cv.create_window((0,0), window=self.res_frame, anchor="nw")
        self.res_frame.bind("<Configure>",
            lambda e: self._res_cv.configure(scrollregion=self._res_cv.bbox("all")))
        self._res_cv.bind("<Configure>",
            lambda e: self._res_cv.itemconfig(_wid, width=e.width))
        def _mw(e): self._res_cv.yview_scroll(int(-1*(e.delta/120)),"units")
        self._res_cv.bind("<MouseWheel>", _mw)
        self.res_frame.bind("<MouseWheel>", _mw)

        def _bind_mw_all(w):
            try: w.bind("<MouseWheel>", _mw)
            except Exception: pass
            for ch in w.winfo_children(): _bind_mw_all(ch)

        def _on_res_configure(e):
            self._res_cv.configure(scrollregion=self._res_cv.bbox("all"))
            _bind_mw_all(self.res_frame)
        self.res_frame.bind("<Configure>", _on_res_configure)
        self._res_cv.bind("<Configure>",
            lambda e: self._res_cv.itemconfig(_wid, width=e.width))

        tk.Label(self.res_frame,
                 text="Введіть дані, оберіть модель і натисніть  ▶ Виконати",
                 font=("Times New Roman",12), fg="#aaa").pack(expand=True, pady=40)

    def _graph_settings(self):
        if self._fig is None:
            messagebox.showinfo("","Спочатку виконайте аналіз."); return
        dlg = tk.Toplevel(self.win); dlg.title("Налаштування графіка регресії")
        dlg.resizable(False, False); dlg.grab_set(); set_icon(dlg)
        rf = ("Times New Roman",11)
        frm = tk.Frame(dlg, padx=16, pady=12); frm.pack()

        tk.Label(frm, text="Заголовок графіка:", font=rf
                 ).grid(row=0, column=0, sticky="w", pady=4)
        tv = tk.StringVar(value=self._graph_title)
        tk.Entry(frm, textvariable=tv, width=36, font=rf
                 ).grid(row=0, column=1, sticky="w", padx=8)

        # Кольори
        color_vars = {}
        color_defs = [
            ("Колір точок:",       "scatter_color", "#4c72b0"),
            ("Колір кривої:",      "line_color",    "#c62828"),
            ("Колір ДІ (смуга):",  "ci_color",      "#c62828"),
        ]
        for ri, (lbl, key, default) in enumerate(color_defs, 1):
            tk.Label(frm, text=lbl, font=rf
                     ).grid(row=ri, column=0, sticky="w", pady=3)
            v = tk.StringVar(value=self.gs.get(key, default))
            color_vars[key] = v
            def _pick(var=v):
                c = colorchooser.askcolor(color=var.get(), parent=dlg)
                if c and c[1]: var.set(c[1])
            tk.Button(frm, text="Обрати колір", command=_pick, font=rf
                      ).grid(row=ri, column=1, sticky="w", padx=8)

        def _apply():
            self._graph_title = tv.get().strip()
            for key, v in color_vars.items():
                self.gs[key] = v.get()
            dlg.destroy()
            # Перебудовуємо графік з поточними даними
            if hasattr(self, '_last_run_args'):
                self._show_result(*self._last_run_args)
        bf = tk.Frame(frm); bf.grid(row=len(color_defs)+2, column=0, columnspan=2, pady=(12,0))
        tk.Button(bf, text="Застосувати", bg="#c62828", fg="white",
                  font=rf, command=_apply).pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Скасувати", font=rf,
                  command=dlg.destroy).pack(side=tk.LEFT)
        center_win(dlg)

    def _save_png(self):
        if self._fig is None:
            messagebox.showinfo("","Спочатку виконайте аналіз."); return
        path = filedialog.asksaveasfilename(
            parent=self.win, defaultextension=".png",
            filetypes=[("PNG","*.png"),("SVG","*.svg")],
            title="Зберегти графік регресії")
        if not path: return
        try:
            self._fig.savefig(path, dpi=150, bbox_inches="tight")
            messagebox.showinfo("Збережено", f"Збережено:\n{path}")
        except Exception as ex:
            messagebox.showerror("Помилка", str(ex))

    # ── Утиліти ──────────────────────────────────────────────
    def _paste(self):
        try: data = self.win.clipboard_get()
        except Exception:
            messagebox.showwarning("","Буфер порожній."); return
        lines_ = [l.strip() for l in data.splitlines() if l.strip()]
        if not lines_: return
        xs, ys = [], []
        for line in lines_:
            # Розбиваємо по Tab, потім по пробілу/комі
            parts = line.replace(",", ".").split("\t")
            if len(parts) == 1:
                parts = line.replace(",", ".").split()
            if len(parts) >= 2:
                xs.append(parts[0].strip())
                ys.append(parts[1].strip())
            elif len(parts) == 1:
                xs.append(parts[0].strip())
        if not xs:
            messagebox.showwarning("","Не вдалося розпізнати дані.\n"
                "Скопіюйте два стовпці (x, y) з Excel."); return
        self.tx.delete("1.0", tk.END)
        self.tx.insert("1.0", "\n".join(xs))
        if ys:
            self.ty.delete("1.0", tk.END)
            self.ty.insert("1.0", "\n".join(ys))

    def _parse_col(self, widget):
        import re
        txt = widget.get("1.0", tk.END).replace(",",".")
        return np.array([float(v) for v in
                         re.findall(r"[-+]?\d*\.?\d+(?:[eE][-+]?\d+)?", txt)], dtype=float)

    def _copy_graph(self):
        if self._fig is None:
            messagebox.showwarning("", "Спочатку виконайте аналіз."); return
        ok, msg = _copy_fig_to_clipboard(self._fig)
        if ok: messagebox.showinfo("", "Графік скопійовано (PNG).\nВставте у Word через Ctrl+V.")
        else:   messagebox.showwarning("", f"Помилка копіювання: {msg}")

    def _show_help(self):
        win = tk.Toplevel(self.win)
        win.title("Довідка — Регресійний аналіз")
        win.geometry("680x620"); set_icon(win)
        frm = tk.Frame(win); frm.pack(fill=tk.BOTH, expand=True, padx=8, pady=6)
        vsb = ttk.Scrollbar(frm, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        txt = tk.Text(frm, wrap="word", font=("Times New Roman",11),
                      yscrollcommand=vsb.set, relief=tk.FLAT, bg="#fafafa",
                      padx=10, pady=8, cursor="arrow")
        txt.pack(fill=tk.BOTH, expand=True)
        vsb.config(command=txt.yview)
        txt.insert("1.0", self.HELP_TEXT.strip())
        txt.configure(state="disabled")
        txt.bind("<MouseWheel>", lambda e: txt.yview_scroll(int(-1*(e.delta/120)), "units"))
        tk.Button(win, text="Закрити", command=win.destroy,
                  font=("Times New Roman",11)).pack(pady=6)

    # ── Виконання аналізу ─────────────────────────────────────
    def _run(self):
        alpha = float(self.alpha_var.get())
        x = self._parse_col(self.tx); y = self._parse_col(self.ty)
        n = min(len(x), len(y)); x = x[:n]; y = y[:n]
        if n < 4:
            messagebox.showwarning("Замало даних",
                "Потрібно ≥ 4 пари значень (x, y)."); return

        model_name = self.model_var.get().split(":")[0].strip()
        result = self._fit_model(model_name, x, y, alpha)
        if result is None: return
        self._show_result(result, x, y, model_name, alpha)

    def _fit_model(self, name, x, y, alpha):
        from scipy.optimize import curve_fit
        n_ = name.strip().lower()
        try:
            if "лінійна" in n_ or n_ == "linear":
                # ── Лінійна ─────────────────────────────────
                X = np.column_stack([np.ones(len(x)), x])
                # lstsq повертає (beta, residuals_sum, rank, sv) — 4 значення
                beta = np.linalg.lstsq(X, y, rcond=None)[0]
                yhat = X @ beta
                params = {"a": beta[0], "b": beta[1]}
                eq = f"y = {fmt(beta[0],4)} + {fmt(beta[1],4)}·x"
                k = 2

            elif "квадратична" in n_ or n_ == "quadratic":
                # ── Квадратична ──────────────────────────────
                X = np.column_stack([np.ones(len(x)), x, x**2])
                beta = np.linalg.lstsq(X, y, rcond=None)[0]
                yhat = X @ beta
                params = {"a": beta[0], "b": beta[1], "c": beta[2]}
                eq = f"y = {fmt(beta[0],4)} + {fmt(beta[1],4)}·x + {fmt(beta[2],4)}·x²"
                k = 3

            elif "кубічна" in n_ or n_ == "cubic":
                # ── Кубічна ──────────────────────────────────
                X = np.column_stack([np.ones(len(x)), x, x**2, x**3])
                beta = np.linalg.lstsq(X, y, rcond=None)[0]
                yhat = X @ beta
                params = {"a": beta[0], "b": beta[1], "c": beta[2], "d": beta[3]}
                eq = f"y = {fmt(beta[0],4)} + {fmt(beta[1],4)}·x + {fmt(beta[2],4)}·x² + {fmt(beta[3],4)}·x³"
                k = 4

            elif "степенева" in n_ or n_ == "power":
                # ── Степенева ────────────────────────────────
                if np.any(x <= 0):
                    messagebox.showwarning("Обмеження моделі",
                        "Степенева модель вимагає x > 0 для всіх спостережень."); return None
                lx = np.log(x); ly = np.log(np.abs(y) + 1e-12)
                X = np.column_stack([np.ones(len(lx)), lx])
                beta = np.linalg.lstsq(X, ly, rcond=None)[0]
                a, b = math.exp(beta[0]), beta[1]
                yhat = a * x**b
                params = {"a": a, "b": b}
                eq = f"y = {fmt(a,4)}·x^{fmt(b,4)}"
                k = 2

            elif "експоненційна" in n_ or n_ == "exponential":
                # ── Експоненційна ─────────────────────────────
                X = np.column_stack([np.ones(len(x)), x])
                ly = np.log(np.abs(y) + 1e-12)
                beta = np.linalg.lstsq(X, ly, rcond=None)[0]
                a, b = math.exp(beta[0]), beta[1]
                yhat = a * np.exp(b * x)
                params = {"a": a, "b": b}
                eq = f"y = {fmt(a,4)}·e^({fmt(b,4)}·x)"
                k = 2

            elif "логарифмічна" in n_ or n_ == "logarithmic":
                # ── Логарифмічна ──────────────────────────────
                if np.any(x <= 0):
                    messagebox.showwarning("Обмеження моделі",
                        "Логарифмічна модель вимагає x > 0 для всіх спостережень."); return None
                X = np.column_stack([np.ones(len(x)), np.log(x)])
                beta = np.linalg.lstsq(X, y, rcond=None)[0]
                yhat = X @ beta
                params = {"a": beta[0], "b": beta[1]}
                eq = f"y = {fmt(beta[0],4)} + {fmt(beta[1],4)}·ln(x)"
                k = 2

            elif "логістична" in n_ or "logistic" in n_:
                # ── Логістична 4-параметрична ─────────────────
                def logistic4(xx, a, b, c, d):
                    return d + (a - d) / (1 + (xx / c) ** b)
                p0 = [float(np.max(y)), 1.0, float(np.median(x)), float(np.min(y))]
                popt, _ = curve_fit(logistic4, x, y, p0=p0, maxfev=15000)
                yhat = logistic4(x, *popt)
                params = {"a": popt[0], "b": popt[1], "c": popt[2], "d": popt[3]}
                eq = (f"y = {fmt(popt[3],4)} + ({fmt(popt[0],4)}−{fmt(popt[3],4)})"
                      f"/(1+(x/{fmt(popt[2],4)})^{fmt(popt[1],4)})")
                k = 4

            else:
                messagebox.showerror("Невідома модель",
                    f"Модель '{name}' не розпізнана.\nОберіть модель зі списку."); return None

            # ── Загальна статистика ───────────────────────────
            residuals = y - yhat
            sse = float(np.sum(residuals**2))
            sst = float(np.sum((y - np.mean(y))**2))
            n_obs = len(x)
            R2     = 1 - sse / sst if sst > 0 else np.nan
            R2_adj = 1 - (1 - R2) * (n_obs - 1) / (n_obs - k - 1) if n_obs > k + 1 else np.nan
            mse    = sse / (n_obs - k) if n_obs > k else np.nan
            rmse   = math.sqrt(mse) if not math.isnan(mse) else np.nan
            # F-тест
            msm = (sst - sse) / k if k > 0 else np.nan
            F   = msm / mse if (not math.isnan(mse) and mse > 1e-12) else np.nan
            p_F = float(1 - f_dist.cdf(F, k, n_obs - k - 1)) \
                  if (not math.isnan(F) and n_obs > k + 1) else np.nan
            # Нормальність залишків
            try:
                _, sw_p = shapiro(residuals) if len(residuals) >= 3 else (np.nan, np.nan)
            except Exception:
                sw_p = np.nan

            return {"equation": eq, "params": params,
                    "R2": R2, "R2_adj": R2_adj,
                    "RMSE": rmse, "F": F, "p_F": p_F, "sw_p": sw_p,
                    "residuals": residuals, "yhat": yhat,
                    "sse": sse, "sst": sst, "n": n_obs, "k": k}

        except Exception as ex:
            messagebox.showerror("Помилка підгонки", str(ex)); return None

    # ── Відображення результатів ──────────────────────────────
    def _show_result(self, r, x, y, model_name, alpha):
        self._last_run_args = (r, x, y, model_name, alpha)
        for w in self.res_frame.winfo_children(): w.destroy()
        # Reset canvas scroll
        self._res_cv.yview_moveto(0)

        p_F_ok  = (not math.isnan(r['p_F'])  and r['p_F']  < alpha)
        sw_ok   = (not math.isnan(r['sw_p']) and r['sw_p'] > alpha)
        ci_pct  = int((1 - alpha) * 100)
        n_pts   = r["n"]; k = r.get("k", 2); rmse_ = r.get("RMSE", np.nan)
        t_crit_ = float(t_dist.ppf(1 - alpha/2, max(1, n_pts - k - 1)))

        def _copy_fig(fig):
            ok, msg = _copy_fig_to_clipboard(fig)
            if ok: messagebox.showinfo("", "Скопійовано. Вставте у Word через Ctrl+V.")
            else:   messagebox.showwarning("", f"Помилка: {msg}")

        def _save_fig(fig, name):
            path = filedialog.asksaveasfilename(
                defaultextension=".png",
                filetypes=[("PNG","*.png"),("SVG","*.svg")],
                title=f"Зберегти {name}")
            if not path: return
            try:
                fig.savefig(path, dpi=150, bbox_inches="tight")
                messagebox.showinfo("Збережено", f"Збережено:\n{path}")
            except Exception as ex:
                messagebox.showerror("Помилка", str(ex))

        # ── РЯД 1: Текст ліво + Графік регресії право ─────────
        row1 = tk.Frame(self.res_frame, height=370)
        row1.pack(fill=tk.X)
        row1.pack_propagate(False)

        # Текстовий звіт (ліворуч у row1)
        txt_f = tk.Frame(row1, bg="#f8f8f8", width=270)
        txt_f.pack(side=tk.LEFT, fill=tk.Y)
        txt_f.pack_propagate(False)

        tk.Label(txt_f, text="РЕЗУЛЬТАТИ РЕГРЕСІЇ",
                 font=("Times New Roman",11,"bold"), fg="#1a4b8c",
                 bg="#f8f8f8", pady=6).pack(anchor="w", padx=8)

        fields = [
            ("Модель:",      model_name),
            ("n:",           str(r['n'])),
            ("α:",           str(alpha)),
            ("Рівняння:",    r['equation']),
            ("R²:",          fmt(r['R2'],4)),
            ("R²adj:",       fmt(r['R2_adj'],4)),
            ("RMSE:",        fmt(r['RMSE'],4)),
            ("F:",           fmt(r['F'],4)),
            ("p (F-тест):",  fmt(r['p_F'],4)),
            ("Значущість:",  "✓ Значуща" if p_F_ok else "✗ Незначуща"),
            (f"ДІ:",         f"{ci_pct}%"),
            ("SW (залишки):", fmt(r['sw_p'],4)),
            ("Нормальність:", "✓ Нормальні" if sw_ok else "⚠ Не норм."),
        ]
        for lbl, val in fields:
            row = tk.Frame(txt_f, bg="#f8f8f8"); row.pack(fill=tk.X, padx=8, pady=1)
            tk.Label(row, text=lbl, font=("Times New Roman",10,"bold"),
                     bg="#f8f8f8", fg="#555", width=14, anchor="w").pack(side=tk.LEFT)
            color = ("#27ae60" if "✓" in val else
                     "#c62828" if ("✗" in val or "⚠" in val) else "#000")
            tk.Label(row, text=val, font=("Times New Roman",10),
                     bg="#f8f8f8", fg=color,
                     wraplength=160, justify="left", anchor="w").pack(side=tk.LEFT)

        # Графік регресії (праворуч у row1)
        if not HAS_MPL:
            messagebox.showwarning("","matplotlib недоступний."); return

        g1_outer = tk.Frame(row1)
        g1_outer.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Toolbar графіка 1
        tb1 = tk.Frame(g1_outer, bg="#f0f0f0", padx=4, pady=2); tb1.pack(fill=tk.X)
        tk.Label(tb1, text="Графік регресії",
                 font=("Times New Roman",10,"bold"), bg="#f0f0f0").pack(side=tk.LEFT, padx=4)
        tk.Button(tb1, text="💾 Зберегти",
                  font=("Times New Roman",9), relief=tk.FLAT, padx=6,
                  command=lambda: _save_fig(fig1,"графік_регресії")).pack(side=tk.RIGHT, padx=2)
        tk.Button(tb1, text="📋 Копіювати",
                  font=("Times New Roman",9), relief=tk.FLAT, padx=6,
                  command=lambda: _copy_fig(fig1)).pack(side=tk.RIGHT, padx=2)
        tk.Button(tb1, text="⚙ Налаштування",
                  font=("Times New Roman",9), relief=tk.FLAT, padx=6,
                  bg="#1a4b8c", fg="white",
                  command=self._graph_settings).pack(side=tk.RIGHT, padx=2)

        # Figure 1
        fig1 = Figure(figsize=(5.2, 4.2), dpi=100)
        ax1  = fig1.add_subplot(111)
        x_sort   = np.sort(x); idx_sort = np.argsort(x)
        ax1.scatter(x, y, s=30,
                    color=self.gs.get("scatter_color","#4c72b0"),
                    zorder=3, label="Спостереження",
                    edgecolors="white", linewidths=0.5)
        ax1.plot(x_sort, r["yhat"][idx_sort],
                 color=self.gs.get("line_color","#c62828"),
                 lw=2, label="Регресійна крива")
        if n_pts > k + 2 and not math.isnan(rmse_):
            try:
                x_pred  = np.linspace(x.min(), x.max(), 300)
                x_mean_ = float(np.mean(x))
                ss_xx   = float(np.sum((x - x_mean_)**2))
                if ss_xx > 0:
                    se_fit = rmse_ * np.sqrt(1/n_pts + (x_pred-x_mean_)**2/ss_xx)
                    yhat_p = np.interp(x_pred, x_sort, r["yhat"][idx_sort])
                    ax1.fill_between(x_pred,
                                     yhat_p - t_crit_*se_fit,
                                     yhat_p + t_crit_*se_fit,
                                     alpha=0.12,
                                     color=self.gs.get("ci_color","#c62828"),
                                     label=f"{ci_pct}% ДІ")
            except Exception: pass
        custom_title = getattr(self, '_graph_title', '')
        ax1.set_title(custom_title if custom_title else f"{model_name}", fontsize=10)
        ax1.set_xlabel("x", fontsize=9); ax1.set_ylabel("y", fontsize=9)
        ax1.legend(fontsize=8)
        r2_str = f"R²={fmt(r['R2'],4)}"
        if not math.isnan(r.get("R2_adj",float("nan"))):
            r2_str += f"  R²adj={fmt(r['R2_adj'],4)}"
        eq = r.get("equation","")
        ax1.text(0.03, 0.97, f"{eq}\n{r2_str}",
                 transform=ax1.transAxes, fontsize=8, va="top",
                 fontfamily="Times New Roman",
                 bbox=dict(boxstyle="round,pad=0.4", facecolor="#eef4ff",
                           edgecolor="#1a4b8c", alpha=0.9, linewidth=1),
                 zorder=5)
        ax1.yaxis.grid(True, linestyle="--", alpha=0.35)
        ax1.spines["top"].set_visible(False); ax1.spines["right"].set_visible(False)
        fig1.tight_layout()
        self._fig = fig1
        embed_figure(fig1, g1_outer)

        # ── РЯД 2: Аналіз залишків ─────────────────────────────
        tk.Frame(self.res_frame, bg="#e0e0e0", height=1).pack(fill=tk.X, pady=4)

        row2 = tk.Frame(self.res_frame); row2.pack(fill=tk.BOTH, expand=False)
        tb2 = tk.Frame(row2, bg="#f0f0f0", padx=4, pady=2); tb2.pack(fill=tk.X)
        tk.Label(tb2, text="Аналіз залишків",
                 font=("Times New Roman",10,"bold"), bg="#f0f0f0").pack(side=tk.LEFT, padx=4)

        # Settings for residuals graph
        def _res_settings():
            dlg = tk.Toplevel(self.win); dlg.title("Налаштування залишків")
            dlg.resizable(False,False); dlg.grab_set(); set_icon(dlg)
            rf2 = ("Times New Roman",11)
            frm = tk.Frame(dlg, padx=14, pady=12); frm.pack()
            tk.Label(frm, text="Заголовок графіка:", font=rf2
                     ).grid(row=0, column=0, sticky="w", pady=4)
            tv = tk.StringVar(value=getattr(self,'_res_title',''))
            tk.Entry(frm, textvariable=tv, width=30, font=rf2
                     ).grid(row=0, column=1, sticky="w", padx=8)
            def _ok():
                self._res_title = tv.get().strip()
                dlg.destroy()
                self._show_result(*self._last_run_args)
            bf = tk.Frame(frm); bf.grid(row=1, column=0, columnspan=2, pady=(10,0))
            tk.Button(bf, text="Застосувати", bg="#c62828", fg="white",
                      font=rf2, command=_ok).pack(side=tk.LEFT, padx=4)
            tk.Button(bf, text="Скасувати", font=rf2,
                      command=dlg.destroy).pack(side=tk.LEFT)
            center_win(dlg)

        tk.Button(tb2, text="💾 Зберегти",
                  font=("Times New Roman",9), relief=tk.FLAT, padx=6,
                  command=lambda: _save_fig(fig2,"аналіз_залишків")).pack(side=tk.RIGHT, padx=2)
        tk.Button(tb2, text="📋 Копіювати",
                  font=("Times New Roman",9), relief=tk.FLAT, padx=6,
                  command=lambda: _copy_fig(fig2)).pack(side=tk.RIGHT, padx=2)
        tk.Button(tb2, text="⚙ Налаштування",
                  font=("Times New Roman",9), relief=tk.FLAT, padx=6,
                  bg="#1a4b8c", fg="white",
                  command=_res_settings).pack(side=tk.RIGHT, padx=2)

        res_arr = np.array(r["residuals"])
        res_title = getattr(self, '_res_title', '') or "Аналіз залишків"
        fig2 = Figure(figsize=(10, 3.5), dpi=100)
        fig2.suptitle(res_title, fontsize=10, y=1.0)

        ax2 = fig2.add_subplot(131)
        ax2.scatter(r["yhat"], res_arr, s=22, color="#dd8452",
                    edgecolors="white", linewidths=0.4, zorder=3)
        ax2.axhline(0, color="#333", lw=0.9, linestyle="--")
        if not math.isnan(rmse_):
            for sgn in [1,-1]:
                ax2.axhline(sgn*rmse_, color="#aaa", lw=0.6, linestyle=":")
            ax2.text(0.98, 0.98, f"±RMSE={fmt(rmse_,3)}",
                     transform=ax2.transAxes, fontsize=7, ha="right", va="top", color="#888")
        ax2.set_xlabel("ŷ", fontsize=9); ax2.set_ylabel("e = y−ŷ", fontsize=9)
        ax2.set_title("Залишки vs ŷ", fontsize=9, pad=4)
        ax2.yaxis.grid(True, linestyle="--", alpha=0.3)
        ax2.spines["top"].set_visible(False); ax2.spines["right"].set_visible(False)

        ax3 = fig2.add_subplot(132)
        ax3.hist(res_arr, bins="auto", color="#4c72b0", edgecolor="white", alpha=0.85)
        try:
            from scipy.stats import norm as _nd
            mu_, sig_ = float(np.mean(res_arr)), float(np.std(res_arr, ddof=1))
            xn = np.linspace(res_arr.min(), res_arr.max(), 100)
            bw = (res_arr.max()-res_arr.min()) / max(1, len(np.histogram_bin_edges(res_arr,"auto"))-1)
            ax3.plot(xn, _nd.pdf(xn,mu_,sig_)*len(res_arr)*bw,
                     color="#c62828", lw=1.5)
        except Exception: pass
        ax3.set_xlabel("Залишок", fontsize=9); ax3.set_ylabel("Частота", fontsize=9)
        ax3.set_title(f"Гістограма  (SW p={fmt(r['sw_p'],4)})", fontsize=9, pad=4)
        ax3.spines["top"].set_visible(False); ax3.spines["right"].set_visible(False)

        ax4 = fig2.add_subplot(133)
        try:
            from scipy.stats import probplot
            (osm,osr),(slope,intercept,rval) = probplot(res_arr, plot=None)
            ax4.plot(osm, osr, "o", color="#4c72b0", markersize=4, alpha=0.8)
            ax4.plot([min(osm),max(osm)],
                     [slope*min(osm)+intercept, slope*max(osm)+intercept],
                     color="#c62828", lw=1.5)
            ax4.set_title(f"Q-Q  (R²={rval**2:.3f})", fontsize=9, pad=4)
        except Exception:
            ax4.set_title("Q-Q", fontsize=9, pad=4)
        ax4.set_xlabel("Теор. квантилі", fontsize=9)
        ax4.set_ylabel("Вибірк. квантилі", fontsize=9)
        ax4.spines["top"].set_visible(False); ax4.spines["right"].set_visible(False)

        fig2.tight_layout()
        self._fig2 = fig2
        embed_figure(fig2, row2)

        # ── Виявлення викидів ─────────────────────────────────
        out_idx, G, _ = detect_outliers_grubbs(r["residuals"])
        if out_idx is not None:
            tk.Label(self.res_frame,
                     text=(f"⚠ Тест Граббса: підозрілий викид — спостереження "
                           f"№{out_idx+1}  (G = {fmt(G,3)}). Перевірте дані."),
                     fg="#c62828", font=("Times New Roman",10),
                     justify="left", padx=8).pack(anchor="w", pady=4)

        out_idx, G, _ = detect_outliers_grubbs(r["residuals"])
        if out_idx is not None:
            tk.Label(self.res_frame,
                     text=(f"⚠ Тест Граббса: підозрілий викид у залишках — "
                           f"спостереження №{out_idx+1}  (G = {fmt(G,3)}).\n"
                           f"   Перевірте це значення у вхідних даних."),
                     fg="#c62828", font=("Times New Roman",11),
                     justify="left", padx=6).pack(anchor="w", pady=2)




# ═══════════════════════════════════════════════════════════════
# SAMPLE SIZE CALCULATOR
# ═══════════════════════════════════════════════════════════════
class SampleSizeWindow:
    """Калькулятор розміру вибірки та статистичної потужності."""

    HELP_TEXT = """
КАЛЬКУЛЯТОР РОЗМІРУ ВИБІРКИ — ІНСТРУКЦІЯ
═════════════════════════════════════════

ДЛЯ ЧОГО ЦЕЙ КАЛЬКУЛЯТОР?
  Перед початком досліду відповідає на питання:
  "Скільки повторностей (r) мені потрібно щоб надійно
  виявити реальну різницю між варіантами?"

  Або навпаки:
  "Якщо я маю r повторностей — яка ймовірність
  що я виявлю різницю якщо вона є?"

ПАРАМЕТРИ:

  Дизайн:
    CRD — повністю рандомізований дослід
    RCBD — рандомізований повний блок
    Split-plot — розщеплені ділянки

  α (рівень значущості):
    Ймовірність хибного виявлення різниці (помилка I роду).
    Стандарт: 0.05 (5%).

  Потужність (1-β):
    Ймовірність виявити реальну різницю якщо вона існує.
    Стандарт: 0.80 (80%). Краще: 0.90 (90%).

  δ (очікувана різниця):
    Мінімальна різниця між варіантами яку важливо виявити.
    В одиницях вашого показника (т/га, см, %).
    Наприклад: δ = 0.5 т/га означає що хочемо виявити різницю ≥ 0.5 т/га.

  σ (стандартне відхилення):
    Варіабельність вашого показника.
    Візьміть з попередніх дослідів або пілотного досліду.
    Або: σ ≈ CV% × Середнє / 100.

  k (кількість варіантів):
    Скільки варіантів (обробок, сортів) у досліді.

  r (кількість повторностей):
    Залиште ПОРОЖНІМ → калькулятор знайде мінімальне r.
    Введіть число → калькулятор розрахує досягнуту потужність.

ПРИКЛАД:
  Дослід з 4 дозами добрива, очікуємо різницю ≥ 0.5 т/га,
  SD з попередніх дослідів = 0.8 т/га.
  Введіть: k=4, δ=0.5, σ=0.8, α=0.05, потужність=0.80.
  Залиште r порожнім.
  Результат покаже скільки повторностей потрібно.

ІНТЕРПРЕТАЦІЯ РЕЗУЛЬТАТУ:
  r = 4 → потрібно 4 повторності на кожен варіант.
  Загальна кількість = k × r ділянок у досліді.
  Досягнута потужність = 0.83 → 83% шанс виявити різницю.

ЯКЩО ПОТРІБНО ЗАБАГАТО ПОВТОРНОСТЕЙ:
  Збільшіть δ (нижня межа практично значущої різниці)
  або зменшіть σ (точніші вимірювання, однорідніші умови).
"""

    def __init__(self, parent):
        self.win = tk.Toplevel(parent)
        self.win.title("Калькулятор розміру вибірки та потужності")
        self.win.geometry("680x700")
        self.win.resizable(True, True)
        set_icon(self.win)
        self._build()

    def _build(self):
        top = tk.Frame(self.win, padx=8, pady=6); top.pack(fill=tk.X)
        tk.Button(top, text="▶ Розрахувати", bg="#c62828", fg="white",
                  font=("Times New Roman",13),
                  command=self._calc).pack(side=tk.LEFT, padx=4)
        tk.Button(top, text="📚 Довідка", bg="#1a4b8c", fg="white",
                  font=("Times New Roman",11),
                  command=self._show_help).pack(side=tk.LEFT, padx=4)

        pfrm = tk.LabelFrame(self.win, text="Параметри досліду",
                             font=("Times New Roman",11,"bold"),
                             padx=12, pady=8)
        pfrm.pack(fill=tk.X, padx=10, pady=(0,6))
        rf = ("Times New Roman",12)

        params = [
            ("Дизайн досліду:",                     None,   "design"),
            ("α — рівень значущості:",              "0.05", "alpha"),
            ("Потужність (1-β):",                   "0.80", "power"),
            ("δ — мінімальна різниця яку виявити:", "",     "delta"),
            ("σ — стандартне відхилення:",          "",     "sigma"),
            ("k — кількість варіантів:",            "3",    "k"),
            ("r — повторностей (порожньо→знайти):", "",     "r"),
        ]
        hints = {
            "alpha": "0.01 / 0.05 / 0.10",
            "power": "0.80 / 0.90",
            "delta": "в одиницях показника (т/га, см...)",
            "sigma": "SD з попередніх дослідів",
            "k":     "кількість варіантів/сортів",
            "r":     "порожньо = автоматичний розрахунок",
        }
        self.vars = {}
        for ri, (label, default, key) in enumerate(params):
            tk.Label(pfrm, text=label, font=rf, anchor="w"
                     ).grid(row=ri, column=0, sticky="w", pady=3)
            if key == "design":
                var = tk.StringVar(value="CRD")
                ttk.Combobox(pfrm, textvariable=var,
                             values=["CRD","RCBD","Split-plot"],
                             state="readonly", width=16,
                             font=rf).grid(row=ri, column=1, sticky="w", padx=6)
            else:
                var = tk.StringVar(value=default or "")
                tk.Entry(pfrm, textvariable=var, width=14,
                         font=rf).grid(row=ri, column=1, sticky="w", padx=6)
                if key in hints:
                    tk.Label(pfrm, text=hints[key],
                             font=("Times New Roman",9), fg="#888"
                             ).grid(row=ri, column=2, sticky="w", padx=4)
            self.vars[key] = var

        res_frm = tk.LabelFrame(self.win, text="Результат",
                                font=("Times New Roman",11,"bold"),
                                padx=8, pady=6)
        res_frm.pack(fill=tk.BOTH, expand=True, padx=10, pady=(0,10))
        vsb = ttk.Scrollbar(res_frm, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        self.res_txt = tk.Text(res_frm, wrap="word",
                               font=("Courier New",11),
                               yscrollcommand=vsb.set,
                               relief=tk.FLAT, bg="#f8f8f8",
                               padx=8, pady=6, cursor="arrow",
                               state="disabled")
        self.res_txt.pack(fill=tk.BOTH, expand=True)
        vsb.config(command=self.res_txt.yview)

    def _set_result(self, text, color="#000000"):
        self.res_txt.configure(state="normal")
        self.res_txt.delete("1.0", tk.END)
        self.res_txt.insert("1.0", text)
        self.res_txt.configure(state="disabled", fg=color)

    def _show_help(self):
        win = tk.Toplevel(self.win)
        win.title("Довідка — Калькулятор вибірки")
        win.geometry("660x640"); set_icon(win)
        frm = tk.Frame(win); frm.pack(fill=tk.BOTH, expand=True, padx=8, pady=6)
        vsb = ttk.Scrollbar(frm, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        txt = tk.Text(frm, wrap="word", font=("Times New Roman",11),
                      yscrollcommand=vsb.set, relief=tk.FLAT,
                      bg="#fafafa", padx=10, pady=8, cursor="arrow")
        txt.pack(fill=tk.BOTH, expand=True); vsb.config(command=txt.yview)
        txt.insert("1.0", self.HELP_TEXT.strip()); txt.configure(state="disabled")
        txt.bind("<MouseWheel>",
                 lambda e: txt.yview_scroll(int(-1*(e.delta/120)),"units"))
        tk.Button(win, text="Закрити", command=win.destroy,
                  font=("Times New Roman",11)).pack(pady=6)

    def _calc(self):
        try:
            alpha  = float(self.vars["alpha"].get())
            power  = float(self.vars["power"].get())
            delta  = float(self.vars["delta"].get())
            sigma  = float(self.vars["sigma"].get())
            k      = int(self.vars["k"].get())
            design = self.vars["design"].get()
            r_str  = self.vars["r"].get().strip()
        except ValueError:
            self._set_result(
                "Заповніть всі числові поля!\n\n"
                "delta i sigma — в одиницях вашого показника.\n"
                "k — ціле число >= 2.", "#c62828")
            return

        if delta <= 0 or sigma <= 0 or k < 2:
            self._set_result("delta та sigma мають бути > 0; k >= 2.", "#c62828"); return
        if not 0 < alpha < 1 or not 0 < power < 1:
            self._set_result("alpha і потужність мають бути між 0 і 1.", "#c62828"); return

        from scipy.stats import ncf
        lines = []; sep = "-" * 44

        if r_str:
            try: r = int(r_str)
            except ValueError:
                self._set_result("r має бути цілим числом.", "#c62828"); return
            if r < 2:
                self._set_result("r має бути >= 2.", "#c62828"); return

            lambda_nc = k * r * (delta**2) / (2 * sigma**2)
            F_crit = float(f_dist.ppf(1-alpha, k-1, k*(r-1)))
            ap = float(1 - ncf.cdf(F_crit, k-1, k*(r-1), lambda_nc))

            lines.append("РЕЖИМ: Розрахунок потужності при заданому r")
            lines.append(sep)
            lines.append(f"Дизайн: {design}  |  k={k}  |  r={r}")
            lines.append(f"delta={delta}  |  sigma={sigma}  |  alpha={alpha}")
            lines.append(sep)
            lines.append(f"Нецентральність lambda = {fmt(lambda_nc,3)}")
            lines.append(f"F критичне (alpha={alpha}) = {fmt(F_crit,3)}")
            lines.append(f"Досягнута потужність (1-beta) = {fmt(ap,4)}")
            lines.append("")
            if ap >= power:
                lines.append(f"OK Потужність ДОСТАТНЯ: {fmt(ap*100,1)}% >= {power*100:.0f}%")
                lines.append(f"При r={r} ви маєте {fmt(ap*100,1)}% шанс")
                lines.append(f"виявити різницю delta>={delta} якщо вона є.")
            else:
                lines.append(f"НЕДОСТАТНЯ: {fmt(ap*100,1)}% < {power*100:.0f}%")
                lines.append("Збільшіть r або delta, або зменшіть sigma.")
            if design == "RCBD":
                lines.append(f"\nRCBD: {r} блоків x {k} варіантів = {k*r} ділянок")
            elif design == "Split-plot":
                lines.append(f"\nSplit-plot: >= {r} блоків для WP фактора")
            else:
                lines.append(f"\nCRD: {k}x{r} = {k*r} ділянок")
        else:
            lines.append("РЕЖИМ: Пошук мінімальних повторностей")
            lines.append(sep)
            lines.append(f"Дизайн: {design}  |  k={k} варіантів")
            lines.append(f"Ціль: alpha={alpha}, потужність>={power}")
            lines.append(f"delta={delta}, sigma={sigma}")
            lines.append(sep)

            found = False
            for r in range(2, 101):
                lambda_nc = k * r * (delta**2) / (2 * sigma**2)
                try:
                    F_crit = float(f_dist.ppf(1-alpha, k-1, k*(r-1)))
                    pwr = float(1 - ncf.cdf(F_crit, k-1, k*(r-1), lambda_nc))
                except Exception: continue
                if pwr >= power:
                    lines.append(f"OK Мінімальне r = {r} повторностей")
                    lines.append(f"   Досягнута потужність: {fmt(pwr*100,1)}%")
                    lines.append(f"   Загальна кількість ділянок: {k}x{r} = {k*r}")
                    if design == "RCBD":
                        lines.append(f"   RCBD: {r} блоків, {k} варіантів у кожному")
                    elif design == "Split-plot":
                        lines.append(f"   Split-plot: >= {r} блоків WP")
                    lines.append("")
                    lines.append(f"{'r':>4}  {'Потужність':>12}  Статус")
                    lines.append("-" * 32)
                    for rr in range(max(2,r-2), min(r+5, 101)):
                        lnc2 = k*rr*(delta**2)/(2*sigma**2)
                        try:
                            fc2 = float(f_dist.ppf(1-alpha, k-1, k*(rr-1)))
                            pw2 = float(1 - ncf.cdf(fc2, k-1, k*(rr-1), lnc2))
                        except Exception: continue
                        mark = " <-- мінімум" if rr == r else ""
                        lines.append(f"{rr:>4}  {pw2*100:>10.1f}%{mark}")
                    found = True; break

            if not found:
                lines.append("Не вдалося знайти r <= 100.")
                lines.append("Спробуйте: збільшити delta або зменшити sigma.")

        self._set_result("\n".join(lines))



# ═══════════════════════════════════════════════════════════════
# CLUSTER ANALYSIS
# ═══════════════════════════════════════════════════════════════
class ClusterWindow:
    """Кластерний аналіз — ієрархічна кластеризація."""

    HELP_TEXT = """
КЛАСТЕРНИЙ АНАЛІЗ — ПОКРОКОВА ІНСТРУКЦІЯ
══════════════════════════════════════════

ЩО ТАКЕ КЛАСТЕРНИЙ АНАЛІЗ?
  Кластерний аналіз групує об'єкти (сорти, зразки, ділянки) так,
  щоб схожі між собою потрапили в один кластер, а несхожі — у різні.
  Результат відображається на ДЕНДРОГРАМІ — деревоподібному графіку.

КОЛИ ВИКОРИСТОВУВАТИ?
  ✓ Класифікація сортів за комплексом показників якості
  ✓ Групування ґрунтових проб за хімічним складом
  ✓ Виявлення природних груп без попередніх гіпотез
  ✓ Як доповнення до PCA для інтерпретації груп

КРОК 1. СТРУКТУРА ТАБЛИЦІ

  Перший стовпець: Назва об'єкта (сорт, зразок, ділянка — текст)
  Решта стовпців: Числові показники (змінні)

  Приклад (4 сорти, 3 показники):
  | Сорт    | Висота | Врожайність | Маса зерна |
  | Сорт А  |  95.3  |    5.8      |    38.2    |
  | Сорт Б  |  88.5  |    4.9      |    35.1    |
  | Сорт В  | 102.4  |    6.8      |    43.7    |
  | Сорт Д  |  91.2  |    5.2      |    36.8    |

  Перейменуйте заголовки (подвійний клік на синій клітинці).
  Мінімум: 2 об'єкти, 1 показник.
  Програма автоматично стандартизує дані (z-оцінки).

КРОК 2. ВИБІР МЕТОДУ ЗЧЕПЛЕННЯ

  Метод зчеплення визначає як вимірюється відстань між кластерами.

  ward (рекомендується для більшості випадків):
    Мінімізує внутрішньокластерну дисперсію.
    Дає компактні, приблизно рівні кластери.
    Найпопулярніший метод у біологічних дослідженнях. ✓

  complete (повне зчеплення):
    Відстань між кластерами = відстань між їх найдальшими об'єктами.
    Дає компактні кластери схожого розміру.
    Добре коли кластери чіткі і компактні.

  average (середнє зчеплення, UPGMA):
    Відстань = середня між усіма парами об'єктів двох кластерів.
    Компроміс між ward і complete.
    Широко використовується у філогенетичному аналізі.

  single (одиночне зчеплення):
    Відстань = відстань між найближчими об'єктами кластерів.
    Схильний до «ефекту ланцюга» — довгих витягнутих кластерів.
    Корисний для виявлення викидів і незвичних груп.

  ЯК ОБРАТИ?
    → Не знаєте яким почати → ward
    → Хочете рівні компактні групи → complete
    → Є підозра на ланцюговий зв'язок → average
    → Шукаєте нетипові об'єкти → single

КРОК 3. ВИБІР КІЛЬКОСТІ КЛАСТЕРІВ k

  k — скільки груп ви хочете отримати.

  ЯК ВИЗНАЧИТИ ПРАВИЛЬНЕ k?

  Спосіб 1: Візуальний аналіз дендрограми (найкращий!)
    Дивіться на дендрограму:
    Де є ВЕЛИКИЙ стрибок у висоті з'єднання між гілками?
    Там і є природна межа кластерів.
    Проведіть уявну горизонтальну лінію нижче цього стрибка →
    кількість вертикальних гілок що її перетинають = k.

  Спосіб 2: Правило великого стрибка
    Висота з'єднання на дендрограмі = «несхожість».
    Найбільший стрибок висоти між двома сусідніми з'єднаннями
    вказує на оптимальне k.

  Спосіб 3: Змістовна логіка
    Якщо ви знаєте що в природі є 3 групи (ранньостиглі,
    середньостиглі, пізньостиглі) → k=3.

  ТИПОВІ ЗНАЧЕННЯ ДЛЯ АГРОНОМІЇ:
    Класифікація сортів: k = 2–5
    Групування ґрунтових проб: k = 3–6
    Екологічні зони: k = 3–8

КРОК 4. ЧИТАННЯ ДЕНДРОГРАМИ

  Горизонтальна вісь: об'єкти (сорти, зразки)
  Вертикальна вісь: відстань (несхожість)

  Об'єкти що з'єднуються НИЗЬКО → дуже схожі
  Об'єкти що з'єднуються ВИСОКО → дуже несхожі
  Різні кольори = різні кластери при обраному k
  Горизонтальна пунктирна лінія = поріг відсікання для k кластерів

  Об'єкт що з'єднується останнім (найвище) →
  найбільш відмінний від усіх інших (потенційний викид!)

КРОК 5. ТАБЛИЦЯ ПРИНАЛЕЖНОСТІ

  Після дендрограми виводиться таблиця:
  кожен об'єкт і номер його кластера (1, 2, 3...).
"""

    def __init__(self, parent, gs):
        self.win = tk.Toplevel(parent)
        self.win.title("Кластерний аналіз")
        self.win.geometry("960x660"); set_icon(self.win)
        self.gs = gs
        self._cl_fig = None
        self._cl_gs  = {
            "font_family":    "Times New Roman",
            "font_size":      9,
            "leaf_font_size": 9,
            "line_color":     "#2176ae",
            "threshold_color":"#c62828",
            "show_threshold": True,
            "figsize_w":      10,
            "figsize_h":      5.5,
        }
        self._build()

    def _build(self):
        # ── Toolbar ──────────────────────────────────────────
        top = tk.Frame(self.win, padx=8, pady=6); top.pack(fill=tk.X)
        tk.Button(top, text="▶ Кластеризувати", bg="#c62828", fg="white",
                  font=("Times New Roman", 13),
                  command=self._run).pack(side=tk.LEFT, padx=4)

        # Параметри
        tk.Label(top, text="Метод:", font=("Times New Roman",12)).pack(side=tk.LEFT, padx=(8,2))
        self.meth_var = tk.StringVar(value="ward")
        ttk.Combobox(top, textvariable=self.meth_var,
                     values=["ward","complete","average","single"],
                     state="readonly", width=12).pack(side=tk.LEFT, padx=2)
        tk.Label(top, text="k:", font=("Times New Roman",12)).pack(side=tk.LEFT, padx=(8,2))
        self.k_var = tk.IntVar(value=3)
        tk.Spinbox(top, from_=2, to=20, textvariable=self.k_var,
                   width=4, font=("Times New Roman",11)).pack(side=tk.LEFT, padx=2)

        # Налаштування — спадне меню
        mb2 = tk.Menubutton(top, text="⚙ Налаштування ▾",
                            font=("Times New Roman",11), relief=tk.RAISED, bd=2)
        mb2.pack(side=tk.LEFT, padx=6)
        sm = tk.Menu(mb2, tearoff=0)
        sm.add_command(label="Додати рядок",      command=self._add_row)
        sm.add_command(label="Видалити рядок",    command=self._del_row)
        sm.add_separator()
        sm.add_command(label="Додати стовпець",   command=self._add_col)
        sm.add_command(label="Видалити стовпець", command=self._del_col)
        sm.add_separator()
        sm.add_command(label="🗑 Очистити таблицю", command=self._clear_table)
        mb2["menu"] = sm

        tk.Button(top, text="Вставити з буфера",
                  font=("Times New Roman",11),
                  command=self._paste).pack(side=tk.LEFT, padx=4)
        tk.Button(top, text="📚 Довідка", bg="#1a4b8c", fg="white",
                  font=("Times New Roman",11),
                  command=self._show_help).pack(side=tk.LEFT, padx=4)

        tk.Label(top,
                 text="Подвійний клік на заголовку → перейменувати",
                 font=("Times New Roman",9), fg="#666").pack(side=tk.LEFT, padx=6)

        # ── Таблиця ─────────────────────────────────────────
        mid = tk.Frame(self.win); mid.pack(fill=tk.BOTH, expand=True, padx=8)
        self.rows_n = 18; self.cols_n = 8
        self._canvas = tk.Canvas(mid)
        self._canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb = ttk.Scrollbar(mid, orient="vertical", command=self._canvas.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        self._canvas.configure(yscrollcommand=sb.set)
        self.inner = tk.Frame(self._canvas)
        self._canvas.create_window((0,0), window=self.inner, anchor="nw")
        self.inner.bind("<Configure>",
                        lambda e: self._canvas.config(scrollregion=self._canvas.bbox("all")))
        self.win.bind("<MouseWheel>",
                      lambda e: self._canvas.yview_scroll(int(-1*(e.delta/120)),"units"))

        default_hdr = ["Назва об'єкта"] + [f"Показник {j}" for j in range(1, self.cols_n)]
        self.header_vars = []; self.header_labels = []
        for j in range(self.cols_n):
            var = tk.StringVar(value=default_hdr[j] if j < len(default_hdr) else f"П{j}")
            self.header_vars.append(var)
            lbl = tk.Label(self.inner, textvariable=var, width=12, cursor="hand2",
                           bg="#1a4b8c", fg="white", relief=tk.RIDGE,
                           font=("Times New Roman",11,"bold"))
            lbl.grid(row=0, column=j, padx=1, pady=1, sticky="nsew")
            lbl.bind("<Double-Button-1>", lambda e, idx=j: self._rename_col(idx))
            self.header_labels.append(lbl)

        self.entries = []
        for i in range(self.rows_n):
            row_ = []
            for j in range(self.cols_n):
                e = tk.Entry(self.inner, width=12, font=("Times New Roman",11),
                             highlightthickness=1, highlightbackground="#c0c0c0")
                e.grid(row=i+1, column=j, padx=1, pady=1)
                row_.append(e)
            self.entries.append(row_)
        _bind_nav(self.entries, self.win)

    # ── Перейменування ────────────────────────────────────────
    def _rename_col(self, idx):
        dlg = tk.Toplevel(self.win); dlg.title("Перейменувати")
        dlg.resizable(False, False); dlg.grab_set()
        tk.Label(dlg, text=f"Назва стовпця {idx+1}:",
                 font=("Times New Roman",12)).pack(padx=16, pady=14)
        var = tk.StringVar(value=self.header_vars[idx].get())
        e = tk.Entry(dlg, textvariable=var, font=("Times New Roman",12), width=26)
        e.pack(padx=16, pady=4); e.select_range(0, tk.END); e.focus_set()
        def apply():
            nm = var.get().strip()
            if nm: self.header_vars[idx].set(nm)
            dlg.destroy()
        tk.Button(dlg, text="OK", bg="#c62828", fg="white",
                  font=("Times New Roman",12), command=apply).pack(pady=(4,14))
        dlg.bind("<Return>", lambda ev: apply()); center_win(dlg)

    # ── Управління таблицею ───────────────────────────────────
    def _add_row(self):
        i = self.rows_n; row_ = []
        for j in range(self.cols_n):
            e = tk.Entry(self.inner, width=12, font=("Times New Roman",11),
                         highlightthickness=1, highlightbackground="#c0c0c0")
            e.grid(row=i+1, column=j, padx=1, pady=1)
            row_.append(e)
        self.entries.append(row_); self.rows_n += 1
        _bind_nav(self.entries, self.win)

    def _del_row(self):
        if not self.entries: return
        for e in self.entries.pop(): e.destroy()
        self.rows_n -= 1

    def _add_col(self):
        ci = self.cols_n; self.cols_n += 1
        var = tk.StringVar(value=f"Показник {ci}")
        self.header_vars.append(var)
        lbl = tk.Label(self.inner, textvariable=var, width=12, cursor="hand2",
                       bg="#1a4b8c", fg="white", relief=tk.RIDGE,
                       font=("Times New Roman",11,"bold"))
        lbl.grid(row=0, column=ci, padx=1, pady=1, sticky="nsew")
        lbl.bind("<Double-Button-1>", lambda e, idx=ci: self._rename_col(idx))
        self.header_labels.append(lbl)
        for i, row_ in enumerate(self.entries):
            e = tk.Entry(self.inner, width=12, font=("Times New Roman",11),
                         highlightthickness=1, highlightbackground="#c0c0c0")
            e.grid(row=i+1, column=ci, padx=1, pady=1)
            row_.append(e)
        _bind_nav(self.entries, self.win)

    def _del_col(self):
        if self.cols_n <= 2: return
        self.header_labels.pop().destroy(); self.header_vars.pop()
        for row_ in self.entries: row_.pop().destroy()
        self.cols_n -= 1

    def _clear_table(self):
        if not messagebox.askyesno("Очистити",
                "Видалити всі дані? (Заголовки залишаться)"): return
        for row in self.entries:
            for e in row: e.delete(0, tk.END)

    # ── Вставка / Довідка ────────────────────────────────────
    def _paste(self):
        try: data = self.win.clipboard_get()
        except Exception:
            messagebox.showwarning("Буфер порожній",
                "Скопіюйте дані з Excel (Ctrl+C) і спробуйте знову."); return
        if not data.strip(): return
        pos = (0, 0)
        w = self.win.focus_get()
        if isinstance(w, tk.Entry):
            for i, row_ in enumerate(self.entries):
                for j, e in enumerate(row_):
                    if e is w: pos=(i,j); break
        r0, c0 = pos
        for ir, line in enumerate(data.splitlines()):
            if not line.strip(): continue
            while r0+ir >= len(self.entries): self._add_row()
            for jc, val in enumerate(line.split("\t")):
                cc = c0+jc
                if cc >= self.cols_n: continue
                self.entries[r0+ir][cc].delete(0,tk.END)
                self.entries[r0+ir][cc].insert(0, val.strip())

    def _show_help(self):
        win = tk.Toplevel(self.win); win.title("Довідка — Кластерний аналіз")
        win.geometry("720x680"); set_icon(win)
        frm = tk.Frame(win); frm.pack(fill=tk.BOTH, expand=True, padx=8, pady=6)
        vsb = ttk.Scrollbar(frm, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        txt = tk.Text(frm, wrap="word", font=("Times New Roman",11),
                      yscrollcommand=vsb.set, relief=tk.FLAT,
                      bg="#fafafa", padx=10, pady=8, cursor="arrow")
        txt.pack(fill=tk.BOTH, expand=True); vsb.config(command=txt.yview)
        txt.insert("1.0", self.HELP_TEXT.strip()); txt.configure(state="disabled")
        txt.bind("<MouseWheel>",
                 lambda e: txt.yview_scroll(int(-1*(e.delta/120)),"units"))
        tk.Button(win, text="Закрити", command=win.destroy,
                  font=("Times New Roman",11)).pack(pady=6)

    # ── Налаштування графіка ──────────────────────────────────
    def _restyle_cluster(self, win, obj_names, Z, k, method, graph_frame):
        dlg = tk.Toplevel(win); dlg.title("Налаштування дендрограми")
        dlg.resizable(False, False); set_icon(dlg); dlg.grab_set()
        gs = self._cl_gs
        frm = tk.Frame(dlg, padx=16, pady=14); frm.pack()
        rb_f = ("Times New Roman",12)

        ff_v  = tk.StringVar(value=gs["font_family"])
        fz_v  = tk.IntVar(value=gs["font_size"])
        lf_v  = tk.IntVar(value=gs["leaf_font_size"])
        fw_v  = tk.DoubleVar(value=gs["figsize_w"])
        fh_v  = tk.DoubleVar(value=gs["figsize_h"])
        st_v  = tk.BooleanVar(value=gs["show_threshold"])
        lc_ref = [gs["line_color"]]
        tc_ref = [gs["threshold_color"]]

        rows_cfg = [
            ("Шрифт:",                  "combo",  ff_v, ["Times New Roman","Arial","Calibri","Georgia"]),
            ("Розмір підписів осей:",   "spin",   fz_v, (6,18)),
            ("Розмір підписів об'єктів:", "spin", lf_v, (5,16)),
            ("Ширина графіка:",         "scale",  fw_v, (5.,20.)),
            ("Висота графіка:",         "scale",  fh_v, (3.,12.)),
            ("Показувати поріг k:",     "check",  st_v, None),
        ]
        btn_refs = {}
        for ri, (lbl, wt, var, opts) in enumerate(rows_cfg):
            tk.Label(frm, text=lbl, font=rb_f).grid(row=ri, column=0, sticky="w", pady=4)
            if wt=="combo":
                ttk.Combobox(frm, textvariable=var, values=opts,
                             state="readonly", width=20).grid(row=ri, column=1, sticky="w", padx=8)
            elif wt=="spin":
                tk.Spinbox(frm, from_=opts[0], to=opts[1], textvariable=var,
                           width=7).grid(row=ri, column=1, sticky="w", padx=8)
            elif wt=="scale":
                tk.Scale(frm, from_=opts[0], to=opts[1], resolution=0.5,
                         orient="horizontal", variable=var,
                         length=160).grid(row=ri, column=1, sticky="w", padx=8)
            elif wt=="check":
                tk.Checkbutton(frm, variable=var).grid(row=ri, column=1, sticky="w", padx=8)

        base_r = len(rows_cfg)
        for ri, (lbl, ref) in enumerate([("Колір ліній дендрограми:", lc_ref),
                                          ("Колір порогової лінії:",    tc_ref)]):
            tk.Label(frm, text=lbl, font=rb_f).grid(row=base_r+ri, column=0, sticky="w", pady=4)
            btn = tk.Button(frm, width=6, relief=tk.SUNKEN, bg=ref[0])
            btn.grid(row=base_r+ri, column=1, sticky="w", padx=8)
            def _pick(r=ref, b=btn):
                ch = colorchooser.askcolor(color=r[0], parent=dlg)
                if ch and ch[1]: r[0]=ch[1]; b.configure(bg=ch[1])
            btn.configure(command=_pick); btn_refs[ri] = btn

        def apply():
            self._cl_gs.update({
                "font_family":    ff_v.get(), "font_size": fz_v.get(),
                "leaf_font_size": lf_v.get(), "figsize_w": fw_v.get(),
                "figsize_h":      fh_v.get(), "show_threshold": st_v.get(),
                "line_color":     lc_ref[0],  "threshold_color": tc_ref[0],
            })
            dlg.destroy()
            self._draw_dendrogram(graph_frame, obj_names, Z, k, method)

        bf = tk.Frame(frm); bf.grid(row=base_r+2, column=0, columnspan=2, pady=(14,0))
        tk.Button(bf, text="OK (застосувати)", bg="#c62828", fg="white",
                  font=rb_f, command=apply).pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Скасувати", font=rb_f, command=dlg.destroy).pack(side=tk.LEFT)
        center_win(dlg)

    # ── Побудова дендрограми ──────────────────────────────────
    def _draw_dendrogram(self, frame, obj_names, Z, k, method):
        from scipy.cluster.hierarchy import dendrogram as _dendro
        for w in frame.winfo_children(): w.destroy()
        gs = self._cl_gs
        fig = Figure(figsize=(10, 6), dpi=100)
        ax  = fig.add_subplot(111)

        # Кольори гілок через color_threshold
        thresh = float(Z[-(k-1), 2]) if k > 1 else float("inf")
        _dendro(Z, labels=obj_names, ax=ax,
                leaf_rotation=90, leaf_font_size=gs["leaf_font_size"],
                color_threshold=thresh if gs["show_threshold"] else 0,
                above_threshold_color=gs["line_color"])

        if gs["show_threshold"] and k > 1:
            ax.axhline(thresh, color=gs["threshold_color"],
                       lw=1.2, linestyle="--",
                       label=f"Поріг k={k}")
            ax.legend(fontsize=gs["font_size"], framealpha=0.7)

        ax.set_title(f"Ієрархічна кластеризація  |  Метод: {method}  |  k = {k}",
                     fontsize=gs["font_size"]+1, fontfamily=gs["font_family"])
        ax.set_ylabel("Відстань (несхожість)",
                      fontsize=gs["font_size"], fontfamily=gs["font_family"])
        ax.tick_params(labelsize=gs["font_size"])
        ax.yaxis.grid(True, linestyle="--", alpha=0.3)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        fig.tight_layout()
        self._cl_fig = fig

        embed_figure(fig, frame)

    # ── Виконання аналізу ─────────────────────────────────────
    def _run(self):
        from scipy.cluster.hierarchy import linkage, fcluster
        from scipy.spatial.distance import pdist

        raw = [[e.get().strip() for e in row] for row in self.entries]
        raw = [r for r in raw if any(v for v in r)]
        if not raw:
            messagebox.showwarning("Немає даних","Введіть дані у таблицю."); return

        obj_names = []; data_matrix = []
        for row in raw:
            nm = row[0].strip() if row else ""
            if not nm: continue
            vals = []
            for v in row[1:]:
                if not v: continue
                try: vals.append(float(v.replace(",",".")))
                except Exception: continue
            if vals:
                obj_names.append(nm); data_matrix.append(vals)

        if len(data_matrix) < 2:
            messagebox.showwarning("Замало об'єктів",
                "Потрібно щонайменше 2 об'єкти з числовими даними.\n"
                "Перший стовпець = назва об'єкта (текст).\n"
                "Решта стовпців = числові показники."); return

        min_cols = min(len(r) for r in data_matrix)
        X = np.array([r[:min_cols] for r in data_matrix], dtype=float)

        from scipy.stats import zscore
        X_std = zscore(X, axis=0, ddof=1); X_std = np.nan_to_num(X_std)

        method = self.meth_var.get()
        k = self.k_var.get()
        k = max(2, min(k, len(obj_names)))

        try:
            Z = linkage(X_std, method=method)
        except Exception as ex:
            messagebox.showerror("Помилка", str(ex)); return

        labels_cl = fcluster(Z, k, criterion='maxclust')

        if not HAS_MPL:
            messagebox.showwarning("","matplotlib недоступний."); return

        # ── Вікно результатів ──────────────────────────────────
        win = tk.Toplevel(self.win)
        win.title("Кластерний аналіз — Результати")
        win.geometry("1060x720"); set_icon(win)

        # Toolbar результатів
        tb = tk.Frame(win, padx=6, pady=5); tb.pack(fill=tk.X)
        graph_frame = tk.Frame(win); # буде pack після tb

        tk.Button(tb, text="📋 Копіювати дендрограму", font=("Times New Roman",11),
                  command=lambda: self._copy_dendro()).pack(side=tk.LEFT, padx=4)
        tk.Button(tb, text="⚙ Налаштування графіка", font=("Times New Roman",11),
                  command=lambda: self._restyle_cluster(
                      win, obj_names, Z, k, method, graph_frame)
                  ).pack(side=tk.LEFT, padx=4)
        tk.Label(tb,
                 text=f"Метод: {method}  |  k = {k}  |  Об'єктів: {len(obj_names)}",
                 font=("Times New Roman",11), fg="#555").pack(side=tk.LEFT, padx=10)

        # Дендрограма
        graph_frame.pack(fill=tk.BOTH, expand=True, padx=4)
        self._draw_dendrogram(graph_frame, obj_names, Z, k, method)

        # Таблиця приналежності (знизу, прокручувана)
        tbl_frame = tk.Frame(win); tbl_frame.pack(fill=tk.X, padx=8, pady=4)
        tk.Label(tbl_frame, text="Приналежність до кластерів:",
                 font=("Times New Roman",11,"bold"), anchor="w").pack(fill=tk.X)
        membership_rows = sorted(
            [[nm, f"Кластер {cl}"] for nm, cl in zip(obj_names, labels_cl)],
            key=lambda r: r[1])
        frm_m, _ = make_tv(tbl_frame, ["Об'єкт","Кластер"], membership_rows)
        frm_m.pack(fill=tk.X)

    def _copy_dendro(self):
        if self._cl_fig is None:
            messagebox.showwarning("","Спочатку виконайте кластеризацію."); return
        ok, msg = _copy_fig_to_clipboard(self._cl_fig)
        if ok: messagebox.showinfo("","Дендрограму скопійовано.\nВставте у Word через Ctrl+V.")
        else:   messagebox.showwarning("",f"Помилка: {msg}")




# ═══════════════════════════════════════════════════════════════
# PCA MODULE
# ═══════════════════════════════════════════════════════════════
class PCAWindow:
    """Аналіз головних компонент (PCA)."""

    HELP_TEXT = """
PCA — ПОКРОКОВА ІНСТРУКЦІЯ
══════════════════════════════════════════

ЩО ТАКЕ PCA?
  PCA (Principal Component Analysis) — Аналіз головних компонент.
  Стискає багато змінних (показників) до кількох «узагальнених»
  показників (головних компонент, ГК), зберігаючи максимум інформації.

КОЛИ ЗАСТОСОВУВАТИ?
  ✓ У вас 5+ показників для кожного об'єкта (сорту, зразка)
  ✓ Хочете виявити природне групування об'єктів
  ✓ Хочете зрозуміти які показники «ходять разом»
  ✓ Як попередній крок перед MANOVA (якщо n ≤ p)
  ✓ Для виявлення викидів (об'єкти далеко від центру biplot)

КРОК 1. СТРУКТУРА ДАНИХ

  Перший стовпець: Мітка об'єкта (назва сорту, зразка — текст)
                   Необов'язково — якщо числове, вважається показником.
  Решта стовпців: Числові показники (змінні).

  Перший рядок: Назви показників (заголовки, синій рядок).

  Приклад (4 сорти, 4 показники):
  | Сорт    | Врожайн. | Висота | Маса зерна | Вміст білку |
  | Сорт А  |   5.8    |  95.3  |    38.2    |    12.5     |
  | Сорт Б  |   4.9    |  88.5  |    35.1    |    14.2     |

  Мінімум: 2 об'єкти, 2 показники.

КРОК 2. ПІДГОТОВКА І ВИКОНАННЯ
  Перейменуйте заголовки (подвійний клік на синій клітинці).
  Введіть дані або вставте з Excel.
  Натисніть «▶ Виконати PCA».

  Програма автоматично СТАНДАРТИЗУЄ дані (z-оцінки)
  щоб показники з різними одиницями мали однаковий вплив.

КРОК 3. SCREE PLOT (Графік відсіювання)

  Стовпчики: % дисперсії пояснений кожною ГК.
  Червона лінія: кумулятивний % (зростаючий).
  Пунктир: 80% поріг.

  Скільки ГК залишити?
  «Правило ліктя»: знайдіть де графік різко стає пологим
  → точка вище = оптимальна кількість ГК.
  Зазвичай ГК1 + ГК2 пояснюють 70-85% → достатньо.
  Власне значення > 1 (Критерій Кайзера) → включайте.

КРОК 4. BIPLOT (ГК1 × ГК2)

  Точки = об'єкти (сорти, зразки):
    Близькі точки = схожі об'єкти за всіма показниками.
    Далекі точки = дуже різні.
    Кластери = природні групи.

  Червоні стрілки = показники (змінні):
    Довга стрілка = показник добре описаний цими ГК.
    Стрілки в одному напрямку = показники корелюють.
    Стрілки протилежних напрямків = обернена кореляція.
    Стрілки під кутом 90° = незалежні показники.
    Об'єкт близько до стрілки = великe значення цього показника.

КРОК 5. ТАБЛИЦЯ НАВАНТАЖЕНЬ (Loadings)

  Теплова карта: як кожен показник пов'язаний з кожною ГК.
  |Навантаження| > 0.5 вважається значущим.
  Темно-зелений: сильна позитивна кореляція з ГК.
  Темно-червоний: сильна негативна кореляція.

КРОК 6. ТАБЛИЦЯ КОМПОНЕНТ

  Власне значення: дисперсія пояснена кожною ГК.
  % дисперсії: відносний внесок.
  Кумулятивний %: наростаючий підсумок.
"""

    def __init__(self, parent, gs):
        self.win = tk.Toplevel(parent)
        self.win.title("Аналіз головних компонент (PCA)")
        self.win.geometry("1000x680"); set_icon(self.win)
        self.gs = gs
        self._pca_fig = None
        self._pca_gs  = {
            "point_color":   "#dd8452",
            "arrow_color":   "#c62828",
            "bar_color":     "#4c72b0",
            "cum_color":     "#c62828",
            "heatmap_cmap":  "RdYlGn",
            "font_family":   "Times New Roman",
            "font_size":     9,
            "point_size":    30,
            "arrow_scale":   0.7,
            "annotate_obj":  True,
            "annotate_var":  True,
        }
        self._build()

    def _build(self):
        # ── Toolbar ──────────────────────────────────────────
        top = tk.Frame(self.win, padx=8, pady=6); top.pack(fill=tk.X)
        tk.Button(top, text="▶ Виконати PCA", bg="#c62828", fg="white",
                  font=("Times New Roman",13),
                  command=self._run).pack(side=tk.LEFT, padx=4)

        # Налаштування — спадне меню
        mb2 = tk.Menubutton(top, text="⚙ Налаштування ▾",
                            font=("Times New Roman",11), relief=tk.RAISED, bd=2)
        mb2.pack(side=tk.LEFT, padx=4)
        sm = tk.Menu(mb2, tearoff=0)
        sm.add_command(label="Додати рядок",      command=self._add_row)
        sm.add_command(label="Видалити рядок",    command=self._del_row)
        sm.add_separator()
        sm.add_command(label="Додати стовпець",   command=self._add_col)
        sm.add_command(label="Видалити стовпець", command=self._del_col)
        sm.add_separator()
        sm.add_command(label="🗑 Очистити таблицю", command=self._clear_table)
        mb2["menu"] = sm

        tk.Button(top, text="Вставити з буфера",
                  font=("Times New Roman",11),
                  command=self._paste).pack(side=tk.LEFT, padx=4)
        tk.Button(top, text="📚 Довідка", bg="#1a4b8c", fg="white",
                  font=("Times New Roman",11),
                  command=self._show_help).pack(side=tk.LEFT, padx=4)

        tk.Label(top,
                 text="Подвійний клік на заголовку → перейменувати показник",
                 font=("Times New Roman",9), fg="#666").pack(side=tk.LEFT, padx=10)

        # ── Таблиця ─────────────────────────────────────────
        mid = tk.Frame(self.win); mid.pack(fill=tk.BOTH, expand=True, padx=8)
        self.rows_n = 18; self.cols_n = 8
        self._canvas = tk.Canvas(mid)
        self._canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb = ttk.Scrollbar(mid, orient="vertical", command=self._canvas.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        self._canvas.configure(yscrollcommand=sb.set)
        self.inner = tk.Frame(self._canvas)
        self._canvas.create_window((0,0), window=self.inner, anchor="nw")
        self.inner.bind("<Configure>",
                        lambda e: self._canvas.config(scrollregion=self._canvas.bbox("all")))
        self.win.bind("<MouseWheel>",
                      lambda e: self._canvas.yview_scroll(int(-1*(e.delta/120)),"units"))

        self.header_labels = []; self.header_vars = []
        default_headers = ["Мітка об'єкта"] + [f"Показник {j}" for j in range(1, self.cols_n)]
        for j in range(self.cols_n):
            var = tk.StringVar(value=default_headers[j] if j < len(default_headers) else f"П{j}")
            self.header_vars.append(var)
            lbl = tk.Label(self.inner, textvariable=var, relief=tk.RIDGE, width=13,
                           bg="#1a4b8c", fg="white", cursor="hand2",
                           font=("Times New Roman",11,"bold"))
            lbl.grid(row=0, column=j, padx=1, pady=1, sticky="nsew")
            lbl.bind("<Double-Button-1>", lambda e, idx=j: self._rename_col(idx))
            self.header_labels.append(lbl)

        self.entries = []
        for i in range(self.rows_n):
            row_ = []
            for j in range(self.cols_n):
                e = tk.Entry(self.inner, width=13, font=("Times New Roman",11),
                             highlightthickness=1, highlightbackground="#c0c0c0")
                e.grid(row=i+1, column=j, padx=1, pady=1)
                row_.append(e)
            self.entries.append(row_)
        _bind_nav(self.entries, self.win)

    # ── Перейменування заголовка ──────────────────────────────
    def _rename_col(self, idx):
        dlg = tk.Toplevel(self.win); dlg.title("Перейменувати")
        dlg.resizable(False, False); dlg.grab_set()
        tk.Label(dlg, text=f"Назва показника {idx+1}:",
                 font=("Times New Roman",12)).pack(padx=16, pady=14)
        var = tk.StringVar(value=self.header_vars[idx].get())
        e = tk.Entry(dlg, textvariable=var, font=("Times New Roman",12), width=26)
        e.pack(padx=16, pady=4); e.select_range(0, tk.END); e.focus_set()
        def apply():
            nm = var.get().strip()
            if nm: self.header_vars[idx].set(nm)
            dlg.destroy()
        tk.Button(dlg, text="OK", bg="#c62828", fg="white",
                  font=("Times New Roman",12), command=apply).pack(pady=(4,14))
        dlg.bind("<Return>", lambda ev: apply()); center_win(dlg)

    # ── Управління таблицею ───────────────────────────────────
    def _add_row(self):
        i = self.rows_n; row_ = []
        for j in range(self.cols_n):
            e = tk.Entry(self.inner, width=13, font=("Times New Roman",11),
                         highlightthickness=1, highlightbackground="#c0c0c0")
            e.grid(row=i+1, column=j, padx=1, pady=1)
            row_.append(e)
        self.entries.append(row_); self.rows_n += 1
        _bind_nav(self.entries, self.win)

    def _del_row(self):
        if not self.entries: return
        for e in self.entries.pop(): e.destroy()
        self.rows_n -= 1

    def _add_col(self):
        ci = self.cols_n; self.cols_n += 1
        var = tk.StringVar(value=f"Показник {ci}")
        self.header_vars.append(var)
        lbl = tk.Label(self.inner, textvariable=var, relief=tk.RIDGE, width=13,
                       bg="#1a4b8c", fg="white", cursor="hand2",
                       font=("Times New Roman",11,"bold"))
        lbl.grid(row=0, column=ci, padx=1, pady=1, sticky="nsew")
        lbl.bind("<Double-Button-1>", lambda e, idx=ci: self._rename_col(idx))
        self.header_labels.append(lbl)
        for i, row_ in enumerate(self.entries):
            e = tk.Entry(self.inner, width=13, font=("Times New Roman",11),
                         highlightthickness=1, highlightbackground="#c0c0c0")
            e.grid(row=i+1, column=ci, padx=1, pady=1)
            row_.append(e)
        _bind_nav(self.entries, self.win)

    def _del_col(self):
        if self.cols_n <= 3: return
        self.header_labels.pop().destroy(); self.header_vars.pop()
        for row_ in self.entries: row_.pop().destroy()
        self.cols_n -= 1

    def _clear_table(self):
        if not messagebox.askyesno("Очистити",
                "Видалити всі дані? (Заголовки залишаться)"): return
        for row in self.entries:
            for e in row: e.delete(0, tk.END)

    # ── Вставка / Довідка ────────────────────────────────────
    def _paste(self):
        try: data = self.win.clipboard_get()
        except Exception:
            messagebox.showwarning("Буфер порожній",
                "Скопіюйте дані з Excel (Ctrl+C) і спробуйте знову."); return
        if not data.strip(): return
        pos = (0, 0)
        w = self.win.focus_get()
        if isinstance(w, tk.Entry):
            for i, row_ in enumerate(self.entries):
                for j, e in enumerate(row_):
                    if e is w: pos=(i,j); break
        r0, c0 = pos
        for ir, line in enumerate(data.splitlines()):
            if not line.strip(): continue
            while r0+ir >= len(self.entries): self._add_row()
            for jc, val in enumerate(line.split("\t")):
                cc = c0+jc
                if cc >= self.cols_n: continue
                self.entries[r0+ir][cc].delete(0,tk.END)
                self.entries[r0+ir][cc].insert(0, val.strip())

    def _show_help(self):
        win = tk.Toplevel(self.win); win.title("Довідка — PCA")
        win.geometry("700x660"); set_icon(win)
        frm = tk.Frame(win); frm.pack(fill=tk.BOTH, expand=True, padx=8, pady=6)
        vsb = ttk.Scrollbar(frm, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        txt = tk.Text(frm, wrap="word", font=("Times New Roman",11),
                      yscrollcommand=vsb.set, relief=tk.FLAT,
                      bg="#fafafa", padx=10, pady=8, cursor="arrow")
        txt.pack(fill=tk.BOTH, expand=True); vsb.config(command=txt.yview)
        txt.insert("1.0", self.HELP_TEXT.strip()); txt.configure(state="disabled")
        txt.bind("<MouseWheel>",
                 lambda e: txt.yview_scroll(int(-1*(e.delta/120)),"units"))
        tk.Button(win, text="Закрити", command=win.destroy,
                  font=("Times New Roman",11)).pack(pady=6)

    # ── Налаштування графіків PCA ─────────────────────────────
    def _restyle_pca(self, callback=None):
        dlg = tk.Toplevel(self.win); dlg.title("Налаштування графіків PCA")
        dlg.resizable(False, False); set_icon(dlg); dlg.grab_set()
        gs = self._pca_gs
        frm = tk.Frame(dlg, padx=16, pady=14); frm.pack()
        rb_f = ("Times New Roman",12)

        ff_v  = tk.StringVar(value=gs["font_family"])
        fz_v  = tk.IntVar(value=gs["font_size"])
        ps_v  = tk.IntVar(value=gs["point_size"])
        sc_v  = tk.DoubleVar(value=gs["arrow_scale"])
        ao_v  = tk.BooleanVar(value=gs["annotate_obj"])
        av_v  = tk.BooleanVar(value=gs["annotate_var"])
        cm_v  = tk.StringVar(value=gs["heatmap_cmap"])
        col_refs = {k: gs[k] for k in ("point_color","arrow_color","bar_color","cum_color")}
        col_btns = {}

        rows_cfg = [
            ("Шрифт:",               "combo",  ff_v, ["Times New Roman","Arial","Calibri","Georgia"]),
            ("Розмір шрифту:",       "spin",   fz_v, (6,18)),
            ("Розмір точок (biplot):","spin",  ps_v, (5,80)),
            ("Масштаб стрілок:",     "scale",  sc_v, (0.2,1.5)),
            ("Підписи об'єктів:",    "check",  ao_v, None),
            ("Підписи змінних:",     "check",  av_v, None),
            ("Палітра теплової карти:","combo",cm_v, ["RdYlGn","coolwarm","RdBu","viridis","plasma"]),
        ]
        for ri, (lbl, wt, var, opts) in enumerate(rows_cfg):
            tk.Label(frm, text=lbl, font=rb_f).grid(row=ri, column=0, sticky="w", pady=4)
            if wt=="combo":
                ttk.Combobox(frm, textvariable=var, values=opts,
                             state="readonly", width=20).grid(row=ri, column=1, sticky="w", padx=8)
            elif wt=="spin":
                tk.Spinbox(frm, from_=opts[0], to=opts[1], textvariable=var,
                           width=7).grid(row=ri, column=1, sticky="w", padx=8)
            elif wt=="scale":
                tk.Scale(frm, from_=opts[0], to=opts[1], resolution=0.05,
                         orient="horizontal", variable=var,
                         length=160).grid(row=ri, column=1, sticky="w", padx=8)
            elif wt=="check":
                tk.Checkbutton(frm, variable=var).grid(row=ri, column=1, sticky="w", padx=8)

        col_names = [("Колір точок (biplot):","point_color"),
                     ("Колір стрілок:","arrow_color"),
                     ("Колір стовпців (scree):","bar_color"),
                     ("Колір кривої кумул.:","cum_color")]
        base_r = len(rows_cfg)
        for ri, (lbl, key) in enumerate(col_names):
            tk.Label(frm, text=lbl, font=rb_f).grid(row=base_r+ri, column=0, sticky="w", pady=4)
            btn = tk.Button(frm, width=6, relief=tk.SUNKEN, bg=col_refs[key])
            btn.grid(row=base_r+ri, column=1, sticky="w", padx=8)
            def _pick(k=key, b=btn, refs=col_refs):
                ch = colorchooser.askcolor(color=refs[k], parent=dlg)
                if ch and ch[1]: refs[k]=ch[1]; b.configure(bg=ch[1])
            btn.configure(command=_pick); col_btns[key] = btn

        def apply():
            self._pca_gs.update({
                "font_family": ff_v.get(), "font_size": fz_v.get(),
                "point_size":  ps_v.get(), "arrow_scale": sc_v.get(),
                "annotate_obj": ao_v.get(), "annotate_var": av_v.get(),
                "heatmap_cmap": cm_v.get(),
                **col_refs
            })
            dlg.destroy()
            if callback: callback()

        bf = tk.Frame(frm); bf.grid(row=base_r+len(col_names), column=0,
                                    columnspan=2, pady=(14,0))
        tk.Button(bf, text="OK", bg="#c62828", fg="white",
                  font=rb_f, command=apply).pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Скасувати", font=rb_f, command=dlg.destroy).pack(side=tk.LEFT)
        center_win(dlg)

    # ── Виконання PCA ─────────────────────────────────────────
    def _run(self):
        raw = [[e.get().strip() for e in row] for row in self.entries]
        raw = [r for r in raw if any(v for v in r)]
        if not raw:
            messagebox.showwarning("Немає даних","Введіть дані у таблицю."); return

        # Визначаємо чи перша колонка = мітки
        has_labels = False
        try: float(raw[0][0].replace(",",".")); has_labels = False
        except ValueError: has_labels = True

        start_col = 1 if has_labels else 0
        obj_names = []; data_rows = []
        for i, row in enumerate(raw):
            nm = row[0].strip() if has_labels else f"Об'єкт {i+1}"
            if not nm: nm = f"Об'єкт {i+1}"
            obj_names.append(nm)
            vals = []
            for v in row[start_col:]:
                if not v: continue
                try: vals.append(float(v.replace(",",".")))
                except Exception: continue
            if vals: data_rows.append(vals)
        obj_names = obj_names[:len(data_rows)]

        if len(data_rows) < 2:
            messagebox.showwarning("Замало об'єктів",
                "Потрібно щонайменше 2 об'єкти (рядки з числовими даними)."); return
        min_c = min(len(r) for r in data_rows)
        if min_c < 2:
            messagebox.showwarning("Замало показників",
                "Потрібно щонайменше 2 числові показники."); return

        # Назви змінних з заголовків
        var_names = []
        for j in range(start_col, self.cols_n):
            if j < len(self.header_vars):
                var_names.append(self.header_vars[j].get().strip() or f"П{j}")
            else:
                var_names.append(f"П{j}")
        var_names = var_names[:min_c]

        X = np.array([r[:min_c] for r in data_rows], dtype=float)

        # Стандартизація
        from scipy.stats import zscore
        X_std = zscore(X, axis=0, ddof=1); X_std = np.nan_to_num(X_std)

        # PCA через власні вектори
        cov_m = np.cov(X_std.T)
        eigenvalues, eigenvectors = np.linalg.eigh(cov_m)
        idx_s = np.argsort(eigenvalues)[::-1]
        eigenvalues = eigenvalues[idx_s]; eigenvectors = eigenvectors[:, idx_s]
        # Тільки невід'ємні власні значення
        eigenvalues = np.maximum(eigenvalues, 0)
        total_var = np.sum(eigenvalues)
        explained  = eigenvalues / total_var * 100 if total_var > 0 else eigenvalues*0
        scores     = X_std @ eigenvectors
        n_comp     = len(eigenvalues)

        if not HAS_MPL:
            messagebox.showwarning("","matplotlib недоступний."); return

        self._show_pca_results(obj_names, var_names, eigenvalues, eigenvectors,
                               explained, scores, n_comp, min_c)

    def _show_pca_results(self, obj_names, var_names, eigenvalues,
                          eigenvectors, explained, scores, n_comp, min_c):
        gs = self._pca_gs
        win = tk.Toplevel(self.win); win.title("PCA — Результати")
        win.geometry("1200x860"); set_icon(win)

        # ── Toolbar ─────────────────────────────────────────
        tb = tk.Frame(win, padx=6, pady=5); tb.pack(fill=tk.X)
        tk.Button(tb, text="📋 Копіювати графіки", font=("Times New Roman",11),
                  command=lambda: self._copy_pca()).pack(side=tk.LEFT, padx=4)
        tk.Button(tb, text="⚙ Налаштування графіків", font=("Times New Roman",11),
                  command=lambda: self._restyle_pca_live(
                      win, obj_names, var_names, eigenvalues, eigenvectors,
                      explained, scores, n_comp, min_c)).pack(side=tk.LEFT, padx=4)

        # ── Зона графіків (фіксована висота) ────────────────
        graph_frame = tk.Frame(win); graph_frame.pack(fill=tk.X)
        self._pca_main_frame = graph_frame

        # Графіки
        fig = Figure(figsize=(10, 6), dpi=100)

        ff  = gs["font_family"]; fz = gs["font_size"]
        pc  = gs["point_color"]; ac = gs["arrow_color"]
        bc  = gs["bar_color"];   cc = gs["cum_color"]
        ps  = gs["point_size"];  sc = gs["arrow_scale"]

        # ── Scree plot ───────────────────────────────────────
        ax1 = fig.add_subplot(131)
        ax1.bar(range(1, n_comp+1), explained[:n_comp], color=bc, alpha=0.8)
        ax1.plot(range(1, n_comp+1), np.cumsum(explained[:n_comp]),
                 "o-", color=cc, markersize=4)
        ax1.set_xlabel("ГК", fontsize=fz, fontfamily=ff)
        ax1.set_ylabel("Пояснена дисперсія (%)", fontsize=fz, fontfamily=ff)
        ax1.set_title("Графік відсіювання (Scree)", fontsize=fz+1, fontfamily=ff)
        ax1.axhline(80, color="gray", lw=0.8, ls="--")
        ax1.yaxis.grid(True, alpha=0.3)
        ax1.spines["top"].set_visible(False); ax1.spines["right"].set_visible(False)

        # ── Biplot ───────────────────────────────────────────
        ax2 = fig.add_subplot(132)
        ax2.scatter(scores[:,0], scores[:,1], s=ps, color=pc, zorder=3,
                    edgecolors="white", linewidths=0.5)
        if gs["annotate_obj"]:
            for i, nm in enumerate(obj_names[:len(scores)]):
                ax2.annotate(nm, (scores[i,0], scores[i,1]),
                             fontsize=max(6, fz-1), alpha=0.85, fontfamily=ff)
        # Стрілки навантажень
        max_score = max(np.max(np.abs(scores[:,0])), np.max(np.abs(scores[:,1])), 1e-6)
        for j in range(min_c):
            lx = eigenvectors[j,0] * max_score * sc
            ly = eigenvectors[j,1] * max_score * sc
            ax2.annotate("", xy=(lx,ly), xytext=(0,0),
                         arrowprops=dict(arrowstyle="->", color=ac, lw=1.3))
            if gs["annotate_var"]:
                nm_j = var_names[j] if j < len(var_names) else f"П{j+1}"
                ax2.text(lx*1.07, ly*1.07, nm_j,
                         fontsize=max(6, fz-1), color=ac, fontfamily=ff)
        ax2.axhline(0, color="#888", lw=0.5); ax2.axvline(0, color="#888", lw=0.5)
        ax2.set_xlabel(f"ГК1 ({fmt(explained[0],1)}%)", fontsize=fz, fontfamily=ff)
        ax2.set_ylabel(f"ГК2 ({fmt(explained[1],1)}%)" if n_comp>1 else "ГК2",
                       fontsize=fz, fontfamily=ff)
        ax2.set_title("Biplot (ГК1 × ГК2)", fontsize=fz+1, fontfamily=ff)
        ax2.spines["top"].set_visible(False); ax2.spines["right"].set_visible(False)

        # ── Теплова карта навантажень ─────────────────────────
        ax3 = fig.add_subplot(133)
        n_show = min(4, n_comp)
        load_mat = eigenvectors[:, :n_show]
        try:    cmap_ = matplotlib.cm.get_cmap(gs["heatmap_cmap"])
        except: cmap_ = matplotlib.cm.get_cmap("RdYlGn")
        im = ax3.imshow(load_mat, cmap=cmap_, vmin=-1, vmax=1, aspect="auto")
        ax3.set_xticks(range(n_show))
        ax3.set_xticklabels([f"ГК{i+1}" for i in range(n_show)],
                            fontsize=fz, fontfamily=ff)
        ax3.set_yticks(range(min_c))
        ax3.set_yticklabels(var_names[:min_c] if var_names else [f"П{j+1}" for j in range(min_c)],
                            fontsize=fz, fontfamily=ff)
        ax3.set_title("Навантаження факторів", fontsize=fz+1, fontfamily=ff)
        for i in range(min_c):
            for j in range(n_show):
                ax3.text(j, i, fmt(load_mat[i,j],2),
                         ha="center", va="center", fontsize=max(6, fz-1), fontfamily=ff)
        fig.colorbar(im, ax=ax3, fraction=0.046, pad=0.04)

        fig.tight_layout()
        self._pca_fig = fig
        self._pca_canvas_frame = tk.Frame(graph_frame)
        self._pca_canvas_frame.pack(fill=tk.X)
        embed_figure(fig, self._pca_canvas_frame)

        # ── Прокручувана текстова частина ────────────────────
        scroll_area = tk.Frame(win); scroll_area.pack(fill=tk.BOTH, expand=True)
        vsb = ttk.Scrollbar(scroll_area, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        txt_canvas = tk.Canvas(scroll_area, yscrollcommand=vsb.set, highlightthickness=0)
        txt_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.config(command=txt_canvas.yview)
        txt_body = tk.Frame(txt_canvas); txt_canvas.create_window((0,0), window=txt_body, anchor="nw")
        txt_body.bind("<Configure>",
                      lambda e: txt_canvas.configure(scrollregion=txt_canvas.bbox("all")))
        win.bind("<MouseWheel>",
                 lambda e: txt_canvas.yview_scroll(int(-1*(e.delta/120)),"units"))

        # ── Таблиця компонент ────────────────────────────────
        # Додаємо пояснення що таке ГК
        tk.Label(txt_body,
                 text="ГК = Головна компонента — «узагальнений показник» "
                      "що об'єднує кілька вихідних змінних. ГК1 пояснює найбільше варіації.",
                 font=("Times New Roman",10), fg="#555", anchor="w"
                 ).pack(fill=tk.X, padx=12, pady=(6,0))
        summary_rows = [[f"ГК{i+1}  (Головна компонента {i+1})",
                         fmt(eigenvalues[i],4),
                         fmt(explained[i],2),
                         fmt(float(np.sum(explained[:i+1])),2),
                         "✓ включити" if eigenvalues[i] >= 1.0 else "розглянути"]
                        for i in range(n_comp)]
        frm_t, _ = make_tv(txt_body,
            ["Компонент","Власне значення (λ)","% дисперсії","Кумулятивний %","Критерій Кайзера (λ≥1)"],
            summary_rows)
        frm_t.pack(fill=tk.X, padx=8, pady=4)

        # ── Таблиця навантажень ──────────────────────────────
        tk.Label(txt_body,
                 text="Навантаження (loadings): кореляція показника з кожною ГК. "
                      "|Навантаження| > 0.5 — значуща роль показника у цій компоненті.",
                 font=("Times New Roman",10), fg="#555", anchor="w"
                 ).pack(fill=tk.X, padx=12, pady=(0,0))
        n_show2 = min(6, n_comp)
        load_headers = ["Показник"] + [f"ГК{i+1}" for i in range(n_show2)]
        load_rows = []
        for j in range(min_c):
            nm_j = var_names[j] if j < len(var_names) else f"П{j+1}"
            load_rows.append([nm_j] + [fmt(eigenvectors[j,k],4) for k in range(n_show2)])
        frm_l, _ = make_tv(txt_body, load_headers, load_rows)
        frm_l.pack(fill=tk.X, padx=8, pady=(0,10))


    def _restyle_pca_live(self, win, obj_names, var_names, eigenvalues,
                           eigenvectors, explained, scores, n_comp, min_c):
        """Відкриває діалог налаштувань і одразу перебудовує графік."""
        self._restyle_pca(callback=lambda: self._rebuild_pca_fig(
            obj_names, var_names, eigenvalues, eigenvectors,
            explained, scores, n_comp, min_c))

    def _rebuild_pca_fig(self, obj_names, var_names, eigenvalues,
                          eigenvectors, explained, scores, n_comp, min_c):
        """Перебудовує лише графік у вже відкритому вікні результатів."""
        if not hasattr(self, '_pca_canvas_frame'): return
        for w in self._pca_canvas_frame.winfo_children(): w.destroy()
        gs = self._pca_gs
        ff=gs["font_family"]; fz=gs["font_size"]
        pc=gs["point_color"]; ac=gs["arrow_color"]
        bc=gs["bar_color"];   cc_=gs["cum_color"]
        ps=gs["point_size"];  sc=gs["arrow_scale"]
        fig = Figure(figsize=(10, 6), dpi=100)
        # Scree
        ax1=fig.add_subplot(131)
        ax1.bar(range(1,n_comp+1),explained[:n_comp],color=bc,alpha=0.8)
        ax1.plot(range(1,n_comp+1),np.cumsum(explained[:n_comp]),"o-",color=cc_,markersize=4)
        ax1.set_xlabel("ГК",fontsize=fz,fontfamily=ff)
        ax1.set_ylabel("Пояснена дисперсія (%)",fontsize=fz,fontfamily=ff)
        ax1.set_title("Графік відсіювання (Scree)",fontsize=fz+1,fontfamily=ff)
        ax1.axhline(80,color="gray",lw=0.8,ls="--")
        ax1.yaxis.grid(True,alpha=0.3)
        ax1.spines["top"].set_visible(False); ax1.spines["right"].set_visible(False)
        # Biplot
        ax2=fig.add_subplot(132)
        ax2.scatter(scores[:,0],scores[:,1],s=ps,color=pc,zorder=3,edgecolors="white",linewidths=0.5)
        if gs["annotate_obj"]:
            for i,nm in enumerate(obj_names[:len(scores)]):
                ax2.annotate(nm,(scores[i,0],scores[i,1]),fontsize=max(6,fz-1),alpha=0.85,fontfamily=ff)
        max_s=max(np.max(np.abs(scores[:,0])),np.max(np.abs(scores[:,1])),1e-6)
        for j in range(min_c):
            lx=eigenvectors[j,0]*max_s*sc; ly=eigenvectors[j,1]*max_s*sc
            ax2.annotate("",xy=(lx,ly),xytext=(0,0),
                         arrowprops=dict(arrowstyle="->",color=ac,lw=1.3))
            if gs["annotate_var"]:
                nm_j=var_names[j] if j<len(var_names) else f"П{j+1}"
                ax2.text(lx*1.07,ly*1.07,nm_j,fontsize=max(6,fz-1),color=ac,fontfamily=ff)
        ax2.axhline(0,color="#888",lw=0.5); ax2.axvline(0,color="#888",lw=0.5)
        ax2.set_xlabel(f"ГК1 ({fmt(explained[0],1)}%)",fontsize=fz,fontfamily=ff)
        ax2.set_ylabel(f"ГК2 ({fmt(explained[1],1)}%)" if n_comp>1 else "ГК2",fontsize=fz,fontfamily=ff)
        ax2.set_title("Biplot (ГК1 × ГК2)",fontsize=fz+1,fontfamily=ff)
        ax2.spines["top"].set_visible(False); ax2.spines["right"].set_visible(False)
        # Loadings heatmap
        ax3=fig.add_subplot(133)
        n_sh=min(4,n_comp)
        lm=eigenvectors[:,:n_sh]
        try: cmap_=matplotlib.cm.get_cmap(gs["heatmap_cmap"])
        except: cmap_=matplotlib.cm.get_cmap("RdYlGn")
        im=ax3.imshow(lm,cmap=cmap_,vmin=-1,vmax=1,aspect="auto")
        ax3.set_xticks(range(n_sh))
        ax3.set_xticklabels([f"ГК{i+1}" for i in range(n_sh)],fontsize=fz,fontfamily=ff)
        ax3.set_yticks(range(min_c))
        ax3.set_yticklabels(var_names[:min_c] if var_names else [f"П{j+1}" for j in range(min_c)],
                            fontsize=fz,fontfamily=ff)
        ax3.set_title("Навантаження факторів",fontsize=fz+1,fontfamily=ff)
        for i in range(min_c):
            for j in range(n_sh):
                ax3.text(j,i,fmt(lm[i,j],2),ha="center",va="center",fontsize=max(6,fz-1),fontfamily=ff)
        fig.colorbar(im,ax=ax3,fraction=0.046,pad=0.04)
        fig.tight_layout()
        self._pca_fig=fig
        embed_figure(fig, self._pca_canvas_frame)

    def _copy_pca(self):
        if self._pca_fig is None:
            messagebox.showwarning("","Спочатку виконайте PCA."); return
        ok, msg = _copy_fig_to_clipboard(self._pca_fig)
        if ok: messagebox.showinfo("","Графіки PCA скопійовано.\nВставте у Word через Ctrl+V.")
        else:   messagebox.showwarning("",f"Помилка: {msg}")




# ═══════════════════════════════════════════════════════════════
# REPEATED MEASURES ANOVA
# ═══════════════════════════════════════════════════════════════
class RepeatedMeasuresWindow:
    """Дисперсійний аналіз повторних вимірювань."""

    HELP_TEXT = """
АНАЛІЗ ПОВТОРНИХ ВИМІРЮВАНЬ — ПОКРОКОВА ІНСТРУКЦІЯ
══════════════════════════════════════════════════

ЩО ТАКЕ ПОВТОРНІ ВИМІРИ?
  Це ситуація коли одні й ті самі об'єкти (рослини, тварини, ділянки,
  дерева) вимірюються КІЛЬКА РАЗІВ:
  • У різні моменти часу (висота пагонів щомісяця)
  • За різних умов (доза А → доза Б → доза В)
  • По фазах вегетації (бутонізація, цвітіння, дозрівання)

  ВІДМІНА від звичайного ANOVA:
  Вимірювання одного об'єкта залежні між собою. Дерево що було
  вищим у травні — залишається відносно вищим у червні.
  Ігнорування цього зв'язку (звичайна ANOVA) завищує помилку
  і знижує шанс виявити реальну різницю.

КОЛИ ВИКОРИСТОВУВАТИ?
  ✓ Динаміка росту рослин або пагонів (вимірюєте ті самі рослини)
  ✓ Зміна вмісту поживних речовин по фазах вегетації
  ✓ Реакція на послідовні обробки (один об'єкт отримує всі обробки)
  ✓ Порівняння до/після (якщо > 2 точок; для 2 точок — парний t-тест)

КРОК 1. СТРУКТУРА ТАБЛИЦІ

  Перший стовпець: Назва суб'єкта (рослина, дерево, ділянка — текст)
  Решта стовпців: Вимірювання у кожній часовій точці (числа)
  Заголовки стовпців (сині): Назви часових точок або умов

  Приклад (ріст пагонів, 5 дерев, 4 вимірювання):
  | Дерево   | Травень | Червень | Липень | Серпень |
  | Дерево 1 |  12.3   |  24.5   |  38.1  |  45.2   |
  | Дерево 2 |  11.8   |  22.9   |  35.7  |  42.8   |
  | Дерево 3 |  13.1   |  26.2   |  40.3  |  48.6   |

  Перейменуйте заголовки часових точок:
  Подвійний клік на синій клітинці → введіть назву.

  Мінімум: 2 суб'єкти, 2 часові точки.

КРОК 2. ВИКОНАННЯ АНАЛІЗУ
  Натисніть «▶ Аналіз».
  Програма автоматично видаляє рядки з пропущеними даними.

КРОК 3. ІНТЕРПРЕТАЦІЯ РЕЗУЛЬТАТІВ

  Таблиця дисперсійного аналізу:

  SS_час — варіація пояснена зміною у часі (те що нас цікавить)
  SS_суб'єкти — варіація між суб'єктами (виноситься окремо!)
  SS_похибка — залишкова варіація

  F(df_час, df_похибка):
    p < α → є значуща динаміка у часі ✓
    p ≥ α → динаміка незначуща

  Partial η² (розмір ефекту):
    < 0.01: дуже слабкий | 0.01–0.06: слабкий
    0.06–0.14: середній  | > 0.14: сильний ← типово для росту

КРОК 4. POST-HOC АНАЛІЗ (після значущого F)

  Виконуються парні t-тести з поправкою Бонферроні.
  Показують ЯКІ САМЕ пари часових точок відрізняються.

  Приклад: «Травень vs Червень: p=0.003 *» → у червні показник
  значущо вищий ніж у травні.

КРОК 5. НОРМАЛЬНІСТЬ РІЗНИЦЬ

  Перевіряється Shapiro-Wilk для різниць між кожною парою точок.
  p > 0.05 → різниці нормальні → результати надійні ✓
  p ≤ 0.05 → розгляньте тест Фрідмана (непараметричний аналог)

КРОК 6. ГРАФІК ДИНАМІКИ (Середні ± СП)

  Показує як середнє значення змінюється у часі.
  Смужки похибок (СП) = стандартна похибка середнього.
  Чим менші смужки → тим точніше визначено середнє.
  S-подібний підйом → типовий ріст рослин.
  Плато → насичення (ріст сповільнився або зупинився).

ПОРАДА:
  Якщо у вас КІЛЬКА ВАРІАНТІВ (сортів, обробок) і ті самі
  об'єкти щороку — це двофакторний Repeated Measures:
  між-суб'єктний фактор (варіант) + всередині-суб'єктний (час).
  Для такого аналізу використовуйте ТРИФАКТОРНУ ANOVA де рік
  є одним з факторів — це загальноприйнята практика.
"""

    def __init__(self, parent, gs):
        self.win = tk.Toplevel(parent)
        self.win.title("Дисперсійний аналіз повторних вимірювань")
        self.win.geometry("940x660"); set_icon(self.win)
        self.gs = gs
        self._rm_fig = None
        self._rm_gs  = {
            "font_family":  "Times New Roman",
            "font_size":    10,
            "line_color":   "#4c72b0",
            "err_color":    "#c62828",
            "marker":       "o",
            "linewidth":    2.0,
            "markersize":   7,
            "show_grid":    True,
        }
        self._build()

    def _build(self):
        try:
            self._build_inner()
        except Exception as e:
            import traceback
            messagebox.showerror("Помилка ініціалізації",
                f"Помилка при побудові вікна:\n{traceback.format_exc()}")

    def _build_inner(self):
        # ── Toolbar ──────────────────────────────────────────
        top = tk.Frame(self.win, padx=8, pady=6); top.pack(fill=tk.X)
        tk.Button(top, text="▶ Аналіз", bg="#c62828", fg="white",
                  font=("Times New Roman", 13),
                  command=self._run).pack(side=tk.LEFT, padx=4)

        # Налаштування — спадне меню
        mb2 = tk.Menubutton(top, text="⚙ Налаштування ▾",
                            font=("Times New Roman",11), relief=tk.RAISED, bd=2)
        mb2.pack(side=tk.LEFT, padx=4)
        sm = tk.Menu(mb2, tearoff=0)
        sm.add_command(label="Додати рядок",       command=self._add_row)
        sm.add_command(label="Видалити рядок",     command=self._del_row)
        sm.add_separator()
        sm.add_command(label="Додати стовпець",    command=self._add_col)
        sm.add_command(label="Видалити стовпець",  command=self._del_col)
        sm.add_separator()
        sm.add_command(label="💾 Зберегти проект", command=self._save_proj)
        sm.add_command(label="📂 Відкрити проект", command=self._load_proj)
        sm.add_separator()
        sm.add_command(label="🗑 Очистити таблицю", command=self._clear_table)
        mb2["menu"] = sm

        tk.Button(top, text="Вставити з буфера",
                  font=("Times New Roman",11),
                  command=self._paste).pack(side=tk.LEFT, padx=4)
        tk.Button(top, text="📚 Довідка", bg="#1a4b8c", fg="white",
                  font=("Times New Roman",11),
                  command=self._show_help).pack(side=tk.LEFT, padx=4)

        tk.Label(top,
                 text="Подвійний клік на заголовку часової точки → перейменувати",
                 font=("Times New Roman",9), fg="#666").pack(side=tk.LEFT, padx=8)

        # ── Таблиця ─────────────────────────────────────────
        mid = tk.Frame(self.win); mid.pack(fill=tk.BOTH, expand=True, padx=8)
        self.rows_n = 20; self.cols_n = 8
        self._canvas = tk.Canvas(mid)
        self._canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb = ttk.Scrollbar(mid, orient="vertical", command=self._canvas.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        self._canvas.configure(yscrollcommand=sb.set)
        self.inner = tk.Frame(self._canvas)
        self._canvas.create_window((0,0), window=self.inner, anchor="nw")
        self.inner.bind("<Configure>",
                        lambda e: self._canvas.config(scrollregion=self._canvas.bbox("all")))
        self.win.bind("<MouseWheel>",
                      lambda e: self._canvas.yview_scroll(int(-1*(e.delta/120)),"units"))

        # Перший заголовок — «Суб'єкт» (фіксований)
        tk.Label(self.inner, text="Суб'єкт", relief=tk.RIDGE, width=13,
                 bg="#444444", fg="white",
                 font=("Times New Roman",11,"bold")).grid(row=0, column=0, padx=1, pady=1, sticky="nsew")

        # Заголовки часових точок (перейменовувані)
        self.col_vars = []; self.col_labels = []
        for j in range(1, self.cols_n):
            var = tk.StringVar(value=f"Точка {j}")
            self.col_vars.append(var)
            lbl = tk.Label(self.inner, textvariable=var, width=12, cursor="hand2",
                           bg="#1a4b8c", fg="white", relief=tk.RIDGE,
                           font=("Times New Roman",11,"bold"))
            lbl.grid(row=0, column=j, padx=1, pady=1, sticky="nsew")
            lbl.bind("<Double-Button-1>", lambda e, idx=j-1: self._rename_time_col(idx))
            self.col_labels.append(lbl)

        self.entries = []
        for i in range(self.rows_n):
            row_ = []
            for j in range(self.cols_n):
                e = tk.Entry(self.inner, width=13 if j==0 else 12,
                             font=("Times New Roman",11))
                e.grid(row=i+1, column=j, padx=1, pady=1)
                row_.append(e)
            self.entries.append(row_)
        _bind_nav(self.entries, self.win)

    # ── Перейменування часової точки ─────────────────────────
    def _rename_time_col(self, idx):
        dlg = tk.Toplevel(self.win); dlg.title("Перейменувати часову точку")
        dlg.resizable(False, False); dlg.grab_set()
        tk.Label(dlg, text=f"Назва часової точки {idx+1}:",
                 font=("Times New Roman",12)).pack(padx=16, pady=14)
        var = tk.StringVar(value=self.col_vars[idx].get())
        e = tk.Entry(dlg, textvariable=var, font=("Times New Roman",12), width=26)
        e.pack(padx=16, pady=4); e.select_range(0, tk.END); e.focus_set()
        def apply():
            nm = var.get().strip()
            if nm: self.col_vars[idx].set(nm)
            dlg.destroy()
        tk.Button(dlg, text="OK", bg="#c62828", fg="white",
                  font=("Times New Roman",12), command=apply).pack(pady=(4,14))
        dlg.bind("<Return>", lambda ev: apply()); center_win(dlg)

    # ── Управління таблицею ───────────────────────────────────
    def _add_row(self):
        i = self.rows_n; row_ = []
        for j in range(self.cols_n):
            e = tk.Entry(self.inner, width=13 if j==0 else 12,
                         font=("Times New Roman",11))
            e.grid(row=i+1, column=j, padx=1, pady=1)
            row_.append(e)
        self.entries.append(row_); self.rows_n += 1
        _bind_nav(self.entries, self.win)

    def _del_row(self):
        if not self.entries: return
        for e in self.entries.pop(): e.destroy()
        self.rows_n -= 1

    def _add_col(self):
        ci = self.cols_n; self.cols_n += 1
        var = tk.StringVar(value=f"Точка {ci}")
        self.col_vars.append(var)
        lbl = tk.Label(self.inner, textvariable=var, width=12, cursor="hand2",
                       bg="#1a4b8c", fg="white", relief=tk.RIDGE,
                       font=("Times New Roman",11,"bold"))
        lbl.grid(row=0, column=ci, padx=1, pady=1, sticky="nsew")
        lbl.bind("<Double-Button-1>", lambda e, idx=ci-1: self._rename_time_col(idx))
        self.col_labels.append(lbl)
        for i, row_ in enumerate(self.entries):
            e = tk.Entry(self.inner, width=12, font=("Times New Roman",11),
                         )
            e.grid(row=i+1, column=ci, padx=1, pady=1)
            row_.append(e)
        _bind_nav(self.entries, self.win)

    def _del_col(self):
        if self.cols_n <= 3: return
        self.col_labels.pop().destroy(); self.col_vars.pop()
        for row_ in self.entries: row_.pop().destroy()
        self.cols_n -= 1

    def _clear_table(self):
        if not messagebox.askyesno("Очистити",
                "Видалити всі дані?\n(Заголовки залишаться)"): return
        for row in self.entries:
            for e in row: e.delete(0, tk.END)

    def _save_proj(self):
        path = filedialog.asksaveasfilename(
            parent=self.win, defaultextension=".sadp",
            filetypes=[("SAD проект","*.sadp"),("JSON","*.json")],
            title="Зберегти проект повторних вимірювань")
        if not path: return
        d = {
            "type": "repeated_measures", "version": APP_VER,
            "col_vars": [v.get() for v in self.col_vars],
            "rows_data": [[e.get() for e in row] for row in self.entries],
        }
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(d, f, ensure_ascii=False, indent=2)
            messagebox.showinfo("Збережено",
                f"Проект збережено:\n{path}\n\n"
                "Щоб додати нові дати: ⚙ → «Додати стовпець»")
        except Exception as ex:
            messagebox.showerror("Помилка збереження", str(ex))

    def _load_proj(self):
        path = filedialog.askopenfilename(
            parent=self.win,
            filetypes=[("SAD проект","*.sadp"),("JSON","*.json")],
            title="Відкрити проект повторних вимірювань")
        if not path: return
        try:
            with open(path, "r", encoding="utf-8") as f:
                d = json.load(f)
        except Exception as ex:
            messagebox.showerror("Помилка відкриття", str(ex)); return
        col_vars = d.get("col_vars", [])
        rows_data = d.get("rows_data", [])
        n_cols_needed = 1 + len(col_vars)
        while self.cols_n < n_cols_needed: self._add_col()
        for i, nm in enumerate(col_vars):
            if i < len(self.col_vars): self.col_vars[i].set(nm)
        while len(self.entries) < len(rows_data): self._add_row()
        for i, row_vals in enumerate(rows_data):
            for j, v in enumerate(row_vals):
                if j < self.cols_n:
                    self.entries[i][j].delete(0, tk.END)
                    self.entries[i][j].insert(0, v)
        messagebox.showinfo("Завантажено",
            "Проект завантажено.\n\n"
            "Щоб додати нові дати: ⚙ → «Додати стовпець»\n"
            "і перейменуйте заголовок подвійним кліком.")

    def _paste(self):
        try: data = self.win.clipboard_get()
        except Exception:
            messagebox.showwarning("Буфер порожній",
                "Скопіюйте дані з Excel (Ctrl+C) і спробуйте знову."); return
        if not data.strip(): return
        pos = (0, 0)
        w = self.win.focus_get()
        if isinstance(w, tk.Entry):
            for i, row_ in enumerate(self.entries):
                for j, e in enumerate(row_):
                    if e is w: pos=(i,j); break
        r0, c0 = pos
        for ir, line in enumerate(data.splitlines()):
            if not line.strip(): continue
            while r0+ir >= len(self.entries): self._add_row()
            for jc, val in enumerate(line.split("\t")):
                cc = c0+jc
                if cc >= self.cols_n: continue
                self.entries[r0+ir][cc].delete(0,tk.END)
                self.entries[r0+ir][cc].insert(0, val.strip())

    def _show_help(self):
        win = tk.Toplevel(self.win)
        win.title("Довідка — Повторні виміри ANOVA")
        win.geometry("720x680"); set_icon(win)
        frm = tk.Frame(win); frm.pack(fill=tk.BOTH, expand=True, padx=8, pady=6)
        vsb = ttk.Scrollbar(frm, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        txt = tk.Text(frm, wrap="word", font=("Times New Roman",11),
                      yscrollcommand=vsb.set, relief=tk.FLAT,
                      bg="#fafafa", padx=10, pady=8, cursor="arrow")
        txt.pack(fill=tk.BOTH, expand=True); vsb.config(command=txt.yview)
        txt.insert("1.0", self.HELP_TEXT.strip()); txt.configure(state="disabled")
        txt.bind("<MouseWheel>",
                 lambda e: txt.yview_scroll(int(-1*(e.delta/120)),"units"))
        tk.Button(win, text="Закрити", command=win.destroy,
                  font=("Times New Roman",11)).pack(pady=6)

    # ── Налаштування графіка ──────────────────────────────────
    def _restyle_rm(self, win, time_names, data_arr, n, ph_results=None, alpha=0.05):
        dlg = tk.Toplevel(win); dlg.title("Налаштування графіка")
        dlg.resizable(False, False); set_icon(dlg); dlg.grab_set()
        gs = self._rm_gs
        frm = tk.Frame(dlg, padx=16, pady=14); frm.pack()
        rb_f = ("Times New Roman",12)

        ff_v  = tk.StringVar(value=gs["font_family"])
        fz_v  = tk.IntVar(value=gs["font_size"])
        lw_v  = tk.DoubleVar(value=gs["linewidth"])
        ms_v  = tk.IntVar(value=gs["markersize"])
        mk_v  = tk.StringVar(value=gs["marker"])
        gr_v  = tk.BooleanVar(value=gs["show_grid"])
        lc_ref = [gs["line_color"]]; ec_ref = [gs["err_color"]]

        rows_cfg = [
            ("Шрифт:",          "combo",  ff_v, ["Times New Roman","Arial","Calibri","Georgia"]),
            ("Розмір шрифту:",  "spin",   fz_v, (7,18)),
            ("Товщина лінії:",  "scale",  lw_v, (0.5,5.0)),
            ("Розмір маркера:", "spin",   ms_v, (3,20)),
            ("Тип маркера:",    "combo",  mk_v, ["o","s","^","D","v","*","+"]),
            ("Показати сітку:", "check",  gr_v, None),
        ]
        for ri, (lbl, wt, var, opts) in enumerate(rows_cfg):
            tk.Label(frm, text=lbl, font=rb_f).grid(row=ri, column=0, sticky="w", pady=4)
            if wt=="combo":
                ttk.Combobox(frm, textvariable=var, values=opts,
                             state="readonly", width=18).grid(row=ri, column=1, sticky="w", padx=8)
            elif wt=="spin":
                tk.Spinbox(frm, from_=opts[0], to=opts[1], textvariable=var,
                           width=7).grid(row=ri, column=1, sticky="w", padx=8)
            elif wt=="scale":
                tk.Scale(frm, from_=opts[0], to=opts[1], resolution=0.1,
                         orient="horizontal", variable=var,
                         length=160).grid(row=ri, column=1, sticky="w", padx=8)
            elif wt=="check":
                tk.Checkbutton(frm, variable=var).grid(row=ri, column=1, sticky="w", padx=8)

        base_r = len(rows_cfg)
        btn_lc = tk.Button(frm, width=6, relief=tk.SUNKEN, bg=lc_ref[0])
        btn_ec = tk.Button(frm, width=6, relief=tk.SUNKEN, bg=ec_ref[0])
        for ri2, (lbl, ref, btn) in enumerate([
            ("Колір лінії:", lc_ref, btn_lc),
            ("Колір смужок похибок:", ec_ref, btn_ec)
        ]):
            tk.Label(frm, text=lbl, font=rb_f).grid(row=base_r+ri2, column=0, sticky="w", pady=4)
            btn.grid(row=base_r+ri2, column=1, sticky="w", padx=8)
            def _pick(r=ref, b=btn):
                ch = colorchooser.askcolor(color=r[0], parent=dlg)
                if ch and ch[1]: r[0]=ch[1]; b.configure(bg=ch[1])
            btn.configure(command=_pick)

        def apply():
            self._rm_gs.update({
                "font_family": ff_v.get(), "font_size": fz_v.get(),
                "linewidth":   lw_v.get(), "markersize": ms_v.get(),
                "marker":      mk_v.get(), "show_grid": gr_v.get(),
                "line_color":  lc_ref[0],  "err_color": ec_ref[0],
            })
            dlg.destroy()
            self._redraw_rm(win, time_names, data_arr, n, ph_results, alpha)

        bf = tk.Frame(frm); bf.grid(row=base_r+2, column=0, columnspan=2, pady=(14,0))
        tk.Button(bf, text="OK (застосувати)", bg="#c62828", fg="white",
                  font=rb_f, command=apply).pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Скасувати", font=rb_f, command=dlg.destroy).pack(side=tk.LEFT)
        center_win(dlg)

    # ── Виконання аналізу ─────────────────────────────────────
    def _run(self):
        # Назви часових точок — лише заповнені стовпці
        all_time = [v.get().strip() or f"Т{i+1}" for i, v in enumerate(self.col_vars)]
        raw = [[e.get().strip() for e in row] for row in self.entries]
        raw = [r for r in raw if any(v for v in r)]
        if not raw:
            messagebox.showwarning("Немає даних","Введіть дані у таблицю."); return

        # Знаходимо реально заповнені числові стовпці (1..cols_n-1)
        filled_t_cols = []
        for j in range(1, self.cols_n):
            for row in raw:
                v = row[j] if j < len(row) else ""
                if v:
                    try: float(v.replace(",",".")); filled_t_cols.append(j); break
                    except ValueError: pass

        if not filled_t_cols:
            messagebox.showwarning("Немає числових даних",
                "Введіть числові значення у стовпці часових точок (2 і далі)."); return

        time_names = [all_time[j-1] for j in filled_t_cols]
        k = len(time_names)

        subjects = []; data_rows = []
        for row in raw:
            subj = row[0].strip() if row and row[0].strip() else f"Суб'єкт {len(subjects)+1}"
            try: float(subj.replace(",",".")); continue
            except ValueError: pass
            vals = []
            for j in filled_t_cols:
                v = row[j].strip() if j < len(row) else ""
                if not v: vals.append(float("nan"))
                else:
                    try: vals.append(float(v.replace(",",".")))
                    except Exception: vals.append(float("nan"))
            if any(not math.isnan(v) for v in vals):
                subjects.append(subj); data_rows.append(vals)

        if len(data_rows) < 2:
            messagebox.showwarning("Замало суб'єктів",
                "Потрібно щонайменше 2 суб'єкти.\n\n"
                "Переконайтесь що:\n"
                "  • Перший стовпець містить назви (текст)\n"
                "  • Числа введені у стовпці 2 і далі"); return

        n_raw = len(data_rows)
        data = np.array(data_rows, dtype=float)
        mask_complete = ~np.any(np.isnan(data), axis=1)
        n_incomplete = int(np.sum(~mask_complete))
        data = data[mask_complete]
        subjects = [s for s, m in zip(subjects, mask_complete) if m]
        n = len(data)
        if n < 2:
            messagebox.showwarning("Замало повних даних",
                f"Після видалення рядків з пропущеними значеннями залишилось {n}.\n"
                "Заповніть пропущені значення або видаліть неповні рядки."); return
        if n_incomplete > 0:
            messagebox.showinfo("Пропущені дані",
                f"Видалено {n_incomplete} рядків з пропущеними значеннями.\n"
                f"Аналіз на {n} суб'єктах.")

        grand_mean = np.mean(data)
        subj_means = np.mean(data, axis=1)
        time_means = np.mean(data, axis=0)
        SS_total = float(np.sum((data - grand_mean)**2))
        SS_subj  = k * float(np.sum((subj_means - grand_mean)**2))
        SS_time  = n * float(np.sum((time_means - grand_mean)**2))
        SS_error = SS_total - SS_subj - SS_time
        df_time = k-1; df_subj = n-1; df_err = (k-1)*(n-1)
        MS_time = SS_time/df_time if df_time > 0 else float("nan")
        MS_err  = SS_error/df_err if df_err > 0 else float("nan")
        F = MS_time/MS_err if (not math.isnan(MS_err) and MS_err > 1e-12) else float("nan")
        p = float(1 - f_dist.cdf(F, df_time, df_err)) if not math.isnan(F) else float("nan")
        eta2_time = SS_time/(SS_time+SS_error) if (SS_time+SS_error) > 0 else float("nan")
        R2 = (SS_time+SS_subj)/SS_total if SS_total > 0 else float("nan")
        alpha = ALPHA

        sw_ps = []
        for j in range(k):
            for jj in range(j+1, k):
                diff = data[:,j] - data[:,jj]
                try: _, p_sw = shapiro(diff)
                except Exception: p_sw = float("nan")
                sw_ps.append(p_sw)
        min_sw = min((pp for pp in sw_ps if not math.isnan(pp)), default=float("nan"))
        norm_ok = not math.isnan(min_sw) and min_sw > 0.05

        from scipy.stats import ttest_rel
        ph_results = {}
        mt = k*(k-1)/2 if k > 1 else 1
        for j in range(k):
            for jj in range(j+1, k):
                st_, p_t_ = ttest_rel(data[:,j], data[:,jj])
                p_adj_ = min(1., float(p_t_)*mt)
                ph_results[(j,jj)] = (float(np.mean(data[:,j]-data[:,jj])), float(st_), p_adj_)

        if not HAS_MPL: messagebox.showwarning("","matplotlib недоступний."); return

        win = tk.Toplevel(self.win)
        win.title("Повторні виміри — Результати")
        win.geometry("1020x760"); set_icon(win)

        tb = tk.Frame(win, padx=6, pady=5); tb.pack(fill=tk.X)
        tk.Button(tb, text="📋 Копіювати графік", font=("Times New Roman",11),
                  command=lambda: self._copy_rm()).pack(side=tk.LEFT, padx=4)
        tk.Button(tb, text="⚙ Налаштування графіка", font=("Times New Roman",11),
                  command=lambda: self._restyle_rm(win, time_names, data, n, ph_results, alpha)
                  ).pack(side=tk.LEFT, padx=4)

        scroll_area = tk.Frame(win); scroll_area.pack(fill=tk.BOTH, expand=True)
        vsb = ttk.Scrollbar(scroll_area, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        sc = tk.Canvas(scroll_area, yscrollcommand=vsb.set, highlightthickness=0)
        sc.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.config(command=sc.yview)
        body = tk.Frame(sc); sc.create_window((0,0), window=body, anchor="nw")
        body.bind("<Configure>", lambda e: sc.configure(scrollregion=sc.bbox("all")))
        win.bind("<MouseWheel>", lambda e: sc.yview_scroll(int(-1*(e.delta/120)),"units"))

        def _head(txt):
            tk.Label(body, text=txt, font=("Times New Roman",12,"bold"),
                     bg="#e8eeff", anchor="w", padx=8, pady=3).pack(fill=tk.X, padx=6, pady=8)
        def _txt(txt, color="#000000"):
            tk.Label(body, text=txt, font=("Times New Roman",11), fg=color,
                     anchor="w", justify="left").pack(fill=tk.X, padx=14, pady=1)

        _head("Дисперсійний аналіз повторних вимірювань")
        _txt(f"Суб'єктів (n): {n}   |   Часових точок (k): {k}   |   α = {alpha}")
        _head("Таблиця дисперсійного аналізу")
        anova_rows = [
            ["Час (within)",fmt(SS_time,4),str(df_time),fmt(MS_time,4),fmt(F,4),fmt(p,4),
             ("✓ значущий" if p < alpha else "✗ незначущий") if not math.isnan(p) else "–"],
            ["Суб'єкти",fmt(SS_subj,4),str(df_subj),"–","–","–","виноситься окремо"],
            ["Похибка",fmt(SS_error,4),str(df_err),fmt(MS_err,4),"–","–",""],
            ["Загальна",fmt(SS_total,4),str(df_time+df_subj+df_err),"–","–","–",""],
        ]
        frm_a, _ = make_tv(body, ["Джерело","SS","df","MS","F","p","Висновок"], anova_rows)
        frm_a.pack(fill=tk.X, padx=8, pady=(2,6))
        _head("Показники якості")
        sig_c = "#1a6b1a" if not math.isnan(p) and p < alpha else "#c62828"
        _txt(f"F({df_time},{df_err}) = {fmt(F,4)},  p = {fmt(p,4)}  "
             f"{'✓ значуща динаміка' if not math.isnan(p) and p < alpha else '✗ динаміка незначуща'}", sig_c)
        _txt(f"Partial η² (час) = {fmt(eta2_time,4)}  →  {eta2_label(eta2_time)}")
        _txt(f"R² = {fmt(R2,4)}  (час + суб'єкти)")
        _txt(f"Shapiro–Wilk (різниці) мін. p = {fmt(min_sw,4)}  "
             f"{'✓ нормальні' if norm_ok else '⚠ ненормальні → розгляньте тест Фрідмана'}",
             "#000" if norm_ok else "#c62828")
        _head("Середні по часових точках")
        means_tbl = [[time_names[j], fmt(float(np.mean(data[:,j])),4),
                      fmt(float(np.std(data[:,j],ddof=1)),4),
                      fmt(float(np.std(data[:,j],ddof=1)/math.sqrt(n)),4)] for j in range(k)]
        frm_m, _ = make_tv(body, ["Часова точка","Середнє","SD","СП (SE)"], means_tbl)
        frm_m.pack(fill=tk.X, padx=8, pady=(2,6))

        # Графік — власний Frame всередині body
        self._rm_graph_frame = tk.Frame(body)
        self._rm_graph_frame.pack(fill=tk.X, padx=8, pady=6)
        self._redraw_rm(win, time_names, data, n, ph_results, alpha)

        if not math.isnan(p) and p < alpha:
            _head(f"Пост-хок порівняння (Бонферроні)")
            _txt(f"Скоригований α = {fmt(alpha,2)} / {int(mt)} пар = {fmt(alpha/mt,4)}   "
                 f"│   * p < {alpha}   │   ** p < {alpha*0.2:.3f}","#555")
            ph_rows = []
            for j in range(k):
                for jj in range(j+1, k):
                    d_, st_, pa_ = ph_results[(j,jj)]
                    mark = "**" if pa_<alpha*0.2 else ("*" if pa_<alpha else "–")
                    ph_rows.append([f"{time_names[j]} vs {time_names[jj]}",
                                    fmt(d_,4), fmt(st_,4), fmt(pa_,4), mark])
            frm_ph, _ = make_tv(body, ["Пара","Різниця","t","p (Bonf.)","Знач."], ph_rows)
            frm_ph.pack(fill=tk.X, padx=8, pady=(2,4))
            _txt(f"* — p < α={alpha} (значуща різниця)   "
                 f"** — p < {alpha*0.2:.3f} (висока значущість)   – — незначуща","#555")
        else:
            _txt("Post-hoc аналіз не виконується при незначущому F-тесті.","#888")

    def _redraw_rm(self, win, time_names, data_arr, n, ph_results=None, alpha=0.05):
        if not hasattr(self,"_rm_graph_frame"): return
        for w in self._rm_graph_frame.winfo_children(): w.destroy()
        gs = self._rm_gs; k = len(time_names)
        fig = Figure(figsize=(10, 6), dpi=100)
        ax  = fig.add_subplot(111)
        means_ = np.mean(data_arr, axis=0)
        ses_   = np.std(data_arr, axis=0, ddof=1) / math.sqrt(n)
        ax.errorbar(range(k), means_, yerr=ses_,
                    fmt=gs["marker"]+"-", capsize=5,
                    color=gs["line_color"], ecolor=gs["err_color"],
                    linewidth=gs["linewidth"], markersize=gs["markersize"], zorder=3)
        if ph_results:
            y_range = float(np.max(means_)-np.min(means_)) if np.max(means_)>np.min(means_) else 1.
            offset = y_range*0.06
            for jj in range(1, k):
                j = jj-1; pk = (j,jj)
                if pk in ph_results and ph_results[pk][2] < alpha:
                    pa = ph_results[pk][2]
                    mark = "**" if pa<alpha*0.2 else "*"
                    xm = (j+jj)/2
                    yb = max(means_[j]+ses_[j], means_[jj]+ses_[jj])+offset
                    ax.plot([j,jj],[yb,yb],color="#555",lw=0.8)
                    ax.plot([j,j],[yb-offset*0.3,yb],color="#555",lw=0.8)
                    ax.plot([jj,jj],[yb-offset*0.3,yb],color="#555",lw=0.8)
                    ax.text(xm, yb+offset*0.1, mark, ha="center", va="bottom",
                            fontsize=gs["font_size"]+1, color="#c62828",
                            fontfamily=gs["font_family"])
        ax.set_xticks(range(k))
        ax.set_xticklabels(time_names, fontsize=gs["font_size"], fontfamily=gs["font_family"])
        ax.set_xlabel("Часова точка / Умова", fontsize=gs["font_size"], fontfamily=gs["font_family"])
        ax.set_ylabel("Середнє ± СП",         fontsize=gs["font_size"], fontfamily=gs["font_family"])
        ax.set_title("Динаміка середніх (Середнє ± СП)",
                     fontsize=gs["font_size"]+1, fontfamily=gs["font_family"])
        if gs["show_grid"]: ax.yaxis.grid(True, linestyle="--", alpha=0.35)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        if ph_results and any(v[2]<alpha for v in ph_results.values()):
            ax.annotate(f"* p<{alpha},  ** p<{alpha*0.2:.3f}  (Бонферроні, між сусідніми точками)",
                        xy=(0.01,0.01), xycoords="axes fraction",
                        fontsize=max(7,gs["font_size"]-1), color="#555", fontfamily=gs["font_family"])
        fig.tight_layout()
        self._rm_fig = fig
        embed_figure(fig, self._rm_graph_frame)


    def _copy_rm(self):
        if self._rm_fig is None:
            messagebox.showwarning("","Спочатку виконайте аналіз."); return
        ok, msg = _copy_fig_to_clipboard(self._rm_fig)
        if ok: messagebox.showinfo("","Графік скопійовано.\nВставте у Word через Ctrl+V.")
        else:   messagebox.showwarning("",f"Помилка: {msg}")





# ═══════════════════════════════════════════════════════════════
# MIXED REPEATED MEASURES  (Split-plot у часі)
# Варіант (between) × Дата (within), кілька повторностей
# ═══════════════════════════════════════════════════════════════
class MixedRepeatedWindow:
    """Змішаний дисперсійний аналіз повторних вимірювань.
    Between-subjects фактор: Варіант/Доза/Сорт (різний між повторностями).
    Within-subjects фактор:  Дата/Час (вимірюється для кожної повторності).
    """

    HELP_TEXT = """
ЗМІШАНИЙ АНАЛІЗ ПОВТОРНИХ ВИМІРЮВАНЬ — ІНСТРУКЦІЯ
══════════════════════════════════════════════════

ДЛЯ ЧОГО ЦЕЙ АНАЛІЗ?
  Коли у вас є КІЛЬКА ВАРІАНТІВ (дози добрива, сорти, обробки)
  і кожен варіант вимірюється КІЛЬКА РАЗІВ у часі.

  Відповідає на три питання одночасно:
  1. Чи відрізняються варіанти між собою? (ефект Варіанту)
  2. Чи є динаміка у часі? (ефект Часу)
  3. Чи по-різному змінюються варіанти у часі?
     (взаємодія Варіант × Час) ← НАЙВАЖЛИВІШЕ!

  Приклад: 4 дози добрива, вимірювання пагонів 4 рази.
  Взаємодія значуща → ефект дози залежить від дати вимірювання.

КРОК 1. ПІДГОТОВКА ДАНИХ

  Якщо у вас кілька дерев/рослин у кожній повторності:
  Спочатку порахуйте СЕРЕДНЄ по деревах для кожної повторності.
  Наприклад: 5 дерев у повторності → одне середнє значення.

  Кожне значення у таблиці = середнє по рослинах повторності.

КРОК 2. СТРУКТУРА ТАБЛИЦІ

  Стовпець 1: Варіант (текст: «Доза 1», «Контроль» тощо)
  Стовпець 2: Повторність (текст або число: «Повт.1», «1» тощо)
  Решта стовпців: Значення по датах/часових точках (числа)

  Синій рядок заголовків = назви дат/часових точок.
  Подвійний клік → перейменувати.

  Приклад (4 варіанти, 4 повторності, 4 дати):
  | Варіант | Повт. | 1.06 | 8.06 | 12.06 | 17.06 |
  | Доза 1  |   1   | 10.2 | 13.4 |  15.8 |  17.9 |
  | Доза 1  |   2   |  9.8 | 12.9 |  15.2 |  17.4 |
  | Доза 1  |   3   | 10.6 | 13.8 |  16.1 |  18.2 |
  | Доза 1  |   4   | 10.1 | 13.1 |  15.5 |  17.7 |
  | Доза 2  |   1   | 11.5 | 14.8 |  17.2 |  19.6 |
  ...

КРОК 3. СТАТИСТИЧНА МОДЕЛЬ

  Це Split-plot у часі. Дві різні помилки:

  Whole-plot error = Варіант × Повторність (помилка між групами)
    Використовується для тесту ефекту Варіанту.

  Sub-plot error = залишок (помилка всередині груп)
    Використовується для тесту Часу і Взаємодії.

  ⚠ Якщо не враховувати цю подвійну структуру — F-значення
  для Варіанту будуть хибними (занижена помилка).

КРОК 4. ІНТЕРПРЕТАЦІЯ

  Ефект Варіанту:
    p < α → варіанти значущо відрізняються (загалом по всіх датах)

  Ефект Часу:
    p < α → є значуща динаміка у часі (загалом по всіх варіантах)

  Взаємодія Варіант × Час:
    p < α → РІЗНА динаміка у різних варіантів!
      Лінії на графіку розходяться або перетинаються.
      Ефект варіанту залежить від дати → аналізуйте по датах окремо.
    p ≥ α → лінії паралельні, ефект варіанту стабільний у часі.

КРОК 5. ГРАФІК

  Кожна лінія = один варіант.
  Вертикальні смужки = ±СП (стандартна похибка середнього).
  Паралельні лінії → взаємодія незначуща.
  Розбіжні/лінії що перетинаються → взаємодія значуща.

КРОК 6. POST-HOC

  Після значущого ефекту Варіанту:
  Парні порівняння між варіантами (Бонферроні).

  Після значущої взаємодії:
  Простий ефект — порівняння варіантів на кожну дату окремо.
"""

    def __init__(self, parent, gs):
        self.win = tk.Toplevel(parent)
        self.win.title("Змішаний аналіз повторних вимірювань")
        self.win.geometry("1060x700"); set_icon(self.win)
        self.gs = gs
        self._fig = None
        self._plot_gs = {
            "font_family": "Times New Roman", "font_size": 10,
            "linewidth": 2.0, "markersize": 7, "marker": "o",
            "show_grid": True, "alpha_fill": 0.12,
            "colors": ["#4c72b0","#dd8452","#55a868","#c44e52",
                       "#8172b2","#937860","#da8bc3","#8c8c8c"],
        }
        self._build()

    def _build(self):
        try:
            self._build_inner()
        except Exception as _be:
            import traceback
            messagebox.showerror("Помилка","Помилка побудови вікна:\n"+traceback.format_exc())

    def _build_inner(self):
        # ── Toolbar ──────────────────────────────────────────
        top = tk.Frame(self.win, padx=8, pady=6); top.pack(fill=tk.X)
        tk.Button(top, text="▶ Аналіз", bg="#c62828", fg="white",
                  font=("Times New Roman",13),
                  command=self._run).pack(side=tk.LEFT, padx=4)

        mb2 = tk.Menubutton(top, text="⚙ Налаштування ▾",
                            font=("Times New Roman",11), relief=tk.RAISED, bd=2)
        mb2.pack(side=tk.LEFT, padx=4)
        sm = tk.Menu(mb2, tearoff=0)
        sm.add_command(label="Додати рядок",       command=self._add_row)
        sm.add_command(label="Видалити рядок",     command=self._del_row)
        sm.add_separator()
        sm.add_command(label="Додати стовпець",    command=self._add_col)
        sm.add_command(label="Видалити стовпець",  command=self._del_col)
        sm.add_separator()
        sm.add_command(label="💾 Зберегти проект", command=self._save_proj)
        sm.add_command(label="📂 Відкрити проект", command=self._load_proj)
        sm.add_separator()
        sm.add_command(label="🗑 Очистити таблицю", command=self._clear)
        mb2["menu"] = sm

        tk.Button(top, text="Вставити з буфера",
                  font=("Times New Roman",11),
                  command=self._paste).pack(side=tk.LEFT, padx=4)
        tk.Button(top, text="📚 Довідка", bg="#1a4b8c", fg="white",
                  font=("Times New Roman",11),
                  command=self._show_help).pack(side=tk.LEFT, padx=4)
        tk.Label(top,
                 text="Стовп.1=Варіант  Стовп.2=Повторність  Решта=Значення по датах",
                 font=("Times New Roman",9), fg="#666").pack(side=tk.LEFT, padx=8)

        # ── Таблиця ─────────────────────────────────────────
        mid = tk.Frame(self.win); mid.pack(fill=tk.BOTH, expand=True, padx=8)
        self.rows_n = 24; self.cols_n = 8
        self._canvas = tk.Canvas(mid)
        self._canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb = ttk.Scrollbar(mid, orient="vertical", command=self._canvas.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        self._canvas.configure(yscrollcommand=sb.set)
        self.inner = tk.Frame(self._canvas)
        self._canvas.create_window((0,0), window=self.inner, anchor="nw")
        self.inner.bind("<Configure>",
                        lambda e: self._canvas.config(scrollregion=self._canvas.bbox("all")))
        self.win.bind("<MouseWheel>",
                      lambda e: self._canvas.yview_scroll(int(-1*(e.delta/120)),"units"))

        # Фіксовані заголовки перших двох стовпців
        for j, txt in enumerate(["Варіант","Повторність"]):
            tk.Label(self.inner, text=txt, width=13, relief=tk.RIDGE,
                     bg="#444444", fg="white",
                     font=("Times New Roman",11,"bold")
                     ).grid(row=0, column=j, padx=1, pady=1, sticky="nsew")

        # Заголовки часових точок (перейменовувані)
        self.time_vars = []; self.time_labels = []
        for j in range(2, self.cols_n):
            var = tk.StringVar(value=f"Дата {j-1}")
            self.time_vars.append(var)
            lbl = tk.Label(self.inner, textvariable=var, width=12, cursor="hand2",
                           bg="#1a4b8c", fg="white", relief=tk.RIDGE,
                           font=("Times New Roman",11,"bold"))
            lbl.grid(row=0, column=j, padx=1, pady=1, sticky="nsew")
            lbl.bind("<Double-Button-1>",
                     lambda e, idx=j-2: self._rename_col(idx))
            self.time_labels.append(lbl)

        self.entries = []
        for i in range(self.rows_n):
            row_ = []
            for j in range(self.cols_n):
                w = 13 if j < 2 else 12
                e = tk.Entry(self.inner, width=w, font=("Times New Roman",11),
                             )
                e.grid(row=i+1, column=j, padx=1, pady=1)
                row_.append(e)
            self.entries.append(row_)
        _bind_nav(self.entries, self.win)

    # ── Перейменування дати ───────────────────────────────────
    def _rename_col(self, idx):
        dlg = tk.Toplevel(self.win); dlg.title("Перейменувати дату")
        dlg.resizable(False, False); dlg.grab_set()
        tk.Label(dlg, text=f"Назва дати/точки {idx+1}:",
                 font=("Times New Roman",12)).pack(padx=16, pady=14)
        var = tk.StringVar(value=self.time_vars[idx].get())
        e = tk.Entry(dlg, textvariable=var, font=("Times New Roman",12), width=24)
        e.pack(padx=16, pady=4); e.select_range(0, tk.END); e.focus_set()
        def apply():
            nm = var.get().strip()
            if nm: self.time_vars[idx].set(nm)
            dlg.destroy()
        tk.Button(dlg, text="OK", bg="#c62828", fg="white",
                  font=("Times New Roman",12), command=apply).pack(pady=(4,14))
        dlg.bind("<Return>", lambda ev: apply()); center_win(dlg)

    # ── Управління таблицею ───────────────────────────────────
    def _add_row(self):
        i = self.rows_n; row_ = []
        for j in range(self.cols_n):
            e = tk.Entry(self.inner, width=13 if j<2 else 12,
                         font=("Times New Roman",11))
            e.grid(row=i+1, column=j, padx=1, pady=1)
            row_.append(e)
        self.entries.append(row_); self.rows_n += 1
        _bind_nav(self.entries, self.win)

    def _del_row(self):
        if not self.entries: return
        for e in self.entries.pop(): e.destroy()
        self.rows_n -= 1

    def _add_col(self):
        ci = self.cols_n; self.cols_n += 1
        var = tk.StringVar(value=f"Дата {ci-1}")
        self.time_vars.append(var)
        lbl = tk.Label(self.inner, textvariable=var, width=12, cursor="hand2",
                       bg="#1a4b8c", fg="white", relief=tk.RIDGE,
                       font=("Times New Roman",11,"bold"))
        lbl.grid(row=0, column=ci, padx=1, pady=1, sticky="nsew")
        lbl.bind("<Double-Button-1>",
                 lambda e, idx=ci-2: self._rename_col(idx))
        self.time_labels.append(lbl)
        for i, row_ in enumerate(self.entries):
            e = tk.Entry(self.inner, width=12, font=("Times New Roman",11),
                         )
            e.grid(row=i+1, column=ci, padx=1, pady=1)
            row_.append(e)
        _bind_nav(self.entries, self.win)

    def _del_col(self):
        if self.cols_n <= 4: return  # мінімум Варіант+Повт+2 дати
        self.time_labels.pop().destroy(); self.time_vars.pop()
        for row_ in self.entries: row_.pop().destroy()
        self.cols_n -= 1

    def _clear(self):
        if not messagebox.askyesno("Очистити","Видалити всі дані?"): return
        for row in self.entries:
            for e in row: e.delete(0, tk.END)

    def _paste(self):
        try: data = self.win.clipboard_get()
        except Exception:
            messagebox.showwarning("","Скопіюйте дані з Excel і спробуйте знову."); return
        if not data.strip(): return
        pos = (0, 0)
        w = self.win.focus_get()
        if isinstance(w, tk.Entry):
            for i, row_ in enumerate(self.entries):
                for j, e in enumerate(row_):
                    if e is w: pos = (i, j); break
        r0, c0 = pos
        for ir, line in enumerate(data.splitlines()):
            if not line.strip(): continue
            while r0+ir >= len(self.entries): self._add_row()
            for jc, val in enumerate(line.split("\t")):
                cc = c0+jc
                if cc >= self.cols_n: continue
                self.entries[r0+ir][cc].delete(0, tk.END)
                self.entries[r0+ir][cc].insert(0, val.strip())
        path = filedialog.asksaveasfilename(
            parent=self.win, defaultextension=".sadp",
            filetypes=[("SAD проект","*.sadp"),("JSON","*.json")],
            title="Зберегти проект Змішаного Repeated Measures")
        if not path: return
        d = {
            "type": "mixed_repeated_measures",
            "version": APP_VER,
            "time_vars": [v.get() for v in self.time_vars],
            "rows_data": [[e.get() for e in row] for row in self.entries],
        }
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(d, f, ensure_ascii=False, indent=2)
            messagebox.showinfo("Збережено",
                f"Проект збережено:\n{path}\n\n"
                "При наступному відкритті завантажте цей файл і\n"
                "додайте нові стовпці (дати) через ⚙ → «Додати стовпець».")
        except Exception as ex:
            messagebox.showerror("Помилка збереження", str(ex))

    def _save_proj(self):
        path = filedialog.asksaveasfilename(
            parent=self.win, defaultextension=".sadp",
            filetypes=[("SAD проект","*.sadp"),("JSON","*.json")],
            title="Зберегти проект Змішаного Repeated Measures")
        if not path: return
        d = {
            "type": "mixed_repeated_measures", "version": APP_VER,
            "time_vars": [v.get() for v in self.time_vars],
            "rows_data": [[e.get() for e in row] for row in self.entries],
        }
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(d, f, ensure_ascii=False, indent=2)
            messagebox.showinfo("Збережено",
                f"Проект збережено:\n{path}\n\n"
                "Щоб додати нові дати: ⚙ → «Додати стовпець»")
        except Exception as ex:
            messagebox.showerror("Помилка збереження", str(ex))

    def _load_proj(self):
        path = filedialog.askopenfilename(
            parent=self.win,
            filetypes=[("SAD проект","*.sadp"),("JSON","*.json")],
            title="Відкрити проект Змішаного Repeated Measures")
        if not path: return
        try:
            with open(path, "r", encoding="utf-8") as f:
                d = json.load(f)
        except Exception as ex:
            messagebox.showerror("Помилка відкриття", str(ex)); return
        time_vars = d.get("time_vars", [])
        rows_data = d.get("rows_data", [])
        # Розширюємо до потрібної кількості стовпців (2 фіксовані + дати)
        n_cols_needed = 2 + len(time_vars)
        while self.cols_n < n_cols_needed: self._add_col()
        # Назви дат
        for i, nm in enumerate(time_vars):
            if i < len(self.time_vars): self.time_vars[i].set(nm)
        # Дані
        while len(self.entries) < len(rows_data): self._add_row()
        for i, row_vals in enumerate(rows_data):
            for j, v in enumerate(row_vals):
                if j < self.cols_n:
                    self.entries[i][j].delete(0, tk.END)
                    self.entries[i][j].insert(0, v)
        messagebox.showinfo("Завантажено",
            "Проект завантажено.\n\n"
            "Щоб додати нові дати вимірювань:\n"
            "  ⚙ → «Додати стовпець» → подвійний клік на заголовку → назва дати.")

    def _paste(self):
        try: data = self.win.clipboard_get()
        except Exception:
            messagebox.showwarning("","Скопіюйте дані з Excel і спробуйте знову."); return
        pos = (0,0)
        w = self.win.focus_get()
        if isinstance(w, tk.Entry):
            for i, row_ in enumerate(self.entries):
                for j, e in enumerate(row_):
                    if e is w: pos=(i,j); break
        r0, c0 = pos
        for ir, line in enumerate(data.splitlines()):
            if not line.strip(): continue
            while r0+ir >= len(self.entries): self._add_row()
            for jc, val in enumerate(line.split("\t")):
                cc = c0+jc
                if cc >= self.cols_n: continue
                self.entries[r0+ir][cc].delete(0,tk.END)
                self.entries[r0+ir][cc].insert(0, val.strip())

    def _show_help(self):
        win = tk.Toplevel(self.win)
        win.title("Довідка — Змішаний Repeated Measures")
        win.geometry("720x680"); set_icon(win)
        frm = tk.Frame(win); frm.pack(fill=tk.BOTH, expand=True, padx=8, pady=6)
        vsb = ttk.Scrollbar(frm, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        txt = tk.Text(frm, wrap="word", font=("Times New Roman",11),
                      yscrollcommand=vsb.set, relief=tk.FLAT,
                      bg="#fafafa", padx=10, pady=8, cursor="arrow")
        txt.pack(fill=tk.BOTH, expand=True); vsb.config(command=txt.yview)
        txt.insert("1.0", self.HELP_TEXT.strip()); txt.configure(state="disabled")
        txt.bind("<MouseWheel>",
                 lambda e: txt.yview_scroll(int(-1*(e.delta/120)),"units"))
        tk.Button(win, text="Закрити", command=win.destroy,
                  font=("Times New Roman",11)).pack(pady=6)

    # ── Зчитування та аналіз ─────────────────────────────────
    def _run(self):
        # Назви дат (лише заповнені)
        all_dates = [v.get().strip() or f"Д{i+1}" for i,v in enumerate(self.time_vars)]

        raw = [[e.get().strip() for e in row] for row in self.entries]
        raw = [r for r in raw if any(v for v in r)]
        if not raw:
            messagebox.showwarning("Немає даних","Введіть дані у таблицю."); return

        # Знаходимо заповнені часові стовпці (індекси 2..)
        filled_t = []
        for j in range(2, self.cols_n):
            for row in raw:
                v = row[j] if j<len(row) else ""
                if v:
                    try: float(v.replace(",",".")); filled_t.append(j); break
                    except ValueError: pass
        if not filled_t:
            messagebox.showwarning("Немає числових даних",
                "Введіть числові значення у стовпці 3 і далі (часові точки)."); return

        time_names = [all_dates[j-2] for j in filled_t]
        k = len(time_names)  # кількість часових точок

        # Зчитуємо рядки: варіант, повторність, значення
        rows_data = []
        for row in raw:
            var_nm = row[0].strip() if row[0].strip() else None
            rep_nm = row[1].strip() if len(row)>1 and row[1].strip() else None
            if not var_nm: continue
            try: float(var_nm.replace(",",".")); continue  # числовий → пропустити
            except ValueError: pass
            vals = []
            for j in filled_t:
                v = row[j].strip() if j<len(row) else ""
                try: vals.append(float(v.replace(",",".")))
                except: vals.append(float("nan"))
            if any(not math.isnan(v) for v in vals):
                rows_data.append({"var": var_nm, "rep": rep_nm or "?", "vals": vals})

        if not rows_data:
            messagebox.showwarning("Немає даних","Не вдалося зчитати жодного рядка."); return

        # Групуємо по варіантах
        var_levels = first_seen([r["var"] for r in rows_data])
        n_vars = len(var_levels)
        if n_vars < 2:
            messagebox.showwarning("Замало варіантів",
                "Потрібно щонайменше 2 варіанти (різні значення у стовпці 1)."); return

        # Будуємо матриці: var_data[var] = np.array(n_reps × k)
        var_data = {}
        for lv in var_levels:
            rr = [r["vals"] for r in rows_data if r["var"]==lv]
            # Видаляємо рядки де є NaN
            rr_clean = [v for v in rr if not any(math.isnan(x) for x in v)]
            if not rr_clean:
                messagebox.showwarning("Замало повних даних",
                    f"Варіант '{lv}' не має жодного рядка з повними даними."); return
            var_data[lv] = np.array([v[:k] for v in rr_clean], dtype=float)

        n_reps = {lv: len(var_data[lv]) for lv in var_levels}
        if min(n_reps.values()) < 2:
            messagebox.showwarning("Замало повторностей",
                "Кожен варіант потребує ≥ 2 повторностей."); return

        alpha = ALPHA

        # ══ Split-plot у часі: правильні дві помилки ══
        # Загальне середнє
        all_data = np.vstack([var_data[lv] for lv in var_levels])
        grand_mean = float(np.mean(all_data))
        N_total = sum(n_reps[lv] for lv in var_levels)  # всього рядків

        # Середні по варіантах (усереднені по часу)
        var_means = {lv: float(np.mean(var_data[lv])) for lv in var_levels}
        # Середні по часу (усереднені по всіх варіантах і повторностях)
        time_means = np.mean(all_data, axis=0)  # shape (k,)

        # ── SS розкладання (Split-plot) ──────────────────────
        # SS_var (between groups): сума (n_rep_i * k) * (mean_i - grand)^2
        SS_var = float(sum(
            n_reps[lv]*k*(var_means[lv]-grand_mean)**2
            for lv in var_levels))

        # SS_whole_error (within groups, between replicates):
        # для кожного варіанту сума по повторностях k*(rep_mean - var_mean)^2
        SS_wp_err = 0.
        for lv in var_levels:
            rep_means_lv = np.mean(var_data[lv], axis=1)  # середнє кожної повт. по часу
            SS_wp_err += k * float(np.sum((rep_means_lv - var_means[lv])**2))

        # SS_time (within subjects, main effect of time)
        SS_time = float(N_total * np.sum((time_means - grand_mean)**2))

        # SS_var_time (interaction)
        SS_inter = 0.
        for lv in var_levels:
            var_time_means = np.mean(var_data[lv], axis=0)  # shape (k,)
            SS_inter += n_reps[lv] * float(
                np.sum((var_time_means - var_means[lv] - time_means + grand_mean)**2))

        # SS_sub_error (within subjects, residual)
        SS_sub_err = 0.
        for lv in var_levels:
            vm_lv = np.mean(var_data[lv], axis=0)  # var×time means
            rep_means_lv = np.mean(var_data[lv], axis=1, keepdims=True)
            vm_lv_rep = var_means[lv]
            for ri in range(n_reps[lv]):
                for ti in range(k):
                    actual = var_data[lv][ri, ti]
                    expected = (var_means[lv] +
                                (np.mean(var_data[lv], axis=0)[ti] - var_means[lv]) +
                                (np.mean(var_data[lv], axis=1)[ri] - var_means[lv]))
                    SS_sub_err += (actual - expected)**2

        # Degrees of freedom
        df_var     = n_vars - 1
        df_wp_err  = sum(n_reps[lv]-1 for lv in var_levels)   # N - n_vars
        df_time    = k - 1
        df_inter   = df_var * df_time
        df_sub_err = df_wp_err * df_time

        # MS
        MS_var     = SS_var    / df_var    if df_var    > 0 else float("nan")
        MS_wp_err  = SS_wp_err / df_wp_err if df_wp_err > 0 else float("nan")
        MS_time    = SS_time   / df_time   if df_time   > 0 else float("nan")
        MS_inter   = SS_inter  / df_inter  if df_inter  > 0 else float("nan")
        MS_sub_err = SS_sub_err/ df_sub_err if df_sub_err > 0 else float("nan")

        # F-тести (два різних знаменники!)
        F_var   = MS_var  / MS_wp_err  if (MS_wp_err  and MS_wp_err  > 1e-12) else float("nan")
        F_time  = MS_time / MS_sub_err if (MS_sub_err and MS_sub_err > 1e-12) else float("nan")
        F_inter = MS_inter/ MS_sub_err if (MS_sub_err and MS_sub_err > 1e-12) else float("nan")

        p_var   = float(1-f_dist.cdf(F_var,   df_var,  df_wp_err))  if not math.isnan(F_var)   else float("nan")
        p_time  = float(1-f_dist.cdf(F_time,  df_time, df_sub_err)) if not math.isnan(F_time)  else float("nan")
        p_inter = float(1-f_dist.cdf(F_inter, df_inter,df_sub_err)) if not math.isnan(F_inter) else float("nan")

        # Partial η²
        def peta2(SS_eff, SS_err):
            return SS_eff/(SS_eff+SS_err) if (SS_eff+SS_err)>0 else float("nan")
        e2_var   = peta2(SS_var,   SS_wp_err)
        e2_time  = peta2(SS_time,  SS_sub_err)
        e2_inter = peta2(SS_inter, SS_sub_err)

        # ── Post-hoc між варіантами (Бонферроні) ─────────────
        from scipy.stats import ttest_ind
        var_pairs = [(var_levels[i],var_levels[j])
                     for i in range(n_vars) for j in range(i+1,n_vars)]
        mt_var = len(var_pairs)
        ph_var = []
        for lv1,lv2 in var_pairs:
            m1 = np.mean(var_data[lv1]); m2 = np.mean(var_data[lv2])
            # t-тест на середніх по часу (агрегований)
            means1 = np.mean(var_data[lv1], axis=1)
            means2 = np.mean(var_data[lv2], axis=1)
            try:
                t_, p_ = ttest_ind(means1, means2)
                p_adj = min(1., float(p_)*mt_var)
            except Exception: t_=float("nan"); p_adj=float("nan")
            mark = "**" if p_adj<alpha*0.2 else ("*" if p_adj<alpha else "–")
            ph_var.append([f"{lv1} vs {lv2}",
                           fmt(float(m1),4), fmt(float(m2),4),
                           fmt(float(m1-m2),4), fmt(float(t_),4),
                           fmt(p_adj,4), mark])

        # ── Простий ефект: порівняння варіантів на кожну дату ─
        simple_rows = []
        for ti, dn in enumerate(time_names):
            col_data = {lv: var_data[lv][:,ti] for lv in var_levels}
            # ANOVA на цій даті
            grand_t = np.mean(np.concatenate(list(col_data.values())))
            ss_b = sum(len(col_data[lv])*(np.mean(col_data[lv])-grand_t)**2 for lv in var_levels)
            ss_w = sum(np.sum((col_data[lv]-np.mean(col_data[lv]))**2) for lv in var_levels)
            dft_b = n_vars-1
            dft_w = sum(len(col_data[lv])-1 for lv in var_levels)
            ms_b = ss_b/dft_b if dft_b>0 else float("nan")
            ms_w = ss_w/dft_w if dft_w>0 else float("nan")
            F_t = ms_b/ms_w if (ms_w and ms_w>1e-12) else float("nan")
            p_t = float(1-f_dist.cdf(F_t,dft_b,dft_w)) if not math.isnan(F_t) else float("nan")
            mark_t = "**" if p_t<alpha*0.2 else ("*" if p_t<alpha else "–")
            simple_rows.append([dn, fmt(F_t,4), f"{dft_b},{dft_w}", fmt(p_t,4), mark_t])

        if not HAS_MPL:
            messagebox.showwarning("","matplotlib недоступний."); return

        self._show_results(
            var_levels, var_data, time_names, n_reps, alpha,
            SS_var, SS_wp_err, SS_time, SS_inter, SS_sub_err,
            df_var, df_wp_err, df_time, df_inter, df_sub_err,
            MS_var, MS_wp_err, MS_time, MS_inter, MS_sub_err,
            F_var, F_time, F_inter,
            p_var, p_time, p_inter,
            e2_var, e2_time, e2_inter,
            ph_var, simple_rows)

    def _show_results(self, var_levels, var_data, time_names, n_reps, alpha,
                      SS_var, SS_wp_err, SS_time, SS_inter, SS_sub_err,
                      df_var, df_wp_err, df_time, df_inter, df_sub_err,
                      MS_var, MS_wp_err, MS_time, MS_inter, MS_sub_err,
                      F_var, F_time, F_inter,
                      p_var, p_time, p_inter,
                      e2_var, e2_time, e2_inter,
                      ph_var, simple_rows):

        win = tk.Toplevel(self.win)
        win.title("Змішаний Repeated Measures — Результати")
        win.geometry("1160x820"); set_icon(win)

        # Toolbar
        tb = tk.Frame(win, padx=6, pady=5); tb.pack(fill=tk.X)
        self._graph_frame = tk.Frame(win)
        tk.Button(tb, text="📋 Копіювати графік", font=("Times New Roman",11),
                  command=lambda: self._copy_fig()).pack(side=tk.LEFT, padx=4)
        tk.Button(tb, text="⚙ Налаштування графіка", font=("Times New Roman",11),
                  command=lambda: self._restyle(win, var_levels, var_data, time_names, alpha)
                  ).pack(side=tk.LEFT, padx=4)

        # Scrollable body
        sa = tk.Frame(win); sa.pack(fill=tk.BOTH, expand=True)
        vsb = ttk.Scrollbar(sa, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        sc = tk.Canvas(sa, yscrollcommand=vsb.set, highlightthickness=0)
        sc.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.config(command=sc.yview)
        body = tk.Frame(sc); sc.create_window((0,0), window=body, anchor="nw")
        body.bind("<Configure>", lambda e: sc.configure(scrollregion=sc.bbox("all")))
        win.bind("<MouseWheel>", lambda e: sc.yview_scroll(int(-1*(e.delta/120)),"units"))

        def _head(t):
            tk.Label(body, text=t, font=("Times New Roman",12,"bold"),
                     bg="#e8eeff", anchor="w", padx=8, pady=3
                     ).pack(fill=tk.X, padx=6, pady=8)
        def _txt(t, color="#000"):
            tk.Label(body, text=t, font=("Times New Roman",11), fg=color,
                     anchor="w", justify="left").pack(fill=tk.X, padx=14, pady=1)
        def _tbl(hdrs, rows):
            f, _ = make_tv(body, hdrs, rows); f.pack(fill=tk.X, padx=8, pady=2)

        _head("Змішаний аналіз повторних вимірювань (Split-plot у часі)")
        _txt(f"Варіантів: {len(var_levels)}   |   Повторностей: {list(n_reps.values())}   "
             f"|   Часових точок: {len(time_names)}   |   α = {alpha}")

        # Основна таблиця ANOVA
        _head("Зведена таблиця дисперсійного аналізу")
        _txt("Варіант тестується через whole-plot error; "
             "Час і Взаємодія — через sub-plot error.", "#555")
        def _row(name, SS, df, MS, F, p, e2, note=""):
            mark = ("**" if p<alpha*0.2 else ("*" if p<alpha else "–")) if not math.isnan(p) else "–"
            return [name, fmt(SS,4), str(df), fmt(MS,4), fmt(F,4),
                    fmt(p,4), mark, fmt(e2,4), eta2_label(e2), note]
        anova_tbl = [
            _row("Варіант (between)",    SS_var,    df_var,    MS_var,    F_var,   p_var,   e2_var,  "÷ WP-error"),
            ["  Whole-plot error",       fmt(SS_wp_err,4), str(df_wp_err), fmt(MS_wp_err,4),"–","–","–","–","",""],
            _row("Час (within)",          SS_time,   df_time,   MS_time,   F_time,  p_time,  e2_time, "÷ SP-error"),
            _row("Варіант × Час",         SS_inter,  df_inter,  MS_inter,  F_inter, p_inter, e2_inter,"÷ SP-error"),
            ["  Sub-plot error",         fmt(SS_sub_err,4),str(df_sub_err),fmt(MS_sub_err,4),"–","–","–","–","",""],
        ]
        _tbl(["Джерело","SS","df","MS","F","p","Знач.","η²","Ефект","Знаменник"],
             anova_tbl)

        # Висновки
        _head("Висновки")
        for label, p_val, e2 in [
            ("Варіант", p_var, e2_var),
            ("Час",     p_time, e2_time),
            ("Варіант × Час (взаємодія)", p_inter, e2_inter)
        ]:
            if not math.isnan(p_val):
                sig = p_val < alpha
                col = "#1a6b1a" if sig else "#c62828"
                txt = (f"✓ {label}: значущий (p={fmt(p_val,4)}, η²={fmt(e2,3)} — {eta2_label(e2)})"
                       if sig else
                       f"✗ {label}: незначущий (p={fmt(p_val,4)})")
                _txt(txt, col)
        if not math.isnan(p_inter) and p_inter < alpha:
            _txt("⚠ Взаємодія значуща → ефект варіанту залежить від дати.\n"
                 "   Аналізуйте прості ефекти (таблиця нижче) і дивіться чи лінії розходяться.",
                 "#c62828")
        elif not math.isnan(p_inter):
            _txt("✓ Взаємодія незначуща → лінії паралельні, ефект варіанту стабільний у часі.",
                 "#1a6b1a")

        # Середні
        _head("Середні значення (Варіант × Дата)")
        means_hdrs = ["Варіант"] + time_names + ["Загальне"]
        means_rows = []
        for lv in var_levels:
            row_ = [lv]
            for ti in range(len(time_names)):
                row_.append(fmt(float(np.mean(var_data[lv][:,ti])),3))
            row_.append(fmt(float(np.mean(var_data[lv])),3))
            means_rows.append(row_)
        _tbl(means_hdrs, means_rows)

        # Графік
        self._graph_frame = tk.Frame(body)
        self._graph_frame.pack(fill=tk.X, padx=8, pady=6)
        self._draw_graph(var_levels, var_data, time_names, alpha)

        # Post-hoc варіанти
        if not math.isnan(p_var) and p_var < alpha:
            _head(f"Пост-хок: порівняння варіантів (Бонферроні, α_скор={fmt(alpha/len(ph_var),4)})")
            _tbl(["Пара","Сер.1","Сер.2","Різниця","t","p (Bonf.)","Знач."], ph_var)
            _txt(f"* p<{alpha}   ** p<{alpha*0.2:.3f}   – незначуща","#555")

        # Прості ефекти (по датах)
        if not math.isnan(p_inter) and p_inter < alpha:
            _head("Простий ефект: порівняння варіантів на кожну дату окремо")
            _txt("Виконується при значущій взаємодії — показує на яких датах варіанти відрізняються.",
                 "#555")
            _tbl(["Дата","F","df","p","Знач."], simple_rows)
            _txt(f"* p<{alpha}   ** p<{alpha*0.2:.3f}   – незначуща","#555")

    def _draw_graph(self, var_levels, var_data, time_names, alpha=0.05):
        for w in self._graph_frame.winfo_children(): w.destroy()
        gs = self._plot_gs
        k = len(time_names)
        colors = gs["colors"]
        fig = Figure(figsize=(10, 6), dpi=100)
        ax  = fig.add_subplot(111)

        for ci, lv in enumerate(var_levels):
            col = colors[ci % len(colors)]
            means_ = np.mean(var_data[lv], axis=0)
            ses_   = np.std(var_data[lv], axis=0, ddof=1) / math.sqrt(len(var_data[lv]))
            ax.errorbar(range(k), means_, yerr=ses_,
                        fmt=gs["marker"]+"-", capsize=5,
                        color=col, ecolor=col,
                        linewidth=gs["linewidth"],
                        markersize=gs["markersize"],
                        label=str(lv), alpha=0.9, zorder=3)
            # Тіньова смуга ±СП
            ax.fill_between(range(k),
                            means_-ses_, means_+ses_,
                            alpha=gs["alpha_fill"], color=col)

        ax.set_xticks(range(k))
        ax.set_xticklabels(time_names,
                           fontsize=gs["font_size"],
                           fontfamily=gs["font_family"])
        ax.set_xlabel("Дата / Часова точка",
                      fontsize=gs["font_size"], fontfamily=gs["font_family"])
        ax.set_ylabel("Середнє ± СП",
                      fontsize=gs["font_size"], fontfamily=gs["font_family"])
        ax.set_title("Динаміка по варіантах (Середнє ± СП)",
                     fontsize=gs["font_size"]+1, fontfamily=gs["font_family"])
        ax.legend(title="Варіант", fontsize=gs["font_size"],
                  title_fontsize=gs["font_size"])
        if gs["show_grid"]:
            ax.yaxis.grid(True, linestyle="--", alpha=0.35)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        fig.tight_layout()
        self._fig = fig

        embed_figure(fig, self._graph_frame)

    def _copy_fig(self):
        if self._fig is None:
            messagebox.showwarning("","Спочатку виконайте аналіз."); return
        ok, msg = _copy_fig_to_clipboard(self._fig)
        if ok: messagebox.showinfo("","Графік скопійовано.\nВставте у Word через Ctrl+V.")
        else:   messagebox.showwarning("",f"Помилка: {msg}")

    def _restyle(self, win, var_levels, var_data, time_names, alpha):
        dlg = tk.Toplevel(win); dlg.title("Налаштування графіка")
        dlg.resizable(False, False); set_icon(dlg); dlg.grab_set()
        gs = self._plot_gs
        frm = tk.Frame(dlg, padx=16, pady=14); frm.pack()
        rb_f = ("Times New Roman",12)
        ff_v = tk.StringVar(value=gs["font_family"])
        fz_v = tk.IntVar(value=gs["font_size"])
        lw_v = tk.DoubleVar(value=gs["linewidth"])
        ms_v = tk.IntVar(value=gs["markersize"])
        mk_v = tk.StringVar(value=gs["marker"])
        gr_v = tk.BooleanVar(value=gs["show_grid"])
        al_v = tk.DoubleVar(value=gs["alpha_fill"])
        rows_cfg = [
            ("Шрифт:",          "combo",  ff_v, ["Times New Roman","Arial","Calibri","Georgia"]),
            ("Розмір шрифту:",  "spin",   fz_v, (7,18)),
            ("Товщина лінії:",  "scale",  lw_v, (0.5,5.0)),
            ("Розмір маркера:", "spin",   ms_v, (3,20)),
            ("Тип маркера:",    "combo",  mk_v, ["o","s","^","D","v","*","+"]),
            ("Показати сітку:", "check",  gr_v, None),
            ("Прозорість тіні:","scale",  al_v, (0.0,0.4)),
        ]
        col_refs = list(gs["colors"])
        col_btns = []
        for ri, (lbl, wt, var, opts) in enumerate(rows_cfg):
            tk.Label(frm, text=lbl, font=rb_f).grid(row=ri, column=0, sticky="w", pady=4)
            if wt=="combo":
                ttk.Combobox(frm, textvariable=var, values=opts,
                             state="readonly", width=18).grid(row=ri, column=1, sticky="w", padx=8)
            elif wt=="spin":
                tk.Spinbox(frm, from_=opts[0], to=opts[1], textvariable=var,
                           width=7).grid(row=ri, column=1, sticky="w", padx=8)
            elif wt=="scale":
                tk.Scale(frm, from_=opts[0], to=opts[1], resolution=0.05,
                         orient="horizontal", variable=var,
                         length=160).grid(row=ri, column=1, sticky="w", padx=8)
            elif wt=="check":
                tk.Checkbutton(frm, variable=var).grid(row=ri, column=1, sticky="w", padx=8)

        # Кольори варіантів
        base_r = len(rows_cfg)
        tk.Label(frm, text="Кольори варіантів:", font=rb_f).grid(
            row=base_r, column=0, sticky="w", pady=4)
        cf = tk.Frame(frm); cf.grid(row=base_r, column=1, sticky="w")
        for ci, lv in enumerate(var_levels[:8]):
            c = col_refs[ci] if ci < len(col_refs) else "#999"
            btn = tk.Button(cf, width=3, relief=tk.SUNKEN, bg=c,
                            text=str(ci+1), font=("Times New Roman",8))
            btn.pack(side=tk.LEFT, padx=2)
            def _pick(idx=ci, b=btn):
                ch = colorchooser.askcolor(color=col_refs[idx], parent=dlg)
                if ch and ch[1]: col_refs[idx]=ch[1]; b.configure(bg=ch[1])
            btn.configure(command=_pick); col_btns.append(btn)

        def apply():
            self._plot_gs.update({
                "font_family": ff_v.get(), "font_size": fz_v.get(),
                "linewidth":   lw_v.get(), "markersize": ms_v.get(),
                "marker":      mk_v.get(), "show_grid": gr_v.get(),
                "alpha_fill":  al_v.get(), "colors": col_refs,
            })
            dlg.destroy()
            self._draw_graph(var_levels, var_data, time_names, alpha)

        bf = tk.Frame(frm); bf.grid(row=base_r+1, column=0, columnspan=2, pady=(14,0))
        tk.Button(bf, text="OK (застосувати)", bg="#c62828", fg="white",
                  font=rb_f, command=apply).pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Скасувати", font=rb_f, command=dlg.destroy).pack(side=tk.LEFT)
        center_win(dlg)


# ═══════════════════════════════════════════════════════════════
# STABILITY ANALYSIS  (Eberhart–Russell + GGE biplot)
# ═══════════════════════════════════════════════════════════════
class StabilityWindow:
    """Аналіз стабільності генотипів (GxE взаємодія)."""

    HELP_TEXT = """
АНАЛІЗ СТАБІЛЬНОСТІ (GxE) — ПОКРОКОВА ІНСТРУКЦІЯ
══════════════════════════════════════════════════

ЩО ТАКЕ АНАЛІЗ СТАБІЛЬНОСТІ?
  Оцінює як стабільно генотипи (сорти) поводяться
  в різних середовищах (роки, локації, умови).

  Два ключових питання:
  1. Який генотип найбільш продуктивний загалом?
  2. Який генотип найбільш СТАБІЛЬНИЙ (мало залежить від умов)?

  Генотип може бути:
  • Стабільним і продуктивним → ідеальний для широкого впровадження
  • Стабільним але низькопродуктивним → стабільний аутсайдер
  • Нестабільним і продуктивним → хороший лише у сприятливих умовах

КРОК 1. СТРУКТУРА ТАБЛИЦІ

  Перший стовпець: Назва генотипу/сорту (текст)
  Решта стовпців: Значення показника у кожному середовищі

  Значення = СЕРЕДНЄ по повторностях для цього генотипу у цьому середовищі.

  Приклад (4 сорти, 4 роки):
  | Генотип  | 2021 | 2022 | 2023 | 2024 |
  | Сорт А   | 5.8  |  6.2 |  5.5 |  6.8 |
  | Сорт Б   | 4.9  |  7.1 |  4.2 |  7.8 |
  | Контроль | 5.2  |  5.4 |  5.1 |  5.6 |

  Перейменуйте заголовки середовищ (подвійний клік).
  Мінімум: 2 генотипи, 2 середовища.

КРОК 2. МЕТОД EBERHART-RUSSELL

  Рівняння: Yij = μi + bi·Ij + δij
  де:
    Yij — врожай сорту i в середовищі j
    μi  — середнє сорту по всіх середовищах
    bi  — коефіцієнт регресії (відгук на умови)
    Ij  — індекс середовища
    δij — відхилення від регресії

  Параметри стабільності:

  bi (коефіцієнт регресії):
    bi = 1.0 → сорт реагує як середній по популяції
    bi > 1.0 → адаптивний/чутливий (краще у сприятливих, гірше у несприятливих)
    bi < 1.0 → консервативний/стабільний (слабко реагує на умови)

  s²d (дисперсія відхилень від регресії):
    s²d ≈ 0 → точна лінійна відповідь, передбачуваний сорт ✓
    s²d > 0 → непередбачувана реакція ✗

  КЛАСИ СТАБІЛЬНОСТІ:
    Стабільний:      bi ≈ 1, s²d ≈ 0  → рекомендований для всіх зон
    Адаптивний:      bi > 1, s²d ≈ 0  → лише для сприятливих умов
    Консервативний:  bi < 1, s²d ≈ 0  → для несприятливих умов
    Нестабільний:    будь-який bi, s²d > 0  → непередбачуваний

КРОК 3. GGE BIPLOT

  GGE = Genotype + Genotype×Environment interaction.
  Двовимірний графік що показує одночасно:
  • Продуктивність генотипів (відстань від центру)
  • Стабільність (чим ближче до кола — тим стабільніший)
  • Адаптацію до конкретних середовищ

  Стрілки = середовища.
  Точки = генотипи.
  Генотип близько до стрілки середовища → добре адаптований до нього.
"""

    def __init__(self, parent, gs):
        self.win = tk.Toplevel(parent)
        self.win.title("Аналіз стабільності (GxE)")
        self.win.geometry("1020x680"); set_icon(self.win)
        self.gs = gs
        self._stab_fig = None
        self._build()

    def _build(self):
        # ── Toolbar ──────────────────────────────────────────
        top = tk.Frame(self.win, padx=8, pady=6); top.pack(fill=tk.X)
        tk.Button(top, text="▶ Аналіз", bg="#c62828", fg="white",
                  font=("Times New Roman",13),
                  command=self._run).pack(side=tk.LEFT, padx=4)

        mb2 = tk.Menubutton(top, text="⚙ Налаштування ▾",
                            font=("Times New Roman",11), relief=tk.RAISED, bd=2)
        mb2.pack(side=tk.LEFT, padx=4)
        sm = tk.Menu(mb2, tearoff=0)
        sm.add_command(label="Додати рядок",       command=self._add_row)
        sm.add_command(label="Видалити рядок",     command=self._del_row)
        sm.add_separator()
        sm.add_command(label="Додати стовпець",    command=self._add_col)
        sm.add_command(label="Видалити стовпець",  command=self._del_col)
        sm.add_separator()
        sm.add_command(label="💾 Зберегти проект", command=self._save_proj)
        sm.add_command(label="📂 Відкрити проект", command=self._load_proj)
        sm.add_separator()
        sm.add_command(label="🗑 Очистити таблицю", command=self._clear)
        mb2["menu"] = sm

        tk.Button(top, text="Вставити з буфера",
                  font=("Times New Roman",11),
                  command=self._paste).pack(side=tk.LEFT, padx=4)
        tk.Button(top, text="📚 Довідка", bg="#1a4b8c", fg="white",
                  font=("Times New Roman",11),
                  command=self._show_help).pack(side=tk.LEFT, padx=4)
        tk.Label(top,
                 text="Подвійний клік на заголовку → перейменувати середовище",
                 font=("Times New Roman",9), fg="#666").pack(side=tk.LEFT, padx=8)

        # ── Таблиця ─────────────────────────────────────────
        mid = tk.Frame(self.win); mid.pack(fill=tk.BOTH, expand=True, padx=8)
        self.rows_n = 16; self.cols_n = 8
        self._canvas = tk.Canvas(mid)
        self._canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb = ttk.Scrollbar(mid, orient="vertical", command=self._canvas.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        self._canvas.configure(yscrollcommand=sb.set)
        self.inner = tk.Frame(self._canvas)
        self._canvas.create_window((0,0), window=self.inner, anchor="nw")
        self.inner.bind("<Configure>",
                        lambda e: self._canvas.config(scrollregion=self._canvas.bbox("all")))
        self.win.bind("<MouseWheel>",
                      lambda e: self._canvas.yview_scroll(int(-1*(e.delta/120)),"units"))

        # Перший заголовок — «Генотип» (фіксований)
        tk.Label(self.inner, text="Генотип", relief=tk.RIDGE, width=14,
                 bg="#444444", fg="white",
                 font=("Times New Roman",11,"bold")
                 ).grid(row=0, column=0, padx=1, pady=1, sticky="nsew")

        # Заголовки середовищ (перейменовувані)
        self.env_vars = []; self.env_labels = []
        for j in range(1, self.cols_n):
            var = tk.StringVar(value=f"E{j}")
            self.env_vars.append(var)
            lbl = tk.Label(self.inner, textvariable=var, width=12, cursor="hand2",
                           bg="#1a4b8c", fg="white", relief=tk.RIDGE,
                           font=("Times New Roman",11,"bold"))
            lbl.grid(row=0, column=j, padx=1, pady=1, sticky="nsew")
            lbl.bind("<Double-Button-1>", lambda e, idx=j-1: self._rename_env(idx))
            self.env_labels.append(lbl)

        self.entries = []
        for i in range(self.rows_n):
            row_ = []
            for j in range(self.cols_n):
                e = tk.Entry(self.inner, width=14 if j==0 else 11,
                             font=("Times New Roman",11))
                e.grid(row=i+1, column=j, padx=1, pady=1)
                row_.append(e)
            self.entries.append(row_)
        _bind_nav(self.entries, self.win)

    # ── Перейменування середовища ─────────────────────────────
    def _rename_env(self, idx):
        dlg = tk.Toplevel(self.win); dlg.title("Перейменувати середовище")
        dlg.resizable(False, False); dlg.grab_set()
        tk.Label(dlg, text=f"Назва середовища {idx+1}:",
                 font=("Times New Roman",12)).pack(padx=16, pady=14)
        var = tk.StringVar(value=self.env_vars[idx].get())
        e = tk.Entry(dlg, textvariable=var, font=("Times New Roman",12), width=24)
        e.pack(padx=16, pady=4); e.select_range(0, tk.END); e.focus_set()
        def apply():
            nm = var.get().strip()
            if nm: self.env_vars[idx].set(nm)
            dlg.destroy()
        tk.Button(dlg, text="OK", bg="#c62828", fg="white",
                  font=("Times New Roman",12), command=apply).pack(pady=(4,14))
        dlg.bind("<Return>", lambda ev: apply()); center_win(dlg)

    # ── Управління таблицею ───────────────────────────────────
    def _add_row(self):
        i = self.rows_n; row_ = []
        for j in range(self.cols_n):
            e = tk.Entry(self.inner, width=14 if j==0 else 11,
                         font=("Times New Roman",11))
            e.grid(row=i+1, column=j, padx=1, pady=1)
            row_.append(e)
        self.entries.append(row_); self.rows_n += 1
        _bind_nav(self.entries, self.win)

    def _del_row(self):
        if not self.entries: return
        for e in self.entries.pop(): e.destroy()
        self.rows_n -= 1

    def _add_col(self):
        ci = self.cols_n; self.cols_n += 1
        var = tk.StringVar(value=f"E{ci}")
        self.env_vars.append(var)
        lbl = tk.Label(self.inner, textvariable=var, width=12, cursor="hand2",
                       bg="#1a4b8c", fg="white", relief=tk.RIDGE,
                       font=("Times New Roman",11,"bold"))
        lbl.grid(row=0, column=ci, padx=1, pady=1, sticky="nsew")
        lbl.bind("<Double-Button-1>", lambda e, idx=ci-1: self._rename_env(idx))
        self.env_labels.append(lbl)
        for i, row_ in enumerate(self.entries):
            e = tk.Entry(self.inner, width=11, font=("Times New Roman",11))
            e.grid(row=i+1, column=ci, padx=1, pady=1)
            row_.append(e)
        _bind_nav(self.entries, self.win)

    def _del_col(self):
        if self.cols_n <= 3: return
        self.env_labels.pop().destroy(); self.env_vars.pop()
        for row_ in self.entries: row_.pop().destroy()
        self.cols_n -= 1

    def _clear(self):
        if not messagebox.askyesno("Очистити","Видалити всі дані?"): return
        for row in self.entries:
            for e in row: e.delete(0, tk.END)

    def _paste(self):
        try: data = self.win.clipboard_get()
        except Exception:
            messagebox.showwarning("","Скопіюйте дані з Excel і спробуйте знову."); return
        if not data.strip(): return
        pos = (0, 0)
        w = self.win.focus_get()
        if isinstance(w, tk.Entry):
            for i, row_ in enumerate(self.entries):
                for j, e in enumerate(row_):
                    if e is w: pos=(i,j); break
        r0, c0 = pos
        for ir, line in enumerate(data.splitlines()):
            if not line.strip(): continue
            while r0+ir >= len(self.entries): self._add_row()
            for jc, val in enumerate(line.split("\t")):
                cc = c0+jc
                if cc >= self.cols_n: continue
                self.entries[r0+ir][cc].delete(0, tk.END)
                self.entries[r0+ir][cc].insert(0, val.strip())

    def _save_proj(self):
        path = filedialog.asksaveasfilename(
            parent=self.win, defaultextension=".sadp",
            filetypes=[("SAD проект","*.sadp"),("JSON","*.json")])
        if not path: return
        d = {"type":"stability","version":APP_VER,
             "env_vars":[v.get() for v in self.env_vars],
             "rows_data":[[e.get() for e in row] for row in self.entries]}
        try:
            with open(path,"w",encoding="utf-8") as f: json.dump(d,f,ensure_ascii=False,indent=2)
            messagebox.showinfo("Збережено", path)
        except Exception as ex: messagebox.showerror("Помилка",str(ex))

    def _load_proj(self):
        path = filedialog.askopenfilename(
            parent=self.win, filetypes=[("SAD проект","*.sadp"),("JSON","*.json")])
        if not path: return
        try:
            with open(path,"r",encoding="utf-8") as f: d=json.load(f)
        except Exception as ex: messagebox.showerror("Помилка",str(ex)); return
        env_vars = d.get("env_vars",[])
        rows_data = d.get("rows_data",[])
        while self.cols_n < 1+len(env_vars): self._add_col()
        for i,nm in enumerate(env_vars):
            if i<len(self.env_vars): self.env_vars[i].set(nm)
        while len(self.entries)<len(rows_data): self._add_row()
        for i,rv in enumerate(rows_data):
            for j,v in enumerate(rv):
                if j<self.cols_n:
                    self.entries[i][j].delete(0,tk.END); self.entries[i][j].insert(0,v)
        messagebox.showinfo("Завантажено","Проект завантажено.")

    def _show_help(self):
        win = tk.Toplevel(self.win); win.title("Довідка — Аналіз стабільності")
        win.geometry("720x680"); set_icon(win)
        frm = tk.Frame(win); frm.pack(fill=tk.BOTH, expand=True, padx=8, pady=6)
        vsb = ttk.Scrollbar(frm, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        txt = tk.Text(frm, wrap="word", font=("Times New Roman",11),
                      yscrollcommand=vsb.set, relief=tk.FLAT,
                      bg="#fafafa", padx=10, pady=8, cursor="arrow")
        txt.pack(fill=tk.BOTH, expand=True); vsb.config(command=txt.yview)
        txt.insert("1.0", self.HELP_TEXT.strip()); txt.configure(state="disabled")
        txt.bind("<MouseWheel>", lambda e: txt.yview_scroll(int(-1*(e.delta/120)),"units"))
        tk.Button(win, text="Закрити", command=win.destroy,
                  font=("Times New Roman",11)).pack(pady=6)



    def _run(self):
        env_names = [v.get().strip() or f"E{i+1}" for i, v in enumerate(self.env_vars)]
        raw = [[e.get().strip() for e in row] for row in self.entries]
        gen_names = []; matrix = []
        for row in raw:
            nm = row[0] if row[0] else f"G{len(gen_names)+1}"
            vals = []
            for v in row[1:len(env_names)+1]:
                if not v: vals.append(np.nan)
                else:
                    try: vals.append(float(v.replace(",",".")))
                    except Exception: vals.append(np.nan)
            if any(not math.isnan(v) for v in vals):
                gen_names.append(nm); matrix.append(vals)

        if len(matrix) < 2: messagebox.showwarning("Замало генотипів","Потрібно ≥ 2 генотипи."); return
        e_count = len(env_names); g_count = len(gen_names)
        data = np.array(matrix, dtype=float)

        # ── Eberhart–Russell regression stability ──────────────
        env_means = np.nanmean(data, axis=0)
        grand_mean = np.nanmean(data)
        env_index = env_means - grand_mean  # environment index

        er_rows = []
        for i in range(g_count):
            y = data[i]; mask = ~np.isnan(y) & ~np.isnan(env_index)
            if np.sum(mask) < 2:
                er_rows.append([gen_names[i], fmt(np.nanmean(y),3), "–", "–", "–"]); continue
            xi = env_index[mask]; yi = y[mask]
            # linear regression of genotype yield on env index
            X_ = np.column_stack([np.ones(len(xi)), xi])
            beta, *_ = np.linalg.lstsq(X_, yi, rcond=None)
            b_i = float(beta[1])  # regression coefficient (stability)
            yhat = X_ @ beta
            ss_dev = float(np.sum((yi - yhat)**2))
            s2d = ss_dev / max(len(yi)-2, 1)  # variance of deviations
            gen_mean = float(np.nanmean(y))
            er_rows.append([gen_names[i], fmt(gen_mean,3), fmt(b_i,4), fmt(s2d,4),
                            "Стабільний (bi≈1, s²d≈0)" if abs(b_i-1)<0.2 and s2d<0.1 else
                            "Адаптивний" if b_i > 1.2 else "Консервативний" if b_i < 0.8 else "Середній"])

        # ── GGE Biplot via SVD ──────────────────────────────────
        # center by environment means
        data_c = data - env_means[np.newaxis, :]
        # replace NaN with 0 for SVD
        data_c = np.nan_to_num(data_c)
        U, S, Vt = np.linalg.svd(data_c, full_matrices=False)
        pc1_g = U[:,0] * S[0]; pc2_g = U[:,1] * S[1] if len(S)>1 else np.zeros(g_count)
        pc1_e = Vt[0,:];       pc2_e = Vt[1,:] if len(S)>1 else np.zeros(e_count)
        var_exp = S**2 / np.sum(S**2) * 100

        if not HAS_MPL: messagebox.showwarning("","matplotlib needed."); return
        win = tk.Toplevel(self.win); win.title("Аналіз стабільності — Результати"); win.geometry("1150x720")

        fig = Figure(figsize=(10, 6), dpi=100)
        # GGE biplot
        ax1 = fig.add_subplot(121)
        ax1.axhline(0, color="k", lw=0.5); ax1.axvline(0, color="k", lw=0.5)
        ax1.scatter(pc1_g, pc2_g, s=80, color="#4c72b0", zorder=3)
        for i, nm in enumerate(gen_names):
            ax1.annotate(nm, (pc1_g[i], pc2_g[i]), fontsize=8, ha='center', va='bottom')
        # environment vectors
        sc = max(np.max(np.abs(pc1_g)), np.max(np.abs(pc2_g)))
        sc_e = sc / max(np.max(np.abs(pc1_e)), max(np.max(np.abs(pc2_e)),1e-10))
        for j, nm in enumerate(env_names):
            ax1.annotate("", xy=(pc1_e[j]*sc_e*0.8, pc2_e[j]*sc_e*0.8), xytext=(0,0),
                         arrowprops=dict(arrowstyle="->", color="#c62828", lw=1.2))
            ax1.text(pc1_e[j]*sc_e*0.85, pc2_e[j]*sc_e*0.85, nm, fontsize=8, color="#c62828")
        ax1.set_xlabel(f"ГК1 ({fmt(var_exp[0],1)}%)"); ax1.set_ylabel(f"ГК2 ({fmt(var_exp[1] if len(var_exp)>1 else 0,1)}%)")
        ax1.set_title("GGE Biplot (Генотип × Середовище)"); ax1.yaxis.grid(True, alpha=0.25)

        # Stability table
        ax2 = fig.add_subplot(122)
        ax2.axis("off")
        col_labels = ["Genotype","Mean","bi","s²d","Stability"]
        tbl = ax2.table(cellText=er_rows, colLabels=col_labels, loc="center", cellLoc="center")
        tbl.auto_set_font_size(False); tbl.set_fontsize(9); tbl.scale(1, 1.4)
        ax2.set_title("Стабільність Eberhart–Russell", pad=14)

        fig.tight_layout()
        embed_figure(fig, win)
        frm, _ = make_tv(win, ["Генотип","Середнє","bi","s²d","Клас стабільності"], er_rows)
        frm.pack(fill=tk.X, padx=8, pady=4)



# ═══════════════════════════════════════════════════════════════
# ANCOVA — Analysis of Covariance
# ═══════════════════════════════════════════════════════════════
class AncovaWindow:
    """ANCOVA — Коваріаційний аналіз."""

    HELP_TEXT = """
ANCOVA — ПОКРОКОВА ІНСТРУКЦІЯ
══════════════════════════════════════════

ЩО ТАКЕ ANCOVA?
  ANCOVA (Коваріаційний аналіз) = ANOVA + контроль неперервної змінної.
  Порівнює групи за залежною змінною (Y), виключаючи вплив
  однієї або кількох коваріат (змінних що ви вимірюєте але не контролюєте).

КОЛИ ВИКОРИСТОВУВАТИ?
  Коли між групами є відмінності у вихідних умовах:
  • Порівняння врожайності сортів, але pH ґрунту різний на ділянках
  • Порівняння приросту, але початкова маса рослин різна
  • Вплив обробки, але температура чи вологість відрізнялась

КРОК 1. СТРУКТУРА ТАБЛИЦІ ДАНИХ

  Стовпці мають бути у такому порядку:
  [Група] | [Коваріата 1] | [Коваріата 2] | ... | [Залежна Y]

  Перейменуйте заголовки (сині/блакитні клітинки зверху):
    Перший стовпець = Назва групи/фактора (текстові мітки!)
    Останній стовпець = Залежна змінна Y (числа)
    Між ними = Коваріати (числа)

  Приклад (порівняння сортів, коваріата = pH):
  | Сорт     | pH ґрунту | Врожайність |
  | Сорт А   |    6.2    |    5.8      |
  | Сорт А   |    5.9    |    5.4      |
  | Сорт Б   |    6.5    |    6.2      |

  Мінімум: 6 спостережень, 2 групи, 2 спостереження в кожній групі.

КРОК 2. РІВЕНЬ ЗНАЧУЩОСТІ α
  Стандарт: 0.05.
  Строже: 0.01 (при множинних порівняннях).

КРОК 3. ВИКОНАННЯ АНАЛІЗУ
  Натисніть «▶ Виконати» та дочекайтесь результатів.

КРОК 4. АВТОМАТИЧНІ ПЕРЕВІРКИ ПЕРЕДУМОВ

  Програма автоматично перевіряє і БЛОКУЄ аналіз при порушеннях:

  ① Паралельність ліній регресії (КЛЮЧОВА ПЕРЕДУМОВА):
    ANCOVA передбачає що вплив коваріати на Y ОДНАКОВИЙ у всіх групах.
    Тест: взаємодія Група×Коваріата.
    p ≥ 0.05 → лінії паралельні → ANCOVA коректна ✓
    p < 0.05 → лінії НЕ паралельні → ANCOVA ЗАБЛОКОВАНА ✗
    (→ використайте звичайну ANOVA з коваріатою як фактором)

  ② Нормальність залишків (Shapiro-Wilk):
    p > 0.05 → залишки нормальні ✓
    p ≤ 0.05 → програма запитає підтвердження

  ③ Однорідність дисперсій (тест Левена):
    p ≥ 0.05 → дисперсії рівні ✓
    p < 0.05 → програма запитає підтвердження

  ④ Мультиколінеарність коваріат (r > 0.95):
    При дуже сильному зв'язку між коваріатами — попередження

КРОК 5. ІНТЕРПРЕТАЦІЯ РЕЗУЛЬТАТІВ

  Таблиця ANCOVA (Тип III SS):
    Джерело «Група» → p < 0.05: групи відрізняються після контролю коваріати
    Джерело «Коваріата» → p < 0.05: коваріата суттєво впливає на Y
    R² → частка варіації Y пояснена всією моделлю

  Скориговані середні (LS Means):
    Це ГОЛОВНИЙ результат ANCOVA!
    Прогнозоване середнє кожної групи за умови що всі групи
    мають ОДНАКОВЕ значення коваріати (= загальне середнє).
    Порівнюйте СКОРИГОВАНІ, а не нескориговані середні!

  Пост-хок (Бонферроні):
    p < 0.05 → пара груп значуще відрізняється за скоригованим середнім

КРОК 6. ГРАФІКИ ЗАЛИШКІВ
  Residuals vs Fitted: точки мають бути хаотично навколо нуля
  QQ-графік залишків: точки мають лежати на прямій
  ⚠ Патерн або вигин → модель порушена
"""

    def __init__(self, parent, gs):
        self.win = tk.Toplevel(parent)
        self.win.title("ANCOVA — Коваріаційний аналіз")
        self.win.geometry("980x680"); set_icon(self.win)
        self.gs = gs; self._build()

    def _build(self):
        # ── Панель інструментів (наш стандарт) ──────────────
        top = tk.Frame(self.win, padx=8, pady=6); top.pack(fill=tk.X)
        tk.Button(top, text="▶ Виконати", bg="#c62828", fg="white",
                  font=("Times New Roman", 13),
                  command=self._run).pack(side=tk.LEFT, padx=4)

        # Налаштування — спадне меню
        mb2 = tk.Menubutton(top, text="⚙ Налаштування ▾",
                            font=("Times New Roman", 11),
                            relief=tk.RAISED, bd=2)
        mb2.pack(side=tk.LEFT, padx=4)
        sm = tk.Menu(mb2, tearoff=0)
        sm.add_command(label="Додати рядок",      command=self._add_row)
        sm.add_command(label="Видалити рядок",    command=self._del_row)
        sm.add_separator()
        sm.add_command(label="🗑 Очистити таблицю", command=self._clear_table)
        mb2["menu"] = sm

        tk.Button(top, text="Вставити з буфера",
                  font=("Times New Roman", 11),
                  command=self._paste).pack(side=tk.LEFT, padx=4)
        tk.Label(top, text="α:", font=("Times New Roman", 12)).pack(side=tk.LEFT, padx=(10, 2))
        self.alpha_var = tk.StringVar(value="0.05")
        ttk.Combobox(top, textvariable=self.alpha_var, values=["0.01","0.05","0.10"],
                     state="readonly", width=7).pack(side=tk.LEFT)
        tk.Button(top, text="📚 Довідка", bg="#1a4b8c", fg="white",
                  font=("Times New Roman", 11),
                  command=self._show_help).pack(side=tk.LEFT, padx=8)

        # ── Інформаційний рядок ──────────────────────────────
        info = tk.Frame(self.win, bg="#f0f4ff", padx=8, pady=4)
        info.pack(fill=tk.X, padx=8, pady=(0, 4))
        tk.Label(info, text=(
            "Порядок стовпців:  [Група/Фактор]  [Коваріата 1]  [Коваріата 2 ...]  [Залежна Y]\n"
            "Заголовки стовпців (блакитні) можна редагувати.  "
            "Перший стовпець — текстові мітки груп.  Решта — числа."),
            font=("Times New Roman", 10), bg="#f0f4ff", justify="left").pack(anchor="w")

        # ── Таблиця даних ────────────────────────────────────
        mid = tk.Frame(self.win); mid.pack(fill=tk.BOTH, expand=True, padx=8)
        self.n_rows = 24; self.n_cols = 6
        canvas = tk.Canvas(mid); canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb = ttk.Scrollbar(mid, orient="vertical", command=canvas.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        canvas.configure(yscrollcommand=sb.set)
        self.inner = tk.Frame(canvas)
        canvas.create_window((0, 0), window=self.inner, anchor="nw")
        self.inner.bind("<Configure>", lambda e: canvas.config(scrollregion=canvas.bbox("all")))
        self.win.bind("<MouseWheel>",
                      lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

        col_hints = ["Група", "Коваріата 1", "Коваріата 2", "Коваріата 3", "Коваріата 4", "Залежна Y"]
        self.header_entries = []
        for j in range(self.n_cols):
            e = tk.Entry(self.inner, width=14, bg="#1a4b8c", fg="white",
                         font=("Times New Roman", 11, "bold"),
                         insertbackground="white")
            e.insert(0, col_hints[j] if j < len(col_hints) else f"Стовп{j+1}")
            e.grid(row=0, column=j, padx=1, pady=1)
            self.header_entries.append(e)
        self.entries = []
        for i in range(self.n_rows):
            row_ = []
            for j in range(self.n_cols):
                e = tk.Entry(self.inner, width=14, font=("Times New Roman", 11),
                             highlightthickness=1, highlightbackground="#c0c0c0")
                e.grid(row=i+1, column=j, padx=1, pady=1)
                row_.append(e)
            self.entries.append(row_)
        _bind_nav(self.entries, self.win)

    def _show_help(self):
        win = tk.Toplevel(self.win)
        win.title("Довідка — ANCOVA")
        win.geometry("700x640"); set_icon(win)
        frm = tk.Frame(win); frm.pack(fill=tk.BOTH, expand=True, padx=8, pady=6)
        vsb = ttk.Scrollbar(frm, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        txt = tk.Text(frm, wrap="word", font=("Times New Roman", 11),
                      yscrollcommand=vsb.set, relief=tk.FLAT,
                      bg="#fafafa", padx=10, pady=8, cursor="arrow")
        txt.pack(fill=tk.BOTH, expand=True)
        vsb.config(command=txt.yview)
        txt.insert("1.0", self.HELP_TEXT.strip())
        txt.configure(state="disabled")
        txt.bind("<MouseWheel>",
                 lambda e: txt.yview_scroll(int(-1*(e.delta/120)), "units"))
        tk.Button(win, text="Закрити", command=win.destroy,
                  font=("Times New Roman", 11)).pack(pady=6)

    def _help(self):
        self._show_help()   # залишаємо для сумісності

    def _paste(self):
        try: data = self.win.clipboard_get()
        except Exception:
            messagebox.showwarning("Буфер порожній",
                "Скопіюйте дані з Excel (Ctrl+C) і спробуйте знову."); return
        rows = [r for r in data.splitlines() if r.strip()]
        for i, line in enumerate(rows[:self.n_rows]):
            for j, val in enumerate(line.split("\t")[:self.n_cols]):
                self.entries[i][j].delete(0, tk.END)
                self.entries[i][j].insert(0, val.strip())

    def _add_row(self):
        i = self.n_rows; row_ = []
        for j in range(self.n_cols):
            e = tk.Entry(self.inner, width=14, font=("Times New Roman", 11),
                         highlightthickness=1, highlightbackground="#c0c0c0")
            e.grid(row=i+1, column=j, padx=1, pady=1)
            row_.append(e)
        self.entries.append(row_); self.n_rows += 1
        _bind_nav(self.entries, self.win)
        self.inner.update_idletasks()

    def _del_row(self):
        if not self.entries: return
        for e in self.entries.pop(): e.destroy()
        self.n_rows -= 1

    def _clear_table(self):
        if not messagebox.askyesno("Очистити таблицю",
                "Видалити всі числові дані?\n(Заголовки залишаться)"):
            return
        for row in self.entries:
            for e in row: e.delete(0, tk.END)



    def _run(self):
        alpha = float(self.alpha_var.get())

        # ── Зчитування заголовків ─────────────────────────────
        headers = [e.get().strip() or f"Col{j+1}"
                   for j, e in enumerate(self.header_entries)]

        # ── Зчитування даних ─────────────────────────────────
        raw = [[e.get().strip() for e in row] for row in self.entries]
        raw = [r for r in raw if any(v for v in r)]
        if not raw:
            messagebox.showwarning("Немає даних",
                "Будь ласка, введіть дані у таблицю."); return

        # Визначаємо які стовпці реально заповнені числами
        # Перший стовпець = Група (текст)
        # Останній заповнений числовий стовпець = Залежна Y
        # Між ними = Коваріати (лише заповнені стовпці)

        # Знаходимо останній стовпець з числовими даними
        n_data_cols = self.n_cols
        # Знаходимо реально заповнені числові стовпці (1..n_cols-1)
        filled_numeric_cols = []
        for j in range(1, self.n_cols):
            has_num = False
            for row in raw:
                v = row[j] if j < len(row) else ""
                if v:
                    try: float(v.replace(",",".")); has_num = True; break
                    except ValueError: pass
            if has_num:
                filled_numeric_cols.append(j)

        if not filled_numeric_cols:
            messagebox.showwarning("Немає числових даних",
                "Не знайдено числових даних у стовпцях 2 і далі.\n"
                "Переконайтесь що введено значення коваріати та залежної змінної."); return

        # Останній заповнений числовий стовпець = залежна Y
        dv_col_idx  = filled_numeric_cols[-1]
        # Всі інші заповнені числові стовпці = коваріати
        cov_col_idxs = filled_numeric_cols[:-1]

        group_col = headers[0]
        dv_col    = headers[dv_col_idx]
        cov_cols  = [headers[j] for j in cov_col_idxs]

        if not cov_col_idxs:
            messagebox.showwarning("Немає коваріат",
                "ANCOVA потребує хоча б одну коваріату.\n\n"
                "Структура таблиці:\n"
                "  Стовпець 1: Група (текстові мітки)\n"
                "  Стовпець 2: Коваріата (числові значення)\n"
                "  Стовпець 3 або далі: Залежна змінна Y\n\n"
                "Якщо у вас одна коваріата — заповніть стовпці 1, 2 і 3.\n"
                "Стовпці 4, 5, 6 залиште порожніми."); return

        # Зчитуємо рядки використовуючи лише знайдені стовпці
        groups = []; cov_data = [[] for _ in cov_col_idxs]; y_data = []
        skipped = 0
        for row in raw:
            while len(row) < self.n_cols: row.append("")
            grp = row[0].strip()
            if not grp: skipped += 1; continue
            # Перевіряємо чи грp не є числом (захист від плутанини)
            try:
                float(grp.replace(",","."))
                skipped += 1; continue  # перший стовпець має бути текстом
            except ValueError:
                pass
            try:
                covs = [float(row[j].replace(",",".")) for j in cov_col_idxs]
                yval = float(row[dv_col_idx].replace(",","."))
            except (ValueError, IndexError):
                skipped += 1; continue
            groups.append(grp)
            for j_idx, cv in enumerate(covs): cov_data[j_idx].append(cv)
            y_data.append(yval)

        if skipped > 0:
            messagebox.showinfo("Пропущені рядки",
                f"Пропущено {skipped} рядків (порожні або нечислові значення).")

        n = len(y_data)
        # ── Guard 1: мінімум спостережень ──
        if n < 6:
            messagebox.showwarning("Замало спостережень",
                f"ANCOVA потребує щонайменше 6 повних спостережень.\n"
                f"Знайдено: {n}.\n\n"
                f"Перевірте:\n"
                f"  • Перший стовпець містить текстові назви груп\n"
                f"  • Числові дані введені у правильні стовпці\n"
                f"  • Немає порожніх клітинок у заповнених рядках"); return

        # ── Guard 2: щонайменше 2 групи ──
        group_levels = first_seen(groups)
        k = len(group_levels)
        if k < 2:
            messagebox.showwarning("Лише одна група",
                "ANCOVA потребує щонайменше 2 групи.\n"
                "Перевірте що перший стовпець містить різні текстові мітки."); return


        from collections import Counter
        grp_counts = Counter(groups)
        min_grp = min(grp_counts.values())
        if min_grp < 2:
            messagebox.showwarning("Занадто мала група",
                f"Each group needs ≥ 2 observations.\n"
                f"Smallest group has {min_grp} observation(s)."); return

        # ── Guard 4: check covariates are numeric (already done in parsing) ──
        # ── Guard 5: check for perfect multicollinearity among covariates ──
        if len(cov_cols) > 1:
            cov_matrix = np.column_stack([np.array(cd) for cd in cov_data])
            corr_matrix = np.corrcoef(cov_matrix.T)
            for i in range(len(cov_cols)):
                for j in range(i+1, len(cov_cols)):
                    if abs(corr_matrix[i,j]) > 0.95:
                        ans = messagebox.askyesno("Висока мультиколінеарність",
                            f"Covariates '{cov_cols[i]}' and '{cov_cols[j]}' are highly correlated\n"
                            f"(r = {corr_matrix[i,j]:.3f}).\n\n"
                            "This may cause unstable estimates (multicollinearity problem).\n"
                            "Consider removing one of the covariates.\n\n"
                            "Continue anyway?")
                        if not ans: return

        y = np.array(y_data, dtype=float)
        covs_arr = [np.array(cd, dtype=float) for cd in cov_data]

        # ── Guard 6: Homogeneity of regression slopes ──
        # Test: fit model with group × covariate interaction
        # If interaction is significant → slopes differ → ANCOVA assumption violated
        slopes_ok = True
        slope_details = []
        for ci, (cov_name, cov_arr) in enumerate(zip(cov_cols, covs_arr)):
            # Build X with intercept + group dummies + covariate + group×covariate
            X_parts = [np.ones(n)]
            g_dummies = []
            for lv in group_levels[1:]:
                d = np.array([1. if g == lv else 0. for g in groups])
                X_parts.append(d); g_dummies.append(d)
            X_parts.append(cov_arr)
            # interaction terms: group_dummy × covariate
            for gd in g_dummies:
                X_parts.append(gd * cov_arr)
            X_int = np.column_stack(X_parts)
            X_no_int = np.column_stack(X_parts[:len(X_parts)-len(g_dummies)])

            _, _, _, sse_full, dfe_full, mse_full = _ols(y, X_int)
            _, _, _, sse_red,  dfe_red,  _        = _ols(y, X_no_int)

            df_int = len(g_dummies)
            ss_int = sse_red - sse_full
            ms_int = ss_int / df_int if df_int > 0 else np.nan
            F_int  = ms_int / mse_full if (not math.isnan(mse_full) and mse_full > 0) else np.nan
            p_int  = float(1 - f_dist.cdf(F_int, df_int, dfe_full)) if not math.isnan(F_int) else np.nan
            slope_details.append((cov_name, F_int, p_int))
            if not math.isnan(p_int) and p_int < alpha:
                slopes_ok = False

        if not slopes_ok:
            failed = [f"'{n}' (F={fmt(F,3)}, p={fmt(p,4)})"
                      for n, F, p in slope_details
                      if not math.isnan(p) and p < alpha]
            ans = messagebox.askyesno(
                "ANCOVA ASSUMPTION VIOLATED — Heterogeneous regression slopes",
                "The homogeneity of regression slopes assumption is VIOLATED for:\n"
                + "\n".join(f"  • {f}" for f in failed) + "\n\n"
                "This means the covariate affects the DV differently in different groups.\n"
                "Standard ANCOVA is NOT appropriate in this situation.\n\n"
                "Options:\n"
                "• Use Johnson-Neyman technique (regions of significance)\n"
                "• Run separate regressions per group\n"
                "• Use interaction ANOVA instead\n\n"
                "Do you want to continue with ANCOVA anyway? (NOT recommended)")
            if not ans: return

        # ── Build ANCOVA model (Type III SS) ──
        # X: intercept + group dummies + covariate(s)
        X_parts = [np.ones(n)]
        ts = {"Intercept": [0]}
        idx_cur = 1
        # group factor
        g_idx = []
        for lv in group_levels[1:]:
            d = np.array([1. if g == lv else 0. for g in groups])
            X_parts.append(d); g_idx.append(idx_cur); idx_cur += 1
        ts["Група"] = g_idx
        # covariates
        for cov_name, cov_arr in zip(cov_cols, covs_arr):
            cov_norm = (cov_arr - np.mean(cov_arr)) / (np.std(cov_arr, ddof=1) + 1e-12)
            X_parts.append(cov_norm)
            ts[f"Covariate: {cov_name}"] = [idx_cur]; idx_cur += 1

        X = np.column_stack(X_parts)
        beta, yhat, residuals, sse, dfe, mse = _ols(y, X)
        sst = float(np.sum((y - np.mean(y))**2))

        # Type III SS for each term
        anova_rows = []
        for term, idx_list in ts.items():
            if term == "Intercept": continue
            keep = [i for i in range(X.shape[1]) if i not in idx_list]
            _, _, _, sse_red, _, _ = _ols(y, X[:, keep])
            ss = float(sse_red - sse)
            df = len(idx_list)
            ms = ss / df if df > 0 else np.nan
            F  = ms / mse if (not math.isnan(mse) and mse > 0) else np.nan
            p  = float(1 - f_dist.cdf(F, df, dfe)) if not math.isnan(F) else np.nan
            mark = sig_mark(p) if not math.isnan(p) else ""
            concl = f"significant {mark}" if mark else ("незнач." if not math.isnan(p) else "–")
            anova_rows.append([term, fmt(ss,4), str(df), fmt(ms,4), fmt(F,4), fmt(p,4), concl])

        anova_rows.append(["Залишок", fmt(sse,4), str(dfe), fmt(mse,4), "", "", ""])
        anova_rows.append(["Загальна",    fmt(sst,4), str(n-1), "", "", "", ""])

        # ── Guard 7: Normality of residuals ──
        try: W_res, p_res = shapiro(residuals) if len(residuals) >= 3 else (np.nan, np.nan)
        except Exception: W_res, p_res = np.nan, np.nan
        if not math.isnan(p_res) and p_res <= alpha:
            ans = messagebox.askyesno("Non-normal residuals",
                f"Residuals do not follow a normal distribution\n"
                f"(Shapiro–Wilk: W={fmt(W_res,4)}, p={fmt(p_res,4)}).\n\n"
                "ANCOVA assumes normally distributed residuals.\n"
                "Consider: data transformation (log/√) before running ANCOVA.\n\n"
                "Continue anyway?")
            if not ans: return

        # ── Guard 8: Homogeneity of variances (Levene) ──
        grp_residuals = defaultdict(list)
        for g, r in zip(groups, residuals): grp_residuals[g].append(r)
        lev_F, lev_p = levene_test(dict(grp_residuals))
        if not math.isnan(lev_p) and lev_p < alpha:
            ans = messagebox.askyesno("Heterogeneous variances (Levene test)",
                f"Levene test: F={fmt(lev_F,4)}, p={fmt(lev_p,4)}\n\n"
                "Variances differ significantly across groups.\n"
                "ANCOVA is somewhat robust to this violation when group sizes are equal,\n"
                "but results may be unreliable with unequal group sizes.\n\n"
                "Continue anyway?")
            if not ans: return

        # ── Adjusted (LS) means ──
        # Compute adjusted means: predict at grand mean of each covariate
        cov_grand_means = [np.mean(ca) for ca in covs_arr]
        adj_means = {}
        for lv in group_levels:
            x_pred = [1.0]  # intercept
            for ref_lv in group_levels[1:]:
                x_pred.append(1.0 if lv == ref_lv else 0.0)
            for ca_mean, ca_arr in zip(cov_grand_means, covs_arr):
                cov_norm_mean = (ca_mean - np.mean(ca_arr)) / (np.std(ca_arr, ddof=1) + 1e-12)
                x_pred.append(cov_norm_mean)
            adj_means[lv] = float(np.dot(beta, x_pred))

        # unadjusted means
        raw_means = {lv: float(np.mean([y_data[i] for i, g in enumerate(groups) if g == lv]))
                     for lv in group_levels}

        R2 = 1 - sse/sst if sst > 0 else np.nan

        self._show_results(anova_rows, adj_means, raw_means, group_levels,
                           residuals, W_res, p_res, lev_F, lev_p,
                           slope_details, R2, mse, dfe, alpha, y, yhat, groups)

    def _show_results(self, anova_rows, adj_means, raw_means, group_levels,
                      residuals, W_res, p_res, lev_F, lev_p,
                      slope_details, R2, mse, dfe, alpha, y, yhat, groups):
        win = tk.Toplevel(self.win); win.title("ANCOVA — Результати")
        win.geometry("1150x760"); set_icon(win)

        # scrollable body
        main = tk.Frame(win); main.pack(fill=tk.BOTH, expand=True)
        vsb = ttk.Scrollbar(main, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        canvas = tk.Canvas(main, yscrollcommand=vsb.set); canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.config(command=canvas.yview)
        body = tk.Frame(canvas); canvas.create_window((0, 0), window=body, anchor="nw")
        body.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        win.bind("<MouseWheel>", lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

        def _head(txt):
            tk.Label(body, text=txt, font=("Times New Roman",12,"bold"),
                     anchor="w").pack(fill=tk.X, padx=10, pady=8)
        def _txt(txt, color="#000000"):
            tk.Label(body, text=txt, font=("Times New Roman",11), fg=color,
                     anchor="w", justify="left").pack(fill=tk.X, padx=10, pady=1)
        def _tbl(headers, rows):
            f, _ = make_tv(body, headers, rows); f.pack(fill=tk.X, padx=10, pady=2)

        _head("ANCOVA — Коваріаційний аналіз")
        _txt(f"R² = {fmt(R2,4)}   |   MSE = {fmt(mse,4)}   |   df_error = {dfe}")

        # Assumption checks
        _head("Assumption Checks")
        norm_color = "#000000" if math.isnan(p_res) or p_res > alpha else "#c62828"
        _txt(f"Нормальність залишків (Shapiro–Wilk):  W={fmt(W_res,4)},  p={fmt(p_res,4)}  "
             f"{'✓ OK' if not math.isnan(p_res) and p_res > alpha else '⚠ VIOLATED'}",
             norm_color)
        lev_color = "#000000" if math.isnan(lev_p) or lev_p >= alpha else "#c62828"
        _txt(f"Однорідність дисперсій (Левен):  F={fmt(lev_F,4)},  p={fmt(lev_p,4)}  "
             f"{'✓ OK' if not math.isnan(lev_p) and lev_p >= alpha else '⚠ VIOLATED'}",
             lev_color)
        for cov_name, F_sl, p_sl in slope_details:
            sl_ok = math.isnan(p_sl) or p_sl >= alpha
            sl_color = "#000000" if sl_ok else "#c62828"
            _txt(f"Однорідність нахилів ({cov_name}):  F={fmt(F_sl,4)},  p={fmt(p_sl,4)}  "
                 f"{'✓ OK' if sl_ok else '⚠ VIOLATED — slopes differ'}",
                 sl_color)

        # ANOVA table
        _head("ANCOVA Table (Type III SS)")
        _tbl(["Джерело","SS","df","MS","F","p","Висновок"], anova_rows)

        # Adjusted means
        _head("Group Means")
        means_rows = [[lv, fmt(raw_means[lv],4), fmt(adj_means[lv],4)]
                      for lv in group_levels]
        _tbl(["Група","Unadjusted Mean","Adjusted Mean (LS Mean)"], means_rows)

        # Pairwise comparisons of adjusted means (Bonferroni t-test on LS means)
        # SE for difference between two LS means ≈ sqrt(MSE * 2/n_harm)
        _head("Pairwise Comparisons of Adjusted Means (Bonferroni)")
        ph_rows = []
        m_tests = len(group_levels) * (len(group_levels)-1) / 2
        n_per_grp = {lv: groups.count(lv) for lv in group_levels}
        for lv1, lv2 in combinations(group_levels, 2):
            n1, n2 = n_per_grp[lv1], n_per_grp[lv2]
            se = math.sqrt(mse * (1/n1 + 1/n2)) if mse > 0 else np.nan
            diff = adj_means[lv1] - adj_means[lv2]
            if math.isnan(se) or se == 0:
                ph_rows.append([f"{lv1} vs {lv2}", fmt(diff,4), "–", "–", "–"]); continue
            t_val = abs(diff) / se
            p_raw = 2 * (1 - float(t_dist.cdf(t_val, dfe)))
            p_adj = min(1., p_raw * m_tests)
            mark  = sig_mark(p_adj)
            ph_rows.append([f"{lv1} vs {lv2}", fmt(diff,4), fmt(t_val,4), fmt(p_adj,4),
                             f"significant {mark}" if mark else "незнач."])
        _tbl(["Порівняння","Різниця","t","p (Bonf.)","Висновок"], ph_rows)

        # Plots
        if HAS_MPL:
            fig = Figure(figsize=(10, 6), dpi=100)
            ax1 = fig.add_subplot(121)
            ax1.scatter(yhat, residuals, s=22, color="#4c72b0", alpha=0.8)
            ax1.axhline(0, color="k", lw=0.8)
            ax1.set_xlabel("Fitted values"); ax1.set_ylabel("Residuals")
            ax1.set_title("Залишки vs Підігнані"); ax1.yaxis.grid(True, alpha=0.3)

            from scipy.stats import probplot
            ax2 = fig.add_subplot(122)
            res_sort = np.sort(residuals)
            rp = probplot(residuals, dist="norm")
            ax2.plot(rp[0][0], rp[0][1], 'o', markersize=4, color="#4c72b0")
            ax2.plot(rp[0][0], rp[1][1] + rp[1][0]*rp[0][0], 'r-', lw=1)
            ax2.set_xlabel("Теоретичні квантилі"); ax2.set_ylabel("Вибіркові квантилі")
            ax2.set_title("QQ-графік залишків"); ax2.yaxis.grid(True, alpha=0.3)
            fig.tight_layout()
            embed_figure(fig, body)


# ═══════════════════════════════════════════════════════════════
# MANOVA — Multivariate Analysis of Variance
# ═══════════════════════════════════════════════════════════════
class ManovaWindow:
    """MANOVA — Багатовимірний дисперсійний аналіз."""

    HELP_TEXT = """
MANOVA — ПОКРОКОВА ІНСТРУКЦІЯ
══════════════════════════════════════════

ЩО ТАКЕ MANOVA?
  MANOVA (Багатовимірний дисперсійний аналіз) перевіряє чи відрізняються
  групи одночасно за КІЛЬКОМА залежними змінними (показниками).

НАВІЩО НЕ КІЛЬКА ANOVA?
  Якщо провести окремі ANOVA для кожного показника:
  При 5 показниках і α=0.05 → ймовірність хоча б одного хибного
  результату = 1-(0.95)⁵ = 23%!
  MANOVA контролює цю сімейну помилку.
  Крім того MANOVA може виявити ефект який окремі ANOVA пропустять
  (коли ефект є у комбінації показників але не в кожному окремо).

КОЛИ ВИКОРИСТОВУВАТИ?
  ✓ Порівняння сортів за комплексом показників якості
    (врожайність + маса + цукристість + кислотність одночасно)
  ✓ Порівняння варіантів обробки за кількома параметрами росту
  ✓ Будь-коли коли у вас 2+ залежних показники і 2+ групи

КРОК 1. СТРУКТУРА ТАБЛИЦІ

  Перший стовпець: Група (текстові мітки: «Сорт А», «Контроль» тощо)
  Решта стовпців: Залежні змінні — по одній на стовпець (числа)

  Приклад (порівняння 3 сортів за 3 показниками):
  | Сорт    | Врожайність | Висота | Маса зерна |
  | Сорт А  |    5.8      |  95.3  |    38.2    |
  | Сорт А  |    6.1      |  98.1  |    40.5    |
  | Сорт Б  |    4.9      |  88.5  |    35.1    |
  | Сорт Б  |    5.2      |  91.2  |    36.8    |
  | Сорт В  |    6.8      | 102.4  |    43.7    |
  | Сорт В  |    7.1      | 105.8  |    45.2    |

  Мінімум: 2 залежних змінних, 2 групи.

КРОК 2. КРИТИЧНА ВИМОГА: n > p У КОЖНІЙ ГРУПІ

  n = кількість спостережень у групі
  p = кількість залежних змінних

  Якщо у групі 3 спостереження і 4 показники → n ≤ p → MANOVA неможлива!
  Програма ЗАБЛОКУЄ аналіз і пояснить що робити.

  Правило: на кожну залежну змінну потрібно щонайменше 10 спостережень.
  Наприклад: 3 ЗЗ → мінімум 10-15 спостережень на групу.

КРОК 3. АВТОМАТИЧНІ ПЕРЕВІРКИ

  Програма перевіряє 5 передумов:

  ① n > p у кожній групі (критична — блокування)
  ② ≥ 2 залежних змінних
  ③ Мультиколінеарність ЗЗ (|r| > 0.90 → попередження)
  ④ Багатовимірна нормальність (тест Мардіа):
     Перевіряє нормальність векторів спостережень одночасно.
     При порушенні → Pillai's Trace є найнадійнішою статистикою.
  ⑤ Однорідність коваріаційних матриць (Box's M тест):
     Аналог тесту Левена але для матриць, а не дисперсій.

КРОК 4. ІНТЕРПРЕТАЦІЯ РЕЗУЛЬТАТІВ

  Чотири тестові статистики:

  Wilks' Lambda (Λ):
    Найпоширеніша. Від 0 до 1. Менше → сильніший ефект.
    Рекомендується при нормальності і рівних коваріаційних матрицях.

  Pillai's Trace (V): ★ НАЙНАДІЙНІША
    Найробустніша до порушень передумов.
    При порушенні нормальності або Box's M → використовуйте її!
    Програма позначає автоматично.

  Hotelling-Lawley Trace (T):
    Потужна коли один ефект домінує над іншими.

  Roy's GCR:
    Найпотужніша але найменш надійна.
    p-значення — верхня межа, не точне.

  Якщо всі 4 статистики дають p < α → впевнений результат ✓
  Якщо результати суперечливі → орієнтуйтесь на Pillai's Trace

КРОК 5. ПРАВИЛЬНА ПОСЛІДОВНІСТЬ ІНТЕРПРЕТАЦІЇ

  1. Перевірте передумови → при порушеннях читайте попередження
  2. Оцініть Pillai's Trace (p < α → групи відрізняються)
  3. ЯКЩО MANOVA ЗНАЧУЩИЙ → переходьте до univariate ANOVA
  4. Univariate ANOVA використовують поправку Бонферроні: α / кількість ЗЗ
     (наприклад при 4 ЗЗ: 0.05/4 = 0.0125)
  5. ЯКЩО MANOVA НЕЗНАЧУЩИЙ → univariate тести НЕ інтерпретуються!

КРОК 6. РОЗМІР ЕФЕКТУ (partial η²)
  Виводиться у univariate результатах:
  < 0.01: дуже слабкий | 0.01-0.06: слабкий
  0.06-0.14: середній  | > 0.14: сильний
"""

    def __init__(self, parent, gs):
        self.win = tk.Toplevel(parent)
        self.win.title("MANOVA — Багатовимірний дисперсійний аналіз")
        self.win.geometry("1020x700"); set_icon(self.win)
        self.gs = gs; self._build()

    def _build(self):
        # ── Панель інструментів (наш стандарт) ──────────────
        top = tk.Frame(self.win, padx=8, pady=6); top.pack(fill=tk.X)
        tk.Button(top, text="▶ Виконати", bg="#c62828", fg="white",
                  font=("Times New Roman", 13),
                  command=self._run).pack(side=tk.LEFT, padx=4)

        # Налаштування — спадне меню
        mb2 = tk.Menubutton(top, text="⚙ Налаштування ▾",
                            font=("Times New Roman", 11),
                            relief=tk.RAISED, bd=2)
        mb2.pack(side=tk.LEFT, padx=4)
        sm = tk.Menu(mb2, tearoff=0)
        sm.add_command(label="Додати рядок",      command=self._add_row)
        sm.add_command(label="Видалити рядок",    command=self._del_row)
        sm.add_separator()
        sm.add_command(label="Додати стовпець",    command=self._add_col)
        sm.add_command(label="Видалити стовпець",  command=self._del_col)
        sm.add_separator()
        sm.add_command(label="🗑 Очистити таблицю", command=self._clear_table)
        mb2["menu"] = sm

        tk.Button(top, text="Вставити з буфера",
                  font=("Times New Roman", 11),
                  command=self._paste).pack(side=tk.LEFT, padx=4)
        tk.Label(top, text="α:", font=("Times New Roman", 12)).pack(side=tk.LEFT, padx=(10, 2))
        self.alpha_var = tk.StringVar(value="0.05")
        ttk.Combobox(top, textvariable=self.alpha_var, values=["0.01","0.05","0.10"],
                     state="readonly", width=7).pack(side=tk.LEFT)
        tk.Button(top, text="📚 Довідка", bg="#1a4b8c", fg="white",
                  font=("Times New Roman", 11),
                  command=self._show_help).pack(side=tk.LEFT, padx=8)

        # ── Інформаційний рядок ──────────────────────────────
        info = tk.Frame(self.win, bg="#f0f4ff", padx=8, pady=4)
        info.pack(fill=tk.X, padx=8, pady=(0, 4))
        tk.Label(info, text=(
            "Порядок стовпців:  [Група/Фактор]  [Залежна змінна 1]  [Залежна змінна 2]  ...\n"
            "Заголовки (блакитні) можна редагувати.  Перший стовпець — текстові мітки груп.  "
            "Мінімум: 1 група + 2 залежних змінних.  Критично: n > p у кожній групі."),
            font=("Times New Roman", 10), bg="#f0f4ff", justify="left").pack(anchor="w")

        # ── Таблиця даних ────────────────────────────────────
        mid = tk.Frame(self.win); mid.pack(fill=tk.BOTH, expand=True, padx=8)
        self.n_rows = 24; self.n_cols = 8
        self._canvas = tk.Canvas(mid)
        self._canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb = ttk.Scrollbar(mid, orient="vertical", command=self._canvas.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        self._canvas.configure(yscrollcommand=sb.set)
        self.inner = tk.Frame(self._canvas)
        self._canvas.create_window((0, 0), window=self.inner, anchor="nw")
        self.inner.bind("<Configure>",
                        lambda e: self._canvas.config(scrollregion=self._canvas.bbox("all")))
        self.win.bind("<MouseWheel>",
                      lambda e: self._canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

        col_hints = ["Група","Показник 1","Показник 2","Показник 3",
                     "Показник 4","Показник 5","Показник 6","Показник 7"]
        self.header_vars = []
        self.header_entries = []   # for compatibility (used in _run via header_entries)
        for j in range(self.n_cols):
            var = tk.StringVar(value=col_hints[j] if j < len(col_hints) else f"Показник {j}")
            self.header_vars.append(var)
            lbl = tk.Label(self.inner, textvariable=var, width=13,
                           bg="#1a4b8c", fg="white", cursor="hand2",
                           font=("Times New Roman",11,"bold"), relief=tk.RIDGE)
            lbl.grid(row=0, column=j, padx=1, pady=1, sticky="nsew")
            lbl.bind("<Double-Button-1>", lambda e, idx=j: self._rename_manova_col(idx))
            self.header_entries.append(lbl)   # dummy for _run compatibility

        self.entries = []
        for i in range(self.n_rows):
            row_ = []
            for j in range(self.n_cols):
                e = tk.Entry(self.inner, width=13, font=("Times New Roman", 11),
                             highlightthickness=1, highlightbackground="#c0c0c0")
                e.grid(row=i+1, column=j, padx=1, pady=1)
                row_.append(e)
            self.entries.append(row_)
        _bind_nav(self.entries, self.win)

    # ── Довідка ───────────────────────────────────────────────
    def _show_help(self):
        win = tk.Toplevel(self.win)
        win.title("Довідка — MANOVA")
        win.geometry("720x660"); set_icon(win)
        frm = tk.Frame(win); frm.pack(fill=tk.BOTH, expand=True, padx=8, pady=6)
        vsb = ttk.Scrollbar(frm, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        txt = tk.Text(frm, wrap="word", font=("Times New Roman", 11),
                      yscrollcommand=vsb.set, relief=tk.FLAT,
                      bg="#fafafa", padx=10, pady=8, cursor="arrow")
        txt.pack(fill=tk.BOTH, expand=True)
        vsb.config(command=txt.yview)
        txt.insert("1.0", self.HELP_TEXT.strip())
        txt.configure(state="disabled")
        txt.bind("<MouseWheel>",
                 lambda e: txt.yview_scroll(int(-1*(e.delta/120)), "units"))
        tk.Button(win, text="Закрити", command=win.destroy,
                  font=("Times New Roman", 11)).pack(pady=6)

    def _help(self):
        self._show_help()

    # ── Управління таблицею ───────────────────────────────────
    def _add_row(self):
        i = self.n_rows; row_ = []
        for j in range(self.n_cols):
            e = tk.Entry(self.inner, width=13, font=("Times New Roman", 11),
                         highlightthickness=1, highlightbackground="#c0c0c0")
            e.grid(row=i+1, column=j, padx=1, pady=1)
            row_.append(e)
        self.entries.append(row_); self.n_rows += 1
        _bind_nav(self.entries, self.win)
        self.inner.update_idletasks()

    def _del_row(self):
        if not self.entries: return
        for e in self.entries.pop(): e.destroy()
        self.n_rows -= 1

    def _add_col(self):
        ci = self.n_cols; self.n_cols += 1
        var = tk.StringVar(value=f"Показник {ci}")
        self.header_vars.append(var)
        lbl = tk.Label(self.inner, textvariable=var, width=13,
                       bg="#1a4b8c", fg="white", cursor="hand2",
                       font=("Times New Roman",11,"bold"), relief=tk.RIDGE)
        lbl.grid(row=0, column=ci, padx=1, pady=1, sticky="nsew")
        lbl.bind("<Double-Button-1>", lambda e, idx=ci: self._rename_manova_col(idx))
        self.header_entries.append(lbl)
        for i, row_ in enumerate(self.entries):
            e = tk.Entry(self.inner, width=13, font=("Times New Roman", 11),
                         highlightthickness=1, highlightbackground="#c0c0c0")
            e.grid(row=i+1, column=ci, padx=1, pady=1)
            row_.append(e)
        _bind_nav(self.entries, self.win)

    def _del_col(self):
        if self.n_cols <= 3: return
        self.header_entries.pop().destroy()
        self.header_vars.pop()
        for row_ in self.entries: row_.pop().destroy()
        self.n_cols -= 1

    def _clear_table(self):
        if not messagebox.askyesno("Очистити таблицю",
                "Видалити всі числові дані?\n(Заголовки залишаться)"):
            return
        for row in self.entries:
            for e in row: e.delete(0, tk.END)

    # ── Вставка з буфера ──────────────────────────────────────
    def _paste(self):
        try: data = self.win.clipboard_get()
        except Exception:
            messagebox.showwarning("Буфер порожній",
                "Скопіюйте дані з Excel (Ctrl+C) і спробуйте знову."); return
        if not data.strip(): return
        for i, line in enumerate(data.splitlines()):
            if not line.strip(): continue
            if i >= len(self.entries): self._add_row()
            for j, val in enumerate(line.split("\t")[:self.n_cols]):
                self.entries[i][j].delete(0, tk.END)
                self.entries[i][j].insert(0, val.strip())



    def _mardia_test(self, X):
        """Mardia's multivariate normality: skewness and kurtosis tests."""
        n, p = X.shape
        if n < p + 1: return np.nan, np.nan, np.nan, np.nan
        X_c = X - np.mean(X, axis=0)
        S = np.cov(X.T, ddof=1)
        try:
            S_inv = np.linalg.pinv(S)
        except Exception: return np.nan, np.nan, np.nan, np.nan
        # Mahalanobis distances
        D = X_c @ S_inv @ X_c.T
        # Mardia skewness
        b1p = float(np.sum(D**3)) / n**2
        k_sk = n * b1p / 6.0
        df_sk = p*(p+1)*(p+2)/6
        p_sk  = float(1 - f_dist.cdf(k_sk/df_sk, df_sk, 1e6)) if df_sk > 0 else np.nan
        # Mardia kurtosis
        b2p = float(np.sum(np.diag(D)**2)) / n
        k_ku = (b2p - p*(p+2)) / math.sqrt(8*p*(p+2)/n)
        from scipy.stats import norm as _norm
        p_ku = float(2 * (1 - _norm.cdf(abs(k_ku))))
        return float(b1p), p_sk, float(b2p), p_ku

    def _box_m_test(self, groups_data, group_levels):
        """Box's M test for homogeneity of covariance matrices (approximate)."""
        k = len(group_levels)
        ns = [len(groups_data[lv]) for lv in group_levels]
        p  = groups_data[group_levels[0]].shape[1]
        covs = []
        for lv in group_levels:
            Xl = groups_data[lv]
            if len(Xl) < 2: return np.nan, np.nan
            covs.append(np.cov(Xl.T, ddof=1))
        # Pooled covariance
        n_tot = sum(ns)
        S_pool = sum((n-1) * C for n, C in zip(ns, covs)) / (n_tot - k)
        try:
            ln_det_pool = np.log(max(np.linalg.det(S_pool), 1e-300))
            ln_dets     = [np.log(max(np.linalg.det(C), 1e-300)) for C in covs]
        except Exception: return np.nan, np.nan
        M = (n_tot - k) * ln_det_pool - sum((n-1)*d for n,d in zip(ns, ln_dets))
        # c1 correction
        c1 = (sum(1/(n-1) for n in ns) - 1/(n_tot-k)) * (2*p**2 + 3*p - 1) / (6*(p+1)*(k-1))
        chi2 = M * (1 - c1)
        df_m  = p*(p+1)*(k-1)/2
        from scipy.stats import chi2 as chi2_dist
        p_m   = float(1 - chi2_dist.cdf(chi2, df_m)) if df_m > 0 else np.nan
        return float(chi2), p_m

    def _run(self):
        alpha = float(self.alpha_var.get())
        # Заголовки з header_vars (tk.StringVar) або header_entries (tk.Entry/Label)
        headers = []
        for j in range(self.n_cols):
            if hasattr(self,'header_vars') and j < len(self.header_vars):
                headers.append(self.header_vars[j].get().strip() or f"Показник {j}")
            else:
                e = self.header_entries[j]
                headers.append(e.get().strip() if hasattr(e,'get') else f"Показник {j}")

        # ── Зчитування даних ─────────────────────────────────
        raw = [[e.get().strip() for e in row] for row in self.entries]
        raw = [r for r in raw if any(v for v in r)]
        if not raw:
            messagebox.showwarning("Немає даних",
                "Будь ласка, введіть дані у таблицю."); return

        groups = []; dv_rows = []; skipped = 0
        for row in raw:
            if len(row) < 2: skipped += 1; continue
            grp = row[0].strip()
            if not grp: skipped += 1; continue
            # Перший стовпець не має бути числом
            try:
                float(grp.replace(",",".")); skipped += 1; continue
            except ValueError: pass
            vals = []
            for v in row[1:]:
                if not v: continue   # пропускаємо порожні, НЕ зупиняємось
                try: vals.append(float(v.replace(",",".")))
                except ValueError: continue
            if len(vals) >= 2:
                groups.append(grp); dv_rows.append(vals)
            else: skipped += 1

        if skipped:
            messagebox.showinfo("Пропущені рядки",
                f"Пропущено {skipped} рядків (порожні або нечислові значення).")

        n = len(dv_rows)
        if n < 4:
            messagebox.showwarning("Замало даних",
                "Потрібно щонайменше 4 повних спостереження."); return

        # Вирівнюємо кількість ЗЗ — беремо мінімальну
        min_dv = min(len(r) for r in dv_rows)
        if min_dv < 2:
            messagebox.showwarning("Замало залежних змінних",
                "Потрібно щонайменше 2 залежних змінних (показники).\n"
                "Переконайтесь що числові дані введені у стовпці 2 і далі."); return

        Y = np.array([r[:min_dv] for r in dv_rows], dtype=float)
        p = min_dv
        dv_names_used = [headers[j+1] if j+1 < len(headers) else f"ЗЗ{j+1}"
                         for j in range(p)]

        group_levels = first_seen(groups)
        k = len(group_levels)

        # ── Guard 1: щонайменше 2 групи ──
        if k < 2:
            messagebox.showwarning("Лише одна група",
                "MANOVA потребує щонайменше 2 групи.\n"
                "Перевірте що перший стовпець містить різні текстові мітки."); return

        # ── Guard 2: n > p у кожній групі (критична!) ──
        groups_data = {}
        for lv in group_levels:
            idx_ = [i for i, g in enumerate(groups) if g == lv]
            groups_data[lv] = Y[idx_]
        for lv in group_levels:
            n_lv = len(groups_data[lv])
            if n_lv <= p:
                messagebox.showerror("ПОРУШЕННЯ: n ≤ p",
                    f"Група '{lv}' має {n_lv} спостережень, але {p} залежних змінних.\n\n"
                    f"MANOVA потребує n > p (спостережень > залежних змінних)\n"
                    f"у КОЖНІЙ групі.\n\n"
                    f"Причина: коваріаційна матриця всередині групи є виродженою\n"
                    f"(singular) і не може бути обернена — MANOVA математично неможлива.\n\n"
                    f"Рішення:\n"
                    f"  • Збільшіть кількість спостережень (рядків) у кожній групі\n"
                    f"    (потрібно мінімум {p+1} спостережень на групу при {p} ЗЗ)\n"
                    f"  • Зменшіть кількість залежних змінних\n"
                    f"  • Спочатку виконайте PCA → аналізуйте головні компоненти")
                return

        # ── Guard 3: мультиколінеарність ЗЗ ──
        corr_Y = np.corrcoef(Y.T)
        high_corr_pairs = []
        for i in range(p):
            for j in range(i+1, p):
                if abs(corr_Y[i,j]) > 0.90:
                    high_corr_pairs.append(
                        (dv_names_used[i], dv_names_used[j], corr_Y[i,j]))
        if high_corr_pairs:
            details = "\n".join(
                f"  • '{a}' та '{b}': r = {c:.3f}" for a,b,c in high_corr_pairs)
            ans = messagebox.askyesno("Висока мультиколінеарність між ЗЗ",
                "Наступні пари залежних змінних сильно корелюють (|r| > 0.90):\n"
                + details + "\n\n"
                "Висока мультиколінеарність знижує потужність MANOVA і може\n"
                "призвести до виродженої коваріаційної матриці.\n\n"
                "Рекомендація: видаліть одну зі змінних або спочатку виконайте PCA.\n\n"
                "Продовжити попри це?")
            if not ans: return

        # ── Guard 4: Багатовимірна нормальність (Мардіа) ──
        Y_res = np.vstack([groups_data[lv] - np.mean(groups_data[lv], axis=0)
                            for lv in group_levels])
        b1p, p_sk, b2p, p_ku = self._mardia_test(Y_res)
        mv_normal = True
        if (not math.isnan(p_sk) and p_sk < alpha) or (not math.isnan(p_ku) and p_ku < alpha):
            mv_normal = False
            ans = messagebox.askyesno(
                "Порушення багатовимірної нормальності (тест Мардіа)",
                f"Тест Мардіа — асиметрія: b1p={fmt(b1p,4)},  p={fmt(p_sk,4)}\n"
                f"Тест Мардіа — ексцес:    b2p={fmt(b2p,4)},  p={fmt(p_ku,4)}\n\n"
                "Передумова багатовимірної нормальності порушена.\n"
                "Найнадійніша статистика у цьому випадку — Pillai's Trace.\n\n"
                "Примітка: MANOVA достатньо робастна при великих вибірках (n > 20 на групу).\n\n"
                "Продовжити? (Pillai's Trace буде позначено як найнадійніша)")
            if not ans: return

        # ── Guard 5: Box's M тест ──
        box_chi2, box_p = self._box_m_test(groups_data, group_levels)
        if not math.isnan(box_p) and box_p < 0.001:
            ans = messagebox.askyesno(
                "Неоднорідність коваріаційних матриць (Box's M)",
                f"Box's M: χ²={fmt(box_chi2,4)},  p={fmt(box_p,6)}\n\n"
                "Коваріаційні матриці значущо відрізняються між групами.\n"
                "Примітка: Box's M дуже чутливий до ненормальності.\n"
                "Якщо p лише трохи < 0.001 — це може бути хибний сигнал.\n\n"
                "Pillai's Trace є найробустнішою при порушенні цієї передумови.\n\n"
                "Продовжити?")
            if not ans: return

        # ══ MANOVA computation ══
        # Between-groups matrix H and Within-groups matrix E
        grand_mean = np.mean(Y, axis=0)
        E = np.zeros((p, p))  # within (error)
        H = np.zeros((p, p))  # between (hypothesis)
        for lv in group_levels:
            Xl = groups_data[lv]
            nl = len(Xl)
            grp_mean = np.mean(Xl, axis=0)
            E += (Xl - grp_mean).T @ (Xl - grp_mean)
            H += nl * np.outer(grp_mean - grand_mean, grp_mean - grand_mean)

        df_h = k - 1        # between df
        df_e = n - k        # within df

        # Eigenvalues of E⁻¹H
        try:
            E_inv = np.linalg.pinv(E)
            EinvH = E_inv @ H
            eigenvalues = np.real(np.linalg.eigvals(EinvH))
            eigenvalues = np.sort(eigenvalues[eigenvalues > 1e-10])[::-1]
        except Exception as ex:
            messagebox.showerror("Помилка обчислення",
                f"Could not compute eigenvalues: {ex}\n"
                "Check for singular covariance matrix (n ≤ p in some group)."); return

        if len(eigenvalues) == 0:
            messagebox.showerror("Немає дійсних власних значень",
                "The covariance matrix is singular. Cannot compute MANOVA.\n"
                "Ensure n > p in every group."); return

        s = min(df_h, p)   # number of non-zero eigenvalues

        # ── Four test statistics ──
        # Wilks' Lambda
        wilks_L = float(np.prod(1 / (1 + eigenvalues[:s])))
        # Approximate F for Wilks
        m_w = df_e + df_h - (p + df_h + 1) / 2
        q_w = math.sqrt((p**2 * df_h**2 - 4) / (p**2 + df_h**2 - 5)) if (p**2 + df_h**2 - 5) > 0 else 1
        df1_w = p * df_h
        df2_w = m_w * q_w - p * df_h / 2 + 1
        F_wilks = ((1 - wilks_L**(1/q_w)) / (wilks_L**(1/q_w))) * (df2_w / df1_w) if (wilks_L > 0 and q_w > 0) else np.nan
        p_wilks = float(1 - f_dist.cdf(F_wilks, df1_w, df2_w)) if not math.isnan(F_wilks) else np.nan

        # Pillai's Trace
        pillai_V = float(np.sum(eigenvalues[:s] / (1 + eigenvalues[:s])))
        # Approximate F for Pillai
        m_p = max(p, df_h)
        F_pillai = (pillai_V / s) / ((s - pillai_V) / s) * ((df_e + df_h - m_p - 1) / m_p) if (s - pillai_V > 0 and m_p > 0) else np.nan
        df1_p = s * m_p; df2_p = s * (df_e + df_h - m_p - 1)
        p_pillai = float(1 - f_dist.cdf(F_pillai, df1_p, df2_p)) if not math.isnan(F_pillai) else np.nan

        # Hotelling-Lawley Trace
        hl_T = float(np.sum(eigenvalues[:s]))
        b_hl = (df_e + df_h - p - 1) * hl_T / s if s > 0 else np.nan
        df1_hl = s * p; df2_hl = s * (df_e + df_h - p - 1)
        F_hl   = b_hl * (df2_hl / (df1_hl * s)) if (not math.isnan(b_hl) and df1_hl > 0 and s > 0) else np.nan
        p_hl   = float(1 - f_dist.cdf(F_hl, df1_hl, df2_hl)) if not math.isnan(F_hl) else np.nan

        # Roy's GCR
        roy_GCR = float(eigenvalues[0]) if len(eigenvalues) > 0 else np.nan
        # Upper bound F for Roy
        F_roy = roy_GCR * df_e / p if p > 0 else np.nan
        p_roy = float(1 - f_dist.cdf(F_roy, p, df_e)) if not math.isnan(F_roy) else np.nan

        # ── Univariate follow-up ANOVAs ──
        univ_rows = []
        bonf_alpha = alpha / p  # Bonferroni correction
        for dv_i, dv_nm in enumerate(dv_names_used):
            y_i = Y[:, dv_i]
            grand_m = np.mean(y_i)
            ss_b = sum(len(groups_data[lv]) * (np.mean(groups_data[lv][:,dv_i]) - grand_m)**2
                       for lv in group_levels)
            ss_w = sum(np.sum((groups_data[lv][:,dv_i] - np.mean(groups_data[lv][:,dv_i]))**2)
                       for lv in group_levels)
            ms_b = ss_b / df_h if df_h > 0 else np.nan
            ms_w = ss_w / df_e if df_e > 0 else np.nan
            F_i  = ms_b / ms_w if (not math.isnan(ms_w) and ms_w > 0) else np.nan
            p_i  = float(1 - f_dist.cdf(F_i, df_h, df_e)) if not math.isnan(F_i) else np.nan
            eta2_i = ss_b / (ss_b + ss_w) if (ss_b + ss_w) > 0 else np.nan
            mark_bonf = "значуще" if (not math.isnan(p_i) and p_i < bonf_alpha) else "незнач."
            univ_rows.append([dv_nm, fmt(F_i,4), fmt(p_i,4),
                              fmt(eta2_i,4), eta2_label(eta2_i),
                              f"α_Bonf = {fmt(bonf_alpha,4)}", mark_bonf])

        self._show_results(wilks_L, F_wilks, p_wilks,
                           pillai_V, F_pillai, p_pillai,
                           hl_T, F_hl, p_hl,
                           roy_GCR, F_roy, p_roy,
                           df1_w, df2_w, df1_p, df2_p, df1_hl, df2_hl,
                           univ_rows, dv_names_used,
                           b1p, p_sk, b2p, p_ku, box_chi2, box_p,
                           alpha, mv_normal, groups_data, group_levels, Y, p)

    def _show_results(self, wilks_L, F_wilks, p_wilks,
                      pillai_V, F_pillai, p_pillai,
                      hl_T, F_hl, p_hl,
                      roy_GCR, F_roy, p_roy,
                      df1_w, df2_w, df1_p, df2_p, df1_hl, df2_hl,
                      univ_rows, dv_names,
                      b1p, p_sk, b2p, p_ku, box_chi2, box_p,
                      alpha, mv_normal, groups_data, group_levels, Y, p):
        win = tk.Toplevel(self.win); win.title("MANOVA — Результати")
        win.geometry("1180x820"); set_icon(win)

        # ── Toolbar результатів ─────────────────────────────
        self._manova_figs = {}   # зберігаємо фігури для копіювання
        self._manova_colors = ["#4c72b0","#dd8452","#55a868","#c44e52","#8172b2","#937860"]

        tb_res = tk.Frame(win, padx=6, pady=5); tb_res.pack(fill=tk.X)
        tk.Button(tb_res, text="📋 Копіювати звіт (текст)", font=("Times New Roman",11),
                  command=lambda: self._copy_manova_text(win)).pack(side=tk.LEFT, padx=4)
        tk.Button(tb_res, text="📋 Копіювати графік 1", font=("Times New Roman",11),
                  command=lambda: self._copy_manova_fig(1)).pack(side=tk.LEFT, padx=4)
        tk.Button(tb_res, text="📋 Копіювати графік 2", font=("Times New Roman",11),
                  command=lambda: self._copy_manova_fig(2)).pack(side=tk.LEFT, padx=4)
        tk.Button(tb_res, text="⚙ Налаштування графіків", font=("Times New Roman",11),
                  command=lambda: self._restyle_manova(win, dv_names, groups_data,
                                                        group_levels, univ_rows,
                                                        alpha, p_pillai)).pack(side=tk.LEFT, padx=4)

        # ── Прокручуване тіло ────────────────────────────────
        main = tk.Frame(win); main.pack(fill=tk.BOTH, expand=True)
        vsb = ttk.Scrollbar(main, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        canvas = tk.Canvas(main, yscrollcommand=vsb.set)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.config(command=canvas.yview)
        self._manova_body = body = tk.Frame(canvas)
        canvas.create_window((0,0), window=body, anchor="nw")
        body.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        win.bind("<MouseWheel>", lambda e: canvas.yview_scroll(int(-1*(e.delta/120)),"units"))

        def _head(txt):
            tk.Label(body, text=txt, font=("Times New Roman",12,"bold"),
                     bg="#e8eeff", anchor="w", padx=8, pady=3
                     ).pack(fill=tk.X, padx=6, pady=10)
        def _txt(txt, color="#000000"):
            tk.Label(body, text=txt, font=("Times New Roman",11), fg=color,
                     anchor="w", justify="left").pack(fill=tk.X, padx=14, pady=1)
        def _tbl(headers, rows):
            f, _ = make_tv(body, headers, rows); f.pack(fill=tk.X, padx=10, pady=2)

        # ── Заголовок ────────────────────────────────────────
        tk.Label(body,
                 text=f"MANOVA — Багатовимірний дисперсійний аналіз    α = {alpha}",
                 font=("Times New Roman",13,"bold"), anchor="w", padx=10, pady=6
                 ).pack(fill=tk.X)

        # ── sig_mark з урахуванням обраного α ─────────────────
        def _sig(p_val):
            """Позначення значущості відповідно до обраного рівня α."""
            if p_val is None or math.isnan(p_val): return ""
            if p_val < alpha * 0.2:   return "**"   # суттєво нижче α (умовно «дуже значущий»)
            if p_val < alpha:         return "*"
            return "–"

        def _sig_color(p_val):
            if p_val is None or math.isnan(p_val): return "#000000"
            return "#1a6b1a" if p_val < alpha else "#000000"

        # ── Перевірка передумов ───────────────────────────────
        _head("Перевірка передумов")

        norm_sk_ok = math.isnan(p_sk) or p_sk > alpha
        norm_ku_ok = math.isnan(p_ku) or p_ku > alpha
        box_ok     = math.isnan(box_p) or box_p >= 0.001

        _txt(f"Тест Мардіа — асиметрія (нормальність):   b1p = {fmt(b1p,4)},  p = {fmt(p_sk,4)}  "
             f"{'✓ ОК' if norm_sk_ok else '⚠ порушено'}",
             "#000000" if norm_sk_ok else "#c62828")
        _txt(f"Тест Мардіа — ексцес (нормальність):      b2p = {fmt(b2p,4)},  p = {fmt(p_ku,4)}  "
             f"{'✓ ОК' if norm_ku_ok else '⚠ порушено'}",
             "#000000" if norm_ku_ok else "#c62828")
        _txt(f"Box's M (однорідність коваріаційних матриць):  χ² = {fmt(box_chi2,4)},  p = {fmt(box_p,6)}  "
             f"{'✓ ОК' if box_ok else '⚠ значущо (Box M чутливий до ненормальності)'}",
             "#000000" if box_ok else "#b07000")
        if not mv_normal:
            _txt("⚠ Багатовимірна нормальність порушена → Pillai's Trace є найнадійнішою статистикою.",
                 "#c62828")

        # ── Тестові статистики MANOVA ─────────────────────────
        _head("Тестові статистики MANOVA")

        # η² (розмір ефекту) для кожної статистики MANOVA
        # Wilks: η² = 1 - Λ^(1/s), де s = min(df_h, p)
        # Pillai: V вже є мірою ефекту (0-1), η² ≈ V/s
        # Hotelling: η² = T/(T+1) (наближення)
        # Roy: η² = θ/(1+θ)
        s_val = min(len(group_levels)-1, p)
        try:
            eta2_wilks = 1 - wilks_L**(1/max(s_val,1)) if not math.isnan(wilks_L) and wilks_L > 0 else np.nan
        except Exception: eta2_wilks = np.nan
        eta2_pillai = pillai_V / max(s_val,1) if not math.isnan(pillai_V) else np.nan
        eta2_hl    = hl_T / (hl_T+1) if not math.isnan(hl_T) and hl_T >= 0 else np.nan
        eta2_roy   = roy_GCR / (roy_GCR+1) if not math.isnan(roy_GCR) and roy_GCR >= 0 else np.nan

        recommended = "Pillai" if not mv_normal else "Wilks"
        manova_rows = [
            ["Wilks' Lambda",
             fmt(wilks_L,6), fmt(F_wilks,4),
             f"{int(df1_w)},{int(df2_w)}" if not math.isnan(df1_w) else "–",
             fmt(p_wilks,4), _sig(p_wilks),
             fmt(eta2_wilks,4), eta2_label(eta2_wilks),
             "★ стандарт" if recommended=="Wilks" else ""],
            ["Pillai's Trace",
             fmt(pillai_V,6), fmt(F_pillai,4),
             f"{int(df1_p)},{int(df2_p)}" if not math.isnan(df1_p) else "–",
             fmt(p_pillai,4), _sig(p_pillai),
             fmt(eta2_pillai,4), eta2_label(eta2_pillai),
             "★ найробустніша" if recommended=="Pillai" else "робастна"],
            ["Hotelling-Lawley",
             fmt(hl_T,6), fmt(F_hl,4),
             f"{int(df1_hl)},{int(df2_hl)}" if not math.isnan(df1_hl) else "–",
             fmt(p_hl,4), _sig(p_hl),
             fmt(eta2_hl,4), eta2_label(eta2_hl), ""],
            ["Roy's GCR",
             fmt(roy_GCR,6), fmt(F_roy,4), "–",
             fmt(p_roy,4), _sig(p_roy),
             fmt(eta2_roy,4), eta2_label(eta2_roy), "верхня межа"],
        ]
        _tbl(["Статистика","Значення","F","df","p",f"Знач.(α={alpha})",
              "partial η²","Сила ефекту","Примітка"], manova_rows)

        # Опис сили ефекту
        _txt("Сила ефекту (partial η²): < 0.01 дуже слабкий | 0.01–0.06 слабкий | "
             "0.06–0.14 середній | > 0.14 сильний", "#555555")

        # ── Висновок по MANOVA ────────────────────────────────
        _head("Висновок")
        if not math.isnan(p_pillai):
            if p_pillai < alpha:
                _txt(f"✓ MANOVA значущий (Pillai p = {fmt(p_pillai,4)} < α = {alpha}):\n"
                     f"  Групи значуще відрізняються за комбінацією залежних змінних.\n"
                     f"  Перейдіть до одновимірних тестів (таблиця нижче).", "#1a6b1a")
            else:
                _txt(f"✗ MANOVA незначущий (Pillai p = {fmt(p_pillai,4)} ≥ α = {alpha}):\n"
                     f"  Немає достатніх підстав вважати що групи відрізняються за\n"
                     f"  комбінацією залежних змінних.\n"
                     f"  Одновимірні тести у цьому випадку НЕ інтерпретуються!", "#c62828")

        # ── Одновимірні тести (follow-up) ─────────────────────
        bonf_alpha = alpha / max(len(dv_names), 1)
        _head(f"Одновимірні тести (Bonferroni α = {fmt(bonf_alpha,4)})")
        _txt("Примітка: ці результати інтерпретуються ЛИШЕ після значущого MANOVA.",
             "#666666")
        _tbl(["Залежна змінна","F","p",f"Знач.(α={fmt(bonf_alpha,4)})",
              "partial η²","Сила ефекту","Висновок"],
             [[r[0], r[1], r[2], _sig(float(r[2]) if r[2] else float("nan")),
               r[3], r[4], r[6]] for r in univ_rows])

        # ── Групові середні ───────────────────────────────────
        _head("Групові середні (Mean ± SD)")
        means_headers = ["Група"] + [f"{nm}\nМ (SD)" for nm in dv_names]
        means_rows = []
        for lv in group_levels:
            arr = groups_data[lv]
            row_ = [lv]
            for j in range(len(dv_names)):
                m  = float(np.mean(arr[:,j]))
                sd = float(np.std(arr[:,j], ddof=1)) if len(arr) > 1 else 0.
                row_.append(f"{fmt(m,3)} ({fmt(sd,3)})")
            means_rows.append(row_)
        _tbl(means_headers, means_rows)

        # ── Графіки: стовпчикова ±SE + профільний ─────────────
        if HAS_MPL and len(dv_names) >= 2:
            n_dv = len(dv_names)
            colors_ = ["#4c72b0","#dd8452","#55a868","#c44e52","#8172b2","#937860"]

            # Графік 1: стовпчикова ±SE для кожної ЗЗ
            fig1 = Figure(figsize=(10, 6), dpi=100)
            for di, dv_nm in enumerate(dv_names):
                ax = fig1.add_subplot(1, n_dv, di+1)
                gm = [float(np.mean(groups_data[lv][:,di])) for lv in group_levels]
                gs_ = [float(np.std(groups_data[lv][:,di],ddof=1) /
                              math.sqrt(len(groups_data[lv])))
                       for lv in group_levels]
                xpos = range(len(group_levels))
                ax.bar(xpos, gm, yerr=gs_, capsize=4,
                       color=[colors_[i % len(colors_)] for i in range(len(group_levels))],
                       alpha=0.85, error_kw={"ecolor":"#333","lw":1.5})
                # позначення значущості univariate
                if univ_rows and di < len(univ_rows):
                    try:
                        p_uv = float(univ_rows[di][2])
                        if p_uv < bonf_alpha:
                            ax.set_title(f"{dv_nm}\n(p={fmt(p_uv,3)}*)", fontsize=8)
                        else:
                            ax.set_title(f"{dv_nm}\n(p={fmt(p_uv,3)})", fontsize=8)
                    except Exception:
                        ax.set_title(dv_nm, fontsize=8)
                else:
                    ax.set_title(dv_nm, fontsize=8)
                ax.set_xticks(list(xpos))
                ax.set_xticklabels(group_levels, rotation=30, ha="right", fontsize=7)
                ax.set_ylabel("Середнє ± СП" if di==0 else "", fontsize=8)
                ax.yaxis.grid(True, alpha=0.3)
                ax.spines["top"].set_visible(False)
                ax.spines["right"].set_visible(False)
            fig1.suptitle("Групові середні (±СП) по залежних змінних", fontsize=10)
            fig1.tight_layout()
            self._manova_figs[1] = fig1
            self._manova_frame1 = tk.Frame(body)
            self._manova_frame1.pack(fill=tk.X, padx=10, pady=4)
            embed_figure(fig1, self._manova_frame1)

            # Графік 2: профільний (нормовані середні)
            fig2 = Figure(figsize=(10, 6), dpi=100)
            ax2 = fig2.add_subplot(111)
            # Нормуємо кожну ЗЗ до [0,1] для порівняння профілів
            all_means = np.array([[float(np.mean(groups_data[lv][:,j]))
                                   for j in range(n_dv)] for lv in group_levels])
            mn_col = all_means.min(axis=0); mx_col = all_means.max(axis=0)
            rng = np.where(mx_col > mn_col, mx_col - mn_col, 1.)
            normed = (all_means - mn_col) / rng

            x_pos = range(n_dv)
            for gi, lv in enumerate(group_levels):
                ax2.plot(list(x_pos), normed[gi], "o-",
                         color=colors_[gi % len(colors_)],
                         label=str(lv), linewidth=2, markersize=7)
            ax2.set_xticks(list(x_pos))
            ax2.set_xticklabels(dv_names, rotation=20, ha="right", fontsize=9)
            ax2.set_ylabel("Нормоване середнє (0–1)", fontsize=9)
            ax2.set_title("Профільний графік груп (нормовані середні по ЗЗ)", fontsize=10)
            ax2.legend(title="Група", fontsize=8, title_fontsize=8)
            ax2.yaxis.grid(True, linestyle="--", alpha=0.4)
            ax2.spines["top"].set_visible(False)
            ax2.spines["right"].set_visible(False)
            fig2.tight_layout()
            self._manova_figs[2] = fig2
            self._manova_frame2 = tk.Frame(body)
            self._manova_frame2.pack(fill=tk.X, padx=10, pady=(0,10))
            embed_figure(fig2, self._manova_frame2)

    # ── Допоміжні методи для результатів MANOVA ───────────────

    def _rebuild_manova_graphs(self, dv_names, groups_data, group_levels,
                                univ_rows, alpha, p_pillai):
        """Перебудовує лише графіки у вже відкритому вікні MANOVA."""
        if not hasattr(self, '_manova_frame1'): return
        bonf_alpha = alpha / max(len(dv_names),1)
        colors_ = self._manova_gs.get("colors", self._manova_colors)
        bar_alpha = self._manova_gs.get("bar_alpha", 0.85)
        lw_ = self._manova_gs.get("lw", 2.0)
        ms_ = self._manova_gs.get("ms", 7)
        ff_ = self._manova_gs.get("font_family","Times New Roman")
        fz_ = self._manova_gs.get("font_size", 9)
        n_dv = len(dv_names)

        # Перебудовуємо графік 1
        for w in self._manova_frame1.winfo_children(): w.destroy()
        fig1 = Figure(figsize=(10, 6), dpi=100)
        for di,dv_nm in enumerate(dv_names):
            ax=fig1.add_subplot(1,n_dv,di+1)
            gm=[float(np.mean(groups_data[lv][:,di])) for lv in group_levels]
            gs_=[float(np.std(groups_data[lv][:,di],ddof=1)/math.sqrt(len(groups_data[lv])))
                 for lv in group_levels]
            xpos=range(len(group_levels))
            ax.bar(xpos,gm,yerr=gs_,capsize=4,
                   color=[colors_[i%len(colors_)] for i in range(len(group_levels))],
                   alpha=bar_alpha,error_kw={"ecolor":"#333","lw":1.5})
            try:
                p_uv=float(univ_rows[di][2]) if univ_rows and di<len(univ_rows) else float("nan")
                mark="*" if p_uv<bonf_alpha else ""
                ax.set_title(f"{dv_nm}\n(p={fmt(p_uv,3)}{mark})",fontsize=fz_,fontfamily=ff_)
            except Exception:
                ax.set_title(dv_nm,fontsize=fz_,fontfamily=ff_)
            ax.set_xticks(list(xpos))
            ax.set_xticklabels(group_levels,rotation=30,ha="right",fontsize=max(6,fz_-1))
            ax.set_ylabel("Середнє ± СП" if di==0 else "",fontsize=fz_)
            ax.yaxis.grid(True,alpha=0.3)
            ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        fig1.suptitle("Групові середні (±СП) по залежних змінних",fontsize=fz_+1)
        fig1.tight_layout()
        self._manova_figs[1]=fig1
        embed_figure(fig1, self._manova_frame1)

        # Перебудовуємо графік 2
        for w in self._manova_frame2.winfo_children(): w.destroy()
        fig2=Figure(figsize=(max(6,n_dv*0.9+2),4),dpi=100)
        ax2=fig2.add_subplot(111)
        all_means=np.array([[float(np.mean(groups_data[lv][:,j])) for j in range(n_dv)]
                             for lv in group_levels])
        mn_col=all_means.min(axis=0); mx_col=all_means.max(axis=0)
        rng=np.where(mx_col>mn_col,mx_col-mn_col,1.)
        normed=(all_means-mn_col)/rng
        for gi,lv in enumerate(group_levels):
            ax2.plot(list(range(n_dv)),normed[gi],"o-",
                     color=colors_[gi%len(colors_)],
                     label=str(lv),linewidth=lw_,markersize=ms_)
        ax2.set_xticks(list(range(n_dv)))
        ax2.set_xticklabels(dv_names,rotation=20,ha="right",fontsize=fz_,fontfamily=ff_)
        ax2.set_ylabel("Нормоване середнє (0–1)",fontsize=fz_,fontfamily=ff_)
        ax2.set_title("Профільний графік груп (нормовані середні по ЗЗ)",fontsize=fz_+1,fontfamily=ff_)
        ax2.legend(title="Група",fontsize=fz_,title_fontsize=fz_)
        ax2.yaxis.grid(True,linestyle="--",alpha=0.4)
        ax2.spines["top"].set_visible(False); ax2.spines["right"].set_visible(False)
        fig2.tight_layout()
        self._manova_figs[2]=fig2
        embed_figure(fig2, self._manova_frame2)


    def _copy_manova_text(self, win):
        """Збирає весь текст зі звіту і копіює у буфер."""
        lines = []
        def _collect(w):
            if isinstance(w, tk.Label):
                t = w.cget("text")
                if t: lines.append(t)
            for ch in w.winfo_children(): _collect(ch)
        try:
            _collect(self._manova_body)
        except Exception: pass
        text = "\n".join(lines)
        win.clipboard_clear(); win.clipboard_append(text)
        messagebox.showinfo("Скопійовано",
            "Текст звіту скопійовано у буфер обміну.\nВставте у Word через Ctrl+V.")

    def _copy_manova_fig(self, n):
        fig = self._manova_figs.get(n)
        if fig is None:
            messagebox.showwarning("","Графік не знайдено. Спочатку виконайте аналіз."); return
        ok, msg = _copy_fig_to_clipboard(fig)
        if ok: messagebox.showinfo("","Графік скопійовано.\nВставте у Word через Ctrl+V.")
        else:   messagebox.showwarning("",f"Помилка: {msg}")

    def _restyle_manova(self, win, dv_names, groups_data, group_levels,
                        univ_rows, alpha, p_pillai):
        """Діалог налаштувань кольорів/розмірів графіків MANOVA."""
        if not hasattr(self, '_manova_gs'):
            self._manova_gs = {
                "colors": ["#4c72b0","#dd8452","#55a868","#c44e52","#8172b2","#937860"],
                "bar_alpha": 0.85, "lw": 2.0, "ms": 7,
                "font_family": "Times New Roman", "font_size": 9,
            }
        dlg = tk.Toplevel(win); dlg.title("Налаштування графіків MANOVA")
        dlg.resizable(False, False); set_icon(dlg); dlg.grab_set()
        gs = self._manova_gs
        frm = tk.Frame(dlg, padx=16, pady=14); frm.pack()
        ff_v  = tk.StringVar(value=gs["font_family"])
        fz_v  = tk.IntVar(value=gs["font_size"])
        al_v  = tk.DoubleVar(value=gs["bar_alpha"])
        lw_v  = tk.DoubleVar(value=gs["lw"])
        ms_v  = tk.IntVar(value=gs["ms"])
        rb_f  = ("Times New Roman",12)
        rows_cfg = [
            ("Шрифт:",         "combo",  ff_v, ["Times New Roman","Arial","Calibri","Georgia"]),
            ("Розмір шрифту:", "spin",   fz_v, (7, 18)),
            ("Прозорість стовпців:", "scale", al_v, (0.3, 1.0)),
            ("Товщина ліній:", "scale",  lw_v, (0.5, 4.0)),
            ("Розмір точок:",  "spin",   ms_v, (3, 20)),
        ]
        for ri, (lbl, wt, var, opts) in enumerate(rows_cfg):
            tk.Label(frm, text=lbl, font=rb_f).grid(row=ri, column=0, sticky="w", pady=5)
            if wt == "combo":
                ttk.Combobox(frm, textvariable=var, values=opts,
                             state="readonly", width=20).grid(row=ri, column=1, sticky="w", padx=8)
            elif wt == "spin":
                tk.Spinbox(frm, from_=opts[0], to=opts[1], textvariable=var,
                           width=7).grid(row=ri, column=1, sticky="w", padx=8)
            elif wt == "scale":
                tk.Scale(frm, from_=opts[0], to=opts[1], resolution=0.05,
                         orient="horizontal", variable=var,
                         length=160).grid(row=ri, column=1, sticky="w", padx=8)
        # Colour pickers for groups
        col_refs = list(gs["colors"])
        tk.Label(frm, text="Кольори груп:", font=rb_f).grid(
            row=len(rows_cfg), column=0, sticky="w", pady=5)
        col_frm = tk.Frame(frm); col_frm.grid(row=len(rows_cfg), column=1, sticky="w")
        col_btns = []
        for ci, grp in enumerate(group_levels[:6]):
            c = col_refs[ci] if ci < len(col_refs) else "#999999"
            btn = tk.Button(col_frm, width=4, relief=tk.SUNKEN, bg=c,
                            text=str(ci+1), font=("Times New Roman",9))
            btn.pack(side=tk.LEFT, padx=2)
            def _pick(idx=ci, b=btn, refs=col_refs):
                ch = colorchooser.askcolor(color=refs[idx], parent=dlg)
                if ch and ch[1]: refs[idx]=ch[1]; b.configure(bg=ch[1])
            btn.configure(command=_pick)
            col_btns.append(btn)

        def apply():
            self._manova_gs.update({
                "colors": list(col_refs), "bar_alpha": al_v.get(),
                "lw": lw_v.get(), "ms": ms_v.get(),
                "font_family": ff_v.get(), "font_size": fz_v.get(),
            })
            self._manova_colors = list(col_refs)
            dlg.destroy()
            # Перебудовуємо графіки одразу
            self._rebuild_manova_graphs(dv_names, groups_data, group_levels,
                                         univ_rows, alpha, p_pillai)
        bf = tk.Frame(frm); bf.grid(row=len(rows_cfg)+1, column=0, columnspan=2, pady=(14,0))
        tk.Button(bf, text="OK", bg="#c62828", fg="white",
                  font=rb_f, command=apply).pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Скасувати", font=rb_f, command=dlg.destroy).pack(side=tk.LEFT)
        center_win(dlg)



# ═══════════════════════════════════════════════════════════════
# ДОВІДКОВА СИСТЕМА
# ═══════════════════════════════════════════════════════════════

HELP_CONTENT = {}

def _fill_help():
    H = HELP_CONTENT
    H["Швидкий старт"] = {"icon":"🚀","short":"З чого почати роботу з програмою","text":"""
ШВИДКИЙ СТАРТ

КРОК 1. Оберіть тип аналізу на головному екрані.
КРОК 2. Введіть дані:
  Ctrl+V - вставити з Excel
  Файл -> Завантажити Excel (.xlsx)
  Вручну (Enter - наступний рядок, стрілки - навігація)
КРОК 3. Натисніть "Аналіз даних".
КРОК 4. Програма автоматично перевірить:
  ✓ Нормальність розподілу (Shapiro-Wilk)
  ✓ Однорідність дисперсій (Левен)
  ✓ Методологічну коректність дій
КРОК 5. Скопіюйте звіт і графіки у Word.

Збереження: Файл -> Зберегти проект (.sadp)
"""}

    H["Введення даних"] = {"icon":"📋","short":"Як правильно вводити та копіювати дані","text":"""
ВВЕДЕННЯ ДАНИХ

СТРУКТУРА ТАБЛИЦІ ДЛЯ ANOVA:
  Ліві колонки = назви рівнів факторів (текст)
  Праві колонки = числові значення повторностей

  Варіант   | Повт.1 | Повт.2 | Повт.3
  Контроль  |  4.2   |  4.5   |  4.1
  Варіант1  |  5.8   |  6.1   |  5.9

СПОСОБИ ВВЕДЕННЯ:
  1. Вручну: клік -> введіть -> Enter (наступний рядок)
  2. Ctrl+V: вставити з Excel
  3. Файл -> Завантажити Excel (.xlsx, .xlsm)
  4. Fill-handle: наведіть на правий нижній кут -> курсор змінюється -> тягніть вниз
  5. Ctrl+C: скопіювати виділений діапазон

⚠ Назви варіантів вводьте ОДНАКОВО (регістр важливий!)
⚠ Десяткова кома автоматично замінюється на крапку
"""}

    H["Дизайни експерименту"] = {"icon":"🧪","short":"CRD, RCBD та Split-plot: коли що обирати","text":"""
ДИЗАЙНИ ПОЛЬОВИХ ЕКСПЕРИМЕНТІВ

CRD (Повна рандомізація):
  ✓ Умови однорідні (лабораторія, вегетаційні горщики)
  ✓ Мала кількість варіантів (2-5)

RCBD (Блочна рандомізація):
  ✓ Є градієнт родючості (схил, зволоження, pH)
  ✓ Польові досліди з багатьма варіантами
  Кожна числова колонка = один блок (повторність)
  ✓ Виключає міжблокову варіацію з помилки -> точніший аналіз

Split-plot (Спліт-плот):
  ✓ Два фактори з РІЗНИМ розміром ділянок
  Головний фактор (WP) = великі ділянки (обробіток ґрунту, норма посіву)
  Підфактор (SP) = малі ділянки (сорти, дози добрив)

  ⚠ УВАГА: два різних похибки!
    Whole-plot error -> для тесту WP фактора
    Sub-plot error (залишок) -> для SP та взаємодії
  ⚠ Застосування RCBD замість Split-plot -> неправильні F-значення!

ЯК ОБРАТИ:
  Немає блоків -> CRD
  Є блоки, однаковий розмір ділянок -> RCBD
  Є блоки, різний розмір ділянок -> Split-plot
"""}

    H["Нормальність розподілу"] = {"icon":"📊","short":"Shapiro-Wilk і дії при порушенні","text":"""
НОРМАЛЬНІСТЬ РОЗПОДІЛУ - SHAPIRO-WILK

Перевіряється нормальність ЗАЛИШКІВ моделі (не сирих даних!).

p > 0.05 -> нормальний -> параметричний аналіз ✓
p <= 0.05 -> ненормальний -> трансформація або непараметричний

⚠ ОБМЕЖЕННЯ:
  n < 8: тест ненадійний
  n > 100: виявляє найменші відхилення -> дивіться QQ-plot!

ДІЇ ПРИ НЕНОРМАЛЬНОСТІ:

1. ТРАНСФОРМАЦІЯ (повернення до параметричних методів):
   ln(x)  - для правоскошених даних, відносних показників
   sqrt(x) - для даних з рахунком (кількість особин, плям)
   log10(x) - для великих значень
   
   Програма: перевіряє допустимість -> виконує -> перевіряє нормальність знову
   ✓ Нормальний -> пропонує параметричний метод
   ✗ Все ще ненормальний -> зупиняється, рекомендує непараметричний

2. НЕПАРАМЕТРИЧНИЙ АНАЛІЗ:
   CRD: Kruskal-Wallis -> Mann-Whitney (post-hoc)
   RCBD: Friedman -> Wilcoxon (парний)
"""}

    H["Однорідність дисперсій"] = {"icon":"⚖️","short":"Тест Левена і наслідки порушення","text":"""
ОДНОРІДНІСТЬ ДИСПЕРСІЙ - ТЕСТ ЛЕВЕНА

Levene (center=median) = варіант Brown-Forsythe - найробустніший.

p >= 0.05 -> умова виконується ✓
p < 0.05 -> дисперсії різняться -> програма блокує і запитує підтвердження

ДІЇ ПРИ ПОРУШЕННІ:
  1. Трансформація (ln або sqrt стабілізують дисперсію)
  2. Для двох груп -> Welch t-test (у модулі t-тест, автоматично)
  3. Непараметричні методи не потребують рівності дисперсій

ANOVA достатньо робастна якщо:
  Розміри груп приблизно рівні
  Відношення max/min дисперсій < 9:1
"""}

    H["Методи порівнянь"] = {"icon":"🔬","short":"НІР, Тьюкі, Дункан, Бонферроні","text":"""
МЕТОДИ МНОЖИННИХ ПОРІВНЯНЬ (POST-HOC)

НІР05 (Fisher Protected LSD):
  ✓ Найпоширеніший в агрономії України
  ✓ Виконується ТІЛЬКИ після значущого F (Protected LSD)
  Рекомендується: <= 6 варіантів

Тест Тьюкі (Tukey HSD):
  ✓ Строгий контроль сімейної помилки (FWER)
  Рекомендується: будь-яка кількість варіантів

Тест Дункана:
  ✓ Справжня степ-даун процедура: alpha_p = 1-(1-alpha)^(p-1)
  ✓ Проміжний між LSD і Тьюкі
  Рекомендується: 5-10 варіантів

Бонферроні:
  Найконсервативніший. p_adj = p x кількість_пар
  Рекомендується: мала кількість запланованих порівнянь

CLD ЛІТЕРИ:
  Однакові літери = немає значущої різниці
  Різні літери = є значуща різниця (p < alpha)

ЗВЕДЕННЯ:
  <= 6 варіантів -> НІР05
  5-10 варіантів -> Дункан
  Будь-яка кількість -> Тьюкі
  Мало запланованих -> Бонферроні
"""}

    H["Непараметричні тести"] = {"icon":"📉","short":"KW, Friedman, MWU, Wilcoxon та розміри ефектів","text":"""
НЕПАРАМЕТРИЧНІ МЕТОДИ

Kruskal-Wallis (аналог ANOVA для CRD):
  Глобальний тест. Розмір ефекту: epsilon^2
  < 0.01 слабкий | 0.01-0.06 середній | > 0.14 сильний
  Post-hoc: Mann-Whitney з Бонферроні (автоматично)

Friedman (аналог RCBD ANOVA):
  Для блочних дизайнів. Розмір ефекту: Kendall W
  ⚠ Вимагає повні блоки (всі варіанти у кожному блоці)!
  Post-hoc: Wilcoxon з Бонферроні

Mann-Whitney U (попарне):
  Розмір ефекту: Cliff delta
  |delta| < 0.147: дуже слабкий | 0.147-0.33: слабкий
  0.33-0.474: середній | > 0.474: сильний

Wilcoxon signed-rank (парний):
  Для пар спостережень або RCBD з 2 варіантами.
  Розмір ефекту: r = Z/sqrt(n)
"""}

    H["Типи SS"] = {"icon":"∑","short":"Типи сум квадратів I-IV","text":"""
ТИПИ СУМ КВАДРАТІВ

При збалансованих даних - всі типи дають однаковий результат.
Різниця виникає при незбалансованих даних.

Тип I (Послідовний):
  Порядок факторів ВАЖЛИВИЙ!
  ⚠ Програма попередить при незбалансованих даних!
  Коли: збалансовані дизайни, регресія з осмисленим порядком

Тип II (Ієрархічний):
  Порядок НЕ важливий. Вища потужність ніж III при відсутності взаємодій.
  Коли: незбалансовані БЕЗ значущих взаємодій

Тип III (Частковий) <- ЗА ЗАМОВЧУВАННЯМ:
  Кожен ефект при всіх інших (включно зі взаємодіями).
  Порядок НЕ важливий. Стандарт SPSS/SAS. ✓
  Коли: більшість випадків, незбалансовані зі взаємодіями

Тип IV:
  Для незбалансованих з ПРОПУЩЕНИМИ КЛІТИНКАМИ.
  Рідкісний у польових дослідах.

РЕКОМЕНДАЦІЯ:
  Збалансований -> будь-який тип
  Незбалансований + взаємодії -> Тип III
  Незбалансований + без взаємодій -> Тип II
  Пропущені клітинки -> Тип IV
"""}

    H["Кореляційний аналіз"] = {"icon":"🔗","short":"Пірсон vs Спірмен, поправки, теплова карта","text":"""
КОРЕЛЯЦІЙНИЙ АНАЛІЗ

СТРУКТУРА ДАНИХ (два варіанти):
  А) Перший рядок = назва показника, нижче - числові значення
  Б) Перша колонка = назва показника, праворуч - значення

МЕТОДИ:
  Авто (рекомендовано): Shapiro-Wilk по кожному показнику
    хоч один ненормальний -> Спірмен
    всі нормальні -> Пірсон

  Пірсон r: лінійний зв'язок, нормальний розподіл
    ⚠ При ненормальних даних -> програма попередить!

  Спірмен rho: монотонний зв'язок, будь-який розподіл

ПОПРАВКИ НА МНОЖИННІ ПОРІВНЯННЯ:
  Бонферроні: строга (FWER), p_adj = p x кількість_пар
  BH (Benjamini-Hochberg FDR): ліберальніша, більш потужна
  -> При > 10 показниках рекомендується BH

ТЕПЛОВА КАРТА - кожна клітинка:
  r = коефіцієнт | p = значущість після поправки | n = кількість пар
  * p < 0.05 | ** p < 0.01

ІНТЕРПРЕТАЦІЯ |r|:
  0.0-0.2: дуже слабкий | 0.2-0.4: слабкий | 0.4-0.6: помірний
  0.6-0.8: сильний | 0.8-1.0: дуже сильний
"""}

    H["Регресійний аналіз"] = {"icon":"📈","short":"7 моделей регресії, R², діагностика","text":"""
РЕГРЕСІЙНИЙ АНАЛІЗ

МОДЕЛІ:
  1. Лінійна:       y = a + b*x
  2. Квадратична:   y = a + b*x + c*x^2  (оптимум добрив)
  3. Кубічна:       y = a + b*x + c*x^2 + d*x^3
  4. Степенева:     y = a * x^b  (x > 0)
  5. Експоненційна: y = a * e^(b*x)
  6. Логарифмічна:  y = a + b*ln(x)  (x > 0)
  7. Логістична 4P: y = d + (a-d)/(1+(x/c)^b)  (S-крива)

ПОКАЗНИКИ ЯКОСТІ:
  R^2   = частка варіації пояснена моделлю (0-1; > 0.9 - добре)
  R^2adj = R^2 з урахуванням кількості параметрів
  RMSE  = середньоквадратична помилка (в одиницях y)
  F-тест = значущість моделі загалом

ДІАГНОСТИКА:
  Residuals vs Fitted: будь-який патерн -> модель неповна!
  Shapiro-Wilk залишків: p > 0.05 -> нормальні ✓
  Grubbs test: автоматичне виявлення викидів у залишках

ВВЕДЕННЯ ДАНИХ:
  Ліве поле: значення x | Праве поле: значення y
  "Вставити дані": два стовпці з Excel (x | y)
"""}

    H["ANCOVA"] = {"icon":"🎛️","short":"Коваріаційний аналіз: контроль зовнішніх факторів","text":"""
ANCOVA - КОВАРІАЦІЙНИЙ АНАЛІЗ

Мета: порівняти групи усуваючи вплив коваріати.

КОВАРІАТА - неперервна змінна яку вимірюєте але не контролюєте:
  Початкова висота рослин, pH грунту, вміст гумусу, опади

СТРУКТУРА:
  [Група] | [Коваріата 1] | [Коваріата 2] | [Y - залежна змінна]

ЗАХИСТ (8 перевірок):
  1. n >= 6 спостережень
  2. >= 2 груп
  3. >= 2 спостережень у кожній групі
  4. Мультиколінеарність коваріат (r > 0.95)
  5. Паралельність ліній регресії <- КЛЮЧОВА!
  6. Нормальність залишків
  7. Однорідність дисперсій
  8. Скориговані середні (LS Means)

⚠ КРИТИЧНА ПЕРЕДУМОВА - ПАРАЛЕЛЬНІСТЬ ЛІНІЙ:
  Тест взаємодії Група x Коваріата:
  p < alpha -> лінії НЕ паралельні -> ANCOVA ЗАБЛОКОВАНА!

РЕЗУЛЬТАТИ:
  Скориговані (LS) Means = середні при однаковій коваріаті у всіх групах
  Порівнюйте скориговані середні, а не нескориговані!
"""}

    H["MANOVA"] = {"icon":"🔢","short":"Багатовимірний аналіз кількох показників","text":"""
MANOVA - БАГАТОВИМІРНИЙ ДИСПЕРСІЙНИЙ АНАЛІЗ

Мета: порівняти групи за КІЛЬКОМА залежними змінними одночасно.

НАВІЩО НЕ КІЛЬКА ANOVA?
  4 ANOVA з alpha=0.05 -> 19% шанс хибної різниці!
  MANOVA контролює сімейну помилку.

⚠ КРИТИЧНА ВИМОГА: n > p У КОЖНІЙ ГРУПІ
  n = кількість спостережень; p = кількість DV
  n <= p в будь-якій групі -> ЖОРСТКЕ БЛОКУВАННЯ!

ЗАХИСТ (7 перевірок):
  1. n > p у кожній групі (блокування)
  2. >= 2 DV
  3. Мультиколінеарність DV (|r| > 0.90)
  4. Багатовимірна нормальність (Mardia)
  5. Однорідність коваріаційних матриць (Box M)
  6. Чотири тестові статистики
  7. Univariate тільки після значущого MANOVA

СТАТИСТИКИ:
  Wilks Lambda: стандартна
  Pillai Trace: найробустніша ★ (при порушеннях)
  Hotelling-Lawley: один домінантний ефект
  Roy GCR: найпотужніша, найменш робастна

ПОСЛІДОВНІСТЬ:
  Перевірте передумови -> Pillai Trace ->
  Якщо значущо -> univariate ANOVAs з Bonferroni (alpha/p)
  Якщо незначущо -> univariate НЕ інтерпретуються!
"""}

    H["Аналіз стабільності"] = {"icon":"🌍","short":"GGE biplot та Eberhart-Russell для GxE","text":"""
АНАЛІЗ СТАБІЛЬНОСТІ - GxE ВЗАЄМОДІЯ

GxE: відносна продуктивність сортів змінюється між середовищами.

СТРУКТУРА:
  Рядки = Генотипи (сорти, лінії)
  Стовпці = Середовища (роки, місця)
  Значення = середня врожайність

EBERHART-RUSSELL:
  bi (коефіцієнт стабільності):
    bi ~= 1.0 -> стабільний (середня реакція)
    bi > 1.2  -> адаптивний (реагує на покращення умов)
    bi < 0.8  -> консервативний (для бідних умов)

  s^2d (дисперсія відхилень):
    ~= 0 -> передбачуваний
    > 0  -> непередбачуваний

  Ідеал: висока середня + bi ~= 1 + s^2d ~= 0

GGE BIPLOT:
  Точки = генотипи
  Стрілки = середовища
  Близько до центру = стабільний генотип
  Близько до стрілки середовища = добре виконує там
"""}

    H["Розмір вибірки"] = {"icon":"🔢","short":"Скільки повторностей потрібно?","text":"""
КАЛЬКУЛЯТОР РОЗМІРУ ВИБІРКИ

Параметри:
  alpha: рівень значущості (зазвичай 0.05)
  Потужність (1-beta): зазвичай 0.80 або 0.90
  delta: мінімальна різниця яку хочете виявити (наприклад 0.5 т/га)
  sigma: стандартне відхилення (з попередніх дослідів)
  k: кількість варіантів

РЕЖИМИ:
  Порожнє поле "повторності" -> мінімальна кількість повторностей
  Введіть повторності -> досягнута потужність тесту

ТИПОВІ ЗНАЧЕННЯ (зернові):
  CV 10-15%, виявити 10% різниці -> зазвичай 3-4 повторності
  CV 20-25% -> 5-7 повторностей

⚠ Планувати з потужністю >= 0.80!
"""}

    H["Проект"] = {"icon":"💾","short":"Збереження та відкриття проектів (.sadp)","text":"""
УПРАВЛІННЯ ПРОЕКТАМИ

Формат: .sadp (JSON-текст, можна відкрити у блокноті)

ЩО ЗБЕРІГАЄТЬСЯ:
  ✓ Кількість та назви факторів
  ✓ Всі введені дані (назви варіантів і числа)
  ✓ Кількість стовпців (повторностей)

ДЛЯ КІЛЬКОХ ПОКАЗНИКІВ ОДНОГО ДОСЛІДУ:
  1. Введіть схему досліду та перший показник
  2. Збережіть: Файл -> Зберегти проект
  3. Аналіз -> збережіть звіт
  4. Видаліть числові дані, залиште назви варіантів
  5. Введіть наступний показник -> аналіз -> звіт
  Або збережіть окремий .sadp для кожного показника.

МЕНЮ ФАЙЛ:
  Зберегти проект (Ctrl+S)
  Відкрити проект (Ctrl+O)
  Очистити таблицю
  Завантажити Excel
"""}

    H["Графіки і налаштування"] = {"icon":"🎨","short":"Boxplot, Venn, теплова карта та їх налаштування","text":"""
ГРАФІЧНИЙ ЗВІТ

BOXPLOT (коробка з вусами):
  Верхній вус: Q3 + 1.5*IQR (або максимум)
  Верхній край коробки: Q3 (75-й перцентиль)
  Лінія: медіана (Q2)
  Нижній край коробки: Q1 (25-й перцентиль)
  Нижній вус: Q1 - 1.5*IQR (або мінімум)
  Кружечки поза вусами: викиди (outliers)

  Літери CLD над коробками:
    Однакові -> немає значущої різниці
    Різні -> є значуща різниця (p < alpha)

ДІАГРАМА ВЕННА (сила впливу):
  Кола = головні ефекти факторів
  Перетини кіл = взаємодії між факторами
  Сума всіх частин = 100%

НАЛАШТУВАННЯ (кнопка ⚙ у вікні графіків):
  Boxplot: шрифт, розмір, кольори коробки/медіани/вусів/викидів
  Venn: шрифт, прозорість, кольори кіл і тексту
  За замовчуванням: APA стиль

КОПІЮВАННЯ:
  "Копіювати PNG" -> 300 dpi -> Ctrl+V у Word
  (Windows only; macOS/Linux - toolbar matplotlib)

ТЕПЛОВА КАРТА (кореляція):
  В клітинці: r / p / n
  Налаштовується через ⚙: colormap, шрифти, кольори
"""}

    H["Позначення у звіті"] = {"icon":"📝","short":"Як читати таблиці і показники у звіті","text":"""
ПОЗНАЧЕННЯ У ЗВІТІ

ЗНАЧУЩІСТЬ:
  **  p < 0.01 (висока значущість)
  *   p < 0.05 (значущо)
  -   p >= 0.05 (не значущо)

ТАБЛИЦЯ ANOVA:
  SS  = сума квадратів
  df  = ступені свободи
  MS  = середній квадрат (SS/df)
  F   = критерій Фішера (MS_ефект / MS_залишок)
  p   = ймовірність (менше = значущіше)

СИЛА ВПЛИВУ (% від SS):
  Частка загальної варіації по кожному джерелу.
  Залишок включений. Сума всіх рядків = 100%.

PARTIAL eta^2 (розмір ефекту):
  eta^2 = SS_ефект / (SS_ефект + SS_залишок)
  < 0.01: дуже слабкий | 0.01-0.06: слабкий
  0.06-0.14: середній | > 0.14: сильний

R^2 (коефіцієнт детермінації):
  Частка варіації пояснена всією моделлю.
  R^2 = 0.90 -> 90% варіації пояснено ✓

CV% (коефіцієнт варіації):
  < 10%: відмінна точність | 10-15%: хороша
  15-20%: задовільна | > 20%: низька

НІР05:
  Якщо різниця > НІР05 -> статистично значуща.
"""}

_fill_help()

# ── Розширення довідки детальними поясненнями ───────────────
def _extend_help():
    H = HELP_CONTENT

    H["Описова статистика"] = {"icon":"📐","short":"Базові показники: середнє, SD, медіана, довірчий інтервал","text":"""
ЩО ТАКЕ ОПИСОВА СТАТИСТИКА?

Перш ніж виконувати будь-який статистичний аналіз, корисно описати ваші дані
за допомогою базових числових характеристик. Це допомагає:
  - Зрозуміти загальний характер даних (велике чи маленьке розкидання?)
  - Виявити можливі помилки у введенні (занадто великі або малі значення)
  - Оцінити чи підходять параметричні методи

ЯК ВВОДИТИ ДАНІ:
  Кожен стовпець = один показник (змінна).
  Перша клітинка стовпця = назва показника (текст).
  Решта клітинок = числові значення.

  Приклад:
  | Врожайність | Висота рослин | Маса 1000 зерен |
  |    4.2      |     95.3      |      38.2       |
  |    5.1      |    102.1      |      41.5       |
  |    4.8      |     98.7      |      39.8       |

ЩО ОЗНАЧАЄ КОЖЕН ПОКАЗНИК:

n (кількість спостережень):
  Кількість числових значень у вибірці. Чим більше n, тим надійніший аналіз.
  Мінімально рекомендовано n >= 5 для більшості тестів.

Середнє (Mean):
  Середньоарифметичне значення. Сума всіх значень поділена на їх кількість.
  Чутливе до викидів: одне дуже велике або мале значення може суттєво змінити середнє.

SD (стандартне відхилення):
  Показує наскільки в середньому значення відхиляються від середнього.
  Велике SD = великий розкид = висока варіабельність.
  Приблизно 68% значень знаходяться в межах Mean ± 1*SD (для нормального розподілу).

СП (стандартна похибка середнього, SE):
  SE = SD / sqrt(n)
  Показує точність оцінки середнього. Чим більше n, тим менше SE.
  Використовується для побудови довірчих інтервалів.

Мін / Макс:
  Найменше та найбільше значення у вибірці.
  Дуже корисні для виявлення помилок введення даних!

Медіана:
  Значення що ділить впорядкований ряд навпіл: 50% значень нижче, 50% вище.
  На відміну від середнього, стійка до викидів.
  Якщо середнє >> медіани -> дані правоскошені.
  Якщо середнє << медіани -> дані лівоскошені.

Q1 (перший квартиль, 25-й перцентиль):
  25% значень нижче цього рівня.

Q3 (третій квартиль, 75-й перцентиль):
  75% значень нижче цього рівня.
  IQR = Q3 - Q1 = міжквартильний розмах (стійка міра варіабельності).

CV% (коефіцієнт варіації):
  CV = SD / Mean x 100%
  Відносна мінливість у відсотках. Дозволяє порівнювати варіабельність
  показників з різними одиницями виміру.
  Для польових дослідів: CV < 10% = відмінна точність.

Асиметрія (Skewness):
  = 0: симетричний розподіл
  > 0: правостороння асиметрія (хвіст праворуч, більшість значень ліворуч)
  < 0: лівостороння асиметрія
  |асиметрія| > 1 = суттєве відхилення від нормального розподілу.

Ексцес (Kurtosis):
  Показує «гостроту» розподілу порівняно з нормальним.
  = 0: нормальний розподіл
  > 0: гостроверхий (більше значень поблизу середнього та більше викидів)
  < 0: пласковерхий

95% Довірчий інтервал (ДІ):
  Діапазон в якому знаходиться «справжнє» середнє генеральної сукупності
  з ймовірністю 95%.
  Якщо 95% ДІ двох груп не перетинаються -> є підстави вважати що середні різні.
  Але для строгого висновку потрібен статистичний тест!

SW p (Shapiro-Wilk p-значення):
  Тест нормальності розподілу.
  p > 0.05: дані відповідають нормальному розподілу (параметричні тести можна)
  p <= 0.05: дані не відповідають (розгляньте непараметричні тести)
"""}

    H["t-тест"] = {"icon":"🔀","short":"Порівняння двох груп: t-тест, Велш, Манн-Уітні, Вілкоксон","text":"""
ЩО ТАКЕ t-ТЕСТ?

t-тест — це статистичний тест для порівняння ДВОХ груп або вибірок.
Відповідає на питання: «Чи є різниця між середніми статистично значущою,
чи це просто випадкові коливання?»

ПРОГРАМА АВТОМАТИЧНО ОБИРАЄ ПРАВИЛЬНИЙ ТЕСТ:
  1. Перевіряє нормальність розподілу (Shapiro-Wilk) для кожної групи
  2. Перевіряє рівність дисперсій (тест Левена) для незалежних вибірок
  3. Обирає відповідний тест:
     Нормальний + рівні дисперсії -> t-тест Стьюдента
     Нормальний + нерівні дисперсії -> t-тест Велша (Welch)
     Ненормальний -> Mann-Whitney U або Wilcoxon

ТРИ РЕЖИМИ:

1. НЕЗАЛЕЖНІ ВИБІРКИ (2 різні групи):
   Порівнюємо дві незалежні групи.
   Приклад: врожайність сорту А проти сорту Б.
   ⚠ Спостереження у групах НЕ пов'язані між собою.

   Введення: Група 1 і Група 2 - числа через Enter або кому.

   Якщо дані нормальні:
     Рівні дисперсії (Левен p >= 0.05) -> t-тест Стьюдента (класичний)
     Нерівні дисперсії (Левен p < 0.05) -> t-тест Велша (Welch) - більш точний!
   Якщо ненормальні -> Mann-Whitney U (непараметричний аналог)

2. ПАРНІ ВИБІРКИ (до/після, або парні вимірювання):
   Ті самі об'єкти вимірюються двічі.
   Приклад: маса рослин до і після обробки; врожайність на тих самих ділянках у 2 роки.
   ⚠ Вимагає ОДНАКОВУ кількість спостережень у обох групах!
   ⚠ Порядок важливий: перше значення Групи 1 пов'язане з першим Групи 2.

   Якщо дані нормальні -> Парний t-тест
   Якщо ненормальні -> Wilcoxon signed-rank

3. ОДНА ВИБІРКА (проти відомого значення):
   Перевіряємо чи середнє вибірки відрізняється від заданого значення.
   Приклад: чи відрізняється врожайність від нормативного показника 5 т/га?

   Введіть відоме середнє μ₀ і значення вибірки.

ЯК ЧИТАТИ РЕЗУЛЬТАТИ:
  t (або U, W) - значення тестової статистики
  p - ймовірність отримати такий або більший результат якщо H₀ вірна
  p < 0.05: різниця значуща ✓
  p >= 0.05: різниця незначуща (але це не означає що груп немає!)

РОЗМІР ЕФЕКТУ для Mann-Whitney (Cliff's delta):
  Показує не лише ЧИ є різниця, але й НАСКІЛЬКИ вона велика.
  |delta| < 0.147: дуже слабкий (практично немає різниці)
  0.147-0.33: слабкий
  0.33-0.474: середній
  > 0.474: сильний (суттєва практична різниця)

ПОРАДА:
  Значущий p ≠ велика різниця! При великих n навіть мізерна різниця буде значущою.
  Завжди оцінюйте розмір ефекту разом з p-значенням.
"""}

    H["Описова статистика — боксплот"] = {"icon":"📦","short":"Як читати боксплот (діаграму коробку з вусами)","text":"""
БОКСПЛОТ (ДІАГРАМА КОРОБКА З ВУСАМИ)

Боксплот — це графічний спосіб відобразити розподіл даних
не залежно від кількості спостережень.

ЯК ЧИТАТИ:

        ╷  <- верхній вус: Q3 + 1.5*IQR
        │     (або максимальне значення якщо воно менше)
    ┌───┐
    │   │  <- верхній край коробки: Q3 (75-й перцентиль)
    │═══│  <- жирна лінія: МЕДІАНА (Q2, 50-й перцентиль)
    │   │  <- нижній край коробки: Q1 (25-й перцентиль)
    └───┘
        │
        ╵  <- нижній вус: Q1 - 1.5*IQR
           (або мінімальне значення якщо воно більше)

    ○   <- окремі точки: ВИКИДИ (outliers)
           значення далі ніж 1.5*IQR від коробки

ЩО ТАКЕ IQR?
  IQR = Q3 - Q1 = міжквартильний розмах.
  Вміщує 50% «середніх» значень.
  Чим більша коробка, тим більший розкид «типових» значень.

ЯК ПОРІВНЮВАТИ БОКСПЛОТИ:
  Коробки не перекриваються -> можлива значуща різниця
  Медіани дуже різні -> явна різниця між групами
  Коробка однієї групи всередині іншої -> групи схожі

ЛІТЕРИ CLD НАД БОКСПЛОТАМИ:
  Це результат пост-хок аналізу після ANOVA.
  Однакові літери -> немає значущої різниці між цими варіантами
  Різні літери -> є значуща різниця (p < alpha)
  Приклад: ab і a не різняться, ab і b не різняться, але a і b можуть різнятися!

ВИКИДИ (outliers):
  Значення далі 1.5*IQR від коробки.
  Можуть бути:
  - Помилками вимірювання (перевірте журнал польового досліду!)
  - Справжніми екстремальними значеннями (посуха, хвороба тощо)
  - Важливою біологічною інформацією
  Не видаляйте викиди без перевірки!
"""}

    H["Кластерний аналіз — детально"] = {"icon":"🌿","short":"Ієрархічна кластеризація та дендрограма","text":"""
КЛАСТЕРНИЙ АНАЛІЗ — ДЕТАЛЬНЕ ПОЯСНЕННЯ

Кластерний аналіз групує об'єкти (сорти, проби, ділянки) так,
щоб схожі між собою опинились в одному кластері.

КОЛИ ЗАСТОСОВУВАТИ:
  ✓ Класифікація сортів за комплексом ознак
  ✓ Групування ґрунтових проб за хімічним складом
  ✓ Виявлення природних груп у даних без попередньої класифікації

ВВЕДЕННЯ ДАНИХ:
  Перший стовпець: назва об'єкта (сорт, зразок тощо)
  Решта стовпців: числові ознаки (показники)
  Перший рядок: назви показників

  Приклад:
  | Сорт    | Висота | Врожайн. | Стійкість |
  | Поліська|  95.3  |   5.8    |    7.2    |
  | Київська| 102.1  |   6.4    |    8.1    |
  | Одеська |  88.5  |   5.1    |    6.8    |

  Програма автоматично СТАНДАРТИЗУЄ дані (z-оцінки) щоб показники
  з різними одиницями виміру мали однаковий вплив.

МЕТОДИ ЗЧЕПЛЕННЯ (linkage):
  ward:      Мінімізує внутрішньокластерну дисперсію.
             Найпопулярніший метод, зазвичай дає найкращі результати. ✓
  complete:  Дистанція між найдальшими об'єктами кластерів.
             Дає компактні кластери однакового розміру.
  average:   Середня дистанція між всіма парами.
             Компроміс між ward і complete.
  single:    Дистанція між найближчими об'єктами.
             Схильний до «ефекту ланцюга» (довгі ланцюжкові кластери).

ЯК ЧИТАТИ ДЕНДРОГРАМУ:
  Вертикальна вісь = відстань (несхожість).
  Гілки зливаються на рівні що відповідає відстані між кластерами.
  Чим вища точка злиття, тим менш схожі ці кластери.

  Щоб отримати k кластерів: проведіть горизонтальну лінію на відповідній висоті
  так щоб перетнути k вертикальних гілок.

ВИБІР КІЛЬКОСТІ КЛАСТЕРІВ k:
  k задається вручну. Як обрати?
  - Дивіться на дендрограму: де є великий «стрибок» у висоті злиття -> там природна межа
  - k = кількість природних груп у вашому досліді (сорти різного типу, ґрунтові зони тощо)
  - Типово для агрономічних досліджень: k = 2-5

КОЛІРНЕ КОДУВАННЯ НА ДЕНДРОГРАМІ:
  Різні кольори = різні кластери.
  Горизонтальна пунктирна лінія показує поріг відсікання для k кластерів.
"""}

    H["PCA — детально"] = {"icon":"🔮","short":"Як читати biplot та scree plot аналізу головних компонент","text":"""
АНАЛІЗ ГОЛОВНИХ КОМПОНЕНТ (PCA) — ДЕТАЛЬНО

Уявіть що у вас є 10 показників для кожного сорту.
Це важко уявити і проаналізувати. PCA стискає цю інформацію
до 2-3 нових «узагальнених» показників (головних компонент)
які описують більшу частину різноманітності у ваших даних.

КОЛИ ЗАСТОСОВУВАТИ:
  ✓ Багато показників (> 5-7) для кожного об'єкта
  ✓ Хочете виявити природне групування об'єктів
  ✓ Хочете зрозуміти які показники «ходять разом»
  ✓ Як крок перед MANOVA якщо n <= p

ВВЕДЕННЯ ДАНИХ:
  Перша колонка: мітка об'єкта (необов'язково, назва сорту тощо)
  Решта колонок: числові показники
  Перший рядок: назви показників

SCREE PLOT (ліворуч):
  Стовпчасти: % дисперсії пояснений кожною компонентою.
  Лінія (червона): кумулятивний % пояснення.
  Горизонтальна пунктирна: 80% поріг.

  Скільки компонент залишити? «Правило ліктя»:
  Знайдіть точку де графік різко змінює нахил (стає плоским) ->
  це і є оптимальна кількість компонент.
  Зазвичай: PC1+PC2 разом пояснюють 70-85% -> достатньо для аналізу.

BIPLOT (посередині):
  Поєднує розташування об'єктів (точки) і змінних (стрілки) на одному графіку.

  Точки (об'єкти/сорти):
    Близькі точки = схожі об'єкти за всім комплексом показників.
    Далекі точки = дуже різні об'єкти.
    Кластери точок = природні групи.

  Стрілки (змінні/показники):
    Довга стрілка = показник добре описується цими компонентами.
    Напрямок стрілки = в якому напрямку зростає цей показник.
    Стрілки в одному напрямку = показники корелюють (ростуть разом).
    Стрілки протилежних напрямків = показники обернено корельовані.
    Стрілки під прямим кутом = показники не пов'язані.

    Об'єкт близький до стрілки = він має відносно більше значення цього показника.

ТЕПЛОВА КАРТА НАВАНТАЖЕНЬ (праворуч):
  Показує як кожен вихідний показник «вкладається» в кожну головну компоненту.
  Великі значення (темно-зелені або темно-червоні) = сильний зв'язок.
  PC1 зазвичай = «загальний розмір» або «загальна продуктивність».
  PC2 = наступна за важливістю вісь незалежна від PC1.
"""}

    H["Повторні виміри ANOVA"] = {"icon":"⏱️","short":"Аналіз динамічних вимірювань одних і тих самих об'єктів","text":"""
ДИСПЕРСІЙНИЙ АНАЛІЗ ПОВТОРНИХ ВИМІРЮВАНЬ

КОЛИ ЗАСТОСОВУВАТИ:
  Одні й ті самі суб'єкти (рослини, тварини, ділянки) вимірюються КІЛЬКА РАЗІВ:
  - У різні моменти часу (висота рослин через кожні 2 тижні)
  - За різних умов (доза добрив A, потім B, потім C)
  - До і після (більш ніж 2 точки -> потрібен повторний ANOVA; якщо 2 точки -> парний t-тест)

  ✓ Динаміка росту рослин
  ✓ Зміна вмісту поживних речовин по фазах вегетації
  ✓ Відповідь на послідовні обробки

  ⚠ ВІДМІНА від звичайного ANOVA: тут спостереження НЕ незалежні
  (одна рослина вимірюється кілька разів -> між вимірами є зв'язок).

ВВЕДЕННЯ ДАНИХ:
  Рядки = суб'єкти (рослини, ділянки тощо)
  Стовпці = часові точки або умови (T1, T2, T3 ...)
  Перший рядок заголовків (синій) = назви часових точок

  Приклад (висота рослин, см):
  | Суб'єкт | Тиждень1 | Тиждень2 | Тиждень3 | Тиждень4 |
  | Ділянка1|   15.2   |   28.4   |   45.1   |   58.3   |
  | Ділянка2|   14.8   |   26.9   |   43.7   |   56.8   |

ЩО ПОКАЗУЮТЬ РЕЗУЛЬТАТИ:

SS (суми квадратів):
  SS_time = варіація пояснена часом (основний ефект)
  SS_subj = варіація між суб'єктами (усувається з помилки -> підвищує точність!)
  SS_error = залишкова варіація

F-тест для «time»:
  p < 0.05: є значуща динаміка (показник змінюється через час)
  p >= 0.05: немає значущої динаміки

Partial η² (розмір ефекту часу):
  Показує яку частку варіації пояснює фактор «час».
  > 0.14: сильний ефект (виразна динаміка).

ГРАФІК MEANS ± SE:
  Показує як середнє значення змінюється у часі.
  SE (смужки похибок) показують точність оцінки середнього.
  Чим менші смужки, тим точніше середнє визначено.

ПОСТ-ХОК (Бонферроні):
  Після значущого F виконуються попарні порівняння часових точок.
  Показує ЯКІ САМЕ пари часових точок відрізняються.
  p_adj = скориговане p (Бонферроні = множення на кількість пар).

НОРМАЛЬНІСТЬ РІЗНИЦЬ:
  Перевіряється нормальність різниць між парами часових точок.
  p > 0.05: нормальний -> результати надійні.
  p <= 0.05: розгляньте непараметричний аналог (тест Фрідмана).
"""}

    H["Аналіз головних компонент"] = H["PCA — детально"]
    H["Кластерний аналіз — пояснення"] = H["Кластерний аналіз — детально"]
    H["t-тест / Манн-Уітні"] = H["t-тест"]

_extend_help()


class HelpWindow:
    def __init__(self, parent, start_topic=None):
        self.win = tk.Toplevel(parent)
        self.win.title("S.A.D. — Довідка")
        self.win.geometry("1000x660"); set_icon(self.win)
        self.current_topic = None
        self._build()
        topic = start_topic or list(HELP_CONTENT.keys())[0]
        self._show_topic(topic)

    def _build(self):
        # Left panel
        left = tk.Frame(self.win, width=220, bg="#f5f5f5", relief=tk.RIDGE, bd=1)
        left.pack(side=tk.LEFT, fill=tk.Y); left.pack_propagate(False)
        tk.Label(left, text="Зміст довідки",
                 font=("Times New Roman",12,"bold"), bg="#1a4b8c", fg="white",
                 pady=8).pack(fill=tk.X)
        # Search
        sf = tk.Frame(left, bg="#f5f5f5"); sf.pack(fill=tk.X, padx=6, pady=6)
        tk.Label(sf, text="Пошук:", bg="#f5f5f5", font=("Times New Roman",11)).pack(side=tk.LEFT)
        self._sv = tk.StringVar(); self._sv.trace_add("write", self._on_search)
        tk.Entry(sf, textvariable=self._sv, font=("Times New Roman",11),
                 relief=tk.FLAT, bg="white").pack(side=tk.LEFT, fill=tk.X, expand=True, padx=4)
        # Topic buttons frame (scrollable)
        self._tf = tk.Frame(left, bg="#f5f5f5"); self._tf.pack(fill=tk.BOTH, expand=True)
        self._btn = {}; self._build_list(list(HELP_CONTENT.keys()))
        # Right panel
        right = tk.Frame(self.win, bg="white"); right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self._title = tk.Label(right, text="", font=("Times New Roman",14,"bold"),
                               bg="#1a4b8c", fg="white", pady=8, padx=10, anchor="w")
        self._title.pack(fill=tk.X)
        tf2 = tk.Frame(right); tf2.pack(fill=tk.BOTH, expand=True, padx=4, pady=4)
        ysb = ttk.Scrollbar(tf2); ysb.pack(side=tk.RIGHT, fill=tk.Y)
        self._txt = tk.Text(tf2, wrap="word", font=("Times New Roman",12),
                            state="disabled", relief=tk.FLAT, bg="white",
                            yscrollcommand=ysb.set, padx=10, pady=8, cursor="arrow")
        self._txt.pack(fill=tk.BOTH, expand=True)
        ysb.config(command=self._txt.yview)
        self._txt.bind("<MouseWheel>",
                       lambda e: self._txt.yview_scroll(int(-1*(e.delta/120)),"units"))
        # Tags
        self._txt.tag_configure("bold", font=("Times New Roman",12,"bold"))
        self._txt.tag_configure("check", foreground="#1a6b1a", font=("Times New Roman",12))
        self._txt.tag_configure("warn",  foreground="#c62828", font=("Times New Roman",12))
        self._txt.tag_configure("normal",font=("Times New Roman",12))
        # Bottom
        bot = tk.Frame(right, bg="#f0f0f0", pady=4); bot.pack(fill=tk.X)
        tk.Button(bot, text="<- Попередня", command=self._prev,
                  font=("Times New Roman",11)).pack(side=tk.LEFT, padx=8)
        tk.Button(bot, text="Наступна ->", command=self._next,
                  font=("Times New Roman",11)).pack(side=tk.LEFT, padx=4)
        tk.Button(bot, text="Закрити", command=self.win.destroy,
                  font=("Times New Roman",11)).pack(side=tk.RIGHT, padx=8)

    def _build_list(self, topics):
        for w in self._tf.winfo_children(): w.destroy()
        self._btn = {}
        for topic in topics:
            info = HELP_CONTENT.get(topic, {})
            frm = tk.Frame(self._tf, bg="#f5f5f5", cursor="hand2")
            frm.pack(fill=tk.X, padx=3, pady=1)
            icon = info.get("icon","•")
            lbl = tk.Label(frm, text=f"{icon}  {topic}",
                           font=("Times New Roman",11,"bold"), bg="#f5f5f5",
                           anchor="w", padx=6, pady=2)
            lbl.pack(fill=tk.X)
            sub = tk.Label(frm, text=f"    {info.get('short','')}",
                           font=("Times New Roman",9), fg="#666", bg="#f5f5f5", anchor="w", padx=6)
            sub.pack(fill=tk.X)
            for w in [frm, lbl, sub]:
                w.bind("<Button-1>", lambda e, t=topic: self._show_topic(t))
                w.bind("<Enter>",  lambda e, f=frm: [c.configure(bg="#dce8ff") for c in [f]+list(f.winfo_children())])
                w.bind("<Leave>",  lambda e, f=frm, t=topic: self._set_bg(f, t))
            self._btn[topic] = frm

    def _set_bg(self, frm, topic):
        bg = "#c8d8ff" if topic == self.current_topic else "#f5f5f5"
        frm.configure(bg=bg)
        for w in frm.winfo_children(): w.configure(bg=bg)

    def _on_search(self, *_):
        q = self._sv.get().strip().lower()
        if not q: self._build_list(list(HELP_CONTENT.keys())); return
        matched = [t for t, info in HELP_CONTENT.items()
                   if q in t.lower() or q in info.get("short","").lower()
                   or q in info.get("text","").lower()]
        self._build_list(matched)

    def _show_topic(self, topic):
        if topic not in HELP_CONTENT: return
        self.current_topic = topic
        info = HELP_CONTENT[topic]
        self._title.configure(text=f"{info.get('icon','')}  {topic}")
        for t, frm in self._btn.items(): self._set_bg(frm, t)
        self._txt.configure(state="normal")
        self._txt.delete("1.0", tk.END)
        for line in info.get("text","").strip().split("\n"):
            s = line.strip()
            if s.startswith("✓") or s.startswith("✔"):
                tag = "check"
            elif s.startswith("⚠") or s.startswith("✗") or "БЛОК" in s or "КРИТ" in s:
                tag = "warn"
            elif line.isupper() and len(line) > 3:
                tag = "bold"
            else:
                tag = "normal"
            self._txt.insert(tk.END, line + "\n", tag)
        self._txt.configure(state="disabled")
        self._txt.yview_moveto(0)

    def _prev(self):
        keys = list(HELP_CONTENT.keys())
        if self.current_topic in keys:
            idx = keys.index(self.current_topic)
            if idx > 0: self._show_topic(keys[idx-1])

    def _next(self):
        keys = list(HELP_CONTENT.keys())
        if self.current_topic in keys:
            idx = keys.index(self.current_topic)
            if idx < len(keys)-1: self._show_topic(keys[idx+1])


def show_help(parent, topic=None):
    HelpWindow(parent, start_topic=topic)

# ═══════════════════════════════════════════════════════════════
# UPDATE MENU — add ANCOVA and MANOVA buttons
# ═══════════════════════════════════════════════════════════════

_SADTk_orig_init = SADTk.__init__


# ═══════════════════════════════════════════════════════════════
# TRIAL DESIGN GENERATOR
# ═══════════════════════════════════════════════════════════════
class TrialDesignWindow:
    """Генератор плану польового досліду — універсальний для всіх культур."""

    # ── Типи культур з налаштуваннями ────────────────────────
    CULTURES = {
        "Зернові / польові культури": {
            "plot_w": 3.0, "plot_l": 10.0, "unit": "ділянка",
            "garden": False,
            "indicators": ["Висота рослин, см", "Маса 1000 зерен, г",
                           "Врожайність, т/га", "Вміст білку, %"],
        },
        "Садівництво (дерева)": {
            "plot_w": 4.0, "plot_l": 5.0, "unit": "дерево",
            "garden": True,
            "row_sp": 4.0, "plant_sp": 5.0,
            "plants_plot": 5, "guard_ends": 1, "guard_rows": 1,
            "indicators": ["Висота дерева, м", "Діаметр крони, м",
                           "Врожайність з дерева, кг", "Маса плоду, г",
                           "Вміст ЦРР, °Brix", "% зав'язування квіток"],
        },
        "Ягідники": {
            "plot_w": 2.0, "plot_l": 0.5, "unit": "кущ",
            "garden": True,
            "row_sp": 2.0, "plant_sp": 0.5,
            "plants_plot": 8, "guard_ends": 2, "guard_rows": 1,
            "indicators": ["Висота куща, см", "Кількість пагонів",
                           "Врожайність з куща, кг", "Маса ягоди, г",
                           "Вміст ЦРР, °Brix"],
        },
        "Овочівництво (відкритий ґрунт)": {
            "plot_w": 2.0, "plot_l": 5.0, "unit": "ділянка",
            "garden": False,
            "indicators": ["Висота рослин, см", "Маса плоду, г",
                           "Врожайність, т/га", "Товарність, %",
                           "Вихід стандарту, %"],
        },
        "Захищений ґрунт (теплиця)": {
            "plot_w": 1.0, "plot_l": 4.0, "unit": "грядка",
            "garden": False,
            "indicators": ["Висота рослин, см", "Кількість плодів/рослину",
                           "Врожайність, кг/м²", "Маса плоду, г"],
        },
        "Виноградарство": {
            "plot_w": 3.0, "plot_l": 1.5, "unit": "кущ",
            "garden": True,
            "row_sp": 3.0, "plant_sp": 1.5,
            "plants_plot": 6, "guard_ends": 2, "guard_rows": 1,
            "indicators": ["Маса грона, г", "Кількість грон/кущ",
                           "Врожайність з куща, кг", "Вміст ЦРР, °Brix",
                           "Кислотність, г/л"],
        },
    }

    DESIGNS = [
        ("crd",   "CRD — Повністю рандомізований",
         "Всі ділянки рівноцінні. Варіанти розміщуються випадково."),
        ("rcbd",  "RCBD — Рандомізовані повні блоки (рекомендується)",
         "Поле ділиться на повторності. Кожна повторність = всі варіанти."),
        ("latin", "Латинський квадрат",
         "Контролює 2 джерела мінливості. k варіантів = k рядів = k стовпців."),
        ("split", "Split-plot — Розщеплені ділянки",
         "Два фактори різного масштабу. WP = великі ділянки, SP = підділянки."),
    ]

    HELP_TEXT = """
ГЕНЕРАТОР ПЛАНУ ПОЛЬОВОГО ДОСЛІДУ
══════════════════════════════════════════════════

ЩО РОБИТЬ?
  Автоматично рандомізує розміщення варіантів
  і формує документи для польової роботи:
  • Польова схема (кольорова карта ділянок)
  • Список рандомізації (порядок закладки)
  • Польовий журнал (таблиця для вимірювань)

══════════════════════════════════════════════════
КРОК 1. ТИП КУЛЬТУРИ
══════════════════════════════════════════════════
  Оберіть тип культури — програма підлаштує:
  • Типові розміри ділянок
  • Стандартні показники для журналу
  • Термінологію (ділянка/дерево/кущ)

══════════════════════════════════════════════════
КРОК 2. ВАРІАНТИ ДОСЛІДУ
══════════════════════════════════════════════════
  Введіть назви варіантів — по одному на рядок.
  Приклади:
    Контроль (без добрив)
    N60P60K60
    N90P60K60
    N120P60K60

  Для Split-plot — введіть також sub-plot варіанти.

══════════════════════════════════════════════════
КРОК 3. ДИЗАЙН ДОСЛІДУ
══════════════════════════════════════════════════

  CRD — Повністю рандомізований:
    Для однорідних умов. Всі ділянки рівноцінні.
    Варіанти розміщуються абсолютно випадково.
    Простий але потребує однорідного фону.

  RCBD — Рандомізовані повні блоки:
    РЕКОМЕНДУЄТЬСЯ для більшості дослідів.
    Поле ділиться на ПОВТОРНОСТІ (блоки).
    Кожна повторність містить всі варіанти.
    Блоки розміщують перпендикулярно до основного
    градієнта мінливості (схил, зрошення, ряди).
    У садівництві: повторність = кілька дерев/ряд.

  Латинський квадрат:
    Контролює ДВА незалежних джерела мінливості.
    k варіантів → k рядів × k стовпців ділянок.
    Кожен варіант — рівно 1 раз у кожному ряду
    і рівно 1 раз у кожному стовпці.
    Рекомендується при k=4-6 варіантів.
    Після досліду → 3-факторна ANOVA у S.A.D.
    (Фактор A = Варіант, B = Рядок, C = Стовпець)

  Split-plot — Розщеплені ділянки:
    Для двох факторів різного масштабу.
    WP (whole-plot): головний фактор — великі ділянки.
    SP (sub-plot): другорядний — всередині WP.
    Приклади:
      Зернові:    WP = обробка ґрунту, SP = сорт
      Садівництво: WP = підщепа, SP = сорт
      Овочі:      WP = спосіб вирощування, SP = сорт

══════════════════════════════════════════════════
КРОК 4. SEED РАНДОМІЗАЦІЇ
══════════════════════════════════════════════════
  Seed — технічний номер жеребкування.
  При однаковому seed → однакова схема.
  ЗБЕРІГАЙТЕ seed у документації досліду!
  Seed ≠ Рік: seed може бути будь-яким числом.

══════════════════════════════════════════════════
КРОК 5. ПОЛЬОВИЙ ЖУРНАЛ
══════════════════════════════════════════════════
  Налаштуйте назви показників (через ";").
  Програма підставляє типові показники для
  обраного типу культури.
  Натисніть "▶ Оновити" після зміни назв.
  Збережіть журнал у Excel для польової роботи.

══════════════════════════════════════════════════
САДІВНИЦТВО — ОСОБЛИВОСТІ
══════════════════════════════════════════════════
  Ділянка = кілька дерев одного варіанту в ряду.
  Повторність = окремий ряд або частина ряду.

  Приклад для яблуні (схема 4×5 м):
    4 сорти × 3 повторності = 12 ділянок
    По 3-5 дерев на ділянку
    Розмір ділянки: 4 м (ширина) × 15 м (5 дерев)

  У RCBD: кожна повторність = 1 ряд дерев,
  де всі 4 сорти розміщені випадково в ряду.

  Показники для кожного дерева усереднюють
  і записують одне середнє значення на ділянку.
"""

    def __init__(self, parent):
        self.win = tk.Toplevel(parent)
        self.win.title("Генератор плану польового досліду")
        self.win.geometry("1160x760")
        self.win.resizable(True, True)
        set_icon(self.win)
        self._plan_data = None
        self._build()

    # ═══════════════════════════════════════════════════════
    # _build — головний інтерфейс
    # ═══════════════════════════════════════════════════════
    def _build(self):
        rf = ("Times New Roman", 11)

        # ── Toolbar ────────────────────────────────────────
        top = tk.Frame(self.win, padx=8, pady=5); top.pack(fill=tk.X)
        tk.Button(top, text="▶ Згенерувати план", bg="#c62828", fg="white",
                  font=("Times New Roman", 13),
                  command=self._generate).pack(side=tk.LEFT, padx=4)
        tk.Button(top, text="💾 Зберегти схему PNG",
                  font=rf, command=self._save_png).pack(side=tk.LEFT, padx=4)
        tk.Button(top, text="💾 Зберегти журнал Excel",
                  font=rf, command=self._save_excel).pack(side=tk.LEFT, padx=4)
        tk.Button(top, text="💾 Зберегти рандомізацію TXT",
                  font=rf, command=self._save_rand_txt).pack(side=tk.LEFT, padx=4)
        tk.Button(top, text="📚 Довідка", bg="#1a4b8c", fg="white",
                  font=rf, command=self._show_help).pack(side=tk.LEFT, padx=4)

        # ── Основна область ────────────────────────────────
        main = tk.Frame(self.win); main.pack(fill=tk.BOTH, expand=True, padx=8, pady=4)

        # ── ЛІВА ПАНЕЛЬ (прокручувана) ─────────────────────
        left_outer = tk.Frame(main, width=360)
        left_outer.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 8))
        left_outer.pack_propagate(False)
        lc = tk.Canvas(left_outer, highlightthickness=0)
        lc.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        lsb = ttk.Scrollbar(left_outer, orient="vertical", command=lc.yview)
        lsb.pack(side=tk.RIGHT, fill=tk.Y)
        lc.configure(yscrollcommand=lsb.set)
        lf = tk.Frame(lc); lc.create_window((0, 0), window=lf, anchor="nw")
        lf.bind("<Configure>", lambda e: lc.configure(scrollregion=lc.bbox("all")))
        left_outer.bind("<MouseWheel>",
                        lambda e: lc.yview_scroll(int(-1*(e.delta/120)), "units"))

        # ─── Тип культури ──────────────────────────────────
        cf = tk.LabelFrame(lf, text="1. Тип культури",
                           font=("Times New Roman", 11, "bold"), padx=8, pady=4)
        cf.pack(fill=tk.X, pady=(0, 4))
        self.culture_var = tk.StringVar(value=list(self.CULTURES.keys())[0])
        self.culture_cb = ttk.Combobox(cf, textvariable=self.culture_var,
                                       values=list(self.CULTURES.keys()),
                                       state="readonly", width=36, font=rf)
        self.culture_cb.pack(fill=tk.X, pady=2)
        self.culture_cb.bind("<<ComboboxSelected>>", self._on_culture)
        self._culture_hint = tk.Label(cf, text="", font=("Times New Roman", 9),
                                      fg="#555", justify="left")
        self._culture_hint.pack(anchor="w")

        # ─── Варіанти ──────────────────────────────────────
        vf = tk.LabelFrame(lf, text="2. Варіанти досліду (один на рядок)",
                           font=("Times New Roman", 11, "bold"), padx=8, pady=4)
        vf.pack(fill=tk.X, pady=(0, 4))
        self.var_text = tk.Text(vf, width=38, height=6, font=rf, wrap="word")
        self.var_text.pack(fill=tk.X, pady=2)
        self.var_text.insert("1.0", "Контроль\nВаріант 1\nВаріант 2\nВаріант 3")

        # ─── Дизайн ────────────────────────────────────────
        df = tk.LabelFrame(lf, text="3. Дизайн досліду",
                           font=("Times New Roman", 11, "bold"), padx=8, pady=4)
        df.pack(fill=tk.X, pady=(0, 4))
        self.design_var = tk.StringVar(value="rcbd")
        for val, label, desc in self.DESIGNS:
            fr = tk.Frame(df); fr.pack(fill=tk.X, pady=1)
            tk.Radiobutton(fr, text=label, variable=self.design_var, value=val,
                           font=("Times New Roman", 11),
                           command=self._on_design).pack(side=tk.LEFT)
        self._design_hint = tk.Label(df, text="", font=("Times New Roman", 9),
                                     fg="#1a4b8c", bg="#eef4ff",
                                     justify="left", wraplength=320, padx=4, pady=3)
        self._design_hint.pack(fill=tk.X, pady=4)

        # Split-plot додатковий фактор
        self.sp_frame = tk.LabelFrame(lf, text="Sub-plot варіанти",
                                      font=("Times New Roman", 11, "bold"),
                                      padx=8, pady=4)
        self.sp_text = tk.Text(self.sp_frame, width=38, height=3, font=rf)
        self.sp_text.pack(fill=tk.X)
        self.sp_text.insert("1.0", "Сорт А\nСорт Б\nСорт В")
        # (показується лише для split-plot)

        # ─── Параметри ─────────────────────────────────────
        pf = tk.LabelFrame(lf, text="4. Параметри",
                           font=("Times New Roman", 11, "bold"), padx=8, pady=4)
        pf.pack(fill=tk.X, pady=(0, 4))
        self._pv = {}
        for ri, (lbl, key, default, hint) in enumerate([
            ("Повторностей:", "reps", "3",
             "Кількість повторностей (рядів)"),
            ("Seed рандомізації:", "seed", "2024",
             "Число для відтворення жеребкування. ≠ рік!"),
        ]):
            tk.Label(pf, text=lbl, font=rf).grid(row=ri, column=0, sticky="w", pady=2)
            v = tk.StringVar(value=default); self._pv[key] = v
            tk.Entry(pf, textvariable=v, width=9, font=rf
                     ).grid(row=ri, column=1, sticky="w", padx=6)
            tk.Label(pf, text=hint, font=("Times New Roman", 9), fg="#666"
                     ).grid(row=ri, column=2, sticky="w")

        # ─── Польові параметри (для зернових/овочів) ────────
        self._field_frame = tk.LabelFrame(lf, text="Розміри ділянки",
                           font=("Times New Roman", 11, "bold"), padx=8, pady=4)
        self._field_frame.pack(fill=tk.X, pady=(0, 4))
        for ri, (lbl, key, default, hint) in enumerate([
            ("Ширина, м:", "pw", "5", "Ширина ділянки"),
            ("Довжина, м:", "pl", "10", "Довжина ділянки"),
        ]):
            tk.Label(self._field_frame, text=lbl, font=rf
                     ).grid(row=ri, column=0, sticky="w", pady=2)
            v = tk.StringVar(value=default); self._pv[key] = v
            tk.Entry(self._field_frame, textvariable=v, width=9, font=rf
                     ).grid(row=ri, column=1, sticky="w", padx=6)
            tk.Label(self._field_frame, text=hint,
                     font=("Times New Roman", 9), fg="#666"
                     ).grid(row=ri, column=2, sticky="w")

        # ─── Садівничі параметри ─────────────────────────────
        self._garden_frame = tk.LabelFrame(lf, text="Параметри садіння",
                            font=("Times New Roman", 11, "bold"), padx=8, pady=4)
        self._gv = {}
        garden_params = [
            ("Схема садіння — між рядами, м:", "row_sp",     "4.0",
             "Відстань між рядами"),
            ("Схема садіння — в ряду, м:",     "plant_sp",   "5.0",
             "Відстань між рослинами в ряду"),
            ("Облікових рослин на ділянку:",   "plants_plot","5",
             "Без захисних — лише облікові"),
            ("Захисних рослин (поч. і кін.):", "guard_ends", "1",
             "Кількість з кожного боку ряду"),
            ("Захисних рядів між варіантами:", "guard_rows", "1",
             "Рядів-буферів між повторностями"),
        ]
        for ri, (lbl, key, default, hint) in enumerate(garden_params):
            tk.Label(self._garden_frame, text=lbl, font=rf
                     ).grid(row=ri, column=0, sticky="w", pady=2)
            v = tk.StringVar(value=default); self._gv[key] = v
            tk.Entry(self._garden_frame, textvariable=v, width=9, font=rf
                     ).grid(row=ri, column=1, sticky="w", padx=6)
            tk.Label(self._garden_frame, text=hint,
                     font=("Times New Roman", 9), fg="#666"
                     ).grid(row=ri, column=2, sticky="w")

        # ─── Паспорт ───────────────────────────────────────
        nf = tk.LabelFrame(lf, text="5. Паспорт досліду",
                           font=("Times New Roman", 11, "bold"), padx=8, pady=4)
        nf.pack(fill=tk.X, pady=(0, 4))
        self._nv = {}
        for ri, (lbl, key) in enumerate([
            ("Назва:", "name"), ("Рік:", "year"),
            ("Місце:", "loc"),  ("Відповідальний:", "resp"),
        ]):
            tk.Label(nf, text=lbl, font=rf).grid(row=ri, column=0, sticky="w", pady=2)
            v = tk.StringVar(value=(str(datetime.now().year) if key=="year" else ""))
            self._nv[key] = v
            tk.Entry(nf, textvariable=v, width=28, font=rf
                     ).grid(row=ri, column=1, sticky="w", padx=6, pady=2)

        # ── ПРАВА ПАНЕЛЬ — вкладки результатів ─────────────
        right = tk.Frame(main); right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.nb = ttk.Notebook(right); self.nb.pack(fill=tk.BOTH, expand=True)

        # Вкладка 1: Схема
        t1 = tk.Frame(self.nb); self.nb.add(t1, text="🗺 Польова схема")
        tk.Label(t1, text="Клітинка = одна ділянка. Кольори = варіанти. Двигайте схему прокруткою.",
                 font=("Times New Roman", 9), fg="#666").pack(anchor="w", padx=4, pady=2)
        self._scheme_cv = tk.Canvas(t1, bg="white")
        s_vsb = ttk.Scrollbar(t1, orient="vertical",
                               command=self._scheme_cv.yview)
        s_hsb = ttk.Scrollbar(t1, orient="horizontal",
                               command=self._scheme_cv.xview)
        s_vsb.pack(side=tk.RIGHT, fill=tk.Y)
        s_hsb.pack(side=tk.BOTTOM, fill=tk.X)
        self._scheme_cv.pack(fill=tk.BOTH, expand=True)
        self._scheme_cv.configure(yscrollcommand=s_vsb.set,
                                   xscrollcommand=s_hsb.set)

        # Вкладка 2: Рандомізація
        t2 = tk.Frame(self.nb); self.nb.add(t2, text="📋 Рандомізація")
        tb2 = tk.Frame(t2); tb2.pack(fill=tk.X, padx=4, pady=3)
        tk.Label(tb2, text="Порядок закладки ділянок:",
                 font=rf).pack(side=tk.LEFT)
        r_vsb = ttk.Scrollbar(t2, orient="vertical")
        r_vsb.pack(side=tk.RIGHT, fill=tk.Y)
        self.rand_txt = tk.Text(t2, font=("Courier New", 10),
                                yscrollcommand=r_vsb.set,
                                state="disabled", wrap="none")
        self.rand_txt.pack(fill=tk.BOTH, expand=True)
        r_vsb.config(command=self.rand_txt.yview)

        # Вкладка 3: Журнал
        t3 = tk.Frame(self.nb); self.nb.add(t3, text="📓 Польовий журнал")
        tb3 = tk.Frame(t3); tb3.pack(fill=tk.X, padx=4, pady=3)
        tk.Label(tb3, text="Показники:", font=rf).pack(side=tk.LEFT, padx=(0,4))
        self.ind_var = tk.StringVar()
        tk.Entry(tb3, textvariable=self.ind_var, width=55,
                 font=("Times New Roman", 10)).pack(side=tk.LEFT, padx=2)
        tk.Button(tb3, text="▶ Оновити", bg="#c62828", fg="white",
                  font=("Times New Roman", 10),
                  command=self._refresh_journal).pack(side=tk.LEFT, padx=4)
        tk.Label(tb3, text="(через крапку з комою)",
                 font=("Times New Roman", 9), fg="#888").pack(side=tk.LEFT)

        j_frame = tk.Frame(t3); j_frame.pack(fill=tk.BOTH, expand=True)
        j_vsb = ttk.Scrollbar(j_frame, orient="vertical")
        j_hsb = ttk.Scrollbar(j_frame, orient="horizontal")
        j_vsb.pack(side=tk.RIGHT, fill=tk.Y)
        j_hsb.pack(side=tk.BOTTOM, fill=tk.X)
        self.journal_tv = ttk.Treeview(j_frame, yscrollcommand=j_vsb.set,
                                       xscrollcommand=j_hsb.set)
        self.journal_tv.pack(fill=tk.BOTH, expand=True)
        j_vsb.config(command=self.journal_tv.yview)
        j_hsb.config(command=self.journal_tv.xview)

        # Ініціалізуємо підказки
        self._on_culture()
        self._on_design()

    # ═══════════════════════════════════════════════════════
    # Обробники змін
    # ═══════════════════════════════════════════════════════
    def _on_culture(self, *_):
        key = self.culture_var.get()
        cfg = self.CULTURES.get(key, {})
        unit = cfg.get("unit", "ділянка")
        is_garden = cfg.get("garden", False)

        if is_garden:
            row_sp    = cfg.get("row_sp",    4.0)
            plant_sp  = cfg.get("plant_sp",  5.0)
            n_plot    = cfg.get("plants_plot",5)
            g_ends    = cfg.get("guard_ends", 1)
            g_rows    = cfg.get("guard_rows", 1)
            self._culture_hint.configure(
                text=f"Одиниця: {unit}  |  Схема: {row_sp}×{plant_sp} м")
            if hasattr(self, '_gv'):
                self._gv["row_sp"].set(str(row_sp))
                self._gv["plant_sp"].set(str(plant_sp))
                self._gv["plants_plot"].set(str(n_plot))
                self._gv["guard_ends"].set(str(g_ends))
                self._gv["guard_rows"].set(str(g_rows))
            # Показати садівничий фрейм, сховати польовий
            if hasattr(self, '_garden_frame'):
                self._field_frame.pack_forget()
                self._garden_frame.pack(fill=tk.X, pady=(0,4),
                                        after=self._field_frame)
        else:
            pw = cfg.get("plot_w", 5); pl = cfg.get("plot_l", 10)
            self._culture_hint.configure(
                text=f"Одиниця: {unit}  |  Типові розміри: {pw}×{pl} м")
            if hasattr(self, '_pv'):
                self._pv["pw"].set(str(pw))
                self._pv["pl"].set(str(pl))
            # Показати польовий фрейм, сховати садівничий
            if hasattr(self, '_garden_frame'):
                self._garden_frame.pack_forget()
                self._field_frame.pack(fill=tk.X, pady=(0,4))

        indicators = cfg.get("indicators", [])
        if hasattr(self, 'ind_var'):
            self.ind_var.set("; ".join(indicators))

    def _on_design(self, *_):
        val = self.design_var.get()
        hints = {
            "crd":   "Однорідний фон. Варіанти розміщуються абсолютно випадково по всьому полю.",
            "rcbd":  "Рекомендується. Поле ділиться на повторності. Кожна повторність = всі варіанти.",
            "latin": "k варіантів = k рядів = k стовпців. Аналіз — 3-факторна ANOVA у S.A.D.",
            "split": "Введіть WP (whole-plot) варіанти вище і SP варіанти нижче.",
        }
        if hasattr(self, '_design_hint'):
            self._design_hint.configure(text=hints.get(val, ""))
        if hasattr(self, 'sp_frame'):
            if val == "split":
                self.sp_frame.pack(fill=tk.X, pady=(0, 4),
                                   after=self.sp_frame.master.winfo_children()[-1]
                                   if self.sp_frame.master.winfo_children() else None)
            else:
                self.sp_frame.pack_forget()

    # ═══════════════════════════════════════════════════════
    # Довідка
    # ═══════════════════════════════════════════════════════
    def _show_help(self):
        win = tk.Toplevel(self.win)
        win.title("Довідка — Генератор плану досліду")
        win.geometry("700x660"); set_icon(win)
        frm = tk.Frame(win); frm.pack(fill=tk.BOTH, expand=True, padx=8, pady=6)
        vsb = ttk.Scrollbar(frm, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        txt = tk.Text(frm, wrap="word", font=("Times New Roman", 11),
                      yscrollcommand=vsb.set, relief=tk.FLAT,
                      bg="#fafafa", padx=10, pady=8, cursor="arrow")
        txt.pack(fill=tk.BOTH, expand=True); vsb.config(command=txt.yview)
        txt.insert("1.0", self.HELP_TEXT.strip()); txt.configure(state="disabled")
        txt.bind("<MouseWheel>",
                 lambda e: txt.yview_scroll(int(-1*(e.delta/120)), "units"))
        tk.Button(win, text="Закрити", command=win.destroy,
                  font=("Times New Roman", 11)).pack(pady=6)

    # ═══════════════════════════════════════════════════════
    # Генерація плану
    # ═══════════════════════════════════════════════════════
    def _generate(self):
        import random

        variants = [v.strip() for v in
                    self.var_text.get("1.0", "end").splitlines() if v.strip()]
        if len(variants) < 2:
            messagebox.showwarning("", "Введіть щонайменше 2 варіанти."); return

        try:
            reps = int(self._pv["reps"].get())
            seed = int(self._pv["seed"].get())
        except ValueError:
            messagebox.showwarning("", "Перевірте числові поля."); return

        cfg = self.CULTURES.get(self.culture_var.get(), {})
        is_garden = cfg.get("garden", False)

        # ── Параметри залежно від режиму ──────────────────────
        if is_garden:
            try:
                row_sp    = float(self._gv["row_sp"].get())
                plant_sp  = float(self._gv["plant_sp"].get())
                n_plot    = int(self._gv["plants_plot"].get())
                g_ends    = int(self._gv["guard_ends"].get())
                g_rows    = int(self._gv["guard_rows"].get())
            except ValueError:
                messagebox.showwarning("", "Перевірте параметри садіння."); return
            pw = row_sp; pl = plant_sp  # для сумісності
        else:
            try:
                pw = float(self._pv["pw"].get())
                pl = float(self._pv["pl"].get())
            except ValueError:
                messagebox.showwarning("", "Перевірте розміри ділянки."); return
            row_sp = pw; plant_sp = pl
            n_plot = 1; g_ends = 0; g_rows = 0

        design = self.design_var.get()
        rng = random.Random(seed)
        k = len(variants)

        if design == "latin" and k > 8:
            messagebox.showwarning("Латинський квадрат",
                f"Максимум 8 варіантів. У вас {k}."); return

        plan = []

        if design == "crd":
            all_p = variants * reps; rng.shuffle(all_p)
            for i, v in enumerate(all_p):
                plan.append({"plot": i+1, "rep": "–",
                             "variant": v, "row": i//k+1, "col": i%k+1})

        elif design == "rcbd":
            pn = 0
            for b in range(1, reps+1):
                bv = variants[:]; rng.shuffle(bv)
                for i, v in enumerate(bv):
                    pn += 1
                    plan.append({"plot": pn, "rep": f"Повт. {b}",
                                "variant": v, "row": b, "col": i+1})

        elif design == "latin":
            reps = k
            base = list(range(k))
            rows_p = [base[:]]
            for _ in range(k-1):
                rows_p.append(rows_p[-1][1:] + [rows_p[-1][0]])
            rng.shuffle(rows_p)
            cp = list(range(k)); rng.shuffle(cp)
            pn = 0
            for r in range(k):
                for c in range(k):
                    pn += 1
                    plan.append({"plot": pn, "rep": f"Рядок {r+1}",
                                "variant": variants[rows_p[r][cp[c]]],
                                "row": r+1, "col": c+1,
                                "col_label": f"Стовп. {c+1}"})

        elif design == "split":
            sp_vars = [v.strip() for v in
                       self.sp_text.get("1.0", "end").splitlines() if v.strip()]
            if len(sp_vars) < 2:
                messagebox.showwarning("", "Введіть щонайменше 2 sub-plot варіанти."); return
            pn = 0
            for b in range(1, reps+1):
                wp_o = variants[:]; rng.shuffle(wp_o)
                for wp in wp_o:
                    sp_o = sp_vars[:]; rng.shuffle(sp_o)
                    for sp in sp_o:
                        pn += 1
                        plan.append({"plot": pn, "rep": f"Повт. {b}",
                                    "variant": f"{wp} / {sp}",
                                    "wp": wp, "sp": sp, "row": b,
                                    "col": wp_o.index(wp)*len(sp_vars)+sp_o.index(sp)+1})

        design_name = {v: l for v, l, _ in self.DESIGNS}.get(design, design)

        # Розраховуємо площу
        if is_garden:
            plot_area = n_plot * row_sp * plant_sp
            total_area = plot_area * len(plan)
            area_msg = (f"Облікових рослин/ділянку: {n_plot}\n"
                       f"Схема садіння: {row_sp}×{plant_sp} м\n"
                       f"Площа ділянки: {plot_area:.1f} м²\n"
                       f"Загальна площа (облікова): {total_area:.0f} м²")
        else:
            plot_area = pw * pl
            total_area = plot_area * len(plan)
            area_msg = f"Площа ділянки: {pw}×{pl} м = {plot_area:.0f} м²\nЗагальна площа: {total_area:.0f} м²"

        self._plan_data = {
            "plan": plan, "variants": variants, "reps": reps,
            "design": design, "design_name": design_name,
            "seed": seed, "k": k, "pw": pw, "pl": pl,
            "is_garden": is_garden,
            "row_sp": row_sp if is_garden else pw,
            "plant_sp": plant_sp if is_garden else pl,
            "n_plot": n_plot, "g_ends": g_ends, "g_rows": g_rows,
            "culture": self.culture_var.get(),
            "unit": cfg.get("unit", "ділянка"),
            "name": self._nv["name"].get(),
            "year": self._nv["year"].get(),
            "loc":  self._nv["loc"].get(),
            "resp": self._nv["resp"].get(),
        }

        self._draw_scheme()
        self._fill_rand()
        self._fill_journal()
        self.nb.select(0)

        messagebox.showinfo("Готово",
            f"План згенеровано!\n\n"
            f"Дизайн: {design_name}\n"
            f"Варіантів: {k}   |   Повторностей: {reps}\n"
            f"Ділянок: {len(plan)}\n"
            f"{area_msg}\n\n"
            f"Seed рандомізації: {seed}\n"
            f"⚠ Збережіть seed у документацію!")
        import random

        variants = [v.strip() for v in
                    self.var_text.get("1.0", "end").splitlines() if v.strip()]
        if len(variants) < 2:
            messagebox.showwarning("", "Введіть щонайменше 2 варіанти."); return

        try:
            reps = int(self._pv["reps"].get())
            seed = int(self._pv["seed"].get())
            pw   = float(self._pv["pw"].get())
            pl   = float(self._pv["pl"].get())
        except ValueError:
            messagebox.showwarning("", "Перевірте числові поля (повторності, seed, розміри)."); return

        design = self.design_var.get()
        rng = random.Random(seed)
        k = len(variants)

        if design == "latin" and k > 8:
            messagebox.showwarning("Латинський квадрат",
                f"Максимум 8 варіантів для ЛК. У вас {k}.\n"
                "Оберіть RCBD для більшої кількості варіантів."); return

        plan = []

        if design == "crd":
            all_p = variants * reps; rng.shuffle(all_p)
            for i, v in enumerate(all_p):
                plan.append({"plot": i+1, "rep": "–",
                             "variant": v, "row": i//k+1, "col": i%k+1})

        elif design == "rcbd":
            pn = 0
            for b in range(1, reps+1):
                bv = variants[:]; rng.shuffle(bv)
                for i, v in enumerate(bv):
                    pn += 1
                    plan.append({"plot": pn, "rep": f"Повт. {b}",
                                "variant": v, "row": b, "col": i+1})

        elif design == "latin":
            reps = k
            base = list(range(k))
            rows_p = [base[:]]
            for _ in range(k-1):
                rows_p.append(rows_p[-1][1:] + [rows_p[-1][0]])
            rng.shuffle(rows_p)
            cp = list(range(k)); rng.shuffle(cp)
            pn = 0
            for r in range(k):
                for c in range(k):
                    pn += 1
                    plan.append({
                        "plot": pn, "rep": f"Рядок {r+1}",
                        "variant": variants[rows_p[r][cp[c]]],
                        "row": r+1, "col": c+1,
                        "col_label": f"Стовп. {c+1}"
                    })

        elif design == "split":
            sp_vars = [v.strip() for v in
                       self.sp_text.get("1.0", "end").splitlines() if v.strip()]
            if len(sp_vars) < 2:
                messagebox.showwarning("", "Введіть щонайменше 2 sub-plot варіанти."); return
            pn = 0
            for b in range(1, reps+1):
                wp_o = variants[:]; rng.shuffle(wp_o)
                for wp in wp_o:
                    sp_o = sp_vars[:]; rng.shuffle(sp_o)
                    for sp in sp_o:
                        pn += 1
                        plan.append({
                            "plot": pn, "rep": f"Повт. {b}",
                            "variant": f"{wp} / {sp}",
                            "wp": wp, "sp": sp,
                            "row": b,
                            "col": (wp_o.index(wp))*len(sp_vars) + sp_o.index(sp) + 1
                        })

        # Зберігаємо план
        design_name = {v: l for v, l, _ in self.DESIGNS}.get(design, design)
        self._plan_data = {
            "plan": plan, "variants": variants, "reps": reps,
            "design": design, "design_name": design_name,
            "seed": seed, "k": k, "pw": pw, "pl": pl,
            "culture": self.culture_var.get(),
            "name":  self._nv["name"].get(),
            "year":  self._nv["year"].get(),
            "loc":   self._nv["loc"].get(),
            "resp":  self._nv["resp"].get(),
        }

        self._draw_scheme()
        self._fill_rand()
        self._fill_journal()
        self.nb.select(0)

        messagebox.showinfo("Готово",
            f"План згенеровано!\n\n"
            f"Дизайн: {design_name}\n"
            f"Варіантів: {k}   |   Повторностей: {reps}\n"
            f"Ділянок: {len(plan)}\n"
            f"Загальна площа: {pw*pl*len(plan):.0f} м²\n\n"
            f"Seed рандомізації: {seed}\n"
            f"⚠ Збережіть seed у документацію досліду!")

    # ═══════════════════════════════════════════════════════
    # Польова схема
    # ═══════════════════════════════════════════════════════
    def _draw_scheme(self):
        if not self._plan_data: return
        cv = self._scheme_cv; cv.delete("all")
        d = self._plan_data; plan = d["plan"]
        is_garden = d.get("is_garden", False)

        if is_garden:
            self._draw_scheme_garden(cv, d, plan)
        else:
            self._draw_scheme_field(cv, d, plan)

    def _draw_scheme_field(self, cv, d, plan):
        """Польова схема — кожна клітинка = ділянка."""
        PALETTES = ["#aed6f1","#a9dfbf","#f9e79f","#f1948a","#d2b4de",
                    "#a3e4d7","#fad7a0","#d5d8dc","#82e0aa","#f0b27a",
                    "#85c1e9","#f7dc6f","#c39bd3","#76d7c4","#f8c471"]
        all_v = list(dict.fromkeys(p["variant"] for p in plan))
        cmap  = {v: PALETTES[i % len(PALETTES)] for i, v in enumerate(all_v)}
        cols_set = sorted(set(p["col"] for p in plan))
        rows_set = sorted(set(p["row"] for p in plan))
        nc = len(cols_set); nr = len(rows_set)
        cw = max(78, min(130, 680 // nc))
        ch = max(44, min(72,  400 // nr))
        pad = 5; x0 = 110; y0 = 55

        title = d.get("name") or "План досліду"
        sub   = f"{d.get('culture','')}  |  {d['design_name']}  |  Seed={d['seed']}  |  {d.get('year','')}"
        cv.create_text(x0+nc*cw//2, 16, text=title,
                       font=("Times New Roman",12,"bold"), fill="#000")
        cv.create_text(x0+nc*cw//2, 34, text=sub,
                       font=("Times New Roman",9), fill="#555")

        rep_map = {}
        for p in plan: rep_map[p["row"]] = p["rep"]
        for i, r in enumerate(rows_set):
            cv.create_text(x0-6, y0+i*ch+ch//2,
                           text=rep_map.get(r, f"Ряд {r}"),
                           anchor="e", font=("Times New Roman",9,"bold"), fill="#333")
        for j, c in enumerate(cols_set):
            lbl = next((p for p in plan if p["col"]==c), {}).get("col_label", f"#{c}")
            cv.create_text(x0+j*cw+cw//2, y0-14,
                           text=lbl, font=("Times New Roman",8), fill="#555")
        for p in plan:
            ci = cols_set.index(p["col"]); ri = rows_set.index(p["row"])
            x1,y1 = x0+ci*cw, y0+ri*ch; x2,y2 = x1+cw-pad, y1+ch-pad
            cv.create_rectangle(x1,y1,x2,y2,
                                fill=cmap.get(p["variant"],"#eee"),
                                outline="#888", width=1)
            cv.create_text(x1+5,y1+6, text=f"№{p['plot']}",
                           anchor="nw", font=("Courier New",7), fill="#555")
            short = p["variant"][:14]+"…" if len(p["variant"])>14 else p["variant"]
            cv.create_text((x1+x2)//2,(y1+y2)//2, text=short,
                           font=("Times New Roman",8), fill="#000", width=cw-10)

        leg_y = y0+nr*ch+16
        cv.create_text(x0, leg_y, text="Легенда:",
                       anchor="w", font=("Times New Roman",10,"bold"))
        cpr = 3
        for i,v in enumerate(all_v):
            lx = x0+(i%cpr)*240; ly = leg_y+18+(i//cpr)*20
            cv.create_rectangle(lx,ly,lx+13,ly+13, fill=cmap[v], outline="#888")
            cv.create_text(lx+17,ly+7, text=v, anchor="w", font=("Times New Roman",9))

        tot_w = x0+nc*cw+20
        tot_h = leg_y+22*(len(all_v)//cpr+2)+10
        cv.configure(scrollregion=(0,0,tot_w,tot_h))

    def _draw_scheme_garden(self, cv, d, plan):
        """
        Садівнича схема — кожна клітинка = одна рослина.
        Захисні рослини на початку/кінці ряду — сірі.
        Захисні ряди між повторностями — штриховані.
        """
        PALETTES = ["#aed6f1","#a9dfbf","#f9e79f","#f1948a","#d2b4de",
                    "#a3e4d7","#fad7a0","#d5d8dc","#82e0aa","#f0b27a",
                    "#85c1e9","#f7dc6f","#c39bd3","#76d7c4","#f8c471"]
        GUARD_COLOR  = "#d0d0d0"   # захисні рослини
        GUARD_STIPPLE = "gray50"   # захисний ряд (штрих)

        all_v = list(dict.fromkeys(p["variant"] for p in plan))
        cmap  = {v: PALETTES[i % len(PALETTES)] for i, v in enumerate(all_v)}

        n_plot   = d.get("n_plot", 5)
        g_ends   = d.get("g_ends", 1)
        g_rows   = d.get("g_rows", 1)
        reps     = d.get("reps",   3)
        k        = d.get("k",      4)   # кількість варіантів
        design   = d.get("design", "rcbd")
        row_sp   = d.get("row_sp",  4.0)
        plant_sp = d.get("plant_sp",5.0)
        unit     = d.get("unit",   "дерево")

        # Загальна кількість рослин у ряду:
        # захисні_поч + [варіант_1 ... варіант_k] (повторяється n_plot) + захисні_кін
        # Для RCBD: 1 ряд = 1 повторність = k варіантів × n_plot рослин
        total_plants_row = g_ends + k * n_plot + g_ends
        # Рядів на схемі: reps повторностей + (reps-1)*g_rows захисних рядів
        total_rows = reps + (reps - 1) * g_rows

        # Розмір клітинки
        cw = max(28, min(60, 700 // total_plants_row))
        ch = max(28, min(55, 500 // total_rows))
        pad = 3; x0 = 120; y0 = 70

        # Заголовок
        title = d.get("name") or "Садівничий дослід"
        sub   = (f"{d.get('culture','')}  |  {d['design_name']}  |  "
                 f"Схема: {row_sp}×{plant_sp} м  |  "
                 f"Рослин/ділянку: {n_plot}  |  Seed={d['seed']}")
        cv.create_text(x0 + total_plants_row*cw//2, 18,
                       text=title, font=("Times New Roman",12,"bold"), fill="#000")
        cv.create_text(x0 + total_plants_row*cw//2, 36,
                       text=sub, font=("Times New Roman",9), fill="#555")

        # Мітки колонок (позиції рослин)
        for ci in range(total_plants_row):
            if ci < g_ends or ci >= total_plants_row - g_ends:
                lbl = "З"   # Захисна
            else:
                pos = ci - g_ends
                var_idx = pos // n_plot
                plant_in_var = pos % n_plot + 1
                lbl = f"#{plant_in_var}"
            cv.create_text(x0+ci*cw+cw//2, y0-14,
                           text=lbl, font=("Times New Roman",7), fill="#777")

        # Рядки (повторності + захисні ряди)
        row_screen = 0   # лічильник рядків на екрані
        for rep_idx in range(reps):
            # Мітка повторності
            row_y = y0 + row_screen * ch
            cv.create_text(x0-6, row_y+ch//2,
                           text=f"Повт.{rep_idx+1}",
                           anchor="e", font=("Times New Roman",8,"bold"), fill="#333")

            # Беремо план для цієї повторності
            rep_plan = [p for p in plan if p["row"] == rep_idx+1]
            # Будуємо порядок варіантів у ряду
            col_to_var = {p["col"]: p["variant"] for p in rep_plan}

            # Малюємо рослини в ряду
            for ci in range(total_plants_row):
                x1 = x0 + ci * cw; y1 = row_y
                x2 = x1 + cw - pad; y2 = y1 + ch - pad

                if ci < g_ends or ci >= total_plants_row - g_ends:
                    # Захисна рослина на початку/кінці
                    cv.create_oval(x1+2, y1+2, x2-2, y2-2,
                                   fill=GUARD_COLOR, outline="#999", width=1)
                    cv.create_text((x1+x2)//2, (y1+y2)//2,
                                   text="З", font=("Times New Roman",7), fill="#888")
                else:
                    pos = ci - g_ends
                    var_col = pos // n_plot + 1  # номер варіанту (стовпець)
                    variant = col_to_var.get(var_col, "")
                    color = cmap.get(variant, "#eee")
                    # Коло = дерево/кущ
                    cv.create_oval(x1+2, y1+2, x2-2, y2-2,
                                   fill=color, outline="#666", width=1)
                    plant_in_var = pos % n_plot + 1
                    cv.create_text((x1+x2)//2, (y1+y2)//2,
                                   text=str(plant_in_var),
                                   font=("Times New Roman",7), fill="#000")

            row_screen += 1

            # Захисні ряди між повторностями
            if rep_idx < reps - 1:
                for gr in range(g_rows):
                    gy = y0 + row_screen * ch
                    cv.create_text(x0-6, gy+ch//2,
                                   text="Захисний", anchor="e",
                                   font=("Times New Roman",7,"italic"), fill="#aaa")
                    for ci in range(total_plants_row):
                        x1 = x0+ci*cw; y1 = gy
                        x2 = x1+cw-pad; y2 = y1+ch-pad
                        cv.create_oval(x1+2, y1+2, x2-2, y2-2,
                                       fill="#e8e8e8", outline="#bbb",
                                       width=1, stipple=GUARD_STIPPLE
                                       if GUARD_STIPPLE else "")
                    row_screen += 1

        # Роздільники між варіантами у рядку (вертикальні лінії)
        for var_i in range(k+1):
            lx = x0 + (g_ends + var_i * n_plot) * cw
            cv.create_line(lx, y0-5, lx, y0+reps*ch + (reps-1)*g_rows*ch + 5,
                           fill="#1a4b8c", width=1, dash=(4,3))

        # Легенда варіантів
        leg_y = y0 + total_rows * ch + 20
        cv.create_text(x0, leg_y, text="Легенда (варіанти):",
                       anchor="w", font=("Times New Roman",10,"bold"))
        cpr = 3
        for i, v in enumerate(all_v):
            lx = x0 + (i%cpr)*260; ly = leg_y+18+(i//cpr)*22
            cv.create_oval(lx, ly, lx+14, ly+14, fill=cmap[v], outline="#666")
            cv.create_text(lx+18, ly+7, text=v, anchor="w",
                           font=("Times New Roman",9))

        leg_y2 = leg_y + 18 + ((len(all_v)-1)//cpr+1)*22 + 8
        # Пояснення символів
        cv.create_oval(x0, leg_y2, x0+14, leg_y2+14,
                       fill=GUARD_COLOR, outline="#999")
        cv.create_text(x0+18, leg_y2+7,
                       text="З — захисна рослина (не обліковується)",
                       anchor="w", font=("Times New Roman",9), fill="#666")
        cv.create_oval(x0, leg_y2+20, x0+14, leg_y2+34,
                       fill="#e8e8e8", outline="#bbb")
        cv.create_text(x0+18, leg_y2+27,
                       text="Захисний ряд між повторностями",
                       anchor="w", font=("Times New Roman",9), fill="#666")
        # Схема садіння
        cv.create_text(x0, leg_y2+50,
                       text=f"Схема садіння: {row_sp} м × {plant_sp} м  |  "
                            f"Облікових {unit}/ділянку: {n_plot}  |  "
                            f"Захисних з кожного боку: {g_ends}",
                       anchor="w", font=("Times New Roman",9), fill="#333")

        tot_w = x0 + total_plants_row*cw + 20
        tot_h = leg_y2 + 70
        cv.configure(scrollregion=(0, 0, tot_w, tot_h))

    # ═══════════════════════════════════════════════════════
    # Список рандомізації
    # ═══════════════════════════════════════════════════════
    def _fill_rand(self):
        if not self._plan_data: return
        d = self._plan_data; plan = d["plan"]
        self.rand_txt.configure(state="normal")
        self.rand_txt.delete("1.0", tk.END)
        lines = [
            "═"*62,
            "     СПИСОК РАНДОМІЗАЦІЇ ПОЛЬОВОГО ДОСЛІДУ",
            "═"*62,
            f"  Назва:          {d.get('name') or '—'}",
            f"  Рік:            {d.get('year','')}",
            f"  Місце:          {d.get('loc') or '—'}",
            f"  Відповідальний: {d.get('resp') or '—'}",
            f"  Культура:       {d.get('culture','')}",
            f"  Дизайн:         {d['design_name']}",
            f"  Варіантів:      {d['k']}",
            f"  Повторностей:   {d['reps']}",
            f"  Ділянок:        {len(plan)}",
            f"  Площа ділянки:  {d['pw']} × {d['pl']} м = {d['pw']*d['pl']:.1f} м²",
            f"  Загальна площа: {d['pw']*d['pl']*len(plan):.0f} м²",
            f"  Seed рандом.:   {d['seed']}  ← зберігайте цей номер!",
            "─"*62,
            f"  {'№':<6}  {'Повторність':<16}  Варіант",
            "─"*62,
        ]
        for p in sorted(plan, key=lambda x: x["plot"]):
            lines.append(f"  {p['plot']:<6}  {p['rep']:<16}  {p['variant']}")
        lines += [
            "─"*62,
            f"  Сформовано: {datetime.now().strftime('%d.%m.%Y  %H:%M')}",
        ]
        self.rand_txt.insert("1.0", "\n".join(lines))
        self.rand_txt.configure(state="disabled")

    # ═══════════════════════════════════════════════════════
    # Польовий журнал
    # ═══════════════════════════════════════════════════════
    def _fill_journal(self):
        if not self._plan_data: return
        plan = self._plan_data["plan"]
        ind_text = self.ind_var.get().strip()
        indicators = [s.strip() for s in ind_text.split(";") if s.strip()]
        if not indicators:
            indicators = ["Показник 1","Показник 2","Показник 3","Показник 4"]

        for item in self.journal_tv.get_children():
            self.journal_tv.delete(item)

        cols = ("№ ділянки","Повторність","Варіант") + tuple(indicators) + ("Примітки",)
        self.journal_tv["columns"] = cols
        self.journal_tv["show"]    = "headings"
        w_map = {"№ ділянки":60,"Повторність":110,"Варіант":180,"Примітки":100}
        for col in cols:
            self.journal_tv.heading(col, text=col)
            w = w_map.get(col, 110)
            self.journal_tv.column(col, width=w,
                                   anchor="center" if w < 130 else "w")

        self.journal_tv.tag_configure("even", background="#f0f4ff")
        self.journal_tv.tag_configure("odd",  background="#ffffff")
        for i, p in enumerate(sorted(plan, key=lambda x: x["plot"])):
            vals = ((p["plot"], p["rep"], p["variant"])
                    + tuple("" for _ in indicators) + ("",))
            self.journal_tv.insert("","end", values=vals,
                                   tags=("even" if i%2==0 else "odd",))

    def _refresh_journal(self):
        if not self._plan_data:
            messagebox.showwarning("","Спочатку згенеруйте план."); return
        self._fill_journal()

    # ═══════════════════════════════════════════════════════
    # Збереження
    # ═══════════════════════════════════════════════════════
    def _save_rand_txt(self):
        if not self._plan_data:
            messagebox.showwarning("","Спочатку згенеруйте план."); return
        path = filedialog.asksaveasfilename(
            parent=self.win, defaultextension=".txt",
            filetypes=[("Текстовий файл","*.txt")],
            title="Зберегти список рандомізації")
        if not path: return
        try:
            with open(path,"w",encoding="utf-8") as f:
                f.write(self.rand_txt.get("1.0",tk.END))
            messagebox.showinfo("Збережено",f"Збережено:\n{path}")
        except Exception as ex:
            messagebox.showerror("Помилка",str(ex))

    def _save_excel(self):
        if not self._plan_data:
            messagebox.showwarning("","Спочатку згенеруйте план."); return
        if not HAS_OPENPYXL:
            messagebox.showerror("","Потрібен openpyxl: pip install openpyxl"); return

        path = filedialog.asksaveasfilename(
            parent=self.win, defaultextension=".xlsx",
            filetypes=[("Excel","*.xlsx")],
            title="Зберегти польовий журнал")
        if not path: return

        try:
            import openpyxl
            from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
            d  = self._plan_data; plan = d["plan"]
            ind_text   = self.ind_var.get().strip()
            indicators = [s.strip() for s in ind_text.split(";") if s.strip()]
            if not indicators:
                indicators = ["Показник 1","Показник 2","Показник 3","Показник 4"]

            wb  = openpyxl.Workbook()
            hfill = PatternFill("solid", fgColor="1A4B8C")
            hfont = Font(color="FFFFFF", bold=True,
                         name="Times New Roman", size=11)
            nfont = Font(name="Times New Roman", size=11)
            bfont = Font(name="Times New Roman", size=11, bold=True)
            ca    = Alignment(horizontal="center", vertical="center",
                              wrap_text=True)
            thin  = Side(style="thin", color="AAAAAA")
            brd   = Border(left=thin, right=thin, top=thin, bottom=thin)

            # ── Лист 1: Польовий журнал ─────────────────────
            ws = wb.active; ws.title = "Польовий журнал"

            # Шапка
            ws.merge_cells("A1:A3")
            ws["A1"] = "ПОЛЬОВИЙ ЖУРНАЛ СПОСТЕРЕЖЕНЬ"
            ws["A1"].font = Font(bold=True, name="Times New Roman", size=13)
            info_rows = [
                f"Дослід: {d.get('name') or '—'}",
                f"Рік: {d.get('year','')}    Місце: {d.get('loc') or '—'}    "
                f"Відповідальний: {d.get('resp') or '—'}",
                f"Культура: {d.get('culture','')}    Дизайн: {d['design_name']}    "
                f"Варіантів: {d['k']}    Повторностей: {d['reps']}    "
                f"Ділянок: {len(plan)}    Площа ділянки: {d['pw']}×{d['pl']} м",
            ]
            for ri, txt in enumerate(info_rows, 1):
                ws.cell(ri, 1, txt).font = nfont if ri > 1 else bfont

            # Таблиця журналу
            hr = len(info_rows) + 2
            j_hdrs = ["№ ділянки","Повторність","Варіант"] + indicators + ["Примітки"]
            for ci, h in enumerate(j_hdrs, 1):
                c = ws.cell(hr, ci, h)
                c.fill = hfill; c.font = hfont
                c.alignment = ca; c.border = brd

            PALETTES_HEX = ["AED6F1","A9DFBF","F9E79F","F1948A","D2B4DE",
                            "A3E4D7","FAD7A0","D5D8DC","82E0AA","F0B27A"]
            all_v = list(dict.fromkeys(p["variant"] for p in plan))
            vcols = {v: PALETTES_HEX[i % len(PALETTES_HEX)]
                     for i, v in enumerate(all_v)}

            for ri, p in enumerate(sorted(plan, key=lambda x: x["plot"])):
                row = hr + 1 + ri
                even = ri % 2 == 0
                rfill = PatternFill("solid",
                                    fgColor="EEF4FF" if even else "FFFFFF")
                for ci, val in enumerate(
                    [p["plot"], p["rep"], p["variant"]]
                    + [""] * len(indicators) + [""], 1
                ):
                    c = ws.cell(row, ci, val)
                    c.font = nfont; c.alignment = ca; c.border = brd
                    if ci <= 3 and even:
                        c.fill = rfill

            # Ширини стовпців
            for ci, w in enumerate(
                [9, 14, 32] + [14]*len(indicators) + [16], 1
            ):
                if ci <= 26:
                    ws.column_dimensions[chr(64+ci)].width = w
            ws.row_dimensions[hr].height = 30

            # ── Лист 2: Рандомізація ────────────────────────
            ws2 = wb.create_sheet("Рандомізація")
            r_hdrs = ["№ ділянки","Повторність","Варіант"]
            for ci, h in enumerate(r_hdrs, 1):
                c = ws2.cell(1, ci, h)
                c.fill = hfill; c.font = hfont
                c.alignment = ca; c.border = brd
            for ri, p in enumerate(sorted(plan, key=lambda x: x["plot"])):
                row = 2 + ri
                fc = PatternFill("solid",
                                 fgColor=vcols.get(p["variant"],"EEEEEE"))
                for ci, val in enumerate(
                    [p["plot"], p["rep"], p["variant"]], 1
                ):
                    c = ws2.cell(row, ci, val)
                    c.font = nfont; c.alignment = ca; c.border = brd
                    if ci == 3: c.fill = fc
            for ci, w in zip([1,2,3],[9,14,36]):
                ws2.column_dimensions[chr(64+ci)].width = w

            wb.save(path)
            messagebox.showinfo("Збережено",
                f"Збережено:\n{path}\n\n"
                "Лист 1 — Польовий журнал\n"
                "Лист 2 — Рандомізація")
        except Exception as ex:
            messagebox.showerror("Помилка збереження", str(ex))

    def _save_png(self):
        if not self._plan_data:
            messagebox.showwarning("","Спочатку згенеруйте план."); return
        if not HAS_MPL:
            messagebox.showwarning("","matplotlib недоступний."); return

        path = filedialog.asksaveasfilename(
            parent=self.win, defaultextension=".png",
            filetypes=[("PNG зображення","*.png")],
            title="Зберегти схему як PNG")
        if not path: return

        d  = self._plan_data; plan = d["plan"]
        all_v = list(dict.fromkeys(p["variant"] for p in plan))
        PALETTES = ["#aed6f1","#a9dfbf","#f9e79f","#f1948a","#d2b4de",
                    "#a3e4d7","#fad7a0","#d5d8dc","#82e0aa","#f0b27a"]
        cmap = {v: PALETTES[i%len(PALETTES)] for i,v in enumerate(all_v)}

        cols_set = sorted(set(p["col"] for p in plan))
        rows_set = sorted(set(p["row"] for p in plan))
        nc = len(cols_set); nr = len(rows_set)

        fig = Figure(figsize=(max(10, nc*1.6+2), max(5, nr*1.1+3)), dpi=130)
        ax  = fig.add_axes([0.09, 0.16, 0.88, 0.72])
        ax.set_xlim(0, nc); ax.set_ylim(0, nr); ax.set_aspect("equal")
        ax.axis("off")

        rep_map = {}
        for p in plan: rep_map[rows_set.index(p["row"])] = p["rep"]

        # Ділянки
        for p in plan:
            ci = cols_set.index(p["col"])
            ri = rows_set.index(p["row"])
            rect = matplotlib.patches.FancyBboxPatch(
                (ci+0.04, nr-ri-0.95), 0.91, 0.89,
                boxstyle="round,pad=0.02",
                facecolor=cmap.get(p["variant"],"#eee"),
                edgecolor="#777", linewidth=0.7)
            ax.add_patch(rect)
            short = (p["variant"][:13]+"…"
                     if len(p["variant"]) > 13 else p["variant"])
            ax.text(ci+0.5, nr-ri-0.5, short, ha="center", va="center",
                    fontsize=6.5, fontfamily="Times New Roman")
            ax.text(ci+0.07, nr-ri-0.1, f"#{p['plot']}",
                    ha="left", va="top", fontsize=5.5, color="#555",
                    fontfamily="Courier New")

        # Мітки рядків і стовпців
        for i,r in enumerate(rows_set):
            ax.text(-0.08, nr-i-0.5, rep_map.get(i,""),
                    ha="right", va="center", fontsize=7,
                    fontfamily="Times New Roman")
        for j,c in enumerate(cols_set):
            p_ = next(p for p in plan if p["col"]==c)
            lbl = p_.get("col_label", f"#{c}")
            ax.text(j+0.5, nr+0.08, lbl, ha="center", va="bottom",
                    fontsize=7, fontfamily="Times New Roman")

        # Заголовок
        name = d.get("name") or "План досліду"
        fig.suptitle(
            f"{name}  |  {d.get('year','')}",
            fontsize=11, fontfamily="Times New Roman", fontweight="bold", y=0.98)
        ax.set_title(
            f"Культура: {d.get('culture','')}  |  Дизайн: {d['design_name']}  "
            f"|  Варіантів: {d['k']}  |  Повторностей: {d['reps']}  "
            f"|  Seed: {d['seed']}",
            fontsize=8, fontfamily="Times New Roman")

        # Легенда
        from matplotlib.patches import Patch
        handles = [Patch(facecolor=cmap[v], edgecolor="#777", label=v)
                   for v in all_v]
        fig.legend(handles=handles, loc="lower center",
                   ncol=min(4, len(all_v)), fontsize=7,
                   framealpha=0.8, bbox_to_anchor=(0.5, 0.01))
        try:
            fig.savefig(path, dpi=150, bbox_inches="tight")
            messagebox.showinfo("Збережено", f"PNG збережено:\n{path}")
            import sys, os
            if sys.platform == "win32": os.startfile(path)
        except Exception as ex:
            messagebox.showerror("Помилка", str(ex))




def _SADTk_new_init(self, root):
    # Ініціалізуємо лише стан (без UI від orig_init)
    self.root = root
    self.table_win = None; self.report_win = None; self.graph_win = None
    self._graph_figs = {}
    self._active_cell = None; self._active_prev = None
    self._sel_anchor = None; self._sel_cells = set(); self._sel_orig = {}
    self._fill_drag = False; self._fill_rows = []; self._fill_cols = []
    self.factor_title_map = {}
    self.graph_settings = dict(DEF_GS)
    self._current_project_path = None
    self._lbf_cache = {}
    if not hasattr(self, '_gs_titles'): self._gs_titles = {}
    self._ordinal_mode = False

    root.geometry("1280x780")
    root.minsize(1100, 680)
    root.configure(bg="#0f1117")
    root.title("S.A.D. — Статистичний аналіз даних")
    set_icon(root)

    # ── Кольорова схема ─────────────────────────────────────
    C = {
        "bg":       "#0f1117",   # основний фон
        "sidebar":  "#161b27",   # бокова панель
        "card":     "#1e2336",   # картка
        "card_hov": "#252d45",   # картка hover
        "accent":   "#4a90d9",   # синій акцент
        "red":      "#c0392b",   # кнопка аналіз
        "text":     "#e8eaf0",   # основний текст
        "sub":      "#8892a4",   # підтекст
        "border":   "#2a3350",   # межі
        "sep":      "#1e2336",   # роздільник
        "green":    "#27ae60",
        "purple":   "#8e44ad",
        "orange":   "#d35400",
        "teal":     "#16a085",
    }

    # ── Статистика використання (зберігається між сесіями) ──
    usage_file = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                              ".sad_usage.json")
    usage = {}
    try:
        if os.path.exists(usage_file):
            with open(usage_file, "r", encoding="utf-8") as _f:
                usage = json.load(_f)
    except Exception: pass

    def _record_usage(key):
        usage[key] = usage.get(key, 0) + 1
        try:
            with open(usage_file, "w", encoding="utf-8") as _f:
                json.dump(usage, _f)
        except Exception: pass

    # ── Визначення всіх аналізів ────────────────────────────
    # key, назва, опис, колір, клас, needs_gs, fn, ключові слова пошуку
    ANALYSES = [
        ("anova1","Однофакторний ANOVA","CRD · RCBD · ЛК",
         "#1a4b8c",None,False,lambda: self.open_table(1),
         "сила впливу нір тьюкі дункан дисперс порівняння варіантів"),
        ("anova2","Двофакторний ANOVA","CRD · RCBD · Split-plot",
         "#1a4b8c",None,False,lambda: self.open_table(2),
         "сила впливу взаємодія факторів нір тьюкі дисперс"),
        ("anova3","Трифакторний ANOVA","Латинський квадрат",
         "#1a4b8c",None,False,lambda: self.open_table(3),
         "латинський квадрат три фактори сила впливу взаємодія"),
        ("anova4","Чотирифакторний ANOVA","Складні дизайни",
         "#1a4b8c",None,False,lambda: self.open_table(4),
         "чотири фактори складний дизайн сила впливу"),
        ("desc","Описова статистика","Mean · SD · Median · CV",
         C["green"],DescriptiveWindow,True,None,
         "середнє медіана дисперсія варіація коефіцієнт cv асиметрія ексцес"),
        ("ttest","t-тест / Манн-Уітні","Порівняння двох груп",
         C["green"],TTestWindow,False,None,
         "дві групи порівняння t критерій непараметричний"),
        ("corr","Кореляційний аналіз","Пірсон · Спірмен · Heat",
         C["accent"],CorrelationWindow,True,None,
         "зв'язок залежність матриця теплова карта пірсон спірмен"),
        ("reg","Регресійний аналіз","7 моделей · R² · p",
         C["purple"],RegressionWindow,True,None,
         "регресія прогноз r квадрат лінійна нелінійна поліном"),
        ("ancova","ANCOVA","Коваріаційний аналіз",
         C["purple"],AncovaWindow,True,None,
         "коваріата контроль змінної ancova"),
        ("manova","MANOVA","Багатовимірний дисп. аналіз",
         C["purple"],ManovaWindow,True,None,
         "кілька залежних змінних wilks pillai bagатовимірний"),
        ("rm","Повторні виміри","Within-subjects ANOVA",
         C["orange"],RepeatedMeasuresWindow,True,None,
         "повторні вимірювання часові точки динаміка within"),
        ("mix","Змішаний RM","Split-plot у часі",
         C["orange"],MixedRepeatedWindow,True,None,
         "кілька варіантів динаміка дати між групами within"),
        ("cluster","Кластерний аналіз","K-means · Ієрархічний",
         C["teal"],ClusterWindow,True,None,
         "групування схожість дендрограма kmeans кластери"),
        ("pca","PCA","Головні компоненти",
         C["teal"],PCAWindow,True,None,
         "головні компоненти зменшення вимірності biplot"),
        ("stab","Аналіз стабільності","Eberhart-Russell · GGE",
         "#8c1a1a",StabilityWindow,True,None,
         "gxe стабільність адаптація сортовипробування eberhart gge"),
        ("sample","Розмір вибірки","Потужність · n · α",
         C["sub"],SampleSizeWindow,False,None,
         "потужність розмір вибірки повторності скільки n alpha"),
        ("trial","Генерація плану","CRD · RCBD · Split-plot",
         C["teal"],TrialDesignWindow,False,None,
         "план рандомізація дослід польовий схема повторності"),
    ]

    def _open(key, cls, needs_gs, custom_fn=None):
        _record_usage(key)
        _refresh_recent()
        root.iconify()  # Згортаємо головне вікно
        if custom_fn:
            custom_fn()
        elif needs_gs:
            w = cls(root, self.graph_settings)
        else:
            w = cls(root)

    # ── HEADER ──────────────────────────────────────────────
    header = tk.Frame(root, bg="#0d1020", height=52)
    header.pack(fill=tk.X, side=tk.TOP)
    header.pack_propagate(False)


    # ── Логотип і назва ───────────────────────────────────
    logo_frm = tk.Frame(header, bg="#0d1020")
    logo_frm.pack(side=tk.LEFT, padx=12, pady=4)

    # Завантажуємо Logo.png, fallback — icon.ico
    def _load_logo(size):
        from PIL import Image, ImageTk
        base = os.path.dirname(os.path.abspath(__file__))
        for fname in ("Logo.png", "logo.png", "icon.ico"):
            p = os.path.join(base, fname)
            if os.path.exists(p):
                img = Image.open(p).convert("RGBA").resize(size, Image.LANCZOS)
                return ImageTk.PhotoImage(img)
        return None

    try:
        _logo_img = _load_logo((44, 44))
        if _logo_img:
            root._logo_img = _logo_img
            tk.Label(logo_frm, image=_logo_img, bg="#0d1020"
                     ).pack(side=tk.LEFT, padx=(0, 10))
    except Exception: pass

    # Назва
    name_f = tk.Frame(logo_frm, bg="#0d1020"); name_f.pack(side=tk.LEFT)
    tk.Label(name_f, text="S.A.D.", bg="#0d1020", fg=C["text"],
             font=("Arial", 18, "bold")).pack(anchor="w")
    tk.Label(name_f, text="Статистичний аналіз даних", bg="#0d1020",
             fg=C["sub"], font=("Arial", 9)).pack(anchor="w")

    # Права частина header — версія, розробник, підтримка
    hr = tk.Frame(header, bg="#0d1020"); hr.pack(side=tk.RIGHT, padx=16)
    def _about():
        dlg = tk.Toplevel(root); dlg.title("Про програму S.A.D.")
        dlg.geometry("480x560"); dlg.resizable(False, False)
        dlg.configure(bg=C["card"]); set_icon(dlg); dlg.grab_set()

        # Логотип у діалозі
        try:
            _li = _load_logo((120, 120))
            if _li:
                dlg._li = _li
                tk.Label(dlg, image=_li, bg=C["card"]
                         ).pack(pady=(20, 4))
        except Exception: pass

        tk.Label(dlg, text="S.A.D.", bg=C["card"], fg=C["text"],
                 font=("Arial", 22, "bold")).pack()
        tk.Label(dlg, text="Статистичний аналіз даних",
                 bg=C["card"], fg=C["sub"], font=("Arial", 12)).pack()

        tk.Frame(dlg, bg=C["border"], height=1).pack(fill=tk.X, padx=30, pady=10)

        info = [
            (f"Версія {APP_VER}",                        C["accent"], 11, "bold"),
            ("Розробник:",                               C["sub"],    9,  "normal"),
            ("Чаплоуцький Андрій Миколайович",           C["text"],   11, "bold"),
            ("Уманський національний університет",       C["sub"],    10, "normal"),
            ("Україна",                     C["sub"],    10, "normal"),
        ]
        for txt, col, sz, weight in info:
            tk.Label(dlg, text=txt, bg=C["card"], fg=col,
                     font=("Arial", sz, weight)).pack(pady=1)

        tk.Frame(dlg, bg=C["border"], height=1).pack(fill=tk.X, padx=30, pady=10)

        tk.Label(dlg, text="Призначення:",
                 bg=C["card"], fg=C["sub"], font=("Arial", 9)).pack()
        tk.Label(dlg,
                 text="Програма для статистичного аналізу\n"
                      "агрономічних та біологічних дослідів.\n"
                      "ANOVA, кореляція, регресія, PCA,\n"
                      "аналіз стабільності GxE та інше.",
                 bg=C["card"], fg=C["text"], font=("Arial", 10),
                 justify="center").pack(pady=4)

        tk.Label(dlg, text="© 2024 – 2025  Всі права захищені",
                 bg=C["card"], fg=C["border"], font=("Arial", 8)).pack(pady=(8, 2))

        tk.Button(dlg, text="Закрити", bg=C["accent"], fg="white",
                  font=("Arial", 11), relief=tk.FLAT, padx=24, pady=5,
                  cursor="hand2", command=dlg.destroy).pack(pady=12)
        dlg.bind("<Return>", lambda e: dlg.destroy())
        center_win(dlg)

    def _support():
        dlg = tk.Toplevel(root); dlg.title("Технічна підтримка")
        dlg.geometry("420x340"); dlg.resizable(False, False)
        dlg.configure(bg=C["card"]); set_icon(dlg); dlg.grab_set()

        tk.Label(dlg, text="📞 Технічна підтримка S.A.D.",
                 bg=C["card"], fg=C["text"],
                 font=("Arial", 13, "bold")).pack(pady=(20, 4))
        tk.Frame(dlg, bg=C["border"], height=1).pack(fill=tk.X, padx=30, pady=8)

        contacts = [
            ("✉ Email:",         "sad.stat.support@gmail.com"),
            ("🌐 Документація:", "docs.sad-stat.com"),
            ("💬 Telegram:",     "@sad_stat_support"),
        ]
        for lbl, val in contacts:
            row = tk.Frame(dlg, bg=C["card"]); row.pack(pady=3)
            tk.Label(row, text=lbl, bg=C["card"], fg=C["sub"],
                     font=("Arial", 10), width=16, anchor="e").pack(side=tk.LEFT)
            tk.Label(row, text=val, bg=C["card"], fg=C["accent"],
                     font=("Arial", 10, "bold"), anchor="w").pack(side=tk.LEFT, padx=8)

        tk.Frame(dlg, bg=C["border"], height=1).pack(fill=tk.X, padx=30, pady=8)
        tk.Label(dlg,
                 text="Ми відповімо протягом 1 робочого дня.\n"
                      "При зверненні вкажіть версію програми\n"
                      f"та опис проблеми. (Версія {APP_VER})",
                 bg=C["card"], fg=C["sub"], font=("Arial", 9),
                 justify="center").pack()
        tk.Button(dlg, text="Закрити", bg=C["accent"], fg="white",
                  font=("Arial", 11), relief=tk.FLAT, padx=24, pady=5,
                  cursor="hand2", command=dlg.destroy).pack(pady=12)
        dlg.bind("<Return>", lambda e: dlg.destroy())
        center_win(dlg)

    # ── Changelog ─────────────────────────────────────────────
    CHANGELOG = [
        (f"v{APP_VER}", "Поточна версія", [
            "Новий темний головний екран з картками аналізів",
            "Бокова панель з пошуком та категоріями",
            "Статистика використання аналізів",
            "Об'єднаний звіт ANOVA (текст + графіки в одному вікні)",
            "Адаптивний розмір графіків під вікно",
            "Генератор плану польового досліду (CRD/RCBD/Split-plot/ЛК)",
            "Латинський квадрат у дисперсійному аналізі",
            "Автовизначення бальних і відсоткових даних",
            "Формула регресії безпосередньо на графіку",
            "Кореляційний аналіз: два графіки в одному вікні",
        ]),
        ("v2.0", "Великий реліз", [
            "ANOVA 1-4 фактори: CRD, RCBD, Split-plot",
            "Кореляційний аналіз (Пірсон, Спірмен, теплова карта)",
            "Регресійний аналіз (7 моделей)",
            "MANOVA, ANCOVA, Повторні виміри",
            "PCA та кластерний аналіз",
            "Аналіз стабільності GxE (Eberhart-Russell, GGE biplot)",
            "Збереження та відкриття проектів",
            "Експорт звітів у Word та Excel",
        ]),
        ("v1.0", "Перший реліз", [
            "Описова статистика",
            "t-тест та критерій Манн-Уітні",
            "Однофакторний дисперсійний аналіз",
            "Базові графіки",
        ]),
    ]

    def _changelog():
        dlg = tk.Toplevel(root); dlg.title("Зміни версій — S.A.D.")
        dlg.geometry("560x520"); dlg.resizable(True, True)
        dlg.configure(bg=C["card"]); set_icon(dlg); dlg.grab_set()

        tk.Label(dlg, text="📋  Зміни версій", bg=C["card"], fg=C["text"],
                 font=("Arial", 14, "bold")).pack(pady=(16, 4))
        tk.Frame(dlg, bg=C["border"], height=1).pack(fill=tk.X, padx=20, pady=6)

        # Прокручуваний список
        outer = tk.Frame(dlg, bg=C["card"]); outer.pack(fill=tk.BOTH, expand=True, padx=16)
        vsb = ttk.Scrollbar(outer, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        cv2 = tk.Canvas(outer, bg=C["card"], highlightthickness=0,
                        yscrollcommand=vsb.set); cv2.pack(fill=tk.BOTH, expand=True)
        vsb.config(command=cv2.yview)
        inner = tk.Frame(cv2, bg=C["card"]); cv2.create_window((0,0), window=inner, anchor="nw")
        inner.bind("<Configure>", lambda e: cv2.configure(scrollregion=cv2.bbox("all")))
        cv2.bind("<MouseWheel>", lambda e: cv2.yview_scroll(int(-1*(e.delta/120)),"units"))

        for ver, tag, items in CHANGELOG:
            # Версія — заголовок
            vh = tk.Frame(inner, bg=C["card"]); vh.pack(fill=tk.X, pady=(10,2))
            tk.Label(vh, text=ver, bg=C["accent"], fg="white",
                     font=("Arial",11,"bold"), padx=10, pady=3
                     ).pack(side=tk.LEFT)
            tk.Label(vh, text=tag, bg=C["card"], fg=C["sub"],
                     font=("Arial",9), padx=8
                     ).pack(side=tk.LEFT, pady=3)
            # Пункти
            for item in items:
                tk.Label(inner, text=f"  ✓  {item}", bg=C["card"], fg=C["text"],
                         font=("Arial",9), anchor="w", justify="left"
                         ).pack(fill=tk.X, padx=8, pady=1)
            tk.Frame(inner, bg=C["border"], height=1).pack(fill=tk.X, padx=8, pady=4)

        tk.Button(dlg, text="Закрити", bg=C["accent"], fg="white",
                  font=("Arial",11), relief=tk.FLAT, padx=24, pady=5,
                  cursor="hand2", command=dlg.destroy).pack(pady=10)
        dlg.bind("<Return>", lambda e: dlg.destroy())
        center_win(dlg)

    def _license():
        dlg = tk.Toplevel(root); dlg.title("Ліцензійна угода — S.A.D.")
        dlg.geometry("600x560"); dlg.resizable(True, True)
        dlg.configure(bg=C["card"]); set_icon(dlg); dlg.grab_set()

        tk.Label(dlg, text="📄  Ліцензійна угода кінцевого користувача",
                 bg=C["card"], fg=C["text"],
                 font=("Arial", 12, "bold")).pack(pady=(16,4))
        tk.Label(dlg, text=f"S.A.D. — Статистичний аналіз даних  |  Версія {APP_VER}",
                 bg=C["card"], fg=C["sub"], font=("Arial",9)).pack()
        tk.Frame(dlg, bg=C["border"], height=1).pack(fill=tk.X, padx=20, pady=8)

        lic_text = f"""ЛІЦЕНЗІЙНА УГОДА КІНЦЕВОГО КОРИСТУВАЧА (EULA)

© 2024–2025  Чаплоуцький Андрій Миколайович
Уманський національний університет, Україна

Прочитайте цю угоду уважно перед використанням програми.
Використовуючи програму, ви погоджуєтесь з умовами цієї угоди.

──────────────────────────────────────────────────────────
1. НАДАННЯ ЛІЦЕНЗІЇ
──────────────────────────────────────────────────────────
Розробник надає вам невиключне, непередаване право на
використання програмного забезпечення S.A.D. на одному
комп'ютері (або відповідно до придбаної ліцензії).

2. ОБМЕЖЕННЯ
──────────────────────────────────────────────────────────
Вам ЗАБОРОНЕНО:
  • Копіювати, розповсюджувати або передавати програму
    третім особам без письмового дозволу розробника
  • Декомпілювати, дисасемблювати або здійснювати
    зворотну розробку програми
  • Здавати програму в оренду або субліцензувати її
  • Видаляти або змінювати повідомлення про авторські права
  • Використовувати програму для надання комерційних послуг
    без укладення окремої угоди з розробником

3. АКАДЕМІЧНЕ ТА НАУКОВЕ ВИКОРИСТАННЯ
──────────────────────────────────────────────────────────
Програма розроблена для використання в наукових
дослідженнях та навчальному процесі. Результати аналізів,
отримані за допомогою S.A.D., можуть публікуватись
у наукових роботах з посиланням на програму.

Рекомендоване посилання:
Чаплоуцький А.М. S.A.D. — Статистичний аналіз даних.
Версія {APP_VER}. Уманський НУС, Україна, 2024. [Комп'ютерна програма]

4. ІНТЕЛЕКТУАЛЬНА ВЛАСНІСТЬ
──────────────────────────────────────────────────────────
Програма та вся документація є інтелектуальною власністю
розробника і захищені законодавством України про авторське
право та міжнародними договорами.

5. ВІДМОВА ВІД ГАРАНТІЙ
──────────────────────────────────────────────────────────
Програма надається «як є» (AS IS). Розробник не гарантує
безперебійну роботу або відсутність помилок. Відповідальність
за результати статистичних аналізів лежить на користувачі.

6. КОРПОРАТИВНЕ ЛІЦЕНЗУВАННЯ
──────────────────────────────────────────────────────────
Для установ, організацій або мереж — зв'яжіться з
розробником для укладення корпоративної ліцензії.
Email: sad.stat.support@gmail.com

7. ПРИПИНЕННЯ ДІЇ ЛІЦЕНЗІЇ
──────────────────────────────────────────────────────────
Ця ліцензія діє до її розірвання. Вона автоматично
припиняється при порушенні будь-якого з умов.

© 2024–2025  Чаплоуцький А.М.  Усі права захищені."""

        outer2 = tk.Frame(dlg, bg=C["card"]); outer2.pack(fill=tk.BOTH, expand=True, padx=16)
        vsb2 = ttk.Scrollbar(outer2, orient="vertical"); vsb2.pack(side=tk.RIGHT, fill=tk.Y)
        txt = tk.Text(outer2, wrap="word", font=("Courier New", 8),
                      bg="#0f1117", fg="#c8cdd8",
                      relief=tk.FLAT, padx=12, pady=8,
                      yscrollcommand=vsb2.set, state="normal", cursor="arrow")
        txt.pack(fill=tk.BOTH, expand=True)
        vsb2.config(command=txt.yview)
        txt.insert("1.0", lic_text)
        txt.configure(state="disabled")
        txt.bind("<MouseWheel>",
                 lambda e: txt.yview_scroll(int(-1*(e.delta/120)),"units"))

        btn_f = tk.Frame(dlg, bg=C["card"]); btn_f.pack(pady=10)
        tk.Button(btn_f, text="✓ Погоджуюсь", bg=C["green"], fg="white",
                  font=("Arial",11), relief=tk.FLAT, padx=20, pady=5,
                  cursor="hand2", command=dlg.destroy).pack(side=tk.LEFT, padx=8)
        tk.Button(btn_f, text="Закрити", bg=C["card"], fg=C["sub"],
                  font=("Arial",11), relief=tk.FLAT, padx=20, pady=5,
                  cursor="hand2", command=dlg.destroy).pack(side=tk.LEFT)
        center_win(dlg)

    # ── Кнопки header ─────────────────────────────────────────
    # Роздільник
    tk.Frame(hr, bg=C["border"], width=1).pack(side=tk.LEFT, fill=tk.Y, padx=8, pady=8)

    btn_style = dict(bg="#0d1020", fg=C["sub"], font=("Arial",9),
                     relief=tk.FLAT, cursor="hand2",
                     activebackground="#161b27", activeforeground=C["text"],
                     padx=8, pady=4)
    for txt, cmd in [
        ("ℹ  Про програму", _about),
        ("📋  Ліцензія",    _license),
        ("📞  Підтримка",   _support),
    ]:
        tk.Button(hr, text=txt, command=cmd, **btn_style
                  ).pack(side=tk.LEFT, padx=2)

    # Роздільник + версія
    tk.Frame(hr, bg=C["border"], width=1).pack(side=tk.LEFT, fill=tk.Y, padx=8, pady=8)
    tk.Label(hr, text=f"v{APP_VER}", bg="#0d1020",
             fg=C["accent"], font=("Arial",9,"bold")).pack(side=tk.LEFT, padx=4)

    # ── MAIN AREA ────────────────────────────────────────────
    body = tk.Frame(root, bg=C["bg"]); body.pack(fill=tk.BOTH, expand=True)

    # ════════════════════════════════════════════════════════
    # БОКОВА ПАНЕЛЬ
    # ════════════════════════════════════════════════════════
    sidebar = tk.Frame(body, bg=C["sidebar"], width=260)
    sidebar.pack(side=tk.LEFT, fill=tk.Y); sidebar.pack_propagate(False)

    # Пошук
    sf = tk.Frame(sidebar, bg=C["sidebar"], pady=8, padx=10)
    sf.pack(fill=tk.X)
    search_var = tk.StringVar()
    search_entry = tk.Entry(sf, textvariable=search_var,
                            bg=C["card"], fg=C["text"], insertbackground=C["text"],
                            relief=tk.FLAT, font=("Arial",10),
                            highlightthickness=1, highlightbackground=C["border"])
    search_entry.pack(fill=tk.X, ipady=5)
    search_entry.insert(0, "🔍  Пошук аналізу...")
    search_entry.config(fg=C["sub"])
    def _search_focus_in(e):
        if search_entry.get().startswith("🔍"):
            search_entry.delete(0, tk.END); search_entry.config(fg=C["text"])
    def _search_focus_out(e):
        if not search_entry.get().strip():
            search_entry.insert(0, "🔍  Пошук аналізу..."); search_entry.config(fg=C["sub"])
    search_entry.bind("<FocusIn>", _search_focus_in)
    search_entry.bind("<FocusOut>", _search_focus_out)

    # Роздільник
    tk.Frame(sidebar, bg=C["border"], height=1).pack(fill=tk.X)

    # Список аналізів у sidebar
    sb_canvas = tk.Canvas(sidebar, bg=C["sidebar"], highlightthickness=0)
    sb_vsb = tk.Scrollbar(sidebar, orient="vertical", command=sb_canvas.yview)
    sb_vsb.pack(side=tk.RIGHT, fill=tk.Y)
    sb_canvas.pack(fill=tk.BOTH, expand=True)
    sb_canvas.configure(yscrollcommand=sb_vsb.set)
    sb_inner = tk.Frame(sb_canvas, bg=C["sidebar"])
    sb_canvas.create_window((0,0), window=sb_inner, anchor="nw")

    def _mw_sidebar(e):
        sb_canvas.yview_scroll(int(-1*(e.delta/120)), "units")
    sb_canvas.bind("<MouseWheel>", _mw_sidebar)
    sb_inner.bind("<MouseWheel>", _mw_sidebar)
    sidebar.bind("<MouseWheel>", _mw_sidebar)
    sf.bind("<MouseWheel>", _mw_sidebar)
    # Bind all sidebar children after make_sidebar
    def _bind_sb_children():
        for w in sb_inner.winfo_children():
            w.bind("<MouseWheel>", _mw_sidebar)
            for ch in w.winfo_children():
                ch.bind("<MouseWheel>", _mw_sidebar)
    sb_inner.bind("<Configure>", lambda e: (
        sb_canvas.configure(scrollregion=sb_canvas.bbox("all")),
        _bind_sb_children()))

    CATEGORIES = [
        ("ANOVA",             ["anova1","anova2","anova3","anova4"]),
        ("Базові методи",     ["desc","ttest"]),
        ("Зв'язок змінних",  ["corr","reg","ancova"]),
        ("Багатовимірні",     ["manova","rm","mix"]),
        ("Багатовимірний ML", ["cluster","pca"]),
        ("Спеціальні",        ["stab","sample","trial"]),
    ]
    _ana_map = {a[0]: a for a in ANALYSES}
    _sb_btns = {}

    def _make_sidebar():
        for w in sb_inner.winfo_children(): w.destroy()
        q = search_var.get().lower().strip()
        if q.startswith("🔍"): q = ""
        for cat_name, keys in CATEGORIES:
            def _matches(k):
                if not q: return True
                a = _ana_map[k]
                return (q in a[1].lower() or q in a[2].lower()
                        or (len(a) > 7 and q in a[7].lower()))
            filtered = [k for k in keys if _matches(k)]
            if not filtered: continue
            tk.Label(sb_inner, text=cat_name.upper(), bg=C["sidebar"],
                     fg=C["sub"], font=("Arial",8,"bold"),
                     anchor="w", padx=12, pady=8
                     ).pack(fill=tk.X)
            for k in filtered:
                a = _ana_map[k]
                col = a[3]
                btn_f = tk.Frame(sb_inner, bg=C["sidebar"]); btn_f.pack(fill=tk.X)
                cnt = usage.get(k,0)
                lbl_txt = f"  {a[1]}"
                b = tk.Label(btn_f, text=lbl_txt, bg=C["sidebar"], fg=C["text"],
                             font=("Arial",10), anchor="w", padx=10, pady=5,
                             cursor="hand2")
                b.pack(side=tk.LEFT, fill=tk.X, expand=True)
                if cnt > 0:
                    tk.Label(btn_f, text=str(cnt), bg=C["sidebar"],
                             fg=C["sub"], font=("Arial",8),
                             padx=6).pack(side=tk.RIGHT)
                # Кольоровий лівий бордер
                border = tk.Frame(btn_f, bg=col, width=3)
                border.place(relx=0, rely=0, relheight=1)
                def _enter(e, f=btn_f, brd=border):
                    f.configure(bg=C["card_hov"])
                    for ch in f.winfo_children(): ch.configure(bg=C["card_hov"])
                    brd.configure(bg=C["accent"])
                def _leave(e, f=btn_f, brd=border, c=col):
                    f.configure(bg=C["sidebar"])
                    for ch in f.winfo_children(): ch.configure(bg=C["sidebar"])
                    brd.configure(bg=c)
                for w2 in [btn_f, b]:
                    w2.bind("<Enter>", _enter)
                    w2.bind("<Leave>", _leave)
                    w2.bind("<Button-1>", lambda e, k2=k, a2=a:
                            _open(k2, a2[4], a2[5], a2[6]))
                _sb_btns[k] = btn_f
            tk.Frame(sb_inner, bg=C["border"], height=1).pack(fill=tk.X, padx=10)

    search_var.trace_add("write", lambda *_: _make_sidebar())
    _make_sidebar()

    # ════════════════════════════════════════════════════════
    # ПРАВА ЧАСТИНА — КАРТКИ
    # ════════════════════════════════════════════════════════
    right = tk.Frame(body, bg=C["bg"]); right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    # Заголовок
    top_r = tk.Frame(right, bg=C["bg"], padx=24, pady=16)
    top_r.pack(fill=tk.X)
    tk.Label(top_r, text="Оберіть тип аналізу", bg=C["bg"],
             fg=C["text"], font=("Arial",18,"bold")).pack(side=tk.LEFT)

    # Кнопки проекту справа
    proj_f = tk.Frame(top_r, bg=C["bg"]); proj_f.pack(side=tk.RIGHT)
    def _load_proj_home():
        path = filedialog.askopenfilename(
            parent=root,
            filetypes=[("SAD проект","*.sadp"),("JSON","*.json"),("All","*.*")],
            title="Відкрити проект S.A.D.")
        if not path: return
        try:
            with open(path,"r",encoding="utf-8") as _f:
                d = json.load(_f)
        except Exception as ex:
            messagebox.showerror("Помилка відкриття", str(ex)); return
        ptype = d.get("type","")
        # ANOVA / головний проект — factors_count є завжди
        if not ptype or "factors_count" in d:
            project_from_dict(self, d)
            messagebox.showinfo("Проект відкрито", f"Проект завантажено:\n{path}")
            return
        # Спеціалізовані проекти
        def _after_open(win_obj):
            root.after(300, lambda w=win_obj: _fill_win(w, d, path))
        def _fill_win(w, dd, pp):
            try:
                rows = dd.get("rows_data",[])
                envs = (dd.get("env_vars") or dd.get("time_vars") or
                        dd.get("col_vars") or [])
                n_need = (1+len(envs)) if envs else len(rows[0]) if rows else 0
                while getattr(w,"cols_n",0) < n_need:
                    if hasattr(w,"_add_col"): w._add_col()
                    else: break
                for i,nm in enumerate(envs):
                    for attr in ["time_vars","col_vars","env_vars"]:
                        vl = getattr(w,attr,[])
                        if i<len(vl): vl[i].set(nm); break
                while len(getattr(w,"entries",[])) < len(rows):
                    if hasattr(w,"_add_row"): w._add_row()
                    else: break
                for i,rv in enumerate(rows):
                    for j,v in enumerate(rv):
                        ents = getattr(w,"entries",[])
                        if i<len(ents) and j<len(ents[i]):
                            ents[i][j].delete(0,tk.END)
                            ents[i][j].insert(0,v)
                messagebox.showinfo("Завантажено", f"Проект завантажено:\n{pp}")
            except Exception as ex:
                messagebox.showerror("Помилка завантаження", str(ex))
        win_map = {
            "stability":              lambda: _after_open(StabilityWindow(root,self.graph_settings)),
            "mixed_repeated_measures":lambda: _after_open(MixedRepeatedWindow(root,self.graph_settings)),
            "repeated_measures":      lambda: _after_open(RepeatedMeasuresWindow(root,self.graph_settings)),
            "correlation":            lambda: _after_open(CorrelationWindow(root,self.graph_settings)),
        }
        for key, fn in win_map.items():
            if key in ptype:
                fn(); return
        messagebox.showinfo("Проект відкрито",
            f"Тип проекту: «{ptype}»\n"
            "Відкрийте відповідний аналіз вручну і\n"
            "скористайтесь «📂 Відкрити проект» у тому вікні.")

    tk.Button(proj_f, text="📂 Відкрити проект", bg=C["card"], fg=C["text"],
              font=("Arial",10), relief=tk.FLAT, padx=12, pady=6,
              cursor="hand2", activebackground=C["card_hov"],
              command=_load_proj_home).pack(side=tk.LEFT, padx=4)

    # Прокручуваний контент
    content_canvas = tk.Canvas(right, bg=C["bg"], highlightthickness=0)
    c_vsb = tk.Scrollbar(right, orient="vertical", command=content_canvas.yview)
    c_vsb.pack(side=tk.RIGHT, fill=tk.Y)
    content_canvas.pack(fill=tk.BOTH, expand=True)
    content_canvas.configure(yscrollcommand=c_vsb.set)
    cf = tk.Frame(content_canvas, bg=C["bg"])
    cf_win = content_canvas.create_window((0,0), window=cf, anchor="nw")

    def _on_cf_configure(e):
        content_canvas.configure(scrollregion=content_canvas.bbox("all"))
        # Прокрутка для всіх нових дочірніх елементів
        def _bind_all(w):
            try: w.bind("<MouseWheel>", _mw_content)
            except Exception: pass
            for ch in w.winfo_children(): _bind_all(ch)
        _bind_all(cf)
    cf.bind("<Configure>", _on_cf_configure)

    content_canvas.bind("<Configure>",
                        lambda e: content_canvas.itemconfig(cf_win, width=e.width))

    def _mw_content(e):
        delta = int(-1*(e.delta/120))
        top, bot = content_canvas.yview()
        if delta < 0 and top <= 0.001: return
        if delta > 0 and bot >= 0.999: return
        content_canvas.yview_scroll(delta, "units")

    def _mw_sidebar_global(e):
        delta = int(-1*(e.delta/120))
        top, bot = sb_canvas.yview()
        if delta < 0 and top <= 0.001: return
        if delta > 0 and bot >= 0.999: return
        sb_canvas.yview_scroll(delta, "units")

    # Глобальне прив'язування через root — найнадійніший спосіб
    def _global_mw(e):
        wx = e.widget
        # Визначаємо чи курсор над правою частиною чи лівою
        try:
            abs_x = e.widget.winfo_rootx()
            sidebar_right = sidebar.winfo_rootx() + sidebar.winfo_width()
            if abs_x >= sidebar_right:
                _mw_content(e)
            else:
                _mw_sidebar_global(e)
        except Exception:
            _mw_content(e)

    root.bind_all("<MouseWheel>", _global_mw)
    right.bind("<MouseWheel>", _mw_content)
    content_canvas.bind("<MouseWheel>", _mw_content)
    cf.bind("<MouseWheel>", _mw_content)

    def _card(parent, key, name, desc, color, cls, needs_gs, custom_fn,
              large=False):
        """Сучасна об'ємна картка аналізу."""
        w = 280 if large else 210
        h = 120 if large else 96
        pad = 16 if large else 12
        name_sz = 13 if large else 11
        desc_sz = 9 if large else 8
        _dark  = _darken(color)
        _light = _lighten(color)

        # Зовнішня рамка — глибока тінь для об'єму
        _darker = _darken(color, 50)
        outer = tk.Frame(parent, bg=_darker,
                         width=w+4, height=h+4, cursor="hand2")
        outer.pack_propagate(False)

        # Середня рамка — бічна тінь
        mid = tk.Frame(outer, bg=_dark, cursor="hand2")
        mid.pack_propagate(False)
        mid.pack(fill=tk.BOTH, expand=True, padx=(1,3), pady=(1,3))

        # Основний фрейм
        frm = tk.Frame(mid, bg=color, cursor="hand2")
        frm.pack_propagate(False)
        frm.pack(fill=tk.BOTH, expand=True)

        # Верхня світла смужка — ефект блиску
        shine = tk.Frame(frm, bg=_light, height=3)
        shine.pack(fill=tk.X, side=tk.TOP)

        # Ліва світла смужка — бічне підсвічування
        left_shine = tk.Frame(frm, bg=_lighten(color, 20), width=2)
        left_shine.pack(side=tk.LEFT, fill=tk.Y)

        # Вміст
        inner = tk.Frame(frm, bg=color, padx=pad, pady=pad-2)
        inner.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        tk.Label(inner, text=name, bg=color, fg="white",
                 font=("Arial", name_sz, "bold"),
                 wraplength=w-pad*2, justify="left", anchor="w"
                 ).pack(anchor="w")
        tk.Label(inner, text=desc, bg=color,
                 font=("Arial", desc_sz),
                 wraplength=w-pad*2, justify="left", anchor="w",
                 fg="#cccccc"
                 ).pack(anchor="w", pady=(3,0))

        cnt = usage.get(key, 0)
        if cnt > 0 and large:
            tk.Label(inner, text=f"↳ використовували {cnt}×",
                     bg=color, fg="#aaaaaa",
                     font=("Arial",7)).pack(anchor="w", pady=(4,0))

        # Hover
        def _e(e):
            outer.configure(bg=_darker)
            mid.configure(bg=_dark)
            frm.configure(bg=_dark)
            shine.configure(bg=_lighten(_dark,20))
            left_shine.configure(bg=_lighten(_dark,15))
            inner.configure(bg=_dark)
            for ch in inner.winfo_children(): ch.configure(bg=_dark)
        def _l(e):
            outer.configure(bg=_darker)
            mid.configure(bg=_dark)
            frm.configure(bg=color)
            shine.configure(bg=_light)
            left_shine.configure(bg=_lighten(color,20))
            inner.configure(bg=color)
            for ch in inner.winfo_children(): ch.configure(bg=color)
        click_cmd = lambda e, k=key, cl=cls, ng=needs_gs, cf2=custom_fn:                     _open(k, cl, ng, cf2)
        for w2 in ([outer, mid, frm, shine, left_shine, inner] +
                   list(inner.winfo_children())):
            w2.bind("<Enter>", _e)
            w2.bind("<Leave>", _l)
            w2.bind("<Button-1>", click_cmd)
        return outer

    def _darken(hex_color, amt=30):
        try:
            h = hex_color.lstrip("#")
            r,g,b = int(h[0:2],16),int(h[2:4],16),int(h[4:6],16)
            return f"#{max(0,r-amt):02x}{max(0,g-amt):02x}{max(0,b-amt):02x}"
        except Exception: return hex_color

    def _lighten(hex_color, amt=40):
        try:
            h = hex_color.lstrip("#")
            r,g,b = int(h[0:2],16),int(h[2:4],16),int(h[4:6],16)
            return f"#{min(255,r+amt):02x}{min(255,g+amt):02x}{min(255,b+amt):02x}"
        except Exception: return hex_color

    def _refresh_recent():
        """Перебудовує секцію «Нещодавні»."""
        for w in cf.winfo_children(): w.destroy()

        padx = 24

        # ── Нещодавні / Часті (великі картки) ──────────────
        recent_keys = sorted(usage.keys(), key=lambda k: -usage.get(k,0))
        recent_keys = [k for k in recent_keys if k in _ana_map][:6]

        if recent_keys:
            sec1 = tk.Frame(cf, bg=C["bg"]); sec1.pack(fill=tk.X, padx=padx, pady=(8,4))
            tk.Label(sec1, text="Нещодавні та часті", bg=C["bg"],
                     fg=C["sub"], font=("Arial",10,"bold")).pack(anchor="w")
            cards_f1 = tk.Frame(cf, bg=C["bg"]); cards_f1.pack(fill=tk.X, padx=padx, pady=4)
            for k in recent_keys:
                a = _ana_map[k]
                c = _card(cards_f1, k, a[1], a[2], a[3], a[4], a[5], a[6], large=True)
                c.pack(side=tk.LEFT, padx=(0,10), pady=4)
            tk.Frame(cf, bg=C["border"], height=1).pack(fill=tk.X, padx=padx, pady=4)

        # ── Всі аналізи по категоріях ───────────────────────
        for cat_name, keys in CATEGORIES:
            sec = tk.Frame(cf, bg=C["bg"]); sec.pack(fill=tk.X, padx=padx, pady=(12,4))
            tk.Label(sec, text=cat_name, bg=C["bg"],
                     fg=C["text"], font=("Arial",12,"bold")).pack(anchor="w")
            row_f = tk.Frame(cf, bg=C["bg"]); row_f.pack(fill=tk.X, padx=padx, pady=4)
            for k in keys:
                a = _ana_map[k]
                c = _card(row_f, k, a[1], a[2], a[3], a[4], a[5], a[6])
                c.pack(side=tk.LEFT, padx=(0,8), pady=4)

        # ── Footer ──────────────────────────────────────────
        footer_f = tk.Frame(cf, bg="#0d1020"); footer_f.pack(fill=tk.X, side=tk.BOTTOM)
        tk.Frame(footer_f, bg=C["border"], height=1).pack(fill=tk.X)
        footer = tk.Frame(footer_f, bg="#0d1020", padx=24, pady=8)
        footer.pack(fill=tk.X)
        tk.Label(footer,
                 text="© 2024–2025  Чаплоуцький А.М.  |  "
                      "Уманський НУ, Україна  |  "
                      "Усі права захищені",
                 bg="#0d1020", fg=C["sub"],
                 font=("Arial", 8)).pack(side=tk.LEFT)
        tk.Label(footer,
                 text=f"S.A.D.  v{APP_VER}",
                 bg="#0d1020", fg=C["border"],
                 font=("Arial", 8)).pack(side=tk.RIGHT)

    _refresh_recent()


SADTk.__init__ = _SADTk_new_init


if __name__ == "__main__":
    root = tk.Tk()
    set_icon(root)
    app = SADTk(root)
    root.mainloop()
