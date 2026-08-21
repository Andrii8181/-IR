# sad_app.py — Головне вікно програми (SADTk)
# -*- coding: utf-8 -*-
from sad_common import *
from sad_correlation import CorrelationWindow, GraphSettingsDlg
from sad_journal_trial import open_indicator_for_anova

# ═══════════════════════════════════════════════════════════════
# GUI — SADTk
# ═══════════════════════════════════════════════════════════════
class SADTk:
    SEL_BG = "#cce5ff"; SEL_ANC = "#99ccff"; ACT_BG = "#fff3c4"

    def __init__(self, root):
        self.root = root
        root.title("S.A.D. — Статистичний аналіз даних")
        root.geometry("1000x580"); set_icon(root)
        try: ttk.Style().theme_use("clam")
        except Exception: pass

        mf = tk.Frame(root, bg="white"); mf.pack(expand=True, fill=tk.BOTH)
        tk.Label(mf, text="S.A.D. — Статистичний аналіз даних",
                 font=("Times New Roman", 20, "bold"), fg="#000000", bg="white").pack(pady=16)

        bf = tk.Frame(mf, bg="white"); bf.pack(pady=8)
        for i, (txt, fc) in enumerate([("Однофакторний аналіз", 1), ("Двофакторний аналіз", 2),
                                        ("Трифакторний аналіз", 3), ("Чотирифакторний аналіз", 4)]):
            tk.Button(bf, text=txt, width=22, height=2, font=("Times New Roman", 13),
                      command=lambda f=fc: self.open_table(f)).grid(row=i // 2, column=i % 2, padx=10, pady=6)

        tk.Button(mf, text="Кореляційний аналіз", width=30, height=2, font=("Times New Roman", 13),
                  bg="#1a4b8c", fg="white",
                  command=self.open_correlation).pack(pady=(4, 10))

        tk.Label(mf, text="Виберіть тип аналізу → Введіть дані → Аналіз даних",
                 font=("Times New Roman", 12), fg="#555555", bg="white").pack(pady=4)

        self.table_win = None; self.report_win = None; self.graph_win = None
        self._graph_figs = {}
        self._active_cell = None; self._active_prev = None
        self._sel_anchor = None; self._sel_cells = set(); self._sel_orig = {}
        self._fill_drag = False; self._fill_rows = []; self._fill_cols = []
        self.factor_title_map = {}
        self.graph_settings = dict(DEF_GS)
        self._current_project_path = None
        self._lbf_cache = {}

    # ── open correlation ──────────────────────────────────────
    def open_correlation(self):
        CorrelationWindow(self.root, self.graph_settings)

    def _show_anova_help(self, parent, fc):
        """Довідка специфічна для кожного типу ANOVA."""
        help_texts = {
            1: """
ОДНОФАКТОРНИЙ ДИСПЕРСІЙНИЙ АНАЛІЗ
═══════════════════════════════════

ЩО РОБИТЬ?
  Порівнює СЕРЕДНІ по 3 і більше групах (варіантах).
  Перевіряє: «Чи хоча б одна група відрізняється від інших?»

  Перевага над t-тестом: при кількох порівняннях одночасно
  не накопичується помилка першого роду.

СТРУКТУРА ТАБЛИЦІ:
  Стовпець 1: Фактор A (назви варіантів/обробок)
  Стовпці 2+: Повторності (числові значення)

  Приклад (3 дози добрива, 4 повторності):
  | Доза    | Повт.1 | Повт.2 | Повт.3 | Повт.4 |
  | Контроль|  5.2   |  4.8   |  5.5   |  4.9   |
  | Доза 1  |  6.1   |  6.4   |  5.9   |  6.3   |
  | Доза 2  |  7.2   |  6.8   |  7.5   |  7.1   |

  Перейменуйте «Фактор A» (подвійний клік) на назву вашого фактора.

ВИБІР ДИЗАЙНУ:
  CRD — ділянки однорідні, варіанти розміщені випадково
  RCBD — є блоки (рельєф, родючість) — ефективніший
  Split-plot — лише для 2+ факторів

ІНТЕРПРЕТАЦІЯ:
  F-тест значущий → є різниця між варіантами
  Переходьте до пост-хок для визначення ЯКИХ САМЕ
""",
            2: """
ДВОФАКТОРНИЙ ДИСПЕРСІЙНИЙ АНАЛІЗ
══════════════════════════════════

ЩО РОБИТЬ?
  Оцінює вплив ДВОХ факторів (A і B) та їх взаємодії (A×B).

  Три питання одночасно:
  1. Чи значущий головний ефект фактора A?
  2. Чи значущий головний ефект фактора B?
  3. Чи є взаємодія A×B? (найважливіше!)

СТРУКТУРА ТАБЛИЦІ:
  Стовпець 1: Фактор A (рівні першого фактора)
  Стовпець 2: Фактор B (рівні другого фактора)
  Стовпці 3+: Повторності (числові значення)

  Приклад (сорт × доза добрива):
  | Сорт    | Доза  | Повт.1 | Повт.2 | Повт.3 |
  | Сорт А  | Доза1 |  5.2   |  4.8   |  5.5   |
  | Сорт А  | Доза2 |  6.1   |  6.4   |  5.9   |
  | Сорт Б  | Доза1 |  4.9   |  5.1   |  4.7   |
  | Сорт Б  | Доза2 |  7.2   |  6.8   |  7.5   |

ВЗАЄМОДІЯ A×B:
  Значуща → ефект фактора A залежить від рівня B.
  Тобто: один сорт краще реагує на добрива, інший — ні.
  При значущій взаємодії — аналізуйте прості ефекти!
  Незначуща → ефекти факторів незалежні.
""",
            3: """
ТРИФАКТОРНИЙ ДИСПЕРСІЙНИЙ АНАЛІЗ
══════════════════════════════════

ЩО РОБИТЬ?
  Оцінює вплив трьох факторів (A, B, C) та їх взаємодій:
  A×B, A×C, B×C, і потрійну взаємодію A×B×C.

СТРУКТУРА ТАБЛИЦІ:
  Стовпець 1: Фактор A
  Стовпець 2: Фактор B
  Стовпець 3: Фактор C
  Стовпці 4+: Повторності

  Кожна унікальна комбінація A×B×C = один рядок.
  При 3 рівнях кожного фактора: 3×3×3 = 27 рядків.

ВЗАЄМОДІЯ A×B×C:
  Значуща → ефект пари факторів залежить від третього.
  Наприклад: вплив добрива і сорту різний в різні роки.

ПОРАДА:
  При 3 факторах рекомендується Тип III SS.
  Кількість рядків = k_A × k_B × k_C (де k = кількість рівнів).

═══════════════════════════════════════════════════════════
ОСОБЛИВИЙ ВИПАДОК: ЛАТИНСЬКИЙ КВАДРАТ (у 3-факторному)
═══════════════════════════════════════════════════════════

Якщо у вас Латинський квадрат — використовуйте саме
3-факторний аналіз з таким призначенням факторів:

  Фактор A = Варіант  (назви обробок: N60, N90, N120...)
  Фактор B = Рядок    (номери рядів поля: 1, 2, 3...)
  Фактор C = Стовпець (номери стовпців поля: 1, 2, 3...)

Структура таблиці (k=4 варіанти, 4×4 квадрат):
  | Варіант | Рядок | Стовп | Значення |
  | N60     |   1   |   1   |   18.4   |
  | N90     |   1   |   2   |   21.3   |
  | N120    |   1   |   3   |   22.8   |
  | N0      |   1   |   4   |   15.2   |
  | N0      |   2   |   1   |   14.8   |
  ...

Кожен рядок таблиці = одна ділянка = одне значення.
Стовпець "Значення" вводьте у колонку «Повт.1».

У вікні параметрів оберіть:  Дизайн → Латинський квадрат

Програма автоматично:
  ✓ Виносить SS_рядки і SS_стовпці з помилки
  ✓ Розраховує правильний df = (k-1)(k-2)
  ✓ Перевіряє що k варіантів = k рядків = k стовпців
""",
            4: """
ЧОТИРИФАКТОРНИЙ ДИСПЕРСІЙНИЙ АНАЛІЗ
═════════════════════════════════════

ЩО РОБИТЬ?
  Оцінює вплив чотирьох факторів (A, B, C, D) та всіх
  можливих взаємодій (4 парних + 4 потрійних + 1 четверна).

СТРУКТУРА ТАБЛИЦІ:
  Стовпці 1-4: Фактори A, B, C, D
  Стовпці 5+: Повторності

  При 2 рівнях кожного: 2×2×2×2 = 16 комбінацій.
  При 3 рівнях кожного: 3×3×3×3 = 81 комбінація!

ВАЖЛИВО:
  Чотирифакторний аналіз вимагає дуже великого досліду.
  Мінімум: 2 повторності × 16 комбінацій = 32 спостереження.
  Рекомендується Тип III SS.
  Потрійні і четверна взаємодії рідко бувають значущими
  і важко інтерпретуються.

ПОРАДА:
  При незначущих вищих взаємодіях — спростіть до трифакторного.
"""
        }
        text = help_texts.get(fc, "Довідка недоступна.")

        win = tk.Toplevel(parent)
        win.title(f"Довідка — {['','Одно','Дво','Три','Чотири'][fc]}факторний ANOVA")
        win.geometry("680x540"); set_icon(win)
        frm = tk.Frame(win); frm.pack(fill=tk.BOTH, expand=True, padx=8, pady=6)
        vsb = ttk.Scrollbar(frm, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        txt = tk.Text(frm, wrap="word", font=("Times New Roman",11),
                      yscrollcommand=vsb.set, relief=tk.FLAT,
                      bg="#fafafa", padx=10, pady=8, cursor="arrow")
        txt.pack(fill=tk.BOTH, expand=True); vsb.config(command=txt.yview)
        txt.insert("1.0", text.strip()); txt.configure(state="disabled")
        txt.bind("<MouseWheel>",
                 lambda e: txt.yview_scroll(int(-1*(e.delta/120)),"units"))
        tk.Button(win, text="Закрити", command=win.destroy,
                  font=("Times New Roman",11)).pack(pady=6)

    # ── factor titles ─────────────────────────────────────────
    def ftitle(self, fk): return self.factor_title_map.get(fk, f"Фактор {fk}")
    def _set_ftitle(self, fk, t): self.factor_title_map[fk] = t.strip() or f"Фактор {fk}"

    # ── project save/load ─────────────────────────────────────
    def save_project(self):
        if not hasattr(self, "entries") or not self.entries:
            messagebox.showwarning("", "Відкрийте таблицю."); return
        path = filedialog.asksaveasfilename(
            parent=self.table_win or self.root, title="Зберегти проект",
            defaultextension=".sadp",
            filetypes=[("SAD проект", "*.sadp"), ("JSON", "*.json"), ("Усі", "*.*")])
        if not path: return
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(project_to_dict(self), f, ensure_ascii=False, indent=2)
            self._current_project_path = path
            messagebox.showinfo("Збережено", f"Проект збережено:\n{path}")
        except Exception as ex: messagebox.showerror("Помилка", str(ex))

    def load_project(self):
        path = filedialog.askopenfilename(
            parent=self.table_win or self.root, title="Відкрити проект",
            filetypes=[("SAD проект", "*.sadp"), ("JSON", "*.json"), ("Усі", "*.*")])
        if not path: return
        try:
            with open(path, "r", encoding="utf-8") as f: d = json.load(f)
            project_from_dict(self, d)
            self._current_project_path = path
            messagebox.showinfo("Відкрито", f"Проект відкрито:\n{path}")
        except Exception as ex: messagebox.showerror("Помилка", str(ex))

    def clear_project(self):
        if not messagebox.askyesno("Очистити", "Очистити всі дані таблиці?"): return
        for row in self.entries:
            for e in row: e.delete(0, tk.END)

    # ── selection ─────────────────────────────────────────────
    def _clear_sel(self):
        # No visual highlighting — just clear tracking state
        self._sel_cells.clear(); self._sel_anchor = None; self._sel_orig.clear()

    def _restore_bg(self, r, c):
        pass   # no-op: no coloring was applied

    def _apply_sel(self, cells):
        pass   # no-op: selection tracking only, no visual highlight

    def _sel_range(self, r1, c1, r2, c2):
        new = {(r, c) for r in range(min(r1, r2), max(r1, r2) + 1)
               for c in range(min(c1, c2), max(c1, c2) + 1)
               if r < len(self.entries) and c < len(self.entries[r])}
        self._sel_cells = new

    def _sel_bounds(self):
        if not self._sel_cells: return None
        rs = [r for r, c in self._sel_cells]; cs = [c for r, c in self._sel_cells]
        return min(rs), min(cs), max(rs), max(cs)

    def _near_br(self, w, mg=6):
        try:
            px = w.winfo_pointerx(); py = w.winfo_pointery()
            x0 = w.winfo_rootx(); y0 = w.winfo_rooty()
            return (x0 + w.winfo_width() - mg <= px <= x0 + w.winfo_width()) and \
                   (y0 + w.winfo_height() - mg <= py <= y0 + w.winfo_height())
        except Exception: return False

    def _sel_handle_cell(self):
        b = self._sel_bounds()
        if b is None: return None
        try: return self.entries[b[2]][b[3]]
        except Exception: return None

    # ── active cell ───────────────────────────────────────────
    def _set_active(self, w):
        if self._active_cell is w: return
        if isinstance(self._active_cell, tk.Entry) and self._active_prev:
            try: self._active_cell.configure(**self._active_prev)
            except Exception: pass
        self._active_cell = w
        if isinstance(w, tk.Entry):
            self._active_prev = {"bg": w.cget("bg"), "highlightthickness": int(w.cget("highlightthickness")),
                                  "highlightbackground": w.cget("highlightbackground"),
                                  "highlightcolor": w.cget("highlightcolor"),
                                  "relief": w.cget("relief"),
                                  "bd": int(w.cget("bd")) if str(w.cget("bd")).isdigit() else 1}
            try: w.configure(bg=self.ACT_BG, highlightthickness=3,
                              highlightbackground="#c62828", highlightcolor="#c62828",
                              relief=tk.SOLID, bd=1)
            except Exception: pass

    # ── bind cell ─────────────────────────────────────────────
    def bind_cell(self, e):
        e.bind("<Return>",               self._on_enter)
        e.bind("<Up>",                   self._on_arrow)
        e.bind("<Down>",                 self._on_arrow)
        e.bind("<Left>",                 self._on_arrow)
        e.bind("<Right>",                self._on_arrow)
        e.bind("<Control-c>",            self._on_copy)
        e.bind("<Control-C>",            self._on_copy)
        e.bind("<Control-v>",            self._on_paste)
        e.bind("<Control-V>",            self._on_paste)
        e.bind("<FocusIn>",              lambda ev: self._set_active(ev.widget))
        e.bind("<ButtonPress-1>",        self._on_press)
        e.bind("<B1-Motion>",            self._on_drag)
        e.bind("<ButtonRelease-1>",      self._on_release)
        e.bind("<Motion>",               self._on_motion)
        e.bind("<Leave>",                lambda ev: ev.widget.configure(cursor=""))
        e.bind("<Shift-ButtonPress-1>",  self._on_shift_click)

    # ── mouse events ──────────────────────────────────────────
    def _on_motion(self, event):
        w = event.widget
        if not isinstance(w, tk.Entry): return
        pos = self._pos(w)
        if not pos: return
        r, c = pos
        if c >= self.factors_count: w.configure(cursor=""); return
        if self._sel_cells and self._sel_handle_cell() is w and self._near_br(w):
            w.configure(cursor="crosshair")
        elif not self._sel_cells and self._near_br(w):
            w.configure(cursor="crosshair")
        else:
            w.configure(cursor="")

    def _on_press(self, event):
        w = event.widget
        if not isinstance(w, tk.Entry): return
        pos = self._pos(w)
        if not pos: return
        r, c = pos
        if c < self.factors_count:
            if self._sel_cells and self._sel_handle_cell() is w and self._near_br(w):
                self._start_fill(use_sel=True); return "break"
            if not self._sel_cells and self._near_br(w):
                self._clear_sel()
                self._sel_anchor = (r, c); self._sel_cells = {(r, c)}
                self._sel_orig[(r, c)] = w.cget("bg"); self._apply_sel({(r, c)})
                self._start_fill(use_sel=False); return "break"
        self._fill_drag = False
        self._clear_sel()
        self._sel_anchor = (r, c); self._sel_cells = {(r, c)}
        self._sel_orig[(r, c)] = w.cget("bg"); self._apply_sel({(r, c)})
        w.focus_set()

    def _on_shift_click(self, event):
        w = event.widget; pos = self._pos(w)
        if not pos or self._sel_anchor is None: return
        ar, ac = self._sel_anchor; r, c = pos
        self._sel_range(ar, ac, r, c); return "break"

    def _on_drag(self, event):
        w = event.widget
        if not isinstance(w, tk.Entry): return
        if self._fill_drag: self._do_fill(event); return "break"
        if self._sel_anchor is None: return
        ar, ac = self._sel_anchor
        pos = self._pos(w)
        if pos:
            r, c = pos
        else:
            py = w.winfo_pointery(); px = w.winfo_pointerx(); r, c = ar, ac
            for ri in range(len(self.entries)):
                for ci in range(len(self.entries[ri])):
                    cell = self.entries[ri][ci]
                    if (cell.winfo_rootx() <= px <= cell.winfo_rootx() + cell.winfo_width() and
                            cell.winfo_rooty() <= py <= cell.winfo_rooty() + cell.winfo_height()):
                        r, c = ri, ci; break
        self._sel_range(ar, ac, r, c)

    def _on_release(self, event):
        if self._fill_drag:
            self._fill_drag = False; self._fill_rows = []; self._fill_cols = []; return "break"

    # ── fill drag ─────────────────────────────────────────────
    def _start_fill(self, use_sel):
        self._fill_drag = True
        if use_sel and self._sel_cells:
            b = self._sel_bounds()
            if b is None: self._fill_drag = False; return
            self._fill_rows = list(range(b[0], b[2] + 1))
            self._fill_cols = list(range(b[1], b[3] + 1))
        elif self._sel_anchor:
            self._fill_rows = [self._sel_anchor[0]]
            self._fill_cols = [self._sel_anchor[1]]
        else:
            self._fill_drag = False

    def _do_fill(self, event):
        w = event.widget
        if not isinstance(w, tk.Entry) or not self._fill_rows or not self._fill_cols: return
        last_src = self._fill_rows[-1]
        py = w.winfo_pointery(); target = last_src
        for rr in range(last_src, len(self.entries)):
            cell = self.entries[rr][self._fill_cols[0]]
            y0 = cell.winfo_rooty()
            if y0 <= py <= y0 + cell.winfo_height(): target = rr; break
        else:
            if py > self.entries[-1][self._fill_cols[0]].winfo_rooty(): target = len(self.entries)
        if target <= last_src: return
        n_src = len(self._fill_rows)
        dst = last_src + 1
        while dst <= target:
            while dst >= len(self.entries): self.add_row()
            src_r = self._fill_rows[(dst - last_src - 1) % n_src]
            for c in self._fill_cols:
                if c >= self.factors_count: break
                self.entries[dst][c].delete(0, tk.END)
                self.entries[dst][c].insert(0, self.entries[src_r][c].get())
            dst += 1
        self.inner.update_idletasks()
        self.canvas.config(scrollregion=self.canvas.bbox("all"))

    # ── copy selection Ctrl+C ─────────────────────────────────
    def _on_copy(self, event=None):
        if not self._sel_cells:
            w = event.widget if event else self.table_win.focus_get()
            if isinstance(w, tk.Entry):
                try: sel = w.get("sel.first", "sel.last")
                except Exception: sel = w.get()
                self.table_win.clipboard_clear(); self.table_win.clipboard_append(sel)
            return "break"
        b = self._sel_bounds()
        if b is None: return "break"
        r1, c1, r2, c2 = b
        lines = []
        for r in range(r1, r2 + 1):
            row = []
            for c in range(c1, c2 + 1):
                try: row.append(self.entries[r][c].get())
                except Exception: row.append("")
            lines.append("\t".join(row))
        self.table_win.clipboard_clear(); self.table_win.clipboard_append("\n".join(lines))
        return "break"

    # ── navigation ────────────────────────────────────────────
    def _on_enter(self, event=None):
        pos = self._pos(event.widget)
        if not pos: return "break"
        i, j = pos; ni = i + 1
        if ni >= len(self.entries): self.add_row()
        self.entries[ni][j].focus_set(); self.entries[ni][j].icursor(tk.END); return "break"

    def _on_arrow(self, event=None):
        pos = self._pos(event.widget)
        if not pos: return "break"
        i, j = pos
        if event.keysym == "Up":    i = max(0, i - 1)
        elif event.keysym == "Down":  i = min(len(self.entries) - 1, i + 1)
        elif event.keysym == "Left":  j = max(0, j - 1)
        elif event.keysym == "Right": j = min(len(self.entries[i]) - 1, j + 1)
        self.entries[i][j].focus_set(); self.entries[i][j].icursor(tk.END); return "break"

    def _on_paste(self, event=None):
        widget = event.widget if event else self.table_win.focus_get()
        if not isinstance(widget, tk.Entry): return "break"
        try: data = self.table_win.clipboard_get()
        except Exception: return "break"
        pos = self._pos(widget)
        if not pos: return "break"
        r0, c0 = pos
        for ir, rt in enumerate([r for r in data.splitlines() if r != ""]):
            for jc, val in enumerate(rt.split("\t")):
                rr = r0 + ir; cc = c0 + jc
                while rr >= len(self.entries): self.add_row()
                if cc >= self.cols: continue
                self.entries[rr][cc].delete(0, tk.END); self.entries[rr][cc].insert(0, val)
        for j in range(self.factors_count):
            _autofit_col(self.entries, j, self.header_labels)
        return "break"

    def _pos(self, widget):
        for i, row in enumerate(self.entries):
            for j, cell in enumerate(row):
                if cell is widget: return i, j
        return None

    def _mk_entry(self, parent):
        return tk.Entry(parent, width=COL_W, fg="#000000", font=("Times New Roman", 12),
                        highlightthickness=1, highlightbackground="#c0c0c0", highlightcolor="#c0c0c0")

    # ── rename factor ─────────────────────────────────────────
    def rename_factor(self, col):
        if col < 0 or col >= self.factors_count: return
        fk = self.factor_keys[col]; old = self.ftitle(fk)
        dlg = tk.Toplevel(self.table_win or self.root)
        dlg.title("Перейменувати фактор"); dlg.resizable(False, False)
        set_icon(dlg); dlg.grab_set()
        tk.Label(dlg, text=f"Назва фактору {fk}:",
                 font=("Times New Roman",12)).pack(padx=16, pady=14)
        var = tk.StringVar(value=old)
        e = tk.Entry(dlg, textvariable=var, font=("Times New Roman",12), width=28)
        e.pack(padx=16, pady=4); e.select_range(0, tk.END); e.focus_set()
        def ok():
            new = var.get().strip()
            if not new: return
            self._set_ftitle(fk, new)
            if col < len(self.header_labels):
                self.header_labels[col].configure(text=new)
            dlg.destroy()
        tk.Button(dlg, text="OK", bg="#c62828", fg="white",
                  font=("Times New Roman",12), command=ok).pack(pady=(4,14))
        dlg.bind("<Return>", lambda ev: ok())
        center_win(dlg)
        dlg.update_idletasks(); center_win(dlg); dlg.bind("<Return>", lambda ev: ok()); dlg.grab_set()

    def show_design_help(self):
        w = tk.Toplevel(self.root); w.title("Пояснення дизайнів"); w.resizable(False, False); set_icon(w)
        frm = tk.Frame(w, padx=16, pady=14); frm.pack()
        txt = ("CRD — Повна рандомізація\nRCBD — Блочна рандомізація\nSplit-plot — Спліт-плот\n\n"
               "SS Тип I: послідовний (порядок важливий)\n"
               "SS Тип II: ієрархічний (без взаємодій)\n"
               "SS Тип III: частковий (з урахуванням всіх ефектів) — рекомендований")
        t = tk.Text(frm, width=55, height=10, wrap="word"); t.insert("1.0", txt)
        t.configure(state="disabled"); t.pack()
        tk.Button(frm, text="OK", width=10, command=w.destroy).pack(pady=(10, 0))
        w.update_idletasks(); center_win(w); w.grab_set()

    def show_about(self):
        messagebox.showinfo("Розробник",
            f"S.A.D. — Статистичний аналіз даних  v{APP_VER}\n"
            "Розробник: Чаплоуцький Андрій Миколайович\n"
            "Уманський національний університет")

    # ══════════════════════════════════════════════════════════
    # OPEN TABLE  — with menu bar
    # ══════════════════════════════════════════════════════════
    def open_table(self, fc):
        if self.table_win and tk.Toplevel.winfo_exists(self.table_win):
            self.table_win.destroy()
        self.factors_count = fc
        self.factor_keys   = ["A", "B", "C", "D"][:fc]
        for fk in self.factor_keys:
            if fk not in self.factor_title_map: self._set_ftitle(fk, f"Фактор {fk}")

        self.table_win = tw = tk.Toplevel(self.root)
        factor_names = {1:"Однофакторний", 2:"Двофакторний",
                        3:"Трифакторний",  4:"Чотирифакторний"}
        # Для 3-факторного — додаємо опис ЛК у довідку
        tw.title(f"S.A.D. — {factor_names.get(fc,str(fc)+'-факторний')} дисперсійний аналіз")
        tw.geometry("1280x720"); set_icon(tw)

        # ── Toolbar ───────────────────────────────────────────
        ctl = tk.Frame(tw, padx=6, pady=4); ctl.pack(fill=tk.X)

        tk.Button(ctl, text="▶ Аналіз", bg="#c62828", fg="white",
                  font=("Times New Roman", 13),
                  command=self.analyze).pack(side=tk.LEFT, padx=4)

        # Налаштування — спадне меню
        mb2 = tk.Menubutton(ctl, text="⚙ Налаштування ▾",
                            font=("Times New Roman", 11),
                            relief=tk.RAISED, bd=2)
        mb2.pack(side=tk.LEFT, padx=4)
        sm2 = tk.Menu(mb2, tearoff=0)
        sm2.add_command(label="Додати рядок",      command=self.add_row)
        sm2.add_command(label="Видалити рядок",    command=self.delete_row)
        sm2.add_separator()
        sm2.add_command(label="Додати стовпчик",   command=self.add_column)
        sm2.add_command(label="Видалити стовпчик", command=self.delete_column)
        sm2.add_separator()
        sm2.add_command(label="💾 Зберегти проект", command=self.save_project)
        sm2.add_command(label="📂 Відкрити проект", command=self.load_project)
        sm2.add_separator()
        sm2.add_command(label="🗑 Очистити таблицю", command=self.clear_project)
        mb2["menu"] = sm2

        tk.Button(ctl, text="Вставити з буфера",
                  font=("Times New Roman", 11),
                  command=self._paste_from_focus).pack(side=tk.LEFT, padx=4)
        tk.Button(ctl, text="📂 Відкрити показник", bg="#1a6b8c", fg="white",
                  font=("Times New Roman", 11),
                  command=lambda: self._open_indicator_from_journal(fc)
                  ).pack(side=tk.LEFT, padx=4)
        tk.Button(ctl, text="📚 Довідка",
                  bg="#1a4b8c", fg="white",
                  font=("Times New Roman", 11),
                  command=lambda: self._show_anova_help(tw, fc)).pack(side=tk.LEFT, padx=4)

        # Підказка
        tk.Label(ctl,
                 text="Подвійний клік на синьому заголовку фактора → перейменувати",
                 font=("Times New Roman",9), fg="#666").pack(side=tk.LEFT, padx=8)

        # ── Table canvas ──────────────────────────────────────
        self.canvas = tk.Canvas(tw)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb = ttk.Scrollbar(tw, orient="vertical", command=self.canvas.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y); self.canvas.configure(yscrollcommand=sb.set)
        self.inner = tk.Frame(self.canvas)
        self.canvas.create_window((0, 0), window=self.inner, anchor="nw")
        # mousewheel scroll
        def _mw(e): self.canvas.yview_scroll(int(-1 * (e.delta / 120)), "units")
        self.canvas.bind_all("<MouseWheel>", _mw)

        self.rows = 12; self.cols = fc + 6
        self.entries = []; self.header_labels = []
        col_names = [self.ftitle(fk) for fk in self.factor_keys] + [f"Повт.{i+1}" for i in range(6)]
        for j, nm in enumerate(col_names):
            is_factor = j < fc
            lbl = tk.Label(self.inner, text=nm, relief=tk.RIDGE, width=COL_W,
                           bg="#1a4b8c" if is_factor else "#2e6b2e",
                           fg="white",
                           cursor="hand2" if is_factor else "arrow",
                           font=("Times New Roman", 12, "bold"))
            lbl.grid(row=0, column=j, padx=2, pady=2, sticky="nsew")
            self.header_labels.append(lbl)
            if is_factor:
                lbl.bind("<Double-Button-1>", lambda e, c=j: self.rename_factor(c))
        for i in range(self.rows):
            row_e = []
            for j in range(self.cols):
                e = self._mk_entry(self.inner); e.grid(row=i + 1, column=j, padx=2, pady=2)
                self.bind_cell(e)
                if j < fc:
                    e.bind("<KeyRelease>",
                           lambda ev, jj=j: _autofit_col(self.entries, jj, self.header_labels))
                row_e.append(e)
            self.entries.append(row_e)
        self.inner.update_idletasks(); self.canvas.config(scrollregion=self.canvas.bbox("all"))
        self.entries[0][0].focus_set()
        tw.bind("<Control-v>", self._on_paste); tw.bind("<Control-V>", self._on_paste)

    # ── table editing ─────────────────────────────────────────
    def _open_indicator_from_journal(self, fc):
        """Завантажує показник із збереженого польового журналу і одразу
        заповнює ПОТОЧНУ ANOVA-таблицю (той самий формат, що й ручне
        заповнення: рядок на комбінацію рівнів факторів, стовпці Повт.N)."""
        result = open_indicator_for_anova(self.root)
        if result is None: return
        factor_cols, rows, record_name = result
        if len(factor_cols) != fc:
            if not messagebox.askyesno("Кількість факторів не збігається",
                    f"У журналі — {len(factor_cols)} фактор(и), а ця таблиця "
                    f"{fc}-факторна. Дані все одно можна вставити (зайві "
                    f"стовпчики факторів залишаться порожні, або значення "
                    f"обріжуться) — продовжити?"):
                return

        by_combo = {}
        for r in rows:
            combo_key = tuple(r[fc_] for fc_ in factor_cols[:fc])
            by_combo.setdefault(combo_key, {})[r["replication"]] = r[record_name]
        combos_sorted = sorted(by_combo.keys())
        all_reps = sorted({rep for v in by_combo.values() for rep in v.keys()})
        n_rep_cols = max(len(all_reps), 1)

        for i, fname in enumerate(factor_cols[:fc]):
            key = self.factor_keys[i]
            self._set_ftitle(key, fname)
            self.header_labels[i].configure(text=self.ftitle(key))

        while len(self.entries) < len(combos_sorted): self.add_row()
        while (self.cols - fc) < n_rep_cols: self.add_column()

        for i, combo in enumerate(combos_sorted):
            for k, lvl_name in enumerate(combo):
                self.entries[i][k].delete(0, tk.END)
                self.entries[i][k].insert(0, str(lvl_name))
            rep_vals = by_combo[combo]
            for ci, rep_num in enumerate(all_reps):
                col = fc + ci
                if col >= self.cols: continue
                v = rep_vals.get(rep_num)
                self.entries[i][col].delete(0, tk.END)
                if v is not None: self.entries[i][col].insert(0, str(v))

        messagebox.showinfo("Дані перенесено",
            f"Дані показника «{record_name}» перенесено в таблицю "
            f"({len(combos_sorted)} комбінацій × {n_rep_cols} повторень). "
            "Перевірте таблицю і натисніть «▶ Аналіз».")

    def add_row(self):
        i = len(self.entries); row_e = []
        for j in range(self.cols):
            e = self._mk_entry(self.inner); e.grid(row=i + 1, column=j, padx=2, pady=2)
            self.bind_cell(e)
            if j < self.factors_count:
                e.bind("<KeyRelease>",
                       lambda ev, jj=j: _autofit_col(self.entries, jj, self.header_labels))
            row_e.append(e)
        self.entries.append(row_e); self.rows += 1
        self.inner.update_idletasks(); self.canvas.config(scrollregion=self.canvas.bbox("all"))

    def delete_row(self):
        if not self.entries: return
        for e in self.entries.pop(): e.destroy()
        self.rows -= 1; self.inner.update_idletasks()
        self.canvas.config(scrollregion=self.canvas.bbox("all"))

    def add_column(self):
        self.cols += 1; ci = self.cols - 1
        nm = f"Повт.{ci - self.factors_count + 1}"
        lbl = tk.Label(self.inner, text=nm, relief=tk.RIDGE, width=COL_W,
                       bg="#f0f0f0", fg="#000000", font=("Times New Roman", 12, "bold"))
        lbl.grid(row=0, column=ci, padx=2, pady=2, sticky="nsew"); self.header_labels.append(lbl)
        for i, row in enumerate(self.entries):
            e = self._mk_entry(self.inner); e.grid(row=i + 1, column=ci, padx=2, pady=2)
            self.bind_cell(e); row.append(e)
        self.inner.update_idletasks(); self.canvas.config(scrollregion=self.canvas.bbox("all"))

    def delete_column(self):
        if self.cols <= self.factors_count + 1: return
        self.header_labels.pop().destroy()
        for row in self.entries: row.pop().destroy()
        self.cols -= 1; self.inner.update_idletasks()
        self.canvas.config(scrollregion=self.canvas.bbox("all"))

    def _paste_from_focus(self):
        w = self.table_win.focus_get()
        if isinstance(w, tk.Entry):
            class _E: widget = w
            self._on_paste(_E())

    def _used_rep(self):
        rep_cols = []
        for c in range(self.factors_count, self.cols):
            for r in range(len(self.entries)):
                s = self.entries[r][c].get().strip()
                if not s: continue
                try: float(s.replace(",", ".")); rep_cols.append(c); break
                except Exception: continue
        return rep_cols

    def collect_long(self, design):
        long = []; rep = self._used_rep()
        if not rep: return long, rep
        for i, row in enumerate(self.entries):
            lvls = [row[k].get().strip() or f"рядок{i+1}" for k in range(self.factors_count)]
            for ic, c in enumerate(rep):
                s = row[c].get().strip()
                if not s: continue
                try: val = float(s.replace(",", "."))
                except Exception: continue
                rec = {"value": val}
                for ki, fk in enumerate(self.factor_keys): rec[fk] = lvls[ki]
                if design in ("rcbd", "split"):
                    rec["BLOCK"] = f"Блок {ic+1}"
                elif design == "latin":
                    # Для латинського квадрату: перший фактор = варіант,
                    # другий фактор = рядок (ROW), третій = стовпець (COL)
                    # якщо factors_count >= 3 — беремо з таблиці
                    if self.factors_count >= 3:
                        rec["ROW"] = lvls[1] if len(lvls) > 1 else f"Ряд {i+1}"
                        rec["COL"] = lvls[2] if len(lvls) > 2 else f"Стовп {ic+1}"
                    else:
                        # factors_count == 1: автогенерація ROW/COL з номерів
                        rec["ROW"] = f"Ряд {i+1}"
                        rec["COL"] = f"Стовп {ic+1}"
                long.append(rec)
        return long, rep

    # ── dialogs ───────────────────────────────────────────────
    def ask_params(self):
        parent = self.table_win or self.root
        dlg = tk.Toplevel(parent); dlg.title("Параметри аналізу")
        dlg.resizable(False, False); set_icon(dlg)
        frm = tk.Frame(dlg, padx=16, pady=14); frm.pack()
        rf = ("Times New Roman", 12)
        rb_f = ("Times New Roman", 13)

        tk.Label(frm, text="Назва показника (необов'язково):",
                 font=rf).grid(row=0, column=0, sticky="w", pady=4)
        e_ind = tk.Entry(frm, width=32, font=rf)
        e_ind.grid(row=0, column=1, pady=4, padx=6)
        tk.Label(frm, text="Одиниці виміру (необов'язково):",
                 font=rf).grid(row=1, column=0, sticky="w", pady=4)
        e_un = tk.Entry(frm, width=32, font=rf)
        e_un.grid(row=1, column=1, pady=4, padx=6)

        # ── Дизайн ───────────────────────────────────────────
        tk.Label(frm, text="Дизайн досліду:", font=("Times New Roman",12,"bold")
                 ).grid(row=2, column=0, columnspan=2, sticky="w", pady=14)

        design_info = tk.Frame(frm); design_info.grid(row=3, column=0, columnspan=2, sticky="w")
        tk.Label(design_info,
                 text=(
                     "CRD — Повністю рандомізований:\n"
                     "  Варіанти розміщені випадково по всіх ділянках.\n"
                     "  Використовується на однорідному фоні.\n\n"
                     "RCBD — Рандомізовані повні блоки:\n"
                     "  Ділянки розбиті на блоки (повторності) однорідних умов.\n"
                     "  Блок = одна повторність. Всередині блоку — випадкове розміщення.\n"
                     "  Рекомендується при неоднорідності фону.\n\n"
                     "Split-plot — Розщеплені ділянки:\n"
                     "  Головний фактор (whole-plot) = великі ділянки,\n"
                     "  Другорядний фактор (sub-plot) = дрібніші, всередині великих.\n"
                     "  Типово: фактор A = обробка всього поля, B = сорт на підділянці."
                 ),
                 font=("Times New Roman",10), justify="left",
                 bg="#f0f4ff", relief=tk.FLAT, padx=8, pady=6
                 ).pack(fill=tk.X)

        dv = tk.StringVar(value="crd")
        df = tk.Frame(frm); df.grid(row=4, column=0, columnspan=2, sticky="w", pady=8)
        for txt, val in [("CRD", "crd"), ("RCBD", "rcbd"),
                          ("Латинський квадрат", "latin"),
                          ("Split-plot (лише параметричний)", "split")]:
            tk.Radiobutton(df, text=txt, variable=dv, value=val,
                           font=rb_f).pack(side=tk.LEFT, padx=8)

        # Підказка для латинського квадрату
        latin_hint = tk.Label(frm,
            text="ℹ ЛАТИНСЬКИЙ КВАДРАТ — інструкція:\n"
                 "  1. Відкрийте 3-ФАКТОРНИЙ аналіз (кнопка «3 фактори» на головній).\n"
                 "  2. Перейменуйте фактори (подвійний клік):\n"
                 "     Фактор A = Варіант  (назви обробок або сортів)\n"
                 "     Фактор B = Рядок    (1, 2, 3... — ряди поля)\n"
                 "     Фактор C = Стовпець (1, 2, 3... — стовпці поля)\n"
                 "  3. Кожен рядок таблиці = одна ділянка (k²  рядків).\n"
                 "     Значення вводьте у колонку «Повт.1».\n"
                 "  4. Вимога: k варіантів = k рядів = k стовпців (k ≥ 3).\n"
                 "  Модель: Y = μ + Варіант + Рядок + Стовпець + ε\n"
                 "  df(похибка) = (k-1)(k-2)",
            font=("Times New Roman",10), fg="#1a4b8c",
            bg="#eef4ff", relief=tk.FLAT, padx=8, pady=6, justify="left")
        latin_hint.grid(row=5, column=0, columnspan=2, sticky="ew", pady=(0,4))
        latin_hint.grid_remove()

        # Тепер реєструємо trace (latin_hint вже існує)
        def _on_design(*_):
            v = dv.get()
            if v == "latin": latin_hint.grid()
            else:            latin_hint.grid_remove()
            if v == "split": sp_frm.grid()
            else:            sp_frm.grid_remove()
        dv.trace_add("write", _on_design)

        mfv = tk.StringVar(value=self.factor_keys[0] if self.factor_keys else "A")
        sp_frm = tk.Frame(frm); sp_frm.grid(row=5, column=0, columnspan=2, sticky="w", pady=(0,4))
        tk.Label(sp_frm, text="Головний фактор (whole-plot):", font=rf).pack(side=tk.LEFT)
        ttk.Combobox(sp_frm, textvariable=mfv, width=6, state="readonly",
                     values=self.factor_keys).pack(side=tk.LEFT, padx=6)
        sp_frm.grid_remove()


        # ── Тип SS ───────────────────────────────────────────
        ttk.Separator(frm, orient="horizontal").grid(
            row=6, column=0, columnspan=2, sticky="ew", pady=8)
        tk.Label(frm, text="Тип суми квадратів (SS):", font=("Times New Roman",12,"bold")
                 ).grid(row=7, column=0, columnspan=2, sticky="w", pady=0)
        tk.Label(frm,
                 text=(
                     "Тип I — Послідовний: кожен фактор після попередніх.\n"
                     "  Порядок факторів важливий. Тільки для збалансованих дизайнів.\n\n"
                     "Тип II — Ієрархічний: кожен фактор після решти головних ефектів\n"
                     "  (без взаємодій). Для незбалансованих даних без взаємодій.\n\n"
                     "Тип III — Частковий ← РЕКОМЕНДУЄТЬСЯ: кожен ефект при всіх інших.\n"
                     "  Стандарт SPSS/SAS. Не залежить від порядку. Взаємодії враховані.\n\n"
                     "Тип IV — Для сильно незбалансованих дизайнів з порожніми клітинками."
                 ),
                 font=("Times New Roman",10), justify="left",
                 bg="#f0f4ff", relief=tk.FLAT, padx=8, pady=4
                 ).grid(row=8, column=0, columnspan=2, sticky="ew")
        ssv = tk.StringVar(value="III")
        ssf = tk.Frame(frm); ssf.grid(row=9, column=0, columnspan=2, sticky="w", pady=4)
        for ss in ["I","II","III","IV"]:
            tk.Radiobutton(ssf, text=f"Тип {ss}", variable=ssv, value=ss,
                           font=rb_f).pack(side=tk.LEFT, padx=8)

        # ── Кнопки ───────────────────────────────────────────
        out = {"ok": False}
        def ok():
            out.update({"ok": True,
                        "indicator": e_ind.get().strip() or "Показник",
                        "units":     e_un.get().strip()  or "–",
                        "design":    dv.get(),
                        "split_main": mfv.get(),
                        "ss_type":   ssv.get()})
            dlg.destroy()

        bf = tk.Frame(frm); bf.grid(row=10, column=0, columnspan=2, pady=(14,0))
        tk.Button(bf, text="▶ Виконати аналіз", width=18,
                  bg="#c62828", fg="white", font=rf, command=ok).pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Скасувати", width=12,
                  font=rf, command=dlg.destroy).pack(side=tk.LEFT)

        dlg.update_idletasks(); center_win(dlg)
        e_ind.focus_set()
        dlg.bind("<Return>", lambda e: ok())
        dlg.grab_set(); parent.wait_window(dlg)
        return out

    def choose_method(self, p_norm, design, n_var):
        parent = self.table_win or self.root
        dlg = tk.Toplevel(parent); dlg.title("Вибір методу")
        dlg.resizable(False, False); set_icon(dlg)
        frm = tk.Frame(dlg, padx=16, pady=14); frm.pack()
        normal = (p_norm is not None) and (not math.isnan(p_norm)) and (p_norm > 0.05)
        rb_f = ("Times New Roman", 13)
        ordinal = getattr(self, '_ordinal_mode', False)

        if ordinal:
            # ── Режим бальних даних: ЛИШЕ непараметричні ────
            tk.Label(frm,
                     text="⚠ Бальна шкала → ЛИШЕ непараметричні методи",
                     fg="#c62828", font=("Times New Roman",12,"bold")
                     ).pack(anchor="w", pady=0)
            tk.Label(frm,
                     text=(
                         "Параметричні методи (НІР, Тьюкі, Дункан) заблоковані.\n"
                         "Причина: бальні дані є порядковими — середнє і\n"
                         "стандартне відхилення методично некоректні для них.\n"
                         "Оберіть непараметричний критерій:"
                     ),
                     fg="#555", font=("Times New Roman",11), justify="left"
                     ).pack(anchor="w", pady=(0,10))
            if design == "crd":
                options = [
                    ("Краскел-Уолліс + Манн-Уітні (Бонферроні)  ← рекомендовано", "kw"),
                ]
            elif n_var == 2:
                options = [("Вілкоксон (парний)  ← рекомендовано для 2 варіантів", "wilcoxon")]
            else:
                options = [("Фрідман + Вілкоксон (Бонферроні)  ← рекомендовано", "friedman")]

        elif normal:
            tk.Label(frm, text="✓ Дані відповідають нормальному розподілу (Shapiro–Wilk).",
                     justify="left", fg="#1a6b1a",
                     font=("Times New Roman",11)).pack(anchor="w", pady=0)
            options = [("НІР₀₅ (LSD)", "lsd"),
                       ("Тест Тьюкі", "tukey"),
                       ("Тест Дункана", "duncan"),
                       ("Бонферроні", "bonferroni")]
        else:
            if design == "split":
                tk.Label(frm,
                         text="Split-plot: лише параметричний.\n"
                              "Залишки не нормальні → аналіз некоректний.\n"
                              "Рекомендація: трансформуйте або оберіть CRD/RCBD.",
                         fg="#c62828", justify="left",
                         font=("Times New Roman",11)).pack(anchor="w")
                options = []
            else:
                tk.Label(frm,
                         text="⚠ Дані НЕ відповідають нормальному розподілу.\n"
                              "Оберіть метод:",
                         fg="#c62828", justify="left",
                         font=("Times New Roman",11)).pack(anchor="w", pady=(0, 8))
                if design == "crd":
                    options = [
                        ("Краскела–Уолліса", "kw"),
                        ("Манна-Уітні", "mw"),
                        ("🔁 arcsin(√p) + параметричний", "arcsin_param"),
                        ("🔁 ln(x) + параметричний",      "log_param"),
                        ("🔁 √x + параметричний",          "sqrt_param"),
                        ("🔁 log₁₀(x) + параметричний",   "log10_param"),
                    ]
                else:
                    if n_var == 2:
                        options = [
                            ("Wilcoxon (парний)", "wilcoxon"),
                            ("🔁 arcsin(√p) + параметричний", "arcsin_param"),
                            ("🔁 ln(x) + параметричний",      "log_param"),
                            ("🔁 √x + параметричний",          "sqrt_param"),
                            ("🔁 log₁₀(x) + параметричний",   "log10_param"),
                        ]
                    else:
                        options = [
                            ("Friedman", "friedman"),
                            ("🔁 arcsin(√p) + параметричний", "arcsin_param"),
                            ("🔁 ln(x) + параметричний",      "log_param"),
                            ("🔁 √x + параметричний",          "sqrt_param"),
                            ("🔁 log₁₀(x) + параметричний",   "log10_param"),
                        ]
        out = {"ok": False, "method": None}
        if not options:
            tk.Button(frm, text="OK", width=10, command=dlg.destroy).pack(pady=(10, 0))
            dlg.update_idletasks(); center_win(dlg); dlg.grab_set(); parent.wait_window(dlg); return out

        var = tk.StringVar(value=options[0][1])
        for txt, val in options:
            tk.Radiobutton(frm, text=txt, variable=var, value=val, font=rb_f).pack(anchor="w", pady=2)

        # ── Пояснення: яка трансформація для якої ситуації ─────
        transform_keys = {"arcsin_param", "log_param", "sqrt_param", "log10_param"}
        if any(val in transform_keys for _, val in options):
            help_frm = tk.Frame(frm, bg="#eef3f8", padx=10, pady=8)
            help_frm.pack(fill=tk.X, pady=(10, 0))
            tk.Label(help_frm, text="ⓘ Яку трансформацію обрати:",
                     bg="#eef3f8", fg="#1a4b8c",
                     font=("Times New Roman", 11, "bold"), anchor="w"
                     ).pack(fill=tk.X)
            transform_help = [
                ("arcsin(√p)", "відсотки, частки, частоти (% ураження, схожості, "
                                "загибелі) — особливо якщо є значення поза 20–80%"),
                ("ln(x)",      "дані з мультиплікативною мінливістю: розкид зростає "
                                "пропорційно значенню (маса, урожайність, сильна "
                                "правостороння скошеність)"),
                ("√x",         "лічильні дані: кількість комах, шкідників, колосків, "
                                "уражених рослин, дефектів (розподіл, близький до "
                                "Пуассонівського)"),
                ("log₁₀(x)",  "те саме, що ln(x), але коли зручніше читати результат "
                                "у порядках величини (напр., концентрації, титри)"),
            ]
            for name, desc in transform_help:
                row = tk.Frame(help_frm, bg="#eef3f8"); row.pack(fill=tk.X, pady=1)
                tk.Label(row, text=name, bg="#eef3f8", fg="#1a4b8c",
                         font=("Times New Roman", 10, "bold"), width=11, anchor="w"
                         ).pack(side=tk.LEFT)
                tk.Label(row, text="— " + desc, bg="#eef3f8", fg="#333",
                         font=("Times New Roman", 10), anchor="w",
                         justify="left", wraplength=340
                         ).pack(side=tk.LEFT, fill=tk.X)
            tk.Label(help_frm,
                     text="Після трансформації середні у звіті повертаються у "
                          "вихідні одиниці; групування (літери) — за трансформованими даними.",
                     bg="#eef3f8", fg="#666", font=("Times New Roman", 9, "italic"),
                     anchor="w", justify="left", wraplength=380
                     ).pack(fill=tk.X, pady=(6, 0))

        def ok():
            out.update({"ok": True, "method": var.get()}); dlg.destroy()
        bf = tk.Frame(frm); bf.pack(pady=(12, 0))
        tk.Button(bf, text="▶ Виконати", width=14, bg="#c62828", fg="white",
                  font=("Times New Roman",12), command=ok).pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Скасувати", width=12,
                  font=("Times New Roman",12), command=dlg.destroy).pack(side=tk.LEFT)
        dlg.update_idletasks(); center_win(dlg); dlg.bind("<Return>", lambda e: ok())
        dlg.grab_set(); parent.wait_window(dlg); return out

    # ════════════════════════════════════════════════════════════
    # ANALYZE
    # ════════════════════════════════════════════════════════════
    def analyze(self):
        try:
            self._analyze_inner()
        except Exception as _ae:
            import traceback
            messagebox.showerror("Помилка аналізу",
                f"Виникла помилка при виконанні аналізу:\n\n"
                f"{str(_ae)}\n\n"
                f"Деталі:\n{traceback.format_exc()[-600:]}")

    def _analyze_inner(self):
        created = datetime.now()
        params = self.ask_params()
        if not params["ok"]: return
        indicator = params["indicator"]; units = params["units"]
        design = params["design"]; split_main = params["split_main"]
        ss_type = params.get("ss_type", "III")

        long, used_rep = self.collect_long(design)
        if not long: messagebox.showwarning("Помилка даних", "Немає числових даних."); return
        values = np.array([r["value"] for r in long], dtype=float)

        # ── Автоматичне виявлення відсоткових даних (arcsin) ───
        is_percent_units = any(u in units.lower() for u in
                               ["%", "відсот", "percent", "частк", "зав'яз",
                                "схожість", "схожіст", "ураж", "виживан",
                                "товарн", "вихід"])
        arcsin_applied = False
        arcsin_was_suggested = False

        if is_percent_units:
            vmin = float(np.min(values)); vmax = float(np.max(values))
            # Перевіряємо чи дані у діапазоні 0-100 (відсотки) або 0-1 (частки)
            in_pct_range  = (vmin >= 0 and vmax <= 100 and vmax > 1)
            in_frac_range = (vmin >= 0 and vmax <= 1)
            is_fraction   = in_pct_range or in_frac_range
            needs_arcsin  = is_fraction and (vmin < 20 or vmax > 80)

            if is_fraction and not needs_arcsin:
                # Всі значення між 20% і 80% — підказка але не наполягаємо
                messagebox.showinfo("Інформація про дані",
                    "Одиниці вимірювання: " + units + "\n"
                    "Діапазон значень: " + fmt(vmin,1) + "% – " + fmt(vmax,1) + "%\n\n"
                    "Всі значення у діапазоні 20–80% \u2192 arcsin трансформація\n"
                    "не є критично необхідною, але дозволена.\n\n"
                    "Аналіз буде виконано без трансформації.\n"
                    "Якщо Shapiro–Wilk покаже ненормальність \u2014\n"
                    "оберіть arcsin(\u221ap) у вікні вибору методу.")

            elif is_fraction and needs_arcsin:
                arcsin_was_suggested = True
                ans = messagebox.askyesno("Виявлено відсоткові дані",
                    "Одиниці вимірювання: " + units + "\n"
                    "Діапазон значень: " + fmt(vmin,1) + "% – " + fmt(vmax,1) + "%\n\n"
                    "Значення поза діапазоном 20–80%\n"
                    "виявляють нерівномірність дисперсій.\n\n"
                    "РЕКОМЕНДОВАНА трансформація: arcsin(\u221ap)\n"
                    "Формула: y = arcsin(\u221a(p/100))\n\n"
                    "Що робить:\n"
                    "\u2022 Вирівнює дисперсії між варіантами\n"
                    "\u2022 Наближає розподіл до нормального\n"
                    "\u2022 Дозволяє коректно застосувати ANOVA\n\n"
                    "Середні у звіті будуть у вихідних %.\n\n"
                    "Застосувати arcsin(\u221ap) трансформацію?")

                if ans:
                    # Застосовуємо arcsin трансформацію
                    def _arcsin_transform(v):
                        p = v / 100.0 if v > 1 else v  # конвертуємо у частки
                        p = max(0.0, min(1.0, p))       # обмежуємо [0,1]
                        return math.asin(math.sqrt(p))

                    long = [dict(r, value=_arcsin_transform(r["value"])) for r in long]
                    values = np.array([r["value"] for r in long], dtype=float)
                    arcsin_applied = True

                    messagebox.showinfo("Трансформацію застосовано",
                        "arcsin(\u221ap) трансформацію застосовано.\n"
                        "Аналіз виконується на трансформованих даних.\n"
                        "У звіті середні наведено у вихідних відсотках.\n"
                        "Літери значущості (НІР) — за трансформованими даними.")

        # ── Автоматичне виявлення бальних (порядкових) даних ──
        is_ordinal_units = any(u in units.lower() for u in
                               ["бал", "score", "rank", "rang", "ранг",
                                "очко", "пункт", "клас", "ступін", "ступен"])
        ordinal_detected = False
        ordinal_forced   = False   # чи примусово переведено на непараметричний

        if is_ordinal_units and not arcsin_applied:
            vmin_o = float(np.min(values)); vmax_o = float(np.max(values))
            # Додаткова перевірка: всі значення цілі і діапазон ≤ 20
            all_int = np.all(values == np.floor(values))
            small_range = (vmax_o - vmin_o) <= 20

            if all_int and small_range:
                ordinal_detected = True
                messagebox.showwarning(
                    "⚠ Виявлено бальну (порядкову) шкалу",
                    f"Одиниці вимірювання: «{units}»\n"
                    f"Діапазон значень: {int(vmin_o)} – {int(vmax_o)} балів\n\n"
                    "МЕТОДИЧНА ВИМОГА:\n"
                    "Бальні шкали є ПОРЯДКОВИМИ (ordinal) даними.\n"
                    "Це означає що:\n"
                    "  • Різниця між балами нерівномірна\n"
                    "  • Розподіл зазвичай ненормальний\n"
                    "  • Параметрична ANOVA методично НЕКОРЕКТНА\n\n"
                    "Програма автоматично застосує:\n"
                    "  ✓ Непараметричний аналіз (Краскел-Уолліс / Фрідман)\n"
                    "  ✓ Медіана [Q1; Q3] замість Mean ± SD у звіті\n"
                    "  ✓ Boxplot або Dot plot для візуалізації\n\n"
                    "Параметричні методи заблоковано для цих даних.\n"
                    "Натисніть OK щоб продовжити з правильним методом.")
                ordinal_forced = True

        # Зберігаємо прапор для choose_method і show_report
        self._ordinal_mode  = ordinal_detected and ordinal_forced
        self._ordinal_units = units

        # ── Мінімальна кількість спостережень ──────────────────
        if len(values) < 6:
            messagebox.showwarning("Замало даних",
                f"Для дисперсійного аналізу потрібно щонайменше 6 спостережень.\n"
                f"Наразі: {len(values)}."); return

        lbf = {f: first_seen([r.get(f) for r in long]) for f in self.factor_keys}

        # ── Перевірка мінімальної кількості рівнів кожного фактора ──
        for f in self.factor_keys:
            if len(lbf[f]) < 2:
                messagebox.showwarning("Помилка даних",
                    f"Фактор «{self.ftitle(f)}» має лише 1 рівень.\n"
                    "Для аналізу потрібно щонайменше 2 рівні."); return

        # ── Перевірка мінімальної кількості повторностей для RCBD/Split ──
        if design in ("rcbd", "split") and len(used_rep) < 2:
            messagebox.showwarning("Помилка дизайну",
                f"Для дизайну {design.upper()} потрібно щонайменше 2 повторності (блоки).\n"
                f"Наразі: {len(used_rep)}."); return

        # ── Перевірка Латинського квадрату ─────────────────────
        if design == "latin":
            if self.factors_count < 3:
                messagebox.showerror("Латинський квадрат — помилка структури",
                    "Для Латинського квадрату відкрийте 3-факторний аналіз:\n\n"
                    "  Фактор A = Варіант (назви варіантів)\n"
                    "  Фактор B = Рядок (номери рядів: 1, 2, 3...)\n"
                    "  Фактор C = Стовпець (номери стовпців: 1, 2, 3...)\n\n"
                    "Кожен рядок таблиці = одна ділянка, одне значення у колонці «Повт.1»"); return
            # Перевіряємо k = n_рядків = n_стовпців
            fk = self.factor_keys
            k_var  = len(lbf.get(fk[0], []))
            k_rows = len(lbf.get(fk[1], [])) if len(fk) > 1 else 0
            k_cols = len(lbf.get(fk[2], [])) if len(fk) > 2 else 0
            if not (k_var == k_rows == k_cols):
                messagebox.showerror("Латинський квадрат — помилка k×k",
                    f"Латинський квадрат вимагає k варіантів = k рядків = k стовпців.\n\n"
                    f"Знайдено:\n"
                    f"  Варіантів (Фактор A): {k_var}\n"
                    f"  Рядків (Фактор B):    {k_rows}\n"
                    f"  Стовпців (Фактор C):  {k_cols}\n\n"
                    f"Перевірте що всі три значення однакові."); return
            if k_var < 3:
                messagebox.showwarning("Замало варіантів",
                    f"Латинський квадрат вимагає щонайменше 3 варіанти (k≥3).\n"
                    f"Знайдено: {k_var}. Рекомендується k=4-6."); return
            if k_var > 8:
                ans = messagebox.askyesno("Попередження: великий квадрат",
                    f"k={k_var} варіантів → {k_var}×{k_var} = {k_var**2} ділянок.\n"
                    f"При k>8 Латинський квадрат стає громіздким.\n"
                    f"Рекомендується RCBD для великої кількості варіантів.\n\n"
                    "Продовжити?")
                if not ans: return

        var_order = first_seen([tuple(r.get(f) for f in self.factor_keys) for r in long])
        v_names = [" | ".join(map(str, k)) for k in var_order]
        n_var = len(var_order)

        # ── Перевірка збалансованості при Тип I SS ─────────────
        if ss_type == "I":
            from collections import Counter
            cell_counts = Counter(tuple(r.get(f) for f in self.factor_keys) for r in long)
            counts = list(cell_counts.values())
            if len(set(counts)) > 1:
                ans = messagebox.askyesno("Увага: незбалансовані дані + Тип I SS",
                    "Дані незбалансовані (різна кількість спостережень у клітинках).\n\n"
                    "При Тип I SS (послідовний) результат залежить від ПОРЯДКУ введення факторів.\n"
                    "Зміна порядку дає інші значення SS та p.\n\n"
                    "Рекомендація: використовуйте Тип III для незбалансованих даних.\n\n"
                    "Продовжити з Тип I SS?")
                if not ans: return

        try:
            if design == "crd":
                res = anova_crd(long, self.factor_keys, lbf, ss_type)
            elif design == "rcbd":
                res = anova_rcbd(long, self.factor_keys, lbf, ss_type=ss_type)
            elif design == "latin":
                # Для ЛК перший ключ = варіант, другий = рядок, третій = стовпець
                fk_var = [self.factor_keys[0]]
                res = anova_latin_square(long, fk_var, lbf, ss_type)
            else:
                if split_main not in self.factor_keys: split_main = self.factor_keys[0]
                res = anova_split(long, self.factor_keys, split_main, ss_type=ss_type)
        except Exception as ex:
            import traceback
            messagebox.showerror("Помилка моделі",
                str(ex) + "\n\nДетальніше:\n" + traceback.format_exc()[-500:])
            return

        residuals = np.array(res.get("residuals", []), dtype=float)
        n_res = len(residuals)
        self._last_residuals = residuals.tolist()  # для вкладки залишків
        try: W, p_norm = shapiro(residuals) if n_res >= 3 else (np.nan, np.nan)
        except Exception: W, p_norm = np.nan, np.nan

        # ── Попередження про обмеження Shapiro-Wilk ────────────
        sw_warning = ""
        if n_res < 8:
            sw_warning = f"\n⚠ Увага: n={n_res} — надто мало для надійного тесту нормальності."
        elif n_res > 100:
            sw_warning = (f"\n⚠ Увага: n={n_res} — при великих вибірках Shapiro–Wilk виявляє\n"
                          "  навіть мінімальні відхилення як значущі. Оцінюйте разом з Q-Q графіком.")

        normal = (not math.isnan(p_norm)) and (p_norm > 0.05)
        if design == "split" and not normal:
            messagebox.showwarning("Split-plot: аналіз неможливий",
                f"Залишки моделі не відповідають нормальному розподілу\n"
                f"(Shapiro–Wilk: W={fmt(W,4)}, p={fmt(p_norm,4)}).\n\n"
                "Split-plot реалізований лише для параметричних методів.\n"
                "Рекомендації:\n"
                "• трансформуйте дані (логарифмування) і повторіть;\n"
                "• або оберіть CRD/RCBD для непараметричного аналізу."
                + sw_warning); return

        choice = self.choose_method(p_norm, design, n_var)
        if not choice or not choice.get("ok"): return
        method = choice["method"]

        log_applied = False
        transform_label = ""
        if method == "arcsin_param" and not arcsin_applied:
            # Користувач обрав arcsin вручну у choose_method
            # (не було автоматичного застосування)
            vmin_ = float(np.min(values)); vmax_ = float(np.max(values))
            if vmin_ < 0 or vmax_ > 100:
                messagebox.showwarning("Помилка трансформації",
                    f"arcsin(√p) застосовна лише для даних у діапазоні 0–100% або 0–1.\n"
                    f"Ваші дані: мін={fmt(vmin_,2)}, макс={fmt(vmax_,2)}.\n\n"
                    "Оберіть інший метод трансформації."); return
            def _asin(v):
                p = v/100.0 if v > 1 else v
                return math.asin(math.sqrt(max(0., min(1., p))))
            long = [dict(r, value=_asin(r["value"])) for r in long]
            values = np.array([r["value"] for r in long], dtype=float)
            arcsin_applied = True
            transform_label = "arcsin(√p)"
            log_applied = True
            # Перераховуємо модель на трансформованих даних
            try:
                if design == "crd":      res = anova_crd(long, self.factor_keys, lbf, ss_type)
                elif design == "rcbd":   res = anova_rcbd(long, self.factor_keys, lbf, ss_type=ss_type)
                elif design == "latin":  res = anova_latin_square(long, self.factor_keys, lbf, ss_type)
                else:                    res = anova_split(long, self.factor_keys, split_main, ss_type=ss_type)
            except Exception as ex: messagebox.showerror("Помилка моделі", str(ex)); return
            residuals = np.array(res.get("residuals",[]), dtype=float)
            try: W, p_norm = shapiro(residuals) if len(residuals)>=3 else (np.nan,np.nan)
            except Exception: W, p_norm = np.nan, np.nan
            # Вибираємо пост-хок метод після arcsin
            messagebox.showinfo("arcsin(√p) застосовано",
                f"Трансформацію arcsin(√p) застосовано.\n"
                f"Shapiro–Wilk після трансформації: W={fmt(W,4)}, p={fmt(p_norm,4)}\n"
                f"{'✓ розподіл нормальний' if not math.isnan(p_norm) and p_norm>0.05 else '⚠ розподіл все ще ненормальний — розгляньте непараметричний метод'}\n\n"
                "Оберіть метод пост-хок порівнянь:")
            choice2 = self.choose_method(p_norm, design, n_var)
            if not choice2["ok"]: return
            method = choice2["method"]
            if method == "arcsin_param": method = "lsd"

        elif method in ("log_param", "sqrt_param", "log10_param"):
            # ── Validate data for chosen transformation ──────────
            if method in ("log_param", "log10_param") and np.any(values <= 0):
                messagebox.showwarning("Трансформація неможлива",
                    "Дані містять нулі або від'ємні значення.\n"
                    "Логарифмування неможливе.\n"
                    "Оберіть √x або непараметричний метод."); return
            if method == "sqrt_param" and np.any(values < 0):
                messagebox.showwarning("Трансформація неможлива",
                    "Дані містять від'ємні значення.\n"
                    "Трансформація √x неможлива."); return

            # ── Apply transformation ─────────────────────────────
            if method == "log_param":
                long = [dict(r, value=math.log(r["value"])) for r in long]
                transform_label = "ln(x)"
            elif method == "sqrt_param":
                long = [dict(r, value=math.sqrt(r["value"])) for r in long]
                transform_label = "√x"
            elif method == "log10_param":
                long = [dict(r, value=math.log10(r["value"])) for r in long]
                transform_label = "log₁₀(x)"

            values = np.array([r["value"] for r in long], dtype=float)
            log_applied = True

            # ── Re-run model on transformed data ─────────────────
            try:
                if design == "crd":      res = anova_crd(long, self.factor_keys, lbf, ss_type)
                elif design == "rcbd":   res = anova_rcbd(long, self.factor_keys, lbf, ss_type=ss_type)
                elif design == "latin":  res = anova_latin_square(long, self.factor_keys, lbf, ss_type)
                else:                    res = anova_split(long, self.factor_keys, split_main, ss_type=ss_type)
            except Exception as ex: messagebox.showerror("Помилка моделі", str(ex)); return

            residuals = np.array(res.get("residuals", []), dtype=float)
            try: W, p_norm = shapiro(residuals) if len(residuals) >= 3 else (np.nan, np.nan)
            except Exception: W, p_norm = np.nan, np.nan

            if math.isnan(p_norm) or p_norm <= 0.05:
                messagebox.showwarning("Трансформація не допомогла",
                    f"Застосовано трансформацію {transform_label}, але залишки\n"
                    f"все одно не відповідають нормальному розподілу\n"
                    f"(Shapiro–Wilk: W={fmt(W,4)}, p={fmt(p_norm,4)}).\n\n"
                    "Параметричний аналіз неможливий.\n"
                    "Оберіть непараметричний метод (Kruskal–Wallis, Mann–Whitney тощо)."); return

            messagebox.showinfo("Трансформація успішна",
                f"Застосовано трансформацію {transform_label}.\n"
                f"Shapiro–Wilk після трансформації:\n"
                f"W={fmt(W,4)},  p={fmt(p_norm,4)}  ✓ нормальний розподіл.\n\n"
                "Оберіть метод парних порівнянь:")
            choice2 = self.choose_method(p_norm, design, n_var)
            if not choice2["ok"]: return
            method = choice2["method"]
            if method in ("log_param", "sqrt_param", "log10_param"): method = "lsd"

        MS_err = res.get("MS_error", np.nan); df_err = res.get("df_error", np.nan)
        MS_wp  = res.get("MS_whole", np.nan); df_wp  = res.get("df_whole", np.nan)
        split_mf = res.get("main_factor", split_main) if design == "split" else None

        vs_ = vstats(long, self.factor_keys)
        v_means = {k: vs_[k][0] for k in vs_}; v_sds = {k: vs_[k][1] for k in vs_}; v_ns = {k: vs_[k][2] for k in vs_}
        means1 = {v_names[i]: v_means.get(var_order[i], np.nan) for i in range(n_var)}
        ns1    = {v_names[i]: v_ns.get(var_order[i], 0)         for i in range(n_var)}
        gv = groups_by(long, tuple(self.factor_keys))
        groups1 = {v_names[i]: gv.get(var_order[i], []) for i in range(n_var)}

        fg = {f: {k[0]: v for k, v in groups_by(long, (f,)).items()} for f in self.factor_keys}
        fm = {f: {lv: float(np.mean(arr)) if arr else np.nan for lv, arr in fg[f].items()} for f in self.factor_keys}
        fn = {f: {lv: len(arr) for lv, arr in fg[f].items()} for f in self.factor_keys}
        fsd= {f: {lv: float(np.std(arr, ddof=1)) if len(arr) >= 2 else (0. if len(arr)==1 else np.nan)
                  for lv, arr in fg[f].items()} for f in self.factor_keys}
        fmed = {f: {lv: median_q(arr)[0] for lv, arr in fg[f].items()} for f in self.factor_keys}
        fq   = {f: {lv: median_q(arr)[1:] for lv, arr in fg[f].items()} for f in self.factor_keys}

        vmed = {var_order[i]: median_q(groups1[v_names[i]])[0] for i in range(n_var)}
        vq   = {var_order[i]: median_q(groups1[v_names[i]])[1:] for i in range(n_var)}
        rkv  = mean_ranks(long, lambda r: " | ".join(str(r.get(f)) for f in self.factor_keys))
        rkf  = {f: mean_ranks(long, lambda r, ff=f: r.get(ff)) for f in self.factor_keys}

        lev_F, lev_p = (np.nan, np.nan)
        if method in ("lsd", "tukey", "duncan", "bonferroni"):
            lev_F, lev_p = levene_test(groups1)
            # ── Блокування при неоднорідних дисперсіях ─────────
            if not math.isnan(lev_p) and lev_p < ALPHA:
                ans = messagebox.askyesno(
                    "Неоднорідність дисперсій (тест Левена)",
                    f"Тест Левена виявив неоднорідність дисперсій\n"
                    f"(F={fmt(lev_F,4)}, p={fmt(lev_p,4)}) — умова ANOVA порушена.\n\n"
                    "Параметричний аналіз при неоднорідних дисперсіях дає\n"
                    "недостовірні F-значення та p-значення.\n\n"
                    "Рекомендації:\n"
                    "• застосуйте трансформацію даних (логарифмування);\n"
                    "• або оберіть непараметричний метод (Kruskal–Wallis).\n\n"
                    "Продовжити параметричний аналіз попри порушення умови?")
                if not ans: return

        kw_H = kw_p = kw_df = kw_eps = np.nan; do_ph = True
        fr_chi = fr_p = fr_df = fr_W = np.nan; wil_s = wil_p = np.nan
        rcbd_ph = []; rcbd_sig = {}
        lf = {f: {lv: "" for lv in lbf[f]} for f in self.factor_keys}
        lnamed = {nm: "" for nm in v_names}
        ph_rows = []; fpt = {}

        if method == "lsd":
            # ── Protected LSD (Fisher): перевірити значущість глобального F ──
            # Знаходимо найменше p-значення серед головних ефектів та взаємодій
            global_p_values = []
            for raw_row in res["table"]:
                nm_, _ss, _df, _ms, _F, pv_ = raw_row
                if any(x in str(nm_) for x in ["Залишок", "Загальна", "Блоки", "WP-error"]):
                    continue
                if pv_ is not None and not (isinstance(pv_, float) and math.isnan(pv_)):
                    global_p_values.append(float(pv_))

            global_F_sig = any(p < ALPHA for p in global_p_values) if global_p_values else False

            if not global_F_sig:
                messagebox.showwarning("Protected LSD: пост-хок заблоковано",
                    "Жоден з ефектів у дисперсійному аналізі не є статистично значущим\n"
                    "(p ≥ 0.05 для всіх факторів та їх взаємодій).\n\n"
                    "Відповідно до принципу Protected LSD (Fisher, 1935),\n"
                    "попарні порівняння можна виконувати ЛИШЕ після значущого F-тесту.\n\n"
                    "Висновок: між варіантами немає статистично значущої різниці.\n"
                    "Звіт сформовано з таблицями описової статистики без пост-хок аналізу.")
                # Продовжуємо — формуємо звіт без літер CLD
            else:
                for f in self.factor_keys:
                    MS_ = MS_wp if (design == "split" and f == split_mf) else MS_err
                    df_ = df_wp if (design == "split" and f == split_mf) else df_err
                    lf[f] = cld(lbf[f], fm[f], lsd_sig(lbf[f], fm[f], fn[f], MS_, df_))
                if design != "split":
                    lnamed = cld(v_names, means1, lsd_sig(v_names, means1, ns1, MS_err, df_err))

        elif method in ("tukey", "duncan", "bonferroni"):
            # ── Перевірка значущості глобального F перед пост-хок ──
            global_p_values_2 = []
            for raw_row in res["table"]:
                nm_, _ss, _df, _ms, _F, pv_ = raw_row
                if any(x in str(nm_) for x in ["Залишок", "Загальна", "Блоки", "WP-error"]):
                    continue
                if pv_ is not None and not (isinstance(pv_, float) and math.isnan(pv_)):
                    global_p_values_2.append(float(pv_))
            global_F_sig_2 = any(p < ALPHA for p in global_p_values_2) if global_p_values_2 else False

            if not global_F_sig_2:
                messagebox.showwarning("Пост-хок заблоковано",
                    "Жоден ефект не є статистично значущим (p ≥ 0.05).\n\n"
                    "Виконання пост-хок порівнянь без значущого F-тесту\n"
                    "призводить до надмірного числа хибнопозитивних результатів.\n\n"
                    "Висновок: між варіантами немає статистично значущої різниці.")
            else:
                if design != "split":
                    ph_rows, sig_ = pairwise_param(v_names, means1, ns1, MS_err, df_err, method)
                    lnamed = cld(v_names, means1, sig_)
                    for f in self.factor_keys:
                        r_, s_ = pairwise_param(lbf[f], fm[f], fn[f], MS_err, df_err, method)
                        fpt[f] = r_; lf[f] = cld(lbf[f], fm[f], s_)
                else:
                    for f in self.factor_keys:
                        MS_ = MS_wp if f == split_mf else MS_err
                        df_ = df_wp if f == split_mf else df_err
                        r_, s_ = pairwise_param(lbf[f], fm[f], fn[f], MS_, df_, method)
                        fpt[f] = r_; lf[f] = cld(lbf[f], fm[f], s_)

        elif method == "kw":
            try:
                smp = [groups1[n] for n in v_names if groups1[n]]
                if len(smp) >= 2:
                    kwr = kruskal(*smp); kw_H = float(kwr.statistic); kw_p = float(kwr.pvalue)
                    kw_df = len(smp) - 1; kw_eps = eps2_kw(kw_H, len(long), len(smp))
            except Exception: pass
            if not math.isnan(kw_p) and kw_p >= ALPHA: do_ph = False
            if do_ph:
                ph_rows, sig_ = pairwise_mw(v_names, groups1)
                lnamed = cld(v_names, {n: float(np.median(groups1[n])) if groups1[n] else np.nan for n in v_names}, sig_)

        elif method == "mw":
            ph_rows, sig_ = pairwise_mw(v_names, groups1)
            lnamed = cld(v_names, {n: float(np.median(groups1[n])) if groups1[n] else np.nan for n in v_names}, sig_)

        elif method == "friedman":
            bnames = first_seen([f"Блок {i+1}" for i in range(len(used_rep))])
            long2  = [dict(r, VARIANT=" | ".join(str(r.get(f)) for f in self.factor_keys)) for r in long]
            mat, _ = rcbd_matrix(long2, v_names, bnames)
            if len(mat) < 2: messagebox.showwarning("", "Потрібні ≥ 2 повних блоки."); return
            try:
                fr = friedmanchisquare(*[np.array(c, dtype=float) for c in zip(*mat)])
                fr_chi = float(fr.statistic); fr_p = float(fr.pvalue)
                fr_df = n_var - 1; fr_W = kendalls_w(fr_chi, len(mat), n_var)
            except Exception: pass
            if not math.isnan(fr_p) and fr_p < ALPHA:
                rcbd_ph, rcbd_sig = pairwise_wilcox(v_names, mat)
                lnamed = cld(v_names, {n: float(np.median(groups1[n])) if groups1[n] else np.nan for n in v_names}, rcbd_sig)

        elif method == "wilcoxon":
            if n_var != 2: messagebox.showwarning("", "Wilcoxon лише для 2 варіантів."); return
            bnames = first_seen([f"Блок {i+1}" for i in range(len(used_rep))])
            long2  = [dict(r, VARIANT=" | ".join(str(r.get(f)) for f in self.factor_keys)) for r in long]
            mat, _ = rcbd_matrix(long2, v_names, bnames)
            if len(mat) < 2: messagebox.showwarning("", "Потрібні ≥ 2 блоки."); return
            arr = np.array(mat, dtype=float)
            try:
                st, p = wilcoxon(arr[:, 0], arr[:, 1], zero_method="wilcox", alternative="two-sided", mode="auto")
                wil_s = float(st); wil_p = float(p)
            except Exception: pass
            if not math.isnan(wil_p) and wil_p < ALPHA:
                lnamed = cld(v_names, {n: float(np.median(groups1[n])) if groups1[n] else np.nan for n in v_names},
                             {(v_names[0], v_names[1]): True})

        lv_var = {var_order[i]: lnamed.get(v_names[i], "") for i in range(n_var)}
        SS_tot = res.get("SS_total", np.nan); SS_err = res.get("SS_error", np.nan)
        R2 = (1 - (SS_err / SS_tot)) if not any(math.isnan(x) for x in [SS_tot, SS_err]) and SS_tot > 0 else np.nan

        cv_r = [[self.ftitle(f), fmt(cv_means([fm[f].get(lv, np.nan) for lv in lbf[f]]), 2)]
                for f in self.factor_keys]
        cv_r.append(["Загальний", fmt(cv_vals(values), 2)])

        nonparam = method in ("mw", "kw", "friedman", "wilcoxon")
        # Якщо arcsin застосовано автоматично — відображаємо у звіті
        if arcsin_applied and not transform_label:
            transform_label = "arcsin(√p)"
            log_applied = True

        def _rn(nm):
            if not isinstance(nm, str) or not nm.startswith("Фактор "): return nm
            rest = nm.replace("Фактор ", "")
            parts = rest.split("×")
            return "×".join(self.ftitle(p) if p in self.factor_keys else p for p in parts)

        anova_rows = []
        for raw_row in res["table"]:
            nm, SSv, dfv, MSv, Fv, pv = raw_row
            df_s = str(int(dfv)) if dfv is not None and not (isinstance(dfv, float) and math.isnan(dfv)) else ""
            nm2 = _rn(nm)
            if any(x in nm2 for x in ["Залишок", "WP-error", "Блоки"]) or nm2 == "Загальна":
                anova_rows.append([nm2, fmt(SSv, 3), df_s, fmt(MSv, 3), "", "", ""])
            else:
                mk = sig_mark(pv); concl = f"різниця {mk}" if mk else "–"
                anova_rows.append([nm2, fmt(SSv, 3), df_s, fmt(MSv, 3), fmt(Fv, 3), fmt(pv, 4), concl])

        eff_rows  = [[_rn(r[0]), r[1]] for r in build_eff_rows(res["table"])]
        pe2_rows  = [[_rn(r[0]), r[1], r[2]] for r in build_pe2_rows(res["table"])]

        # cache for graph redraw
        self._lbf_cache = lbf

        self.show_report(
            created=created, indicator=indicator, units=units, design=design,
            arcsin_applied=arcsin_applied,
            ordinal_mode=getattr(self, '_ordinal_mode', False),
            ss_type=ss_type, method=method, log_applied=log_applied,
            transform_label=transform_label,
            n_var=n_var, n_rep=len(used_rep), n_obs=len(long),
            split_mf=split_mf, W=W, p_norm=p_norm,
            lev_F=lev_F, lev_p=lev_p,
            kw_H=kw_H, kw_p=kw_p, kw_df=kw_df, kw_eps=kw_eps, do_ph=do_ph,
            fr_chi=fr_chi, fr_p=fr_p, fr_df=fr_df, fr_W=fr_W,
            wil_s=wil_s, wil_p=wil_p,
            anova_rows=anova_rows, eff_rows=eff_rows, pe2_rows=pe2_rows,
            cv_r=cv_r, R2=R2,
            lf=lf, lv_var=lv_var, lbf=lbf,
            fm=fm, fsd=fsd, fmed=fmed, fq=fq, fn=fn,
            var_order=var_order, v_names=v_names,
            v_means=v_means, v_sds=v_sds, v_ns=v_ns,
            vmed=vmed, vq=vq, rkv=rkv, rkf=rkf,
            groups1=groups1, ph_rows=ph_rows, fpt=fpt,
            rcbd_ph=rcbd_ph, nonparam=nonparam, res=res,
        )
        self.show_graphs(long, lf, indicator, units, eff_rows, pe2_rows,
                         parent_win=self.report_win if (
                             self.report_win and
                             tk.Toplevel.winfo_exists(self.report_win)) else None)

    # ════════════════════════════════════════════════════════════
    # REPORT WINDOW
    # ════════════════════════════════════════════════════════════
    def show_report(self, **kw):
        if self.report_win and tk.Toplevel.winfo_exists(self.report_win):
            self.report_win.destroy()
        self.report_win = rw = tk.Toplevel(self.table_win or self.root)
        rw.title("Звіт ANOVA — " + kw.get("indicator",""))
        maximize_win(rw)
        set_icon(rw)

        # ── Бокова панель + контент ──────────────────────────
        main_frame = tk.Frame(rw); main_frame.pack(fill=tk.BOTH, expand=True)
        sidebar = tk.Frame(main_frame, width=190, bg="#2c3e50")
        sidebar.pack(side=tk.LEFT, fill=tk.Y); sidebar.pack_propagate(False)
        content = tk.Frame(main_frame); content.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        tk.Label(sidebar, text="ЗВІТ", bg="#2c3e50", fg="#ecf0f1",
                 font=("Times New Roman",12,"bold"), pady=12).pack(fill=tk.X)

        self._active_panel = None
        self._active_rpt_btn = None

        def _show_panel(frame, btn):
            if self._active_panel: self._active_panel.pack_forget()
            if self._active_rpt_btn:
                self._active_rpt_btn.configure(bg="#2c3e50", fg="#bdc3c7")
            frame.pack(fill=tk.BOTH, expand=True)
            self._active_panel = frame
            btn.configure(bg="#c62828", fg="white")
            self._active_rpt_btn = btn

        def _sidebar_btn(text, tooltip):
            fr = tk.Frame(sidebar, bg="#2c3e50"); fr.pack(fill=tk.X)
            b = tk.Button(fr, text=f"  {text}", bg="#2c3e50", fg="#bdc3c7",
                          font=("Times New Roman",11), relief=tk.FLAT,
                          anchor="w", padx=12, pady=6,
                          activebackground="#c62828", activeforeground="white")
            b.pack(fill=tk.X)
            tk.Label(fr, text=f"    {tooltip}", bg="#2c3e50", fg="#7f8c8d",
                     font=("Times New Roman",8), anchor="w").pack(fill=tk.X)
            tk.Frame(sidebar, bg="#3d5166", height=1).pack(fill=tk.X)
            return b

        # ── Панель 1: Текстовий звіт ─────────────────────────
        rpt_frame = tk.Frame(content)

        # Toolbar звіту
        rpt_tb = tk.Frame(rpt_frame, padx=6, pady=4); rpt_tb.pack(fill=tk.X)
        self._report_buf = []
        def copy_all():
            rw.clipboard_clear()
            rw.clipboard_append("\n".join(self._report_buf))
            messagebox.showinfo("","Текстовий звіт скопійовано у буфер.\nВставте у Word через Ctrl+V.")
        tk.Button(rpt_tb, text="📋 Копіювати звіт",
                  font=("Times New Roman",11), command=copy_all).pack(side=tk.LEFT, padx=4)

        # Scrollable body звіту
        outer = tk.Frame(rpt_frame); outer.pack(fill=tk.BOTH, expand=True)
        vsb = ttk.Scrollbar(outer, orient="vertical"); vsb.pack(side=tk.RIGHT, fill=tk.Y)
        hsb = ttk.Scrollbar(outer, orient="horizontal"); hsb.pack(side=tk.BOTTOM, fill=tk.X)
        cv = tk.Canvas(outer, yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        cv.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.config(command=cv.yview); hsb.config(command=cv.xview)
        body = tk.Frame(cv); cv.create_window((0,0), window=body, anchor="nw")
        def _mw(e): cv.yview_scroll(int(-1*(e.delta/120)),"units")
        body.bind("<Configure>",
                  lambda e: (cv.configure(scrollregion=cv.bbox("all"))))
        rw.bind("<MouseWheel>", _mw); cv.bind("<MouseWheel>", _mw)

        # Кнопки бокової панелі
        b_rpt = _sidebar_btn("📄 Текстовий звіт",  "ANOVA таблиця, НІР, висновки")
        b_rpt.configure(command=lambda: _show_panel(rpt_frame, b_rpt))

        buf = self._report_buf
        def _txt(s):
            tk.Label(body, text=s, font=("Times New Roman", 12), fg="#000000",
                     justify="left", anchor="w", wraplength=1100).pack(fill=tk.X, padx=12, pady=1)
            buf.append(s)
        def _head(s):
            tk.Label(body, text=s, font=("Times New Roman", 13, "bold"), fg="#000000",
                     justify="left", anchor="w").pack(fill=tk.X, padx=12, pady=8)
            buf.append("\n" + s)
        def _sep():
            ttk.Separator(body, orient="horizontal").pack(fill=tk.X, padx=12, pady=4)
            buf.append("-" * 80)
        def _table(headers, rows, min_col=90):
            frm, tv_widget = make_tv(body, headers, rows, min_col)
            frm.pack(fill=tk.X, padx=12, pady=(2, 8))
            # bind mousewheel on treeview and its frame so scroll works
            frm.bind("<MouseWheel>", _mw)
            tv_widget.bind("<MouseWheel>", _mw)
            # plain text for clipboard: TAB-separated (Word auto-formats as table)
            buf.append("\t".join(str(h) for h in headers))
            buf.append("\t".join("-" * max(3, len(str(h))) for h in headers))
            for row in rows:
                buf.append("\t".join("" if v is None else str(v) for v in row))
            buf.append("")

        d = kw
        _txt("З В І Т   С Т А Т И С Т И Ч Н О Г О   А Н А Л І З У")
        _txt(f"Показник: {d['indicator']}   |   Одиниці: {d['units']}")
        _txt(f"Дата: {d['created'].strftime('%d.%m.%Y %H:%M')}")
        _sep()
        design_lbl = {"crd": "CRD (повна рандомізація)", "rcbd": "RCBD (блочна рандомізація)", "split": "Split-plot"}[d['design']]
        _txt(f"Дизайн: {design_lbl}   |   Тип SS: {d['ss_type']}   |   Варіантів: {d['n_var']}   |   Повт.: {d['n_rep']}   |   Спостережень: {d['n_obs']}")
        if d['design'] == "split": _txt(f"Головний фактор (WP): {d['split_mf']}")
        if d.get('arcsin_applied'):
            _txt("⚠ Застосовано трансформацію arcsin(√p) для відсоткових даних. "
                 "Середні у звіті наведено у вихідних відсотках. "
                 "Літери значущості (НІР) визначено за трансформованими даними.")
        elif d.get('ordinal_mode'):
            _txt("ℹ БАЛЬНА ШКАЛА: у звіті наведено медіану [Q1; Q3] замість Mean ± SD. "
                 "Параметричні методи заблоковані. "
                 "Для візуалізації використовуйте Boxplot або Dot plot.")
        elif d['log_applied']:
            tl = d.get('transform_label', 'ln(x)')
            _txt(f"⚠ Застосовано трансформацію {tl}. Середні у звіті — у трансформованій шкалі.")
        method_lbl = {"lsd": "НІР₀₅ (LSD)", "tukey": "Тест Тьюкі", "duncan": "Тест Дункана",
                      "bonferroni": "Бонферроні", "kw": "Kruskal–Wallis", "mw": "Mann–Whitney",
                      "friedman": "Friedman", "wilcoxon": "Wilcoxon"}.get(d['method'], "")
        _txt(f"Метод: {method_lbl}")
        _txt("** — p<0.01; * — p<0.05; різні літери → істотна різниця.")
        _sep()
        _txt(f"Shapiro–Wilk (залишки): {norm_txt(d['p_norm'])}   W={fmt(d['W'], 4)}   p={fmt(d['p_norm'], 4)}")

        if d['method'] == "kw" and not math.isnan(d['kw_p']):
            c_ = ("різниця " + sig_mark(d['kw_p'])) if d['kw_p'] < ALPHA else "–"
            _txt(f"Kruskal–Wallis:  H={fmt(d['kw_H'],4)}  df={d['kw_df']}  p={fmt(d['kw_p'],4)}  {c_}   ε²={fmt(d['kw_eps'],4)}")
        if d['method'] == "friedman" and not math.isnan(d['fr_p']):
            c_ = ("різниця " + sig_mark(d['fr_p'])) if d['fr_p'] < ALPHA else "–"
            _txt(f"Friedman:  χ²={fmt(d['fr_chi'],4)}  df={d['fr_df']}  p={fmt(d['fr_p'],4)}  {c_}   Kendall's W={fmt(d['fr_W'],4)}")
        if d['method'] == "wilcoxon" and not math.isnan(d['wil_p']):
            c_ = ("різниця " + sig_mark(d['wil_p'])) if d['wil_p'] < ALPHA else "–"
            _txt(f"Wilcoxon:  W={fmt(d['wil_s'],4)}  p={fmt(d['wil_p'],4)}  {c_}")

        if not d['nonparam']:
            if not math.isnan(d['lev_p']):
                lc = "умова виконується" if d['lev_p'] >= ALPHA else f"умова порушена {sig_mark(d['lev_p'])}"
                _txt(f"Тест Левена (однорідність дисперсій):  F={fmt(d['lev_F'],4)}  p={fmt(d['lev_p'],4)}  {lc}")
            _sep()
            _head("ТАБЛИЦЯ 1. Дисперсійний аналіз (ANOVA)")
            _table(["Джерело", "SS", "df", "MS", "F", "p", "Висновок"], d['anova_rows'])
            _head("ТАБЛИЦЯ 2. Сила впливу факторів (% від SS)")
            _table(["Джерело", "%"], d['eff_rows'])
            _head("ТАБЛИЦЯ 3. Розмір ефекту (partial η²)")
            _table(["Джерело", "partial η²", "Сила ефекту"], d['pe2_rows'])
            _head("ТАБЛИЦЯ 4. Коефіцієнт варіації (CV, %)")
            _table(["Елемент", "CV, %"], d['cv_r'])
            _txt(f"Коефіцієнт детермінації R² = {fmt(d['R2'], 4)}")
            tno = 5
            if d['method'] == "lsd":
                nir_r = [[k, fmt(v, 4)] for k, v in d['res'].get("NIR05", {}).items()]
                if nir_r:
                    _head(f"ТАБЛИЦЯ {tno}. НІР₀₅"); _table(["Елемент", "НІР₀₅"], nir_r); tno += 1

            for f in self.factor_keys:
                _head(f"ТАБЛИЦЯ {tno}. Середні по фактору: {self.ftitle(f)}")
                rows_f = [[str(lv), fmt(d['fm'][f].get(lv, np.nan), 3),
                           fmt(d['fsd'][f].get(lv, np.nan), 3),
                           d['lf'][f].get(lv, "") or "–"]
                          for lv in d['lbf'][f]]
                _table([self.ftitle(f), "Середнє", "± SD", "Літери CLD"], rows_f); tno += 1

            _head(f"ТАБЛИЦЯ {tno}. Середні по варіантах")
            rows_v = [[nm, fmt(d['v_means'].get(d['var_order'][i], np.nan), 3),
                       fmt(d['v_sds'].get(d['var_order'][i], np.nan), 3),
                       d['lv_var'].get(d['var_order'][i], "") or "–"]
                      for i, nm in enumerate(d['v_names'])]
            _table(["Варіант", "Середнє", "± SD", "Літери CLD"], rows_v); tno += 1

            if d['design'] != "split":
                if d['method'] in ("tukey", "duncan", "bonferroni") and d['ph_rows']:
                    _head(f"ТАБЛИЦЯ {tno}. Парні порівняння варіантів")
                    _table(["Пара", "p", "Висновок"], d['ph_rows']); tno += 1
            else:
                for f in self.factor_keys:
                    rr = d['fpt'].get(f, [])
                    if rr:
                        _head(f"ТАБЛИЦЯ {tno}. Парні порівняння: {self.ftitle(f)}")
                        _table(["Пара", "p", "Висновок"], rr); tno += 1
        else:
            tno = 1
            for f in self.factor_keys:
                _head(f"ТАБЛИЦЯ {tno}. Описова (непараметрична): {self.ftitle(f)}")
                rows = [[str(lv), str(d['fn'][f].get(lv, 0)), fmt(d['fmed'][f].get(lv, np.nan), 3),
                         f"{fmt(d['fq'][f].get(lv,(np.nan,np.nan))[0],3)}–{fmt(d['fq'][f].get(lv,(np.nan,np.nan))[1],3)}",
                         fmt(d['rkf'][f].get(lv, np.nan), 2)]
                        for lv in d['lbf'][f]]
                _table([self.ftitle(f), "n", "Медіана", "Q1–Q3", "Сер. ранг"], rows); tno += 1
            _head(f"ТАБЛИЦЯ {tno}. Описова (непараметрична): варіанти")
            rows = [[d['v_names'][i], str(d['v_ns'].get(d['var_order'][i], 0)),
                     fmt(d['vmed'].get(d['var_order'][i], np.nan), 3),
                     f"{fmt(d['vq'].get(d['var_order'][i],(np.nan,np.nan))[0],3)}–{fmt(d['vq'].get(d['var_order'][i],(np.nan,np.nan))[1],3)}",
                     fmt(d['rkv'].get(d['v_names'][i], np.nan), 2)]
                    for i in range(d['n_var'])]
            _table(["Варіант", "n", "Медіана", "Q1–Q3", "Сер. ранг"], rows); tno += 1
            if d['method'] == "kw":
                if not d['do_ph']: _txt("Kruskal–Wallis p ≥ 0.05 → пост-хок не виконувався.")
                elif d['ph_rows']:
                    _head(f"ТАБЛИЦЯ {tno}. Парні порівняння (MWU + Bonferroni, Cliff's δ)")
                    _table(["Пара", "U", "p(Bonf.)", "Висновок", "δ", "Ефект"], d['ph_rows'])
            if d['method'] == "mw" and d['ph_rows']:
                _head(f"ТАБЛИЦЯ {tno}. Парні порівняння (MWU + Bonferroni, Cliff's δ)")
                _table(["Пара", "U", "p(Bonf.)", "Висновок", "δ", "Ефект"], d['ph_rows'])
            if d['method'] == "friedman":
                if not math.isnan(d['fr_p']) and d['fr_p'] >= ALPHA:
                    _txt("Friedman p ≥ 0.05 → пост-хок не виконувався.")
                elif d['rcbd_ph']:
                    _head(f"ТАБЛИЦЯ {tno}. Парні порівняння (Wilcoxon + Bonferroni)")
                    _table(["Пара", "W", "p(Bonf.)", "Висновок", "r"], d['rcbd_ph'])
        _sep()
        _txt(f"Звіт сформовано: {d['created'].strftime('%d.%m.%Y, %H:%M')}")

        # Показуємо текстовий звіт як першу вкладку
        _show_panel(rpt_frame, b_rpt)

        # Зберігаємо sidebar/content/show_panel для show_graphs
        self._rpt_sidebar = sidebar
        self._rpt_content = content
        self._rpt_show_panel = _show_panel
        self._rpt_sidebar_btn = _sidebar_btn

    # ════════════════════════════════════════════════════════════
    # GRAPHICAL REPORT  — sidebar tabs
    # ════════════════════════════════════════════════════════════
    def show_graphs(self, long, letters_factor, indicator, units,
                    eff_rows, pe2_rows, parent_win=None):
        if not HAS_MPL:
            messagebox.showwarning("", "matplotlib недоступний."); return

        # Якщо є вікно звіту — вбудовуємось у його бокову панель
        if (hasattr(self, '_rpt_sidebar') and
                self.report_win and tk.Toplevel.winfo_exists(self.report_win)):
            gw = self.report_win
            sidebar   = self._rpt_sidebar
            content   = self._rpt_content
            _show_panel = self._rpt_show_panel
            _sidebar_btn = self._rpt_sidebar_btn
            # Роздільник між звітом і графіками
            tk.Frame(sidebar, bg="#1a3a4a", height=2).pack(fill=tk.X)
            tk.Label(sidebar, text="ГРАФІКИ", bg="#2c3e50", fg="#ecf0f1",
                     font=("Times New Roman",11,"bold"),
                     pady=8).pack(fill=tk.X)
        else:
            # Відкриваємо окреме вікно
            if self.graph_win and tk.Toplevel.winfo_exists(self.graph_win):
                self.graph_win.destroy()
            self.graph_win = gw = tk.Toplevel(self.table_win or self.root)
            gw.title(f"Графічний звіт — {indicator}")
            maximize_win(gw)
            set_icon(gw)
            main = tk.Frame(gw); main.pack(fill=tk.BOTH, expand=True)
            sidebar = tk.Frame(main, width=190, bg="#2c3e50")
            sidebar.pack(side=tk.LEFT, fill=tk.Y); sidebar.pack_propagate(False)
            content = tk.Frame(main); content.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            tk.Label(sidebar, text="ГРАФІКИ", bg="#2c3e50", fg="#ecf0f1",
                     font=("Times New Roman",11,"bold"), pady=12).pack(fill=tk.X)
            self._active_panel = None; self._active_rpt_btn = None
            def _show_panel(frame, btn):
                if self._active_panel: self._active_panel.pack_forget()
                if self._active_rpt_btn:
                    self._active_rpt_btn.configure(bg="#2c3e50", fg="#bdc3c7")
                frame.pack(fill=tk.BOTH, expand=True)
                self._active_panel = frame
                btn.configure(bg="#c62828", fg="white")
                self._active_rpt_btn = btn
            def _sidebar_btn(text, tooltip):
                fr = tk.Frame(sidebar, bg="#2c3e50"); fr.pack(fill=tk.X)
                b = tk.Button(fr, text=f"  {text}", bg="#2c3e50", fg="#bdc3c7",
                              font=("Times New Roman",11), relief=tk.FLAT,
                              anchor="w", padx=12, pady=6,
                              activebackground="#c62828", activeforeground="white")
                b.pack(fill=tk.X)
                tk.Label(fr, text=f"    {tooltip}", bg="#2c3e50", fg="#7f8c8d",
                         font=("Times New Roman",8), anchor="w").pack(fill=tk.X)
                tk.Frame(sidebar, bg="#3d5166", height=1).pack(fill=tk.X)
                return b

        # Зберігаємо дані для перебудови
        self._g_long = long; self._g_lf = letters_factor
        self._g_ind  = indicator; self._g_units = units
        self._g_eff  = eff_rows; self._g_pe2 = pe2_rows
        self._graph_figs = {}
        if not hasattr(self, '_gs_titles'): self._gs_titles = {}
        self._lbf_cache = {f: list(letters_factor.get(f, {}).keys())
                           for f in self.factor_keys}

        ordinal = getattr(self, '_ordinal_mode', False)

        # Якщо вбудовуємось у вікно звіту — використовуємо його sidebar і content
        if (hasattr(self, '_rpt_sidebar') and
                self.report_win and tk.Toplevel.winfo_exists(self.report_win)):
            sidebar      = self._rpt_sidebar
            content      = self._rpt_content
            _show_panel  = self._rpt_show_panel
            gw = self.report_win
        else:
            # Окреме вікно (запасний варіант)
            if self.graph_win and tk.Toplevel.winfo_exists(self.graph_win):
                self.graph_win.destroy()
            self.graph_win = gw = tk.Toplevel(self.table_win or self.root)
            gw.title(f"Графічний звіт — {indicator}")
            maximize_win(gw)
            set_icon(gw)
            outer_f = tk.Frame(gw); outer_f.pack(fill=tk.BOTH, expand=True)
            sidebar = tk.Frame(outer_f, width=195, bg="#2c3e50")
            sidebar.pack(side=tk.LEFT, fill=tk.Y); sidebar.pack_propagate(False)
            content = tk.Frame(outer_f)
            content.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            tk.Label(sidebar, text="ГРАФІКИ", bg="#2c3e50", fg="#ecf0f1",
                     font=("Times New Roman",11,"bold"), pady=10).pack(fill=tk.X)
            self._active_panel = None; self._active_rpt_btn = None
            def _show_panel(frame, btn):
                if self._active_panel: self._active_panel.pack_forget()
                if self._active_rpt_btn:
                    self._active_rpt_btn.configure(bg="#2c3e50", fg="#bdc3c7")
                frame.pack(fill=tk.BOTH, expand=True)
                self._active_panel = frame
                btn.configure(bg="#c62828", fg="white")
                self._active_rpt_btn = btn

        gs = self.graph_settings

        def _make_btn(lbl, tooltip):
            fr = tk.Frame(sidebar, bg="#2c3e50"); fr.pack(fill=tk.X)
            b = tk.Button(fr, text=f"  {lbl}", bg="#2c3e50", fg="#bdc3c7",
                          font=("Times New Roman",11), relief=tk.FLAT,
                          anchor="w", padx=10, pady=5,
                          activebackground="#c62828", activeforeground="white")
            b.pack(fill=tk.X)
            tk.Label(fr, text=f"    {tooltip}", bg="#2c3e50", fg="#7f8c8d",
                     font=("Times New Roman",8), anchor="w").pack(fill=tk.X)
            tk.Frame(sidebar, bg="#3d5166", height=1).pack(fill=tk.X)
            return b

        if ordinal:
            self._ordinal_graph_var = tk.StringVar(value="Boxplot")
            ord_frame = tk.Frame(content)
            b_box = _make_btn("📦 Boxplot", "Медіана + квартилі")
            b_dot = _make_btn("● Dot plot", "Точки + медіана")
            def _show_ord(val, btn):
                self._ordinal_graph_var.set(val)
                _show_panel(ord_frame, btn)
                self._rebuild_ordinal_graph(ord_frame, long, letters_factor,
                                            indicator, units)
            b_box.configure(command=lambda: _show_ord("Boxplot", b_box))
            b_dot.configure(command=lambda: _show_ord("Dot plot", b_dot))
            self._rebuild_ordinal_graph(ord_frame, long, letters_factor,
                                        indicator, units)
            # НЕ показуємо автоматично — чекаємо кліку користувача
            return

        # ── Звичайні вкладки ────────────────────────────────────
        tab_defs = [
            ("bp",   "📦 Boxplot",           "Розподіл даних"),
            ("bar",  "📊 Середні ± SE",       "Стовпчики з планками"),
            ("int",  "↗ Взаємодія",           "Профіль взаємодії"),
            ("line", "📈 Динаміка рівнів",    "Лінійний по рівнях"),
            ("hist", "〰 Залишки",            "Гістограма + Q-Q"),
            ("vn",   "🎯 Сила впливу",        "% від SS"),
            ("pe",   "💡 Розмір ефекту",      "Partial η²"),
        ]

        frames = {}
        for key, lbl, tooltip in tab_defs:
            f = tk.Frame(content); frames[key] = f
            b = _make_btn(lbl, tooltip)
            b.configure(command=lambda k=key, bt=b: _show_panel(frames[k], bt))

        tk.Label(sidebar, text=f"{indicator}\n{units}",
                 bg="#2c3e50", fg="#95a5a6",
                 font=("Times New Roman",8), justify="center",
                 wraplength=180).pack(side=tk.BOTTOM, pady=6)

        # Будуємо всі графіки (але не показуємо жодного)
        self._build_bp_tab(  frames["bp"],   long, letters_factor,
                             indicator, units, gs, gw)
        self._build_bar_tab( frames["bar"],  long, letters_factor,
                             indicator, units, gs, gw)
        self._build_int_tab( frames["int"],  long, letters_factor,
                             indicator, units, gs, gw)
        self._build_line_tab(frames["line"], long, letters_factor,
                             indicator, units, gs, gw)
        self._build_hist_tab(frames["hist"], long, gs, indicator, units, gw)
        self._build_vn_tab(  frames["vn"],   eff_rows, gs, gw)
        self._build_pe_tab(  frames["pe"],   pe2_rows, gs, gw)
        # НЕ показуємо жодну вкладку автоматично



    # ── Toolbar кожної вкладки з PNG і налаштуваннями ─────────
    def _tab_toolbar(self, frame, fig_key, rebuild_fn=None, settings_fn=None):
        tb = tk.Frame(frame, bg="#f0f0f0", padx=4, pady=4)
        tb.pack(fill=tk.X, side=tk.BOTTOM)
        tk.Button(tb, text="💾 Зберегти PNG",
                  font=("Times New Roman",10),
                  command=lambda: self._save_fig_png(fig_key)
                  ).pack(side=tk.LEFT, padx=4)
        tk.Button(tb, text="📋 Копіювати",
                  font=("Times New Roman",10),
                  command=lambda: self._copy_fig(fig_key)
                  ).pack(side=tk.LEFT, padx=4)
        if settings_fn:
            tk.Button(tb, text="⚙ Налаштування",
                      font=("Times New Roman",10),
                      bg="#1a4b8c", fg="white",
                      command=settings_fn
                      ).pack(side=tk.LEFT, padx=4)
        return tb

    def _save_fig_png(self, key):
        fig = self._graph_figs.get(key)
        if fig is None: messagebox.showwarning("","Графік відсутній."); return
        path = filedialog.asksaveasfilename(
            defaultextension=".png",
            filetypes=[("PNG","*.png"),("SVG","*.svg")],
            title="Зберегти графік")
        if not path: return
        try:
            fig.savefig(path, dpi=150, bbox_inches="tight")
            messagebox.showinfo("Збережено", f"Збережено:\n{path}")
        except Exception as ex:
            messagebox.showerror("Помилка", str(ex))

    def _settings_dialog(self, gw, key, rebuild_fn, extra_params=None):
        """Уніфікований діалог налаштувань графіка з заголовком + специфічні опції."""
        if not hasattr(self, '_gs_titles'): self._gs_titles = {}
        dlg = tk.Toplevel(gw or self.report_win or self.root)
        dlg.title("Налаштування графіка"); dlg.resizable(False, False)
        set_icon(dlg); dlg.grab_set()
        rf = ("Times New Roman",12)
        frm = tk.Frame(dlg, padx=16, pady=12); frm.pack()

        # Заголовок графіка
        tk.Label(frm, text="Заголовок графіка:", font=rf
                 ).grid(row=0, column=0, sticky="w", pady=4)
        cur = self._gs_titles.get(key, "")
        tv = tk.StringVar(value=cur)
        tk.Entry(frm, textvariable=tv, width=36, font=rf
                 ).grid(row=0, column=1, sticky="w", padx=8)

        # Шрифт і розмір
        gs = self.graph_settings
        tk.Label(frm, text="Шрифт:", font=rf
                 ).grid(row=1, column=0, sticky="w", pady=4)
        ff_v = tk.StringVar(value=gs.get("font_family","Times New Roman"))
        ttk.Combobox(frm, textvariable=ff_v,
                     values=["Times New Roman","Arial","Calibri","Georgia"],
                     state="readonly", width=18
                     ).grid(row=1, column=1, sticky="w", padx=8)

        tk.Label(frm, text="Розмір шрифту:", font=rf
                 ).grid(row=2, column=0, sticky="w", pady=4)
        fz_v = tk.IntVar(value=gs.get("font_size",10))
        tk.Spinbox(frm, from_=7, to=18, textvariable=fz_v, width=7
                   ).grid(row=2, column=1, sticky="w", padx=8)

        # Специфічні параметри для цього графіка
        extra_vars = {}
        row_offset = 3
        if extra_params:
            for ri, (lbl, key2, default, wtype, opts) in enumerate(extra_params):
                tk.Label(frm, text=lbl, font=rf
                         ).grid(row=row_offset+ri, column=0, sticky="w", pady=4)
                var = tk.StringVar(value=str(gs.get(key2, default)))
                extra_vars[key2] = var
                if wtype == "combo":
                    ttk.Combobox(frm, textvariable=var, values=opts,
                                 state="readonly", width=16
                                 ).grid(row=row_offset+ri, column=1, sticky="w", padx=8)
                elif wtype == "color":
                    def _pick(v=var):
                        c = colorchooser.askcolor(color=v.get(), parent=dlg)
                        if c and c[1]: v.set(c[1])
                    tk.Button(frm, text="Обрати колір",
                              command=_pick, font=rf
                              ).grid(row=row_offset+ri, column=1, sticky="w", padx=8)
                elif wtype == "check":
                    bv = tk.BooleanVar(value=bool(gs.get(key2, default)))
                    extra_vars[key2] = bv
                    tk.Checkbutton(frm, variable=bv
                                   ).grid(row=row_offset+ri, column=1, sticky="w", padx=8)
                else:
                    tk.Entry(frm, textvariable=var, width=10, font=rf
                             ).grid(row=row_offset+ri, column=1, sticky="w", padx=8)

        def _apply():
            # Зберігаємо заголовок
            self._gs_titles[key] = tv.get().strip()
            # Зберігаємо налаштування шрифту
            self.graph_settings["font_family"] = ff_v.get()
            self.graph_settings["font_size"]   = fz_v.get()
            # Специфічні
            for k2, v2 in extra_vars.items():
                self.graph_settings[k2] = v2.get()
            dlg.destroy()
            if rebuild_fn: rebuild_fn()

        bf = tk.Frame(frm); bf.grid(row=row_offset+len(extra_params or [])+1,
                                    column=0, columnspan=2, pady=(12,0))
        tk.Button(bf, text="✓ Застосувати", bg="#c62828", fg="white",
                  font=rf, command=_apply).pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Скасувати", font=rf,
                  command=dlg.destroy).pack(side=tk.LEFT)
        center_win(dlg)

    # ── TAB 1: Boxplot ─────────────────────────────────────────
    def _build_bp_tab(self, frame, long, lf, indicator, units, gs, gw=None):
        for w in frame.winfo_children(): w.destroy()

        def _rebuild():
            self._build_bp_tab(frame, long, lf, indicator, units,
                               self.graph_settings, gw)
        def _settings():
            self._settings_dialog(gw, "bp", _rebuild, extra_params=[
                ("Колір боксів:", "box_color", "#aed6f1", "color", None),
                ("Колір медіани:", "median_color", "#c62828", "color", None),
                ("Показати сітку:", "show_grid", True, "check", None),
            ])
        self._tab_toolbar(frame, "bp", _rebuild, _settings)
        plot_f = tk.Frame(frame); plot_f.pack(fill=tk.BOTH, expand=True)

        fp = {"fontsize": gs["font_size"], "fontfamily": gs["font_family"]}
        ff = gs["font_family"]; fz = gs["font_size"]
        title = self._gs_titles.get("bp", f"{indicator}, {units}")
        fig = Figure(figsize=(10, 6), dpi=100); ax = fig.add_subplot(111)
        positions=[]; data=[]; xlbls=[]; let_list=[]; fcentres=[]
        x=1.; gap=1.
        for f in self.factor_keys:
            lvls = list(lf.get(f,{}).keys()) or first_seen(
                [r.get(f) for r in long if r.get(f)])
            if not lvls: continue
            sx=x
            for lv in lvls:
                arr=[float(r["value"]) for r in long if r.get(f)==lv
                     and not math.isnan(float(r.get("value",float("nan"))))]
                data.append(arr); positions.append(x)
                xlbls.append(str(lv)); let_list.append((f,lv)); x+=1.
            fcentres.append(((sx+x-1)/2., self.ftitle(f))); x+=gap
        if data:
            bp=ax.boxplot(data, positions=positions, widths=0.6,
                          showfliers=True, patch_artist=True)
            for p in bp["boxes"]:
                p.set(facecolor=gs.get("box_color","#aed6f1"), alpha=0.85)
            for m in bp["medians"]:
                m.set(color=gs.get("median_color","#c62828"), linewidth=2)
            for w in bp["whiskers"]+bp["caps"]:
                w.set(color=gs.get("whisker_color","#555"), linewidth=1.2)
            for fl in bp["fliers"]:
                fl.set(markerfacecolor=gs.get("flier_color","#c62828"),
                       marker="o", markersize=4)
            ax.set_xticks(positions)
            ax.set_xticklabels(xlbls, rotation=30, ha="right",
                               fontfamily=ff, fontsize=max(7,fz-1))
            allv=[v for a in data for v in a]
            if len(allv)>1 and max(allv)>min(allv):
                off=0.04*(max(allv)-min(allv))
                for i,(f_,lv_) in enumerate(let_list):
                    lt=lf.get(f_,{}).get(lv_,"")
                    if lt and data[i]:
                        ax.text(positions[i], max(data[i])+off, lt,
                                ha="center", va="bottom", **fp)
            for cx,fnm in fcentres:
                ax.text(cx,-0.22,fnm,ha="center",va="top",
                        transform=ax.get_xaxis_transform(),**fp)
            fig.subplots_adjust(bottom=0.28,top=0.91,left=0.08,right=0.98)
        ax.set_title(title, **fp); ax.set_ylabel(units, **fp)
        if gs.get("show_grid", True):
            ax.yaxis.grid(True, linestyle="--", alpha=0.35)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        self._graph_figs["bp"]=fig
        embed_figure(fig, plot_f)

    # ── TAB 2: Середні ± SE ────────────────────────────────────
    def _build_bar_tab(self, frame, long, lf, indicator, units, gs, gw=None):
        for w in frame.winfo_children(): w.destroy()

        def _rebuild():
            self._build_bar_tab(frame, long, lf, indicator, units,
                                self.graph_settings, gw)
        def _settings():
            self._settings_dialog(gw, "bar", _rebuild, extra_params=[
                ("Колір стовпців (усі стовпці):", "bar_color", "#4c72b0", "color", None),
                ("Показати сітку:", "show_grid", True, "check", None),
            ])
        self._tab_toolbar(frame, "bar", _rebuild, _settings)
        plot_f = tk.Frame(frame); plot_f.pack(fill=tk.BOTH, expand=True)

        fp = {"fontsize": gs["font_size"], "fontfamily": gs["font_family"]}
        ff = gs["font_family"]; fz = gs["font_size"]
        user_bar_color = gs.get("bar_color", "#4c72b0")
        title = self._gs_titles.get("bar", f"{indicator}, {units}")
        fig = Figure(figsize=(10, 6), dpi=100); ax = fig.add_subplot(111)
        positions=[]; means=[]; ses=[]; xlbls=[]; let_list=[]; fcentres=[]
        bar_colors=[]; x=1.; gap=1.; ci=0
        for f in self.factor_keys:
            lvls = list(lf.get(f,{}).keys()) or first_seen(
                [r.get(f) for r in long if r.get(f)])
            if not lvls: continue
            sx=x
            for lv in lvls:
                arr=[float(r["value"]) for r in long if r.get(f)==lv
                     and not math.isnan(float(r.get("value",float("nan"))))]
                n=len(arr); m=float(np.mean(arr)) if arr else 0.
                se=float(np.std(arr,ddof=1)/math.sqrt(n)) if n>1 else 0.
                means.append(m); ses.append(se); positions.append(x)
                xlbls.append(str(lv)); let_list.append((f,lv))
                bar_colors.append(user_bar_color)
                x+=1.; ci+=1
            fcentres.append(((sx+x-1)/2., self.ftitle(f))); x+=gap
        if means:
            ax.bar(positions, means, yerr=ses, capsize=4, width=0.65,
                   color=bar_colors, edgecolor="white", linewidth=0.8,
                   error_kw={"ecolor":"#444","lw":1.2,"capthick":1.2})
            allv=[m+se for m,se in zip(means,ses)]
            lowv=[m-se for m,se in zip(means,ses)]
            if allv and means:
                off=max(0.02*(max(allv)-min(means)) if len(allv)>1 else 0.3, 0.3)
                for i,(f_,lv_) in enumerate(let_list):
                    lt=lf.get(f_,{}).get(lv_,"")
                    if lt:
                        ax.text(positions[i], means[i]+ses[i]+off, lt,
                                ha="center", va="bottom", **fp)
                # Y від розумного мінімуму — не від нуля
                data_min = min(lowv) if lowv else 0
                data_max = max(allv) if allv else 1
                spread = data_max - data_min if data_max != data_min else 1
                ymin = data_min - spread * 0.12
                ymax = data_max + spread * 0.18
                ax.set_ylim(ymin, ymax)
            ax.set_xticks(positions)
            ax.set_xticklabels(xlbls, rotation=30, ha="right",
                               fontfamily=ff, fontsize=max(7,fz-1))
            for cx,fnm in fcentres:
                ax.text(cx,-0.22,fnm,ha="center",va="top",
                        transform=ax.get_xaxis_transform(),**fp)
            fig.subplots_adjust(bottom=0.28,top=0.91,left=0.08,right=0.98)
        ax.set_title(title, **fp); ax.set_ylabel(units, **fp)
        if gs.get("show_grid", True):
            ax.yaxis.grid(True, linestyle="--", alpha=0.35)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        self._graph_figs["bar"]=fig
        embed_figure(fig, plot_f)

    # ── TAB 3: Взаємодія ───────────────────────────────────────
    def _build_int_tab(self, frame, long, lf, indicator, units, gs, gw=None):
        for w in frame.winfo_children(): w.destroy()

        def _rebuild():
            self._build_int_tab(frame, long, lf, indicator, units,
                                self.graph_settings, gw)
        def _settings():
            self._settings_dialog(gw, "int", _rebuild, extra_params=[
                ("Товщина лінії:", "line_width", 1.8, "entry", None),
                ("Маркер:", "marker_style", "o", "combo",
                 ["o","s","^","D","v","*"]),
                ("Показати сітку:", "show_grid", True, "check", None),
            ])
        self._tab_toolbar(frame, "int", _rebuild, _settings)
        plot_f = tk.Frame(frame); plot_f.pack(fill=tk.BOTH, expand=True)

        fp = {"fontsize": gs["font_size"], "fontfamily": gs["font_family"]}
        colors_list = ["#4c72b0","#dd8452","#55a868","#c44e52",
                       "#8172b2","#937860","#da8bc3","#8c8c8c"]
        title = self._gs_titles.get("int",
            "Графік взаємодії факторів (профіль середніх)")
        fkeys = self.factor_keys
        fig = Figure(figsize=(10, 6), dpi=100); ax = fig.add_subplot(111)
        lw = float(gs.get("line_width", 1.8))
        mk = gs.get("marker_style", "o")
        if len(fkeys) >= 2:
            f1, f2 = fkeys[0], fkeys[1]
            lvls1 = list(lf.get(f1,{}).keys()) or first_seen(
                [r.get(f1) for r in long if r.get(f1)])
            lvls2 = list(lf.get(f2,{}).keys()) or first_seen(
                [r.get(f2) for r in long if r.get(f2)])
            for gi, lv2 in enumerate(lvls2):
                means_=[float(np.mean([r["value"] for r in long
                    if r.get(f1)==lv1 and r.get(f2)==lv2
                    and not math.isnan(r.get("value",float("nan")))] or [float("nan")]))
                    for lv1 in lvls1]
                ax.plot(range(len(lvls1)), means_, marker=mk, label=str(lv2),
                        color=colors_list[gi%len(colors_list)],
                        linewidth=lw, markersize=7)
            ax.set_xticks(range(len(lvls1)))
            ax.set_xticklabels([str(l) for l in lvls1],
                               rotation=20, ha="right", **fp)
            ax.set_xlabel(self.ftitle(f1), **fp)
            ax.legend(title=self.ftitle(f2), fontsize=fp["fontsize"]-1,
                      title_fontsize=fp["fontsize"]-1)
        elif len(fkeys) == 1:
            f1 = fkeys[0]
            lvls1 = list(lf.get(f1,{}).keys()) or first_seen(
                [r.get(f1) for r in long if r.get(f1)])
            means_=[float(np.mean([r["value"] for r in long if r.get(f1)==lv
                and not math.isnan(r.get("value",float("nan")))] or [float("nan")]))
                for lv in lvls1]
            ax.plot(range(len(lvls1)), means_, marker=mk,
                    color=colors_list[0], linewidth=lw, markersize=8)
            ax.set_xticks(range(len(lvls1)))
            ax.set_xticklabels([str(l) for l in lvls1], **fp)
        else:
            ax.text(0.5,0.5,"Потрібно ≥ 2 фактори",
                    ha="center",va="center",transform=ax.transAxes,**fp)
            ax.axis("off")
        ax.set_title(title, **fp); ax.set_ylabel(units, **fp)
        if gs.get("show_grid", True):
            ax.yaxis.grid(True, linestyle="--", alpha=0.35)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        fig.tight_layout()
        self._graph_figs["int"]=fig
        embed_figure(fig, plot_f)

    # ── TAB 4: Динаміка по рівнях ──────────────────────────────
    def _build_line_tab(self, frame, long, lf, indicator, units, gs, gw=None):
        for w in frame.winfo_children(): w.destroy()

        def _rebuild():
            self._build_line_tab(frame, long, lf, indicator, units,
                                 self.graph_settings, gw)
        def _settings():
            self._settings_dialog(gw, "line", _rebuild, extra_params=[
                ("Товщина лінії:", "line_width", 1.8, "entry", None),
                ("Маркер:", "marker_style", "o", "combo",
                 ["o","s","^","D","v","*"]),
                ("Показати сітку:", "show_grid", True, "check", None),
            ])
        self._tab_toolbar(frame, "line", _rebuild, _settings)
        plot_f = tk.Frame(frame); plot_f.pack(fill=tk.BOTH, expand=True)

        fp = {"fontsize": gs["font_size"], "fontfamily": gs["font_family"]}
        colors_list = ["#4c72b0","#dd8452","#55a868","#c44e52",
                       "#8172b2","#937860","#da8bc3","#8c8c8c"]
        title = self._gs_titles.get("line",
            "Середні ± SE по рівнях кожного фактора")
        fkeys = self.factor_keys; n = len(fkeys)
        lw = float(gs.get("line_width", 1.8))
        mk = gs.get("marker_style", "o")
        fig = Figure(figsize=(10, 6), dpi=100)
        if n == 0:
            ax = fig.add_subplot(111)
            ax.text(0.5,0.5,"Немає факторів",ha="center",va="center")
            ax.axis("off")
        else:
            for fi, f in enumerate(fkeys):
                ax = fig.add_subplot(1, n, fi+1)
                lvls = list(lf.get(f,{}).keys()) or first_seen(
                    [r.get(f) for r in long if r.get(f)])
                means_=[]; ses_=[]
                for lv in lvls:
                    arr=[r["value"] for r in long if r.get(f)==lv
                         and not math.isnan(r.get("value",float("nan")))]
                    means_.append(float(np.mean(arr)) if arr else float("nan"))
                    ses_.append(float(np.std(arr,ddof=1)/math.sqrt(len(arr)))
                                if len(arr)>1 else 0.)
                ax.errorbar(range(len(lvls)), means_, yerr=ses_,
                           marker=mk, color=colors_list[fi%len(colors_list)],
                           linewidth=lw, markersize=7, capsize=4, ecolor="#555")
                ax.set_xticks(range(len(lvls)))
                ax.set_xticklabels([str(l) for l in lvls],
                                   rotation=20, ha="right", **fp)
                ax.set_title(self.ftitle(f), **fp)
                ax.set_ylabel(units if fi==0 else "", **fp)
                if gs.get("show_grid", True):
                    ax.yaxis.grid(True, linestyle="--", alpha=0.35)
                ax.spines["top"].set_visible(False)
                ax.spines["right"].set_visible(False)
                for i,lv in enumerate(lvls):
                    lt=lf.get(f,{}).get(lv,"")
                    if lt and not math.isnan(means_[i]):
                        ax.text(i, means_[i]+ses_[i]+0.3, lt,
                                ha="center", va="bottom", **fp)
            fig.suptitle(title, **fp)
            fig.tight_layout()
        self._graph_figs["line"]=fig
        embed_figure(fig, plot_f)

    # ── TAB 5: Залишки ─────────────────────────────────────────
    def _build_hist_tab(self, frame, long, gs, indicator, units, gw=None):
        for w in frame.winfo_children(): w.destroy()

        def _rebuild():
            self._build_hist_tab(frame, long, self.graph_settings,
                                 indicator, units, gw)
        def _settings():
            self._settings_dialog(gw, "hist", _rebuild, extra_params=[
                ("Колір гістограми:", "hist_color", "#4c72b0", "color", None),
                ("Колір точок Q-Q:", "qq_point_color", "#1a4b8c", "color", None),
            ])
        self._tab_toolbar(frame, "hist", _rebuild, _settings)
        plot_f = tk.Frame(frame); plot_f.pack(fill=tk.BOTH, expand=True)

        fp = {"fontsize": gs["font_size"], "fontfamily": gs["font_family"]}
        title = self._gs_titles.get("hist", "Аналіз залишків")
        fig = Figure(figsize=(10, 6), dpi=100)
        residuals = getattr(self, '_last_residuals', None)
        if residuals and len(residuals) > 2:
            res = np.array(residuals)
            ax1 = fig.add_subplot(121)
            counts, bin_edges, _patches = ax1.hist(
                res, bins="auto",
                color=gs.get("hist_color","#4c72b0"),
                edgecolor="white", alpha=0.85)
            mu_ = float(np.mean(res)); sigma_ = float(np.std(res, ddof=1))
            if sigma_ > 0:
                from scipy.stats import norm as _norm
                xs = np.linspace(res.min(), res.max(), 200)
                bin_w = bin_edges[1] - bin_edges[0]
                ax1.plot(xs, _norm.pdf(xs, mu_, sigma_) * len(res) * bin_w,
                         color=gs.get("median_color","#c62828"), lw=1.8,
                         label="Теоретична нормальна крива")
                ax1.legend(fontsize=max(7, gs["font_size"]-2), frameon=False)
            ax1.set_title("Гістограма залишків", **fp)
            ax1.set_xlabel("Залишок", **fp); ax1.set_ylabel("Частота", **fp)
            ax1.yaxis.grid(True, linestyle="--", alpha=0.3)
            ax1.spines["top"].set_visible(False); ax1.spines["right"].set_visible(False)
            from scipy.stats import probplot
            ax2 = fig.add_subplot(122)
            (osm,osr),(slope,intercept,r)=probplot(res, plot=None)
            ax2.plot(osm, osr, "o",
                     color=gs.get("qq_point_color","#1a4b8c"),
                     markersize=5, alpha=0.85, zorder=3)
            ax2.plot([min(osm),max(osm)],
                     [slope*min(osm)+intercept, slope*max(osm)+intercept],
                     color=gs.get("median_color","#c62828"), lw=1.5, zorder=2)
            ax2.set_title(f"Q-Q графік (R²={r**2:.3f})", **fp)
            ax2.set_xlabel("Теоретичні квантилі", **fp)
            ax2.set_ylabel("Вибіркові квантилі", **fp)
            ax2.spines["top"].set_visible(False); ax2.spines["right"].set_visible(False)
        else:
            ax = fig.add_subplot(111)
            ax.text(0.5,0.5,"Залишки недоступні.\nВиконайте аналіз і відкрийте знову.",
                    ha="center",va="center",transform=ax.transAxes,**fp)
            ax.axis("off")
        fig.suptitle(title, **fp); fig.tight_layout()
        self._graph_figs["hist"]=fig
        embed_figure(fig, plot_f)

    # ── TAB 6: Сила впливу ─────────────────────────────────────
    def _build_vn_tab(self, frame, eff_rows, gs, gw=None):
        for w in frame.winfo_children(): w.destroy()

        def _rebuild():
            self._build_vn_tab(frame, self._g_eff, self.graph_settings, gw)
        def _settings():
            self._settings_dialog(gw, "vn", _rebuild, extra_params=[
                ("Колір головних ефектів:", "vn_main_color", "#1a4b8c",
                 "color", None),
                ("Колір взаємодій:", "vn_inter_color", "#c62828",
                 "color", None),
            ])
        self._tab_toolbar(frame, "vn", _rebuild, _settings)
        plot_f = tk.Frame(frame); plot_f.pack(fill=tk.BOTH, expand=True)

        fp = {"fontsize": gs["font_size"], "fontfamily": gs["font_family"]}
        title = self._gs_titles.get("vn", "Сила впливу факторів (% від суми SS)")
        fig = Figure(figsize=(10, 6), dpi=100); ax = fig.add_subplot(111)
        valid = [(str(nm), float(pct)) for nm,pct in eff_rows
                 if pct and not math.isnan(float(pct)) and float(pct)>0]
        if valid:
            valid.sort(key=lambda x: x[1])
            labels_ = [v[0] for v in valid]
            values_ = [v[1] for v in valid]
            main_c  = gs.get("vn_main_color","#1a4b8c")
            inter_c = gs.get("vn_inter_color","#c62828")
            colors_ = [inter_c if "×" in l else main_c for l in labels_]
            bars = ax.barh(range(len(labels_)), values_,
                           color=colors_, edgecolor="white", height=0.6)
            for i,(bar,val) in enumerate(zip(bars,values_)):
                ax.text(val+0.3, i, f"{val:.1f}%", va="center", **fp)
            ax.set_yticks(range(len(labels_)))
            ax.set_yticklabels(labels_, **fp)
            ax.set_xlabel("% від суми SS", **fp)
            ax.xaxis.grid(True, linestyle="--", alpha=0.35)
            from matplotlib.patches import Patch
            ax.legend(handles=[
                Patch(color=main_c,  label="Головний ефект"),
                Patch(color=inter_c, label="Взаємодія"),
            ], fontsize=fp["fontsize"]-1, loc="lower right")
        else:
            ax.text(0.5,0.5,"Немає даних",ha="center",va="center",
                    transform=ax.transAxes,**fp); ax.axis("off")
        ax.set_title(title, **fp)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        fig.tight_layout()
        self._graph_figs["vn"]=fig
        embed_figure(fig, plot_f)

    # ── TAB 7: Розмір ефекту ───────────────────────────────────
    def _build_pe_tab(self, frame, pe2_rows, gs, gw=None):
        for w in frame.winfo_children(): w.destroy()

        def _rebuild():
            self._build_pe_tab(frame, self._g_pe2, self.graph_settings, gw)
        def _settings():
            self._settings_dialog(gw, "pe", _rebuild, extra_params=[
                ("Колір головних ефектів:", "pe_main_color", "#1a6b1a",
                 "color", None),
                ("Колір взаємодій:", "pe_inter_color", "#c62828",
                 "color", None),
            ])
        self._tab_toolbar(frame, "pe", _rebuild, _settings)
        plot_f = tk.Frame(frame); plot_f.pack(fill=tk.BOTH, expand=True)

        fp = {"fontsize": gs["font_size"], "fontfamily": gs["font_family"]}
        title = self._gs_titles.get("pe", "Розмір ефекту (partial η²)")
        fig = Figure(figsize=(10, 6), dpi=100); ax = fig.add_subplot(111)
        valid = [(str(nm), float(pct)) for nm,pct,_ in pe2_rows
                 if pct and not math.isnan(float(pct)) and float(pct)>0]
        if valid:
            valid.sort(key=lambda x: x[1])
            labels_ = [v[0] for v in valid]
            values_ = [v[1] for v in valid]
            main_c  = gs.get("pe_main_color","#1a6b1a")
            inter_c = gs.get("pe_inter_color","#c62828")
            colors_ = [inter_c if "×" in l else main_c for l in labels_]
            ax.barh(range(len(labels_)), values_,
                    color=colors_, edgecolor="white", height=0.6)
            for i,val in enumerate(values_):
                strength = ("дуже слабкий" if val<0.01 else
                            "слабкий" if val<0.06 else
                            "середній" if val<0.14 else "сильний")
                ax.text(val+0.002, i, f"η²={val:.3f} ({strength})",
                        va="center", **fp)
            ax.set_yticks(range(len(labels_)))
            ax.set_yticklabels(labels_, **fp)
            ax.set_xlabel("partial η²", **fp)
            for thresh, lbl, col in [(0.01,"мала","#aaa"),
                                     (0.06,"середня","#888"),
                                     (0.14,"велика","#555")]:
                ax.axvline(thresh, color=col, lw=0.8, linestyle="--")
                ax.text(thresh, len(labels_)-0.5, lbl,
                        color=col, fontsize=max(7,fp["fontsize"]-2),
                        ha="center", va="bottom")
            ax.xaxis.grid(True, linestyle="--", alpha=0.25)
        else:
            ax.text(0.5,0.5,"Немає даних",ha="center",va="center",
                    transform=ax.transAxes,**fp); ax.axis("off")
        ax.set_title(title, **fp)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        fig.tight_layout()
        self._graph_figs["pe"]=fig
        embed_figure(fig, plot_f)

    def _tab_settings(self, *args, **kwargs): pass  # замінено _settings_dialog

    def _open_gs(self, gw, long, lf, ind, units, eff, pe2):
        dlg = GraphSettingsDlg(gw, self.graph_settings)
        gw.wait_window(dlg)
        if dlg.result: self.graph_settings = dlg.result


        plot_f = tk.Frame(frame); plot_f.pack(fill=tk.BOTH, expand=True)
        self._tab_toolbar(frame, "bp",
            on_settings=lambda: self._tab_settings(
                frame, "bp", indicator, units,
                rebuild=lambda: self._build_bp_tab(
                    frame, long, lf, indicator, units,
                    self.graph_settings)))
        fp = {"fontsize": gs["font_size"], "fontfamily": gs["font_family"]}
        ff = gs["font_family"]; fz = gs["font_size"]
        fig = Figure(figsize=(10, 6), dpi=100); ax = fig.add_subplot(111)
        positions=[]; data=[]; xlbls=[]; let_list=[]; fcentres=[]
        x=1.; gap=1.
        for f in self.factor_keys:
            lvls = list(lf.get(f,{}).keys()) or first_seen(
                [r.get(f) for r in long if r.get(f)])
            if not lvls: continue
            sx=x
            for lv in lvls:
                arr=[float(r["value"]) for r in long if r.get(f)==lv
                     and not math.isnan(float(r.get("value",float("nan"))))]
                data.append(arr); positions.append(x)
                xlbls.append(str(lv)); let_list.append((f,lv)); x+=1.
            fcentres.append(((sx+x-1)/2., self.ftitle(f))); x+=gap
        title = getattr(self, '_gs_titles', {}).get("bp", f"{indicator}, {units}")
        if data:
            bp=ax.boxplot(data,positions=positions,widths=0.6,showfliers=True,patch_artist=True)
            for p in bp["boxes"]: p.set(facecolor=gs["box_color"])
            for m in bp["medians"]: m.set(color=gs["median_color"],linewidth=2)
            for w in bp["whiskers"]+bp["caps"]: w.set(color=gs["whisker_color"])
            for fl in bp["fliers"]: fl.set(markerfacecolor=gs["flier_color"],marker="o",markersize=4)
            ax.set_xticks(positions)
            ax.set_xticklabels(xlbls,rotation=90,fontfamily=ff,fontsize=max(7,fz-1))
            allv=[v for a in data for v in a]
            if len(allv)>1:
                off=max(0.01*(max(allv)-min(allv)),0.3)
                for i,(f_,lv_) in enumerate(let_list):
                    lt=(lf.get(f_,{})).get(lv_,"")
                    if lt and data[i]:
                        ax.text(positions[i],max(data[i])+off,lt,
                                ha="center",va="bottom",**fp)
            for cx,fnm in fcentres:
                ax.text(cx,-0.22,fnm,ha="center",va="top",
                        transform=ax.get_xaxis_transform(),**fp)
            fig.subplots_adjust(bottom=0.32,top=0.91,left=0.08,right=0.98)
        ax.set_title(title,**fp); ax.set_ylabel(units,**fp)
        ax.yaxis.grid(True,linestyle="--",alpha=0.35)
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        self._graph_figs["bp"]=fig
        embed_figure(fig, plot_f)
        fig.set_size_inches(plot_f.winfo_width()/100 or 10,
                            plot_f.winfo_height()/100 or 5)

    # ── TAB 2: Середні ± SE ────────────────────────────────────
    def _open_gs(self, gw, long, lf, ind, units, eff, pe2):
        dlg = GraphSettingsDlg(gw, self.graph_settings)
        gw.wait_window(dlg)
        if dlg.result:
            self.graph_settings = dlg.result



    def _copy_fig(self, key):
        fig = self._graph_figs.get(key)
        if fig is None: messagebox.showwarning("","Графік відсутній."); return
        ok, msg = _copy_fig_to_clipboard(fig)
        if ok: messagebox.showinfo("","Графік скопійовано (PNG).\nВставте у Word через Ctrl+V.")
        else:  messagebox.showwarning("",f"Помилка: {msg}")



# ═══════════════════════════════════════════════════════════════
# EXPORT  — Word (.docx) and PDF via reportlab
# ═══════════════════════════════════════════════════════════════
def export_report_docx(text_lines, tables, filepath):
    """Export plain-text + table data to .docx using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("Встановіть python-docx:\n  pip install python-docx")
    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Times New Roman'; style.font.size = Pt(12)
    for section in doc.sections:
        section.top_margin = Cm(2); section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2)
    for item in text_lines:
        kind = item.get("kind", "text")
        if kind == "heading":
            p = doc.add_heading(item["text"], level=item.get("level", 2))
            p.runs[0].font.name = 'Times New Roman'
        elif kind == "table":
            headers = item["headers"]; rows = item["rows"]
            tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
            tbl.style = 'Table Grid'
            hrow = tbl.rows[0]
            for j, h in enumerate(headers):
                hrow.cells[j].text = str(h)
                run = hrow.cells[j].paragraphs[0].runs[0]
                run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(11)
            for i, row in enumerate(rows):
                for j, val in enumerate(row):
                    cell = tbl.rows[i+1].cells[j]
                    cell.text = "" if val is None else str(val)
                    cell.paragraphs[0].runs[0].font.name = 'Times New Roman' if cell.paragraphs[0].runs else None
            doc.add_paragraph()
        else:
            p = doc.add_paragraph(item.get("text", ""))
            p.runs[0].font.name = 'Times New Roman' if p.runs else None
    doc.save(filepath)

def export_report_pdf(text_lines, filepath):
    """Export plain text to PDF using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
    except ImportError:
        raise RuntimeError("Встановіть reportlab:\n  pip install reportlab")
    doc = SimpleDocTemplate(filepath, pagesize=A4,
                            rightMargin=2*cm, leftMargin=2.5*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle('h1', parent=styles['Heading2'], fontName='Times-Roman', fontSize=12, spaceAfter=4)
    normal = ParagraphStyle('n', parent=styles['Normal'], fontName='Times-Roman', fontSize=11, spaceAfter=2)
    story = []
    for item in text_lines:
        kind = item.get("kind","text")
        if kind == "heading":
            story.append(Paragraph(item["text"], h1))
        elif kind == "table":
            headers = item["headers"]; rows = item["rows"]
            data = [headers] + [[("" if v is None else str(v)) for v in r] for r in rows]
            t = Table(data, repeatRows=1)
            t.setStyle(TableStyle([
                ('FONTNAME',(0,0),(-1,0),'Times-Roman'), ('FONTSIZE',(0,0),(-1,0),10),
                ('FONTNAME',(0,1),(-1,-1),'Times-Roman'), ('FONTSIZE',(0,1),(-1,-1),9),
                ('GRID',(0,0),(-1,-1),0.5,colors.black),
                ('BACKGROUND',(0,0),(-1,0),colors.HexColor('#e8e8e8')),
            ]))
            story.append(t); story.append(Spacer(1,6))
        else:
            txt = item.get("text","").replace("\n","<br/>")
            if txt.strip(): story.append(Paragraph(txt, normal))
    doc.build(story)


